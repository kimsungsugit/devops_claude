#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
두 DOCX 파일 비교 분석 스크립트
- 우리가 생성한 UDS vs 레퍼런스 SUDS
"""

from docx import Document
from docx.oxml.ns import qn
import re
from collections import defaultdict

# 파일 경로
UDS_PATH = r"D:/Project/devops/260105/backend/reports/uds_local/uds_spec_generated_expanded_20260325_094858.docx"
SUDS_PATH = r"D:/Project/devops/260105/docs/(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"

def load_doc(path):
    try:
        doc = Document(path)
        print(f"  [OK] 로드 성공: {path.split('/')[-1]}")
        return doc
    except Exception as e:
        print(f"  [ERROR] 로드 실패: {e}")
        return None

def get_paragraph_style(para):
    return para.style.name if para.style else "Unknown"

def get_headings(doc):
    headings = []
    for para in doc.paragraphs:
        style = get_paragraph_style(para)
        if style.startswith("Heading"):
            headings.append((style, para.text.strip()))
    return headings

def get_table_info(table, idx):
    rows = table.rows
    cols = table.columns
    row_count = len(rows)
    col_count = len(cols) if cols else 0

    # 헤더 행 (첫 번째 행)
    header_cells = []
    if row_count > 0:
        for cell in rows[0].cells:
            text = cell.text.strip()
            header_cells.append(text)

    # 두 번째 행 (데이터 샘플)
    second_row = []
    if row_count > 1:
        for cell in rows[1].cells:
            text = cell.text.strip()[:50]  # 50자 제한
            second_row.append(text)

    return {
        "idx": idx,
        "rows": row_count,
        "cols": col_count,
        "headers": header_cells,
        "sample_row": second_row
    }

def analyze_structure(doc, name):
    print(f"\n{'='*60}")
    print(f"  {name} 구조 분석")
    print(f"{'='*60}")

    # 기본 통계
    paras = doc.paragraphs
    tables = doc.tables
    sections = doc.sections

    print(f"\n[기본 통계]")
    print(f"  총 단락 수: {len(paras)}")
    print(f"  총 테이블 수: {len(tables)}")
    print(f"  총 섹션 수: {len(sections)}")

    # 제목/커버 정보 (첫 10개 단락)
    print(f"\n[문서 시작부 (첫 15개 비어있지 않은 단락)]")
    count = 0
    for para in paras:
        if para.text.strip():
            style = get_paragraph_style(para)
            print(f"  [{style}] {para.text.strip()[:100]}")
            count += 1
            if count >= 15:
                break

    # 헤딩 구조
    headings = get_headings(doc)
    print(f"\n[헤딩 구조 - 총 {len(headings)}개]")
    for i, (style, text) in enumerate(headings[:30]):  # 최대 30개
        level = style.replace("Heading ", "H")
        indent = "  " * (int(style[-1]) - 1) if style[-1].isdigit() else ""
        print(f"  {indent}[{level}] {text[:80]}")
    if len(headings) > 30:
        print(f"  ... (총 {len(headings)}개 중 30개만 표시)")

    return {
        "para_count": len(paras),
        "table_count": len(tables),
        "section_count": len(sections),
        "headings": headings
    }

def analyze_tables(doc, name, max_tables=None):
    tables = doc.tables
    print(f"\n{'='*60}")
    print(f"  {name} 테이블 분석")
    print(f"{'='*60}")
    print(f"\n[전체 테이블 목록 - 총 {len(tables)}개]")

    display_count = min(len(tables), max_tables or len(tables))

    for i, table in enumerate(tables[:display_count]):
        info = get_table_info(table, i+1)
        print(f"\n  테이블 #{i+1}: {info['rows']}행 x {info['cols']}열")
        print(f"    헤더: {info['headers']}")
        if info['sample_row'] and info['sample_row'] != info['headers']:
            print(f"    샘플: {info['sample_row']}")

    if len(tables) > display_count:
        print(f"\n  ... (총 {len(tables)}개 중 {display_count}개만 표시)")

    return [get_table_info(t, i+1) for i, t in enumerate(tables)]

def find_function_sections(doc, name):
    """함수 항목 구조 파악"""
    print(f"\n{'='*60}")
    print(f"  {name} 함수 항목 구조 분석")
    print(f"{'='*60}")

    paras = doc.paragraphs
    tables = doc.tables

    # 함수처럼 보이는 헤딩 찾기 (번호 패턴 포함)
    func_headings = []
    for para in paras:
        style = get_paragraph_style(para)
        text = para.text.strip()
        if not text:
            continue
        # 함수명 패턴: 알파벳+숫자+괄호 또는 특정 헤딩 레벨
        if style.startswith("Heading") and (
            re.search(r'\b[A-Za-z_][A-Za-z0-9_]*\s*\(', text) or
            re.search(r'^\d+\.\d+', text) or
            re.search(r'[A-Z][a-z]+[A-Z]', text)  # CamelCase
        ):
            func_headings.append((style, text))

    print(f"\n[함수처럼 보이는 헤딩 - 총 {len(func_headings)}개]")
    for style, text in func_headings[:20]:
        print(f"  [{style}] {text[:80]}")
    if len(func_headings) > 20:
        print(f"  ... ({len(func_headings)}개 중 20개만 표시)")

    # 테이블 헤더 패턴 수집
    print(f"\n[테이블 헤더 패턴 분석]")
    header_patterns = defaultdict(int)
    for table in tables:
        if table.rows:
            headers = tuple(cell.text.strip() for cell in table.rows[0].cells)
            header_patterns[headers] += 1

    print(f"  고유 헤더 패턴 수: {len(header_patterns)}")
    for headers, count in sorted(header_patterns.items(), key=lambda x: -x[1])[:15]:
        print(f"  (x{count}) {list(headers)}")

    return func_headings, header_patterns

def compare_headings(uds_headings, suds_headings):
    print(f"\n{'='*60}")
    print(f"  헤딩 비교 분석")
    print(f"{'='*60}")

    uds_texts = set(text for _, text in uds_headings)
    suds_texts = set(text for _, text in suds_headings)

    only_in_suds = suds_texts - uds_texts
    only_in_uds = uds_texts - suds_texts
    common = uds_texts & suds_texts

    print(f"\n[공통 헤딩: {len(common)}개]")
    for text in sorted(common)[:20]:
        print(f"  - {text[:70]}")

    print(f"\n[레퍼런스(SUDS)에만 있는 헤딩: {len(only_in_suds)}개]")
    for text in sorted(only_in_suds)[:30]:
        print(f"  - {text[:70]}")

    print(f"\n[우리 UDS에만 있는 헤딩: {len(only_in_uds)}개]")
    for text in sorted(only_in_uds)[:30]:
        print(f"  - {text[:70]}")

def analyze_sample_functions(doc, name, sample_count=3):
    """샘플 함수 상세 분석"""
    print(f"\n{'='*60}")
    print(f"  {name} 샘플 함수 상세 분석")
    print(f"{'='*60}")

    # 문서 요소를 순서대로 순회
    body = doc.element.body
    current_heading = None
    current_level = None
    func_blocks = []  # (heading, tables)
    current_tables = []

    # XML body에서 순서대로 요소 처리
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # 단락 처리
            para_style = ""
            style_elem = child.find('.//' + qn('w:pStyle'))
            if style_elem is not None:
                para_style = style_elem.get(qn('w:val'), '')

            # 텍스트 추출
            texts = []
            for t in child.findall('.//' + qn('w:t')):
                if t.text:
                    texts.append(t.text)
            para_text = ''.join(texts).strip()

            if para_style.startswith('Heading') and para_text:
                # 이전 함수 블록 저장
                if current_heading and current_tables:
                    func_blocks.append((current_level, current_heading, current_tables))
                    current_tables = []

                current_heading = para_text
                current_level = para_style

        elif tag == 'tbl':
            # 테이블 처리
            if current_heading:
                current_tables.append(child)

    # 마지막 블록 저장
    if current_heading and current_tables:
        func_blocks.append((current_level, current_heading, current_tables))

    print(f"\n테이블을 포함한 섹션 수: {len(func_blocks)}")

    # 함수처럼 보이는 블록만 필터링
    func_like = [
        (lvl, h, tbls) for lvl, h, tbls in func_blocks
        if re.search(r'[A-Za-z_][A-Za-z0-9_]*\s*\(|[A-Z][a-z]+[A-Z]|\bfunction\b|\bvoid\b|\bint\b', h, re.IGNORECASE)
    ]

    print(f"함수처럼 보이는 섹션 수: {len(func_like)}")

    # 샘플 출력
    for i, (level, heading, tbl_elems) in enumerate(func_like[:sample_count]):
        print(f"\n  --- 샘플 {i+1}: [{level}] {heading[:60]} ---")
        print(f"  포함 테이블 수: {len(tbl_elems)}")

        for j, tbl_elem in enumerate(tbl_elems):
            # XML에서 직접 테이블 데이터 추출
            rows = tbl_elem.findall('.//' + qn('w:tr'))
            print(f"  테이블 {j+1}: {len(rows)}행")
            if rows:
                # 헤더 행
                header_cells = []
                for cell in rows[0].findall('.//' + qn('w:tc')):
                    texts = []
                    for t in cell.findall('.//' + qn('w:t')):
                        if t.text:
                            texts.append(t.text)
                    header_cells.append(''.join(texts).strip())
                print(f"    헤더: {header_cells}")

                # 데이터 행들
                for row_idx, row in enumerate(rows[1:4], 1):
                    row_cells = []
                    for cell in row.findall('.//' + qn('w:tc')):
                        texts = []
                        for t in cell.findall('.//' + qn('w:t')):
                            if t.text:
                                texts.append(t.text)
                        row_cells.append(''.join(texts).strip()[:40])
                    print(f"    행{row_idx}: {row_cells}")

    return func_like

def check_special_elements(doc, name):
    """특수 요소 분석"""
    print(f"\n{'='*60}")
    print(f"  {name} 특수 요소 분석")
    print(f"{'='*60}")

    # 스타일 종류
    styles_used = set()
    for para in doc.paragraphs:
        styles_used.add(get_paragraph_style(para))

    print(f"\n[사용된 단락 스타일 종류: {len(styles_used)}개]")
    for s in sorted(styles_used):
        print(f"  - {s}")

    # 특수 텍스트 패턴 탐색
    special_patterns = {
        "함수 프로토타입": r'\b\w+\s*\([^)]*\)\s*;',
        "번호 목록": r'^\s*\d+\.',
        "특수 기호/태그": r'<[^>]+>|\[.*?\]|\{.*?\}',
        "버전/날짜": r'v\d+\.\d+|\d{4}[-/]\d{2}[-/]\d{2}',
    }

    print(f"\n[특수 텍스트 패턴]")
    for pattern_name, pattern in special_patterns.items():
        matches = []
        for para in doc.paragraphs:
            if re.search(pattern, para.text):
                matches.append(para.text.strip()[:60])
        print(f"  {pattern_name}: {len(matches)}개")
        for m in matches[:3]:
            print(f"    예: {m}")

def main():
    print("=" * 70)
    print("  DOCX 파일 비교 분석")
    print("=" * 70)

    print("\n[파일 로드 중...]")
    uds_doc = load_doc(UDS_PATH)
    suds_doc = load_doc(SUDS_PATH)

    if not uds_doc or not suds_doc:
        print("\n[ERROR] 파일 로드 실패. 경로를 확인하세요.")
        return

    # 1. 전체 구조 비교
    print("\n" + "#" * 70)
    print("# 1. 전체 구조 비교")
    print("#" * 70)
    uds_info = analyze_structure(uds_doc, "우리 UDS")
    suds_info = analyze_structure(suds_doc, "레퍼런스 SUDS")

    print(f"\n{'='*60}")
    print(f"  구조 비교 요약")
    print(f"{'='*60}")
    print(f"  {'항목':<20} {'우리 UDS':>15} {'레퍼런스 SUDS':>20}")
    print(f"  {'-'*55}")
    print(f"  {'총 단락 수':<20} {uds_info['para_count']:>15} {suds_info['para_count']:>20}")
    print(f"  {'총 테이블 수':<20} {uds_info['table_count']:>15} {suds_info['table_count']:>20}")
    print(f"  {'총 섹션 수':<20} {uds_info['section_count']:>15} {suds_info['section_count']:>20}")
    print(f"  {'헤딩 수':<20} {len(uds_info['headings']):>15} {len(suds_info['headings']):>20}")

    # 2. 테이블 구조 비교
    print("\n" + "#" * 70)
    print("# 2. 테이블 구조 비교")
    print("#" * 70)
    uds_tables = analyze_tables(uds_doc, "우리 UDS", max_tables=10)
    suds_tables = analyze_tables(suds_doc, "레퍼런스 SUDS", max_tables=10)

    # 3. 함수 항목 비교
    print("\n" + "#" * 70)
    print("# 3. 함수 항목 구조 분석")
    print("#" * 70)
    uds_funcs, uds_headers = find_function_sections(uds_doc, "우리 UDS")
    suds_funcs, suds_headers = find_function_sections(suds_doc, "레퍼런스 SUDS")

    # 샘플 함수 상세 분석
    print("\n" + "#" * 70)
    print("# 3-2. 샘플 함수 상세 내용 비교")
    print("#" * 70)
    uds_func_blocks = analyze_sample_functions(uds_doc, "우리 UDS", sample_count=3)
    suds_func_blocks = analyze_sample_functions(suds_doc, "레퍼런스 SUDS", sample_count=3)

    # 4. 헤딩 비교
    print("\n" + "#" * 70)
    print("# 4. 헤딩/섹션 비교")
    print("#" * 70)
    compare_headings(uds_info['headings'], suds_info['headings'])

    # 5. 특수 요소 분석
    print("\n" + "#" * 70)
    print("# 5. 특수 요소 분석")
    print("#" * 70)
    check_special_elements(uds_doc, "우리 UDS")
    check_special_elements(suds_doc, "레퍼런스 SUDS")

    # 6. 레퍼런스 특이 항목 심층 분석
    print("\n" + "#" * 70)
    print("# 6. 레퍼런스(SUDS) 특이 항목 심층 분석")
    print("#" * 70)
    print("\n[SUDS 전체 헤딩 목록]")
    for style, text in suds_info['headings']:
        level = style.replace("Heading ", "H")
        indent = "  " * (int(style[-1]) - 1) if style and style[-1].isdigit() else ""
        print(f"  {indent}[{level}] {text[:80]}")

    print("\n[SUDS 전체 테이블 헤더 목록]")
    for i, table in enumerate(suds_doc.tables):
        info = get_table_info(table, i+1)
        print(f"  테이블#{i+1} ({info['rows']}행x{info['cols']}열): {info['headers']}")

    print("\n\n분석 완료!")

if __name__ == "__main__":
    main()
