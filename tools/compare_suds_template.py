import json
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document


def _iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", child)
        elif child.tag.endswith("}tbl"):
            yield ("tbl", child)


def _build_element_maps(doc: Document) -> Tuple[Dict[Any, Any], Dict[Any, Any]]:
    p_map = {para._p: para for para in doc.paragraphs}
    tbl_map = {table._tbl: table for table in doc.tables}
    return p_map, tbl_map


def _heading_level(style_name: str) -> int:
    if not style_name or not style_name.startswith("Heading"):
        return 0
    parts = [p for p in style_name.split() if p.isdigit()]
    if parts:
        return int(parts[0])
    for ch in style_name:
        if ch.isdigit():
            return int(ch)
    return 1


def analyze_docx(path: Path) -> Dict[str, Any]:
    doc = Document(str(path))
    headings: List[Tuple[int, str]] = []
    heading_stack: List[str] = []
    tables: List[Dict[str, Any]] = []
    paras: List[Dict[str, Any]] = []

    p_map, tbl_map = _build_element_maps(doc)

    for kind, node in _iter_blocks(doc):
        if kind == "p":
            para = p_map.get(node)
            if not para:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            style = str(getattr(para.style, "name", "") or "")
            level = _heading_level(style)
            if level > 0:
                while len(heading_stack) >= level:
                    heading_stack.pop()
                heading_stack.append(text)
                headings.append((level, text))
            else:
                paras.append({"text": text, "style": style, "ctx": " > ".join(heading_stack)})
        elif kind == "tbl":
            table = tbl_map.get(node)
            if not table or not table.rows:
                continue
            header_rows = []
            for row in table.rows[:2]:
                header_rows.append([cell.text.strip() for cell in row.cells])
            header_keywords = {
                "file name",
                "version",
                "date",
                "note",
                "macro",
                "type",
                "define",
                "description",
                "parameter",
                "component",
                "function",
                "comment",
                "data name",
                "data type",
                "value range",
                "reset",
            }
            def _data_like(row_text: str) -> bool:
                return any(
                    token in row_text
                    for token in [
                        "#define",
                        "0x",
                        "(",
                        ")",
                        "swufn_",
                        "swcom_",
                    ]
                )
            if len(header_rows) == 2:
                first = " ".join([c.lower() for c in header_rows[0] if c]).strip()
                second = " ".join([c.lower() for c in header_rows[1] if c]).strip()
                def _is_header(text: str) -> bool:
                    return any(k in text for k in header_keywords) and not _data_like(text)
                if _is_header(first) and not _is_header(second):
                    header_rows = header_rows[:1]
            tables.append(
                {
                    "ctx": " > ".join(heading_stack),
                    "cols": len(table.columns),
                    "header_rows": header_rows,
                }
            )
    return {"headings": headings, "tables": tables, "paras": paras}


def _normalize_headers(rows: List[List[str]]) -> List[List[str]]:
    normalized: List[List[str]] = []
    for row in rows:
        out_row = []
        for cell in row:
            base = cell.lower().strip()
            base = re.sub(r"\d+", "#", base)
            out_row.append(base)
        while out_row and out_row[-1] == "":
            out_row.pop()
        if out_row and len(set(out_row)) == 1:
            out_row = [out_row[0]]
        normalized.append(out_row)
    normalized = [
        row
        for row in normalized
        if not any(re.search(r"swufn_#", cell) for cell in row)
    ]
    return normalized


def compare_docs(base: Dict[str, Any], gen: Dict[str, Any]) -> Dict[str, Any]:
    base_headings = [">".join([str(lvl), text]) for lvl, text in base["headings"]]
    gen_headings = [">".join([str(lvl), text]) for lvl, text in gen["headings"]]
    base_heading_set = set(base_headings)
    gen_heading_set = set(gen_headings)

    base_tables_by_ctx = {}
    for idx, tbl in enumerate(base["tables"]):
        key = f"{tbl['ctx']}::#{idx}"
        base_tables_by_ctx[key] = tbl

    gen_tables_by_ctx = {}
    for idx, tbl in enumerate(gen["tables"]):
        key = f"{tbl['ctx']}::#{idx}"
        gen_tables_by_ctx[key] = tbl

    base_ctxs = set(base_tables_by_ctx.keys())
    gen_ctxs = set(gen_tables_by_ctx.keys())

    header_mismatches = []
    for key in sorted(base_ctxs & gen_ctxs):
        b = base_tables_by_ctx[key]
        g = gen_tables_by_ctx[key]
        def _all_empty(rows: List[List[str]]) -> bool:
            for row in rows:
                for cell in row:
                    if cell.strip():
                        return False
            return True
        if _all_empty(b["header_rows"]):
            continue
        def _has_dynamic_ids(rows: List[List[str]]) -> bool:
            for row in rows:
                for cell in row:
                    if re.search(r"swufn_", cell, flags=re.I):
                        return True
            return False
        if _has_dynamic_ids(b["header_rows"]) or _has_dynamic_ids(g["header_rows"]):
            continue
        if b["cols"] != g["cols"] or _normalize_headers(b["header_rows"]) != _normalize_headers(g["header_rows"]):
            header_mismatches.append(
                {
                    "ctx": b["ctx"],
                    "base_cols": b["cols"],
                    "gen_cols": g["cols"],
                    "base_header_rows": b["header_rows"],
                    "gen_header_rows": g["header_rows"],
                }
            )

    return {
        "heading_missing_in_generated": sorted(base_heading_set - gen_heading_set),
        "heading_extra_in_generated": sorted(gen_heading_set - base_heading_set),
        "table_missing_in_generated": sorted(base_ctxs - gen_ctxs),
        "table_extra_in_generated": sorted(gen_ctxs - base_ctxs),
        "table_header_mismatches": header_mismatches,
        "base_table_count": len(base["tables"]),
        "gen_table_count": len(gen["tables"]),
        "base_heading_count": len(base_headings),
        "gen_heading_count": len(gen_headings),
    }


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    template_path = repo_root / "docs" / "(HDPDM01_SUDS)_template_clean.docx"
    if not template_path.exists():
        template_path = repo_root / "docs" / "(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"
    report_dir = repo_root / "backend" / "reports" / "uds_local"
    candidates = sorted(report_dir.glob("uds_spec_generated_expanded_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        candidates = sorted(report_dir.glob("uds_local_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError(f"No generated UDS docx found under {report_dir}")
    generated_path = candidates[0]
    out_path = repo_root / "backend" / "reports" / "uds_local" / "compare_template_report.json"

    base = analyze_docx(template_path)
    gen = analyze_docx(generated_path)
    report = compare_docs(base, gen)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
