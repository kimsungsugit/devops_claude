from __future__ import annotations

import json
import re
from pathlib import Path

import docx  # type: ignore


def analyze(path: Path) -> dict:
    doc = docx.Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style
        and str(getattr(p.style, "name", "")).startswith("Heading")
        and p.text.strip()
    ]
    tables = len(doc.tables)
    words = sum(len(re.findall(r"\S+", t)) for t in paras)
    return {
        "path": str(path),
        "paras": len(paras),
        "headings": len(headings),
        "tables": tables,
        "words": words,
        "heading_samples": headings[:15],
    }


def main() -> None:
    base = Path(
        r"D:\Project\devops\260105\docs\(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"
    )
    generated = Path(
        r"D:\Project\devops\260105\backend\reports\uds_local\uds_spec_generated_expanded.docx"
    )
    data = {"base": analyze(base), "generated": analyze(generated)}
    print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
