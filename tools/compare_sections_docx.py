from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx import Document  # type: ignore


def _iter_blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", child)
        elif child.tag.endswith("}tbl"):
            yield ("tbl", child)


def _build_maps(doc: Document):
    p_map = {para._p: para for para in doc.paragraphs}
    t_map = {table._tbl: table for table in doc.tables}
    return p_map, t_map


def _heading_level(style_name: str) -> int:
    if not style_name or not style_name.startswith("Heading"):
        return 0
    for token in style_name.split():
        if token.isdigit():
            return int(token)
    for ch in style_name:
        if ch.isdigit():
            return int(ch)
    return 1


def _collect_sections(doc: Document, prefix: str) -> List[Tuple[str, List[str]]]:
    p_map, t_map = _build_maps(doc)
    sections: List[Tuple[str, List[str]]] = []
    current_title = ""
    current_lines: List[str] = []
    for kind, node in _iter_blocks(doc):
        if kind == "p":
            para = p_map.get(node)
            if not para:
                continue
            text = (para.text or "").strip()
            if not text:
                continue
            style = str(getattr(para.style, "name", "") or "")
            level = _heading_level(style)
            if level > 0:
                if current_title:
                    sections.append((current_title, current_lines))
                current_title = text
                current_lines = []
                continue
            if current_title:
                current_lines.append(text)
        elif kind == "tbl":
            table = t_map.get(node)
            if not table or not current_title:
                continue
            # summarize table header
            header = []
            if table.rows:
                header = [c.text.strip() for c in table.rows[0].cells]
            current_lines.append(f"[TABLE cols={len(table.columns)} header={header}]")
    if current_title:
        sections.append((current_title, current_lines))
    if prefix:
        return [(t, lines) for t, lines in sections if t.startswith(prefix)]
    return sections


def _norm(text: str) -> str:
    return " ".join(text.strip().split()).lower()


def _filter_titles(sections: List[Tuple[str, List[str]]], titles: List[str]) -> List[Tuple[str, List[str]]]:
    targets = {_norm(t): t for t in titles}
    out = []
    for t, lines in sections:
        nt = _norm(t)
        if nt in targets:
            out.append((t, lines))
    return out


def main() -> None:
    template = Path(r"D:\Project\devops\260105\docs\(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx")
    report_dir = Path(r"D:\Project\devops\260105\backend\reports\uds_local")
    candidates = sorted(
        report_dir.glob("uds_spec_generated_expanded_*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(f"No generated docx in {report_dir}")
    generated = candidates[0]
    out = Path(r"D:\Project\devops\260105\backend\reports\uds_local\compare_sections_1_2.txt")

    template_doc = Document(str(template))
    gen_doc = Document(str(generated))

    tpl_sections = _collect_sections(template_doc, "")
    gen_sections = _collect_sections(gen_doc, "")

    targets = [
        "Introduction",
        "Purpose",
        "Scope",
        "Terms, Abbreviations and Definitions",
        "Reference",
        "Software Unit Design",
        "Common Macro Definition",
        "Type Definition",
        "Parameter Definition",
        "Version Information",
        "Software Unit Structure",
    ]

    tpl_filtered = _filter_titles(tpl_sections, targets)
    gen_filtered = _filter_titles(gen_sections, targets)

    lines: List[str] = []
    lines.append("== TEMPLATE TITLES ==")
    for t, _ in tpl_sections[:40]:
        lines.append(t)
    lines.append("")
    lines.append("== GENERATED TITLES ==")
    for t, _ in gen_sections[:40]:
        lines.append(t)
    lines.append("")
    lines.append("== TEMPLATE ==")
    for title, content in tpl_filtered:
        lines.append(f"[{title}]")
        lines.append(f"lines={len(content)}")
        lines.extend(content[:40])
        if len(content) > 40:
            lines.append("...truncated...")
        lines.append("")
    lines.append("== GENERATED ==")
    for title, content in gen_filtered:
        lines.append(f"[{title}]")
        lines.append(f"lines={len(content)}")
        lines.extend(content[:40])
        if len(content) > 40:
            lines.append("...truncated...")
        lines.append("")
    out.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
