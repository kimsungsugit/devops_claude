from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx import Document  # type: ignore


def _heading_level(style_name: str) -> int:
    if not style_name or not style_name.startswith("Heading"):
        return 0
    for token in style_name.split():
        if token.isdigit():
            return int(token)
    for ch in style_name:
        if ch.isdigit():
            return int(ch)
    return 1


def _list_headings(doc: Document) -> List[Tuple[int, str]]:
    out: List[Tuple[int, str]] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = str(getattr(p.style, "name", "") or "")
        level = _heading_level(style)
        if level > 0:
            out.append((level, text))
    return out


def main() -> None:
    template = Path(r"D:\Project\devops\260105\docs\(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx")
    generated = Path(r"D:\Project\devops\260105\backend\reports\uds_local\uds_spec_generated_expanded_20260210_112926.docx")
    out = Path(r"D:\Project\devops\260105\backend\reports\uds_local\docx_headings.txt")

    tpl_headings = _list_headings(Document(str(template)))
    gen_headings = _list_headings(Document(str(generated)))

    lines: List[str] = []
    lines.append("== TEMPLATE HEADINGS ==")
    for lvl, text in tpl_headings[:200]:
        lines.append(f"{lvl} {text}")
    if len(tpl_headings) > 200:
        lines.append("...truncated...")
    lines.append("")
    lines.append("== GENERATED HEADINGS ==")
    for lvl, text in gen_headings[:200]:
        lines.append(f"{lvl} {text}")
    if len(gen_headings) > 200:
        lines.append("...truncated...")
    out.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
