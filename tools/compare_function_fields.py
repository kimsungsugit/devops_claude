from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List

from docx import Document  # type: ignore


def _extract_function_tables(doc: Document) -> Dict[str, Dict[str, str]]:
    result: Dict[str, Dict[str, str]] = {}
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = [c.text.strip() for c in table.rows[0].cells]
        if not header_cells or "[ Function Information ]" not in header_cells[0]:
            continue
        fn_id = ""
        if len(table.rows) > 1:
            row1 = [c.text.strip() for c in table.rows[1].cells]
            for cell in row1:
                m = re.search(r"(SwUFn_\d+)", cell)
                if m:
                    fn_id = m.group(1)
                    break
        if not fn_id:
            continue
        fields: Dict[str, str] = {}
        for row in table.rows[2:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            label = cells[0].strip()
            if not label:
                continue
            value = ""
            for cell in cells[2:]:
                if cell:
                    value = cell
                    break
            fields[label] = value
        result[fn_id] = fields
    return result


def main() -> None:
    template = Path(r"D:\Project\devops\260105\docs\(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx")
    report_dir = Path(r"D:\Project\devops\260105\backend\reports\uds_local")
    candidates = sorted(
        report_dir.glob("uds_spec_generated_expanded_*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(f"No generated docx in {report_dir}")
    generated = candidates[0]

    tpl_map = _extract_function_tables(Document(str(template)))
    gen_map = _extract_function_tables(Document(str(generated)))

    fields = ["Related ID", "Input Parameters", "Output Parameters"]
    mismatches: List[str] = []
    for fn_id, tpl_fields in tpl_map.items():
        gen_fields = gen_map.get(fn_id)
        if not gen_fields:
            continue
        for f in fields:
            tval = (tpl_fields.get(f) or "").strip()
            gval = (gen_fields.get(f) or "").strip()
            if tval and gval and tval != gval:
                mismatches.append(f"{fn_id} {f} tpl='{tval}' gen='{gval}'")
                if len(mismatches) >= 20:
                    break
        if len(mismatches) >= 20:
            break
    out = report_dir / "function_field_mismatches.txt"
    out.write_text("\n".join(mismatches) if mismatches else "No mismatches found.", encoding="utf-8")


if __name__ == "__main__":
    main()
