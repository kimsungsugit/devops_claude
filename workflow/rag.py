# /app/workflow/rag.py
# -*- coding: utf-8 -*-
# RAG Knowledge Base (v30.4: directory-backed, atomic writes)

from __future__ import annotations

import json
import os
import re
import sqlite3
from collections import OrderedDict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Iterable, Tuple

import logging
import time as _time

import numpy as np

_rag_logger = logging.getLogger("workflow.rag")
_RAG_PERF_LOG = str(os.environ.get("DEVOPS_RAG_PERF_LOG", "0")).strip().lower() in ("1", "true", "yes")

try:  # optional HTTP embedder
    import requests  # type: ignore
except Exception:  # pragma: no cover
    requests = None  # type: ignore

try:
    import psycopg2  # type: ignore
except Exception:  # pragma: no cover
    psycopg2 = None  # type: ignore

import config


def _cosine(a: np.ndarray, b: np.ndarray) -> float:
    if a.size == 0 or b.size == 0:
        return 0.0
    na = np.linalg.norm(a)
    nb = np.linalg.norm(b)
    if na == 0.0 or nb == 0.0:
        return 0.0
    return float(np.dot(a, b) / (na * nb))


def _normalize_message(msg: str) -> str:
    msg = msg or ""
    # 컴파일 로그 등의 노이즈 제거
    msg = re.sub(r"/app/[^\s]+", "<PATH>", msg)
    msg = re.sub(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}\.\d+Z", "<TIME>", msg)
    msg = re.sub(r"\s+", " ", msg)
    return msg.strip()


def _split_paths(val: Any) -> List[str]:
    if not val:
        return []
    if isinstance(val, list):
        return [str(v).strip() for v in val if str(v).strip()]
    s = str(val or "").strip()
    if not s:
        return []
    parts = [x.strip() for x in s.replace("\n", ",").replace(";", ",").split(",")]
    return [p for p in parts if p]


def _read_text_from_file(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext in (".txt", ".md", ".csv", ".log", ".json", ".xml", ".yaml", ".yml"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext in (".html", ".htm"):
            try:
                from bs4 import BeautifulSoup  # type: ignore
                html = path.read_text(encoding="utf-8", errors="ignore")
                return BeautifulSoup(html, "html.parser").get_text("\n")
            except Exception:
                return path.read_text(encoding="utf-8", errors="ignore")
        if ext in (".pdf",):
            try:
                try:
                    import pdfplumber  # type: ignore
                except Exception:
                    pdfplumber = None  # type: ignore
                texts = []
                if pdfplumber:
                    with pdfplumber.open(str(path)) as pdf:
                        for idx, page in enumerate(pdf.pages, start=1):
                            page_text = page.extract_text() or ""
                            if page_text:
                                texts.append(f"=== Page {idx} ===")
                                texts.append(page_text)
                            try:
                                tables = page.extract_tables() or []
                            except Exception:
                                tables = []
                            for t in tables:
                                rows = []
                                for row in t:
                                    if not row:
                                        continue
                                    cells = [str(c or "").strip() for c in row]
                                    if any(cells):
                                        rows.append(" | ".join(cells))
                                if rows:
                                    texts.append("=== Table ===")
                                    texts.extend(rows)
                    return "\n".join(texts)
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(str(path))
                texts = []
                for idx, p in enumerate(reader.pages, start=1):
                    try:
                        page_text = p.extract_text() or ""
                    except Exception:
                        page_text = ""
                    if page_text:
                        texts.append(f"=== Page {idx} ===")
                        texts.append(page_text)
                return "\n".join(texts)
            except Exception:
                return ""
        if ext in (".docx",):
            try:
                import docx  # type: ignore
                doc = docx.Document(str(path))
                lines = []
                for p in doc.paragraphs:
                    text = (p.text or "").strip()
                    if not text:
                        continue
                    style = str(getattr(p, "style", "") or "")
                    style_name = ""
                    try:
                        style_name = p.style.name  # type: ignore
                    except Exception:
                        style_name = str(style)
                    if "Heading" in style_name:
                        level = re.findall(r"\d+", style_name)
                        prefix = "#" * int(level[0]) if level else "##"
                        lines.append(f"{prefix} {text}")
                    elif "TOC" in style_name or "Table of Contents" in style_name:
                        lines.append(f"TOC: {text}")
                    else:
                        lines.append(text)
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        lines.append("=== Table ===")
                        lines.extend(rows)
                return "\n".join(lines)
            except Exception:
                return ""
    except Exception:
        return ""
    return ""


def _chunk_text(text: str, *, chunk_size: int = 1200, overlap: int = 200) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    if chunk_size <= 0:
        return [text]
    out: List[str] = []
    step = max(1, chunk_size - max(0, overlap))
    for i in range(0, len(text), step):
        out.append(text[i : i + chunk_size])
    return [c for c in out if c.strip()]


REQ_ID_PATTERN = re.compile(r"\b(?:REQ|SDS|SW|SWS|SWR|SRS|SWC|REQS)\s*[-_:]?\s*[A-Za-z0-9_.-]+\b")


def _extract_req_ids_from_text(text: str) -> List[str]:
    if not text:
        return []
    ids = [m.group(0).replace(" ", "").strip() for m in REQ_ID_PATTERN.finditer(text)]
    uniq: List[str] = []
    seen = set()
    for rid in ids:
        if not rid or rid in seen:
            continue
        seen.add(rid)
        uniq.append(rid)
    return uniq


def _chunk_by_req_ids(text: str, *, chunk_size: int, overlap: int) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    matches = list(REQ_ID_PATTERN.finditer(text))
    if len(matches) < 2:
        return _chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    chunks: List[str] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        seg = text[start:end].strip()
        if not seg:
            continue
        if len(seg) > chunk_size * 2:
            chunks.extend(_chunk_text(seg, chunk_size=chunk_size, overlap=overlap))
        else:
            chunks.append(seg)
    return [c for c in chunks if c.strip()]


def _chunk_docx_by_heading(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    try:
        import docx  # type: ignore
    except Exception:
        return []
    try:
        doc = docx.Document(str(path))
    except Exception:
        return []
    sections: List[Tuple[str, List[str]]] = []
    current_title = ""
    current_lines: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = p.style.name  # type: ignore
        except Exception:
            style_name = str(getattr(p, "style", "") or "")
        if "Heading" in style_name:
            if current_lines:
                sections.append((current_title, current_lines))
            current_title = text
            current_lines = []
            continue
        current_lines.append(text)
    if current_lines:
        sections.append((current_title, current_lines))
    chunks: List[str] = []
    for title, lines in sections:
        block = "\n".join([title.strip(), "\n".join(lines).strip()]).strip()
        if not block:
            continue
        chunks.extend(_chunk_by_req_ids(block, chunk_size=chunk_size, overlap=overlap))
    return [c for c in chunks if c.strip()]


def _chunk_xlsx_rows(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    try:
        import pandas as pd  # type: ignore
    except Exception:
        return []
    chunks: List[str] = []
    try:
        sheets = pd.read_excel(str(path), sheet_name=None)
    except Exception:
        return []
    for sheet_name, df in sheets.items():
        if df is None:
            continue
        try:
            records = df.fillna("").to_dict(orient="records")
        except Exception:
            continue
        for idx, row in enumerate(records):
            payload = {"sheet": sheet_name, "row_index": idx + 1, "data": row}
            text = json.dumps(payload, ensure_ascii=False)
            chunks.extend(_chunk_text(text, chunk_size=chunk_size, overlap=overlap))
    return [c for c in chunks if c.strip()]


def _chunk_c_by_function(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    """AST-based chunking for C/H files: one chunk per function definition."""
    try:
        from workflow.code_parser.c_parser import parse_c_project
    except ImportError:
        return []
    try:
        src = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return []
    chunks: List[str] = []
    func_pattern = re.compile(
        r"(?:^|\n)"
        r"((?:/\*[\s\S]*?\*/\s*|//[^\n]*\n\s*)*)"
        r"((?:static\s+|inline\s+|extern\s+|const\s+)*\w[\w\s*]+\s+\w+\s*\([^)]*\)\s*\{)",
        re.MULTILINE,
    )
    last_end = 0
    for m in func_pattern.finditer(src):
        start = m.start()
        brace_count = 0
        body_start = src.index("{", m.start(2))
        pos = body_start
        while pos < len(src):
            if src[pos] == "{":
                brace_count += 1
            elif src[pos] == "}":
                brace_count -= 1
                if brace_count == 0:
                    func_text = src[start:pos + 1].strip()
                    if len(func_text) > chunk_size:
                        func_text = func_text[:chunk_size]
                    if func_text:
                        chunks.append(f"[{path.name}]\n{func_text}")
                    last_end = pos + 1
                    break
            pos += 1
    if not chunks:
        return _chunk_text(src, chunk_size=chunk_size, overlap=overlap)
    return chunks


def _chunk_source_file(
    path: Path,
    *,
    chunk_size: int,
    overlap: int,
    max_chunks: int,
) -> List[str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        chunks = _chunk_docx_by_heading(path, chunk_size=chunk_size, overlap=overlap)
    elif ext == ".xlsx":
        chunks = _chunk_xlsx_rows(path, chunk_size=chunk_size, overlap=overlap)
    elif ext == ".pdf":
        text = _read_text_from_file(path)
        chunks = _chunk_by_req_ids(text, chunk_size=chunk_size, overlap=overlap)
    elif ext in {".c", ".h", ".cpp", ".hpp"}:
        chunks = _chunk_c_by_function(path, chunk_size=chunk_size, overlap=overlap)
    else:
        text = _read_text_from_file(path)
        chunks = _chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    return chunks[:max(1, int(max_chunks))]


def _read_and_chunk_file(
    path: Path,
    *,
    chunk_size: int = 1200,
    overlap: int = 200,
    max_chunks: int = 12,
) -> List[str]:
    if not path or not Path(path).exists():
        return []
    return _chunk_source_file(
        Path(path),
        chunk_size=int(chunk_size or 1200),
        overlap=int(overlap or 0),
        max_chunks=int(max_chunks or 1),
    )


class KnowledgeBase:
    """
    디렉터리 기반 RAG 저장소

    - base_dir/ 아래에 여러 개의 JSON 엔트리 파일 생성
      예) kb_store/kb_20251203T075959123456Z.json
    - 각 파일에는 단일 dict 엔트리 저장
    - 저장 시 항상 tmp 파일에 먼저 쓰고 rename 으로 교체
    """

    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.base_dir.mkdir(parents=True, exist_ok=True)
        self.data: List[Dict[str, Any]] = []
        force_pg = bool(getattr(config, "FORCE_PGVECTOR", False))
        force_pg_strict = bool(getattr(config, "FORCE_PGVECTOR_STRICT", False))
        self.storage = str(getattr(config, "KB_STORAGE", "sqlite") or "sqlite").strip().lower()
        if force_pg:
            self.storage = "pgvector"
        self.db_path = self.base_dir / "kb_index.sqlite"
        self._db_ok = False
        self._db_has_source_file = False
        self.pg_dsn = (
            str(getattr(config, "PGVECTOR_DSN", "") or "").strip()
            or str(getattr(config, "PGVECTOR_URL", "") or "").strip()
            or str(os.environ.get("PGVECTOR_DSN", "") or "").strip()
            or str(os.environ.get("PGVECTOR_URL", "") or "").strip()
        )
        self._pg_ok = False
        self._embed_cache: "OrderedDict[str, List[float]]" = OrderedDict()
        self._embed_cache_max = int(getattr(config, "KB_EMBED_CACHE_MAX", 1000))
        self._max_entries = int(getattr(config, "KB_MAX_ENTRIES", 5000))
        self._db_has_error_count = False
        self._db_has_project_root = False
        self._db_has_metadata = False

        if self.storage == "pgvector":
            if psycopg2 is None or not self.pg_dsn:
                self._pg_ok = False
                if not force_pg:
                    self.storage = "sqlite"
                elif force_pg_strict:
                    raise RuntimeError("pgvector required but psycopg2/dsn not configured")
            else:
                self._pg_ok = self._init_pgvector()
                if force_pg and force_pg_strict and not self._pg_ok:
                    raise RuntimeError("pgvector required but initialization failed")

        if self.storage == "sqlite":
            self._db_ok = self._init_db()
        self._load_all()
        self._ingest_sources_once()

    # ---------------- 내부 유틸 ----------------

    def _iter_entry_files(self):
        for p in self.base_dir.glob("*.json"):
            if p.name.endswith(".tmp") or p.name.endswith(".bak"):
                continue
            yield p

    def _write_atomic(self, path: Path, payload: Any) -> None:
        tmp = path.with_suffix(path.suffix + ".tmp")
        tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)

    def _ensure_shape(self, raw: Dict[str, Any], source_file: str) -> Dict[str, Any]:
        d = dict(raw)
        error_raw = str(d.get("error_raw", ""))
        error_clean = _normalize_message(d.get("error_clean") or error_raw)
        d["error_raw"] = error_raw
        d["error_clean"] = error_clean
        d["fix"] = str(d.get("fix", ""))
        tags = d.get("tags") or []
        if not isinstance(tags, list):
            tags = [str(tags)]
        d["tags"] = [str(t) for t in tags]
        role = d.get("role")
        d["role"] = str(role) if role is not None else None
        stage = d.get("stage")
        d["stage"] = str(stage) if stage is not None else None
        context = d.get("context")
        d["context"] = str(context) if context is not None else ""
        category = d.get("category") or d.get("kb_category") or ""
        d["category"] = str(category).strip() if category else "general"
        vec = d.get("vector") or []
        if isinstance(vec, list):
            d["vector"] = [float(x) for x in vec]
        else:
            d["vector"] = []
        d["weight"] = float(d.get("weight", 1.0))
        d["apply_count"] = int(d.get("apply_count", 0))
        d["timestamp"] = d.get("timestamp") or datetime.utcnow().isoformat()
        d["source_file"] = source_file
        if "id" not in d:
            d["id"] = os.path.splitext(os.path.basename(source_file))[0]
        metadata = d.get("metadata")
        if not isinstance(metadata, dict):
            metadata = {}
        d["metadata"] = metadata
        return d

    def _external_index_path(self) -> Path:
        return self.base_dir / "kb_external_index.json"

    def _load_external_index(self) -> Dict[str, str]:
        try:
            p = self._external_index_path()
            if p.exists():
                obj = json.loads(p.read_text(encoding="utf-8"))
                if isinstance(obj, dict):
                    return {str(k): str(v) for k, v in obj.items()}
        except Exception:
            return {}
        return {}

    def _save_external_index(self, idx: Dict[str, str]) -> None:
        try:
            self._write_atomic(self._external_index_path(), idx)
        except Exception:
            pass

    def add_document(
        self,
        title: str,
        content: str,
        *,
        category: str,
        tags: Optional[List[str]] = None,
        source_file: Optional[str] = None,
    ) -> None:
        text = (content or "").strip()
        if not text:
            return
        vec = self._get_embedding(text)
        req_ids = _extract_req_ids_from_text(text)
        entry: Dict[str, Any] = {
            "id": self._new_id(),
            "error_raw": str(title or "")[:200],
            "error_clean": str(title or "")[:200],
            "fix": text,
            "tags": tags or [],
            "role": "rag",
            "stage": "rag",
            "context": text,
            "category": str(category or "general"),
            "vector": vec,
            "weight": 1.0,
            "apply_count": 0,
            "timestamp": datetime.utcnow().isoformat(),
            "source_file": source_file or "",
            "metadata": {
                "req_ids": req_ids,
                "source_type": str(category or "general"),
            },
        }
        self._append_new_entry(entry, write_to_disk=True)

    def _init_db(self) -> bool:
        try:
            conn = sqlite3.connect(str(self.db_path))
            cur = conn.cursor()
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS kb_entries (
                    id TEXT PRIMARY KEY,
                    error_raw TEXT,
                    error_clean TEXT,
                    fix TEXT,
                    tags TEXT,
                    role TEXT,
                    stage TEXT,
                    category TEXT,
                    context TEXT,
                    vector TEXT,
                    weight REAL,
                    apply_count INTEGER,
                    timestamp TEXT,
                    source_file TEXT,
                    error_count INTEGER,
                    project_root TEXT,
                    metadata TEXT
                )
                """
            )
            try:
                cur.execute("PRAGMA table_info(kb_entries)")
                cols = [row[1] for row in cur.fetchall()]
                if "source_file" not in cols:
                    cur.execute("ALTER TABLE kb_entries ADD COLUMN source_file TEXT")
                if "error_count" not in cols:
                    cur.execute("ALTER TABLE kb_entries ADD COLUMN error_count INTEGER")
                if "project_root" not in cols:
                    cur.execute("ALTER TABLE kb_entries ADD COLUMN project_root TEXT")
                if "metadata" not in cols:
                    cur.execute("ALTER TABLE kb_entries ADD COLUMN metadata TEXT")
                self._db_has_source_file = True
                self._db_has_error_count = True
                self._db_has_project_root = True
                self._db_has_metadata = True
            except Exception:
                self._db_has_source_file = False
                self._db_has_error_count = False
                self._db_has_project_root = False
                self._db_has_metadata = False
            conn.commit()
            conn.close()
            return True
        except Exception:
            return False

    def _pg_connect(self):
        if psycopg2 is None or not self.pg_dsn:
            return None
        try:
            return psycopg2.connect(self.pg_dsn)
        except Exception:
            return None

    def _vector_to_str(self, vec: List[float]) -> Optional[str]:
        if not vec:
            return None
        try:
            return "[" + ",".join(f"{float(v):.6f}" for v in vec) + "]"
        except Exception:
            return None

    def _init_pgvector(self) -> bool:
        conn = self._pg_connect()
        if conn is None:
            return False
        try:
            cur = conn.cursor()
            cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS kb_entries (
                    id TEXT PRIMARY KEY,
                    error_raw TEXT,
                    error_clean TEXT,
                    fix TEXT,
                    tags JSONB,
                    role TEXT,
                    stage TEXT,
                    category TEXT,
                    context TEXT,
                    vector VECTOR(64),
                    weight REAL,
                    apply_count INTEGER,
                    timestamp TEXT,
                    source_file TEXT,
                    error_count INTEGER,
                    project_root TEXT,
                    metadata JSONB
                )
                """
            )
            try:
                cur.execute("ALTER TABLE kb_entries ADD COLUMN IF NOT EXISTS source_file TEXT")
                cur.execute("ALTER TABLE kb_entries ADD COLUMN IF NOT EXISTS error_count INTEGER")
                cur.execute("ALTER TABLE kb_entries ADD COLUMN IF NOT EXISTS project_root TEXT")
                cur.execute("ALTER TABLE kb_entries ADD COLUMN IF NOT EXISTS metadata JSONB")
                self._db_has_metadata = True
            except Exception:
                self._db_has_metadata = False
            conn.commit()
            conn.close()
            return True
        except Exception:
            try:
                conn.close()
            except Exception:
                pass
            return False

    def _pg_upsert(self, entry: Dict[str, Any]) -> None:
        if not self._pg_ok:
            return
        conn = self._pg_connect()
        if conn is None:
            return
        try:
            cur = conn.cursor()
            vec_str = self._vector_to_str(entry.get("vector") or [])
            cur.execute(
                """
                INSERT INTO kb_entries
                (id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp, source_file, error_count, project_root, metadata)
                VALUES (%s, %s, %s, %s, %s::jsonb, %s, %s, %s, %s, %s::vector, %s, %s, %s, %s, %s, %s, %s::jsonb)
                ON CONFLICT (id) DO UPDATE SET
                    error_raw = EXCLUDED.error_raw,
                    error_clean = EXCLUDED.error_clean,
                    fix = EXCLUDED.fix,
                    tags = EXCLUDED.tags,
                    role = EXCLUDED.role,
                    stage = EXCLUDED.stage,
                    category = EXCLUDED.category,
                    context = EXCLUDED.context,
                    vector = EXCLUDED.vector,
                    weight = EXCLUDED.weight,
                    apply_count = EXCLUDED.apply_count,
                    timestamp = EXCLUDED.timestamp,
                    source_file = EXCLUDED.source_file,
                    error_count = EXCLUDED.error_count,
                    project_root = EXCLUDED.project_root,
                    metadata = EXCLUDED.metadata
                """,
                (
                    entry.get("id"),
                    entry.get("error_raw"),
                    entry.get("error_clean"),
                    entry.get("fix"),
                    json.dumps(entry.get("tags") or [], ensure_ascii=False),
                    entry.get("role"),
                    entry.get("stage"),
                    entry.get("category"),
                    entry.get("context"),
                    vec_str,
                    float(entry.get("weight", 1.0)),
                    int(entry.get("apply_count", 0)),
                    entry.get("timestamp"),
                    entry.get("source_file") or "",
                    int(entry.get("error_count", 0)),
                    entry.get("project_root") or "",
                    json.dumps(entry.get("metadata") or {}, ensure_ascii=False),
                ),
            )
            conn.commit()
            conn.close()
        except Exception:
            try:
                conn.close()
            except Exception:
                pass

    def _pg_search(
        self,
        q_vec: List[float],
        query: str,
        top_k: int,
        *,
        tags: Optional[List[str]] = None,
        role: Optional[str] = None,
        stage: Optional[str] = None,
        categories: Optional[List[str]] = None,
        req_ids: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        if not self._pg_ok:
            return []
        conn = self._pg_connect()
        if conn is None:
            return []
        if not q_vec:
            return []
        q_vec_str = self._vector_to_str(q_vec)
        if not q_vec_str:
            return []

        clauses: List[str] = []
        params: List[Any] = []
        if role:
            clauses.append("role = %s")
            params.append(role)
        if stage:
            clauses.append("stage = %s")
            params.append(stage)
        if categories:
            clauses.append("category = ANY(%s)")
            params.append(categories)

        where_sql = ""
        if clauses:
            where_sql = " WHERE " + " AND ".join(clauses)

        sql = (
            "SELECT id, error_raw, error_clean, fix, tags, role, stage, category, context, "
            "weight, apply_count, timestamp, source_file, error_count, project_root, metadata, "
            "(1 - (vector <=> %s::vector)) AS score "
            f"FROM kb_entries{where_sql} "
            "ORDER BY vector <=> %s::vector NULLS LAST "
            "LIMIT %s"
        )
        params_with_vec = list(params) + [q_vec_str, q_vec_str, int(max(1, top_k * 5))]

        results: List[Dict[str, Any]] = []
        project_boost = float(getattr(config, "RAG_PROJECT_BOOST", 0.0))
        recency_days = float(getattr(config, "RAG_RECENCY_DAYS", 0))
        recency_boost = float(getattr(config, "RAG_RECENCY_BOOST", 0.0))
        apply_boost = float(getattr(config, "RAG_APPLY_COUNT_BOOST", 0.0))
        error_boost = float(getattr(config, "RAG_ERROR_COUNT_BOOST", 0.0))
        try:
            cur = conn.cursor()
            cur.execute(sql, params_with_vec)
            rows = cur.fetchall()
            conn.close()
        except Exception:
            try:
                conn.close()
            except Exception:
                pass
            return []

        norm_tags = [str(t) for t in (tags or []) if str(t).strip()]
        req_ids = [str(x).strip() for x in (req_ids or []) if str(x).strip()]
        exact_boost = float(getattr(config, "RAG_EXACT_MATCH_BOOST", 0.4))
        for row in rows:
            ent_tags = []
            try:
                ent_tags = json.loads(row[4] or "[]")
                if not isinstance(ent_tags, list):
                    ent_tags = []
            except Exception:
                ent_tags = []

            if norm_tags:
                hit = len(set(ent_tags).intersection(norm_tags))
                if hit == 0:
                    continue

            metadata = {}
            try:
                metadata = row[15] or {}
            except Exception:
                metadata = {}
            if isinstance(metadata, str):
                try:
                    metadata = json.loads(metadata)
                except Exception:
                    metadata = {}
            item = {
                "id": row[0],
                "error_raw": row[1],
                "error_clean": row[2],
                "fix": row[3],
                "tags": ent_tags,
                "role": row[5],
                "stage": row[6],
                "category": row[7] or "general",
                "context": row[8] or "",
                "weight": float(row[9] or 1.0),
                "apply_count": int(row[10] or 0),
                "timestamp": row[11] or "",
                "source_file": row[12] or "",
                "error_count": int(row[13] or 0),
                "project_root": row[14] or "",
                "metadata": metadata if isinstance(metadata, dict) else {},
                "score": float(row[16] or 0.0),
            }
            score = float(item.get("score") or 0.0)
            if recency_days > 0 and item.get("timestamp"):
                try:
                    ts = datetime.fromisoformat(str(item.get("timestamp")))
                    delta_days = (datetime.utcnow() - ts).total_seconds() / 86400.0
                    if delta_days < recency_days:
                        score += recency_boost * (1.0 - (delta_days / recency_days))
                except Exception:
                    pass
            if project_boost > 0:
                project_root = str(item.get("project_root") or "")
                if project_root and project_root in query:
                    score += project_boost
            if apply_boost > 0:
                score += apply_boost * float(item.get("apply_count") or 0)
            if error_boost > 0:
                score += error_boost * float(item.get("error_count") or 0)
            if req_ids:
                hay = " ".join(
                    [
                        str(item.get("error_raw") or ""),
                        str(item.get("error_clean") or ""),
                        str(item.get("context") or ""),
                        str(item.get("source_file") or ""),
                        json.dumps(item.get("metadata") or {}, ensure_ascii=False),
                    ]
                )
                if any(rid in hay for rid in req_ids):
                    score += exact_boost
            item["score"] = float(score)
            results.append(item)

        results.sort(key=lambda x: x.get("score", 0.0), reverse=True)
        return results[:top_k]

    def _db_upsert(self, entry: Dict[str, Any]) -> None:
        if not self._db_ok or self.storage != "sqlite":
            return
        try:
            conn = sqlite3.connect(str(self.db_path))
            cur = conn.cursor()
            if self._db_has_source_file:
                if self._db_has_metadata:
                    cur.execute(
                        """
                        INSERT OR REPLACE INTO kb_entries
                        (id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp, source_file, error_count, project_root, metadata)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            entry.get("id"),
                            entry.get("error_raw"),
                            entry.get("error_clean"),
                            entry.get("fix"),
                            json.dumps(entry.get("tags") or [], ensure_ascii=False),
                            entry.get("role"),
                            entry.get("stage"),
                            entry.get("category"),
                            entry.get("context"),
                            json.dumps(entry.get("vector") or [], ensure_ascii=False),
                            float(entry.get("weight", 1.0)),
                            int(entry.get("apply_count", 0)),
                            entry.get("timestamp"),
                            entry.get("source_file") or "",
                            int(entry.get("error_count", 0)),
                            entry.get("project_root") or "",
                            json.dumps(entry.get("metadata") or {}, ensure_ascii=False),
                        ),
                    )
                else:
                    cur.execute(
                        """
                        INSERT OR REPLACE INTO kb_entries
                        (id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp, source_file, error_count, project_root)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            entry.get("id"),
                            entry.get("error_raw"),
                            entry.get("error_clean"),
                            entry.get("fix"),
                            json.dumps(entry.get("tags") or [], ensure_ascii=False),
                            entry.get("role"),
                            entry.get("stage"),
                            entry.get("category"),
                            entry.get("context"),
                            json.dumps(entry.get("vector") or [], ensure_ascii=False),
                            float(entry.get("weight", 1.0)),
                            int(entry.get("apply_count", 0)),
                            entry.get("timestamp"),
                            entry.get("source_file") or "",
                            int(entry.get("error_count", 0)),
                            entry.get("project_root") or "",
                        ),
                    )
            else:
                cur.execute(
                    """
                    INSERT OR REPLACE INTO kb_entries
                    (id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        entry.get("id"),
                        entry.get("error_raw"),
                        entry.get("error_clean"),
                        entry.get("fix"),
                        json.dumps(entry.get("tags") or [], ensure_ascii=False),
                        entry.get("role"),
                        entry.get("stage"),
                        entry.get("category"),
                        entry.get("context"),
                        json.dumps(entry.get("vector") or [], ensure_ascii=False),
                        float(entry.get("weight", 1.0)),
                        int(entry.get("apply_count", 0)),
                        entry.get("timestamp"),
                    ),
                )
            conn.commit()
            conn.close()
        except Exception:
            pass

    def _db_load_all(self) -> List[Dict[str, Any]]:
        if not self._db_ok or self.storage != "sqlite":
            return []
        out: List[Dict[str, Any]] = []
        try:
            conn = sqlite3.connect(str(self.db_path))
            cur = conn.cursor()
            if self._db_has_source_file:
                if self._db_has_metadata:
                    cur.execute(
                        """
                        SELECT id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp, source_file, error_count, project_root, metadata
                        FROM kb_entries
                        """
                    )
                else:
                    cur.execute(
                        """
                        SELECT id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp, source_file, error_count, project_root
                        FROM kb_entries
                        """
                    )
            else:
                cur.execute(
                    """
                    SELECT id, error_raw, error_clean, fix, tags, role, stage, category, context, vector, weight, apply_count, timestamp
                    FROM kb_entries
                    """
                )
            rows = cur.fetchall()
            conn.close()
            for row in rows:
                tags = []
                vector = []
                try:
                    tags = json.loads(row[4] or "[]")
                except Exception:
                    tags = []
                try:
                    vector = json.loads(row[9] or "[]")
                except Exception:
                    vector = []
                source_file = ""
                error_count = 0
                project_root = ""
                metadata = {}
                if self._db_has_source_file and len(row) > 13:
                    source_file = row[13] or ""
                if self._db_has_error_count and len(row) > 14:
                    try:
                        error_count = int(row[14] or 0)
                    except Exception:
                        error_count = 0
                if self._db_has_project_root and len(row) > 15:
                    project_root = row[15] or ""
                if self._db_has_metadata and len(row) > 16:
                    try:
                        metadata = json.loads(row[16] or "{}")
                    except Exception:
                        metadata = {}
                out.append(
                    {
                        "id": row[0],
                        "error_raw": row[1],
                        "error_clean": row[2],
                        "fix": row[3],
                        "tags": tags if isinstance(tags, list) else [],
                        "role": row[5],
                        "stage": row[6],
                        "category": row[7] or "general",
                        "context": row[8] or "",
                        "vector": vector if isinstance(vector, list) else [],
                        "weight": float(row[10] or 1.0),
                        "apply_count": int(row[11] or 0),
                        "timestamp": row[12] or "",
                        "source_file": source_file,
                        "error_count": error_count,
                        "project_root": project_root,
                        "metadata": metadata if isinstance(metadata, dict) else {},
                    }
                )
        except Exception:
            return []
        return out

    def _ingest_sources_once(self) -> None:
        src_dir = str(getattr(config, "KB_SOURCES_DIR", "") or os.environ.get("KB_SOURCES_DIR", "") or "").strip()
        if not src_dir:
            return
        base = Path(src_dir).expanduser().resolve()
        if not base.exists() or not base.is_dir():
            return
        index_path = self.base_dir / "kb_ingest_index.json"
        seen: Dict[str, Any] = {}
        try:
            if index_path.exists():
                seen = json.loads(index_path.read_text(encoding="utf-8"))
                if not isinstance(seen, dict):
                    seen = {}
        except Exception:
            seen = {}

        updated = False
        for fp in base.rglob("*.json"):
            try:
                rel = fp.relative_to(base).as_posix()
                stat = fp.stat()
                sig = f"{stat.st_mtime_ns}:{stat.st_size}"
                if seen.get(rel) == sig:
                    continue
                payload = json.loads(fp.read_text(encoding="utf-8"))
                entries = payload if isinstance(payload, list) else [payload]
                cat = fp.parent.name if fp.parent != base else "general"
                for ent in entries:
                    if not isinstance(ent, dict):
                        continue
                    ent.setdefault("category", cat)
                    shaped = self._ensure_shape(ent, fp.name)
                    self._append_new_entry(shaped, write_to_disk=True)
                seen[rel] = sig
                updated = True
            except Exception:
                continue

        if updated:
            try:
                index_path.write_text(json.dumps(seen, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception:
                pass

    def _load_all(self) -> None:
        started = _time.perf_counter()
        started = _time.perf_counter()
        if self.storage == "pgvector" and self._pg_ok:
            self.data.clear()
            return
        # 1) 레거시 단일 JSON 파일 → 디렉터리로 마이그레이션 (최초 1회)
        legacy = self.base_dir.parent / "knowledge_base.json"
        has_entries = any(True for _ in self._iter_entry_files())
        if legacy.exists() and not has_entries:
            try:
                raw = json.loads(legacy.read_text(encoding="utf-8"))
                if isinstance(raw, dict):
                    raw = raw.get("data") or raw.get("entries") or []
                if isinstance(raw, list):
                    for idx, ent in enumerate(raw):
                        if not isinstance(ent, dict):
                            continue
                        tmp = self._ensure_shape(ent, f"legacy_{idx}.json")
                        self._append_new_entry(tmp, write_to_disk=True)
                # 백업 후 원본 rename
                legacy.rename(legacy.with_suffix(legacy.suffix + ".bak"))
            except Exception:
                # 실패해도 그냥 무시, 이후 현재 디렉터리만 사용
                pass

        # 2) 디렉터리 내 엔트리 로드
        self.data.clear()
        if self.storage == "sqlite" and self._db_ok:
            self.data = self._db_load_all()
            if self.data:
                return

        for fp in sorted(self._iter_entry_files(), key=lambda p: p.name):
            try:
                txt = fp.read_text(encoding="utf-8")
                if not txt.strip():
                    continue
                obj = json.loads(txt)
                if isinstance(obj, list):
                    objs = obj
                else:
                    objs = [obj]
                for ent in objs:
                    if not isinstance(ent, dict):
                        continue
                    shaped = self._ensure_shape(ent, fp.name)
                    self._db_upsert(shaped)
                    self.data.append(shaped)
            except Exception:
                # 손상된 파일은 건너뛰기
                continue

    def _new_id(self) -> str:
        return "kb_" + datetime.utcnow().strftime("%Y%m%dT%H%M%S%fZ")

    def _append_new_entry(self, entry: Dict[str, Any], write_to_disk: bool = True) -> None:
        if "id" not in entry:
            entry["id"] = self._new_id()
        file_name = f"{entry['id']}.json"
        if write_to_disk:
            path = self.base_dir / file_name
            self._write_atomic(path, entry)
        if self.storage == "pgvector":
            self._pg_upsert(entry)
        else:
            self._db_upsert(entry)
        self.data.append(entry)

    def _get_embedding(self, text: str) -> List[float]:
        text = text.strip()
        if not text:
            return []
        if text in self._embed_cache:
            vec = self._embed_cache.pop(text)
            self._embed_cache[text] = vec
            return vec
        # 1) 외부 임베딩 API 사용 시 (with retry)
        embed_url = os.environ.get("KB_EMBED_URL")
        if requests is not None and embed_url:
            last_err = None
            for attempt in range(2):
                try:
                    resp = requests.post(
                        embed_url,
                        json={"text": text},
                        timeout=5,
                    )
                    resp.raise_for_status()
                    vec = resp.json().get("vector") or []
                    out = [float(v) for v in vec]
                    if self._embed_cache_max > 0:
                        self._embed_cache[text] = out
                        while len(self._embed_cache) > self._embed_cache_max:
                            self._embed_cache.popitem(last=False)
                    return out
                except Exception as e:
                    last_err = e
                    if attempt == 0:
                        _rag_logger.warning("Embedding API failed (attempt 1), retrying: %s", e)
                        _time.sleep(1.5)
            _rag_logger.warning(
                "Embedding API failed after retries, using local fallback: %s", last_err
            )
        # 2) 로컬 sentence-transformers 모델
        try:
            if not hasattr(self, "_st_model"):
                from sentence_transformers import SentenceTransformer
                self._st_model = SentenceTransformer("all-MiniLM-L6-v2")
                _rag_logger.info("Loaded local embedding model: all-MiniLM-L6-v2")
            if hasattr(self, "_st_model") and self._st_model is not None:
                vec = self._st_model.encode(text).tolist()
                out = [float(v) for v in vec]
                if self._embed_cache_max > 0:
                    self._embed_cache[text] = out
                    while len(self._embed_cache) > self._embed_cache_max:
                        self._embed_cache.popitem(last=False)
                return out
        except ImportError:
            self._st_model = None
        except Exception as e:
            _rag_logger.warning("Local embedding model failed: %s", e)
            self._st_model = None
        # 3) 로컬 난수 기반 임베딩 (동일 입력 -> 동일 벡터 위해 seed 고정)
        seed = abs(hash(text)) % (2**32)
        rng = np.random.default_rng(seed)
        out = rng.normal(size=64).astype(float).tolist()
        if self._embed_cache_max > 0:
            self._embed_cache[text] = out
            while len(self._embed_cache) > self._embed_cache_max:
                self._embed_cache.popitem(last=False)
        return out

    # ---------------- 퍼블릭 API ----------------

    def search(
        self,
        error_msg: str,
        top_k: int = 3,
        *,
        tags: Optional[List[str]] = None,
        role: Optional[str] = None,
        stage: Optional[str] = None,
        categories: Optional[List[str]] = None,
        category: Optional[str] = None,
    ) -> List[Dict[str, Any]]:
        """
        에러 메시지를 기반으로 과거 성공 패턴 검색
        """
        started = _time.perf_counter()
        if self.storage == "pgvector" and self._pg_ok:
            norm_cats = [str(c) for c in (categories or []) if str(c).strip()]
            if category:
                norm_cats.append(str(category))
            query = _normalize_message(error_msg)
            req_ids = _extract_req_ids_from_text(query)
            q_vec = self._get_embedding(query)
            rows = self._pg_search(
                q_vec,
                query=query,
                top_k=top_k,
                tags=tags,
                role=role,
                stage=stage,
                categories=norm_cats or None,
                req_ids=req_ids,
            )
            if _RAG_PERF_LOG:
                _rag_logger.info(
                    "rag_search storage=%s entries=%d query_chars=%d top_k=%d hits=%d elapsed_ms=%.1f",
                    self.storage,
                    len(self.data),
                    len(query),
                    top_k,
                    len(rows),
                    (_time.perf_counter() - started) * 1000.0,
                )
            return rows

        if not self.data:
            return []

        query = _normalize_message(error_msg)
        if not query:
            return []

        q_vec = np.array(self._get_embedding(query), dtype=float)
        if q_vec.size == 0:
            return []

        norm_tags = [str(t) for t in (tags or []) if str(t).strip()]
        norm_categories = [str(c) for c in (categories or []) if str(c).strip()]
        if category:
            norm_categories.append(str(category))
        role = str(role) if role else None
        stage = str(stage) if stage else None

        results: List[Dict[str, Any]] = []
        project_boost = float(getattr(config, "RAG_PROJECT_BOOST", 0.0))
        recency_days = float(getattr(config, "RAG_RECENCY_DAYS", 0))
        recency_boost = float(getattr(config, "RAG_RECENCY_BOOST", 0.0))
        apply_boost = float(getattr(config, "RAG_APPLY_COUNT_BOOST", 0.0))
        error_boost = float(getattr(config, "RAG_ERROR_COUNT_BOOST", 0.0))
        req_ids = _extract_req_ids_from_text(query)
        exact_boost = float(getattr(config, "RAG_EXACT_MATCH_BOOST", 0.4))
        for idx, ent in enumerate(self.data):
            v = ent.get("vector") or self._get_embedding(ent["error_clean"])
            v_arr = np.array(v, dtype=float)
            score = _cosine(q_vec, v_arr) * float(ent.get("weight", 1.0))
            if role and ent.get("role") == role:
                score += 0.15
            if stage and ent.get("stage") == stage:
                score += 0.1
            if norm_categories:
                ent_cat = str(ent.get("category") or "")
                if ent_cat in norm_categories:
                    score += 0.12
                else:
                    continue
            if norm_tags:
                ent_tags = set(ent.get("tags") or [])
                hit = len(ent_tags.intersection(norm_tags))
                if hit:
                    score += 0.05 * hit
            if recency_days > 0 and ent.get("timestamp"):
                try:
                    ts = datetime.fromisoformat(str(ent.get("timestamp")))
                    delta_days = (datetime.utcnow() - ts).total_seconds() / 86400.0
                    if delta_days < recency_days:
                        score += recency_boost * (1.0 - (delta_days / recency_days))
                except Exception:
                    pass
            if project_boost > 0 and ent.get("project_root"):
                if str(ent.get("project_root")) in error_msg:
                    score += project_boost
            if apply_boost > 0:
                try:
                    score += apply_boost * float(ent.get("apply_count") or 0)
                except Exception:
                    pass
            if error_boost > 0:
                try:
                    score += error_boost * float(ent.get("error_count") or 0)
                except Exception:
                    pass
            if req_ids:
                hay = " ".join(
                    [
                        str(ent.get("error_raw") or ""),
                        str(ent.get("error_clean") or ""),
                        str(ent.get("context") or ""),
                        str(ent.get("source_file") or ""),
                        json.dumps(ent.get("metadata") or {}, ensure_ascii=False),
                    ]
                )
                if any(rid in hay for rid in req_ids):
                    score += exact_boost
            if score <= 0.0:
                continue
            item = dict(ent)
            item["index"] = idx
            item["score"] = score
            results.append(item)

        results.sort(key=lambda x: x["score"], reverse=True)
        rows = results[:top_k]
        if _RAG_PERF_LOG:
            _rag_logger.info(
                "rag_search storage=%s entries=%d query_chars=%d top_k=%d hits=%d elapsed_ms=%.1f",
                self.storage,
                len(self.data),
                len(query),
                top_k,
                len(rows),
                (_time.perf_counter() - started) * 1000.0,
            )
        return rows

    def learn(
        self,
        error_msg: str,
        fix_pattern: str,
        tags: Optional[List[str]] = None,
        success: bool = True,
        *,
        role: Optional[str] = None,
        stage: Optional[str] = None,
        context: Optional[str] = None,
        category: Optional[str] = None,
        project_root: Optional[str] = None,
    ) -> None:
        """
        성공한 수정 패턴을 지식 베이스에 저장
        """
        if not success:
            # 실패한 패턴은 기본적으로 저장하지 않음
            return

        ctx = _normalize_message(error_msg)
        if not ctx:
            return

        # 중복 패턴 간단 필터링 (동일 error_clean + 유사 fix)
        for ent in self.data:
            if ent.get("error_clean") == ctx and ent.get("fix") == fix_pattern:
                if role and ent.get("role") not in (None, role):
                    continue
                if stage and ent.get("stage") not in (None, stage):
                    continue
                # 이미 동일 패턴 존재 → weight만 조금 올리고 종료
                ent["weight"] = float(ent.get("weight", 1.0)) + 0.1
                ent["apply_count"] = int(ent.get("apply_count", 0)) + 1
                ent["error_count"] = int(ent.get("error_count", 0)) + 1
                if project_root:
                    ent["project_root"] = project_root
                path = self.base_dir / ent["source_file"]
                self._write_atomic(path, ent)
                if self.storage == "pgvector":
                    self._pg_upsert(ent)
                else:
                    self._db_upsert(ent)
                return

        vec = self._get_embedding(ctx)
        entry: Dict[str, Any] = {
            "id": self._new_id(),
            "error_raw": error_msg,
            "error_clean": ctx,
            "fix": fix_pattern,
            "tags": tags or [],
            "role": str(role) if role else None,
            "stage": str(stage) if stage else None,
            "context": str(context) if context else "",
            "category": str(category).strip() if category else "general",
            "vector": vec,
            "weight": 1.0,
            "apply_count": 1,
            "error_count": 1,
            "timestamp": datetime.utcnow().isoformat(),
            "source_file": "",  # _append_new_entry 에서 채움
            "project_root": project_root or "",
        }
        self._append_new_entry(entry, write_to_disk=True)
        self._enforce_max_entries()

    def _enforce_max_entries(self) -> None:
        if self._max_entries <= 0 or len(self.data) <= self._max_entries:
            return
        sorted_entries = sorted(
            self.data,
            key=lambda e: (float(e.get("weight", 1.0)), e.get("timestamp", "")),
        )
        to_remove = sorted_entries[: len(self.data) - self._max_entries]
        remove_ids = {e.get("id") for e in to_remove}
        for ent in to_remove:
            sf = ent.get("source_file")
            if sf:
                p = self.base_dir / sf
                try:
                    if p.exists():
                        p.unlink()
                except OSError:
                    pass
        self.data = [e for e in self.data if e.get("id") not in remove_ids]

    def stats(self) -> Dict[str, Any]:
        by_category: Dict[str, int] = {}
        by_source: Dict[str, int] = {}
        source_latest: Dict[str, str] = {}
        for ent in self.data:
            cat = str(ent.get("category") or "general")
            by_category[cat] = by_category.get(cat, 0) + 1
            src = str(ent.get("source_file") or "unknown")
            by_source[src] = by_source.get(src, 0) + 1
            ts = str(ent.get("timestamp") or "")
            if ts:
                prev = source_latest.get(src, "")
                if not prev or ts > prev:
                    source_latest[src] = ts
        source_list = [
            {"source": k, "count": v, "last_ts": source_latest.get(k, "")}
            for k, v in by_source.items()
        ]
        source_list.sort(key=lambda x: x.get("count", 0), reverse=True)
        category_list = [{"category": k, "count": v} for k, v in by_category.items()]
        category_list.sort(key=lambda x: x.get("count", 0), reverse=True)
        return {
            "total": len(self.data),
            "by_category": by_category,
            "by_source": by_source,
            "source_list": source_list,
            "category_list": category_list,
        }

    def feedback(self, index: int, positive: bool = True) -> None:
        """
        나중에 UI에서 thumbs-up/down 같은 피드백 연결용 훅
        """
        if index < 0 or index >= len(self.data):
            return
        ent = self.data[index]
        delta = 0.2 if positive else -0.2
        ent["weight"] = max(0.1, float(ent.get("weight", 1.0)) + delta)
        ent["apply_count"] = int(ent.get("apply_count", 0)) + (1 if positive else 0)
        path = self.base_dir / ent["source_file"]
        self._write_atomic(path, ent)
        if self.storage == "pgvector":
            self._pg_upsert(ent)
        else:
            self._db_upsert(ent)


def _collect_files_from_paths(
    paths: Iterable[str],
    *,
    exts: Optional[Tuple[str, ...]] = None,
    globs: Optional[List[str]] = None,
    max_files: int = 200,
) -> List[Path]:
    files: List[Path] = []
    for p in paths:
        try:
            path = Path(p).expanduser()
        except Exception:
            continue
        if path.is_file():
            files.append(path)
            continue
        if path.is_dir():
            if globs:
                for g in globs:
                    for hit in path.glob(g):
                        if hit.is_file():
                            files.append(hit)
            else:
                for hit in path.rglob("*"):
                    if hit.is_file():
                        files.append(hit)
    if exts:
        files = [f for f in files if f.suffix.lower() in exts]
    # dedup + cap
    uniq: Dict[str, Path] = {}
    for f in files:
        uniq[str(f.resolve())] = f
    return list(uniq.values())[: max(1, int(max_files))]


def _infer_vectorcast_tags(path: Path) -> List[str]:
    name = path.name.lower()
    tags = ["vectorcast"]
    if "ut" in name or "unit" in name:
        tags.append("ut")
    if "it" in name or "integration" in name:
        tags.append("it")
    if "coverage" in name:
        tags.append("coverage")
    return tags


def ingest_external_sources(kb: KnowledgeBase, cfg: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    cfg = cfg or {}
    enabled = bool(getattr(config, "RAG_INGEST_ENABLE", True))
    if not enabled:
        return {"enabled": False, "reason": "disabled"}

    max_files = int(cfg.get("rag_ingest_max_files") or getattr(config, "RAG_INGEST_MAX_FILES", 200))
    max_chunks = int(
        cfg.get("rag_ingest_max_chunks") or getattr(config, "RAG_INGEST_MAX_CHUNKS_PER_FILE", 12)
    )
    chunk_size = int(cfg.get("rag_chunk_size") or getattr(config, "RAG_CHUNK_SIZE", 1200))
    overlap = int(cfg.get("rag_chunk_overlap") or getattr(config, "RAG_CHUNK_OVERLAP", 200))

    vc_paths = _split_paths(cfg.get("vc_reports_paths") or getattr(config, "VC_REPORTS_PATHS", ""))
    uds_paths = _split_paths(cfg.get("uds_spec_paths") or getattr(config, "UDS_SPEC_PATHS", ""))
    req_paths = _split_paths(cfg.get("req_docs_paths") or getattr(config, "REQ_DOCS_PATHS", ""))
    code_paths = _split_paths(cfg.get("codebase_paths") or getattr(config, "CODEBASE_PATHS", ""))

    idx = kb._load_external_index()
    updated = 0
    skipped = 0

    def _ingest(
        paths: List[str],
        *,
        category: str,
        tags: Any,
        exts: Tuple[str, ...],
        globs: Optional[List[str]] = None,
    ):
        nonlocal updated, skipped, idx
        files = _collect_files_from_paths(paths, exts=exts, globs=globs, max_files=max_files)
        for fp in files:
            try:
                sig = f"{fp.stat().st_mtime_ns}:{fp.stat().st_size}"
            except Exception:
                sig = ""
            key = f"{category}:{fp.as_posix()}"
            if sig and idx.get(key) == sig:
                skipped += 1
                continue

            chunks = _chunk_source_file(
                fp,
                chunk_size=chunk_size,
                overlap=overlap,
                max_chunks=max_chunks,
            )
            if not chunks:
                skipped += 1
                continue
            use_tags: List[str] = []
            try:
                if callable(tags):
                    use_tags = list(tags(fp))
                elif isinstance(tags, list):
                    use_tags = list(tags)
            except Exception:
                use_tags = []

            for i, ch in enumerate(chunks):
                title = f"{category}:{fp.name}#{i+1}"
                kb.add_document(
                    title=title,
                    content=ch,
                    category=category,
                    tags=use_tags,
                    source_file=str(fp),
                )
                updated += 1
            if sig:
                idx[key] = sig

    if vc_paths:
        _ingest(
            vc_paths,
            category="vectorcast",
            tags=_infer_vectorcast_tags,
            exts=(".html", ".htm", ".csv", ".txt", ".log"),
        )
    if uds_paths:
        _ingest(
            uds_paths,
            category="uds",
            tags=["uds"],
            exts=(".pdf", ".docx", ".txt", ".md", ".xlsx"),
        )
    if req_paths:
        _ingest(
            req_paths,
            category="requirements",
            tags=["requirements"],
            exts=(".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx"),
        )
    if code_paths:
        _ingest(
            code_paths,
            category="code",
            tags=["code"],
            exts=(".c", ".h", ".cpp", ".hpp"),
            globs=list(getattr(config, "CODE_RAG_GLOBS", [])) or None,
        )

    if updated or skipped:
        kb._save_external_index(idx)
    return {"enabled": True, "updated": updated, "skipped": skipped}


def ingest_uds_reference(
    kb: KnowledgeBase,
    *,
    ref_suds_path: Optional[str] = None,
    function_details: Optional[Dict[str, Any]] = None,
    globals_info_map: Optional[Dict[str, Any]] = None,
    req_map: Optional[Dict[str, Any]] = None,
) -> Dict[str, int]:
    """Ingest UDS-specific data into RAG KB with specialized categories."""
    counts = {"uds_description": 0, "uds_globals": 0, "uds_requirements": 0, "uds_reference": 0}

    if ref_suds_path:
        rpath = Path(ref_suds_path)
        if rpath.exists() and rpath.suffix.lower() == ".docx":
            chunks = _chunk_source_file(rpath, chunk_size=1500, overlap=300, max_chunks=50)
            for i, ch in enumerate(chunks):
                kb.add_document(
                    title=f"uds_reference:{rpath.name}#{i+1}",
                    content=ch,
                    category="uds_description",
                    tags=["uds_description", "reference"],
                    source_file=str(rpath),
                )
                counts["uds_reference"] += 1

    if isinstance(function_details, dict):
        batch_lines: List[str] = []
        for fid, info in function_details.items():
            if not isinstance(info, dict):
                continue
            desc = str(info.get("description") or "").strip()
            dsrc = str(info.get("description_source") or "").strip()
            if dsrc in {"comment", "sds", "reference"} and desc and len(desc) > 10:
                fname = str(info.get("name") or "").strip()
                proto = str(info.get("prototype") or "").strip()[:100]
                batch_lines.append(f"{fname}: {desc} [{proto}]")
                if len(batch_lines) >= 20:
                    kb.add_document(
                        title=f"uds_description:batch_{counts['uds_description']}",
                        content="\n".join(batch_lines),
                        category="uds_description",
                        tags=["uds_description"],
                    )
                    counts["uds_description"] += 1
                    batch_lines = []
        if batch_lines:
            kb.add_document(
                title=f"uds_description:batch_{counts['uds_description']}",
                content="\n".join(batch_lines),
                category="uds_description",
                tags=["uds_description"],
            )
            counts["uds_description"] += 1

    if isinstance(globals_info_map, dict) and globals_info_map:
        globals_lines: List[str] = []
        for gname, ginfo in globals_info_map.items():
            if not isinstance(ginfo, dict):
                continue
            gtype = str(ginfo.get("type") or "").strip()
            gfile = Path(str(ginfo.get("file") or "")).name
            gdesc = str(ginfo.get("desc") or "").strip()
            globals_lines.append(f"{gname} ({gtype}) [{gfile}] {gdesc}".strip())
            if len(globals_lines) >= 30:
                kb.add_document(
                    title=f"uds_globals:batch_{counts['uds_globals']}",
                    content="\n".join(globals_lines),
                    category="uds_globals",
                    tags=["uds_globals", "globals"],
                )
                counts["uds_globals"] += 1
                globals_lines = []
        if globals_lines:
            kb.add_document(
                title=f"uds_globals:batch_{counts['uds_globals']}",
                content="\n".join(globals_lines),
                category="uds_globals",
                tags=["uds_globals", "globals"],
            )
            counts["uds_globals"] += 1

    if isinstance(req_map, dict) and req_map:
        req_lines: List[str] = []
        for rid, rinfo in req_map.items():
            if isinstance(rinfo, dict):
                rdesc = str(rinfo.get("description") or rinfo.get("text") or "").strip()
                req_lines.append(f"{rid}: {rdesc[:200]}")
            elif isinstance(rinfo, str):
                req_lines.append(f"{rid}: {rinfo[:200]}")
            if len(req_lines) >= 25:
                kb.add_document(
                    title=f"uds_requirements:batch_{counts['uds_requirements']}",
                    content="\n".join(req_lines),
                    category="uds_requirements",
                    tags=["uds_requirements", "requirements"],
                )
                counts["uds_requirements"] += 1
                req_lines = []
        if req_lines:
            kb.add_document(
                title=f"uds_requirements:batch_{counts['uds_requirements']}",
                content="\n".join(req_lines),
                category="uds_requirements",
                tags=["uds_requirements", "requirements"],
            )
            counts["uds_requirements"] += 1

    _rag_logger.info("UDS reference ingestion: %s", counts)
    return counts


def ingest_runtime_summary(kb: KnowledgeBase, summary: Dict[str, Any], report_dir: Path) -> Dict[str, Any]:
    if not bool(getattr(config, "RAG_INGEST_RUNTIME_ENABLE", True)):
        return {"enabled": False, "reason": "disabled"}
    updated = 0

    def _add_runtime(category: str, title: str, content: str, source_file: str) -> None:
        nonlocal updated
        kb.add_document(
            title=title,
            content=content[: int(getattr(config, "RAG_CONTEXT_MAX_CHARS", 4000))],
            category=category,
            tags=[category, "runtime"],
            source_file=source_file,
        )
        updated += 1

    build = summary.get("build", {}) if isinstance(summary.get("build"), dict) else {}
    coverage = summary.get("coverage", {}) if isinstance(summary.get("coverage"), dict) else {}
    tests = summary.get("tests", {}) if isinstance(summary.get("tests"), dict) else {}

    build_ctx = json.dumps(
        {
            "reason": build.get("reason"),
            "data": build.get("data", {}),
        },
        ensure_ascii=False,
    )
    _add_runtime("build", "runtime:build", build_ctx, "runtime:build")

    tests_ctx = json.dumps(
        {
            "enabled": tests.get("enabled"),
            "mode": tests.get("mode"),
            "results": tests.get("results", [])[:50],
            "execution": tests.get("execution", {}),
        },
        ensure_ascii=False,
    )
    _add_runtime("tests", "runtime:tests", tests_ctx, "runtime:tests")

    cov_ctx = json.dumps(
        {
            "enabled": coverage.get("enabled"),
            "ok": coverage.get("ok"),
            "line_rate": coverage.get("line_rate"),
            "line_rate_pct": coverage.get("line_rate_pct"),
            "branch_rate": coverage.get("branch_rate"),
            "branch_rate_pct": coverage.get("branch_rate_pct"),
            "threshold": coverage.get("threshold"),
            "reason": coverage.get("reason") or coverage.get("parse_error"),
        },
        ensure_ascii=False,
    )
    _add_runtime("coverage", "runtime:coverage", cov_ctx, "runtime:coverage")

    return {"ok": True, "updated": updated, "report_dir": str(report_dir)}


def get_kb(report_dir: Path) -> KnowledgeBase:
    """
    report_dir 기준 RAG 저장소 인스턴스 반환

    기존: report_dir / "knowledge_base.json"
    변경: report_dir / config.KB_DIR_NAME (디렉터리)
    """
    kb_dir_name = getattr(config, "KB_DIR_NAME", "kb_store")
    global_dir = str(getattr(config, "KB_GLOBAL_DIR", "") or "").strip()
    if global_dir:
        base_dir = Path(global_dir).expanduser().resolve()
    else:
        base_dir = Path(report_dir) / kb_dir_name
    return KnowledgeBase(base_dir)
