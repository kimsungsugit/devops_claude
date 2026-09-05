# workflow/rag/chunker.py
# -*- coding: utf-8 -*-
# RAG chunking / file reading utilities (extracted from workflow.rag)

from __future__ import annotations

import json
import os
import re
import tempfile
import threading
from collections import OrderedDict
from pathlib import Path
from typing import List, Optional, Tuple

# ── 대용량 문서 텍스트 추출 캐시 ────────────────────────────────────────────
#
# 실측(2026-08-03): UDS 생성 한 번이 **같은 38.8MB 참조 SUDS 를 세 경로에서 각각**
# 읽는다 — `docx_builder` 의 함수정보 파싱(7.4s), `helpers/uds.py` 의 SwCom diff,
# `local.py`/`jenkins.py` 의 AI 예시 텍스트(`_read_text_from_file` 단독 **11.7초**).
# 서로를 모르는 세 호출처라 아무도 중복을 못 본다.
#
# 반환값이 **불변 문자열**이라 이 저장소가 겪은 캐시 사고(공유 가변 구조를 호출자가
# 제자리 변경)와 무관하다. 무효화는 `(정규화 경로, mtime_ns, size)` 로 한다 —
# 내용 해시가 아니라 파일 신원이므로 stale 은 파일이 안 바뀐 경우뿐이다.
_TEXT_CACHE: "OrderedDict[Tuple[str, int, int], str]" = OrderedDict()
_TEXT_CACHE_MAX = 4
_TEXT_CACHE_MIN_BYTES = 1_000_000  # 작은 파일은 캐시 이득이 없고 항목만 밀어낸다
_TEXT_CACHE_LOCK = threading.Lock()


def _text_cache_key(path: Path) -> Optional[Tuple[str, int, int]]:
    """캐시 자격과 키. 자격 미달이면 `None`(= 캐시 안 함)."""
    try:
        st = path.stat()
        resolved = os.path.normcase(str(path.resolve()))
    except OSError:
        return None
    if st.st_size < _TEXT_CACHE_MIN_BYTES:
        return None
    # ⚠ `tempfile.NamedTemporaryFile` 은 temp **루트 직속**에 파일을 만들고 OS 가 그
    #   경로를 재사용한다. 업로드본을 tmp 로 받아 읽는 호출처가 여럿이라
    #   (`local.py:1042` 등) 경로·크기가 겹치면 남의 파일 내용을 돌려줄 수 있다.
    #   루트 **직속**만 뺀다 — 하위 디렉터리(pytest tmp_path 등)는 매번 새로 만들어져
    #   재사용되지 않으므로, 넓게 막으면 캐시가 검증 불가능해지기만 한다.
    try:
        if os.path.normcase(str(path.resolve().parent)) == os.path.normcase(
            str(Path(tempfile.gettempdir()).resolve())
        ):
            return None
    except OSError:
        return None
    return (resolved, st.st_mtime_ns, st.st_size)


def _read_text_from_file(path: Path) -> str:
    """파일에서 텍스트를 뽑는다. 큰 문서는 파일 신원 기준으로 캐시한다."""
    key = _text_cache_key(path)
    if key is None:
        return _read_text_uncached(path)
    with _TEXT_CACHE_LOCK:
        hit = _TEXT_CACHE.get(key)
        if hit is not None:
            _TEXT_CACHE.move_to_end(key)
            return hit
    # 락 밖에서 읽는다 — 11초짜리 추출을 락 안에서 하면 다른 요청이 통째로 막힌다.
    # 동시 miss 는 중복 작업일 뿐 결과를 바꾸지 않는다(문자열은 불변).
    text = _read_text_uncached(path)
    with _TEXT_CACHE_LOCK:
        _TEXT_CACHE[key] = text
        _TEXT_CACHE.move_to_end(key)
        while len(_TEXT_CACHE) > _TEXT_CACHE_MAX:
            _TEXT_CACHE.popitem(last=False)
    return text


# 이 파서가 실제로 본문을 뽑을 수 있는 확장자 — **단일 출처**.
# 아래 분기와 반드시 함께 움직여야 한다(test_req_doc_format_gate.py 가 대조한다).
# 호출부가 "읽을 수 있는가"를 미리 알아야 하는 이유: cloudium 에서는 본문 추출 전에
# worker IPC 로 바이트를 통째로 끌어온다. 못 읽을 형식을 걸러내지 않으면 수십 MB를
# 받아 놓고 ""를 돌려주고, 사용자는 "본문 0자"라는 **원인과 무관한** 사유를 본다.
SUPPORTED_TEXT_EXTS = frozenset({
    ".txt", ".md", ".csv", ".log", ".json", ".xml", ".yaml", ".yml",
    ".html", ".htm",
    ".pdf",
    ".docx",
})


def _read_text_uncached(path: Path) -> str:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_TEXT_EXTS:
        return ""
    try:
        if ext in (".txt", ".md", ".csv", ".log", ".json", ".xml", ".yaml", ".yml"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext in (".html", ".htm"):
            try:
                from bs4 import BeautifulSoup  # type: ignore
                html = path.read_text(encoding="utf-8", errors="ignore")
                return BeautifulSoup(html, "html.parser").get_text("\n")
            except Exception:
                return path.read_text(encoding="utf-8", errors="ignore")
        if ext in (".pdf",):
            try:
                try:
                    import pdfplumber  # type: ignore
                except Exception:
                    pdfplumber = None  # type: ignore
                texts = []
                if pdfplumber:
                    with pdfplumber.open(str(path)) as pdf:
                        for idx, page in enumerate(pdf.pages, start=1):
                            page_text = page.extract_text() or ""
                            if page_text:
                                texts.append(f"=== Page {idx} ===")
                                texts.append(page_text)
                            try:
                                tables = page.extract_tables() or []
                            except Exception:
                                tables = []
                            for t in tables:
                                rows = []
                                for row in t:
                                    if not row:
                                        continue
                                    cells = [str(c or "").strip() for c in row]
                                    if any(cells):
                                        rows.append(" | ".join(cells))
                                if rows:
                                    texts.append("=== Table ===")
                                    texts.extend(rows)
                    return "\n".join(texts)
                from pypdf import PdfReader  # type: ignore
                reader = PdfReader(str(path))
                texts = []
                for idx, p in enumerate(reader.pages, start=1):
                    try:
                        page_text = p.extract_text() or ""
                    except Exception:
                        page_text = ""
                    if page_text:
                        texts.append(f"=== Page {idx} ===")
                        texts.append(page_text)
                return "\n".join(texts)
            except Exception:
                return ""
        if ext in (".docx",):
            try:
                import docx  # type: ignore
                doc = docx.Document(str(path))
                lines = []
                for p in doc.paragraphs:
                    text = (p.text or "").strip()
                    if not text:
                        continue
                    style = str(getattr(p, "style", "") or "")
                    style_name = ""
                    try:
                        style_name = p.style.name  # type: ignore
                    except Exception:
                        style_name = str(style)
                    if "Heading" in style_name:
                        level = re.findall(r"\d+", style_name)
                        prefix = "#" * int(level[0]) if level else "##"
                        lines.append(f"{prefix} {text}")
                    elif "TOC" in style_name or "Table of Contents" in style_name:
                        lines.append(f"TOC: {text}")
                    else:
                        lines.append(text)
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        lines.append("=== Table ===")
                        lines.extend(rows)
                return "\n".join(lines)
            except Exception:
                return ""
    except Exception:
        return ""
    return ""


def _chunk_text(text: str, *, chunk_size: int = 1200, overlap: int = 200) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    if chunk_size <= 0:
        return [text]
    out: List[str] = []
    step = max(1, chunk_size - max(0, overlap))
    for i in range(0, len(text), step):
        out.append(text[i : i + chunk_size])
    return [c for c in out if c.strip()]


REQ_ID_PATTERN = re.compile(r"\b(?:REQ|SDS|SW|SWS|SWR|SRS|SWC|REQS)\s*[-_:]?\s*[A-Za-z0-9_.-]+\b")


def _extract_req_ids_from_text(text: str) -> List[str]:
    if not text:
        return []
    ids = [m.group(0).replace(" ", "").strip() for m in REQ_ID_PATTERN.finditer(text)]
    uniq: List[str] = []
    seen = set()
    for rid in ids:
        if not rid or rid in seen:
            continue
        seen.add(rid)
        uniq.append(rid)
    return uniq


def _chunk_by_req_ids(text: str, *, chunk_size: int, overlap: int) -> List[str]:
    text = (text or "").strip()
    if not text:
        return []
    matches = list(REQ_ID_PATTERN.finditer(text))
    if len(matches) < 2:
        return _chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    chunks: List[str] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        seg = text[start:end].strip()
        if not seg:
            continue
        if len(seg) > chunk_size * 2:
            chunks.extend(_chunk_text(seg, chunk_size=chunk_size, overlap=overlap))
        else:
            chunks.append(seg)
    return [c for c in chunks if c.strip()]


def _chunk_docx_by_heading(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    try:
        import docx  # type: ignore
    except Exception:
        return []
    try:
        doc = docx.Document(str(path))
    except Exception:
        return []
    sections: List[Tuple[str, List[str]]] = []
    current_title = ""
    current_lines: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = p.style.name  # type: ignore
        except Exception:
            style_name = str(getattr(p, "style", "") or "")
        if "Heading" in style_name:
            if current_lines:
                sections.append((current_title, current_lines))
            current_title = text
            current_lines = []
            continue
        current_lines.append(text)
    if current_lines:
        sections.append((current_title, current_lines))
    chunks: List[str] = []
    for title, lines in sections:
        block = "\n".join([title.strip(), "\n".join(lines).strip()]).strip()
        if not block:
            continue
        chunks.extend(_chunk_by_req_ids(block, chunk_size=chunk_size, overlap=overlap))
    return [c for c in chunks if c.strip()]


def _chunk_xlsx_rows(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    try:
        import pandas as pd  # type: ignore
    except Exception:
        return []
    chunks: List[str] = []
    try:
        sheets = pd.read_excel(str(path), sheet_name=None)
    except Exception:
        return []
    for sheet_name, df in sheets.items():
        if df is None:
            continue
        try:
            records = df.fillna("").to_dict(orient="records")
        except Exception:
            continue
        for idx, row in enumerate(records):
            payload = {"sheet": sheet_name, "row_index": idx + 1, "data": row}
            text = json.dumps(payload, ensure_ascii=False)
            chunks.extend(_chunk_text(text, chunk_size=chunk_size, overlap=overlap))
    return [c for c in chunks if c.strip()]


def _chunk_c_by_function(path: Path, *, chunk_size: int, overlap: int) -> List[str]:
    """AST-based chunking for C/H files: one chunk per function definition."""
    try:
        from workflow.code_parser.c_parser import parse_c_project
    except ImportError:
        return []
    try:
        src = path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return []
    chunks: List[str] = []
    func_pattern = re.compile(
        r"(?:^|\n)"
        r"((?:/\*[\s\S]*?\*/\s*|//[^\n]*\n\s*)*)"
        r"((?:static\s+|inline\s+|extern\s+|const\s+)*\w[\w\s*]+\s+\w+\s*\([^)]*\)\s*\{)",
        re.MULTILINE,
    )
    last_end = 0
    for m in func_pattern.finditer(src):
        start = m.start()
        brace_count = 0
        body_start = src.find("{", m.start(2))
        if body_start == -1:
            continue
        pos = body_start
        while pos < len(src):
            if src[pos] == "{":
                brace_count += 1
            elif src[pos] == "}":
                brace_count -= 1
                if brace_count == 0:
                    func_text = src[start:pos + 1].strip()
                    if len(func_text) > chunk_size:
                        func_text = func_text[:chunk_size]
                    if func_text:
                        chunks.append(f"[{path.name}]\n{func_text}")
                    last_end = pos + 1
                    break
            pos += 1
    if not chunks:
        return _chunk_text(src, chunk_size=chunk_size, overlap=overlap)
    return chunks


def _chunk_source_file(
    path: Path,
    *,
    chunk_size: int,
    overlap: int,
    max_chunks: int,
) -> List[str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        chunks = _chunk_docx_by_heading(path, chunk_size=chunk_size, overlap=overlap)
    elif ext == ".xlsx":
        chunks = _chunk_xlsx_rows(path, chunk_size=chunk_size, overlap=overlap)
    elif ext == ".pdf":
        text = _read_text_from_file(path)
        chunks = _chunk_by_req_ids(text, chunk_size=chunk_size, overlap=overlap)
    elif ext in {".c", ".h", ".cpp", ".hpp"}:
        chunks = _chunk_c_by_function(path, chunk_size=chunk_size, overlap=overlap)
    else:
        text = _read_text_from_file(path)
        chunks = _chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    return chunks[:max(1, int(max_chunks))]


def _read_and_chunk_file(
    path: Path,
    *,
    chunk_size: int = 1200,
    overlap: int = 200,
    max_chunks: int = 12,
) -> List[str]:
    if not path or not Path(path).exists():
        return []
    return _chunk_source_file(
        Path(path),
        chunk_size=int(chunk_size or 1200),
        overlap=int(overlap or 0),
        max_chunks=int(max_chunks or 1),
    )
