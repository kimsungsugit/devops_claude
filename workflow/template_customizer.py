# workflow/template_customizer.py
"""DOCX Template Customization - detect and map placeholders in user templates.

Scans uploaded DOCX templates for {{placeholder}} patterns and provides
a mapping interface for the UDS generation pipeline.
"""

from __future__ import annotations

import re
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Set

logger = logging.getLogger(__name__)

_PLACEHOLDER_PAT = re.compile(r"\{\{([^}]+)\}\}")

KNOWN_UDS_PLACEHOLDERS = {
    "project_name", "overview", "requirements", "interfaces",
    "uds_frames", "notes", "software_unit_design", "unit_structure",
    "global_data", "interface_functions", "internal_functions",
    "function_table", "function_details", "generated_at",
    "call_map", "global_vars", "static_vars", "macro_defs",
}


def scan_template_placeholders(docx_path: str) -> Dict[str, Any]:
    """Scan a DOCX template and return all {{placeholder}} tokens found."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed", "placeholders": []}

    doc = Document(docx_path)
    found: Set[str] = set()
    locations: List[Dict[str, str]] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text or ""
        for m in _PLACEHOLDER_PAT.finditer(text):
            token = m.group(1).strip()
            found.add(token)
            locations.append({
                "placeholder": token,
                "location": "paragraph",
                "index": i,
                "context": text[:100],
            })

    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text or ""
                for m in _PLACEHOLDER_PAT.finditer(text):
                    token = m.group(1).strip()
                    found.add(token)
                    locations.append({
                        "placeholder": token,
                        "location": "table",
                        "table": table_idx,
                        "row": row_idx,
                        "cell": cell_idx,
                        "context": text[:100],
                    })

    known_found = found & KNOWN_UDS_PLACEHOLDERS
    custom_found = found - KNOWN_UDS_PLACEHOLDERS

    return {
        "file": Path(docx_path).name,
        "total_placeholders": len(found),
        "known_placeholders": sorted(known_found),
        "custom_placeholders": sorted(custom_found),
        "locations": locations,
        "is_valid_template": len(known_found) >= 2,
    }


def build_placeholder_mapping(
    template_info: Dict[str, Any],
    user_mapping: Optional[Dict[str, str]] = None,
) -> Dict[str, str]:
    """Build a complete placeholder -> value-source mapping.

    Known placeholders map to standard UDS fields. Custom placeholders
    can be mapped via user_mapping or left as empty strings.
    """
    mapping: Dict[str, str] = {}

    for ph in template_info.get("known_placeholders", []):
        mapping[ph] = f"uds.{ph}"

    for ph in template_info.get("custom_placeholders", []):
        if user_mapping and ph in user_mapping:
            mapping[ph] = user_mapping[ph]
        else:
            mapping[ph] = ""

    return mapping
