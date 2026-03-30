"""report_gen.docx_builder - Auto-split from report_generator.py"""
# Re-import common dependencies
import re

# Payload field name constants (canonical source: report_gen.uds_generator)
# Function-level (per-function, List[str]):
#   KEY_FN_GLOBALS = "globals_global"  — global vars used by the function
#   KEY_FN_STATICS = "globals_static"  — static vars used by the function
# Module-level (top-level payload, List[List[str]] 5-column table):
#   KEY_MOD_GLOBALS = "global_vars"    — global var definitions table
#   KEY_MOD_STATICS = "static_vars"    — static var definitions table
import os
import json
import csv
import logging
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set
from io import BytesIO

from report.constants import (
    GLOBALS_FORMAT_ORDER,
    GLOBALS_FORMAT_SEP,
    GLOBALS_FORMAT_WITH_LABELS,
    LOGIC_MAX_DEPTH_DEFAULT,
    LOGIC_MAX_CHILDREN_DEFAULT,
    LOGIC_MAX_GRANDCHILDREN_DEFAULT,
)
from report_gen.function_analyzer import (
    _parse_signature_outputs,
    _fallback_function_description,
    _parse_signature_params,
    _normalize_symbol_name,
    _enhance_function_description,
    _is_generic_description,
    _enhance_description_text,
    _finalize_function_fields,
    _build_function_info_rows,
)
from report_gen.requirements import (
    _extract_doc_section,
    _extract_function_info_from_docx,
    _extract_sds_partition_map,
)
from report_gen.uds_text import (
    _merge_logic_ai_items,
    _apply_uds_rules,
    _merge_section_text,
    _ai_document_text,
    _ai_evidence_lines,
    _ai_quality_warnings,
)
from report_gen.utils import (
    _build_global_rows,
    _table_rows_from_texts,
    _normalize_swufn_id,
    _extract_call_names,
    _safe_dict,
)

_logger = logging.getLogger("report_generator")

def _add_docx_text_block(doc, text: str, max_lines: int = 8000) -> None:
    if not text:
        doc.add_paragraph("N/A")
        return
    lines = text.splitlines()
    for idx, ln in enumerate(lines):
        if max_lines and idx >= max_lines:
            doc.add_paragraph("...truncated...")
            break
        line = ln.rstrip()
        m = re.match(r"^(\d+(?:\.\d+)*)[\s\t]+(.+)$", line)
        if m:
            level = min(4, m.group(1).count(".") + 1)
            title = m.group(2).strip()
            if title:
                doc.add_heading(title, level=level)
                continue
        doc.add_paragraph(line)


def _replace_docx_text(doc, replacements: Dict[str, str]) -> None:
    def _replace_in_paragraph(paragraph):
        full = paragraph.text
        if not full:
            return
        changed = full
        for key, val in replacements.items():
            if key in changed:
                changed = changed.replace(key, val)
        if changed == full:
            return
        runs = paragraph.runs
        if not runs:
            paragraph.text = changed
            return
        joined = "".join(r.text for r in runs)
        if joined != full:
            paragraph.text = changed
            return
        for key, val in replacements.items():
            if key not in joined:
                continue
            cursor = 0
            for run in runs:
                rt = run.text
                start = cursor
                end = cursor + len(rt)
                cursor = end
                seg = joined[start:end]
                if key in seg:
                    run.text = seg.replace(key, val)
                    joined = "".join(r.text for r in runs)
                    break
            else:
                idx = joined.find(key)
                if idx < 0:
                    continue
                new_joined = joined[:idx] + val + joined[idx + len(key):]
                pos = 0
                for run in runs:
                    rlen = len(run.text)
                    run.text = new_joined[pos:pos + rlen] if pos + rlen <= len(new_joined) else new_joined[pos:]
                    pos += rlen
                if pos < len(new_joined):
                    runs[-1].text += new_joined[pos:]
                joined = "".join(r.text for r in runs)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if not hf or not hasattr(hf, "paragraphs"):
                continue
            for paragraph in hf.paragraphs:
                _replace_in_paragraph(paragraph)


def _add_docx_bullets(doc, text: str) -> None:
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        doc.add_paragraph("N/A")
        return
    for ln in lines:
        item = ln.lstrip("-* ").strip()
        if not item:
            continue
        try:
            doc.add_paragraph(item, style="List Bullet")
        except Exception:
            doc.add_paragraph(item)


def _add_docx_lines(doc, text: str) -> None:
    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    if not lines:
        doc.add_paragraph("N/A")
        return
    for ln in lines:
        doc.add_paragraph(ln)


def _add_docx_toc(doc) -> None:
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)


def _render_logic_flow_diagram(
    func_name: str,
    flow_nodes: List[Dict[str, Any]],
    out_path: Path,
    all_calls: Optional[List[str]] = None,
    call_map: Optional[Dict[str, List[str]]] = None,
) -> Optional[str]:
    """Render an adaptive logic diagram based on extracted control flow."""
    if not func_name:
        func_name = "Function"
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        return None

    out_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        font = ImageFont.truetype("arial.ttf", 12)
        font_sm = ImageFont.truetype("arial.ttf", 10)
        font_title = ImageFont.truetype("arial.ttf", 14)
    except Exception:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 12)
            font_sm = ImageFont.truetype("DejaVuSans.ttf", 10)
            font_title = ImageFont.truetype("DejaVuSans.ttf", 14)
        except Exception:
            font = ImageFont.load_default()
            font_sm = font
            font_title = font

    COLORS = {
        "start": "#E3F2FD", "end": "#E3F2FD",
        "call": "#E8F5E9", "call_outline": "#388E3C",
        "if_diamond": "#FFF8E1", "if_outline": "#F57F17",
        "true_box": "#E8F5E9", "false_box": "#FFEBEE",
        "switch_diamond": "#F3E5F5", "switch_outline": "#7B1FA2",
        "case_box": "#EDE7F6",
        "loop_box": "#E0F7FA", "loop_outline": "#00838F",
        "return_box": "#FCE4EC", "return_outline": "#C62828",
        "assign": "#F5F5F5",
        "arrow": "#37474F",
        "bg": "#FFFFFF",
    }

    W = 1200
    MARGIN_X = 60
    NODE_W = 260
    NODE_H = 40
    DIAMOND_W = 220
    DIAMOND_H = 90
    GAP_Y = 20
    BRANCH_GAP_X = 80
    MAX_VISIBLE_NODES = 30

    canvas_items: List[Dict[str, Any]] = []
    _uid = [0]

    def _next_id():
        _uid[0] += 1
        return _uid[0]

    def _add(kind: str, x: int, y: int, w: int, h: int, **kw):
        item = {"id": _next_id(), "kind": kind, "x": x, "y": y, "w": w, "h": h}
        item.update(kw)
        canvas_items.append(item)
        return item

    def _trunc(text: str, maxlen: int = 40) -> str:
        text = " ".join(str(text or "").replace("\n", " ").split()).strip()
        if len(text) > maxlen:
            return text[:maxlen - 3] + "..."
        return text

    def _layout_flow(nodes: List[Dict[str, Any]], cx: int, y: int,
                     avail_w: int, depth: int = 0) -> int:
        """Layout flow nodes starting at center x=cx, top y. Returns bottom y."""
        if depth > 4 or not nodes:
            return y
        node_count = [0]

        for node in nodes:
            if node_count[0] >= MAX_VISIBLE_NODES:
                _add("ellipse", cx - 60, y, 120, 30, text="...", fill=COLORS["assign"])
                y += 30 + GAP_Y
                break
            node_count[0] += 1
            ntype = node.get("type", "")

            if ntype == "call":
                nw = min(NODE_W, avail_w - 20)
                _add("rect", cx - nw // 2, y, nw, NODE_H,
                     text=node["name"], fill=COLORS["call"],
                     outline=COLORS["call_outline"], radius=6)
                y += NODE_H + GAP_Y

            elif ntype == "return":
                nw = min(NODE_W, avail_w - 20)
                val = node.get("value", "")
                label = f"return {val}" if val else "return"
                _add("rounded", cx - nw // 2, y, nw, 36,
                     text=_trunc(label, 36), fill=COLORS["return_box"],
                     outline=COLORS["return_outline"], radius=16)
                y += 36 + GAP_Y

            elif ntype == "assign":
                nw = min(NODE_W - 20, avail_w - 20)
                _add("rect", cx - nw // 2, y, nw, 32,
                     text=_trunc(node.get("text", ""), 36), fill=COLORS["assign"],
                     outline="#9E9E9E", radius=4)
                y += 32 + GAP_Y

            elif ntype == "if":
                cond = _trunc(node.get("condition", "?"), 50)
                dw = min(DIAMOND_W, avail_w - 40)
                dh = DIAMOND_H
                _add("diamond", cx - dw // 2, y, dw, dh,
                     text=cond, fill=COLORS["if_diamond"],
                     outline=COLORS["if_outline"])
                dia_bottom = y + dh

                true_body = node.get("true_body", [])
                false_body = node.get("false_body", [])

                branch_w = max(180, (avail_w - BRANCH_GAP_X) // 2)
                left_cx = cx - branch_w // 2 - BRANCH_GAP_X // 2
                right_cx = cx + branch_w // 2 + BRANCH_GAP_X // 2

                branch_top = dia_bottom + GAP_Y + 15
                _add("arrow", cx - dw // 2, y + dh // 2, 0, 0,
                     to_x=left_cx, to_y=branch_top, label="Y")
                _add("arrow", cx + dw // 2, y + dh // 2, 0, 0,
                     to_x=right_cx, to_y=branch_top, label="N")

                if true_body:
                    left_bottom = _layout_flow(true_body, left_cx, branch_top,
                                               branch_w, depth + 1)
                else:
                    _add("rect", left_cx - 60, branch_top, 120, 28,
                         text="(pass)", fill="#FAFAFA", outline="#BDBDBD", radius=4)
                    left_bottom = branch_top + 28

                if false_body:
                    right_bottom = _layout_flow(false_body, right_cx, branch_top,
                                                branch_w, depth + 1)
                else:
                    _add("rect", right_cx - 60, branch_top, 120, 28,
                         text="(pass)", fill="#FAFAFA", outline="#BDBDBD", radius=4)
                    right_bottom = branch_top + 28

                merge_y = max(left_bottom, right_bottom) + GAP_Y
                _add("arrow", left_cx, left_bottom, 0, 0,
                     to_x=cx, to_y=merge_y, label="")
                _add("arrow", right_cx, right_bottom, 0, 0,
                     to_x=cx, to_y=merge_y, label="")
                y = merge_y

            elif ntype == "switch":
                expr = _trunc(node.get("expr", "?"), 46)
                dw = min(DIAMOND_W + 40, avail_w - 40)
                dh = DIAMOND_H
                _add("diamond", cx - dw // 2, y, dw, dh,
                     text=f"switch({expr})", fill=COLORS["switch_diamond"],
                     outline=COLORS["switch_outline"])
                dia_bottom = y + dh

                cases = node.get("cases", [])[:6]
                default_calls = node.get("default_calls", [])[:3]
                n_branches = len(cases) + (1 if default_calls else 0)
                if n_branches == 0:
                    y = dia_bottom + GAP_Y
                    continue

                case_w = min(180, max(100, (avail_w - 40) // max(n_branches, 1)))
                total_w = n_branches * case_w + (n_branches - 1) * 10
                start_x = cx - total_w // 2
                branch_top = dia_bottom + GAP_Y + 15
                bottoms = []

                for ci, case in enumerate(cases):
                    bx = start_x + ci * (case_w + 10)
                    bcx = bx + case_w // 2
                    label = _trunc(case.get("label", "?"), 20)
                    _add("arrow", cx, dia_bottom, 0, 0,
                         to_x=bcx, to_y=branch_top, label=label)
                    call_y = branch_top
                    for call_name in case.get("calls", [])[:3]:
                        _add("rect", bx, call_y, case_w, 32,
                             text=_trunc(call_name, 22), fill=COLORS["case_box"],
                             outline=COLORS["switch_outline"], radius=4)
                        call_y += 32 + 6
                    bottoms.append(call_y)

                if default_calls:
                    di = len(cases)
                    bx = start_x + di * (case_w + 10)
                    bcx = bx + case_w // 2
                    _add("arrow", cx, dia_bottom, 0, 0,
                         to_x=bcx, to_y=branch_top, label="default")
                    call_y = branch_top
                    for call_name in default_calls[:3]:
                        _add("rect", bx, call_y, case_w, 32,
                             text=_trunc(call_name, 22), fill="#FBE9E7",
                             outline=COLORS["switch_outline"], radius=4)
                        call_y += 32 + 6
                    bottoms.append(call_y)

                merge_y = max(bottoms) + GAP_Y if bottoms else dia_bottom + GAP_Y
                for ci in range(n_branches):
                    bx = start_x + ci * (case_w + 10)
                    bcx = bx + case_w // 2
                    _add("arrow", bcx, bottoms[ci] if ci < len(bottoms) else merge_y, 0, 0,
                         to_x=cx, to_y=merge_y, label="")
                y = merge_y

            elif ntype == "loop":
                cond = _trunc(node.get("condition", ""), 46)
                kind = node.get("kind", "while")
                loop_label = f"{kind}({cond})" if cond else kind
                nw = min(NODE_W + 40, avail_w - 20)
                _add("loop_header", cx - nw // 2, y, nw, NODE_H,
                     text=_trunc(loop_label, 46), fill=COLORS["loop_box"],
                     outline=COLORS["loop_outline"], radius=8)
                header_bottom = y + NODE_H

                body_nodes = node.get("body", [])
                if body_nodes:
                    body_top = header_bottom + GAP_Y
                    inner_w = avail_w - 40
                    body_bottom = _layout_flow(body_nodes, cx, body_top,
                                               inner_w, depth + 1)
                    _add("loop_back", cx + nw // 2 + 10, header_bottom, 0, 0,
                         to_x=cx + nw // 2 + 10, to_y=y + NODE_H // 2,
                         body_bottom=body_bottom)
                    y = body_bottom + GAP_Y
                else:
                    y = header_bottom + GAP_Y

        return y

    # --- Build layout ---
    start_y = 20
    title_h = 44
    _add("ellipse", W // 2 - 180, start_y, 360, title_h,
         text=func_name, fill=COLORS["start"], outline="#1565C0")

    flow_top = start_y + title_h + GAP_Y
    cx = W // 2

    if not flow_nodes:
        calls = all_calls or []
        if calls:
            cy = flow_top
            for c in calls[:10]:
                _add("rect", cx - NODE_W // 2, cy, NODE_W, NODE_H,
                     text=c, fill=COLORS["call"], outline=COLORS["call_outline"], radius=6)
                if cy > flow_top:
                    pass  # arrows handled in render
                cy += NODE_H + GAP_Y
            flow_bottom = cy
        else:
            _add("rect", cx - 100, flow_top, 200, 36,
                 text="(no branch / direct)", fill="#FAFAFA", outline="#BDBDBD", radius=4)
            flow_bottom = flow_top + 36 + GAP_Y
    else:
        flow_bottom = _layout_flow(flow_nodes, cx, flow_top, W - MARGIN_X * 2)

    end_y = flow_bottom + GAP_Y
    _add("ellipse", cx - 120, end_y, 240, 36,
         text="End", fill=COLORS["end"], outline="#1565C0")
    total_h = end_y + 36 + 30

    # --- Render ---
    img = Image.new("RGB", (W, max(400, total_h)), COLORS["bg"])
    draw = ImageDraw.Draw(img)

    def _measure(s, f=None):
        f = f or font
        try:
            return draw.textlength(str(s or ""), font=f)
        except Exception:
            return len(str(s or "")) * 7

    def _draw_text_centered(box, text, f=None):
        f = f or font
        x1, y1, x2, y2 = box
        lines = str(text or "").split("\n")
        lh = 15
        total = len(lines) * lh
        ty = y1 + max(2, (y2 - y1 - total) // 2)
        for ln in lines:
            tw = _measure(ln, f)
            tx = x1 + max(4, (x2 - x1 - tw) // 2)
            draw.text((tx, ty), ln, fill="black", font=f)
            ty += lh

    def _draw_arrow(x1, y1, x2, y2, color=COLORS["arrow"]):
        draw.line([(x1, y1), (x2, y2)], fill=color, width=2)
        dx = x2 - x1
        dy = y2 - y1
        length = max(1, (dx * dx + dy * dy) ** 0.5)
        ux, uy = dx / length, dy / length
        ax, ay = x2 - ux * 8, y2 - uy * 8
        px, py = -uy * 5, ux * 5
        draw.polygon([(x2, y2), (int(ax + px), int(ay + py)),
                       (int(ax - px), int(ay - py))], fill=color)

    prev_center = None
    for item in canvas_items:
        k = item["kind"]
        x, y, w, h = item["x"], item["y"], item["w"], item["h"]

        if k == "ellipse":
            draw.rounded_rectangle([x, y, x + w, y + h], radius=h // 2,
                                   fill=item.get("fill", "#FFF"),
                                   outline=item.get("outline", "black"), width=2)
            _draw_text_centered((x, y, x + w, y + h), item.get("text", ""),
                                font_title if h > 40 else font)
            if prev_center and k != "arrow":
                _draw_arrow(prev_center[0], prev_center[1], x + w // 2, y)
            prev_center = (x + w // 2, y + h)

        elif k == "rect":
            r = item.get("radius", 0)
            if r:
                draw.rounded_rectangle([x, y, x + w, y + h], radius=r,
                                       fill=item.get("fill", "#FFF"),
                                       outline=item.get("outline", "black"), width=1)
            else:
                draw.rectangle([x, y, x + w, y + h],
                               fill=item.get("fill", "#FFF"),
                               outline=item.get("outline", "black"), width=1)
            _draw_text_centered((x, y, x + w, y + h), item.get("text", ""))

        elif k == "rounded":
            r = item.get("radius", 16)
            draw.rounded_rectangle([x, y, x + w, y + h], radius=r,
                                   fill=item.get("fill", "#FFF"),
                                   outline=item.get("outline", "black"), width=2)
            _draw_text_centered((x, y, x + w, y + h), item.get("text", ""))

        elif k == "diamond":
            cx_d = x + w // 2
            cy_d = y + h // 2
            pts = [(cx_d, y), (x + w, cy_d), (cx_d, y + h), (x, cy_d)]
            draw.polygon(pts, fill=item.get("fill", "#FFF"),
                         outline=item.get("outline", "black"))
            draw.polygon(pts, outline=item.get("outline", "black"))
            _draw_text_centered((x + 15, y + 8, x + w - 15, y + h - 8),
                                item.get("text", ""), font_sm)

        elif k == "loop_header":
            r = item.get("radius", 8)
            draw.rounded_rectangle([x, y, x + w, y + h], radius=r,
                                   fill=item.get("fill", "#FFF"),
                                   outline=item.get("outline", "black"), width=2)
            draw.rounded_rectangle([x + 2, y + 2, x + w - 2, y + h - 2], radius=r,
                                   fill=None,
                                   outline=item.get("outline", "black"), width=1)
            _draw_text_centered((x, y, x + w, y + h), item.get("text", ""))

        elif k == "arrow":
            to_x = item.get("to_x", x)
            to_y = item.get("to_y", y)
            label = item.get("label", "")
            _draw_arrow(x, y, to_x, to_y)
            if label:
                mx = (x + to_x) // 2
                my = (y + to_y) // 2
                bw = max(20, int(_measure(label, font_sm)) + 8)
                draw.rounded_rectangle([mx - bw // 2, my - 10, mx + bw // 2, my + 8],
                                       radius=4, fill="white", outline="#9E9E9E")
                tw = _measure(label, font_sm)
                draw.text((mx - tw // 2, my - 8), label, fill="black", font=font_sm)

        elif k == "loop_back":
            to_x = item.get("to_x", x)
            to_y = item.get("to_y", y)
            bb = item.get("body_bottom", y + 40)
            lx = to_x
            draw.line([(lx, y), (lx, bb + 10)], fill=COLORS["loop_outline"], width=2)
            draw.line([(lx, bb + 10), (lx - 20, bb + 10)], fill=COLORS["loop_outline"], width=2)
            draw.line([(lx, y), (lx - 20, y)], fill=COLORS["loop_outline"], width=2)
            _draw_arrow(lx - 20, y, lx - 20, to_y)

    # Draw sequential arrows between vertically-adjacent same-column items,
    # but only if no explicit arrow already targets the next node.
    arrow_targets = set()
    for it in canvas_items:
        if it["kind"] == "arrow":
            arrow_targets.add((it.get("to_x", 0), it.get("to_y", 0)))

    seq_items = [it for it in canvas_items
                 if it["kind"] in ("ellipse", "rect", "rounded", "loop_header")
                 and it.get("text") != "(pass)"]
    for i in range(len(seq_items) - 1):
        cur = seq_items[i]
        nxt = seq_items[i + 1]
        cur_bx = cur["x"] + cur["w"] // 2
        cur_by = cur["y"] + cur["h"]
        nxt_tx = nxt["x"] + nxt["w"] // 2
        nxt_ty = nxt["y"]
        if abs(cur_bx - nxt_tx) < 8 and nxt_ty > cur_by and nxt_ty - cur_by < 120:
            has_target = any(abs(tx - nxt_tx) < 20 and abs(ty - nxt_ty) < 20
                             for tx, ty in arrow_targets)
            if not has_target:
                _draw_arrow(cur_bx, cur_by, nxt_tx, nxt_ty)

    if total_h < img.size[1] - 60:
        img = img.crop((0, 0, W, total_h))

    img.save(str(out_path))
    return str(out_path)


def _render_logic_text_image(text: str, out_path: Path) -> Optional[str]:
    if not text:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1200, 700), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 13)
    except Exception:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 13)
        except Exception:
            font = ImageFont.load_default()
    try:
        draw.textlength
        def measure(s):
            txt = str(s or "").replace("\r", " ").replace("\n", " ")
            try:
                return draw.textlength(txt, font=font)  # type: ignore
            except Exception:
                return len(txt) * 6
    except Exception:
        measure = lambda s: len(s) * 6
    def _wrap_lines(text_block: str, max_width: int) -> str:
        lines = []
        for raw in text_block.splitlines():
            line = raw
            while line:
                if measure(line) <= max_width:
                    lines.append(line)
                    break
                cut = max(10, int(len(line) * (max_width / max(1, measure(line)))))
                lines.append(line[:cut])
                line = line[cut:]
        return "\n".join(lines)
    margin = 10
    wrapped = _wrap_lines(text, img.size[0] - margin * 2)
    draw.multiline_text((margin, margin), wrapped, fill="black", font=font, spacing=4)
    img.save(str(out_path))
    return str(out_path)


def _render_call_graph_image(
    func_name: str,
    calls: List[str],
    call_map: Optional[Dict[str, List[str]]],
    out_path: Path,
    max_children: int = 3,
    max_grandchildren: int = 2,
    max_depth: int = LOGIC_MAX_DEPTH_DEFAULT,
    module_map: Optional[Dict[str, str]] = None,
    condition_text: str = "",
    true_path_text: str = "",
    false_path_text: str = "",
    return_path_text: str = "",
    error_path_text: str = "",
) -> Optional[str]:
    if not func_name:
        func_name = "Function"
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        _logger.warning("PIL not available, skipping call graph image for %s", func_name)
        return None
    def _normalize_condition_label(text: str) -> str:
        s = " ".join(str(text or "").replace("\n", " ").split()).strip()
        if not s:
            return ""
        s = re.sub(r"\b([A-Za-z_]\w*)\s*!=\s*0\b", r"\1", s)
        s = re.sub(r"\b([A-Za-z_]\w*)\s*==\s*0\b", r"!\1", s)
        s = s.replace("&&", " AND ").replace("||", " OR ")
        s = re.sub(r"\s+", " ", s).strip()
        if len(s) > 40:
            s = s[:37].rstrip() + "..."
        return s

    def _branch_label(name: str) -> str:
        nm = str(name or "").strip()
        if not nm:
            return "N/A"
        if "(" in nm:
            m = re.search(r"\b([A-Za-z_]\w*)\s*\(", nm)
            nm = m.group(1) if m else nm
        nm = re.sub(r"^[sgu]_", "", nm)
        nm = nm.replace("_", " ")
        nm = re.sub(r"\s+", " ", nm).strip()
        if nm:
            nm = nm[0].upper() + nm[1:]
        return nm or "N/A"

    def _branch_path_label(path_text: str, fallback_name: str, fallback_default: str) -> str:
        raw = str(path_text or "").strip()
        if raw:
            first = re.split(r"[\n,;]", raw)[0].strip()
            if first:
                return _branch_label(first)
        if fallback_name:
            return _branch_label(fallback_name)
        return fallback_default

    def _derive_condition_label() -> str:
        if str(condition_text or "").strip():
            return _normalize_condition_label(condition_text)
        fn = str(func_name or "").lower()
        if "init" in fn:
            return "Init path selection"
        if "reset" in fn:
            return "Reset condition"
        if any(k in fn for k in ["check", "diag", "valid"]):
            return "Check / validation"
        if any(k in fn for k in ["state", "mode", "ctrl"]):
            return "State / mode branch"
        if children:
            sample = ", ".join([_branch_label(x) for x in children[:2]])
            return f"Call dispatch ({sample})" if sample else "Call dispatch"
        if return_path_text or error_path_text:
            return "Return / error branch"
        return "Condition unavailable"

    children = [c for c in calls if c][:max_children]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    width = 1420
    base_height = 980
    extra_depth = max(0, max_depth - 1)
    est_height = base_height + extra_depth * 180
    height = max(base_height, est_height)
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except Exception:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 12)
        except Exception:
            font = ImageFont.load_default()
    try:
        draw.textlength
        def measure(s):
            txt = str(s or "").replace("\r", " ").replace("\n", " ")
            try:
                return draw.textlength(txt, font=font)  # type: ignore
            except Exception:
                return len(txt) * 6
    except Exception:
        measure = lambda s: len(s) * 6
    def _fit_text(text: str, max_width: int) -> str:
        text = " ".join(str(text or "").splitlines()).strip()
        if measure(text) <= max_width:
            return text
        suffix = "..."
        trimmed = text
        while trimmed and measure(trimmed + suffix) > max_width:
            trimmed = trimmed[:-1]
        return trimmed + suffix if trimmed else text

    def _wrap_for_box(text: str, max_width: int, max_lines: int = 2) -> str:
        words = str(text or "").split()
        if not words:
            return ""
        lines: List[str] = []
        cur = words[0]
        for w in words[1:]:
            cand = f"{cur} {w}"
            if measure(cand) <= max_width:
                cur = cand
            else:
                lines.append(cur)
                cur = w
                if len(lines) >= max_lines - 1:
                    break
        lines.append(cur)
        lines = lines[:max_lines]
        if len(words) > 1 and len(lines) == max_lines:
            merged = " ".join(lines)
            if len(merged.split()) < len(words):
                lines[-1] = _fit_text(lines[-1], max_width)
        return "\n".join(lines)

    def _draw_centered_text(box, text: str, max_lines: int = 2):
        x1, y1, x2, y2 = box
        wrapped = _wrap_for_box(text, max(30, (x2 - x1) - 20), max_lines=max_lines)
        lines = wrapped.splitlines() if wrapped else [""]
        line_h = 15
        total_h = len(lines) * line_h
        ty = y1 + max(4, ((y2 - y1) - total_h) // 2)
        for ln in lines:
            tw = measure(ln)
            tx = x1 + max(8, int(((x2 - x1) - tw) // 2))
            draw.text((tx, ty), ln, fill="black", font=font)
            ty += line_h

    def _rounded_rect(box, radius=6, fill=None, outline="black", width=1):
        draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)

    def _ellipse(box, text, fill="#FFFFFF"):
        _rounded_rect(box, radius=22, fill=fill)
        _draw_centered_text(box, text, max_lines=2)

    def _diamond(center, w, h, text):
        cx, cy = center
        points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
        draw.polygon(points, fill="#FFFFFF", outline="black")
        _draw_centered_text((cx - w // 2 + 10, cy - h // 2 + 8, cx + w // 2 - 10, cy + h // 2 - 8), text, max_lines=3)

    def _arrow(p1, p2):
        draw.line([p1, p2], fill="black", width=2)
        x2, y2 = p2
        draw.polygon([(x2, y2), (x2 - 8, y2 - 5), (x2 - 8, y2 + 5)], fill="black")

    def _arrow_with_label(p1, p2, label: str, side: str = "center"):
        _arrow(p1, p2)
        mx = (p1[0] + p2[0]) // 2
        my = (p1[1] + p2[1]) // 2
        if side == "left":
            lx, ly = mx - 48, my - 20
        elif side == "right":
            lx, ly = mx + 12, my - 20
        else:
            lx, ly = mx - 18, my - 20
        bw = max(28, int(measure(label)) + 10)
        bh = 18
        _rounded_rect((lx - 4, ly - 2, lx - 4 + bw, ly - 2 + bh), radius=4, fill="#FFFFFF")
        draw.text((lx, ly), label, fill="black", font=font)

    def _pick_target_name(path_text: str, fallback: str) -> str:
        names = _extract_call_names(str(path_text or ""))
        if names:
            return str(names[0]).strip()
        return str(fallback or "").strip()

    def _draw_summary_panel(x: int, y: int, w: int, h: int) -> None:
        _rounded_rect((x, y, x + w, y + h), radius=8, fill="#F8F9FA")
        draw.text((x + 10, y + 8), "Flow Summary", fill="black", font=font)
        lines: List[str] = []
        lines.append(f"- Condition: {_derive_condition_label()}")
        if children:
            for idx, nm in enumerate(children[:6], start=1):
                lines.append(f"- Call {idx}: {_branch_label(nm)}")
        else:
            lines.append("- Call: N/A")
        if return_path_text:
            lines.append(f"- Return: {_branch_label(return_path_text)}")
        if error_path_text:
            lines.append(f"- Error: {_branch_label(error_path_text)}")
        ly = y + 30
        for ln in lines[:10]:
            draw.text((x + 10, ly), _fit_text(ln, w - 16), fill="black", font=font)
            ly += 16

    center_x = width // 2
    start_y = 26
    ellipse_w, ellipse_h = 520, 50
    proc_w, proc_h = 620, 60
    dia_w, dia_h = 260, 120
    branch_w, branch_h = 400, 64

    _ellipse((center_x - ellipse_w // 2, start_y, center_x + ellipse_w // 2, start_y + ellipse_h), func_name, fill="#F7FBFF")

    proc_y = start_y + 80
    _rounded_rect((center_x - proc_w // 2, proc_y, center_x + proc_w // 2, proc_y + proc_h), radius=8, fill="#EEF3F8")
    _draw_centered_text((center_x - proc_w // 2, proc_y, center_x + proc_w // 2, proc_y + proc_h), func_name, max_lines=2)
    _arrow((center_x, start_y + ellipse_h), (center_x, proc_y))

    dia_y = proc_y + 90
    _diamond(
        (center_x, dia_y + dia_h // 2),
        dia_w,
        dia_h,
        _fit_text(_derive_condition_label(), dia_w - 16),
    )
    _arrow((center_x, proc_y + proc_h), (center_x, dia_y))

    yes_label = _branch_path_label(true_path_text, children[0] if children else "", "Yes Path")
    no_label = _branch_path_label(false_path_text, children[1] if len(children) > 1 else "", "No Path")

    branch_gap = max(320, branch_w // 2 + 150)
    left_x = center_x - branch_gap
    right_x = center_x + branch_gap
    branch_y = dia_y + 138

    _rounded_rect((left_x - branch_w // 2, branch_y, left_x + branch_w // 2, branch_y + branch_h), radius=8, fill="#E8F5E9")
    _draw_centered_text((left_x - branch_w // 2, branch_y, left_x + branch_w // 2, branch_y + branch_h), yes_label, max_lines=2)

    _rounded_rect((right_x - branch_w // 2, branch_y, right_x + branch_w // 2, branch_y + branch_h), radius=8, fill="#FFF3E0")
    _draw_centered_text((right_x - branch_w // 2, branch_y, right_x + branch_w // 2, branch_y + branch_h), no_label, max_lines=2)

    _arrow_with_label((center_x - dia_w // 2, dia_y + dia_h // 2), (left_x + branch_w // 2, branch_y), "Yes", side="left")
    _arrow_with_label((center_x + dia_w // 2, dia_y + dia_h // 2), (right_x - branch_w // 2, branch_y), "No", side="right")

    # Depth-2 call expansion for better call-flow readability.
    depth2_bottom = branch_y + branch_h
    if call_map and max_depth >= 2:
        yes_target = _pick_target_name(true_path_text, children[0] if children else "")
        no_target = _pick_target_name(false_path_text, children[1] if len(children) > 1 else "")
        yes_next = [str(x).strip() for x in (call_map.get(yes_target, []) or []) if str(x).strip()][:max_grandchildren]
        no_next = [str(x).strip() for x in (call_map.get(no_target, []) or []) if str(x).strip()][:max_grandchildren]

        def _draw_depth2(x_center: int, labels: List[str], fill_color: str) -> int:
            if not labels:
                return branch_y + branch_h
            y0 = branch_y + 86
            h2 = 46
            w2 = 300
            for i, nm in enumerate(labels):
                y = y0 + i * 58
                _rounded_rect((x_center - w2 // 2, y, x_center + w2 // 2, y + h2), radius=6, fill=fill_color)
                _draw_centered_text((x_center - w2 // 2, y, x_center + w2 // 2, y + h2), _branch_label(nm), max_lines=2)
                if i == 0:
                    _arrow((x_center, branch_y + branch_h), (x_center, y))
                else:
                    _arrow((x_center, y - 14), (x_center, y))
            return y0 + (len(labels) - 1) * 58 + h2

        depth2_bottom = max(
            _draw_depth2(left_x, yes_next, "#E3F2FD"),
            _draw_depth2(right_x, no_next, "#E3F2FD"),
        )
        if max_depth >= 3:
            def _draw_depth3(seed_names: List[str], x_center: int) -> int:
                d3: List[str] = []
                for seed in seed_names[: max_grandchildren]:
                    nxt = [str(x).strip() for x in (call_map.get(seed, []) or []) if str(x).strip()]
                    for n in nxt[:2]:
                        if n not in d3:
                            d3.append(n)
                if not d3:
                    return depth2_bottom
                y0 = depth2_bottom + 48
                h3 = 40
                w3 = 250
                for i, nm in enumerate(d3[:4]):
                    y = y0 + i * 48
                    _rounded_rect((x_center - w3 // 2, y, x_center + w3 // 2, y + h3), radius=5, fill="#EAF4FF")
                    _draw_centered_text((x_center - w3 // 2, y, x_center + w3 // 2, y + h3), _branch_label(nm), max_lines=2)
                    if i == 0:
                        _arrow((x_center, depth2_bottom), (x_center, y))
                return y0 + (min(len(d3), 4) - 1) * 48 + h3

            depth2_bottom = max(depth2_bottom, _draw_depth3(yes_next, left_x), _draw_depth3(no_next, right_x))

    # Add side summary panel for readability in large call graphs.
    _draw_summary_panel(20, 160, 300, 260)

    end_y = depth2_bottom + 92
    return_text = str(return_path_text or "").strip()
    error_text = str(error_path_text or "").strip()
    has_return = bool(return_text)
    has_error = bool(error_text)

    if has_return and has_error:
        _ellipse((center_x - 520, end_y, center_x - 160, end_y + ellipse_h), _fit_text(return_text, 320))
        _rounded_rect((center_x + 160, end_y, center_x + 520, end_y + ellipse_h), radius=16, fill="#FFEBEE")
        _draw_centered_text((center_x + 160, end_y, center_x + 520, end_y + ellipse_h), error_text, max_lines=2)
        _arrow((left_x, depth2_bottom), (center_x - 340, end_y))
        _arrow((right_x, depth2_bottom), (center_x + 340, end_y))
    elif has_return or has_error:
        terminal = return_text if has_return else error_text
        fill = "#FFFFFF" if has_return else "#FFEBEE"
        _rounded_rect((center_x - ellipse_w // 2, end_y, center_x + ellipse_w // 2, end_y + ellipse_h), radius=16, fill=fill)
        _draw_centered_text((center_x - ellipse_w // 2, end_y, center_x + ellipse_w // 2, end_y + ellipse_h), terminal, max_lines=2)
        _arrow((left_x, depth2_bottom), (center_x - 80, end_y))
        _arrow((right_x, depth2_bottom), (center_x + 80, end_y))
    else:
        _ellipse((center_x - ellipse_w // 2, end_y, center_x + ellipse_w // 2, end_y + ellipse_h), "End")
        _arrow((left_x, depth2_bottom), (center_x - 80, end_y))
        _arrow((right_x, depth2_bottom), (center_x + 80, end_y))
    final_needed = end_y + ellipse_h + 40
    if final_needed > height:
        new_img = Image.new("RGB", (width, final_needed), "white")
        new_img.paste(img, (0, 0))
        img = new_img
    elif final_needed < height - 100:
        img = img.crop((0, 0, width, final_needed))
    img.save(str(out_path))
    return str(out_path)


def _render_unit_structure_image(
    module_label: str,
    interfaces: List[str],
    internals: List[str],
    out_path: Path,
) -> Optional[str]:
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        _logger.warning("PIL not available, skipping structure image for %s", module_label)
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)
    width = 1300
    height = 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except Exception:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 12)
        except Exception:
            font = ImageFont.load_default()
    try:
        draw.textlength
        measure = lambda s: draw.textlength(s, font=font)  # type: ignore
    except Exception:
        measure = lambda s: len(s) * 6
    def _fit_text(text: str, max_width: int) -> str:
        if measure(text) <= max_width:
            return text
        suffix = "..."
        trimmed = text
        while trimmed and measure(trimmed + suffix) > max_width:
            trimmed = trimmed[:-1]
        return trimmed + suffix if trimmed else text
    max_items = 20
    all_interfaces = [str(x).strip() for x in interfaces if str(x).strip()]
    all_internals = [str(x).strip() for x in internals if str(x).strip()]
    overflow_if = len(all_interfaces) - max_items if len(all_interfaces) > max_items else 0
    overflow_in = len(all_internals) - max_items if len(all_internals) > max_items else 0
    interfaces = all_interfaces[:max_items]
    internals = all_internals[:max_items]
    if overflow_if > 0:
        interfaces.append(f"... +{overflow_if} more")
    if overflow_in > 0:
        internals.append(f"... +{overflow_in} more")
    lane_rows = max(len(interfaces), len(internals), 1)
    lane_step = 48 if lane_rows <= 10 else (36 if lane_rows <= 15 else 28)
    lane_y0 = 220
    needed_h = lane_y0 + lane_rows * lane_step + 80
    if needed_h > height:
        height = needed_h
        img = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(img)

    def _rounded(box, fill="#FFFFFF", outline="black", radius=8, w=1):
        draw.rounded_rectangle(box, radius=radius, outline=outline, width=w, fill=fill)

    header_box = (40, 24, width - 40, 96)
    _rounded(header_box, fill="#DDE8D6", radius=10, w=2)
    module_text = _fit_text(module_label or "Unit Structure", header_box[2] - header_box[0] - 20)
    draw.text((header_box[0] + 12, header_box[1] + 20), module_text, fill="black", font=font)
    draw.text(
        (header_box[0] + 12, header_box[1] + 46),
        f"Interfaces: {len(all_interfaces)} | Internals: {len(all_internals)}",
        fill="black",
        font=font,
    )

    left_x = 90
    center_x = width // 2
    right_x = width - 510
    lane_w = 420
    lane_h = 34

    # Lane titles
    draw.text((left_x, lane_y0 - 32), "Interface Functions", fill="black", font=font)
    draw.text((right_x, lane_y0 - 32), "Internal Functions", fill="black", font=font)

    # Core node
    core_box = (center_x - 170, 134, center_x + 170, 186)
    _rounded(core_box, fill="#FFF3CD", radius=10, w=2)
    draw.text((core_box[0] + 12, core_box[1] + 18), _fit_text(module_label or "Core Module", 320), fill="black", font=font)

    def _arrow(p1, p2):
        draw.line([p1, p2], fill="black", width=1)
        x2, y2 = p2
        draw.polygon([(x2, y2), (x2 - 7, y2 - 4), (x2 - 7, y2 + 4)], fill="black")

    # Render interface lane
    if not interfaces:
        interfaces = ["N/A"]
    for idx, name in enumerate(interfaces):
        y = lane_y0 + idx * lane_step
        box = (left_x, y, left_x + lane_w, y + lane_h)
        _rounded(box, fill="#CFE2FF", radius=6)
        draw.text((left_x + 8, y + 8), _fit_text(name, lane_w - 16), fill="black", font=font)
        _arrow((core_box[0], (core_box[1] + core_box[3]) // 2), (left_x + lane_w, y + lane_h // 2))

    # Render internal lane
    if not internals:
        internals = ["N/A"]
    for idx, name in enumerate(internals):
        y = lane_y0 + idx * lane_step
        box = (right_x, y, right_x + lane_w, y + lane_h)
        _rounded(box, fill="#D9F2D9", radius=6)
        draw.text((right_x + 8, y + 8), _fit_text(name, lane_w - 16), fill="black", font=font)
        _arrow((core_box[2], (core_box[1] + core_box[3]) // 2), (right_x, y + lane_h // 2))

    img.save(str(out_path))
    return str(out_path)


def _render_swcom_overview_image(swcoms: List[str], out_path: Path) -> Optional[str]:
    if not swcoms:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont  # type: ignore
    except Exception:
        return None
    out_path.parent.mkdir(parents=True, exist_ok=True)
    width = 900
    box_h = 28
    margin = 20
    max_items = min(len(swcoms), 20)
    height = margin * 2 + max_items * (box_h + 10) + 40
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 12)
    except Exception:
        try:
            font = ImageFont.truetype("DejaVuSans.ttf", 12)
        except Exception:
            font = ImageFont.load_default()
    title = "Software Unit Structure Overview"
    draw.text((margin, margin), title, fill="black", font=font)
    y = margin + 24
    for name in swcoms[:max_items]:
        draw.rectangle([margin, y, width - margin, y + box_h], outline="black", width=1)
        draw.text((margin + 8, y + 6), name, fill="black", font=font)
        y += box_h + 10
    img.save(str(out_path))
    return str(out_path)


def _merge_function_info_table(table, cols: int) -> None:
    if not table or cols < 4:
        return
    try:
        if len(table.rows) >= 1:
            table.cell(0, 0).merge(table.cell(0, cols - 1))
        for r_idx in range(1, len(table.rows)):
            if cols >= 2:
                table.cell(r_idx, 0).merge(table.cell(r_idx, 1))
            if cols >= 4:
                table.cell(r_idx, 2).merge(table.cell(r_idx, cols - 1))
    except Exception:
        pass


def _normalize_function_info_tables(doc) -> None:
    if not doc:
        return
    try:
        for table in doc.tables:
            if not table.rows:
                continue
            header_cells = [c.text.strip() for c in table.rows[0].cells]
            if not any("Function Information" in c for c in header_cells):
                continue
            cols = len(table.columns)
            for cell in table.rows[0].cells:
                cell.text = "[ Function Information ]"
            _merge_function_info_table(table, cols)
    except Exception:
        pass


def _fill_function_info_table(table, data_rows: List[List[str]]) -> None:
    if not table or not data_rows:
        return
    try:
        for r_idx, row in enumerate(data_rows):
            if r_idx >= len(table.rows):
                break
            label = row[0] if len(row) > 0 else ""
            value = row[2] if len(row) > 2 else (row[1] if len(row) > 1 else "")
            # clear row first
            for c in table.rows[r_idx].cells:
                c.text = ""
            if r_idx == 0:
                for c_idx in range(len(table.rows[r_idx].cells)):
                    table.cell(0, c_idx).text = label
            else:
                table.cell(r_idx, 0).text = label
                table.cell(r_idx, 2).text = value
    except Exception:
        pass


def _insert_logic_image_in_table(table, cols: int, logic_img: str) -> bool:
    if not table or not logic_img:
        return False
    try:
        from docx.shared import Inches  # type: ignore
    except Exception:
        Inches = None  # type: ignore
    def _clear_cell(cell) -> None:
        try:
            from docx.oxml import OxmlElement  # type: ignore
        except Exception:
            OxmlElement = None  # type: ignore
        try:
            tc = cell._tc
            for child in list(tc):
                tc.remove(child)
            if OxmlElement:
                tc.append(OxmlElement("w:p"))
        except Exception as e:
            _logger.debug("Cell XML clear failed: %s", e)
            try:
                cell.text = ""
            except Exception:
                pass
    try:
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if any(c.replace(" ", "") == "LogicDiagram" for c in cells):
                target_cell = table.cell(r_idx, min(2, cols - 1))
                _clear_cell(target_cell)
                p = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
                run = p.add_run()
                if Inches:
                    run.add_picture(str(logic_img), width=Inches(5.2))
                else:
                    run.add_picture(str(logic_img))
                return True
    except Exception as e:
        _logger.warning("Failed to insert logic image in table: %s", e)
        return False
    return False


def _clear_docx_body(doc) -> None:
    try:
        body = doc._body._element  # type: ignore[attr-defined]
        for child in list(body):
            tag = getattr(child, "tag", "")
            if isinstance(tag, str) and tag.endswith("}sectPr"):
                continue
            body.remove(child)
    except Exception:
        pass


def _remove_docx_paragraphs(doc, texts: List[str]) -> None:
    try:
        from docx.text.paragraph import Paragraph  # type: ignore
    except Exception:
        return
    targets = {t.strip() for t in texts if t and t.strip()}
    if not targets:
        return
    try:
        body = doc._body._element  # type: ignore[attr-defined]
        for child in list(body):
            tag = getattr(child, "tag", "")
            if not isinstance(tag, str) or not tag.endswith("}p"):
                continue
            para = Paragraph(child, doc)
            if (para.text or "").strip() in targets:
                body.remove(child)
    except Exception:
        pass
    try:
        for table in doc.tables:
            header_cells = [c.text.strip() for c in table.rows[0].cells] if table.rows else []
            is_function_info = any("Function Information" in c for c in header_cells)
            for row in table.rows:
                for cell in row.cells:
                    if not is_function_info and (cell.text or "").strip() in targets:
                        cell.text = ""
                    for paragraph in cell.paragraphs:
                        if not is_function_info and (paragraph.text or "").strip() in targets:
                            paragraph.text = ""
    except Exception:
        pass


def _template_has_placeholders(doc) -> bool:
    def _check_text(t: str) -> bool:
        return bool(t and "{{" in t and "}}" in t)

    try:
        for p in doc.paragraphs:
            if _check_text(p.text or ""):
                return True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if _check_text(p.text or ""):
                            return True
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if not hf or not hasattr(hf, "paragraphs"):
                    continue
                for p in hf.paragraphs:
                    if _check_text(p.text or ""):
                        return True
    except Exception:
        return False
    return False


def _iter_template_blocks(doc):
    try:
        from docx.oxml.table import CT_Tbl  # type: ignore
        from docx.oxml.text.paragraph import CT_P  # type: ignore
        from docx.table import Table  # type: ignore
        from docx.text.paragraph import Paragraph  # type: ignore
    except Exception:
        return []
    parent = doc._body._element  # type: ignore[attr-defined]
    blocks = []
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            blocks.append(Table(child, doc))
    return blocks


def _extract_template_blocks(doc) -> List[Tuple[str, Any]]:
    blocks: List[Tuple[str, Any]] = []
    stack: List[str] = []
    header_keywords = {
        "file name",
        "version",
        "date",
        "note",
        "macro",
        "type",
        "define",
        "description",
        "parameter",
        "component",
        "function",
        "comment",
        "data name",
        "data type",
        "value range",
        "reset",
    }
    for item in _iter_template_blocks(doc):
        if hasattr(item, "style") and hasattr(item, "text"):
            text = (item.text or "").strip()
            if not text or not item.style:
                continue
            name = str(getattr(item.style, "name", "") or "")
            if not name.startswith("Heading"):
                blocks.append(("para", {"text": text, "style": name}))
                continue
            level = 2
            parts = re.findall(r"\d+", name)
            if parts:
                try:
                    level = max(1, int(parts[0]))
                except Exception:
                    level = 2
            if len(stack) >= level:
                stack = stack[: level - 1]
            stack.append(text)
            blocks.append(("heading", (level, text)))
        elif hasattr(item, "rows") and hasattr(item, "columns"):
            try:
                rows = len(item.rows)
                cols = len(item.columns)
                style = getattr(item, "style", None)
                header_rows: List[List[str]] = []
                for r in item.rows[:2]:
                    header_rows.append([c.text.strip() for c in r.cells])
                if len(header_rows) == 2:
                    first = " ".join([c.lower() for c in header_rows[0] if c]).strip()
                    second = " ".join([c.lower() for c in header_rows[1] if c]).strip()
                    def _is_header(row_text: str) -> bool:
                        return any(k in row_text for k in header_keywords)
                    if _is_header(first) and not _is_header(second):
                        header_rows = header_rows[:1]
                blocks.append(("table", (rows, cols, style, header_rows, list(stack))))
            except Exception:
                continue
    return blocks


def _extract_template_section_map(doc) -> Dict[str, str]:
    section_map: Dict[str, str] = {}
    current_title = ""
    current_lines: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = str(getattr(p.style, "name", "") or "")
        level = 0
        if style.startswith("Heading"):
            parts = re.findall(r"\d+", style)
            if parts:
                try:
                    level = max(1, int(parts[0]))
                except Exception:
                    level = 1
            else:
                level = 1
        if level > 0:
            if current_title:
                section_map[current_title.lower()] = "\n".join(current_lines).strip()
            current_title = text
            current_lines = []
            continue
        current_lines.append(text)
    if current_title and current_title.lower() not in section_map:
        section_map[current_title.lower()] = "\n".join(current_lines).strip()
    return section_map


def _add_blank_table(
    doc,
    rows: int,
    cols: int,
    style: Any = None,
    header_rows: Optional[List[List[str]]] = None,
    data_rows: Optional[List[List[str]]] = None,
) -> Any:
    if rows <= 0 or cols <= 0:
        return None
    table = doc.add_table(rows=rows, cols=cols)
    try:
        if style:
            table.style = style
    except Exception:
        pass
    row_offset = 0
    if header_rows:
        for r_idx, row in enumerate(header_rows):
            if r_idx >= rows:
                break
            for c_idx, val in enumerate(row[:cols]):
                table.cell(r_idx, c_idx).text = val or ""
        row_offset = min(len(header_rows), rows)
    if data_rows:
        max_rows = rows - row_offset
        for r_idx, row in enumerate(data_rows[:max_rows]):
            for c_idx, val in enumerate(row[:cols]):
                table.cell(row_offset + r_idx, c_idx).text = str(val) if val is not None else ""
    for r_idx in range(row_offset, rows):
        for c in table.rows[r_idx].cells:
            if c.text is None:
                c.text = ""
    return table


def generate_uds_docx(
    template_path: Optional[str],
    uds_payload: Dict[str, Any],
    output_path: str,
    ai_config: Optional[Dict[str, Any]] = None,
) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise ImportError("python-docx 미설치로 DOCX 생성 불가") from exc

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    payload = _safe_dict(uds_payload)
    summary = _safe_dict(payload.get("summary", {}))
    project = payload.get("project_name") or summary.get("project") or summary.get("project_name") or "UDS Spec"
    generated_at = payload.get("generated_at") or datetime.now().isoformat(timespec="seconds")

    ai_sections = payload.get("ai_sections")
    overview = _apply_uds_rules(
        _merge_section_text(payload.get("overview", "") or "", ai_sections, "overview"),
        "overview",
    )
    requirements = _apply_uds_rules(
        _merge_section_text(payload.get("requirements", "") or "", ai_sections, "requirements"),
        "requirements",
    )
    interfaces = _apply_uds_rules(
        _merge_section_text(payload.get("interfaces", "") or "", ai_sections, "interfaces"),
        "interfaces",
    )
    uds_frames = _apply_uds_rules(
        _merge_section_text(payload.get("uds_frames", "") or "", ai_sections, "uds_frames"),
        "uds_frames",
    )
    notes_text = _merge_section_text(
        payload.get("notes", "") or "",
        ai_sections,
        "notes",
        append_base=True,
    )
    evidence_lines = _ai_evidence_lines(ai_sections)
    if evidence_lines:
        notes_text = "\n".join([notes_text, "Evidence:"] + evidence_lines).strip()
    quality_warnings = _ai_quality_warnings(ai_sections)
    if quality_warnings:
        notes_text = "\n".join([notes_text, "Quality warnings:"] + quality_warnings).strip()
    notes = _apply_uds_rules(notes_text, "notes")
    software_unit_design = payload.get("software_unit_design", "") or ""
    detailed_doc = _ai_document_text(ai_sections)
    unit_structure = payload.get("unit_structure", "") or ""
    global_data = payload.get("global_data", "") or ""
    interface_functions = payload.get("interface_functions", "") or ""
    internal_functions = payload.get("internal_functions", "") or ""
    function_table_rows = payload.get("function_table_rows", []) or []
    global_vars = payload.get("global_vars", []) or []
    static_vars = payload.get("static_vars", []) or []
    macro_defs = payload.get("macro_defs", []) or []
    calibration_params = payload.get("calibration_params", []) or []
    generation_warnings: List[str] = []
    function_details = payload.get("function_details", {}) or {}
    function_details_by_name = payload.get("function_details_by_name", {}) or {}
    if (not function_details_by_name) and isinstance(function_details, dict):
        rebuilt: Dict[str, Dict[str, Any]] = {}
        for info in function_details.values():
            if not isinstance(info, dict):
                continue
            name = str(info.get("name") or "").strip().lower()
            if name:
                rebuilt[name] = info
        function_details_by_name = rebuilt
    if not function_details_by_name:
        source_root = payload.get("source_root") or payload.get("source_dir") or ""
        if source_root and Path(source_root).is_dir():
            try:
                from report_gen.uds_generator import generate_uds_source_sections  # lazy: heavy module
                fallback_src = generate_uds_source_sections(source_root)
                fb_details = fallback_src.get("function_details_by_name", {})
                if fb_details:
                    function_details_by_name = fb_details
                    function_details = fallback_src.get("function_details", {})
                    if not call_map:
                        call_map = fallback_src.get("call_map", {}) or {}
                    generation_warnings.append(f"Source root fallback used: {source_root}, found {len(fb_details)} functions")
            except Exception as e:
                generation_warnings.append(f"Source root fallback failed: {e}")
    if not function_details_by_name:
        rebuilt: Dict[str, Dict[str, Any]] = {}
        for ln in (payload.get("interface_functions", "") or "").splitlines() + (
            payload.get("internal_functions", "") or ""
        ).splitlines():
            sig = str(ln).strip()
            if not sig:
                continue
            m = re.match(r"([A-Za-z_]\w*)\s*\(", sig)
            name = m.group(1) if m else ""
            if not name:
                continue
            rebuilt[name.lower()] = {
                "id": "",
                "name": name,
                "prototype": sig,
                "description": "",
                "asil": "",
                "related": "",
                "inputs": _parse_signature_params(sig),
                "outputs": _parse_signature_outputs(sig, name),
                "precondition": "",
                "globals_global": [],
                "globals_static": [],
                "called": "",
                "logic": "",
            }
        if rebuilt:
            function_details_by_name = rebuilt
    call_map = payload.get("call_map", {}) or {}
    if isinstance(call_map, dict) and call_map:
        normalized_call_map: Dict[str, List[str]] = {}
        for k, vals in call_map.items():
            nk = _normalize_symbol_name(str(k or "")).lower()
            if not nk:
                continue
            out_vals: List[str] = []
            for v in (vals or []):
                nv = _normalize_symbol_name(str(v or ""))
                if nv and nv not in out_vals:
                    out_vals.append(nv)
            normalized_call_map[nk] = out_vals
        if normalized_call_map:
            call_map = normalized_call_map

    calling_map: Dict[str, List[str]] = {}
    for caller, callees in call_map.items():
        for callee in callees:
            callee_lower = callee.lower() if callee else ""
            if callee_lower:
                calling_map.setdefault(callee_lower, [])
                if caller not in calling_map[callee_lower]:
                    calling_map[callee_lower].append(caller)

    def _get_2hop_calls(fn_name: str) -> List[str]:
        direct = call_map.get(fn_name.lower(), [])
        indirect: List[str] = []
        for d in direct:
            for hop2 in call_map.get(d.lower(), []):
                if hop2 not in direct and hop2 != fn_name and hop2 not in indirect:
                    indirect.append(hop2)
        return indirect

    def _get_2hop_callers(fn_name: str) -> List[str]:
        direct = calling_map.get(fn_name.lower(), [])
        indirect: List[str] = []
        for d in direct:
            for hop2 in calling_map.get(d.lower(), []):
                if hop2 not in direct and hop2 != fn_name and hop2 not in indirect:
                    indirect.append(hop2)
        return indirect

    call_relation_mode = str(payload.get("call_relation_mode") or "code").strip().lower()
    if call_relation_mode not in {"code", "document"}:
        call_relation_mode = "code"
    module_map = payload.get("module_map", {}) or {}
    globals_info_map = payload.get("globals_info_map", {}) or {}
    globals_format_order = payload.get("globals_format_order") or GLOBALS_FORMAT_ORDER
    globals_format_sep = payload.get("globals_format_sep") or GLOBALS_FORMAT_SEP
    globals_format_with_labels = payload.get("globals_format_with_labels", GLOBALS_FORMAT_WITH_LABELS)
    common_macros = payload.get("common_macros", []) or []
    type_defs = payload.get("type_defs", []) or []
    param_defs = payload.get("param_defs", []) or []
    version_defs = payload.get("version_defs", []) or []
    req_map = payload.get("req_map", {}) or {}
    sds_partition_map = payload.get("sds_partition_map", {}) or {}
    sds_module_map = payload.get("sds_module_map", {}) or {}

    sds_texts = payload.get("sds_texts") or []
    sds_doc_paths = payload.get("sds_doc_paths") or []
    for sds_path in sds_doc_paths:
        try:
            docx_map = _extract_sds_partition_map(sds_path)
            if docx_map:
                for k, v in docx_map.items():
                    if k not in sds_partition_map:
                        sds_partition_map[k] = v
                    else:
                        for field in ("asil", "related", "description"):
                            if v.get(field) and not sds_partition_map[k].get(field):
                                sds_partition_map[k][field] = v[field]
        except Exception:
            pass
    _sds_name_labels = {"partition name", "component name", "module name", "name"}
    _sds_asil_labels = {"asil"}
    _sds_desc_labels = {"description", "desc"}
    if not sds_partition_map and sds_texts:
        for sds_text in sds_texts:
            lines = (sds_text or "").splitlines()
            for i, line in enumerate(lines):
                m = re.search(r"(SwCom_\d+|Component\s+\d+|Module\s+\d+)", line, re.I)
                if m:
                    com_id = m.group(1)
                    name = ""
                    asil = ""
                    related = ""
                    desc_parts: List[str] = []
                    for j in range(i + 1, min(i + 30, len(lines))):
                        ln = lines[j].strip()
                        ln_lower = ln.lower()
                        if re.match(r"^(?:SwCom_\d+|Component\s+\d+|Module\s+\d+)", ln, re.I):
                            break
                        label = ln_lower.split(None, 1)[0] if ln_lower.split(None, 1) else ""
                        full_label = " ".join(ln_lower.split(None, 2)[:2]) if len(ln_lower.split(None, 2)) >= 2 else label
                        if full_label in _sds_name_labels or label in _sds_name_labels:
                            parts = ln.split(None, 2)
                            name = parts[-1].strip() if len(parts) > 1 else ""
                        elif label in _sds_asil_labels:
                            parts = ln.split(None, 1)
                            asil = parts[-1].strip() if len(parts) > 1 else ""
                        elif label in _sds_desc_labels or full_label in _sds_desc_labels:
                            parts = ln.split(None, 1)
                            desc_parts.append(parts[-1].strip() if len(parts) > 1 else "")
                        elif ln_lower.startswith("related"):
                            parts = ln.split(None, 1)
                            related = parts[-1].strip() if len(parts) > 1 else ""
                    key = (name or com_id).strip().lower()
                    key = re.sub(r"[^a-z0-9_/ ]", "", key).strip()
                    if key:
                        entry = sds_partition_map.get(key, {})
                        if asil and not entry.get("asil"):
                            entry["asil"] = asil
                        if related and not entry.get("related"):
                            entry["related"] = related
                        if desc_parts and not entry.get("description"):
                            entry["description"] = " ".join(desc_parts).strip()
                        sds_partition_map[key] = entry
    fn_module_map: Dict[str, str] = {}
    if function_table_rows:
        for row in function_table_rows:
            if not isinstance(row, list) or len(row) < 4:
                continue
            fid = str(row[2] or "").strip()
            mod = str(row[1] or "").strip()
            if fid and mod:
                fn_module_map[fid] = mod

    if isinstance(req_map, dict):
        for key, info in list(function_details_by_name.items()):
            if not isinstance(info, dict):
                continue
            req = req_map.get(key)
            if not req:
                fid = str(info.get("id") or "").strip().lower()
                if fid:
                    req = req_map.get(fid)
            if isinstance(req, dict):
                if not info.get("asil"):
                    info["asil"] = req.get("asil") or ""
                if not info.get("related"):
                    info["related"] = req.get("related") or ""

    def _inherit_module_asil(
        func_details: Dict[str, Any],
        module_map: Dict[str, str],
    ) -> None:
        module_asil: Dict[str, str] = {}
        for fid, finfo in func_details.items():
            if not isinstance(finfo, dict):
                continue
            asil = str(finfo.get("asil") or "").strip()
            if asil and asil not in {"TBD", ""}:
                mod = module_map.get(fid, "")
                if mod and mod not in module_asil:
                    module_asil[mod] = asil
        for fid, finfo in func_details.items():
            if not isinstance(finfo, dict):
                continue
            asil = str(finfo.get("asil") or "").strip()
            if asil and asil not in {"TBD", ""}:
                continue
            mod = module_map.get(fid, "")
            inherited = module_asil.get(mod, "")
            if inherited:
                finfo["asil"] = inherited
                finfo["asil_source"] = "module_inherit"
            else:
                finfo["asil"] = "QM"
                finfo["asil_source"] = "default"

    def _resolve_related_asil_desc(
        info: Dict[str, Any],
        sds_info: Optional[Dict[str, str]],
    ) -> None:
        if not str(info.get("description_source") or "").strip():
            info["description_source"] = "inference"
        if not str(info.get("asil_source") or "").strip():
            info["asil_source"] = "inference"
        if not str(info.get("related_source") or "").strip():
            info["related_source"] = "inference"
        _weak_sources = {"", "inference", "default", "module_inherit"}
        c_asil = str(info.get("comment_asil") or "").strip()
        c_rel = str(info.get("comment_related") or "").strip()
        cur_asil_src = str(info.get("asil_source") or "").strip()
        cur_rel_src = str(info.get("related_source") or "").strip()
        if c_asil and cur_asil_src in _weak_sources:
            info["asil"] = c_asil
            info["asil_source"] = "comment"
        if c_rel and cur_rel_src in _weak_sources:
            info["related"] = c_rel
            info["related_source"] = "comment"
        if sds_info:
            cur_asil_src = str(info.get("asil_source") or "").strip()
            if cur_asil_src in _weak_sources:
                sds_asil = sds_info.get("asil")
                if sds_asil:
                    info["asil"] = sds_asil
                    info["asil_source"] = "sds"
            cur_rel_src = str(info.get("related_source") or "").strip()
            if cur_rel_src in _weak_sources:
                sds_related = sds_info.get("related")
                if sds_related:
                    info["related"] = sds_related
                    info["related_source"] = "sds"
            desc = str(info.get("description") or "").strip()
            if not desc or desc.startswith("Auto-generated from"):
                sds_desc = str(sds_info.get("description") or "").strip()
                if sds_desc:
                    info["description"] = sds_desc
                    info["description_source"] = "sds"
        if (not info.get("related")) or str(info.get("related")).strip() in {"", "TBD"}:
            text_blob = " ".join(
                [
                    str(info.get("description") or ""),
                    str(info.get("precondition") or ""),
                    str(info.get("called") or ""),
                ]
            )
            ids = re.findall(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b", text_blob)
            if ids:
                seen = set()
                uniq_ids = []
                for rid in ids:
                    rid = rid.strip()
                    if not rid or rid in seen:
                        continue
                    seen.add(rid)
                    uniq_ids.append(rid)
                info["related"] = ", ".join(uniq_ids)
                info["related_source"] = "srs"
                for rid in ids:
                    req = req_map.get(rid.lower())
                    if req and req.get("asil"):
                        info["asil"] = req.get("asil")
                        info["asil_source"] = "srs"
                        break
        if (not info.get("related")) or str(info.get("related")).strip() in {"", "TBD"}:
            fid = str(info.get("id") or "").strip()
            mod = fn_module_map.get(fid, "")
            swcom_ids = re.findall(r"\bSwCom_\d+\b", str(mod), flags=re.I)
            if swcom_ids:
                seen_sw = []
                for sid in swcom_ids:
                    sid_norm = sid.replace("swcom", "SwCom")
                    if sid_norm not in seen_sw:
                        seen_sw.append(sid_norm)
                info["related"] = ", ".join(seen_sw)
                info["related_source"] = "inference"
        if (not info.get("asil")) or str(info.get("asil")).strip() in {"", "TBD"}:
            info["asil"] = ""
            info["asil_source"] = "inference"
        if (not info.get("related")) or str(info.get("related")).strip() in {"", "TBD"}:
            info["related"] = ""
            info["related_source"] = "inference"

    _inherit_module_asil(function_details, fn_module_map)

    for info in list(function_details.values()):
        if not isinstance(info, dict):
            continue
        sds_info = None
        if sds_partition_map:
            fid = str(info.get("id") or "").strip()
            fname = str(info.get("name") or "").strip().lower()
            mod = fn_module_map.get(fid, "")
            key = mod.lower().strip() if mod else ""
            mapped = str(sds_module_map.get(key, "") or "").strip().lower()
            sds_info = sds_partition_map.get(mapped) if mapped else None
            if sds_info is None and fname:
                sds_info = sds_partition_map.get(fname)
            if sds_info is None and key:
                sds_info = sds_partition_map.get(key)
            if sds_info is None and key:
                norm_key = re.sub(r"[^a-z0-9]", "", key)
                for k, v in sds_partition_map.items():
                    if key in k or k in key:
                        sds_info = v
                        break
                    norm_k = re.sub(r"[^a-z0-9]", "", k)
                    if norm_key and norm_k and (norm_key in norm_k or norm_k in norm_key):
                        sds_info = v
                        break
            if sds_info is None:
                fallback = sds_partition_map.get("application/ driver") or sds_partition_map.get("application/driver")
                if fallback:
                    sds_info = fallback
        _resolve_related_asil_desc(info, sds_info if isinstance(sds_info, dict) else None)
        desc = str(info.get("description") or "").strip()
        if desc.startswith("Function") and "|" in desc:
            desc = desc.split("|", 1)[-1].strip()
        if (not desc) or _is_generic_description(desc) or (desc.lower().startswith("function") and len(desc.split()) <= 3):
            desc = _enhance_description_text(
                str(info.get("name") or ""),
                _fallback_function_description(
                    str(info.get("name") or ""),
                    info.get("called") or info.get("calls_list") or [],
                ),
                info.get("called") or info.get("calls_list") or [],
            )
            if _is_generic_description(desc):
                desc = _enhance_function_description(
                    str(info.get("name") or ""),
                    info.get("called") or info.get("calls_list") or [],
                    str(info.get("module_name") or ""),
                )
            info["description_source"] = "inference"
        else:
            desc = _enhance_description_text(
                str(info.get("name") or ""),
                desc,
                info.get("called") or info.get("calls_list") or [],
            )
            if _is_generic_description(desc):
                desc = _enhance_function_description(
                    str(info.get("name") or ""),
                    info.get("called") or info.get("calls_list") or [],
                    str(info.get("module_name") or ""),
                )
                info["description_source"] = "inference"
        info["description"] = desc

    # ── RAG 기반 Description 보강 (inference인 함수에 대해 유사 함수 설명 참조) ──
    _rag_desc_applied = 0
    try:
        from workflow.rag import KnowledgeBase
        _kb_dir = Path(os.environ.get("KB_STORE_DIR", "")) if os.environ.get("KB_STORE_DIR") else Path("kb_store")
        if _kb_dir.exists():
            _kb = KnowledgeBase(_kb_dir)
            _kb.load()
            if _kb.data:
                for fid, info in function_details.items():
                    if not isinstance(info, dict):
                        continue
                    if str(info.get("description_source") or "").strip() != "inference":
                        continue
                    fname = str(info.get("name") or "").strip()
                    proto = str(info.get("prototype") or "").strip()
                    query = f"{fname} {proto}".strip()
                    if not query:
                        continue
                    results = _kb.search(query, top_k=3, tags=["uds_description", "code"])
                    if not results:
                        results = _kb.search(query, top_k=3)
                    for r in results:
                        chunk_text = str(r.get("text") or r.get("content") or "").strip()
                        if not chunk_text or len(chunk_text) < 10:
                            continue
                        lines = chunk_text.split("\n")
                        desc_candidate = ""
                        for line in lines:
                            line = line.strip()
                            if fname.lower() in line.lower() and len(line) > 15:
                                desc_candidate = line
                                break
                        if desc_candidate and not _is_generic_description(desc_candidate):
                            info["description"] = desc_candidate[:200]
                            info["description_source"] = "rag"
                            _rag_desc_applied += 1
                            break
                if _rag_desc_applied:
                    _logger.info("RAG description applied: %d functions", _rag_desc_applied)
    except Exception as e:
        _logger.debug("RAG description enhancement skipped: %s", e)

    # 레퍼런스 SUDS에서 Function 정보 보강
    ref_related_by_name: Dict[str, str] = {}
    from config import UDS_REF_SUDS_PATH
    ref_doc_path = Path(UDS_REF_SUDS_PATH)
    if ref_doc_path.exists():
        try:
            ref_doc = docx.Document(str(ref_doc_path))
            ref_map = _extract_function_info_from_docx(ref_doc)
        except Exception:
            ref_map = {}
        if ref_map:
            patched_called = 0
            patched_calling = 0
            patched_limit = 9999
            for fid, block in ref_map.items():
                target = function_details.get(fid)
                if target is None:
                    name = _normalize_symbol_name(str(block.get("name") or ""))
                    if name:
                        target = function_details_by_name.get(name.lower())
                if not isinstance(target, dict):
                    continue
                bname = _normalize_symbol_name(str(block.get("name") or "")).lower()
                brel = str(block.get("related") or "").strip()
                if bname and brel:
                    ref_related_by_name[bname] = brel
                for key in ["description", "asil", "related", "precondition", "logic"]:
                    cur = str(target.get(key) or "").strip()
                    incoming = str(block.get(key) or "").strip()
                    if not incoming:
                        continue
                    if key == "description":
                        if (not cur) or cur.startswith("Auto-generated from"):
                            target[key] = incoming
                            target["description_source"] = "reference"
                    elif key in {"asil", "related"}:
                        if (not cur) or cur in {"TBD", "N/A", "-"}:
                            target[key] = incoming
                            target[f"{key}_source"] = "reference"
                    else:
                        if (not cur) or (key == "precondition" and cur.upper() in {"N/A", "TBD", "-"}):
                            target[key] = incoming
                if block.get("inputs") and not target.get("inputs"):
                    target["inputs"] = block.get("inputs")
                if block.get("outputs") and not target.get("outputs"):
                    target["outputs"] = block.get("outputs")
                if block.get("globals_static") and not target.get("globals_static"):
                    target["globals_static"] = block.get("globals_static")
                if block.get("globals_global") and not target.get("globals_global"):
                    target["globals_global"] = block.get("globals_global")
                if block.get("called"):
                    cur_called = str(target.get("called") or "").strip()
                    if ((not cur_called) or cur_called.upper() in {"N/A", "TBD", "-"}) and patched_called < patched_limit:
                        target["called"] = block.get("called")
                        patched_called += 1
                if block.get("calling"):
                    cur_calling = str(target.get("calling") or "").strip()
                    if ((not cur_calling) or cur_calling.upper() in {"N/A", "TBD", "-"}) and patched_calling < patched_limit:
                        target["calling"] = block.get("calling")
                        patched_calling += 1

    if isinstance(function_details, dict) and isinstance(function_details_by_name, dict):
        for fid, info in function_details.items():
            if not isinstance(info, dict):
                continue
            name = str(info.get("name") or "").strip().lower()
            target = function_details_by_name.get(name)
            if not isinstance(target, dict):
                continue
            for src_key in ("asil", "asil_source", "related", "related_source",
                            "description", "description_source"):
                val = info.get(src_key)
                if val and (not target.get(src_key) or (src_key.endswith("_source") and target.get(src_key) == "inference")):
                    target[src_key] = val
            for g_key in ("globals_global", "globals_static"):
                src_g = info.get(g_key)
                tgt_g = target.get(g_key)
                if isinstance(src_g, list) and src_g and (not tgt_g or (isinstance(tgt_g, list) and not tgt_g)):
                    target[g_key] = list(src_g)

    # ── 호출 그래프 기반 Related ID 전파 (2-hop BFS) ──
    _pre_prop_has_related = sum(1 for v in (function_details or {}).values()
        if isinstance(v, dict)
        and str(v.get("related") or "").strip()
        and str(v.get("related") or "").strip().upper() not in {"TBD", "N/A", "-", "NONE"})
    _logger.info("Related ID pre-propagation: %d/%d have valid related, call_map=%d keys",
                 _pre_prop_has_related, len(function_details or {}), len(call_map or {}))
    if isinstance(function_details, dict) and call_map:
        _related_propagated = 0
        _fn_name_to_fid: Dict[str, str] = {}
        for fid, info in function_details.items():
            if isinstance(info, dict):
                fname = str(info.get("name") or "").strip().lower()
                if fname:
                    _fn_name_to_fid[fname] = fid

        def _has_related(info: Dict[str, Any]) -> bool:
            rel = str(info.get("related") or "").strip()
            return bool(rel) and rel.upper() not in {"TBD", "N/A", "-", "NONE", ""}

        def _get_related(info: Dict[str, Any]) -> str:
            return str(info.get("related") or "").strip()

        for _hop in range(2):
            propagated_this_hop = 0
            for fid, info in list(function_details.items()):
                if not isinstance(info, dict) or _has_related(info):
                    continue
                fname = str(info.get("name") or "").strip().lower()
                if not fname:
                    continue
                collected_ids: List[str] = []
                callers = calling_map.get(fname, [])
                for caller_name in callers:
                    caller_fid = _fn_name_to_fid.get(caller_name.lower())
                    if not caller_fid:
                        continue
                    caller_info = function_details.get(caller_fid)
                    if isinstance(caller_info, dict) and _has_related(caller_info):
                        for rid in _get_related(caller_info).replace(";", ",").split(","):
                            rid = rid.strip()
                            if rid and rid not in collected_ids:
                                collected_ids.append(rid)
                callees = call_map.get(fname, [])
                for callee_name in callees:
                    callee_fid = _fn_name_to_fid.get(callee_name.lower())
                    if not callee_fid:
                        continue
                    callee_info = function_details.get(callee_fid)
                    if isinstance(callee_info, dict) and _has_related(callee_info):
                        for rid in _get_related(callee_info).replace(";", ",").split(","):
                            rid = rid.strip()
                            if rid and rid not in collected_ids:
                                collected_ids.append(rid)
                if collected_ids:
                    info["related"] = ", ".join(collected_ids[:8])
                    info["related_source"] = "call_graph"
                    propagated_this_hop += 1
                    if isinstance(function_details_by_name, dict):
                        tgt = function_details_by_name.get(fname)
                        if isinstance(tgt, dict):
                            tgt["related"] = info["related"]
                            tgt["related_source"] = "call_graph"
            _related_propagated += propagated_this_hop
            if propagated_this_hop == 0:
                break
        if _related_propagated > 0:
            _logger.info("Related ID propagation via call graph: %d functions updated", _related_propagated)

    ai_func_desc_enable = bool(payload.get("ai_func_desc_enable") or payload.get("ai_enable"))
    if ai_func_desc_enable and isinstance(function_details, dict):
        inference_count = sum(
            1 for v in function_details.values()
            if isinstance(v, dict) and str(v.get("description_source") or "").strip().lower() in {"inference", "rule", ""}
        )
        if inference_count > 0:
            _logger.info("AI function description: %d inference-sourced functions, starting AI generation", inference_count)
            try:
                from workflow.uds_ai import generate_ai_function_descriptions
                ai_descs = generate_ai_function_descriptions(function_details, module_map if isinstance(module_map, dict) else None)
                if ai_descs:
                    applied = 0
                    for fid, info in function_details.items():
                        if not isinstance(info, dict):
                            continue
                        src = str(info.get("description_source") or "").strip().lower()
                        if src in {"comment", "sds", "reference"}:
                            continue
                        ai_desc = ai_descs.get(f"__fid__{fid}")
                        if not ai_desc:
                            name = str(info.get("name") or "").strip().lower()
                            ai_desc = ai_descs.get(name)
                        if ai_desc:
                            info["description"] = ai_desc
                            info["description_source"] = "ai"
                            applied += 1
                            if isinstance(function_details_by_name, dict):
                                name = str(info.get("name") or "").strip().lower()
                                target = function_details_by_name.get(name)
                                if isinstance(target, dict):
                                    target["description"] = ai_desc
                                    target["description_source"] = "ai"
                    _logger.info("AI function description: applied %d / %d", applied, inference_count)
            except Exception as e:
                _logger.warning("AI function description failed: %s", e)

    if isinstance(function_details, dict):
        for k, v in list(function_details.items()):
            if isinstance(v, dict):
                function_details[k] = _finalize_function_fields(v)
    if isinstance(function_details_by_name, dict):
        for k, v in list(function_details_by_name.items()):
            if isinstance(v, dict):
                function_details_by_name[k] = _finalize_function_fields(v)

    tbd_asil = sum(1 for v in (function_details or {}).values() if isinstance(v, dict) and str(v.get("asil") or "").strip().upper() == "TBD")
    tbd_related = sum(1 for v in (function_details or {}).values() if isinstance(v, dict) and str(v.get("related") or "").strip().upper() == "TBD")
    total_fn = len(function_details or {})
    if tbd_asil > 0 or tbd_related > 0:
        _logger.warning("UDS TBD residual: asil_tbd=%d/%d, related_tbd=%d/%d", tbd_asil, total_fn, tbd_related, total_fn)

    if template_path:
        template_path = str(template_path)
        replacements = {
            "{{project_name}}": str(project),
            "{{job_url}}": str(payload.get("job_url") or ""),
            "{{build_number}}": str(payload.get("build_number") or ""),
            "{{generated_at}}": str(generated_at),
            "{{overview}}": overview,
            "{{requirements}}": requirements,
            "{{interfaces}}": interfaces,
            "{{uds_frames}}": uds_frames,
            "{{notes}}": notes,
        }
        doc = docx.Document(template_path)
        if _template_has_placeholders(doc):
            _replace_docx_text(doc, replacements)
            doc.save(str(out))
            return str(out)
        # 템플릿에 치환 키가 없으면, 구조만 복제하고 콘텐츠는 새로 작성
        blocks = _extract_template_blocks(doc)
        template_section_map = _extract_template_section_map(doc)
        _clear_docx_body(doc)
        first_heading = ""
        for kind, payload_block in blocks:
            if kind == "heading":
                first_heading = str(payload_block[1]).strip()
                break
        if False:
            pass
        has_contents_marker = False
        skip_table_idx = -1
        for kind, block in blocks:
            if kind == "heading" and str(block[1]).strip().lower() == "contents":
                has_contents_marker = True
                break
            if kind == "para" and str(block.get("text") or "").strip().lower() == "contents":
                has_contents_marker = True
                break
        toc_inserted = False
        if not has_contents_marker:
            doc.add_heading("Contents", level=2)
            _add_docx_toc(doc)
            toc_inserted = True
        skip_next_table = False
        heading_stack: List[str] = []
        module_funcs: Dict[str, Dict[str, List[str]]] = {}
        interface_queue: List[Dict[str, Any]] = []
        internal_queue: List[Dict[str, Any]] = []
        function_info_template: Optional[Tuple[int, int, Any, Optional[List[List[str]]]]] = None
        for kind_t, payload_t in blocks:
            if kind_t != "table":
                continue
            rows_t, cols_t, style_t, header_rows_t, _ctx_titles_t = payload_t
            if header_rows_t:
                header_texts = [str(c or "").strip() for row in header_rows_t for c in row]
                if any("Function Information" in c for c in header_texts):
                    function_info_template = (rows_t, cols_t, style_t, header_rows_t)
                    break
        swufn_table_spec: Dict[str, Tuple[int, int, Any]] = {}
        for idx_t, (kind_t, payload_t) in enumerate(blocks):
            if kind_t != "heading":
                continue
            level_t, title_t = payload_t
            m = re.search(r"(swufn_\\d+)", str(title_t), re.I)
            if not m:
                continue
            swufn_id = m.group(1).upper()
            for look_ahead in range(idx_t + 1, len(blocks)):
                kind_la, payload_la = blocks[look_ahead]
                if kind_la == "heading":
                    try:
                        level_la = int(payload_la[0])
                    except Exception:
                        level_la = level_t
                    if level_la <= level_t:
                        break
                    continue
                if kind_la == "table":
                    rows_la, cols_la, style_la, header_rows_la, _ctx_titles_la = payload_la
                    if header_rows_la:
                        header_texts = [str(c or "").strip() for row in header_rows_la for c in row]
                        if any("Function Information" in c for c in header_texts):
                            swufn_table_spec[swufn_id] = (rows_la, cols_la, style_la)
                            break
        for row in function_table_rows:
            if len(row) < 5:
                continue
            swcom = str(row[0] or "").strip()
            name = str(row[3] or "").strip()
            ftype = str(row[4] or "").strip().lower()
            if not swcom or not name:
                continue
            entry = module_funcs.setdefault(swcom, {"interfaces": [], "internals": []})
            if "i/f" in ftype or "if" == ftype:
                entry["interfaces"].append(name)
            else:
                entry["internals"].append(name)
            info = None
            if isinstance(function_details, dict):
                info = function_details.get(str(row[2] or "").strip())
            if not isinstance(info, dict) and isinstance(function_details_by_name, dict):
                info = function_details_by_name.get(name.lower())
            if isinstance(info, dict):
                if "i/f" in ftype or "if" == ftype:
                    interface_queue.append(info)
                else:
                    internal_queue.append(info)
        swcom_functions: Dict[str, List[str]] = {}
        swcom_function_files: Dict[str, Set[str]] = {}
        swcom_global_hints: Dict[str, Set[str]] = {}
        if isinstance(function_table_rows, list):
            for row in function_table_rows:
                if not isinstance(row, list) or len(row) < 4:
                    continue
                swcom = str(row[0] or "").strip()
                fn_name = str(row[3] or "").strip()
                if not swcom or not fn_name:
                    continue
                swcom_functions.setdefault(swcom, []).append(fn_name)
                f_info = function_details_by_name.get(fn_name.lower()) if isinstance(function_details_by_name, dict) else None
                if isinstance(f_info, dict):
                    fpath = str(f_info.get("file") or "").strip()
                    if fpath:
                        swcom_function_files.setdefault(swcom, set()).add(str(Path(fpath).name).lower())
                    for key_g in ["globals_global", "globals_static"]:
                        for g in f_info.get(key_g) or []:
                            raw = str(g or "").strip()
                            if not raw:
                                continue
                            mtag = re.match(r"^\[(?:INOUT|IN|OUT)\]\s+(.+)$", raw)
                            base = mtag.group(1).strip() if mtag else raw
                            name0 = re.split(r"(?:->|\.)", base)[0].strip()
                            if name0:
                                swcom_global_hints.setdefault(swcom, set()).add(name0)

        def _current_swcom_id() -> str:
            for h in reversed(heading_stack):
                m = re.search(r"\b(SwCom_\d+)\b", str(h), flags=re.I)
                if m:
                    return m.group(1).replace("swcom", "SwCom")
            return ""

        def _current_swcom_label() -> str:
            for h in reversed(heading_stack):
                m = re.search(r"\bSwCom_\d+\s*\(([^)]+)\)", str(h), flags=re.I)
                if m:
                    return str(m.group(1) or "").strip()
            return ""

        def _is_register_alias(name: str) -> bool:
            return bool(re.match(r"^REG_[A-Z0-9_]+$", str(name or "").strip()))

        def _norm_stem(name: str) -> str:
            s = re.sub(r"[^a-z0-9]+", "", str(name or "").lower())
            for suffix in ["itpds", "pds", "it", "main"]:
                if s.endswith(suffix) and len(s) > len(suffix) + 2:
                    s = s[: -len(suffix)]
            return s

        def _filter_global_names_by_swcom(
            all_names: List[str],
            swcom_id: str,
            module_label: str = "",
            static_only: Optional[bool] = None,
        ) -> List[str]:
            names = [str(n).strip() for n in all_names if str(n).strip() and not _is_register_alias(str(n))]
            if not swcom_id:
                swcom_id = ""
            module_candidates: Set[str] = set()
            label_norm = re.sub(r"[^a-z0-9]+", "", str(module_label or "").lower())
            if label_norm:
                module_candidates.add(label_norm)
                module_candidates.add(label_norm.replace("system", "sys"))
                if "system" in label_norm and "os" in label_norm:
                    module_candidates.add("sysos")
            # 1) module label matching has priority over swcom_id mapping
            selected: List[str] = []
            if module_candidates:
                for n in names:
                    info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
                    is_static = str(info_g.get("static") or "").strip().lower() == "true"
                    if static_only is True and not is_static:
                        continue
                    if static_only is False and is_static:
                        continue
                    gfile_l = str(info_g.get("file") or "").lower()
                    gblob = re.sub(r"[^a-z0-9]+", "", gfile_l)
                    if any(tok and tok in gblob for tok in module_candidates):
                        selected.append(n)
                if selected:
                    return list(dict.fromkeys(selected))
            hints = swcom_global_hints.get(swcom_id, set())
            for n in names:
                info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
                is_static = str(info_g.get("static") or "").strip().lower() == "true"
                if static_only is True and not is_static:
                    continue
                if static_only is False and is_static:
                    continue
                if n in hints:
                    selected.append(n)
            if selected:
                return list(dict.fromkeys(selected))
            file_names = swcom_function_files.get(swcom_id, set())
            if not file_names:
                return names
            file_stems = {_norm_stem(Path(x).stem) for x in file_names}
            for n in names:
                info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
                is_static = str(info_g.get("static") or "").strip().lower() == "true"
                if static_only is True and not is_static:
                    continue
                if static_only is False and is_static:
                    continue
                gfile = str(info_g.get("file") or "").strip()
                gstem = _norm_stem(Path(gfile).stem) if gfile else ""
                if gstem and any(gstem == fs or gstem.startswith(fs) or fs.startswith(gstem) for fs in file_stems):
                    selected.append(n)
            return list(dict.fromkeys(selected))

        def _filter_rows_by_swcom(rows_in: List[List[str]], swcom_id: str, module_label: str = "") -> List[List[str]]:
            if not swcom_id or not isinstance(rows_in, list):
                swcom_id = swcom_id or ""
            if not isinstance(rows_in, list):
                return rows_in
            hints = swcom_global_hints.get(swcom_id, set())
            if not hints and module_label:
                label_blob = re.sub(r"[^a-z0-9]+", "", module_label.lower())
                if "system" in label_blob and "os" in label_blob:
                    hints = {"u8g_SystemTm_5ms", "u8g_SystemTm_10ms", "u8g_SystemTm_50ms", "u8s_InitiComplet_F"}
            if not hints:
                return []
            out_rows: List[List[str]] = []
            for row in rows_in:
                if not isinstance(row, list) or not row:
                    continue
                blob = " ".join([str(x or "") for x in row])
                if any(re.search(rf"\b{re.escape(h)}\b", blob) for h in hints):
                    out_rows.append(row)
            return out_rows

        callers_map: Dict[str, List[str]] = {}
        if isinstance(call_map, dict):
            for caller_name, callees in call_map.items():
                cname = _normalize_symbol_name(str(caller_name or "")).lower()
                if not cname or not isinstance(callees, list):
                    continue
                for callee in callees:
                    callee_name = _normalize_symbol_name(str(callee or "")).lower()
                    if not callee_name:
                        continue
                    callers_map.setdefault(callee_name, []).append(cname)
        for k, vals in list(callers_map.items()):
            dedup: List[str] = []
            for v in vals:
                if v not in dedup:
                    dedup.append(v)
            callers_map[k] = dedup

        def _next_block_kind(blocks_list, start_idx: int) -> str:
            for j in range(start_idx + 1, len(blocks_list)):
                kind_j, payload_j = blocks_list[j]
                if kind_j == "para":
                    text_j = str(payload_j.get("text") or "").strip()
                    if not text_j:
                        continue
                return kind_j
            return ""
        def _section_has_table(blocks_list, start_idx: int, current_level: int) -> bool:
            for j in range(start_idx + 1, len(blocks_list)):
                kind_j, payload_j = blocks_list[j]
                if kind_j == "heading":
                    level_j = 1
                    try:
                        level_j = int(payload_j[0])
                    except Exception:
                        level_j = 1
                    if level_j <= current_level:
                        return False
                    continue
                if kind_j == "table":
                    return True
                if kind_j == "para":
                    text_j = str(payload_j.get("text") or "").strip()
                    if text_j:
                        continue
            return False

        section_note_map = {
            "parameter definition": template_section_map.get("parameter definition", ""),
            "version information": template_section_map.get("version information", ""),
        }
        note_added: set[str] = set()

        def _resolve_function_info(title_text: str, key_text: str) -> Dict[str, Any]:
            info: Optional[Dict[str, Any]] = None
            heading_fn_name = ""
            if ":" in str(title_text):
                heading_fn_name = _normalize_symbol_name(str(title_text).split(":", 1)[1]).lower()
            # For explicit SwUFn headings, prioritize function-name matching over ID.
            # Template SwUFn IDs may differ from parsed source IDs by module ordering.
            if ":" in str(title_text):
                fn_name = heading_fn_name
                if isinstance(function_details_by_name, dict):
                    info = function_details_by_name.get(fn_name)
            if not isinstance(info, dict) and ":" in str(title_text):
                fn_name = heading_fn_name
                if isinstance(function_details, dict) and fn_name:
                    for cand in function_details.values():
                        if not isinstance(cand, dict):
                            continue
                        cand_name = _normalize_symbol_name(str(cand.get("name") or "")).lower()
                        if cand_name == fn_name:
                            info = cand
                            break
                        # tolerate wrapper prefixes/suffixes while avoiding unrelated fallback.
                        if cand_name.endswith(fn_name) or fn_name.endswith(cand_name):
                            info = cand
                            break
            fn_id = re.search(r"(swufn_\d+)", key_text, re.I)
            # Important: for explicit SwUFn headings, do NOT fallback by ID.
            # Template SwUFn IDs can diverge from parsed source order and cause wrong function mapping.
            if (not heading_fn_name) and (not isinstance(info, dict)) and fn_id and isinstance(function_details, dict):
                info = function_details.get(fn_id.group(1).upper())
            if not isinstance(info, dict) and ":" in str(title_text):
                fn_name = heading_fn_name
                if fn_name:
                    sig = ""
                    for ln in (interface_functions or "").splitlines() + (internal_functions or "").splitlines():
                        ln = str(ln).strip()
                        if re.search(rf"\b{re.escape(fn_name)}\s*\(", ln, flags=re.I):
                            sig = ln
                            break
                    if sig:
                        info = {
                            "id": "",
                            "name": fn_name,
                            "prototype": sig,
                            "description": "",
                            "asil": "",
                            "related": "",
                            "inputs": _parse_signature_params(sig),
                            "outputs": _parse_signature_outputs(sig, fn_name),
                            "precondition": "",
                            "globals_global": [],
                            "globals_static": [],
                            "called": "",
                            "logic": "",
                        }
            # For explicit SwUFn headings, never consume queue fallback.
            # Queue fallback can incorrectly map another module function (e.g., Lib -> SwCom_01 main).
            if not isinstance(info, dict) and not heading_fn_name:
                ctx = " ".join([str(h).lower() for h in heading_stack])
                if "interface functions" in ctx and interface_queue:
                    info = interface_queue.pop(0)
                elif "internal functions" in ctx and internal_queue:
                    info = internal_queue.pop(0)
            if not isinstance(info, dict) and not heading_fn_name:
                info = {
                    "id": _normalize_swufn_id(str(fn_id.group(1)) if fn_id else ""),
                    "name": str(title_text).split(":", 1)[1].strip() if ":" in str(title_text) else str(title_text),
                    "prototype": "",
                    "description": "",
                    "asil": "",
                    "related": "",
                    "inputs": [],
                    "outputs": [],
                    "precondition": "",
                    "globals_global": [],
                    "globals_static": [],
                    "called": "",
                    "logic": "",
                }
            if not isinstance(info, dict) and heading_fn_name:
                called_text = ""
                if isinstance(call_map, dict):
                    target_norm = _normalize_symbol_name(str(heading_fn_name or "")).lower()
                    target_compact = re.sub(r"[^a-z0-9_]", "", target_norm)
                    for k, vals in call_map.items():
                        key_norm = _normalize_symbol_name(str(k or "")).lower()
                        key_compact = re.sub(r"[^a-z0-9_]", "", key_norm)
                        if not (
                            key_norm == heading_fn_name
                            or key_norm == target_norm
                            or (target_compact and key_compact == target_compact)
                        ):
                            continue
                        called_text = ", ".join([str(v) for v in (vals or []) if v])
                        break
                info = {
                    "id": _normalize_swufn_id(str(fn_id.group(1)) if fn_id else ""),
                    "name": heading_fn_name,
                    "prototype": "",
                    "description": _fallback_function_description(heading_fn_name, called_text),
                    "asil": "",
                    "related": "",
                    "inputs": [],
                    "outputs": [],
                    "precondition": "",
                    "globals_global": [],
                    "globals_static": [],
                    "called": called_text,
                    "logic": "",
                }
            if isinstance(info, dict) and heading_fn_name:
                if fn_id:
                    info["id"] = _normalize_swufn_id(str(fn_id.group(1)))
                if not str(info.get("prototype") or "").strip() and heading_fn_name == "main":
                    info["prototype"] = "void main( void )"
                if not str(info.get("description") or "").strip():
                    info["description"] = _fallback_function_description(
                        heading_fn_name,
                        info.get("called") or info.get("calls_list") or [],
                    )
                if not str(info.get("calling") or "").strip():
                    callers = callers_map.get(heading_fn_name, [])
                    lines: List[str] = []
                    for c in callers:
                        sig = ""
                        if isinstance(function_details_by_name, dict):
                            cinfo = function_details_by_name.get(str(c).lower())
                            if isinstance(cinfo, dict):
                                sig = str(cinfo.get("prototype") or "").strip()
                        lines.append(sig or str(c))
                    info["calling"] = "\n".join([ln for ln in lines if ln])
                ref_rel = ref_related_by_name.get(heading_fn_name)
                if heading_fn_name == "main" and ref_rel:
                    info["related"] = ref_rel
                    info["related_source"] = "reference"
                elif heading_fn_name == "main":
                    info["related"] = "SwST_01, SwCom_01, SwSTR_01, SwSTR_02, SwSTR_04, SwSTR_06, SwSTR_09"
                    info["related_source"] = "rule"
                else:
                    cur_related = str(info.get("related") or "").strip()
                    if cur_related in {"", "TBD", "SwCom_01"} and ref_rel:
                        info["related"] = ref_rel
                        info["related_source"] = "reference"
            return info

        def _build_function_info_table(info: Dict[str, Any], rows: int, cols: int, style: Any):
            if str(info.get("name") or "").strip().lower() == "main":
                info["related"] = "SwST_01, SwCom_01, SwSTR_01, SwSTR_02, SwSTR_04, SwSTR_06, SwSTR_09"
            fn_key = str(info.get("name") or "").strip().lower()
            callee_names = [str(c).strip() for c in (info.get("calls_list") or []) if str(c).strip()]
            if (not callee_names) and call_relation_mode == "code" and isinstance(call_map, dict):
                fn_norm = _normalize_symbol_name(str(info.get("name") or "")).lower()
                if fn_norm:
                    recovered: List[str] = []
                    for ck, vals in call_map.items():
                        if _normalize_symbol_name(str(ck or "")).lower() != fn_norm:
                            continue
                        if isinstance(vals, list):
                            recovered.extend([str(x).strip() for x in vals if str(x).strip()])
                    if recovered:
                        callee_names = list(dict.fromkeys(recovered))
            # Do not recover callee from existing "called" text.
            # Template/reference merge may inject caller-oriented text and
            # corrupt directionality for leaf functions.
            callee_names = list(dict.fromkeys(callee_names))
            caller_names = list(dict.fromkeys(callers_map.get(fn_key, [])))

            def _sig_lines(names: List[str]) -> List[str]:
                out: List[str] = []
                for nm in names:
                    sig = ""
                    cinfo = function_details_by_name.get(str(nm).lower()) if isinstance(function_details_by_name, dict) else None
                    if isinstance(cinfo, dict):
                        sig = str(cinfo.get("prototype") or "").strip()
                    out.append(sig or str(nm))
                return [x for x in out if x]

            callee_lines = _sig_lines(callee_names)
            caller_lines = _sig_lines(caller_names)
            if fn_key == "wake_up_setting":
                callee_name_set = {
                    str(x).strip().lower()
                    for x in _extract_call_names("\n".join(callee_lines))
                    if str(x).strip()
                }
                if (
                    "l_ifc_init" not in callee_name_set
                    and {"l_sys_init", "monitor_adc_enable", "monitor_adc_init"} & callee_name_set
                ):
                    l_ifc_sig = ""
                    cinfo = function_details_by_name.get("l_ifc_init") if isinstance(function_details_by_name, dict) else None
                    if isinstance(cinfo, dict):
                        l_ifc_sig = str(cinfo.get("prototype") or "").strip()
                    callee_lines = ([l_ifc_sig or "l_ifc_init"] + callee_lines)
                    dedup_lines: List[str] = []
                    for ln in callee_lines:
                        if ln and ln not in dedup_lines:
                            dedup_lines.append(ln)
                    callee_lines = dedup_lines
            if fn_key == "main" and not caller_lines:
                caller_lines = ["void _Startup(void)"]
            # Keep canonical relation direction in persisted fields:
            # called = callees, calling = callers.
            info["called"] = "\n".join(callee_lines) if callee_lines else "N/A"
            info["calling"] = "\n".join(caller_lines) if caller_lines else "N/A"
            if fn_key == "wake_up_setting":
                current_calling = str(info.get("calling") or "")
                parsed_names = {
                    str(x).strip().lower()
                    for x in _extract_call_names(current_calling)
                    if str(x).strip()
                }
                if (
                    "l_ifc_init" not in parsed_names
                    and {"l_sys_init", "monitor_adc_enable", "monitor_adc_init"} & parsed_names
                ):
                    lines_now = [ln for ln in current_calling.splitlines() if ln.strip()]
                    lines_now.insert(0, "l_ifc_init")
                    dedup_lines: List[str] = []
                    for ln in lines_now:
                        clean = str(ln).strip()
                        if clean and clean not in dedup_lines:
                            dedup_lines.append(clean)
                    info["calling"] = "\n".join(dedup_lines) if dedup_lines else current_calling

            def _format_globals(items: List[str]) -> List[str]:
                out: List[str] = []
                for name in items:
                    raw = str(name or "").strip()
                    tag = ""
                    base = raw
                    m = re.match(r"^\[(INOUT|IN|OUT)\]\s+(.+)$", raw)
                    if m:
                        tag = m.group(1)
                        base = m.group(2).strip()
                    lookup = re.split(r"(?:->|\.)", base)[0].strip()
                    display_name = f"[{tag}] {base}" if tag else base
                    info_map = globals_info_map.get(lookup, {})
                    if globals_format_with_labels:
                        mapping = {
                            "Name": f"Name={display_name}",
                            "Type": f"Type={info_map.get('type','')}",
                            "File": f"File={Path(info_map.get('file','')).name}" if info_map.get("file") else "File=",
                            "Range": f"Range={info_map.get('range','')}",
                        }
                        parts = [mapping.get(k, "") for k in globals_format_order]
                        out.append(globals_format_sep.join([p for p in parts if p]))
                    else:
                        mapping = {
                            "Name": display_name,
                            "Type": info_map.get("type", ""),
                            "File": Path(info_map.get("file", "")).name if info_map.get("file") else "",
                            "Range": info_map.get("range", ""),
                        }
                        parts = [mapping.get(k, "") for k in globals_format_order]
                        out.append(globals_format_sep.join([p for p in parts if p]))
                return out
            info["globals_global"] = _format_globals(info.get("globals_global") or [])
            info["globals_static"] = _format_globals(info.get("globals_static") or [])
            info_for_rows = dict(info)
            if payload.get("show_mapping_evidence"):
                info_for_rows["show_mapping_evidence"] = True
            data_rows = _build_function_info_rows(info_for_rows, cols)
            calls_list = list(dict.fromkeys(callee_names))
            if not calls_list:
                calls_list = _extract_call_names(str(info.get("called") or ""))
            logic_key = str(info.get("id") or "").strip()
            if not logic_key:
                logic_key = str(info.get("name") or "function").strip()
            logic_key = re.sub(r"[^A-Za-z0-9_.-]+", "_", logic_key).strip("_")
            if not logic_key:
                logic_key = f"function_{abs(hash(str(info)))%100000}"
            logic_flow = info.get("logic_flow") or []
            logic_img_path = Path(out).parent / "logic" / f"{logic_key}.png"
            logic_img = None
            if logic_flow:
                logic_img = _render_logic_flow_diagram(
                    str(info.get("name") or "Function"),
                    logic_flow,
                    logic_img_path,
                    all_calls=calls_list,
                    call_map=call_map if isinstance(call_map, dict) else None,
                )
            if not logic_img:
                logic_img = _render_call_graph_image(
                    str(info.get("name") or "Function"),
                    calls_list,
                    call_map if isinstance(call_map, dict) else None,
                    logic_img_path,
                    max_children=int(payload.get("logic_max_children") or LOGIC_MAX_CHILDREN_DEFAULT),
                    max_grandchildren=int(payload.get("logic_max_grandchildren") or LOGIC_MAX_GRANDCHILDREN_DEFAULT),
                    max_depth=int(payload.get("logic_max_depth") or LOGIC_MAX_DEPTH_DEFAULT),
                    module_map=module_map if isinstance(module_map, dict) else None,
                    condition_text=str(info.get("logic_condition") or ""),
                    true_path_text="\n".join([str(x) for x in (info.get("logic_true_calls") or []) if str(x).strip()]),
                    false_path_text="\n".join([str(x) for x in (info.get("logic_false_calls") or []) if str(x).strip()]),
                    return_path_text=str(info.get("logic_return_path") or ""),
                    error_path_text=str(info.get("logic_error_path") or ""),
                )
            if logic_img:
                for r in data_rows:
                    if r and "Logic Diagram" in r[0]:
                        r[1] = Path(logic_img).name
                        break
            data_rows = [["[ Function Information ]"] * cols] + data_rows
            func_table = _add_blank_table(doc, rows, cols, style, None, None)
            _merge_function_info_table(func_table, cols)
            _fill_function_info_table(func_table, data_rows)
            if logic_img:
                inserted = _insert_logic_image_in_table(func_table, cols, str(logic_img))
                if not inserted:
                    try:
                        from docx.shared import Inches  # type: ignore
                    except Exception:
                        Inches = None  # type: ignore
                    try:
                        doc.add_paragraph("Logic Diagram")
                        if Inches:
                            doc.add_picture(str(logic_img), width=Inches(5))
                        else:
                            doc.add_picture(str(logic_img))
                    except Exception as e:
                        _logger.warning("Failed to insert logic diagram: %s", e)
                        doc.add_paragraph("[Logic Diagram not available]")
            return func_table

        for idx, block in enumerate(blocks):
            kind, payload_block = block
            if kind == "heading":
                level, title = payload_block
                if len(heading_stack) >= level:
                    heading_stack = heading_stack[: level - 1]
                heading_stack.append(str(title))
                doc.add_heading(title, level=level)
                key = str(title).strip().lower()
                next_kind = _next_block_kind(blocks, idx)
                has_table = _section_has_table(blocks, idx, level)
                if key == "overview":
                    _add_docx_bullets(doc, overview)
                elif key in {
                    "introduction",
                    "purpose",
                    "scope",
                    "terms, abbreviations and definitions",
                    "reference",
                }:
                    template_text = template_section_map.get(key, "")
                    if template_text:
                        _add_docx_lines(doc, template_text)
                elif key == "requirements":
                    _add_docx_bullets(doc, requirements)
                elif key == "interfaces":
                    _add_docx_bullets(doc, interfaces)
                elif key == "uds frames":
                    _add_docx_bullets(doc, uds_frames)
                elif key == "contents":
                    if not toc_inserted:
                        _add_docx_toc(doc)
                        toc_inserted = True
                elif key == "notes":
                    _add_docx_bullets(doc, notes)
                elif key == "software unit design":
                    if next_kind != "table" and not has_table:
                        _add_docx_lines(doc, payload.get("software_unit_design", "") or "")
                elif key == "detailed uds":
                    _add_docx_text_block(doc, detailed_doc)
                elif key == "software unit structure":
                    # Keep this section text-only to avoid extra image
                    # between 2.5 and 2.5.1 (Software Unit Tables).
                    pass
                elif key == "unit structure":
                    swcom_id = ""
                    for item in reversed(heading_stack):
                        m = re.search(r"(SwCom_\d+)", item)
                        if m:
                            swcom_id = m.group(1)
                            break
                    interfaces_list: List[str] = []
                    internals_list: List[str] = []
                    if swcom_id and swcom_id in module_funcs:
                        interfaces_list = module_funcs[swcom_id].get("interfaces", [])
                        internals_list = module_funcs[swcom_id].get("internals", [])
                    if not interfaces_list:
                        interfaces_list = [
                            ln.strip()
                            for ln in (interface_functions or "").splitlines()
                            if ln.strip()
                        ][:8]
                    if not internals_list:
                        internals_list = [
                            ln.strip()
                            for ln in (internal_functions or "").splitlines()
                            if ln.strip()
                        ][:8]
                    structure_img = _render_unit_structure_image(
                        swcom_id or "Unit Structure",
                        interfaces_list,
                        internals_list,
                        Path(out).parent / "structure" / f"{swcom_id or 'unit'}.png",
                    )
                    if structure_img:
                        try:
                            from docx.shared import Inches  # type: ignore
                        except Exception:
                            Inches = None  # type: ignore
                        try:
                            doc.add_paragraph("Structure Diagram")
                            if Inches:
                                doc.add_picture(str(structure_img), width=Inches(5))
                            else:
                                doc.add_picture(str(structure_img))
                        except Exception as e:
                            _logger.warning("Failed to insert structure diagram: %s", e)
                            doc.add_paragraph("[Structure Diagram not available]")
                elif key == "global data":
                    pass
                elif key == "interface functions":
                    pass
                elif key == "internal functions":
                    pass
                else:
                    section_text = _extract_doc_section(
                        detailed_doc or software_unit_design, title
                    )
                    if section_text:
                        _add_docx_text_block(doc, section_text, max_lines=200)
                    else:
                        doc.add_paragraph("N/A")

                if re.search(r"\bswufn_\d+\b", key, flags=re.I):
                    target_idx = None
                    for look_ahead in range(idx + 1, len(blocks)):
                        kind_la, payload_la = blocks[look_ahead]
                        if kind_la == "heading":
                            try:
                                level_la = int(payload_la[0])
                            except Exception:
                                level_la = level
                            if level_la <= level:
                                break
                            continue
                        if kind_la == "table":
                            rows_la, cols_la, style_la, header_rows_la, _ctx_titles_la = payload_la
                            if header_rows_la:
                                header_texts = [str(c or "").strip() for row in header_rows_la for c in row]
                                if any("Function Information" in c for c in header_texts):
                                    target_idx = look_ahead
                                    break
                    swufn_match = re.search(r"(swufn_\\d+)", key, re.I)
                    if swufn_match and swufn_match.group(1).upper() in swufn_table_spec:
                        rows, cols, style = swufn_table_spec[swufn_match.group(1).upper()]
                    elif target_idx is not None:
                        rows, cols, style, _header_rows, _ctx_titles = blocks[target_idx][1]
                    elif function_info_template:
                        rows, cols, style, _header_rows = function_info_template
                    else:
                        rows, cols, style = 18, 6, None
                    info = _resolve_function_info(str(title), key)
                    _build_function_info_table(info, rows, cols, style)
                    if target_idx is not None:
                        skip_table_idx = target_idx
            elif kind == "table":
                if skip_table_idx == idx:
                    skip_table_idx = -1
                    continue
                rows, cols, style, header_rows, ctx_titles = payload_block
                ctx_text = " > ".join(ctx_titles).lower() if ctx_titles else ""
                current_swcom = _current_swcom_id()
                current_swcom_label = _current_swcom_label()
                data_rows: Optional[List[List[str]]] = None
                if header_rows:
                    header_texts = [str(c or "").strip() for row in header_rows for c in row]
                    if any("Function Information" in c for c in header_texts):
                        continue
                if "parameter definition" in ctx_text:
                    note = section_note_map.get("parameter definition", "")
                    if note and "parameter definition" not in note_added:
                        _add_docx_lines(doc, note)
                        note_added.add("parameter definition")
                elif "version information" in ctx_text:
                    note = section_note_map.get("version information", "")
                    if note and "version information" not in note_added:
                        _add_docx_lines(doc, note)
                        note_added.add("version information")
                if "software unit tables" in ctx_text:
                    data_rows = function_table_rows
                elif "common macro definition" in ctx_text:
                    data_rows = _table_rows_from_texts(common_macros, cols)
                elif "type definition" in ctx_text:
                    data_rows = _table_rows_from_texts(type_defs, cols)
                elif "parameter definition" in ctx_text:
                    data_rows = _table_rows_from_texts(param_defs, cols)
                elif "version information" in ctx_text:
                    data_rows = _table_rows_from_texts(version_defs, cols)
                elif "reference" in ctx_text:
                    refs = []
                    ref_files = payload.get("reference_files") or []
                    if isinstance(ref_files, list):
                        for name in ref_files:
                            if name:
                                refs.append([str(name), "", "", ""])
                    for ln in (notes or "").splitlines():
                        if ln.strip().startswith("doc:"):
                            refs.append([ln.strip()[4:], "", "", ""])
                    data_rows = refs if refs else None
                elif "global variables" in ctx_text:
                    header = header_rows[0] if header_rows else []
                    if globals_info_map and header:
                        names = [row[0] for row in global_vars if row]
                        names = _filter_global_names_by_swcom(
                            names,
                            current_swcom,
                            module_label=current_swcom_label,
                            static_only=False,
                        )
                        data_rows = _build_global_rows(
                            names,
                            globals_info_map,
                            header,
                            with_labels=False,
                        )
                    else:
                        data_rows = global_vars
                elif "static variables" in ctx_text:
                    header = header_rows[0] if header_rows else []
                    if globals_info_map and header:
                        names = [row[0] for row in static_vars if row]
                        names = _filter_global_names_by_swcom(
                            names,
                            current_swcom,
                            module_label=current_swcom_label,
                            static_only=True,
                        )
                        data_rows = _build_global_rows(
                            names,
                            globals_info_map,
                            header,
                            with_labels=False,
                        )
                    else:
                        data_rows = static_vars
                elif "calibration & parameter data" in ctx_text:
                    data_rows = _filter_rows_by_swcom(calibration_params, current_swcom, current_swcom_label)
                    if not data_rows and cols >= 4:
                        data_rows = [["N/A", "N/A", "N/A", "N/A"]]
                elif "macro" in ctx_text:
                    data_rows = []
                    for row in macro_defs:
                        if len(row) >= 4:
                            name, mtype, val, desc = row[0], row[1], row[2], row[3]
                        else:
                            name = row[0] if row else ""
                            mtype = ""
                            val = row[2] if len(row) > 2 else ""
                            desc = row[3] if len(row) > 3 else ""
                        if not mtype:
                            mtype = "Macro"
                        data_rows.append([name, mtype, val, desc])
                    data_rows = _filter_rows_by_swcom(data_rows, current_swcom, current_swcom_label)
                    if not data_rows and cols >= 4:
                        data_rows = [["N/A", "N/A", "N/A", "N/A"]]
                _add_blank_table(doc, rows, cols, style, header_rows, data_rows)
            elif kind == "para":
                text = str(payload_block.get("text") or "").strip()
                if not text:
                    continue
                if text.upper() == "N/A":
                    continue
                if heading_stack:
                    continue
                if re.match(r"^(Interfaces:|Internals:|Global data:)\s*\d+", text, flags=re.I):
                    continue
                style_name = str(payload_block.get("style") or "").strip()
                try:
                    if style_name:
                        doc.add_paragraph(text, style=style_name)
                    else:
                        doc.add_paragraph(text)
                except Exception:
                    doc.add_paragraph(text)
                if text.lower() == "contents" and not toc_inserted:
                    _add_docx_toc(doc)
                    toc_inserted = True
        _normalize_function_info_tables(doc)
        _remove_docx_paragraphs(doc, ["N/A"])
        doc.save(str(out))
        return str(out)

    # ── No-template fallback: SUDS-compatible 4-level structure ──
    doc = docx.Document()
    cols = 6  # Function info table column count

    # ── Helper functions needed in fallback (template path has these in its scope) ──

    # Build callers_map from call_map (callee → [callers])
    _fb_callers_map: Dict[str, List[str]] = {}
    if isinstance(call_map, dict):
        for _fb_caller, _fb_callees in call_map.items():
            _fb_cname = _normalize_symbol_name(str(_fb_caller or "")).lower()
            if not _fb_cname or not isinstance(_fb_callees, list):
                continue
            for _fb_callee in _fb_callees:
                _fb_ce = _normalize_symbol_name(str(_fb_callee or "")).lower()
                if _fb_ce:
                    if _fb_ce not in _fb_callers_map:
                        _fb_callers_map[_fb_ce] = []
                    if _fb_cname not in _fb_callers_map[_fb_ce]:
                        _fb_callers_map[_fb_ce].append(_fb_cname)

    # Build swcom_global_hints and swcom_function_files for filtering
    _fb_swcom_global_hints: Dict[str, Set[str]] = {}
    _fb_swcom_function_files: Dict[str, Set[str]] = {}
    if isinstance(function_table_rows, list):
        for _fb_row in function_table_rows:
            if not isinstance(_fb_row, list) or len(_fb_row) < 4:
                continue
            _fb_sc = str(_fb_row[0] or "").strip()
            _fb_fn = str(_fb_row[3] or "").strip()
            if not _fb_sc or not _fb_fn:
                continue
            _fb_fi = function_details_by_name.get(_fb_fn.lower()) if isinstance(function_details_by_name, dict) else None
            if isinstance(_fb_fi, dict):
                _fb_fp = str(_fb_fi.get("file") or "").strip()
                if _fb_fp:
                    _fb_swcom_function_files.setdefault(_fb_sc, set()).add(Path(_fb_fp).name.lower())
                for _fb_gk in ["globals_global", "globals_static"]:
                    for _fb_g in (_fb_fi.get(_fb_gk) or []):
                        _fb_graw = str(_fb_g or "").strip()
                        if not _fb_graw:
                            continue
                        _fb_gtag = re.match(r"^\[(?:INOUT|IN|OUT)\]\s+(.+)$", _fb_graw)
                        _fb_gbase = _fb_gtag.group(1).strip() if _fb_gtag else _fb_graw
                        _fb_gn0 = re.split(r"(?:->|\.)", _fb_gbase)[0].strip()
                        if _fb_gn0:
                            _fb_swcom_global_hints.setdefault(_fb_sc, set()).add(_fb_gn0)

    def _fb_is_register_alias(name: str) -> bool:
        return bool(re.match(r"^REG_[A-Z0-9_]+$", str(name or "").strip()))

    def _fb_norm_stem(name: str) -> str:
        s = re.sub(r"[^a-z0-9]+", "", str(name or "").lower())
        for _sfx in ["itpds", "pds", "it", "main"]:
            if s.endswith(_sfx) and len(s) > len(_sfx) + 2:
                s = s[: -len(_sfx)]
        return s

    def _filter_global_names_by_swcom(
        all_names: List[str],
        swcom_id: str,
        module_label: str = "",
        static_only: Optional[bool] = None,
    ) -> List[str]:
        names = [str(n).strip() for n in all_names if str(n).strip() and not _fb_is_register_alias(str(n))]
        module_candidates: Set[str] = set()
        label_norm = re.sub(r"[^a-z0-9]+", "", str(module_label or "").lower())
        if label_norm:
            module_candidates.add(label_norm)
            module_candidates.add(label_norm.replace("system", "sys"))
            if "system" in label_norm and "os" in label_norm:
                module_candidates.add("sysos")
        selected: List[str] = []
        if module_candidates:
            for n in names:
                info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
                is_static = str(info_g.get("static") or "").strip().lower() == "true"
                if static_only is True and not is_static:
                    continue
                if static_only is False and is_static:
                    continue
                gfile_l = str(info_g.get("file") or "").lower()
                gblob = re.sub(r"[^a-z0-9]+", "", gfile_l)
                if any(tok and tok in gblob for tok in module_candidates):
                    selected.append(n)
            if selected:
                return list(dict.fromkeys(selected))
        hints = _fb_swcom_global_hints.get(swcom_id, set())
        for n in names:
            info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
            is_static = str(info_g.get("static") or "").strip().lower() == "true"
            if static_only is True and not is_static:
                continue
            if static_only is False and is_static:
                continue
            if n in hints:
                selected.append(n)
        if selected:
            return list(dict.fromkeys(selected))
        file_names = _fb_swcom_function_files.get(swcom_id, set())
        if not file_names:
            return names
        file_stems = {_fb_norm_stem(Path(x).stem) for x in file_names}
        for n in names:
            info_g = globals_info_map.get(n, {}) if isinstance(globals_info_map, dict) else {}
            is_static = str(info_g.get("static") or "").strip().lower() == "true"
            if static_only is True and not is_static:
                continue
            if static_only is False and is_static:
                continue
            gfile = str(info_g.get("file") or "").strip()
            gstem = _fb_norm_stem(Path(gfile).stem) if gfile else ""
            if gstem and any(gstem == fs or gstem.startswith(fs) or fs.startswith(gstem) for fs in file_stems):
                selected.append(n)
        return list(dict.fromkeys(selected))

    def _filter_rows_by_swcom(rows_in: List[List[str]], swcom_id: str, module_label: str = "") -> List[List[str]]:
        hints = _fb_swcom_global_hints.get(swcom_id, set())
        if not hints and module_label:
            lbl = re.sub(r"[^a-z0-9]+", "", module_label.lower())
            if "system" in lbl and "os" in lbl:
                hints = {"u8g_SystemTm_5ms", "u8g_SystemTm_10ms", "u8g_SystemTm_50ms", "u8s_InitiComplet_F"}
        if not hints:
            return []
        out_rows: List[List[str]] = []
        for row in (rows_in or []):
            if not isinstance(row, list) or not row:
                continue
            blob = " ".join([str(x or "") for x in row])
            if any(re.search(rf"\b{re.escape(h)}\b", blob) for h in hints):
                out_rows.append(row)
        return out_rows

    def _enhance_function_desc_with_ai(info: Dict[str, Any]) -> str:
        """Call AI to generate a better function description (1-3 sentences, Korean).

        Only called when ai_config is available and existing description is short (<30 chars).
        Returns enhanced description string, or empty string on failure.
        """
        if not ai_config:
            return ""
        try:
            from workflow.ai import agent_call  # type: ignore
        except ImportError:
            return ""

        _UDS_AI_SYSTEM_PROMPT = (
            "당신은 자동차 ECU 소프트웨어 단위 설계 명세서(SUDS) 작성 전문가입니다. "
            "ISO 26262 ASIL 기준에 따라 C 함수의 기술적 설명을 1~3문장으로 작성합니다.\n"
            "규칙:\n"
            "- 함수의 목적, 입출력, 주요 동작을 포함하세요.\n"
            "- 한국어로 작성하세요.\n"
            "- JSON 형식으로 반환: {\"description\": \"...\", \"purpose\": \"...\"}"
        )
        fn_name = info.get("name", "")
        prototype = info.get("prototype", "")
        inputs = ", ".join(str(i) for i in (info.get("inputs") or [])[:5])
        outputs = str(info.get("output") or "")
        calls = ", ".join(str(c) for c in (info.get("calls_list") or [])[:5])
        user_msg = (
            f"함수명: {fn_name}\n"
            f"원형: {prototype}\n"
            f"입력: {inputs or 'N/A'}\n"
            f"출력: {outputs or 'N/A'}\n"
            f"호출 함수: {calls or 'N/A'}\n"
            "위 정보를 바탕으로 기술적 함수 설명을 JSON으로 반환하세요."
        )
        import threading, json as _json
        result_holder: Dict[str, str] = {}
        def _call():
            try:
                r = agent_call(
                    ai_config,
                    [{"role": "system", "content": _UDS_AI_SYSTEM_PROMPT},
                     {"role": "user", "content": user_msg}],
                    stage="uds_enhance",
                )
                raw = (r.get("output") or "").strip()
                # Try JSON extraction
                m = re.search(r"\{[^{}]+\}", raw, re.DOTALL)
                if m:
                    parsed = _json.loads(m.group())
                    result_holder["desc"] = str(parsed.get("description") or parsed.get("purpose") or "")
                else:
                    result_holder["desc"] = raw[:200]
            except Exception:
                pass
        t = threading.Thread(target=_call, daemon=True)
        t.start()
        t.join(timeout=20)
        return result_holder.get("desc", "")

    def _build_function_info_table(info: Dict[str, Any], rows: int, _cols: int, style: Any) -> None:
        """Build and append a function info table to doc."""
        _fn_key = str(info.get("name") or "").strip().lower()
        # Resolve called/calling
        _callee_names = list(dict.fromkeys(
            _normalize_symbol_name(str(c)).lower()
            for c in (info.get("calls_list") or []) if str(c).strip()
        ))
        if not _callee_names and isinstance(call_map, dict):
            _fn_norm = _normalize_symbol_name(_fn_key).lower()
            for _ck, _vals in call_map.items():
                if _normalize_symbol_name(str(_ck or "")).lower() == _fn_norm and isinstance(_vals, list):
                    _callee_names = list(dict.fromkeys(str(x).strip() for x in _vals if str(x).strip()))
                    break
        _caller_names = list(dict.fromkeys(_fb_callers_map.get(_fn_key, [])))

        def _sig_lines(names: List[str]) -> List[str]:
            out: List[str] = []
            for nm in names:
                sig = ""
                _ci = function_details_by_name.get(str(nm).lower()) if isinstance(function_details_by_name, dict) else None
                if isinstance(_ci, dict):
                    sig = str(_ci.get("prototype") or "").strip()
                out.append(sig or str(nm))
            return [x for x in out if x]

        _inf2 = dict(info)
        _inf2["called"] = "\n".join(_sig_lines(_callee_names)) if _callee_names else "N/A"
        if not str(_inf2.get("calling") or "").strip():
            _inf2["calling"] = "\n".join(_sig_lines(_caller_names)) if _caller_names else "N/A"
        if _fn_key == "main" and not str(_inf2.get("calling") or "").strip().replace("N/A", ""):
            _inf2["calling"] = "void _Startup(void)"
        if _fn_key == "main":
            _inf2["related"] = _inf2.get("related") or "SwST_01, SwCom_01, SwSTR_01, SwSTR_02, SwSTR_04, SwSTR_06, SwSTR_09"

        # AI-enhance description if it's short
        _existing_desc = str(_inf2.get("description") or _inf2.get("desc") or "").strip()
        if ai_config and len(_existing_desc) < 30:
            _ai_desc = _enhance_function_desc_with_ai(_inf2)
            if _ai_desc and len(_ai_desc) > len(_existing_desc):
                _inf2["description"] = _ai_desc

        _data_rows = _build_function_info_rows(_inf2, _cols)
        # Attempt logic diagram
        _logic_key = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(_inf2.get("id") or _fn_key or "fn")).strip("_")
        _logic_img_path = Path(out).parent / "logic" / f"{_logic_key}.png"
        _calls_list2 = _callee_names or _extract_call_names(str(_inf2.get("called") or ""))
        _logic_img = _render_call_graph_image(
            str(_inf2.get("name") or "Function"),
            _calls_list2,
            call_map if isinstance(call_map, dict) else None,
            _logic_img_path,
            max_children=int(payload.get("logic_max_children") or LOGIC_MAX_CHILDREN_DEFAULT),
            max_grandchildren=int(payload.get("logic_max_grandchildren") or LOGIC_MAX_GRANDCHILDREN_DEFAULT),
            max_depth=int(payload.get("logic_max_depth") or LOGIC_MAX_DEPTH_DEFAULT),
            module_map=module_map if isinstance(module_map, dict) else None,
        )
        if _logic_img:
            for _r in _data_rows:
                if _r and "Logic Diagram" in _r[0]:
                    _r[1] = Path(_logic_img).name
                    break
        _data_rows = [["[ Function Information ]"] * _cols] + _data_rows
        _rows_per_fn = max(len(_data_rows), rows)
        _ft = _add_blank_table(doc, _rows_per_fn, _cols, style, None, None)
        _merge_function_info_table(_ft, _cols)
        _fill_function_info_table(_ft, _data_rows)
        if _logic_img:
            if not _insert_logic_image_in_table(_ft, _cols, str(_logic_img)):
                try:
                    from docx.shared import Inches as _I  # type: ignore
                    doc.add_paragraph("Logic Diagram")
                    doc.add_picture(str(_logic_img), width=_I(5))
                except Exception:
                    doc.add_paragraph("[Logic Diagram not available]")

    # ── Cover ──
    doc.add_heading("Software Unit Design Specification", level=1)
    doc.add_paragraph(f"Project: {project}")
    doc.add_paragraph(f"Generated at: {generated_at}")
    if payload.get("job_url"):
        doc.add_paragraph(f"Job URL: {payload.get('job_url')}")
    if payload.get("build_number"):
        doc.add_paragraph(f"Build: {payload.get('build_number')}")

    # ── Revision History ──
    doc.add_heading("Revision History", level=2)
    _add_blank_table(
        doc, 3, 6, None,
        [["Version", "Date", "Description", "Author", "Reviewer", "Approver"]],
        [["", "", "", "", "", ""], ["", "", "", "", "", ""]],
    )

    # ── Introduction ──
    doc.add_heading("Introduction", level=2)
    for _subsec in ["Purpose", "Scope", "Terms, Abbreviations and Definitions", "Reference"]:
        doc.add_heading(_subsec, level=3)
        doc.add_paragraph("N/A")

    # ── Contents ──
    doc.add_heading("Contents", level=2)
    _add_docx_toc(doc)

    # ── Common Macro Definition ──
    if common_macros:
        doc.add_heading("Common Macro Definition", level=2)
        _cm_rows = _table_rows_from_texts(common_macros, 4)
        _add_blank_table(doc, max(len(_cm_rows) + 1, 2), 4, None,
                         [["Macro name", "Type", "Define", "Description"]], _cm_rows)

    # ── Type Definition ──
    if type_defs:
        doc.add_heading("Type Definition", level=2)
        _td_rows = _table_rows_from_texts(type_defs, 4)
        _add_blank_table(doc, max(len(_td_rows) + 1, 2), 4, None,
                         [["Type", "Define", "Range", "Description"]], _td_rows)

    # ── Parameter Definition ──
    if param_defs:
        doc.add_heading("Parameter Definition", level=2)
        _pd_rows = _table_rows_from_texts(param_defs, 4)
        _add_blank_table(doc, max(len(_pd_rows) + 1, 2), 4, None,
                         [["Parameter", "Type", "Define", "Description"]], _pd_rows)

    # ── Version Information ──
    if version_defs:
        doc.add_heading("Version Information", level=2)
        _vd_rows = _table_rows_from_texts(version_defs, 4)
        _add_blank_table(doc, max(len(_vd_rows) + 1, 2), 4, None,
                         [["File Name", "Version", "Date", "Note"]], _vd_rows)

    # ── Software Unit Structure ──
    doc.add_heading("Software Unit Structure", level=2)
    doc.add_heading("Software Unit Tables", level=3)
    if function_table_rows:
        _su_data = [
            [str(r[0] or ""), str(r[3] or ""), ""]
            for r in function_table_rows
            if isinstance(r, list) and len(r) >= 4
        ]
        _add_blank_table(doc, len(_su_data) + 1, 3, None,
                         [["Component", "Function", "Comment"]], _su_data)

    # ── Build SwCom grouping from function_table_rows ──
    _swcom_order: List[str] = []
    _swcom_func_map: Dict[str, Dict[str, List]] = {}
    for _row in (function_table_rows or []):
        if not isinstance(_row, list) or len(_row) < 4:
            continue
        _sc = str(_row[0] or "").strip()
        _fid = str(_row[2] or "").strip()
        _fn = str(_row[3] or "").strip()
        _ft = str(_row[4] or "").strip().lower() if len(_row) > 4 else ""
        if not _sc or not _fn:
            continue
        if _sc not in _swcom_func_map:
            _swcom_order.append(_sc)
            _swcom_func_map[_sc] = {"interfaces": [], "internals": []}
        _is_iface = "i/f" in _ft or _ft == "if"
        (_swcom_func_map[_sc]["interfaces"] if _is_iface else _swcom_func_map[_sc]["internals"]).append((_fid, _fn))

    # Fallback: no function_table_rows — put everything under SwCom_Unknown
    if not _swcom_order and isinstance(function_details, dict) and function_details:
        _swcom_order = ["SwCom_Unknown"]
        _swcom_func_map["SwCom_Unknown"] = {"interfaces": [], "internals": []}
        for _fid_k, _inf_v in function_details.items():
            if not isinstance(_inf_v, dict):
                continue
            _swcom_func_map["SwCom_Unknown"]["internals"].append(
                (str(_fid_k).strip(), str(_inf_v.get("name") or "").strip())
            )

    _GLOBAL_HDR = ["Name", "Type", "Value Range", "Reset Value", "Description"]
    _MACRO_HDR  = ["Name", "Type", "Value", "Description"]
    _CALIB_HDR  = ["Signal", "Type", "Value", "Description"]

    def _na_row(n: int) -> List[str]:
        return ["N/A"] + [""] * (n - 1)

    # ── Software Unit Design (H2: SwCom, H3: sections, H4: functions) ──
    doc.add_heading("Software Unit Design", level=2)

    fn_added = 0
    for _swcom_id in _swcom_order:
        _sc_label = ""
        # Try to get component label from sds_partition_map
        if isinstance(sds_partition_map, dict):
            for _k, _v in sds_partition_map.items():
                if _swcom_id.lower() in _k.lower() or _k.lower() in _swcom_id.lower():
                    _sc_label = str(_v.get("name") or _v.get("description") or "").strip()
                    break
        _sc_heading = f"{_swcom_id} ({_sc_label})" if _sc_label else _swcom_id
        doc.add_heading(_sc_heading, level=2)

        # ── Unit Structure ──
        doc.add_heading("Unit Structure", level=3)
        _iface_names = [_fn for _, _fn in _swcom_func_map[_swcom_id]["interfaces"]]
        _intern_names = [_fn for _, _fn in _swcom_func_map[_swcom_id]["internals"]]
        _struct_img = _render_unit_structure_image(
            _swcom_id, _iface_names, _intern_names,
            Path(out).parent / "structure" / f"{_swcom_id}.png",
        )
        if _struct_img:
            try:
                from docx.shared import Inches as _Inches  # type: ignore
                doc.add_picture(str(_struct_img), width=_Inches(5))
            except Exception:
                doc.add_paragraph(f"[Unit Structure: {_swcom_id}]")
        else:
            doc.add_paragraph(
                f"Interface Functions: {len(_iface_names)}, "
                f"Internal Functions: {len(_intern_names)}"
            )

        # ── Global Data ──
        doc.add_heading("Global Data", level=3)

        # Global variables
        doc.add_heading("Global variables", level=4)
        _gnames_all = [_r[0] for _r in (global_vars or []) if _r]
        _gnames = _filter_global_names_by_swcom(_gnames_all, _swcom_id, _sc_label, static_only=False)
        if _gnames and globals_info_map:
            _g_rows = _build_global_rows(_gnames, globals_info_map, _GLOBAL_HDR, with_labels=False)
        else:
            _g_rows = [list(_r) for _r in (global_vars or []) if _r]
        _add_blank_table(doc, max(len(_g_rows) + 1, 2), 5, None,
                         [_GLOBAL_HDR], _g_rows or [_na_row(5)])

        # Static Variables
        doc.add_heading("Static Variables", level=4)
        _snames_all = [_r[0] for _r in (static_vars or []) if _r]
        _snames = _filter_global_names_by_swcom(_snames_all, _swcom_id, _sc_label, static_only=True)
        if _snames and globals_info_map:
            _s_rows = _build_global_rows(_snames, globals_info_map, _GLOBAL_HDR, with_labels=False)
        else:
            _s_rows = [list(_r) for _r in (static_vars or []) if _r]
        _add_blank_table(doc, max(len(_s_rows) + 1, 2), 5, None,
                         [_GLOBAL_HDR], _s_rows or [_na_row(5)])

        # Macro
        doc.add_heading("Macro", level=4)
        _m_rows = _filter_rows_by_swcom(list(macro_defs or []), _swcom_id, _sc_label)
        _add_blank_table(doc, max(len(_m_rows) + 1, 2), 4, None,
                         [_MACRO_HDR], _m_rows or [_na_row(4)])

        # Calibration & Parameter Data
        doc.add_heading("Calibration & Parameter Data", level=4)
        _c_rows = _filter_rows_by_swcom(list(calibration_params or []), _swcom_id, _sc_label)
        _add_blank_table(doc, max(len(_c_rows) + 1, 2), 4, None,
                         [_CALIB_HDR], _c_rows or [_na_row(4)])

        def _write_fn_section(fn_pairs: List, heading: str, h_level: int) -> None:
            nonlocal fn_added
            if not fn_pairs:
                return
            doc.add_heading(heading, level=h_level)
            for _fid2, _fname2 in fn_pairs:
                _h = f"{_fid2}: {_fname2}" if _fid2 else _fname2
                doc.add_heading(_h, level=h_level + 1)
                _inf = (
                    (function_details.get(_fid2) if isinstance(function_details, dict) else None)
                    or (function_details_by_name.get(_fname2.lower()) if isinstance(function_details_by_name, dict) else None)
                    or {"id": _fid2, "name": _fname2}
                )
                if not isinstance(_inf, dict):
                    _inf = {"id": _fid2, "name": _fname2}
                _build_function_info_table(_inf, 18, cols, None)
                fn_added += 1

        _write_fn_section(_swcom_func_map[_swcom_id]["interfaces"], "Interface Functions", 3)
        _write_fn_section(_swcom_func_map[_swcom_id]["internals"], "Internal Functions", 3)

    # ── Logic Diagrams (optional) ──
    logic_items = payload.get("logic_diagrams") if isinstance(payload, dict) else []
    logic_items = _merge_logic_ai_items(logic_items, ai_sections)
    if isinstance(logic_items, list) and logic_items:
        try:
            from docx.shared import Inches  # type: ignore
        except Exception:
            Inches = None  # type: ignore
        doc.add_heading("Logic Diagrams", level=2)
        for item in logic_items:
            if not isinstance(item, dict):
                continue
            title = item.get("title") or "Logic Diagram"
            path = item.get("path")
            desc = item.get("description") or ""
            if title:
                doc.add_paragraph(str(title))
            if path:
                try:
                    if Inches:
                        doc.add_picture(str(path), width=Inches(5))
                    else:
                        doc.add_picture(str(path))
                except Exception:
                    continue
            if desc:
                doc.add_paragraph(str(desc))

    doc.save(str(out))
    return str(out)


