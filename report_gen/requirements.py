"""report_gen.requirements - Auto-split from report_generator.py"""
# Re-import common dependencies
import csv
import json
import logging
import re
from io import StringIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from report_gen.function_analyzer import _normalize_symbol_name
from report_gen.source_parser import (
    _extract_comment_lines,
    _scan_source_function_names,
    _scan_source_requirement_ids,
)
from report_gen.trace_integrity import build_integrity_audit
from report_gen.utils import (
    _dedupe_multiline_text,
    _normalize_asil_value,
    _normalize_related_ids,
    _normalize_swcom_label,
)

_logger = logging.getLogger("report_generator")

# SwUFn/SwIFn 단위·통합 함수 ID 패턴 — SITS/VectorCAST 2-hop bridge에서 행마다
# 재사용(7천+ 행 루프 재컴파일 방지, reviewer INFO 권고).
_SWUFN_RE = re.compile(r"Sw[UI]Fn_\d+", re.IGNORECASE)

# 설계-ID bridge(SRS→SDS→UDS) 대상 namespace — SwFn/SwSTR/SwST/SwTK만.
# _normalize_req_id(대문자화) 출력에 매칭. SwCom(loose·컴포넌트 fan-out 중앙 101)과
# SwUFn(함수 자기 ID)은 구조적 제외. SRS 요구(SwTR/SwEI/SwTSR 등)도 접두가 달라 불매치.
# 앵커 ^…$로 부분매칭 차단. 실측(KJPDS02_PV): 이 집합으로 UDS 밴드 43→64(+21), fan-out
# 중앙 22(현행 44보다 낮음). SwCom 포함 시 reach 불변(64)·fan-out 중앙 56으로 악화.
_DESIGN_ID_BRIDGE_RE = re.compile(r"^SW(?:FN|STR|ST|TK)_\d+$")

# 설계-ID bridge 방어 필터(deep-review C1): UDS Function Information 표의 값 셀이 라벨을
# echo하는 junk 표에서 func_name='Name'/func_id='ID' 등이 harvest돼 source_ids에 섞이면
# bridge가 이를 '함수'로 요구에 부착(over-trace). 파서 echo 가드가 1차 차단하나, 문서군
# 편차에 대비해 조립부에서도 필드 라벨을 요구 추적 함수에서 제외한다(방어심층).
_UDS_FUNC_JUNK = frozenset({
    "id", "name", "type", "asil", "reuse", "prototype", "description",
    "no", "related id", "cyber security",
})

# 미추적 VectorCAST subprogram 의미 분류용 — ISR/인터럽트/부트 핸들러 등
# 'SRS 추적 대상이 아닌 게 당연한' 인프라 함수를 식별(트리 미추적 루트의 isr 버킷).
# 정밀도 우선(reviewer WARNING): interrupt/exception/trap/fault 같은 무경계 부분일치는
# 안전 관련 일반 함수(Fault_Monitor, Entrapment_Detect=끼임감지, Bootstrap_Init 등)를
# ISR로 오분류해 검토 신호를 숨기므로 제외. 모든 토큰을 앵커/경계로만 매치한다
# (_irq도 _irq$/_irq_로 앵커 — config_irqd 류 무경계 부분일치 제거, 라운드 재검증 I4).
_ISR_RE = re.compile(r"(_isr$|_isr_|\bisr\b|_irq$|_irq_|_handler$|\bnmi\b)", re.I)

# isr 버킷은 warn=false('추적 대상 아님이 정상인 인프라')다. 그러나 _handler$ 앵커는
# Safety_Fault_Handler·Diag_Trap_Handler 같은 안전 핸들러도 매치하므로(라운드 재검증 W4),
# 안전/진단 토큰을 가진 함수는 isr로 침묵 강등하지 않고 vcast_only로 둬 검토 신호를 보존한다.
# 라운드111 보강: ISO 26262-5 하드웨어 자가진단(CPU/레지스터/스택/메모리 무결성 시험)과
# 방어적 가드 연산(overflow-protected arithmetic)이 누락돼 있었다 — s_StackGuardCheck·
# s_CPUInstructionTest·s_RegisterCheck·s_MCUErrorCheck·*_ECCerror·ROM/RAM/EEPROM_Test·
# *_Guarded(11개) 등 20개 ASIL 안전기제가 amber 검토우선에 안 걸렸다. 실데이터(1005함수)에서
# 신규 매치 20개 전부 진짜 안전기제(거짓양성 0) 확인 후 토큰 추가. 토큰은 충분히 구체적이라
# (registercheck/cpuinstr/rom_test 등) 일반 함수 오탐 위험 낮음.
# 라운드112 보강: ISO 26262-3/-6 통신·클록·범위 안전기제가 추가로 under-flag돼 있었다 —
# E2E AC/CRC 프로파일 체크(AUTOSAR E2E protection), CRC 무결성, lib *bit_rangecheck(방어적
# 범위검사), CPU PLL/OSC/clock status 감시가 빠져 있었다. 실데이터(2088함수 universe)에서
# 신규 매치 23개 전부 진짜 안전/무결성 기제(거짓양성 0) 확인 후 토큰 추가.
#  (?<![a-z])e2e : 'writE2Eeprom'(write2eeprom) 속 'e2e' substring 오탐 차단, e2e_*/e2eprofile 유지.
#  crc/rangecheck/cpupll/cpuosc/pllstatus/oscstatus/failurecheck : 모두 안전·무결성 의미가
#  내재된 구체 토큰이라 경계 없이도 오탐 없음(linfailchecktimer 등 진짜 TP 포착).
_SAFETY_TOKEN_RE = re.compile(
    # 워드 경계 정밀화(라운드111 reviewer): substring 오탐 차단.
    #  (?<!de)fault : 'default'(handleDefault) 속 'fault' 차단 (s_MotorBattShortRun_HandleDefault).
    #  (?<![a-z])trap : 'strap'/'bootstrap' 차단, '_trap'/'Diag_Trap_Handler'는 유지.
    #  eeprom_?test|(?<![a-z])rom/ram_?test : 'from_test'/'paramtest'/'histogram_test' 차단,
    #     메모리 자가시험 ROM/RAM/EEPROM_Test 유지.
    r"((?<!de)fault|diag|safety|monitor|watchdog|wdg|(?<![a-z])trap|brake|steer|airbag|crash|asil"
    r"|guard|selftest|self_test|stackover|stackchk|cpuinstr|instructiontest|registercheck"
    r"|register_check|mcuerror|eeprom_?test|(?<![a-z])rom_?test|(?<![a-z])ram_?test|eccerr|integrit"
    # 라운드112: 통신(E2E/CRC)·클록(PLL/OSC)·범위·실패감시 안전기제.
    r"|(?<![a-z])e2e|crc|rangecheck|range_check|pllstatus|oscstatus|clockcheck|clockstatus"
    r"|cpustatus|cpupll|cpuosc|failurecheck|failcheck)",
    re.I,
)

# 미추적 VectorCAST FAIL 판정 — 프론트(TraceUnmappedRoot failTotal/failN)의
# /^(fail|failed|false|0)$/i 와 동일. dedup 시 FAIL을 PASS보다 우선 보존(W2).
_RESULT_FAIL_RE = re.compile(r"^(fail|failed|false|0)$", re.I)

# SDS 컴포넌트명 → 함수명 bridge 키 정규화 (라운드 109 fix).
# SDS 추출(sds/extract-mapping)이 함수명에 C 시그니처 조각('s_systemhashcalculate( void'),
# 배열 첨자('u8g_..._partnoinfo[10]'), 표 파싱 아티팩트('33\tswcom_35: ...\t115')를 붙여
# 와서, 정확매칭 bridge(sds_func_to_reqs)가 실제 함수를 못 찾고 SRS 추적을 silent 누락한다.
# 실데이터 KJPDS02: 's_systemhashcalculate( void'가 14개 SRS 요구사항에 귀속됐으나 노이즈로
# 미추적(suts_tested)으로 오분류 → 정규화하면 정상 추적 복구. fuzzy 아님: 괄호 '이전'만,
# 첨자만 제거하므로 'mcu 이상 감지(레지스터 미지원)' 같은 한글 설명문은 키가 그대로 비매칭.
_SDS_TABLE_ARTIFACT_RE = re.compile(r"^\d+\s*\t")   # 선행 '행번호+탭'
_SDS_ARRAY_SUBSCRIPT_RE = re.compile(r"\[[^\]]*\]")  # 배열 첨자 [10]/[]
_C_IDENT_RE = re.compile(r"[a-z_][a-z0-9_]*\Z")      # C 식별자(소문자화 후)

# §C(공백분리 반환형): SDS가 component_id에 함수 프로토타입을 통째로 실어보내면
# ('void f( void )', 'static u16 s_calc( void )') 반환형/저장한정자가 함수명 앞에 공백으로
# 분리돼 붙는다. 마지막 토큰(함수명)만 채택하되 **선행이 전부 타입/한정자이고 그중 실제
# 타입이 최소 1개**일 때만 한다(_sds_comp_key 본문). 'reset counter'/'do something(x)'의 끝
# 단어를 함수명으로 오채택하면 곧바로 over-trace(설명문→SRS 거짓추적)이므로, 열린 last-token
# 채택을 금지하고 이 닫힌 화이트리스트로 게이트한다. 붙은 반환형('u16s_foo')은 §A 담당(별개).
#
# ⚠ 화이트리스트는 **자연어 영어 단어와 겹치지 않는 것만** 담는다(deep-review W1, 재현·실증):
# 기본 C 타입 int/long/short/char/double/float/signed/unsigned/bool 과 벤더 typedef
# word/byte/dword/boolean 은 'long delay( ms )'·'double click( x )'·'byte order( swap )' 같은
# 설계 라벨의 첫 단어와 겹쳐, 괄호가 뒤따르면 함수 delay/click/order 에 거짓 추적된다(auto·
# register 를 같은 이유로 뺐던 가드 논리를 이 부류 전체로 확장). 이들을 뺀 결과 순수 시그니처
# 'int foo()'는 못 벗기나(under-trace, 안전측), §C 이득은 이 저장소 실데이터 0건이라 손실 없고
# over-trace 위험만 제거된다. 자연어와 안 겹치는 것(void·헝가리안 u8~f64·stdint)만 타입으로
# 남긴다. struct/union/enum도 제외: 진짜 반환형이면 태그(struct Foo)를 동반해 lead에 비화이트
# 토큰이 생겨 어차피 차단(복구 이득 0)이고, 단독이면 자연어('union select')로 over-trace만
# 노출한다(deep-review 후속). '타입 토큰 최소 1개' 규칙이 한정자 단독('static delay( ms )')도 차단.
_C_TYPE_TOKEN = frozenset({           # 실제 반환 '타입'(자연어 비충돌만) — 벗기려면 최소 1개 필수.
    "void",
    "u8", "u16", "u32", "u64", "s8", "s16", "s32", "s64", "f32", "f64",
    "uint8_t", "uint16_t", "uint32_t", "uint64_t",
    "int8_t", "int16_t", "int32_t", "int64_t",
    "size_t", "ssize_t", "uintptr_t", "intptr_t",
})
_C_QUALIFIER_TOKEN = frozenset({      # 저장/타입 한정자 — 타입에 선행하나 단독으론 불충분.
    "static", "const", "volatile", "inline", "extern",
})
_C_RETURN_QUALIFIER = _C_TYPE_TOKEN | _C_QUALIFIER_TOKEN   # 벗기기 허용 선행 토큰(합집합)


def _sds_comp_key(comp: Any) -> str:
    """SDS component_id를 함수명 bridge용 정규화 키(lower)로 변환.

    추출 노이즈(시그니처 조각·배열 첨자·표 아티팩트)를 제거해 실제 C 식별자와
    정확매칭되게 한다. 정규화 후 **순수 C 식별자가 아니면 빈 문자열을 반환**해 키를
    버린다(reviewer W1): 공백/콜론/한글/선행숫자가 남은 설명문 컴포넌트('power operation
    disable', 'mcu 이상 감지', 'swcom_35: bootloader\\t115')는 함수명과 절대 매칭되면
    안 되므로 dict 오염·거짓 bridge 표면을 원천 차단한다. None은 초입에서, 숫자('123')는
    식별자 검증에서 걸러진다. 반환형이 공백분리된 프로토타입('void f( void )')은 선행이
    전부 C 반환형/한정자일 때만 함수명을 채택한다(§C, _C_RETURN_QUALIFIER).
    """
    if comp is None:                 # str(None)='None'→'none'이 유효 식별자라 아래 검증을
        return ""                    # 통과하는 것을 초입 차단('none' 키 오염 방지)
    s = str(comp).strip().lower()
    if not s:
        return ""
    s = _SDS_TABLE_ARTIFACT_RE.sub("", s)
    _had_paren = "(" in s            # 함수 시그니처 문맥 표지(§C 벗기기 전제)
    s = s.split("(", 1)[0]            # 'name( void' / 'name(void)' → 'name'
    s = _SDS_ARRAY_SUBSCRIPT_RE.sub("", s).strip()
    # §C: 공백으로 분리된 반환형/저장한정자 접두 제거 ('void f( x )' → 'f'). 3중 게이트로
    # over-trace를 차단한다(deep-review W1): ① 원본에 '('가 있던 시그니처 문맥(_had_paren)
    # ② 말단(함수명 후보) 제외 선행 토큰이 전부 타입/한정자이고 그중 실제 타입이 최소 1개
    # ③ 말단 토큰 자신은 타입 키워드가 아님('unsigned int( x )'→'int' 방지). 이로써 자연어
    # 라벨('auto close'·'long delay( ms )'·'static delay( ms )')이 함수명으로 새지 않는다.
    # 포인터 '*'는 토큰서 제거.
    if _had_paren and " " in s:
        _toks = [t for t in (p.strip("*") for p in s.split()) if t]
        _lead = _toks[:-1]
        if (len(_toks) >= 2
                and _toks[-1] not in _C_RETURN_QUALIFIER
                and all(t in _C_RETURN_QUALIFIER for t in _lead)
                and any(t in _C_TYPE_TOKEN for t in _lead)):
            s = _toks[-1]
    # 선행 언더스코어 정규화(라운드112): 링커/컴파일러 내부 표기('_entrypoint')와 설계 표기
    # ('entrypoint')의 차이로 bridge가 끊기는 것을 막는다. 실데이터 검증: SDS 키 중 선행 '_'는
    # 0개라 strip은 기존 매칭을 절대 깨지 않고(순수 가산) '_entrypoint'->'entrypoint' 1건만 새로
    # 연결한다. 추적성 목적상 선행 '_'는 의미 구분자가 아니므로 양변에서 동일 제거.
    s = s.lstrip("_")
    return s if _C_IDENT_RE.match(s) else ""


# 반환형 헝가리안 접두사 보정 (라운드 111 fix).
# SDS는 함수명에 반환형을 붙여(u16s_MotorSpdCtrl_AutoOpen) 표기하나, 테스트/VectorCAST는
# 생략한 형(s_MotorSpdCtrl_AutoOpen)을 쓴다. 이 불일치로 정확매칭 bridge가 끊겨 실데이터에서
# 도어 모터제어 4함수(s_motorspdctrl_{auto,assist}{open,close})가 각 11개 SRS 요구사항 추적을
# 잃고 있었다. 저장클래스 접두사(s/g/l) '직전'의 반환형 토큰만 제거한다 — lookahead로 s/g/l_가
# 뒤따를 때만 매칭하므로 'u8_foo'(저장클래스 없음)는 건드리지 않아 '_foo' 같은 오염을 막는다.
_RET_TYPE_PREFIX_RE = re.compile(r"^(?:u8|u16|u32|s8|s16|s32)(?=[sgl]_)")


def _strip_ret_type_prefix(key: str) -> str:
    """SDS 함수키의 선행 반환형 토큰 제거(없으면 원본). 'u16s_x'->'s_x', 'u8g_y'->'g_y'."""
    return _RET_TYPE_PREFIX_RE.sub("", key)


# ── ISO 26262 미추적 함수 계층(layer) 분류 (라운드112) ──
# SRS 미추적 VectorCAST/SUTS 함수를 'SwDS가 그 단위를 어느 계층에서 명세·추적해야 하는가'로
# 분류해, '애플리케이션 설계 공백(=실제 추적성 finding)'과 '정당한 범위 경계(BSW/부트/라이브러리)'를
# 정직히 구분 노출한다. 이 분류는 보고 hint일 뿐 in_uds/sds_reqs/safety 같은 1차 추적 판정을
# 바꾸지 않는다(순수 additive). 안전성은 직교 신호이므로 layer가 아니라 _SAFETY_TOKEN_RE 기반
# safety 플래그로 따로 표시한다(가드연산 같은 안전기제도 도메인은 APP일 수 있으므로).
# 분류 원칙(보수): 인프라 토큰은 함수명 core(선행 _·반환형·저장클래스 제거) 선두 앵커 위주로 잡아
# 애플리케이션 함수(중간에 eeprom 등 포함)를 인프라로 잘못 삼켜 공백을 숨기는 것을 막는다.
# 불확실하면 APP_LEAF(=검토 대상)로 떨어뜨려 안전측(공백 미은닉)으로 편향한다.
# 단, 구별성이 높은(앱과 충돌 없는) 인프라 토큰은 비앵커로도 추가한다 — 선두 앵커만으론
# 중간 위치 토큰(diageeprom·write2eeprom·pp1_..._pwm 등)을 놓쳐 인프라를 APP로 과대집계했다.
# 추가 토큰은 KJPDS02_PV 실데이터로 검증됨(이동 103건 전부 APP→인프라 단방향, 인프라→APP 역행 0,
# app_design_gap delta 0 — 이동분 전부 in_uds=True라 진짜갭 은닉 없음. 구조적으로 NEW_APP⊆OLD_APP
# (제거 브랜치 0)라 app_design_gap=|APP∩¬uds|은 단조 비증가, deep-review 3중 확증).
# s_buzzerstateflashing·s_antipinch·doorstatectrl 등 앱 도메인은 APP 유지. ⚠단, 비앵커 eeprom은
# 'Ap_*Eeprom*' 앱 함수 2건(ResetEepromParamState 등)도 BOOT로 재라벨한다(EEPROM 조작이라 방어가능,
# in_uds=True 무영향). bare eeprom/reprog의 잔여 latent 위험 = 미래 in_uds=False 앱 함수. _isr은
# 저장소 _ISR_RE 규율대로 '_isr$|_isr_' 앵커(앱 술어 _IsReady/_IsReset 오매치 배제). bare flash/pin/lin 금지.
_LAYER_CORE_PREFIX_RE = re.compile(r"^(?:u8|u16|u32|s8|s16|s32)?[sgl]_")
# 앱 도메인 네이밍 앵커 — 인프라 키워드를 **이름 중간에** 품은 앱 함수를 지켜낸다.
# 실측 근거(KJPDS02_PV build_125, 함수 정의 955개): `Ap_` 접두는 APP 계층 전용 규약이다 —
# APP 디렉터리 밖에 `Ap_*` **정의는 0건**이고 SYSTEM/IF에 보이는 것은 전부 호출이었다.
# 이 앵커가 없으면 `s_Ap_PreviousCtrl_ResetEepromParams`(Sources/APP/Ap_DoorPreCtrl_PDS.c)가
# 중간 'eeprom' 때문에 BOOT_REPROG로 삼켜진다. 앵커 도입 시 이동은 그 1건뿐이고
# BSW 216·LIB 44는 불변이라, 인프라 복구(4e449dc)를 되돌리지 않는다.
_LAYER_APP_ANCHOR_RE = re.compile(r"^ap_")
_LAYER_BOOT_RE = re.compile(
    # entrypoint는 선두 앵커(^) — 'validate_entrypoint'/'check_entrypoint_valid' 같은 APP
    # 함수를 BOOT로 오삼켜 공백을 숨기지 않도록(라운드112 W2). 실 부트 엔트리는 core가 정확히
    # 'entrypoint'라 ^로 충분(실데이터 '_entrypoint'는 lstrip 후 core='entrypoint').
    r"(^sf_|secureflash|bootload|^boot|^reprog|reprog|^clearreprog|^chkprog|^checkprog|^setreprog"
    r"|^getreprog|^eep|eeprom|^syseepromctrl|^entrypoint|jump_main|chkappisvalid|copy_shadow"
    r"|backupsector|fccob|^linuds|linudsreprog|^writeblock$|^writeword$"
    # 비앵커 eeprom/reprog + 플래시 프로그래밍(pflash/ftmrz/erase*sector) — diageeprom·
    # write2eeprom·maininitreprog·플래시섹터소거를 BOOT로 복구(실데이터 검증, +32).
    r"|pflash|ftmrz|eraseflashsector|erasesectorinternal)",
    re.I,
)
_LAYER_BSW_RE = re.compile(
    r"(^adc|^pwm|^spi|^gpio|^port|^timer|^tim[0-9]|systick|^rti|^cpu|_isr$|_isr_|^pt[0-9]|^dma|^clock|^osc|^pll"
    r"|^lin|^can|^uart|^sci|^pp[0-9]|^pj[0-9]|^pad[0-9]|^pl[0-9]|^pe_"
    # ^tim[0-9]/systick/^rti(타이머) · ^cpu(예외·인터럽트·PLL/OSC 핸들러) · _isr(ISR) ·
    # ^pp/pj/pad/pl[0-9](포트핀 I/O) · ^pe_(Processor Expert init) · driveic(모터 드라이브 IC)
    # 추가 — 저수준 드라이버/핸들러를 BSW로 복구(실데이터 검증, +69). 앱 도메인 오삼킴 0.
    r"|drv8706|iim20670|driveic|^drvin|^drvout|spictrl|lintp|lin_lld|sbcm|^sbc|^hw_|^mcu_|^reg_)",
    re.I,
)
_LAYER_LIB_RE = re.compile(
    # stdutil(copydata/dataset = memcpy/memset류 표준 유틸) 추가(실데이터 검증, +2).
    r"(sha256|^aes|crc32|^lib_|_lib_|movingaverage|_conv$|_conv_|slope_conv|^math|^util|stdutil)",
    re.I,
)


def _classify_unmapped_layer(names: List[str]) -> str:
    """미추적 함수명 후보(정규화 lower 배열)를 ISO 26262 SwDS 계층으로 분류.

    반환: 'TEST_ARTIFACT' | 'BOOT_REPROG' | 'BSW_DRIVER' | 'LIB_UTIL' | 'APP_LEAF'.
    우선순위 TEST>BOOT>BSW>LIB>APP(기본값). names가 비면 APP_LEAF.
    """
    # 방어적 소문자화(라운드112 I2): docstring 계약(lower 배열)을 보장하고, _C_IDENT_RE가
    # [a-z_] 기반이라 대소문자 혼재 입력이 거짓 TEST_ARTIFACT로 분류되는 latent edge를 막는다.
    cand = [n.lower() for n in (names or []) if n]
    if not cand:
        return "APP_LEAF"
    first = cand[0]
    # 순수 C 식별자가 아니거나 VectorCAST range-test 산출물은 추적 대상 함수가 아님.
    if not _C_IDENT_RE.match(first.lstrip("_")) or re.match(r"^(range$|<<)", first, re.I):
        return "TEST_ARTIFACT"
    core_list = [_LAYER_CORE_PREFIX_RE.sub("", n.lstrip("_")) for n in cand]
    cores = " ".join(core_list)
    # 앱 도메인 앵커가 인프라 키워드보다 우선한다(위 _LAYER_APP_ANCHOR_RE 주석의 실측 근거).
    # 후보 **전부**가 앱 앵커일 때만 — 하나라도 인프라 이름이 섞이면 인프라 판정 기회를 남긴다
    # (같은 요구에 앱·드라이버 함수가 함께 걸린 경우 인프라 쪽이 더 강한 신호다).
    if all(_LAYER_APP_ANCHOR_RE.match(c) for c in core_list):
        return "APP_LEAF"
    if _LAYER_BOOT_RE.search(cores):
        return "BOOT_REPROG"
    if _LAYER_BSW_RE.search(cores):
        return "BSW_DRIVER"
    if _LAYER_LIB_RE.search(cores):
        return "LIB_UTIL"
    return "APP_LEAF"


_REQ_ID_PAT = re.compile(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b", re.I)

def _extract_requirements_from_comments(text: str) -> List[str]:
    results: List[str] = []
    for ln in _extract_comment_lines(text):
        m = re.search(r"(REQ|Requirement|요구사항)\s*[:\-]\s*(.+)", ln, flags=re.I)
        if m:
            results.append(m.group(2).strip())
    return results


def _extract_table_section(lines: List[str], header: str, stop_headers: List[str], max_rows: int) -> List[str]:
    header_idx = None
    for i, ln in enumerate(lines):
        if header.lower() in ln.lower():
            header_idx = i
            break
    if header_idx is None:
        return []
    rows: List[str] = []
    for ln in lines[header_idx + 1 :]:
        if not ln.strip():
            if rows:
                break
            continue
        if any(h.lower() in ln.lower() for h in stop_headers):
            break
        rows.append(ln.strip())
        if len(rows) >= max_rows:
            break
    return rows


def _normalize_table_row(row: str) -> List[str]:
    if not row:
        return []
    parts = re.split(r"\s{2,}|\t+", row.strip())
    return [p.strip() for p in parts if p.strip()]


def _extract_function_blocks(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    lines = [ln.rstrip() for ln in text.splitlines()]
    blocks: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    state = ""
    current_swcom = ""
    for ln in lines:
        line = ln.strip()
        if re.match(r"^SwCom_\d+\b", line):
            current_swcom = line
            continue
        m_header = re.search(r"\b(SwUFn_\d+)\s*:\s*(.+)$", line)
        if m_header:
            if current:
                blocks.append(current)
                current = {}
            current["header"] = line
            current["id"] = m_header.group(1).strip()
            current["name"] = m_header.group(2).strip()
            if current_swcom:
                current["swcom"] = current_swcom
            state = ""
            continue
        if re.search(r"\bSwUFn_\d+", line):
            if current:
                blocks.append(current)
                current = {}
            current["header"] = line
            if current_swcom:
                current["swcom"] = current_swcom
            state = ""
            continue
        if not current:
            continue
        if line.startswith("["):
            state = line.lower()
            continue
        if not line:
            continue
        if line.startswith("ID"):
            current["id"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Name"):
            current["name"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Prototype"):
            current["prototype"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Description"):
            current["description"] = line.split(None, 1)[-1].strip()
        elif line.startswith("ASIL"):
            current["asil"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Related ID"):
            current["related"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Called Function"):
            current["called"] = line.split(None, 1)[-1].strip()
        elif line.startswith("Calling Function"):
            current["calling"] = line.split(None, 1)[-1].strip()
        elif line.startswith("사용 전역변수"):
            current["globals"] = line.split(None, 1)[-1].strip()
        elif line.startswith("선행조건"):
            current["precondition"] = line.split(None, 1)[-1].strip()
        elif "input param" in state:
            current.setdefault("inputs", []).append(line)
        elif "output param" in state:
            current.setdefault("outputs", []).append(line)
        elif "logic diagram" in state:
            current["logic"] = "present"
    if current:
        blocks.append(current)
    return blocks


def _docx_to_text(doc) -> str:
    lines: List[str] = []
    try:
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = (p.text or "").strip()
                        if text:
                            lines.append(text)
    except Exception:
        pass
    return "\n".join(lines)


def _extract_function_info_from_docx(doc) -> Dict[str, Dict[str, Any]]:
    result: Dict[str, Dict[str, Any]] = {}

    def _norm_label(raw: str) -> str:
        s = re.sub(r"\s+", " ", raw).strip().lower()
        s = re.sub(r"[\[\]()（）]", "", s).strip()
        return s

    def _is_param_header_row(cells: List[str]) -> bool:
        norms = [c.strip().lower() for c in cells]
        return "no" in norms and ("name" in norms or "type" in norms)

    def _parse_param_row(cells: List[str]) -> str:
        if len(cells) < 3:
            return ""
        no_val = cells[0].strip()
        if not no_val or not no_val[0].isdigit():
            return ""
        name = cells[1].strip() if len(cells) > 1 else ""
        ptype = cells[2].strip() if len(cells) > 2 else ""
        vrange = cells[3].strip() if len(cells) > 3 else ""
        reset = cells[4].strip() if len(cells) > 4 else ""
        desc = cells[5].strip() if len(cells) > 5 else ""
        if not name or name.upper() in {"N/A", "-", "NONE"}:
            return ""
        parts = [name]
        if ptype and ptype.upper() not in {"N/A", "-"}:
            parts[0] = f"{name} : {ptype}"
        if vrange and vrange.upper() not in {"N/A", "-"}:
            parts.append(f"range: {vrange}")
        if reset and reset.upper() not in {"N/A", "-"}:
            parts.append(f"reset: {reset}")
        entry = parts[0]
        if len(parts) > 1:
            entry += " (" + ", ".join(parts[1:]) + ")"
        return entry

    try:
        for table in doc.tables:
            if not table.rows:
                continue
            header = [c.text.strip() for c in table.rows[0].cells]
            if not header:
                continue
            header_joined = " ".join(header)
            if "Function Information" not in header_joined and "[ Function Information ]" not in header_joined:
                continue
            fn_id = ""
            if len(table.rows) > 1:
                for cell in table.rows[1].cells:
                    m = re.search(r"(SwUFn_\d+)", cell.text or "")
                    if m:
                        fn_id = m.group(1)
                        break
            if not fn_id:
                continue
            info: Dict[str, Any] = {"id": fn_id}
            last_label_norm = ""
            collecting_params = ""
            skip_next_header = False
            for row in table.rows[2:]:
                cells = [c.text.strip() for c in row.cells]
                if not cells:
                    continue
                label = cells[0].strip()
                label_norm = _norm_label(label) if label else ""

                # Extract value from non-label cells BEFORE any label checks
                # so that value is available when processing input/output parameter rows.
                values: List[str] = []
                value_seen: Set[str] = set()
                for cell in cells[1:]:
                    cval = str(cell or "").strip()
                    if not cval or cval == label:
                        continue
                    if cval in value_seen:
                        continue
                    value_seen.add(cval)
                    values.append(cval)
                value = "\n".join([v for v in values if v]).strip()
                if not value and values:
                    value = values[-1].strip()

                if skip_next_header and _is_param_header_row(cells):
                    skip_next_header = False
                    continue

                if collecting_params and label and label[0].isdigit():
                    entry = _parse_param_row(cells)
                    if entry:
                        key = "inputs" if collecting_params == "input" else "outputs"
                        direction = "[IN]" if collecting_params == "input" else "[OUT]"
                        info.setdefault(key, []).append(f"{direction} {entry}")
                    continue

                if label_norm in {"input parameters", "[ input parameters ]"}:
                    if value and value.upper() not in {"N/A", "TBD", "-", "NONE"}:
                        params = [ln.strip() for ln in value.splitlines() if ln.strip() and ln.strip().upper() not in {"N/A", "-"}]
                        if params:
                            for p in params:
                                info.setdefault("inputs", []).append(f"[IN] {p}" if not p.startswith("[") else p)
                    collecting_params = "input"
                    last_label_norm = "input parameters"
                    skip_next_header = True
                    continue
                if label_norm in {"output parameters", "[ output parameters ]"}:
                    if value and value.upper() not in {"N/A", "TBD", "-", "NONE"}:
                        params = [ln.strip() for ln in value.splitlines() if ln.strip() and ln.strip().upper() not in {"N/A", "-"}]
                        if params:
                            for p in params:
                                info.setdefault("outputs", []).append(f"[OUT] {p}" if not p.startswith("[") else p)
                    collecting_params = "output"
                    last_label_norm = "output parameters"
                    skip_next_header = True
                    continue

                if collecting_params and label_norm and not label[0].isdigit():
                    collecting_params = ""

                if not label_norm and last_label_norm and value:
                    if last_label_norm in {"description", "called function", "calling function"}:
                        prev = str(info.get({
                            "description": "description",
                            "called function": "called",
                            "calling function": "calling",
                        }.get(last_label_norm, ""), "") or "").strip()
                        joined = "\n".join([x for x in [prev, value] if x]).strip()
                        joined = _dedupe_multiline_text(joined)
                        if last_label_norm == "description":
                            info["description"] = joined
                            if joined.strip():
                                info["description_source"] = "reference"
                        elif last_label_norm == "called function":
                            info["called"] = joined
                        elif last_label_norm == "calling function":
                            info["calling"] = joined
                        continue
                    if last_label_norm in {
                        "used globals global", "used globals (global)", "used global variable global", "used global variableglobal", "used global variables global", "used global variablesglobal",
                        "used globals static", "used globals (static)", "used global variable static", "used global variablestatic", "used global variables static", "used global variablesstatic",
                        "사용 전역변수", "사용 전역 변수",
                    }:
                        key = "globals"
                        if "static" in last_label_norm:
                            key = "globals_static"
                        elif "global" in last_label_norm:
                            key = "globals_global"
                        prev_list = list(info.get(key) or [])
                        prev_list.extend([ln.strip() for ln in value.splitlines() if ln.strip()])
                        info[key] = prev_list
                        continue
                if not label_norm:
                    continue
                last_label_norm = label_norm
                if label_norm == "name":
                    info["name"] = _normalize_symbol_name(value)
                elif label_norm == "prototype":
                    info["prototype"] = value
                elif label_norm == "description":
                    info["description"] = _dedupe_multiline_text(value)
                    if value.strip():
                        info["description_source"] = "reference"
                elif label_norm == "asil":
                    info["asil"] = _normalize_asil_value(value)
                elif label_norm == "related id":
                    info["related"] = _normalize_related_ids(value)
                elif label_norm in {"precondition", "선행조건"}:
                    info["precondition"] = _dedupe_multiline_text(value, na_to_empty=True) or "N/A"
                elif label_norm == "called function":
                    info["called"] = value
                elif label_norm == "calling function":
                    info["calling"] = value
                elif label_norm in {"used globals global", "used globals (global)", "used global variable global", "used global variableglobal", "used global variables global", "used global variablesglobal"}:
                    info["globals_global"] = [ln.strip() for ln in value.splitlines() if ln.strip()]
                elif label_norm in {"used globals static", "used globals (static)", "used global variable static", "used global variablestatic", "used global variables static", "used global variablesstatic"}:
                    info["globals_static"] = [ln.strip() for ln in value.splitlines() if ln.strip()]
                elif label_norm in {"사용 전역변수", "사용 전역 변수"}:
                    all_vars = [ln.strip() for ln in value.splitlines() if ln.strip()]
                    info["globals"] = all_vars
                    from config import GLOBAL_VAR_PREFIXES, STATIC_VAR_PREFIXES
                    for var in all_vars:
                        v_stripped = var.split(",")[0].strip().split(":")[0].strip()
                        if any(v_stripped.startswith(p) for p in STATIC_VAR_PREFIXES):
                            info.setdefault("globals_static", []).append(var)
                        elif any(v_stripped.startswith(p) for p in GLOBAL_VAR_PREFIXES):
                            info.setdefault("globals_global", []).append(var)
                        else:
                            info.setdefault("globals_global", []).append(var)
                elif label_norm == "logic diagram":
                    info["logic"] = value
            result[fn_id] = info
    except Exception:
        return result
    return result


# 손상 임베드 파트(깨진 이미지 등)로 python-docx의 Package.open이 BadZipFile을 던지는 docx를
# 위해, 읽히는 멤버만 재압축해 여는 resilient 로더. python-docx는 문서를 열 때 모든 파트를
# eager 로드하므로 임베드 이미지 1개만 깨져도 본문(document.xml)이 멀쩡한데 전체가 안 열린다.
# (라운드110: KJPDS02 SDS v2.03이 깨진 image4.png 등 32개로 sds_pairs=0 되던 회귀.)
# 정상 문서는 첫 시도에서 바로 열려 fast-path, 예외 시에만 재압축 우회 → 투명·저위험.
_DOCX_PNG_1x1 = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
    "0000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
)
_DOCX_IMG_EXT = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".emf", ".wmf", ".tif", ".tiff")


def _safe_docx_open(source: Any) -> Any:
    """python-docx로 docx를 연다. 손상 파트로 실패하면 읽히는 멤버만 재압축 후 재시도.

    손상 이미지는 유효한 1x1 PNG로 대체해 relationship을 유지한다(표/문단 추출에는
    이미지가 불필요). document.xml이 정상이면 매핑 추출이 그대로 동작한다.
    """
    import docx  # type: ignore

    try:
        return docx.Document(source)
    except Exception:
        pass
    import io
    import zipfile

    if isinstance(source, (bytes, bytearray)):
        raw = bytes(source)
    elif hasattr(source, "read"):          # file-like(BytesIO 등) — 첫 시도가 소비했을 수 있어 되감기
        try:
            source.seek(0)
        except Exception:
            pass
        raw = source.read()
    else:
        raw = Path(source).read_bytes()
    src = zipfile.ZipFile(io.BytesIO(raw))
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for info in src.infolist():
            try:
                blob = src.read(info.filename)
            except Exception:
                # 손상 멤버: 이미지면 유효한 1x1 PNG로 대체(rels 유지), 그 외 빈 바이트
                blob = _DOCX_PNG_1x1 if info.filename.lower().endswith(_DOCX_IMG_EXT) else b""
            zout.writestr(info, blob)
    buf.seek(0)
    return docx.Document(buf)


# ── SDS 파티션 엔트리의 kind 체계 ──────────────────────────────────────────────
# 예전엔 'component'/'function' 2종뿐이었고 `_add_entry` 의 기본값이 'component' 라,
# 11개 호출지점 중 인터페이스 함수 행(1곳)만 명시하고 나머지 10곳이 전부 컴포넌트로
# 등록됐다. 그 결과 SDS 밴드에 상태명('initial'/'standby'/'auto close')과 설계ID
# ('SwFn_'/'SwSTR_'/'SwST_')가 섞여 실 컴포넌트 33개가 201개로 부풀었다(HDPDM01 실측).
#
# 랭크는 "같은 키가 여러 표에서 다른 성격으로 등장할 때 무엇이 이기는가"를 정한다.
# component 가 모두를 이기는 것은 종전과 동일하고, function 이 약한 kind 를 이기는 것이
# 신규다(예전엔 약한 kind 가 먼저 걸리면 함수 승격이 막혔다).
_SDS_KIND_RANK = {
    "component": 4,       # SwCom ID·이름 표 — 밴드에 세는 유일한 kind
    "function": 3,        # SwCom 인터페이스 함수 — 브리지·영향분석 ASIL 전용
    "design_element": 1,  # 설계 블록의 name(상태명: initial/standby/auto close)
    "design_id": 1,       # SwFn_/SwSTR_/SwST_ — 설계 ID 이지 컴포넌트 아님
    "table_row": 1,       # 범용 표 폴백 — 헤더만 맞은 미확인 행
    "heading": 1,         # 문단 heading 잔재
}
# ⚠ **긍정형** 화이트리스트여야 한다. 예전의 부정형(`kind != 'function'`)이면 새 kind 가
# 전부 그대로 밴드를 통과해 이 분류가 무효가 된다.
_SDS_BAND_KINDS = frozenset({"component"})
# kind 없는 레거시 엔트리(구 캐시·외부 주입)는 분류 이전 동작을 그대로 보존한다.
_SDS_DEFAULT_KIND = "component"

# SwCom ID 추출 — 표기 혼재(`SwCom_7` / `Sw Com 07` / `SwCom-7`)를 제로패딩 canonical 로.
# `report_gen/utils.py::_normalize_swcom_label` 과 같은 관용 규칙이되, 그쪽은 주변 텍스트를
# 보존하는 **라벨** 정규화라 목적이 다르다(`Door Control(SwCom_14)` → 같은 문자열).
_SWCOM_ID_RE = re.compile(r"\bSw\s*Com\s*[_\s-]?\s*(\d{1,3})\b", re.I)


def _canonical_swcom_id(text: str) -> str:
    """`SwCom_7` / `Sw Com 07` → `SwCom_07`. SwCom ID 가 없으면 빈 문자열."""
    m = _SWCOM_ID_RE.search(str(text or ""))
    return f"SwCom_{int(m.group(1)):02d}" if m else ""


def _extract_sds_partition_map(doc_path: str) -> Dict[str, Dict[str, str]]:
    try:
        pass  # type: ignore
    except Exception:
        return {}
    if not doc_path:
        return {}
    path = Path(doc_path)
    if not path.exists():
        return {}
    try:
        doc = _safe_docx_open(str(path))
    except Exception:
        return {}
    mapping: Dict[str, Dict[str, str]] = {}
    # component_description 폴백 출처 추적 — 같은 함수가 여러 SwCom 표에 서로 다른 설명으로
    # 등장하면(이 프로젝트가 실측한 SwCom echo/교차참조 과다추출 패턴) 표 순서로 임의 상속하는
    # 오귀속을 막는다: 두 번째 '다른' 설명을 만나면 모호로 표시하고 component_description을 제거
    # (honest miss — 잘못된 컴포넌트 맥락보다 없는 편이 안전, deep-review Warning).
    _comp_desc_src: Dict[str, str] = {}
    _comp_desc_ambiguous: set = set()

    def _collapse_adjacent_duplicates(values: List[str]) -> List[str]:
        result: List[str] = []
        for value in [str(v or "").strip() for v in values]:
            if not value:
                result.append("")
                continue
            if result and result[-1] == value:
                continue
            result.append(value)
        return result

    def _add_entry(name: str, asil: str, related: str, desc: str, kind: str = "component",
                   canonical: str = "") -> None:
        key = str(name or "").strip().lower()
        if not key:
            return
        entry = mapping.get(key, {})
        # canonical: 같은 SwCom 을 가리키는 ID 키(`swcom_14`)와 이름 키(`door control`)를
        # 한 컴포넌트로 접기 위한 표시/집계용 앵커. 조회 키(원 키)는 그대로 살려 둔다 —
        # Related ID 가 컴포넌트를 이름으로만 참조한 요구의 링크가 끊기면 안 된다.
        if canonical and not entry.get("canonical"):
            entry["canonical"] = canonical
        if asil and not entry.get("asil"):
            entry["asil"] = asil
        if related and not entry.get("related"):
            entry["related"] = related
        if desc and not entry.get("description"):
            entry["description"] = desc
        # kind 승격은 `_SDS_KIND_RANK` 비교로 한다 — 같은 키가 여러 표에서 다른 성격으로
        # 등장할 때 강한 쪽이 이긴다. SDS 밴드/링크 집계는 'component' 만 세어 함수
        # fan-out 과대표기(요구사항당 16→413)를 막고, 함수는 SUTS/VCAST 브리지에 그대로 쓰인다.
        # ⚠ 예전 코드는 인자 `kind` 를 무시하고 "function" 을 하드코딩했다 — kind 가 2종일
        #    때만 우연히 맞았고, 약한 kind 가 먼저 걸리면 함수 승격이 막혔다.
        _new = kind if kind in _SDS_KIND_RANK else _SDS_DEFAULT_KIND
        _cur = entry.get("kind")
        if _cur is None or _SDS_KIND_RANK[_new] > _SDS_KIND_RANK.get(_cur, 0):
            entry["kind"] = _new
        mapping[key] = entry

    def _find_col(norm_headers: List[str], keywords: List[str]) -> int:
        for kw in keywords:
            for i, h in enumerate(norm_headers):
                if kw == h or kw in h:
                    return i
        return -1

    swcom_asil_map: Dict[str, str] = {}

    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        header_norm = [h.lower() for h in header]
        header_joined = " ".join(header_norm)

        if "software component information" in header_joined:
            sc_data: Dict[str, str] = {}
            func_rows: List[Dict[str, str]] = []
            in_interface = False
            iface_header: List[str] = []
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                compact_cells = _collapse_adjacent_duplicates(cells)
                first = cells[0].lower() if cells else ""
                last_val = ""
                for c in reversed(cells):
                    cv = c.strip()
                    if cv and cv != cells[0].strip():
                        last_val = cv
                        break
                if not last_val:
                    last_val = cells[-1].strip() if cells else ""

                if first in {"sc id", "sc_id"}:
                    m = re.search(r"SwCom_\d+", last_val, re.I)
                    if m:
                        sc_data["id"] = m.group(0)
                elif first in {"sc name", "sc_name"}:
                    sc_data["name"] = last_val
                elif first in {"sc description", "sc_description"}:
                    sc_data["description"] = last_val[:500]
                elif first == "asil":
                    sc_data["asil"] = last_val
                elif first in {"related id", "related_id"}:
                    sc_data["related"] = last_val
                elif "sw component interface" in first or "component interface" in first:
                    in_interface = True
                    continue
                elif first == "no" and ("name" in " ".join(cells).lower()):
                    in_interface = True
                    iface_header = [c.lower() for c in compact_cells]
                    continue
                elif "software component design" in first or "component design" in first:
                    in_interface = False
                    continue

                if in_interface and first and first[0].isdigit():
                    fname = ""
                    fdesc = ""
                    name_idx = -1
                    desc_idx = -1
                    if iface_header:
                        for i, header_name in enumerate(iface_header):
                            h = str(header_name or "").strip().lower()
                            if name_idx < 0 and h == "name":
                                name_idx = i
                            if desc_idx < 0 and "description" in h:
                                desc_idx = i
                    if name_idx >= 0 and name_idx < len(compact_cells):
                        fname = compact_cells[name_idx].strip()
                    if desc_idx >= 0 and desc_idx < len(compact_cells):
                        fdesc = compact_cells[desc_idx].strip()
                    if not fname:
                        for cv in compact_cells[1:]:
                            token = cv.strip()
                            if token and not token[0].isdigit():
                                fname = token
                                break
                    if not fdesc:
                        for cv in reversed(compact_cells[1:]):
                            token = cv.strip()
                            if token and token != fname and not re.fullmatch(r"(?:static\s+)?(?:void|u8|u16|u32|u64|s8|s16|s32|s64|enum)(?:\s*\(\s*void\s*\))?", token, re.I):
                                fdesc = token
                                break
                    if fname:
                        func_rows.append({"name": fname, "desc": fdesc})

            sc_id = sc_data.get("id", "")
            sc_name = sc_data.get("name", "")
            sc_asil = sc_data.get("asil", "") or swcom_asil_map.get(sc_id.lower(), "")
            sc_related = sc_data.get("related", "")
            sc_desc = sc_data.get("description", "")
            # ID 키와 이름 키는 **같은 표의 같은 행**에서 나온 같은 컴포넌트다 → 같은 canonical.
            sc_canon = _canonical_swcom_id(sc_id)
            if sc_id:
                _add_entry(sc_id, sc_asil, sc_related, sc_desc, kind="component", canonical=sc_canon)
            if sc_name:
                _add_entry(sc_name, sc_asil, sc_related, sc_desc, kind="component", canonical=sc_canon)
            for fr in func_rows:
                fn = fr["name"].rstrip("()").strip()
                _add_entry(fn, sc_asil, sc_related, fr.get("desc", ""), kind="function")
                # 인터페이스 행 description이 비어도 소속 SwCom 설명을 폴백 필드로 부착 —
                # 영향분석 SDS 카드가 함수 소속 컴포넌트 맥락을 표시할 수 있게(description 불변,
                # 신규 필드라 밴드 집계·SUTS/VCAST 브리지 등 기존 소비처엔 무영향).
                fk = fn.strip().lower()
                if sc_desc and fk and fk in mapping:
                    _capped = sc_desc[:500]
                    if fk in _comp_desc_ambiguous:
                        pass  # 이미 다중 SwCom 충돌 확정 — 붙이지 않음
                    elif fk not in _comp_desc_src:
                        _comp_desc_src[fk] = _capped
                        mapping[fk]["component_description"] = _capped
                    elif _comp_desc_src[fk] != _capped:
                        # 다른 SwCom이 다른 설명 부여 → 모호. 오귀속 방지로 제거 후 확정.
                        _comp_desc_ambiguous.add(fk)
                        mapping[fk].pop("component_description", None)
            continue

        idx_comp_id = _find_col(header_norm, ["comp id", "component id", "swcom"])
        idx_comp_name = _find_col(header_norm, ["component name", "comp name"])
        idx_comp_asil = _find_col(header_norm, ["asil", "safety level", "safety"])
        if idx_comp_id >= 0 and idx_comp_name >= 0 and idx_comp_asil >= 0:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if idx_comp_id >= len(cells):
                    continue
                cid = cells[idx_comp_id]
                cname = cells[idx_comp_name] if idx_comp_name < len(cells) else ""
                casil = cells[idx_comp_asil] if idx_comp_asil < len(cells) else ""
                if cid and casil:
                    swcom_asil_map[cid.lower()] = casil
                    cid_canon = _canonical_swcom_id(cid)
                    _add_entry(cid, casil, "", "", kind="component", canonical=cid_canon)
                    if cname:
                        _add_entry(cname, casil, "", "", kind="component", canonical=cid_canon)
            continue

        idx_name = _find_col(header_norm, [
            "partition name", "component name", "module name", "name",
            "function", "function name", "sw component", "swcom",
        ])
        idx_asil = _find_col(header_norm, [
            "asil", "safety level", "safety", "safety class", "integrity level",
        ])
        idx_rel = _find_col(header_norm, [
            "related id", "related", "requirement", "req id", "trace",
            "traceability", "parent id",
        ])
        idx_desc = _find_col(header_norm, [
            "description", "desc", "function description", "purpose",
        ])
        if idx_name < 0:
            attr_idx = next((idx for idx, col in enumerate(header_norm) if col.startswith("attribute")), -1)
            cont_idx = next((idx for idx, col in enumerate(header_norm) if col == "contents"), -1)
            if attr_idx >= 0 and cont_idx >= 0:
                block: Dict[str, str] = {}
                for row in table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if attr_idx < len(cells) and cont_idx < len(cells):
                        block[cells[attr_idx].lower()] = cells[cont_idx]
                bid = block.get("id", "")
                if bid and re.match(r"Sw\w+_\d+", bid):
                    # Attribute/Contents 세로 블록 — `name` 은 상태명('initial'/'standby'/
                    # 'auto close'), `id` 는 설계ID(SwFn_/SwSTR_/SwST_). 둘 다 컴포넌트가
                    # 아니므로 밴드에서 제외한다(HDPDM01 실측 201개 중 131개가 여기서 왔다).
                    _add_entry(
                        block.get("name", bid),
                        block.get("asil", ""),
                        block.get("related id", block.get("related", "")),
                        block.get("description", "")[:500],
                        kind="design_element",
                    )
                    _add_entry(
                        bid,
                        block.get("asil", ""),
                        block.get("related id", block.get("related", "")),
                        block.get("description", "")[:500],
                        kind="design_id",
                    )
            continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if idx_name >= len(cells):
                continue
            # 범용 표 폴백 — 헤더 이름만 맞은 미확인 행이라 컴포넌트로 단정할 수 없다.
            _add_entry(
                cells[idx_name],
                cells[idx_asil] if idx_asil >= 0 and idx_asil < len(cells) else "",
                cells[idx_rel] if idx_rel >= 0 and idx_rel < len(cells) else "",
                cells[idx_desc] if idx_desc >= 0 and idx_desc < len(cells) else "",
                kind="table_row",
            )

    _asil_pat = re.compile(r"\bASIL[\s\-_]*([A-D](?:\s*\([A-D]\))?)\b|\bQM\b", re.I)
    _module_heading_pat = re.compile(
        r"^(?:\d+\.?\d*\.?\s+)?(?:Module|Component|Partition|Software\s+Unit|SwCom|SW\s*Component)\s*[:\-_]?\s*(.+)",
        re.I,
    )
    _swcom_pat = re.compile(r"\bSwCom[_\s-]*(\d+)\b", re.I)
    current_module = ""
    current_asil = ""
    desc_buffer: List[str] = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        heading_m = _module_heading_pat.match(txt)
        swcom_m = _swcom_pat.search(txt)
        is_heading = hasattr(para, "style") and para.style and "heading" in str(para.style.name or "").lower()
        if heading_m or is_heading or swcom_m:
            if current_module and desc_buffer:
                _add_entry(current_module, "", "", " ".join(desc_buffer).strip(), kind="heading")
                desc_buffer = []
            candidate = heading_m.group(1).strip() if heading_m else txt
            candidate = re.sub(r"^\d+\.?\d*\.?\s*", "", candidate).strip()
            if candidate:
                current_module = candidate
                current_asil = ""
            continue
        if current_module:
            asil_m = _asil_pat.search(txt)
            if asil_m:
                asil_val = "QM" if asil_m.group(0).strip().upper().startswith("QM") else asil_m.group(1)[0].upper()
                _add_entry(current_module, asil_val, "", "", kind="heading")
                current_asil = asil_val
            req_ids = re.findall(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b", txt)
            if req_ids:
                _add_entry(current_module, "", ", ".join(req_ids), "", kind="heading")
            if not asil_m and not req_ids and len(txt) > 10 and not txt.startswith(("Table", "Figure")):
                desc_buffer.append(txt)
    if current_module and desc_buffer:
        _add_entry(current_module, "", "", " ".join(desc_buffer).strip()[:500], kind="heading")

    return mapping


# SDS Related ID 안의 요구/설계 ID 토큰. 내부 공백 허용("SwRS_ 001") — 두 라우터가
# 각자 복제해 쓰던 정규식을 여기로 단일화했다.
_SDS_RELATED_ID_RE = re.compile(r"Sw[A-Za-z]{2,}\s*_\s*\d+|Sy[A-Za-z]{2,}\s*_\s*\d+")


def build_sds_component_maps(partition_map: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """SDS 파티션 맵 → 요구사항별 컴포넌트 맵. **SDS 밴드 판정의 단일 출처.**

    `backend/routers/jenkins.py`(Jenkins 모드)와 `backend/routers/local.py`(로컬 모드)가
    같은 판정을 복제하고 있던 것을 합쳤다. 복제 시절엔 한쪽만 고쳐져 모드 간 SDS 컴포넌트
    수가 갈릴 수 있었다 — `scripts/_ratchet_core.py` 가 ruff/eslint ratchet 에서 같은
    문제를 해결한 선례를 따른다.

    반환 키:
      ``req_to_comps``        요구ID → [컴포넌트+함수 **전체**]. 브리지 3종
                              (``sds_func_to_reqs``/``design_to_reqs``/``comp_to_reqs``)이
                              여기 의존하므로 **절대 줄이지 말 것**.
      ``req_to_design_comps`` 요구ID → [실 SwCom/모듈만]. SDS 밴드 집계·표시용.
                              같은 컴포넌트의 ID 키·이름 키는 canonical 로 접힌다.
      ``req_to_folded_comps`` 요구ID → [접기로 사라진 **원 키**]. 소비처가 차집합
                              (``sds_functions``)을 낼 때 여기 실린 키를 함께 빼야 한다.
      ``component_asil``      컴포넌트/함수명(lower) → ASIL.
      ``comp_set``            ``req_to_comps`` 에 실린 distinct 키.
      ``design_comp_set``     ``req_to_design_comps`` 에 실린 distinct 표시 라벨.

    ⚠ ``related`` 없는 엔트리는 두 맵 모두에서 탈락한다. 이 게이트 때문에 SDS 이름 집합은
    **불완전**하며(하한선), Jenkins 경로의 ``_all_sds_keys``·``unmapped_sds_name_hit`` 은
    그 하한선 위에서 계산된다. 로컬 경로는 ``generate_uds_traceability_matrix`` 를 호출하지
    않아 이 인과가 성립하지 않는다. SDS 이름 완전 집합이 필요해지면 여기를 고친다.

    ⚠ ``component_asil`` 만 게이트 **밖**에서 전 엔트리를 순회한다 — 요구의 component_id 가
    related 없는 SwCom 정의 행을 가리킬 수 있기 때문(P5 ASIL 결합).

    통합 전 두 라우터의 실측 차이는 ``if req_ids:`` 가드 유무 하나뿐이었고, 가드가 없는 쪽도
    빈 리스트면 루프가 무동작이라 **맵 결과는 동치**다. 여기서는 가드 있는 쪽(jenkins)을 채택해
    ``comp_set`` 의 의미("요구ID를 실제로 하나 이상 낳은 엔트리")를 보존한다.
    """
    req_to_comps: Dict[str, List[str]] = {}
    req_to_design_comps: Dict[str, List[str]] = {}
    req_to_folded_comps: Dict[str, List[str]] = {}
    comp_set: Set[str] = set()
    design_comp_set: Set[str] = set()
    component_asil: Dict[str, str] = {}

    for comp_key, info in (partition_map or {}).items():
        if not isinstance(info, dict):
            continue
        casil = str(info.get("asil") or "").strip()
        if casil:
            component_asil[str(comp_key).strip().lower()] = casil

    for comp_key, info in (partition_map or {}).items():
        if not isinstance(info, dict):
            continue
        related = info.get("related", "")
        if not related:
            continue
        req_ids = [_normalize_req_id(rid) for rid in _SDS_RELATED_ID_RE.findall(str(related))]
        if not req_ids:
            continue
        # ⚠ **긍정형** 화이트리스트. 예전의 부정형(`kind != "function"`)이면 상태명·설계ID·
        # 표 폴백 잔재가 전부 그대로 밴드를 통과한다.
        # kind 부재(구 캐시·외부 주입) **와** 미등록 kind 문자열은 둘 다 `_SDS_DEFAULT_KIND`
        # 로 흡수한다 — `_add_entry` 의 fail-safe 와 같은 규칙이어야 한다. 여기만 미지 kind 를
        # 탈락시키면, 나중에 `_add_entry` 에 kind 를 추가하고 `_SDS_KIND_RANK` 등록을 잊었을 때
        # 그 엔트리가 **경고 없이 밴드에서 사라진다**(밴드 제외는 등록이라는 명시 행위여야 한다).
        _kind = str(info.get("kind") or "").strip().lower()
        if _kind not in _SDS_KIND_RANK:
            _kind = _SDS_DEFAULT_KIND
        is_band = _kind in _SDS_BAND_KINDS
        # 표시 라벨 = canonical(`SwCom_14`) 우선. 없으면 원 키 그대로 — SwCom 표에 없는
        # 고아 이름이 조용히 사라지지 않게.
        canon = str(info.get("canonical") or "").strip()
        display = canon or comp_key
        comp_set.add(comp_key)
        if is_band:
            design_comp_set.add(display)
        for rid in req_ids:
            bucket = req_to_comps.setdefault(rid, [])
            if comp_key not in bucket:
                bucket.append(comp_key)
            if is_band:
                dbucket = req_to_design_comps.setdefault(rid, [])
                if display not in dbucket:
                    dbucket.append(display)
                # ⚠ 접기로 사라진 **원 키**를 반드시 남긴다. 이게 없으면 소비처의 차집합
                #   (sds_components 를 뺀 나머지 = sds_functions)이 `swcom_14` 와
                #   `door control` 을 **둘 다 함수로** 밀어 넣어 이중 계상된다.
                #   canonical 이 원 키와 다르면 원 키 자신도 folded 대상이다(대소문자 불일치).
                if canon and canon != comp_key:
                    fbucket = req_to_folded_comps.setdefault(rid, [])
                    if comp_key not in fbucket:
                        fbucket.append(comp_key)

    return {
        "req_to_comps": req_to_comps,
        "req_to_design_comps": req_to_design_comps,
        "req_to_folded_comps": req_to_folded_comps,
        "component_asil": component_asil,
        "comp_set": comp_set,
        "design_comp_set": design_comp_set,
    }


# 거친 입도 임계 — 실 컴포넌트의 이 비율을 **초과**해 연결되면 변별력 없음으로 본다.
_SDS_COARSE_RATIO = 0.4
# 총 컴포넌트가 이보다 적으면 미적용 — 3개 중 2개 같은 소규모에서 비율이 무의미해진다.
_SDS_COARSE_MIN_COMPONENTS = 5


def annotate_sds_coarse(rows: Optional[List[Dict[str, Any]]]) -> Tuple[int, int]:
    """행별 ``sds_coarse`` 플래그를 채우고 ``(실 컴포넌트 총수, coarse 행 수)`` 반환.

    한 요구사항이 실 설계 컴포넌트의 40%를 넘게 참조하면 "이 요구가 어느 설계에
    대응하는가"를 좁히지 못한다(예: SwCom Related-ID 에 SwTR 을 통째로 나열). 결함이
    아니라 **문서 입도 신호**다.

    분모는 반드시 **정화 후 실 컴포넌트**여야 한다. 상태명·설계ID·목차줄까지 세던 동안
    분모가 6배로 부풀어(HDPDM01 201 vs 실 33) 임계가 80.4 로 떠 있었고, 최대 행이
    34개라 이 지표는 **한 번도 발화하지 않았다**(0/63 → 정화 후 8/63).

    Jenkins 경로(`generate_uds_traceability_matrix`)와 local 경로(`local.py`)가 각자
    세면 같은 문서가 모드에 따라 다른 값을 내므로 단일 출처로 둔다 —
    `build_sds_component_maps` 와 같은 이유.
    """
    _rows = [r for r in (rows or []) if isinstance(r, dict)]
    total = len({
        str(c).strip().lower()
        for r in _rows for c in (r.get("sds_components") or []) if str(c).strip()
    })
    coarse_count = 0
    if total >= _SDS_COARSE_MIN_COMPONENTS:
        threshold = _SDS_COARSE_RATIO * total
        for r in _rows:
            is_coarse = len(r.get("sds_components") or []) > threshold
            r["sds_coarse"] = is_coarse
            if is_coarse:
                coarse_count += 1
    return total, coarse_count


def _load_component_map() -> Dict[str, Dict[str, str]]:
    # report_gen/docs/ → 프로젝트 루트 docs/ 순으로 탐색
    candidates = [
        Path(__file__).resolve().parent / "docs" / "component_map.json",
        Path(__file__).resolve().parent.parent / "docs" / "component_map.json",
    ]
    path = None
    for c in candidates:
        try:
            if c.exists():
                path = c
                break
        except Exception:
            continue
    if not path:
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    if not isinstance(data, list):
        return {}
    mapping: Dict[str, Dict[str, str]] = {}
    for row in data:
        if not isinstance(row, dict):
            continue
        file_name = str(row.get("file") or "").strip()
        component = str(row.get("component") or "").strip()
        verify = str(row.get("verify") or "").strip().upper()
        structure = str(row.get("structure") or "").strip()
        if not file_name or not component:
            continue
        component = _normalize_swcom_label(component)
        entry = {
            "component": component,
            "verify": verify,
            "structure": structure,
        }
        mapping[file_name] = entry
        mapping[Path(file_name).stem] = entry
        # Fuzzy: _it_PDS ↔ _PDS 변환 매칭 (소스 파일명 규칙 차이 대응)
        stem = Path(file_name).stem
        if "_it_" in stem:
            alt = stem.replace("_it_", "_")
            mapping[alt] = entry
            mapping[alt + Path(file_name).suffix] = entry
        elif "_PDS" in stem:
            alt = stem.replace("_PDS", "_it_PDS")
            mapping[alt] = entry
            mapping[alt + Path(file_name).suffix] = entry
        # 경로 기반 엔트리: "PDS64_FBL/Sources/main.c" 같은 상대경로 포함 키
        if "/" in file_name or "\\" in file_name:
            # 경로 구분자 통일
            norm_path = file_name.replace("\\", "/")
            mapping[norm_path] = entry
    return mapping


def _build_req_map_from_texts(texts: List[str]) -> Dict[str, Dict[str, str]]:
    mapping: Dict[str, Dict[str, str]] = {}
    for txt in texts:
        for block in _extract_requirement_blocks(txt):
            name = str(block.get("name") or "").strip()
            rid = str(block.get("id") or "").strip()
            asil = str(block.get("asil") or "").strip()
            related = str(block.get("related_ids") or block.get("related") or "").strip()
            if name:
                mapping[name.lower()] = {"asil": asil, "related": related or rid}
            if rid:
                mapping[rid.lower()] = {"asil": asil, "related": related or rid}
        current_id = ""
        for raw in txt.splitlines():
            line = raw.strip()
            if not line:
                continue
            m = re.search(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b", line)
            if m:
                current_id = m.group(1)
            asil_match = re.search(
                r"\bASIL\b(?:\s*[:|\-]\s*|\s+)?((?:ASIL-)?(?:A|B|C|D)|QM)\b",
                line,
                re.I,
            )
            if m and asil_match:
                mapping[m.group(1).lower()] = {
                    "asil": _normalize_asil_value(asil_match.group(1)),
                    "related": mapping.get(m.group(1).lower(), {}).get("related", m.group(1)),
                }
            if current_id and line.lower().startswith("related id"):
                related_val = line.split(None, 2)[-1].strip() if " " in line else ""
                mapping[current_id.lower()] = {
                    "asil": mapping.get(current_id.lower(), {}).get("asil", ""),
                    "related": related_val,
                }
            if current_id and asil_match and not m:
                mapping[current_id.lower()] = {
                    "asil": _normalize_asil_value(asil_match.group(1)),
                    "related": mapping.get(current_id.lower(), {}).get("related", current_id),
                }
    return mapping


def _build_req_map_from_doc_paths(doc_paths: List[str], texts: Optional[List[str]] = None) -> Dict[str, Dict[str, str]]:
    mapping: Dict[str, Dict[str, str]] = {}

    def _merge_entry(key: str, asil: str = "", related: str = "") -> None:
        norm_key = str(key or "").strip().lower()
        if not norm_key:
            return
        entry = mapping.get(norm_key, {})
        asil_norm = _normalize_asil_value(asil)
        related_norm = _normalize_related_ids(related)
        if asil_norm and not entry.get("asil"):
            entry["asil"] = asil_norm
        if related_norm and not entry.get("related"):
            entry["related"] = related_norm
        mapping[norm_key] = entry

    def _is_req_id(value: str) -> bool:
        return bool(re.match(r"^Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+$", str(value or "").strip(), re.I))

    def _table_rows(table: Any) -> List[List[str]]:
        rows: List[List[str]] = []
        for row in table.rows:
            cells: List[str] = []
            for cell in row.cells:
                parts = [p.text.strip() for p in cell.paragraphs if (p.text or "").strip()]
                cell_text = "\n".join(parts).strip() if parts else (cell.text or "").strip()
                cells.append(cell_text)
            rows.append(cells)
        return rows

    try:
        import docx  # type: ignore
    except Exception:
        docx = None  # type: ignore

    if docx:
        for raw_path in doc_paths or []:
            path = Path(str(raw_path or "").strip())
            if not path.exists() or path.suffix.lower() != ".docx":
                continue
            try:
                doc = _safe_docx_open(str(path))
            except Exception:
                continue

            for table in doc.tables:
                rows = _table_rows(table)
                if not rows:
                    continue
                header = [str(c or "").strip().lower() for c in rows[0]]

                attr_idx = next((idx for idx, col in enumerate(header) if col.startswith("attribute")), -1)
                cont_idx = next((idx for idx, col in enumerate(header) if col == "contents"), -1)
                if attr_idx >= 0 and cont_idx >= 0:
                    block: Dict[str, str] = {}
                    for row in rows[1:]:
                        if attr_idx >= len(row) or cont_idx >= len(row):
                            continue
                        label = str(row[attr_idx] or "").strip().lower()
                        value = str(row[cont_idx] or "").strip()
                        if label and value:
                            block[label] = value
                    rid = str(block.get("id") or "").strip()
                    if _is_req_id(rid):
                        asil = block.get("asil", "")
                        related = block.get("related id", block.get("related", rid))
                        _merge_entry(rid, asil, related)
                        name = str(block.get("name") or "").strip()
                        if name:
                            _merge_entry(name, asil, related)
                    continue

                header_joined = " ".join(header)
                if not header_joined:
                    continue

                id_idx = -1
                for idx, col in enumerate(header):
                    if col == "id" or col.endswith(" id"):
                        id_idx = idx
                        break
                asil_idx = next((idx for idx, col in enumerate(header) if col == "asil" or "asil" in col), -1)
                related_idx = next(
                    (
                        idx for idx, col in enumerate(header)
                        if col in {"related id", "related", "parent id", "traceability", "trace"}
                        or "related id" in col
                    ),
                    -1,
                )
                name_idx = next((idx for idx, col in enumerate(header) if col == "name" or col.endswith(" name")), -1)
                if id_idx < 0:
                    continue
                for row in rows[1:]:
                    if id_idx >= len(row):
                        continue
                    rid = str(row[id_idx] or "").strip()
                    if not _is_req_id(rid):
                        continue
                    asil = row[asil_idx] if asil_idx >= 0 and asil_idx < len(row) else ""
                    related = row[related_idx] if related_idx >= 0 and related_idx < len(row) else rid
                    _merge_entry(rid, asil, related)
                    if name_idx >= 0 and name_idx < len(row):
                        name = str(row[name_idx] or "").strip()
                        if name:
                            _merge_entry(name, asil, related)

    text_map = _build_req_map_from_texts(texts or [])
    for key, value in text_map.items():
        if key not in mapping:
            mapping[key] = dict(value)
            continue
        if value.get("asil") and not mapping[key].get("asil"):
            mapping[key]["asil"] = value["asil"]
        if value.get("related") and not mapping[key].get("related"):
            mapping[key]["related"] = value["related"]
    return mapping


def _build_uds_asil_map(uds_doc_paths: Optional[List[str]]) -> Dict[str, str]:
    """SwUDS(v3.02 세로 key-value 표) 문서에서 {함수명(lower): ASIL 등급} 직독.

    impact 파이프라인의 검증된 extract_function_asil_from_kv_tables를 재사용한다(권위 문서
    직독 — SDS 이름 휴리스틱 매칭보다 정확). 이 파서·상위 파서가 path.exists()/open()으로
    로컬 fs 직접 접근하므로, 호출부가 cloudium U: 경로를 로컬 tmp로 변환해 넘긴 뒤에만 유효.
    다중 UDS·동명 함수는 max-merge(안전측 — 낮은 등급 채택은 under-report). 파싱 실패는
    graceful skip(추적성 본류 비차단). lazy import로 계층 위생 + iso26262 모듈 로드 실패 격리.
    """
    out: Dict[str, str] = {}
    for p in (uds_doc_paths or []):
        try:
            path = Path(str(p))
            if not path.exists() or path.suffix.lower() != ".docx":
                continue
            from backend.services.iso26262_doc_asil_extractor import (
                extract_function_asil_from_kv_tables,
            )
            parsed = extract_function_asil_from_kv_tables(path.read_bytes())
            for k, v in (parsed or {}).items():
                nk = str(k).strip().lower()
                if not nk or not v:
                    continue
                if nk not in out or _ASIL_RANK.get(str(v).upper(), -1) > _ASIL_RANK.get(str(out[nk]).upper(), -1):
                    out[nk] = v
        except Exception:
            continue
    return out


def enrich_function_details_with_docs(
    function_details: Dict[str, Dict[str, Any]],
    function_table_rows: Optional[List[List[Any]]] = None,
    *,
    req_doc_paths: Optional[List[str]] = None,
    sds_doc_paths: Optional[List[str]] = None,
    uds_doc_paths: Optional[List[str]] = None,
) -> Dict[str, Dict[str, Any]]:
    if not isinstance(function_details, dict) or not function_details:
        return function_details

    req_paths = [str(p).strip() for p in (req_doc_paths or []) if str(p).strip()]
    sds_paths = [str(p).strip() for p in (sds_doc_paths or []) if str(p).strip()]
    uds_paths = [str(p).strip() for p in (uds_doc_paths or []) if str(p).strip()]
    req_map = _build_req_map_from_doc_paths(req_paths) if req_paths else {}
    # SwUDS 문서 직독 ASIL(권위) — 함수별 blank를 SDS 이름매칭보다 먼저 채운다(under-report 방지).
    uds_name_asil = _build_uds_asil_map(uds_paths)

    sds_map: Dict[str, Dict[str, str]] = {}
    for path in sds_paths:
        for key, value in _extract_sds_partition_map(path).items():
            if key not in sds_map:
                sds_map[key] = dict(value)
                continue
            for field in ("asil", "related", "description"):
                if value.get(field) and not sds_map[key].get(field):
                    sds_map[key][field] = value[field]

    fid_to_swcom: Dict[str, str] = {}
    if isinstance(function_table_rows, list):
        for row in function_table_rows:
            if not isinstance(row, list) or len(row) < 4:
                continue
            swcom = str(row[0] or "").strip()
            fid = str(row[2] or "").strip()
            if swcom and fid:
                fid_to_swcom[fid] = swcom

    def _normalize_key(value: str) -> str:
        return re.sub(r"[^a-z0-9]", "", str(value or "").lower())

    def _extract_related_ids(value: str) -> Set[str]:
        return {m.group(1).upper() for m in _REQ_ID_PAT.finditer(str(value or ""))}

    def _tokenize_text(value: str) -> List[str]:
        words = re.split(r"[^a-z0-9]+", str(value or "").lower())
        return [w for w in words if len(w) >= 3]

    def _prototype_candidates(info: Dict[str, Any]) -> List[str]:
        prototype = str(info.get("prototype") or "").strip()
        if not prototype:
            return []
        parts: List[str] = [prototype]
        m = re.match(r"^\s*([A-Za-z_]\w*(?:\s*\*+)?)\s+([A-Za-z_]\w*)\s*\((.*)\)\s*$", prototype)
        if m:
            ret_type = m.group(1).strip()
            fn_name = m.group(2).strip()
            params = m.group(3).strip()
            parts.extend([ret_type, fn_name])
            if params and params.lower() != "void":
                for chunk in params.split(","):
                    token = str(chunk).strip()
                    if token:
                        parts.append(token)
        return [p for p in parts if p]

    def _module_candidates(info: Dict[str, Any], fid: str) -> List[str]:
        candidates: List[str] = []
        func_name = str(info.get("name") or "").strip()
        if func_name:
            candidates.append(func_name)
            stripped = re.sub(r"^[gs]_", "", func_name, flags=re.I)
            if stripped != func_name:
                candidates.append(stripped)
            words = re.sub(r"([a-z])([A-Z])", r"\1 \2", stripped.replace("_", " "))
            words = re.sub(r"\bctrl\b", "control", words, flags=re.I)
            words = re.sub(r"\bdiag\b", "diagnostic", words, flags=re.I)
            if words.strip():
                candidates.append(words)
        module_name = str(info.get("module_name") or "").strip()
        if module_name:
            candidates.append(module_name)
            base = re.sub(r"_pds$", "", module_name, flags=re.I)
            candidates.append(base)
            tokenized = re.sub(r"([a-z])([A-Z])", r"\1 \2", base.replace("_", " "))
            tokenized = re.sub(r"\bctrl\b", "control", tokenized, flags=re.I)
            tokenized = re.sub(r"\bdiag\b", "diagnostic", tokenized, flags=re.I)
            tokenized = re.sub(r"\bprev(?:ious)?\b", "previous", tokenized, flags=re.I)
            words = [w for w in tokenized.split() if w.lower() not in {"ap", "drv", "sys", "pds", "main", "func"}]
            if words:
                candidates.append(" ".join(words))
        swcom = fid_to_swcom.get(fid, "")
        if swcom:
            candidates.append(swcom)
        return [c for c in dict.fromkeys([c.strip() for c in candidates if c and c.strip()])]

    def _find_sds_info(info: Dict[str, Any], fid: str) -> Tuple[Optional[str], Optional[Dict[str, str]], str]:
        candidates = _module_candidates(info, fid)
        swcom_direct_fallback: Tuple[Optional[str], Optional[Dict[str, str]], str] = (None, None, "")
        for candidate in candidates:
            direct = sds_map.get(candidate.lower())
            if direct:
                if candidate.lower().startswith("swcom_"):
                    swcom_direct_fallback = (candidate.lower(), direct, "direct")
                    continue
                return candidate.lower(), direct, "direct"

        norm_candidates = [_normalize_key(c) for c in candidates]
        for candidate, nc in zip(candidates, norm_candidates):
            if not nc:
                continue
            for key, value in sds_map.items():
                nk = _normalize_key(key)
                if nk and nc == nk:
                    if str(key).lower().startswith("swcom_"):
                        swcom_direct_fallback = (key, value, "normalized_exact")
                        continue
                    return key, value, "normalized_exact"

        related_ids = _extract_related_ids(str(info.get("related") or info.get("comment_related") or ""))
        proto_texts = _prototype_candidates(info)
        proto_tokens: Set[str] = set()
        for text in proto_texts + candidates:
            proto_tokens.update(_tokenize_text(text))
        best_key: Optional[str] = None
        best_value: Optional[Dict[str, str]] = None
        best_score = 0
        for key, value in sds_map.items():
            if str(key or "").lower().startswith("swcom_"):
                continue
            score = 0
            sds_related_ids = _extract_related_ids(str(value.get("related") or ""))
            overlap = related_ids & sds_related_ids
            if overlap:
                score += len(overlap) * 10
            sds_tokens = set(_tokenize_text(key))
            sds_tokens.update(_tokenize_text(str(value.get("description") or "")))
            token_overlap = proto_tokens & sds_tokens
            if token_overlap:
                score += min(len(token_overlap), 6) * 2
            if score > best_score and (overlap or len(token_overlap) >= 2):
                best_key = key
                best_value = value
                best_score = score
        if best_key and best_value:
            return best_key, best_value, "related_prototype"

        # Containment matching is intentionally strict to avoid generic terms
        # like "Lib" or "Main" matching unrelated SDS rows.
        for candidate, nc in zip(candidates, norm_candidates):
            if len(nc) < 6:
                continue
            cand_words = [w for w in re.split(r"[^a-z0-9]+", candidate.lower()) if len(w) >= 4]
            if not cand_words:
                continue
            for key, value in sds_map.items():
                nk = _normalize_key(key)
                if len(nk) < 6:
                    continue
                key_words = [w for w in re.split(r"[^a-z0-9]+", key.lower()) if len(w) >= 4]
                if not key_words:
                    continue
                if nc in nk or nk in nc:
                    overlap = set(cand_words) & set(key_words)
                    if overlap:
                        return key, value, "normalized_overlap"
        if swcom_direct_fallback[0] and swcom_direct_fallback[1]:
            return swcom_direct_fallback
        return None, None, ""

    for fid, info in function_details.items():
        if not isinstance(info, dict):
            continue
        sds_key, sds_info, sds_match_mode = _find_sds_info(info, str(fid))
        if sds_info:
            sds_scope = "swcom" if str(sds_key or "").lower().startswith("swcom_") else "function"
            if sds_scope == "swcom":
                mapping_confidence = 0.55
            elif sds_match_mode == "direct":
                mapping_confidence = 0.95
            elif sds_match_mode == "normalized_exact":
                mapping_confidence = 0.85
            elif sds_match_mode == "related_prototype":
                mapping_confidence = 0.8
            else:
                mapping_confidence = 0.7
            info["sds_match_key"] = sds_key or ""
            info["sds_match_mode"] = sds_match_mode
            info["sds_match_scope"] = sds_scope
            info["mapping_confidence"] = mapping_confidence
            current_related = str(info.get("related") or "").strip()
            current_asil = str(info.get("asil") or "").strip().upper()
            if (not current_related) or current_related in {"TBD", "N/A", "-"}:
                if sds_info.get("related"):
                    info["related"] = sds_info["related"]
                    info["related_source"] = "sds"
            if (not current_asil) or current_asil in {"TBD", "N/A", "-"}:
                if sds_info.get("asil"):
                    info["asil"] = _normalize_asil_value(sds_info["asil"])
                    info["asil_source"] = "sds"
            desc = str(info.get("description") or "").strip()
            if (not desc or desc.lower().startswith("function")) and sds_info.get("description"):
                info["description"] = sds_info["description"]
                info["description_source"] = "sds"
            elif str(info.get("description_source") or "").strip() in {"", "inference"}:
                info["description_source"] = "sds_match"

        related = str(info.get("related") or "").strip()
        matched_req_with_asil = False
        matched_req_without_asil = False
        if related and req_map:
            for match in _REQ_ID_PAT.finditer(related):
                req = req_map.get(match.group(1).lower())
                if not isinstance(req, dict):
                    continue
                req_asil_raw = str(req.get("asil") or "").strip()
                if req_asil_raw:
                    matched_req_with_asil = True
                else:
                    matched_req_without_asil = True
                asil = _normalize_asil_value(req.get("asil", ""))
                cur_asil = str(info.get("asil") or "").strip().upper()
                if asil and ((not cur_asil) or cur_asil in {"TBD", "N/A", "-"}):
                    info["asil"] = asil
                    info["asil_source"] = "srs"
                    break
        cur_asil = str(info.get("asil") or "").strip().upper()
        if (
            related
            and matched_req_without_asil
            and not matched_req_with_asil
            and ((not cur_asil) or cur_asil in {"TBD", "N/A", "-"})
        ):
            info["asil"] = "QM"
            info["asil_source"] = "srs_default_qm"

        # SwUDS 문서 직독 ASIL — 등급 낮추기 절대 없음(상향만). 단 소스 유래(asil_source="comment",
        # C @asil Doxygen)와 문서 유래(sds 이름휴리스틱·srs·inference·blank)를 구분한다:
        #  · 문서 유래가 UDS보다 낮으면 → UDS로 상향(under-report 해소). SDS 오매칭·SRS 미스보다
        #    UDS 문서 직독이 권위. UDS 문서에만 명시된 안전함수(보안접근 등)를 회복.
        #  · 소스 코드 @asil(c_source 권위, backend/services/CLAUDE.md "c_source > swuds")이 UDS보다
        #    낮으면 → 등급 유지 + asil_doc_conflict 표면화(자동 상향 안 함). 문서>코드 불일치는 수동
        #    검토 신호이지 코드 등급을 문서에 맞춰 올릴 사안 아님 — impact _enrich_asil_from_uds와 lockstep.
        #    이로써 같은 함수가 impact 탭·추적성 탭에서 모순 표시되지 않는다(deep-review W1).
        # SDS direct/exact가 UDS보다 높은 경우(예 sf_* secure flash)는 max라 SDS 우위 보존([B]).
        if uds_name_asil:
            _nm = str(info.get("name") or "").strip().lower()
            _uds_a = uds_name_asil.get(_nm) if _nm else None
            if _uds_a:
                _cur = str(info.get("asil") or "").strip().upper()
                if _ASIL_RANK.get(str(_uds_a).upper(), -1) > _ASIL_RANK.get(_cur, -1):
                    if str(info.get("asil_source") or "").strip() == "comment":
                        info["asil_doc_conflict"] = f"source={_cur or '(빈)'}<uds={_uds_a}"
                    else:
                        info["asil"] = _uds_a
                        info["asil_source"] = "uds"

    return function_details


def _split_doc_function_blocks(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    blocks: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    for ln in text.splitlines():
        line = ln.strip()
        m = re.match(r"^(SwUFn_\d+)\s*:\s*(.+)$", line)
        if m:
            if current:
                blocks.append(current)
            current = {"id": m.group(1), "title": f"{m.group(1)}: {m.group(2).strip()}", "lines": []}
            continue
        if current:
            current["lines"].append(line)
    if current:
        blocks.append(current)
    return blocks


def _collect_section_lines(lines: List[str], header: str) -> List[str]:
    results: List[str] = []
    collecting = False
    for ln in lines:
        line = ln.strip()
        if not line:
            if collecting and results:
                continue
        if line.startswith(header):
            collecting = True
            tail = line[len(header) :].strip()
            if tail:
                results.append(tail)
            continue
        if collecting:
            if (
                line.startswith("[")
                or line.startswith("ID")
                or line.startswith("Name")
                or line.startswith("Prototype")
                or line.startswith("Description")
                or line.startswith("ASIL")
                or line.startswith("Related ID")
                or line.startswith("선행조건")
                or line.startswith("사용 전역변수")
                or line.startswith("Called Function")
                or line.startswith("Calling Function")
            ):
                collecting = False
                continue
            if re.match(r"^SwUFn_\d+", line):
                collecting = False
                continue
            if line:
                results.append(line)
    return results


def _extract_state_tokens(lines: List[str]) -> List[str]:
    states: List[str] = []
    for ln in lines:
        for token in re.findall(r"\bST_[A-Za-z0-9_]+\b", ln):
            if token not in states:
                states.append(token)
    return states


def _extract_requirements_from_doc(text: str) -> List[str]:
    if not text:
        return []
    lines = [ln.rstrip() for ln in text.splitlines()]
    blocks: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    collecting_desc = False
    desc_lines: List[str] = []
    stop_keys = {
        "Rationale",
        "Priority",
        "Status",
        "Risk",
        "Reuse",
        "Verification criteria",
        "System State",
        "Software State",
        "Type",
    }

    def _flush() -> None:
        nonlocal current, desc_lines, collecting_desc
        if not current:
            return
        if desc_lines:
            current["description"] = " ".join(desc_lines).strip()
        blocks.append(current)
        current = {}
        desc_lines = []
        collecting_desc = False

    for raw in lines:
        line = raw.strip()
        if not line:
            if collecting_desc:
                collecting_desc = False
            continue

        m = re.search(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b", line)
        if line.startswith("ID") and m:
            _flush()
            current = {"id": m.group(1)}
            continue
        if re.match(r"^Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+[:\s]", line):
            _flush()
            current = {"id": m.group(1) if m else line.split()[0]}
            if ":" in line:
                current["name"] = line.split(":", 1)[1].strip()
            continue

        if not current:
            continue

        if line.startswith("Name"):
            current["name"] = line.split(None, 1)[-1].strip()
            continue
        if line.startswith("Description"):
            desc = line.split(None, 1)[-1].strip()
            if desc and desc != "Description":
                desc_lines.append(desc)
            collecting_desc = True
            continue
        asil_line_match = re.match(
            r"^(?:ASIL|Safety\s*Level|Safety\s*Class|Integrity\s*Level)[\s\-_:]*(.*)$",
            line, re.I,
        )
        if asil_line_match:
            collecting_desc = False
            asil_val = asil_line_match.group(1).strip()
            if not asil_val:
                asil_val = line.split(None, 1)[-1].strip() if len(line.split(None, 1)) > 1 else ""
            if asil_val:
                norm = re.match(r"(?:ASIL[\s\-_]*)?([A-D])\s*(?:\([A-D]\))?|QM", asil_val, re.I)
                current["asil"] = norm.group(0).strip() if norm else asil_val
            continue
        related_match = re.match(r"^(?:Related\s*ID|Related\s*Req|Trace(?:ability)?|Parent\s*ID)[\s:]*(.*)$", line, re.I)
        if related_match:
            collecting_desc = False
            related_val = related_match.group(1).strip().lstrip(":").strip()
            if related_val:
                current["related_id"] = related_val
            continue
        if any(line.startswith(k) for k in stop_keys):
            collecting_desc = False
            continue
        if collecting_desc:
            desc_lines.append(line)

    _flush()
    results: List[str] = []
    for block in blocks:
        rid = block.get("id") or ""
        name = block.get("name") or ""
        desc = block.get("description") or ""
        asil = block.get("asil") or ""
        related = block.get("related_id") or ""
        parts = []
        if rid:
            parts.append(rid)
        if name:
            parts.append(name)
        if desc:
            parts.append(f"- {desc}")
        if asil:
            parts.append(f"[ASIL:{asil}]")
        if related:
            parts.append(f"[Related:{related}]")
        if parts:
            results.append(" ".join(parts))
    return results


def _extract_requirements_fallback(text: str, max_items: int = 200) -> List[str]:
    if not text:
        return []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    hits: List[str] = []
    req_keywords = re.compile(
        r"\b(shall|must|should|requirement|requirements|specification|spec)\b",
        re.I,
    )
    ko_keywords = re.compile(r"(요구|요건|필수|해야|기능|명세)")
    id_keywords = re.compile(r"\b(REQ|SRS|SDS|SR|SWR|SYS|SW)-?\d+\b", re.I)
    for ln in lines:
        if req_keywords.search(ln) or ko_keywords.search(ln) or id_keywords.search(ln):
            hits.append(ln)
            if len(hits) >= max_items:
                break
    return hits


def _extract_doc_section(text: str, title: str) -> str:
    if not text or not title:
        return ""
    title_clean = re.escape(title.strip())
    pattern = re.compile(rf"^\s*\d+(?:\.\d+)*\s+{title_clean}\s*$", re.I)
    lines = text.splitlines()
    start = None
    for idx, ln in enumerate(lines):
        if pattern.match(ln.strip()):
            start = idx + 1
            break
    if start is None:
        return ""
    collected: List[str] = []
    for ln in lines[start:]:
        if re.match(r"^\s*\d+(?:\.\d+)*\s+\S+", ln.strip()):
            break
        if ln.strip():
            collected.append(ln.rstrip())
    return "\n".join(collected).strip()


def _extract_requirement_blocks(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    lines = [ln.rstrip() for ln in text.splitlines()]
    blocks: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    collecting_desc = False
    desc_lines: List[str] = []
    stop_keys = {
        "Rationale",
        "Priority",
        "Status",
        "Risk",
        "Reuse",
        "Related ID",
        "Verification criteria",
        "ASIL",
        "System State",
        "Software State",
        "Type",
    }

    def _flush() -> None:
        nonlocal current, desc_lines, collecting_desc
        if not current:
            return
        if desc_lines:
            current["description"] = " ".join(desc_lines).strip()
        # 표 셀 파이프 구분자 정제 — docx 표가 '| Battery Power Source' 형태로 추출돼
        # name/description에 '| ' 잡음이 남던 것 제거(라운드110). 헤딩형은 이미 깨끗.
        for _k in ("name", "description"):
            if current.get(_k):
                current[_k] = re.sub(r"\s*\|\s*", " ", str(current[_k])).strip()
        if current.get("name") == "":
            current.pop("name", None)
        blocks.append(current)
        current = {}
        desc_lines = []
        collecting_desc = False

    id_re = re.compile(r"\b(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\b")
    for raw in lines:
        line = raw.strip()
        if not line:
            if collecting_desc:
                collecting_desc = False
            continue

        m = id_re.search(line)
        # 마크다운 헤딩형 요구사항 정의 '#### SwTR_0101: Auto Close' — 깨끗한 제목 포착.
        # 기존 '^Sw...' 분기는 '#' 접두를 못 잡아 헤딩의 정상 제목을 놓치고, 표 파편의
        # 빈/'| ' 잡음 name만 잡혔다(요구사항 제목이 빈/잡음으로 표시되던 회귀, 라운드110).
        m_head = re.match(r"^#{1,6}\s*(Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+)\s*[:：]\s*(.+)$", line)
        if m_head:
            _flush()
            current = {"id": m_head.group(1), "name": m_head.group(2).strip()}
            continue
        if line.startswith("ID") and m:
            _flush()
            current = {"id": m.group(1)}
            continue
        if re.match(r"^Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+[:\s]", line):
            _flush()
            current = {"id": m.group(1) if m else line.split()[0]}
            if ":" in line:
                current["name"] = line.split(":", 1)[1].strip()
            continue

        if not current:
            continue

        if line.startswith("Name"):
            current["name"] = line.split(None, 1)[-1].strip()
            continue
        if line.startswith("Description"):
            desc = line.split(None, 1)[-1].strip()
            if desc and desc != "Description":
                desc_lines.append(desc)
            collecting_desc = True
            continue
        if line.startswith("Related ID"):
            current["related_ids"] = line.split(None, 1)[-1].strip()
            continue
        if any(line.startswith(k) for k in stop_keys):
            collecting_desc = False
            continue
        if collecting_desc:
            desc_lines.append(line)

    _flush()
    return blocks


def generate_uds_requirements_preview(texts: List[str]) -> Dict[str, Any]:
    items: List[Dict[str, Any]] = []
    for txt in texts:
        items.extend(_extract_requirement_blocks(txt))
    seen = set()
    uniq: List[Dict[str, Any]] = []
    for item in items:
        key = (item.get("id") or "", item.get("name") or "", item.get("description") or "")
        if key in seen:
            continue
        seen.add(key)
        uniq.append(item)
    counts: Dict[str, int] = {}
    for item in uniq:
        rid = str(item.get("id") or "")
        m = re.match(r"^(Sw[A-Za-z]+)_\d+", rid)
        if m:
            counts[m.group(1)] = counts.get(m.group(1), 0) + 1
    return {"items": uniq, "counts": counts}


def generate_uds_requirements_mapping(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    mappings: List[Dict[str, Any]] = []
    for item in items:
        rid = str(item.get("id") or "")
        if rid and not (rid.startswith("SwTR_") or rid.startswith("SwTSR_")):
            continue
        related = str(item.get("related_ids") or "")
        swcom = re.findall(r"\bSwCom_\d+\b", related)
        swfn = re.findall(r"\bSwFn_\d+\b", related)
        if not swcom and not swfn:
            continue
        mappings.append(
            {
                "requirement_id": rid,
                "requirement_name": item.get("name") or "",
                "related_swcom": swcom,
                "related_swfn": swfn,
            }
        )
    return mappings


def _extract_doc_function_names(texts: List[str]) -> List[str]:
    names: set[str] = set()
    for txt in texts:
        if not txt:
            continue
        for name in re.findall(r"\b[gs]_[A-Za-z0-9_]+\b", txt):
            names.add(name)
    return sorted(names)


def generate_uds_function_mapping(texts: List[str], source_root: str) -> Dict[str, Any]:
    doc_funcs = _extract_doc_function_names(texts)
    source_info = _scan_source_function_names(source_root)
    source_funcs = set(source_info.get("names") or [])
    matched = [fn for fn in doc_funcs if fn in source_funcs]
    missing = [fn for fn in doc_funcs if fn not in source_funcs]

    fuzzy_matched: List[Dict[str, str]] = []
    still_missing: List[str] = []
    source_lower_map = {s.lower(): s for s in source_funcs}
    for fn in missing:
        fn_lower = fn.lower()
        if fn_lower in source_lower_map:
            fuzzy_matched.append({"doc": fn, "source": source_lower_map[fn_lower], "method": "case_insensitive"})
            continue
        fn_stripped = re.sub(r"^(g_|s_|static\s+)", "", fn)
        found = False
        for sfn in source_funcs:
            sfn_stripped = re.sub(r"^(g_|s_|static\s+)", "", sfn)
            if fn_stripped and sfn_stripped and fn_stripped.lower() == sfn_stripped.lower():
                fuzzy_matched.append({"doc": fn, "source": sfn, "method": "prefix_stripped"})
                found = True
                break
        if not found:
            still_missing.append(fn)

    return {
        "doc_functions": doc_funcs,
        "matched": matched,
        "fuzzy_matched": fuzzy_matched,
        "missing": still_missing,
        "source_scanned": int(source_info.get("scanned") or 0),
    }


def _normalize_trace_mapping_entry(entry: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    rid = str(entry.get("requirement_id") or entry.get("requirement") or entry.get("req_id") or "").strip()
    if not rid:
        return None
    raw_sources = (
        entry.get("source_ids")
        or entry.get("source_id")
        or entry.get("source")
        or entry.get("sources")
        or []
    )
    source_ids: List[str] = []
    if isinstance(raw_sources, str):
        source_ids = [s.strip() for s in raw_sources.split(",") if s.strip()]
    elif isinstance(raw_sources, list):
        source_ids = [str(s).strip() for s in raw_sources if str(s).strip()]
    else:
        source_ids = [str(raw_sources).strip()] if str(raw_sources).strip() else []
    if not source_ids:
        return None
    return {"requirement_id": rid, "source_ids": source_ids}


def _parse_traceability_json(text: str) -> List[Dict[str, Any]]:
    try:
        data = json.loads(text)
    except Exception:
        return []
    items: List[Dict[str, Any]] = []
    if isinstance(data, dict) and isinstance(data.get("mappings"), list):
        data = data.get("mappings")
    if isinstance(data, list):
        for raw in data:
            if not isinstance(raw, dict):
                continue
            item = _normalize_trace_mapping_entry(raw)
            if item:
                items.append(item)
        return items
    if isinstance(data, dict):
        for rid, src in data.items():
            item = _normalize_trace_mapping_entry({"requirement_id": rid, "source_ids": src})
            if item:
                items.append(item)
    return items


def _parse_traceability_csv(text: str) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    try:
        reader = csv.DictReader(StringIO(text))
    except Exception:
        return []
    if not reader.fieldnames:
        return []
    for row in reader:
        if not isinstance(row, dict):
            continue
        item = _normalize_trace_mapping_entry(row)
        if item:
            items.append(item)
    return items


def _parse_traceability_text(text: str) -> List[Dict[str, Any]]:
    if not text:
        return []
    trimmed = text.strip()
    if not trimmed:
        return []
    if trimmed.startswith("{") or trimmed.startswith("["):
        items = _parse_traceability_json(trimmed)
        if items:
            return items
    return _parse_traceability_csv(trimmed)


def generate_uds_traceability_mapping(
    items: List[Dict[str, Any]],
    mapping_texts: List[str],
    function_details: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    req_ids = sorted({str(x.get("id") or "").strip() for x in items if str(x.get("id") or "").strip()})
    mappings: Dict[str, List[str]] = {}
    for text in mapping_texts:
        for entry in _parse_traceability_text(text):
            rid = entry["requirement_id"]
            srcs = entry["source_ids"]
            if rid not in mappings:
                mappings[rid] = []
            for src in srcs:
                if src not in mappings[rid]:
                    mappings[rid].append(src)
    if function_details:
        for fid, info in function_details.items():
            related = str(info.get("comment_related") or info.get("related") or "").strip()
            fn_name = str(info.get("name") or "").strip()
            if related and related.upper() not in {"TBD", "N/A", "-"}:
                for rid in re.findall(r"Sw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK|Com)_\d+", related):
                    src_label = fn_name or fid
                    mappings.setdefault(rid, [])
                    if src_label not in mappings[rid]:
                        mappings[rid].append(src_label)
    req_id_set = set(req_ids)
    if req_ids:
        req_lower_map = {r.lower(): r for r in req_ids}
        for rid_key in list(mappings.keys()):
            if rid_key not in req_id_set:
                canonical = req_lower_map.get(rid_key.lower())
                if canonical:
                    mappings.setdefault(canonical, []).extend(mappings.pop(rid_key))
    mapped_req_ids = [rid for rid in req_ids if rid in mappings]
    missing_req_ids = [rid for rid in req_ids if rid not in mappings]
    extra_mapping = [rid for rid in mappings.keys() if rid not in req_id_set]
    source_ids: List[str] = []
    for srcs in mappings.values():
        for src in srcs:
            if src not in source_ids:
                source_ids.append(src)
    mapping_pairs = [
        {"requirement_id": rid, "source_ids": mappings[rid]} for rid in mapped_req_ids
    ]
    return {
        "total_requirements": len(req_ids),
        "mapped_requirements": mapped_req_ids,
        "missing_requirements": missing_req_ids,
        "extra_mapping": extra_mapping,
        "mapping_pairs": mapping_pairs,
        "total_sources": len(source_ids),
    }


def _normalize_req_id(rid: str) -> str:
    """Normalize requirement ID: remove all whitespace, uppercase for consistent matching."""
    rid = "".join(rid.split())  # remove internal whitespace too (e.g., "SwRS_ 001" → "SWRS_001")
    if not rid:
        return rid
    return rid.upper()


def _normalize_vcast_rows(rows: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    out: Dict[str, List[Dict[str, Any]]] = {}
    seen_keys: set = set()  # (rid, testcase, source) dedup key
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        testcase = row.get("testcase") or row.get("subprogram") or ""
        source = row.get("source") or ""
        result = row.get("result") or ""
        # Deduplicate: same requirement + testcase + source + subprogram = skip.
        # subprogram을 키에 포함(reviewer INFO): 서로 다른 함수가 동일 testcase 명을
        # 가지고 같은 SRS req로 bridge될 때 한 건이 silent drop되던 경계 케이스 방지.
        # 비-vcast 행은 subprogram이 비어 키에 영향 없음(STS/SUTS/SITS 동작 불변).
        dedup_key = (rid, testcase, source, str(row.get("subprogram") or ""))
        if dedup_key in seen_keys:
            continue
        seen_keys.add(dedup_key)
        out.setdefault(rid, []).append(
            {
                "testcase": testcase,
                "result": result,
                "unit": row.get("unit") or "",
                "report": row.get("report") or "",
                "source": source,
                # subprogram(=시험 대상 함수, vcast 산출물에서만 채워짐) 보존 — 함수중심 그래프가
                # 함수↔VectorCAST를 잇는 키. STS/SUTS/SITS는 비어 무영향(프론트 매칭은 unit 사용).
                "subprogram": row.get("subprogram") or "",
                "trace_type": row.get("trace_type") or "direct",
                "confidence": row.get("confidence") if row.get("confidence") not in (None, "") else ("exact" if source in ("STS", "SUTS", "SITS") else "fuzzy"),
            }
        )
    return out


# ASIL 등급 순위(ISO 26262: QM<A<B<C<D) — 요구사항 ASIL은 연결 설계요소 중 최고로 도출.
_ASIL_RANK = {"QM": 0, "A": 1, "B": 2, "C": 3, "D": 4}


def _asil_max_of(raw_values: List[str]) -> str:
    """주어진 ASIL 문자열들('A','QM','A, B' 등) 중 최고 등급을 canonical 토큰으로 반환.

    _normalize_asil_value로 정규화 후 토큰별 순위 비교. 인식 불가/빈 값은 무시.
    하나도 인식 못 하면 빈 문자열(graceful — ASIL 미상).
    """
    best = ""
    best_rank = -1
    for v in raw_values:
        for tok in _normalize_asil_value(v).split(","):
            tok = tok.strip().upper()
            r = _ASIL_RANK.get(tok)
            if r is not None and r > best_rank:
                best_rank = r
                best = tok
    return best


def generate_uds_traceability_matrix(
    items: List[Dict[str, Any]],
    mapping_pairs: Optional[List[Dict[str, Any]]] = None,
    vcast_rows: Optional[List[Dict[str, Any]]] = None,
    sds_pairs: Optional[List[Dict[str, Any]]] = None,
    sits_rows: Optional[List[Dict[str, Any]]] = None,
    uds_function_ids: Optional[List[str]] = None,
    component_asil: Optional[Dict[str, str]] = None,
    hsis_pairs: Optional[List[Dict[str, Any]]] = None,
    uds_function_asil: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    # ASIL 결합(P5) — {컴포넌트/함수명(lower): ASIL}. SDS 추출(component_asil)에서 전달.
    # 요구사항별 ASIL = 연결된 SDS 컴포넌트·UDS 함수의 ASIL 중 최고(QM<A<B<C<D).
    # hiMA는 ASIL을 셀에 비노출 → 우리 차별점. 데이터 없으면 빈 문자열(graceful).
    comp_asil_map: Dict[str, str] = {}
    for _k, _v in (component_asil or {}).items():
        _kk = str(_k or "").strip().lower()
        _vv = str(_v or "").strip()
        if _kk and _vv:
            comp_asil_map[_kk] = _vv
    # SwUDS 문서 직독 함수 ASIL을 요구사항 ASIL 도출 소스에 max-merge. SDS 컴포넌트 ASIL만으론
    # UDS 함수 단위 안전등급(보안접근·슬립감지 등)이 누락돼 요구사항 ASIL이 under-report될 수 있다
    # (extract-mapping이 uds_function_asil로 전달). 같은 함수명에 SDS/UDS 등급이 다르면 max —
    # 낮은 등급 채택은 under-report(위험), over-report는 안전측. (라이브: kjpds02_pv 함수 ASIL 갭)
    for _k, _v in (uds_function_asil or {}).items():
        _kk = str(_k or "").strip().lower()
        _vv = str(_v or "").strip()
        if _kk and _vv and (_kk not in comp_asil_map or _ASIL_RANK.get(_vv.upper(), -1) > _ASIL_RANK.get(str(comp_asil_map[_kk]).upper(), -1)):
            comp_asil_map[_kk] = _vv

    # Build original→normalized ID mapping to preserve display IDs
    raw_ids = sorted({str(x.get("id") or "").strip() for x in items if str(x.get("id") or "").strip()})
    logger = logging.getLogger(__name__)
    norm_to_raw: Dict[str, str] = {}
    # 정합성 감사(trace_integrity): 한 canonical로 붕괴하는 raw 철자 전체를 보존한다.
    # 기존 first-wins(norm_to_raw)는 표시용 1개만 남기고 나머지를 log로만 흘려 삼키므로,
    # 여기서 전 변형을 모아 build_integrity_audit가 '정규화 충돌'을 표면화하게 한다(additive).
    norm_to_raws: Dict[str, List[str]] = {}
    for rid in raw_ids:
        norm = _normalize_req_id(rid)
        if norm in norm_to_raw and norm_to_raw[norm] != rid:
            logger.warning("Duplicate requirement ID after normalization: '%s' and '%s' both normalize to '%s'", norm_to_raw[norm], rid, norm)
        if norm not in norm_to_raw:
            norm_to_raw[norm] = rid  # keep first occurrence for display
        norm_to_raws.setdefault(norm, []).append(rid)

    req_ids = sorted(norm_to_raw.keys())

    # 요구사항 표시명(name) — 한 ID에 파편 item이 여러 개(빈 헤더행/제목행)면 가장 정보량
    # 많은(파이프 정제 후 비지 않고 긴) name을 채택. SRS 표 추출이 ID당 다중 행을 만들어
    # 첫 행(빈 name)이 표시되던 문제 해소(라운드110). 방어적으로 '| ' 잔여 잡음도 정제.
    name_map: Dict[str, str] = {}
    # SyRS 상위 추적(요구→상위 시스템요구) — items의 raw related_ids에서 Sy* 토큰을 보존 소비.
    # builder 내부 _normalize_req_id(requirements.py:1804)는 Sy 보존(공백제거+대문자)이라 collapse 안 됨.
    # → SR→SyRS→SwRS 체인의 상위 링크를 surface(배지/밴드). 신규 추출 아님(이미 들어온 raw 소비).
    syrs_map: Dict[str, List[str]] = {}
    _SY_RE = re.compile(r"Sy[A-Za-z]{2,}_\d+")
    for x in items:
        rid_n = _normalize_req_id(str(x.get("id") or "").strip())
        if not rid_n:
            continue
        nm = re.sub(r"\s*\|\s*", " ", str(x.get("name") or "")).strip()
        if len(nm) > len(name_map.get(rid_n, "")):
            name_map[rid_n] = nm
        raw_rel = str(x.get("related_ids") or x.get("related") or "")
        if raw_rel:
            existing = syrs_map.setdefault(rid_n, [])
            for tok in _SY_RE.findall(raw_rel):
                tok = tok.upper()
                if tok not in existing:
                    existing.append(tok)

    # ── UDS function mapping (requirement → source functions) ──
    mapping_pairs = mapping_pairs or []
    map_lookup: Dict[str, List[str]] = {}
    for row in mapping_pairs:
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        srcs = row.get("source_ids") or []
        if isinstance(srcs, str):
            srcs = [s.strip() for s in srcs.split(",") if s.strip()]
        elif isinstance(srcs, list):
            srcs = [str(s).strip() for s in srcs if str(s).strip()]
        else:
            srcs = []
        existing = map_lookup.get(rid, [])
        for s in srcs:
            if s not in existing:
                existing.append(s)
        map_lookup[rid] = existing

    # ── SDS component mapping (requirement → design components) ──
    sds_lookup: Dict[str, List[str]] = {}
    for row in (sds_pairs or []):
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        comps = row.get("component_ids") or []
        if isinstance(comps, str):
            comps = [c.strip() for c in comps.split(",") if c.strip()]
        elif isinstance(comps, list):
            comps = [str(c).strip() for c in comps if str(c).strip()]
        else:
            comps = []
        existing = sds_lookup.get(rid, [])
        for c in comps:
            if c not in existing:
                existing.append(c)
        sds_lookup[rid] = existing

    # ── HSIS interface mapping (requirement → interface signals) ──
    # 시스템 레벨 인터페이스 밴드(design-arm). HSIS extract 엔드포인트가 Related ID(Sy*)를
    # SW namespace로 평탄화해 hsis_pairs[{requirement_id, hsis_signals:[HSI_xx/SW변수]}]로 전달.
    hsis_lookup: Dict[str, List[str]] = {}
    for row in (hsis_pairs or []):
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        sigs = row.get("hsis_signals") or []
        if isinstance(sigs, str):
            sigs = [s.strip() for s in sigs.split(",") if s.strip()]
        elif isinstance(sigs, list):
            sigs = [str(s).strip() for s in sigs if str(s).strip()]
        else:
            sigs = []
        existing = hsis_lookup.get(rid, [])
        for s in sigs:
            if s not in existing:
                existing.append(s)
        hsis_lookup[rid] = existing

    # 실 설계 컴포넌트만(인터페이스 함수 제외) — 행 sds_components 표시/집계용. sds_lookup(전체)은
    # 브리지(sds_func_to_reqs)·ASIL 롤업에 그대로 쓰여 동작 불변. design_component_ids 부재(구버전
    # sds_pairs/클라이언트) 시 component_ids로 폴백 → 회귀 없음(분리 전과 동일 동작).
    sds_comp_lookup: Dict[str, List[str]] = {}
    for row in (sds_pairs or []):
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        dcomps = row.get("design_component_ids")
        if dcomps is None:
            dcomps = row.get("component_ids") or []
        if isinstance(dcomps, str):
            dcomps = [c.strip() for c in dcomps.split(",") if c.strip()]
        elif isinstance(dcomps, list):
            dcomps = [str(c).strip() for c in dcomps if str(c).strip()]
        else:
            dcomps = []
        dexisting = sds_comp_lookup.get(rid, [])
        for c in dcomps:
            if c not in dexisting:
                dexisting.append(c)
        sds_comp_lookup[rid] = dexisting

    # canonical 접기로 sds_components 에서 사라진 **원 키**. 차집합(sds_functions)에서 함께
    # 빼지 않으면 `swcom_14`·`door control` 이 둘 다 인터페이스 함수로 이중 계상된다.
    # 구 응답·구 캐시엔 이 필드가 없다 → 빈 set → 접기 이전 동작 그대로(회귀 없음).
    sds_folded_lookup: Dict[str, List[str]] = {}
    for row in (sds_pairs or []):
        if not isinstance(row, dict):
            continue
        rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        if not rid:
            continue
        fcomps = row.get("folded_component_ids") or []
        if isinstance(fcomps, str):
            fcomps = [c.strip() for c in fcomps.split(",") if c.strip()]
        elif isinstance(fcomps, list):
            fcomps = [str(c).strip() for c in fcomps if str(c).strip()]
        else:
            fcomps = []
        fexisting = sds_folded_lookup.get(rid, [])
        for c in fcomps:
            if c not in fexisting:
                fexisting.append(c)
        sds_folded_lookup[rid] = fexisting

    # ── Test rows: merge STS/SUTS/VectorCAST + SITS ──
    all_test_rows = list(vcast_rows or [])
    for row in (sits_rows or []):
        if isinstance(row, dict):
            all_test_rows.append(row)

    # Build function→requirement reverse mapping from UDS mapping_pairs
    func_to_reqs: Dict[str, List[str]] = {}
    for mp in mapping_pairs:
        if not isinstance(mp, dict):
            continue
        mp_rid = _normalize_req_id(str(mp.get("requirement_id") or ""))
        if not mp_rid:
            continue
        srcs = mp.get("source_ids") or []
        if isinstance(srcs, str):
            srcs = [s.strip() for s in srcs.split(",") if s.strip()]
        for fn in srcs:
            fn_lower = str(fn).strip().lower()
            if fn_lower and fn_lower not in func_to_reqs:
                func_to_reqs[fn_lower] = []
            if fn_lower and mp_rid not in func_to_reqs[fn_lower]:
                func_to_reqs[fn_lower].append(mp_rid)

    # Build SDS component→requirement reverse mapping
    # SITS TCs reference SwCom_XX (SDS components), not SwTR_XXXX directly.
    # ISO 26262 T5: SRS → SDS(SwCom) → SITS TC
    comp_to_reqs: Dict[str, List[str]] = {}
    for rid, comps in sds_lookup.items():
        for comp in comps:
            comp_norm = _normalize_req_id(comp)
            if comp_norm not in comp_to_reqs:
                comp_to_reqs[comp_norm] = []
            if rid not in comp_to_reqs[comp_norm]:
                comp_to_reqs[comp_norm].append(rid)

    req_id_set = set(req_ids)

    # ── SDS 함수명 bridge (사용자 결정: "SDS로 bridge") ──
    # UDS는 함수를 설계레벨 ID(SwSTR 등)에, SUTS/SITS는 단위/통합 ID에 추적해서 SRS
    # 요구사항(SwTR/SwEI 등)과 직접 안 맞는다. 그런데 SDS의 component_ids에는 SwCom_XX
    # 뿐 아니라 함수명(g_sysoptionctrl 등)이 들어 있어 "SRS요구사항→함수명"을 제공한다.
    # 이를 역으로 "함수명(lower)→[SRS요구사항]"으로 만들어 UDS 함수·테스트 unit을 SRS
    # 행에 연결한다. (component description/SwCom 키도 들어가지만 함수명 키만 실제 매칭)
    # 키는 _sds_comp_key로 정규화 — 추출 노이즈('( void'/'[10]'/표아티팩트)가 실제 함수의
    # SRS 추적을 끊는 것을 방지(라운드 109 fix). 매트릭스에 있는 SRS 요구사항만 시드.
    sds_func_to_reqs: Dict[str, List[str]] = {}
    # 전체 SDS 멤버십(req_id_set 미필터) — 미추적 함수가 '설계엔 명세됐는지'(SDS 닿는지)
    # 판별용. matrix 밖 req(SwFn/SwST 등)에만 귀속한 함수까지 포함해 'SDS 연동 but SRS 미추적'
    # 신호를 줄 수 있다. 단 정확매칭이라 거짓양성 없음(fuzzy 미사용).
    sds_all_func_to_reqs: Dict[str, List[str]] = {}
    # 모든 SDS comp 키 선스캔 — 반환형 접두사 alias가 별도 SDS 키와 충돌하면(서로 다른 함수
    # 오인 위험) alias를 생략하기 위함(라운드111). 예: u16g_drvin_motorspeed의 base
    # g_drvin_motorspeed가 이미 별도 SDS 키면 alias 추가하지 않고 기존 정확매칭에 맡긴다.
    _all_sds_keys = {
        k for comps in sds_lookup.values() for c in comps if (k := _sds_comp_key(c))
    }
    # 반환형 접두사 alias 안전 집합(라운드111): base가 ① 기존 SDS 키가 아니고 ② 단 하나의
    # 접두사형에서만 파생될 때만 허용. 2+ 접두사형(예: u8g_doorctrl_slipchkspd·
    # s8g_doorctrl_slipchkspd — unsigned8/signed8 반환형이 다른 별개 함수일 수 있음)이 같은
    # base로 모이면, 그 base로 들어오는 테스트 함수에 서로 다른 함수의 req가 union돼 거짓연결될
    # 위험이 있으므로 alias를 만들지 않는다(over-trace > under-trace 위험, 충돌 보수 처리).
    _prefixed_base_count: Dict[str, int] = {}
    for _k in _all_sds_keys:
        _b = _strip_ret_type_prefix(_k)
        if _b != _k:
            _prefixed_base_count[_b] = _prefixed_base_count.get(_b, 0) + 1
    _alias_safe = {b for b, c in _prefixed_base_count.items() if c == 1 and b not in _all_sds_keys}
    # ── 원시 SDS 이름 인덱스(대칭 정규화) — 라운드113 ──────────────────────────────
    # ⚠ 이건 '추적'이 아니라 '이름 존재' 인덱스다. 바로 위 alias 보정은 SDS 키 쪽에만
    # _strip_ret_type_prefix 를 걸고, 접두사형이 2개 이상 충돌하면 alias 를 아예 포기한다
    # (_alias_safe — over-trace 보수 처리). 그래서 표기 규약이 어긋난 함수는 브리지 두 맵
    # (sds_func_to_reqs / sds_all_func_to_reqs)이 **동시에** 미스해 'SDS 미명세'로 거짓 보고된다.
    # 두 맵은 같은 루프·같은 키·같은 소스에서 in_matrix 게이트 하나만 다르므로 한쪽이 놓치면
    # 다른 쪽도 놓친다(unmapped_sds_linked 가 구조적으로 0 인 이유 — 집계부 주석 참조).
    # 여기서 **양변에 동일 정규화**를 걸어 그 실패 모드만 별도 채널로 노출한다.
    # 절대 sds_reqs/traced/covered/in_uds 에 먹이지 않는다(순수 보고용).
    _sds_name_index: Dict[str, List[str]] = {}   # strip(key) → [원 SDS 키...]
    _sds_core_index: Dict[str, List[str]] = {}   # core(key)  → [원 SDS 키...] (저장클래스까지 제거)
    # ⚠ sorted() 필수 — _all_sds_keys 는 set 이고 CPython 문자열 해시는 프로세스마다
    # randomize 되므로, 정렬 없이 순회하면 아래 리스트 순서가 재기동마다 바뀐다. 그러면
    # sds_name_hits 의 표시 캡이 '임의의 N개'를 조용히 고르게 되어 CSV 감사 증빙과
    # localStorage 영속 매트릭스가 같은 입력에도 달라진다(베이스라인 diff에 유령 변경).
    for _k in sorted(_all_sds_keys):
        _sds_name_index.setdefault(_strip_ret_type_prefix(_k), []).append(_k)
        _sds_core_index.setdefault(_LAYER_CORE_PREFIX_RE.sub("", _k), []).append(_k)
    for rid_srs, comps in sds_lookup.items():
        in_matrix = rid_srs in req_id_set
        for comp in comps:
            key = _sds_comp_key(comp)
            if not key:
                continue
            # 반환형 헝가리안 접두사 불일치 보정(라운드111): SDS 'u16s_X' ↔ 테스트 's_X'.
            # exact 키는 보존(양쪽 prefixed 매칭 불변), 안전한 base만 alias 추가 등록.
            keys = [key]
            alias = _strip_ret_type_prefix(key)
            if alias != key and alias in _alias_safe:
                keys.append(alias)
            for kk in keys:
                all_lst = sds_all_func_to_reqs.setdefault(kk, [])
                if rid_srs not in all_lst:
                    all_lst.append(rid_srs)
                if in_matrix:
                    lst = sds_func_to_reqs.setdefault(kk, [])
                    if rid_srs not in lst:
                        lst.append(rid_srs)

    # UDS 함수 전체 (lower→원형 display) — source_ids bridge용
    uds_all_funcs: Dict[str, str] = {}
    for mp in mapping_pairs:
        if not isinstance(mp, dict):
            continue
        srcs = mp.get("source_ids") or []
        if isinstance(srcs, str):
            srcs = [s.strip() for s in srcs.split(",") if s.strip()]
        for fn in (srcs or []):
            f = str(fn).strip()
            if f:
                uds_all_funcs.setdefault(f.lower(), f)
    # extract-mapping이 전달한 전체 UDS 함수 인벤토리(설계 req 참조 없는 함수 포함)로 보강.
    # mapping_pairs는 설계 req를 참조하는 함수(~5%)만 담으므로, 이것 없이는 SDS→UDS bridge가
    # 대다수 함수를 못 찾아 "UDS 함수" 컬럼이 비게 된다(54/1005 누락 케이스).
    for fn in (uds_function_ids or []):
        f = str(fn).strip()
        if f:
            uds_all_funcs.setdefault(f.lower(), f)

    # 설계-ID bridge(SRS→SDS→UDS): UDS 함수의 Related ID 설계ID(SwFn/SwSTR/SwST/SwTK)를
    # SDS의 '설계ID→SRS요구' 매핑으로 이어 UDS 밴드에 부착한다. comp_to_reqs(SwCom·SITS
    # 전용)의 형제 — 여기선 SwCom 제외(loose, reach 불변·fan-out만 악화). map_lookup은
    # 설계ID(정규화)→[함수명+자기 SwUFn ID], sds_lookup은 SRS요구→[component_ids(설계ID·함수명)].
    design_to_reqs: Dict[str, List[str]] = {}          # 설계ID(정규화) → [SRS 요구], SwCom 제외
    for rid_srs, comps in sds_lookup.items():
        if rid_srs not in req_id_set:                  # 매트릭스 SRS 요구에만 시드(설계→설계 노이즈 차단)
            continue
        for comp in comps:
            dkey = _normalize_req_id(comp)
            if not _DESIGN_ID_BRIDGE_RE.match(dkey):    # SwFn/SwSTR/SwST/SwTK만; SwCom/SwUFn/함수명 배제
                continue
            _dlst = design_to_reqs.setdefault(dkey, [])
            if rid_srs not in _dlst:
                _dlst.append(rid_srs)
    design_bridge_funcs: Dict[str, List[str]] = {}     # SRS 요구 → [UDS 함수 display]
    # 역방향 맵: _sds_comp_key(함수명) → [SRS 요구] via 설계ID. SUTS/VectorCAST test-row가
    # 시험 함수의 UDS 설계ID로 SRS에 닿게 한다(UDS 밴드와 동일 체인의 test-arm 확장 —
    # 시험이 함수 F를 검증, F가 SwFn_05를 구현, SwFn_05가 요구 R을 실현 → 시험이 R 검증).
    func_design_reqs: Dict[str, List[str]] = {}
    for _did, _fns in map_lookup.items():
        _reqs = design_to_reqs.get(_did)
        if not _reqs:
            continue
        for _fn in _fns:
            _fs = str(_fn).strip()
            if _SWUFN_RE.match(_fs):                     # 함수 자기 SwUFn ID 값 제외 — 밴드엔 함수명만
                continue
            if _fs.lower() in _UDS_FUNC_JUNK:            # junk 필드 라벨 제외(C1 방어심층)
                continue
            _fdisp = uds_all_funcs.get(_fs.lower(), _fs)
            _fkey = _sds_comp_key(_fs)                   # SUTS unit/VCAST resolved 함수명 조회 키공간
            for _rid in _reqs:
                _blst = design_bridge_funcs.setdefault(_rid, [])
                if _fdisp not in _blst:
                    _blst.append(_fdisp)
                if _fkey:
                    _flst = func_design_reqs.setdefault(_fkey, [])
                    if _rid not in _flst:
                        _flst.append(_rid)

    # SITS/VectorCAST 2-hop bridge용: SUTS가 제공하는 SwUFn(단위함수 ID) → 함수명 맵.
    # SITS/vcast는 testcase·subprogram에 SwUFn ID를 박아두지만 함수명/SRS ID가 없다.
    # SUTS의 SwUFn↔함수명(unit)으로 함수명을 얻은 뒤 SDS 함수명 bridge로 SRS에 연결.
    # 한 SwUFn이 여러 함수명에 대응(리팩터·분할)할 수 있어 List로 누적(reviewer WARNING:
    # first-wins setdefault는 대체 함수명을 버려 silent under-trace — 안전상 더 위험한 방향).
    swufn_to_func: Dict[str, List[str]] = {}
    for row in all_test_rows:
        if isinstance(row, dict) and row.get("source") == "SUTS":
            srid = _normalize_req_id(str(row.get("requirement_id") or ""))
            sfn = str(row.get("unit") or "").strip().lower()
            if srid and sfn:
                _fns = swufn_to_func.setdefault(srid, [])
                if sfn not in _fns:
                    _fns.append(sfn)

    # Enrich SUTS/SITS rows with reverse mappings
    enriched_rows: List[Dict[str, Any]] = []
    # VectorCAST 추적 가시성(reviewer WARNING): 입력 vcast 행 중 SRS 요구사항에
    # bridge된 행 수를 집계해, "의도적 미추적(부트로더/ISR)"과 "bridge 파손"을
    # 구분할 신호를 summary로 노출한다. 미추적 행은 매트릭스에서 빠지므로 카운트만.
    vcast_input_rows = 0
    vcast_traced_rows = 0
    # §G: bridge(seen_vc)로 traced된 subprogram 집합 — 같은 subprogram이 다른 행에서 bridge
    # 실패해 미추적 목록에 잔류한 것을 loop 후 일괄 revoke(행 순서 무관). never-revoke 이중집계 해소.
    _traced_subs: Set[str] = set()
    # 미추적(SRS 미연결) VectorCAST subprogram 목록 — 역방향 추적성 공백 가시화.
    # 시험은 했으나 이 SRS 요구사항에 안 닿는 함수(일부 보안/안전 관련이라 의미 있음).
    # 트리 뷰의 'SRS 미추적 시험 포함' 토글이 이 목록을 의미 3버킷으로 묶어 보여준다.
    # 카운트(vcast_*)는 행 기준 그대로 두고, 목록만 distinct subprogram으로 dedup한다.
    unmapped_vcast: List[Dict[str, Any]] = []
    _unmapped_idx: Dict[str, int] = {}   # sub_lower → unmapped_vcast 인덱스 (FAIL 우선 머지용, W2)
    for row in all_test_rows:
        if not isinstance(row, dict):
            continue
        orig_rid = _normalize_req_id(str(row.get("requirement_id") or ""))
        source = row.get("source", "")
        unit = str(row.get("unit") or "").strip().lower()
        subprogram = str(row.get("subprogram") or "").strip()

        # Skip rows with no useful data
        # (VectorCAST 행은 requirement_id/unit 없이 subprogram만 들고 오므로 포함)
        if not orig_rid and not unit and not subprogram:
            continue

        # Keep original row only if its requirement_id is a valid matrix requirement
        if orig_rid in req_id_set:
            enriched_rows.append({**row, "trace_type": "direct"})
        # else: non-matching ID (e.g. SwCom_XX) ��� skip original, only add via reverse mapping

        # SDS component reverse mapping: SwCom_XX → [SwTR_XXXX, ...]
        if orig_rid and orig_rid not in req_id_set and source in ("SITS",):
            comp_mapped = comp_to_reqs.get(orig_rid, [])
            for mrid in comp_mapped:
                enriched_rows.append({**row, "requirement_id": mrid, "trace_type": "indirect"})

        # Function name reverse mapping: unit → [SwTR_XXXX, ...]
        # UDS 기반(func_to_reqs)은 SwSTR 등 설계레벨이라 SRS 행과 안 맞을 수 있어,
        # SDS 함수명 bridge(sds_func_to_reqs, → SRS 요구사항)를 함께 사용한다.
        if unit and source in ("SUTS", "SITS"):
            mapped_rids = list(func_to_reqs.get(unit, []))
            # SDS 키는 _sds_comp_key로 정규화돼 있으므로 조회 키도 동일 정규화해야 한다
            # (라운드112): 선행 '_'('_entrypoint') 등으로 raw 키가 어긋나면 정당한 SRS 추적이
            # silent 누락된다. func_to_reqs는 raw lowercase 키 규약이라 그대로 둔다.
            for r in sds_func_to_reqs.get(_sds_comp_key(unit), []):
                if r not in mapped_rids:
                    mapped_rids.append(r)
            # 설계-ID bridge: 시험 함수(unit)의 UDS Related ID 설계ID로도 SRS에 연결.
            # 함수명 bridge가 SDS에 이름 없는 함수를 놓치던 것을 보완(SUTS 43→64).
            for r in func_design_reqs.get(_sds_comp_key(unit), []):
                if r not in mapped_rids:
                    mapped_rids.append(r)
            for mrid in mapped_rids:
                # 유효한 SRS 요구사항(req_id_set)으로만 간접 추적 추가 — SwSTR 등 노이즈 배제
                if mrid != orig_rid and mrid in req_id_set:
                    enriched_rows.append({**row, "requirement_id": mrid, "trace_type": "indirect"})

        # SITS 2-hop bridge: testcase에 박힌 SwUFn → (SUTS)함수명 → (SDS)SRS 요구사항.
        # SITS는 함수명/SRS ID 컬럼이 없어 unit bridge가 안 걸리므로 별도 처리.
        if source == "SITS":
            seen_sits: set = set()
            for swufn in _SWUFN_RE.findall(str(row.get("testcase") or "")):
                for fn in swufn_to_func.get(_normalize_req_id(swufn), []):
                    for mrid in sds_func_to_reqs.get(_sds_comp_key(fn), []):
                        if mrid in req_id_set and mrid != orig_rid and mrid not in seen_sits:
                            seen_sits.add(mrid)
                            enriched_rows.append({**row, "requirement_id": mrid, "trace_type": "indirect"})

        # VectorCAST 함수기반 추적: vcast 행은 subprogram(거의 SwUFn ID)만 들고 온다.
        # UDS 매핑은 설계레벨(SwSTR)이라 SRS에 직접 안 닿으므로, SUTS/SITS와 동일한
        # 함수명→SRS bridge로 간접 연결한다. 두 경로 모두 시도:
        #   (1) subprogram이 함수명이면 SDS 함수명 bridge 직접 매칭
        #   (2) subprogram/testcase의 SwUFn → (SUTS)함수명 → (SDS)SRS 2-hop
        # (vcast subprogram의 ~98%가 SwUFn ID라 (2)가 주 경로. 매칭 안 되는 행은
        #  부트로더/ISR 등 SRS 추적 대상이 아니므로 자연히 매트릭스에서 제외된다.)
        if source == "VectorCAST" and subprogram:
            vcast_input_rows += 1
            seen_vc: set = set()
            sub_lower = subprogram.lower()
            # 함수명 bridge(sds_func_to_reqs) + 설계-ID bridge(func_design_reqs) 병행 —
            # 시험 함수의 UDS 설계ID로도 SRS에 연결(VectorCAST 43→64). seen_vc가 중복 차단.
            for _map in (sds_func_to_reqs, func_design_reqs):
                for mrid in _map.get(_sds_comp_key(sub_lower), []):
                    if mrid in req_id_set and mrid not in seen_vc:
                        seen_vc.add(mrid)
                        enriched_rows.append({**row, "requirement_id": mrid, "trace_type": "indirect"})
            hay = subprogram + " " + str(row.get("testcase") or "")
            for swufn in _SWUFN_RE.findall(hay):
                for fn in swufn_to_func.get(_normalize_req_id(swufn), []):
                    for _map in (sds_func_to_reqs, func_design_reqs):
                        for mrid in _map.get(_sds_comp_key(fn), []):
                            if mrid in req_id_set and mrid not in seen_vc:
                                seen_vc.add(mrid)
                                enriched_rows.append({**row, "requirement_id": mrid, "trace_type": "indirect"})
            # bridge(seen_vc) 또는 direct requirement_id(orig_rid, 2367서 이미 trace_type:"direct"
            # 로 traced)로 SRS에 연결된 행은 traced로 집계한다. direct rid 행이 subprogram bridge에
            # 실패해도 여기서 흡수돼 untraced(unmapped_vcast) 이중집계를 막는다(§F).
            if seen_vc or orig_rid in req_id_set:
                vcast_traced_rows += 1
                if seen_vc:
                    _traced_subs.add(sub_lower)   # §G: bridge traced subprogram 수집(revoke용)
            else:
                # SRS 미연결 — 역방향 추적 공백 목록에 distinct subprogram만 수집.
                # SUTS bridge(swufn_to_func)로 함수명을 해석해 의미 분류한다:
                #   isr        : 부트로더/ISR/핸들러 등 추적 대상 아님이 당연한 인프라
                #   suts_tested: SUTS 단위시험이 존재(함수명 해석됨) — '시험했으나 미명세', 검토 가치 ↑
                #   vcast_only : SUTS 참조도 없음(VectorCAST 단독 커버리지)
                res_str = str(row.get("result") or "")
                if sub_lower not in _unmapped_idx:
                    resolved: List[str] = []
                    for swufn in _SWUFN_RE.findall(hay):
                        for fn in swufn_to_func.get(_normalize_req_id(swufn), []):
                            if fn and fn not in resolved:
                                resolved.append(fn)
                    # 분류 우선순위: SUTS 단위시험 존재(resolved)는 강한 '시험했으나 미명세'
                    # 신호 → ISR 이름 휴리스틱보다 우선한다(reviewer WARNING). 이름만 ISR이고
                    # 단위시험 없는 것만 isr(인프라)로 본다. 단, 이름이 ISR 패턴이어도 안전/진단
                    # 토큰(Safety_Fault_Handler 등)을 가지면 isr로 침묵 강등하지 않는다(재검증 W4).
                    if resolved:
                        category = "suts_tested"
                    elif _ISR_RE.search(subprogram) and not _SAFETY_TOKEN_RE.search(subprogram):
                        category = "isr"
                    else:
                        category = "vcast_only"
                    # SDS(설계) 멤버십 — subprogram·해석된 함수명이 SDS 컴포넌트로 명세돼 있나.
                    # SRS엔 미추적이라도 '설계엔 닿는'(SDS 연동) 함수면 매트릭스 밖 req를 노출.
                    # 정확매칭(정규화 키)이므로 거짓양성 없음. 비면 프론트는 'SDS 미명세'로 표기
                    # → 'SRS·SDS 모두 미명세' 추적성 공백을 정직히 가시화(라운드 109).
                    # 후보를 아래 이름 대조와 **공유**한다 — hit ⊇ linked 단조성(항등식
                    # variant == hit − linked)을 코드 수준에서 보장하기 위함. _uds_cands 는
                    # SwUFn ID 를 제외해 후보가 좁으므로 여기 쓰면 항등식이 음수가 될 수 있다.
                    _sds_cands = [sub_lower, *resolved]
                    sds_reqs: List[str] = []
                    for cand in _sds_cands:
                        for r in sds_all_func_to_reqs.get(_sds_comp_key(cand), []):
                            if r not in sds_reqs:
                                sds_reqs.append(r)
                    # UDS(단위설계) 인벤토리 존재 여부 — 사용자 질문("SDS 미추적이어도 UDS엔
                    # 연동됐나"). SRS 역추적이 끊긴 함수라도 UDS 단위설계엔 명세돼 있으면
                    # '시험+단위설계 완료, SDS 아키텍처 roll-up만 누락'(정당한 입도차)으로,
                    # UDS에도 없으면 '진짜 설계 갭'으로 구분 가시화한다. SwUFn ID 자기-메아리
                    # (인벤토리가 ID도 포함)는 신호가 아니므로, SUTS로 해석된 실제 함수명 +
                    # subprogram이 함수명인 경우(SwUFn ID 아님)만 UDS 인벤토리와 대조한다.
                    uds_funcs: List[str] = []
                    _uds_cands = list(resolved)
                    if not _SWUFN_RE.search(subprogram):
                        _uds_cands.append(sub_lower)
                    for cand in _uds_cands:
                        disp = uds_all_funcs.get(cand) or uds_all_funcs.get(_sds_comp_key(cand))
                        if disp and disp not in uds_funcs:
                            uds_funcs.append(disp)
                    # 원시 SDS 이름 대조(라운드113) — 요구ID 브리지가 아니라 '이름이 SDS에
                    # 있는가'만 본다. 강한 신호부터 티어를 소진하고 첫 히트에서 멈춘다:
                    #   exact      : 정확 키 일치(기존 sds_reqs 와 동치 — 신규 정보 없음)
                    #   ret_prefix : 반환형 접두사만 다름(u16s_X ↔ s_X). SDS 실측 충돌률 0.5%
                    #   core       : 저장클래스까지 다름(s_X ↔ g_X). 가장 약한 신호 — 별개 함수일
                    #                수 있으므로 티어를 항상 함께 노출해 '힌트'로만 쓰게 한다
                    _sds_hits: List[str] = []
                    _sds_match = ""
                    for _tier in ("exact", "ret_prefix", "core"):
                        for cand in _sds_cands:
                            _k = _sds_comp_key(cand)
                            if not _k:
                                continue
                            if _tier == "exact":
                                _raws = [_k] if _k in _all_sds_keys else []
                            elif _tier == "ret_prefix":
                                _raws = _sds_name_index.get(_strip_ret_type_prefix(_k), [])
                            else:
                                _raws = _sds_core_index.get(_LAYER_CORE_PREFIX_RE.sub("", _k), [])
                            for _raw in _raws:
                                if _raw not in _sds_hits:
                                    _sds_hits.append(_raw)
                        if _sds_hits:
                            _sds_match = _tier
                            break
                    _unmapped_idx[sub_lower] = len(unmapped_vcast)
                    unmapped_vcast.append({
                        "subprogram": subprogram,
                        "result": res_str,
                        "testcase": str(row.get("testcase") or ""),
                        "unit": str(row.get("unit") or ""),
                        "resolved_funcs": resolved,
                        "category": category,
                        # SDS 설계에 명세된 SRS 요구사항(매트릭스 밖 포함) — 비면 'SDS 미명세'.
                        "sds_reqs": sds_reqs,
                        # ── 라운드113: 원시 SDS **이름** 대조(요구ID 브리지와 독립) ──
                        # sds_reqs 가 비었는데 여기 값이 있으면 = 'SDS는 이 함수를 명세하는데 요구ID
                        # 연결만 끊김'(표기 규약 드리프트) → 설계 공백이 아니라 브리지·명명 결함이다.
                        # ⚠ 한계: 이름 출처가 sds_lookup(=component_ids)이라 파서가 Related ID 없는
                        #   SDS 엔트리를 버린다(backend/routers/jenkins.py:4596-4598). 따라서
                        #   '히트=SDS에 있다'는 참이지만 '미히트=SDS에 없다'는 결론지을 수 없다(불완전 집합).
                        # ⚠ 표시 캡(8). 총량을 함께 실어 절단을 표면화한다 — 캡만 두고 총량을
                        #   숨기면 UI·CSV가 '전부'라고 오독한다(열 상한 침묵 손실과 같은 결함 클래스).
                        "sds_name_hits": _sds_hits[:8],
                        "sds_name_hits_total": len(_sds_hits),
                        "sds_name_match": _sds_match,          # "" | "exact" | "ret_prefix" | "core"
                        "sds_name_ambiguous": len(_sds_hits) > 1,
                        # UDS 단위설계 인벤토리에 존재하는 정규 함수명(비면 '단위설계 미명세' = 진짜 갭).
                        "uds_funcs": uds_funcs,
                        "in_uds": bool(uds_funcs),
                        # 안전/진단 토큰 보유 — 버킷(isr/vcast_only)과 무관하게 프론트에서
                        # amber로 강조해 백워드 추적성 검토 신호를 보존한다(재검증 W4 가시화).
                        # ★subprogram은 대개 SwUFn ID(안전토큰 無)라 해석된 함수명(resolved)·
                        # UDS 정규명(uds_funcs)에도 적용해야 ASIL 자가진단·가드 함수가 잡힌다
                        # (라운드111: subprogram만 검사하던 누락 — s_StackGuardCheck 등 미플래그).
                        "safety": bool(
                            _SAFETY_TOKEN_RE.search(subprogram)
                            or any(_SAFETY_TOKEN_RE.search(str(f)) for f in resolved)
                            or any(_SAFETY_TOKEN_RE.search(str(f)) for f in uds_funcs)
                        ),
                        # ISO 26262 SwDS 계층(라운드112) — 추적 공백을 정직히 분리: APP_LEAF/
                        # BSW_DRIVER/BOOT_REPROG/LIB_UTIL/TEST_ARTIFACT. 보고 hint이며 안전성은
                        # 직교(safety 플래그). 입력은 해석된 실제 함수명(_uds_cands) — SwUFn ID
                        # 자기-메아리를 피하고 도메인 토큰이 실린 진짜 이름으로 분류한다.
                        # §H: subprogram이 SwUFn ID인데 함수명 해석(resolved)이 없으면 UDS 대조
                        # 입력(_uds_cands)이 비어 _classify가 APP_LEAF로 기본화되고, 진짜 설계갭
                        # (unmapped_app_design_gap = APP_LEAF ∩ not in_uds)을 부풀린다. 실제로는
                        # '분류 불가'(SwUFn↔함수명 미해결)이므로 UNRESOLVED 버킷으로 빼 오염을 막는다.
                        "layer": (
                            "UNRESOLVED"
                            if (_SWUFN_RE.search(subprogram) and not resolved)
                            else _classify_unmapped_layer(_uds_cands)
                        ),
                    })
                else:
                    # worst-case 집계: 동일 subprogram의 후속 행이 FAIL이면 기존 항목 result를
                    # FAIL로 격상(재검증 W2: PASS 선행 시 FAIL이 silent 손실돼 트리의 미추적
                    # FAIL 카운트가 과소표시되는 것을 방지). 프론트 failTotal과 동일 판정.
                    _existing = unmapped_vcast[_unmapped_idx[sub_lower]]
                    if _RESULT_FAIL_RE.match(res_str) and not _RESULT_FAIL_RE.match(str(_existing["result"])):
                        _existing["result"] = res_str

    # §G: 같은 subprogram이 다른 행에서 bridge로 traced됐으면 미추적 목록에서 제거한다(never-
    # revoke 이중집계 해소). 행 순서 무관하게 loop 후 일괄 필터 — traced 행이 unmapped 행보다
    # 먼저 와도 정확. _unmapped_idx는 이후 미사용이라 인덱스 무효화 무해.
    if _traced_subs:
        unmapped_vcast = [u for u in unmapped_vcast
                          if str(u["subprogram"]).lower() not in _traced_subs]

    # 의미 버킷 우선순위로 정렬 — 잘림/상단 노출 시 신호(suts_tested)가 먼저 보이도록.
    _UNMAPPED_ORDER = {"suts_tested": 0, "isr": 1, "vcast_only": 2}
    unmapped_vcast.sort(key=lambda x: (_UNMAPPED_ORDER.get(x["category"], 3), x["subprogram"].lower()))

    vcast_map = _normalize_vcast_rows(enriched_rows)

    matrix: List[Dict[str, Any]] = []
    mapped_source_count = 0
    mapped_sds_count = 0
    mapped_hsis_count = 0
    mapped_test_count = 0
    total_pass = 0
    total_fail = 0
    total_tests = 0
    source_stats: Dict[str, int] = {}  # source → count of mappings

    for rid in req_ids:
        tests = vcast_map.get(rid, [])
        test_ids = [t.get("testcase") for t in tests if t.get("testcase")]
        src_direct = list(map_lookup.get(rid, []))  # 진짜 UDS RelatedID(직접 추적, confidence=direct)
        src_list = list(src_direct)
        # SDS 함수명 bridge: 이 SRS 요구사항에 SDS가 귀속한 함수 중 UDS에 존재하는 것을
        # UDS 추적(source_ids)으로 채운다 (UDS가 SwSTR로 추적해 직접 안 붙던 문제 해소).
        # 조회 키는 _sds_comp_key로 정규화 — SDS 키 공간과 일치(라운드112 W1: 다른 4개 bridge
        # 조회 사이트와 동일하게, 선행 '_'('_entrypoint') 등 정규화 차이로 끊기는 것 방지).
        # 여기서 append되는 항목은 SDS 경유 '추정'이라 link_table에서 indirect로 라벨(정직화).
        for flower, fdisp in uds_all_funcs.items():
            if rid in sds_func_to_reqs.get(_sds_comp_key(flower), []) and fdisp not in src_list:
                src_list.append(fdisp)
        # 설계-ID bridge(2b) 주입: name-bridge가 놓친 요구사항의 UDS 함수를 SwFn/SwSTR/
        # SwST/SwTK 경유로 부착. SDS 경유 '추정'이라 src_direct(직접)엔 안 넣는다 →
        # source_ids_direct 비어 유지 → link_table에서 indirect 라벨(over-trace 안전판).
        for _bf in design_bridge_funcs.get(rid, []):
            if _bf not in src_list:
                src_list.append(_bf)
        sds_list = sds_lookup.get(rid, [])              # 전체(함수 포함) — ASIL 롤업·내부용
        sds_comp_list = sds_comp_lookup.get(rid, [])    # 실 설계 컴포넌트만 — 표시/집계(추적 정화)
        _scomp_set = set(sds_comp_list)
        _folded_set = set(sds_folded_lookup.get(rid, []))   # canonical 로 접힌 원 키
        sds_func_list = [c for c in sds_list
                         if c not in _scomp_set and c not in _folded_set]  # 인터페이스 함수(별도)
        hsis_list = hsis_lookup.get(rid, [])            # HSIS 인터페이스 신호(시스템 레벨 design-arm)
        # ASIL 결합(P5) — 요구사항 ASIL = 연결된 SDS 컴포넌트·UDS 함수의 최고 ASIL.
        # 컴포넌트/함수명(lower)으로 comp_asil_map 조회. 맵 없거나 매칭 0이면 ''(graceful).
        # ASIL 키는 전체(sds_list) 사용 — 함수가 부모 ASIL 상속이라 max 불변(분리 전과 동일).
        row_asil = ""
        if comp_asil_map:
            _asil_keys = [str(c).strip().lower() for c in sds_list]
            _asil_keys += [str(s).strip().lower() for s in src_list]
            row_asil = _asil_max_of([comp_asil_map[k] for k in _asil_keys if k in comp_asil_map])
        if src_list:
            mapped_source_count += 1
        if sds_comp_list:
            mapped_sds_count += 1
        if hsis_list:
            mapped_hsis_count += 1
        if test_ids:
            mapped_test_count += 1

        # Per-row test result stats
        row_pass = sum(1 for t in tests if t.get("result", "").lower() in ("pass", "passed", "true", "1"))
        row_fail = sum(1 for t in tests if t.get("result", "").lower() in ("fail", "failed", "false", "0"))
        total_pass += row_pass
        total_fail += row_fail
        total_tests += len(tests)

        # Track data sources
        for t in tests:
            src = t.get("source") or "unknown"
            source_stats[src] = source_stats.get(src, 0) + 1

        # ISO 26262 추적성: 문서 계층별 분류 + 직접/간접 구분
        sts_tests = [t for t in tests if t.get("source") == "STS"]
        suts_tests = [t for t in tests if t.get("source") == "SUTS"]
        sits_tests = [t for t in tests if t.get("source") == "SITS"]
        # 시스템 레벨 시험(SyTS/SyITS) — 시스템 레벨 검증 증거. 결정1 재정의로 SW covered에는 미포함
        # (참고·V-model 시스템 쌍·밴드 카운트로만 노출). 아래 syts_tests/syits_tests 필드로 분리 보존.
        syts_tests = [t for t in tests if t.get("source") == "SyTS"]
        syits_tests = [t for t in tests if t.get("source") == "SyITS"]
        # VectorCAST 실행추적(fuzzy, indirect) — V-model 통계에 별도 노출(reviewer INFO:
        # vcast 기여가 *_indirect 카운트에 안 잡혀 audit 불가하던 문제 해소).
        vcast_tests_row = [t for t in tests if t.get("source") == "VectorCAST"]

        # Direct vs Indirect trace counts
        sts_direct = [t for t in sts_tests if t.get("trace_type") != "indirect"]
        suts_direct = [t for t in suts_tests if t.get("trace_type") != "indirect"]
        suts_indirect = [t for t in suts_tests if t.get("trace_type") == "indirect"]
        sits_direct = [t for t in sits_tests if t.get("trace_type") != "indirect"]
        sits_indirect = [t for t in sits_tests if t.get("trace_type") == "indirect"]

        # Derive confidence with ISO 26262 semantics
        has_direct = len(sts_direct) > 0 or len(suts_direct) > 0 or len(sits_direct) > 0
        has_indirect = len(suts_indirect) > 0 or len(sits_indirect) > 0
        has_fuzzy = any(t.get("confidence") == "fuzzy" or t.get("source") == "VectorCAST" for t in tests)
        row_confidence = None
        if tests:
            if has_direct and not has_fuzzy:
                row_confidence = "direct"
            elif has_indirect and not has_direct and not has_fuzzy:
                row_confidence = "indirect"
            elif has_fuzzy and not has_direct and not has_indirect:
                row_confidence = "fuzzy"
            else:
                row_confidence = "mixed"

        matrix.append(
            {
                "requirement_id": norm_to_raw.get(rid, rid),
                # 요구사항 표시명(제목) — 프론트 표/트리에서 ID 옆에 노출(라운드110).
                "requirement_name": name_map.get(rid, ""),
                # ASIL(P5) — 연결 설계요소 최고 등급(QM<A<B<C<D). 데이터 없으면 ''(graceful).
                "asil": row_asil,
                # T1: SRS→SDS (아키텍처 추적) — 실 설계 컴포넌트만(인터페이스 함수 fan-out 제외, 추적 정화)
                "sds_components": sds_comp_list,
                # 인터페이스 함수(투명성·드릴다운). SDS 밴드 집계엔 미포함, SUTS/VCAST 브리지엔 사용됨.
                "sds_functions": sds_func_list,
                # 인터페이스 밴드(시스템 레벨 design-arm) — HSIS 신호(HSI_xx/SW변수). SwEI 등 인터페이스 요구의 realization.
                "hsis_signals": hsis_list,
                # 상위 추적 — 이 요구가 유도된 상위 시스템 요구(SyRS: SyTR/SyTSR/SyEI…). SR→SyRS→SwRS 체인.
                "syrs_parents": syrs_map.get(rid, []),
                # T2: SDS→UDS (상세 설계 추적)
                "source_ids": src_list,
                # 직접 UDS RelatedID만(나머지 source_ids는 SDS 브리지 추정) — link_table confidence 정직화용
                "source_ids_direct": src_direct,
                # T3: SRS→STS (SW 테스트 추적)
                "sts_tests": sts_tests,
                "sts_count": len(sts_tests),
                "sts_direct": len(sts_direct),
                # T4: UDS→SUTS (단위 테스트 추적)
                "suts_tests": suts_tests,
                "suts_count": len(suts_tests),
                "suts_direct": len(suts_direct),
                "suts_indirect": len(suts_indirect),
                # T5: SDS→SITS (통합 테스트 추적)
                "sits_tests": sits_tests,
                "sits_count": len(sits_tests),
                "sits_direct": len(sits_direct),
                "sits_indirect": len(sits_indirect),
                # 시스템 레벨 시험(SyTS/SyITS) — 시스템 레벨 검증 증거(결정1: SW covered 미포함, 참고 보존)
                "syts_tests": syts_tests,
                "syts_count": len(syts_tests),
                "syits_tests": syits_tests,
                "syits_count": len(syits_tests),
                # VectorCAST 실행추적 (전부 indirect/fuzzy)
                "vcast_count": len(vcast_tests_row),
                # 기존 호환
                "tests": tests,
                "test_ids": test_ids,
                "test_count": len(tests),
                "pass_count": row_pass,
                "fail_count": row_fail,
                "confidence": row_confidence,
            }
        )
    # ── ID 정합성 감사(trace_integrity, additive) ──
    # hiMA exact-match가 silent하게 오인하는 클래스(정규화 충돌·SRS에 없는 대상 참조·
    # placeholder)를 명시 finding으로 표면화. 본 매트릭스 로직/기존 키 불변, 새 'integrity'만 가산.
    _ref: Dict[str, Dict[str, str]] = {"UDS": {}, "SDS": {}}
    _rel: Dict[str, List[str]] = {"UDS": [], "SDS": []}
    for _mp in mapping_pairs:
        if not isinstance(_mp, dict):
            continue
        _raw = str(_mp.get("requirement_id") or "").strip()
        if _raw:
            _ref["UDS"].setdefault(_normalize_req_id(_raw), _raw)
        _srcs = _mp.get("source_ids") or []
        if isinstance(_srcs, str):
            _srcs = [s.strip() for s in _srcs.split(",") if s.strip()]
        for _s in (_srcs or []):
            _ss = str(_s).strip()
            if _ss:
                _rel["UDS"].append(_ss)
    for _sp in (sds_pairs or []):
        if not isinstance(_sp, dict):
            continue
        _raw = str(_sp.get("requirement_id") or "").strip()
        if _raw:
            _ref["SDS"].setdefault(_normalize_req_id(_raw), _raw)
        _comps = _sp.get("component_ids") or []
        if isinstance(_comps, str):
            _comps = [c.strip() for c in _comps.split(",") if c.strip()]
        for _c in (_comps or []):
            _cc = str(_c).strip()
            if _cc:
                _rel["SDS"].append(_cc)
    _total_sds_comps, _coarse_count = annotate_sds_coarse(matrix)

    try:
        integrity = build_integrity_audit(req_ids, norm_to_raws, _ref, _rel)
    except Exception as _intg_exc:  # best-effort: 감사 실패가 매트릭스 생성을 막지 않도록
        logger.debug("Integrity audit skipped: %s", _intg_exc)
        # 폴백 stats shape를 build_integrity_audit 정상경로와 정확히 일치시킨다
        # (suspect/foreign 키 포함) — 소비자의 '?? default'/destructure 함정 방지.
        integrity = {
            "id_collisions": [], "dangling_refs": {}, "dangling_by_namespace": {},
            "dangling_layer_summary": {},
            "placeholder_ids": {},
            "stats": {"collision_count": 0, "collision_affected_raw": 0,
                      "dangling_count": 0, "dangling_suspect_count": 0,
                      "dangling_foreign_count": 0, "placeholder_count": 0, "clean": True},
        }
    return {
        "total_requirements": len(req_ids),
        "rows": matrix,
        "summary": {
            "requirement_count": len(req_ids),
            "mapped_sds_count": mapped_sds_count,
            # 추적 정화: 실 설계 컴포넌트 총수(함수 fan-out 제외) + 거친입도(>40% 연결) 요구사항 수
            "total_sds_components": _total_sds_comps,
            "sds_coarse_count": _coarse_count,
            "mapped_source_count": mapped_source_count,
            # 시스템 레벨 인터페이스(HSIS) 연결된 요구사항 수
            "mapped_hsis_count": mapped_hsis_count,
            "mapped_test_count": mapped_test_count,
            "total_tests": total_tests,
            "total_pass": total_pass,
            "total_fail": total_fail,
            "source_stats": source_stats,
            # ISO 26262 V-Model 추적성 통계
            "mapped_sts_count": sum(1 for r in matrix if r.get("sts_count")),
            "mapped_suts_count": sum(1 for r in matrix if r.get("suts_count")),
            "mapped_sits_count": sum(1 for r in matrix if r.get("sits_count")),
            # 시스템 레벨 상위(SyRS) 연결 요구사항 수 — 상위 추적 provenance(다운스트림 커버리지 분모 제외, _UPSTREAM_BANDS)
            "mapped_syrs_count": sum(1 for r in matrix if r.get("syrs_parents")),
            # 시스템 레벨 시험(SyTS/SyITS) 연결 요구사항 수
            "mapped_syts_count": sum(1 for r in matrix if r.get("syts_count")),
            "mapped_syits_count": sum(1 for r in matrix if r.get("syits_count")),
            "mapped_sts_direct": sum(1 for r in matrix if r.get("sts_direct")),
            "mapped_suts_direct": sum(1 for r in matrix if r.get("suts_direct")),
            "mapped_suts_indirect": sum(1 for r in matrix if r.get("suts_indirect")),
            "mapped_sits_direct": sum(1 for r in matrix if r.get("sits_direct")),
            "mapped_sits_indirect": sum(1 for r in matrix if r.get("sits_indirect")),
            "mapped_vcast_count": sum(1 for r in matrix if r.get("vcast_count")),
            # VectorCAST bridge 가시성 — 입력 행 중 SRS에 연결된 행 수 / 미연결(부트로더·ISR 등)
            "vcast_input_rows": vcast_input_rows,
            "vcast_traced_rows": vcast_traced_rows,
            "vcast_untraced_rows": vcast_input_rows - vcast_traced_rows,
            # 미추적 목록(distinct subprogram)의 의미 버킷별 개수 — 트리 미추적 루트 뱃지용.
            "unmapped_vcast_count": len(unmapped_vcast),
            "unmapped_suts_tested": sum(1 for u in unmapped_vcast if u["category"] == "suts_tested"),
            "unmapped_vcast_only": sum(1 for u in unmapped_vcast if u["category"] == "vcast_only"),
            "unmapped_isr": sum(1 for u in unmapped_vcast if u["category"] == "isr"),
            # 버킷과 무관하게 안전/진단 토큰을 가진 미추적 함수 수 — 프론트 amber 강조·뱃지용(W4).
            "unmapped_safety": sum(1 for u in unmapped_vcast if u.get("safety")),
            # SRS 미추적이지만 SDS 설계엔 명세된(역방향 부분추적) 함수 수 — 프론트 'SDS:<req>' 뱃지용.
            # ⚠ 이 지표는 구조적으로 거의 항상 0 이다(라운드113 정정). sds_reqs 는
            #   sds_all_func_to_reqs 를, traced 판정은 sds_func_to_reqs 를 조회하는데 두 맵은 같은
            #   루프·같은 키·같은 소스에서 in_matrix 게이트 하나만 다르게 만들어진다. 따라서 값이
            #   채워지려면 '그 함수의 SDS 요구가 전부 매트릭스 밖'이라는 극단 조건이 필요하고,
            #   키 정규화가 어긋나면 두 맵이 **동시에** 미스한다. 0 은 "설계가 이 함수들을 명세
            #   안 함"을 **의미하지 않는다** — 코드가 보장하는 건 '매트릭스 밖 요구에만 귀속된
            #   미추적 함수가 0' 뿐이다. 실제 SDS 명세 여부는 아래 unmapped_sds_name_hit 을 볼 것.
            #   (구 주석 "설계가 이 함수들을 명세 안 함"은 과한 결론이라 정정. 라운드 109→113)
            "unmapped_sds_linked": sum(1 for u in unmapped_vcast if u.get("sds_reqs")),
            # 미추적 함수 중 **SDS 이름 집합(정규화 키)에 이름이 존재**하는 수(요구ID 연결 무관, 라운드113).
            # 이름 집합이 Related ID 있는 엔트리만 담으므로 하한선(lower bound)이다 — 미히트를
            # '미명세'로 읽지 말 것.
            # ⚠ SDS component_ids 에 SwUFn 같은 **시험 ID**가 섞여 있으면 exact 티어가 그 ID 자기-메아리를
            #   히트로 센다(함수명 일치가 아님). 단 그 경우 sds_reqs 도 함께 채워지므로 아래 신규 신호
            #   (variant = 이름은 있고 요구ID는 없음)에는 들어가지 않는다 — 오염은 hit 총계에만 국한.
            "unmapped_sds_name_hit": sum(1 for u in unmapped_vcast if u.get("sds_name_hits")),
            # ★신규 신호: 이름은 SDS에 있는데 요구ID 브리지는 끊긴 수 = '설계 공백'이 아니라 '브리지 결함'.
            #  항등식: unmapped_sds_name_variant == unmapped_sds_name_hit − unmapped_sds_linked
            #  (sds_reqs 비공백 ⟹ 반드시 name_hit — exact 티어가 같은 키공간을 보므로 단조. 테스트 고정)
            "unmapped_sds_name_variant": sum(
                1 for u in unmapped_vcast if u.get("sds_name_hits") and not u.get("sds_reqs")
            ),
            # SRS 미추적이지만 UDS 단위설계엔 존재하는 함수 수 — '시험+단위설계 완료, SDS
            # 아키텍처 roll-up만 누락'(정당한 입도차). KJPDS02 실데이터 661/662.
            # (unmapped_sds_linked와 동일 패턴: 캐시 trace_summary.json·감사·문서화용 집계이며,
            #  프론트 루트 뱃지는 unmapped_vcast list의 in_uds로 직접 재계산해 표시·카운트 동기 보장.)
            "unmapped_uds_linked": sum(1 for u in unmapped_vcast if u.get("in_uds")),
            # UDS에도 없는(단위설계 미명세) 미추적 함수 수 — 진짜 설계 공백(검토 우선순위 ↑).
            "unmapped_design_gap": sum(1 for u in unmapped_vcast if not u.get("in_uds")),
            # ★진짜 '실 finding' = layer축(app_leaf)이 아니라 design축(미설계)이다. app_leaf 중 대다수는
            #  단위설계(UDS)+단위시험(SUTS)까지 된 leaf helper로 부모 SwUFn 아래 roll-up되는 '정당한
            #  입도차'(unmapped_uds_linked)일 뿐. 진짜 검토 대상 = APP_LEAF ∩ 미설계(UDS에도 없음).
            #  (실측 KJPDS02_PV: app_leaf 413 중 미설계 ~3 — layer만 보면 ~100배 과대 집계.)
            "unmapped_app_design_gap": sum(
                1 for u in unmapped_vcast if u.get("layer") == "APP_LEAF" and not u.get("in_uds")
            ),
            # TEST_ARTIFACT 중 ¬uds (design_gap 하위집합) — 프론트 배너 design축 분해 전용(app_design_gap과 동형).
            # ⚠ 아래 unmapped_layer_test_artifact(전체)와 구분: 그건 6버킷 합==count 불변식용이고, 이건 design_gap
            #  정확 분해용이다. _classify는 cand[0]만 보고 TEST_ARTIFACT를 판정하나 in_uds는 _uds_cands 전체를
            #  보므로 TEST_ARTIFACT∩in_uds가 이론상 가능(예: 첫 해석명이 시그니처형이라 non-ident인데 다른 후보가
            #  UDS 매칭). 그때 전체 카운트를 design_gap에서 빼면 이중차감→음수 클램프로 범위경계 갭을 은닉하므로,
            #  ¬uds로 필터한 이 필드를 배너 차감(boundaryGap = design_gap − app_gap − unresolved − 이 값)에 쓴다.
            "unmapped_test_artifact_gap": sum(
                1 for u in unmapped_vcast if u.get("layer") == "TEST_ARTIFACT" and not u.get("in_uds")
            ),
            # ISO 26262 SwDS 계층(라운드112) — '코드 종류' 축(app/BSW/boot/lib/test). ⚠ layer는 갭 판정이
            # 아니라 분류 힌트다: app_leaf ∩ in_uds = 설계+시험된 leaf roll-up(정당한 입도차),
            # app_leaf ∩ 미설계 = 진짜 finding(위 unmapped_app_design_gap). BSW/BOOT/LIB = 정당한 범위경계.
            "unmapped_layer_app_leaf": sum(1 for u in unmapped_vcast if u.get("layer") == "APP_LEAF"),
            "unmapped_layer_bsw_driver": sum(1 for u in unmapped_vcast if u.get("layer") == "BSW_DRIVER"),
            "unmapped_layer_boot_reprog": sum(1 for u in unmapped_vcast if u.get("layer") == "BOOT_REPROG"),
            "unmapped_layer_lib_util": sum(1 for u in unmapped_vcast if u.get("layer") == "LIB_UTIL"),
            "unmapped_layer_test_artifact": sum(1 for u in unmapped_vcast if u.get("layer") == "TEST_ARTIFACT"),
            # §H: 분류 불가(SwUFn↔함수명 미해결) — layer 정합식(6버킷 합 == unmapped_vcast_count)의
            # 6번째 정식 버킷. app_design_gap(APP_LEAF ∩ 미설계)엔 미포함이라 진짜 갭 과대집계를 막는다.
            "unmapped_layer_unresolved": sum(1 for u in unmapped_vcast if u.get("layer") == "UNRESOLVED"),
        },
        # 역방향 추적성 공백 — '시험은 했으나 이 SRS에 안 닿는' VectorCAST subprogram 전체 목록.
        # 트리 뷰의 'SRS 미추적 시험 포함' 토글이 의미 3버킷으로 묶어 보여준다.
        "unmapped_vcast": unmapped_vcast,
        # ID 정합성 감사(trace_integrity, additive) — 정규화 충돌·dangling·placeholder finding.
        "integrity": integrity,
        "has_sds_mapping": any(r.get("sds_components") for r in matrix),
        "has_source_mapping": any(r.get("source_ids") for r in matrix),
        "has_tests": any(r.get("test_count") for r in matrix),
    }


def generate_uds_requirements_compare(
    items: List[Dict[str, Any]],
    source_root: str,
) -> Dict[str, Any]:
    req_ids = sorted({_normalize_req_id(str(x.get("id") or "")) for x in items if str(x.get("id") or "").strip()})
    source_ids = [_normalize_req_id(sid) for sid in _scan_source_requirement_ids(source_root)]
    source_set = set(source_ids)
    req_set = set(req_ids)
    matched = [rid for rid in req_ids if rid in source_set]
    missing = [rid for rid in req_ids if rid not in source_set]
    source_only = [rid for rid in source_ids if rid not in req_set]
    return {
        "total_requirements": len(req_ids),
        "matched": matched,
        "missing": missing,
        "source_only": source_only,
        "source_scanned": len(source_ids),
    }


def generate_uds_requirements_from_docs(texts: List[str]) -> str:
    lines: List[str] = []
    for txt in texts:
        lines.extend(_extract_requirements_from_doc(txt))
        if not lines:
            lines.extend(_extract_requirements_fallback(txt))
    seen = set()
    uniq: List[str] = []
    for ln in lines:
        if ln in seen:
            continue
        seen.add(ln)
        uniq.append(ln)
    return "\n".join(uniq).strip()
