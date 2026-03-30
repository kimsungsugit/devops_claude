"""report_gen.validation - Auto-split from report_generator.py"""
# Re-import common dependencies
import re

# Payload field name constants (canonical source: report_gen.uds_generator)
# Function-level (per-function, List[str]):
#   KEY_FN_GLOBALS = "globals_global"  — global vars used by the function
#   KEY_FN_STATICS = "globals_static"  — static vars used by the function
# Legacy: older sidecar JSONs may use bare "globals" key → fall back to it when
# reading (see _extract_payload_function_details / row.get("globals_global") or row.get("globals"))
import os
import json
import csv
import logging
import time
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set
from datetime import datetime

from report_gen.docx_builder import _iter_template_blocks
from report_gen.function_analyzer import (
    _normalize_symbol_name,
    _is_generic_description,
    _classify_description_quality,
)
from report_gen.requirements import _extract_function_info_from_docx
from report_gen.utils import _extract_call_names, _safe_dict

_logger = logging.getLogger("report_generator")


def _load_uds_payload_for_docx(docx_path: str) -> Dict[str, Any]:
    docx = Path(docx_path)
    candidates = [
        docx.with_suffix(".docx.payload.full.json"),
        docx.with_suffix(".payload.full.json"),
    ]
    for cand in candidates:
        try:
            if cand.exists():
                data = json.loads(cand.read_text(encoding="utf-8"))
                if isinstance(data, dict):
                    return data
        except Exception:
            continue
    return {}


def _valid_call_names(value: Any) -> List[str]:
    raw_items = value if isinstance(value, list) else _extract_call_names(str(value or ""))
    names: List[str] = []
    skip = {"if", "for", "while", "switch", "return", "sizeof", "case", "else"}
    for item in raw_items or []:
        text = _normalize_symbol_name(str(item or ""))
        if not text:
            continue
        if text.lower() in skip:
            continue
        if not re.match(r"^[A-Za-z_]\w*$", text):
            continue
        if text.isupper() and len(text) <= 6:
            continue
        if text in names:
            continue
        names.append(text)
    return names


def _payload_function_details_by_name(payload: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    by_name: Dict[str, Dict[str, Any]] = {}
    for key in ("function_details_by_name", "function_details"):
        details = payload.get(key) or {}
        if not isinstance(details, dict):
            continue
        for raw_name, info in details.items():
            if not isinstance(info, dict):
                continue
            name = _normalize_symbol_name(str(info.get("name") or raw_name or "")).lower()
            if name and name not in by_name:
                by_name[name] = info
    return by_name


def _has_doc_output_slot(prototype: str) -> bool:
    proto = " ".join(str(prototype or "").split())
    if not proto or "(" not in proto or ")" not in proto:
        return False
    params = proto.split("(", 1)[1].rsplit(")", 1)[0].strip()
    if not params or params.lower() == "void":
        return False
    for part in [p.strip() for p in params.split(",") if p.strip()]:
        lower = part.lower()
        if "(*" in part:
            continue
        if "*" in part or "[" in part:
            if "const" not in lower:
                return True
        if re.search(r"\b(struct|union)\b", part, re.I) and "const" not in lower:
            return True
    return False


def _has_doc_input_slot(prototype: str) -> bool:
    proto = " ".join(str(prototype or "").split())
    if not proto or "(" not in proto or ")" not in proto:
        return False
    params = proto.split("(", 1)[1].rsplit(")", 1)[0].strip()
    if not params or params.lower() == "void":
        return False
    return True

def validate_uds_docx_structure(docx_path: str) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "docx_path": docx_path,
        "ok": False,
        "table_count": 0,
        "image_count": 0,
        "swufn_heading_count": 0,
        "function_info_table_count": 0,
        "logic_row_count": 0,
        "logic_with_image_count": 0,
        "top_headers": [],
        "issues": [],
    }
    try:
        import docx  # type: ignore
    except Exception:
        result["issues"].append("python-docx is not installed")
        return result
    path = Path(docx_path)
    if not path.exists():
        result["issues"].append(f"docx not found: {docx_path}")
        return result
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        result["issues"].append(f"failed to open docx: {exc}")
        return result
    result["table_count"] = len(doc.tables)
    result["image_count"] = len(doc.inline_shapes)
    header_counter: Counter[str] = Counter()
    for para in doc.paragraphs:
        style = str(getattr(para.style, "name", "") or "")
        text = (para.text or "").strip()
        if style.startswith("Heading") and re.search(r"\bSwUFn_\d+\b", text, flags=re.I):
            result["swufn_heading_count"] += 1
    for table in doc.tables:
        if not table.rows:
            continue
        header_key = "|".join([(c.text or "").strip() for c in table.rows[0].cells])
        header_counter[header_key] += 1
        first_row = [c.text.strip() for c in table.rows[0].cells]
        if any("Function Information" in cell for cell in first_row):
            result["function_info_table_count"] += 1
            for r_idx, row in enumerate(table.rows):
                row_cells = [c.text.strip() for c in row.cells]
                if any("Logic Diagram" in c for c in row_cells):
                    result["logic_row_count"] += 1
                    target_col = min(2, max(0, len(row.cells) - 1))
                    try:
                        xml = table.cell(r_idx, target_col)._tc.xml
                    except Exception:
                        xml = ""
                    if "w:drawing" in xml or "v:imagedata" in xml:
                        result["logic_with_image_count"] += 1
    result["top_headers"] = [
        {"header": h, "count": c}
        for h, c in header_counter.most_common(10)
    ]
    if result["function_info_table_count"] != result["swufn_heading_count"]:
        result["issues"].append(
            f"SwUFn headings({result['swufn_heading_count']}) != FunctionInfo tables({result['function_info_table_count']})"
        )
    if result["logic_row_count"] != result["logic_with_image_count"]:
        result["issues"].append(
            f"Logic rows({result['logic_row_count']}) != rows with image({result['logic_with_image_count']})"
        )
    result["ok"] = len(result["issues"]) == 0
    return result


def generate_uds_validation_report(docx_path: str, out_path: str) -> str:
    report = validate_uds_docx_structure(docx_path)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    lines: List[str] = []
    lines.append("# UDS Validation Report")
    lines.append("")
    lines.append(f"- Docx: `{report.get('docx_path')}`")
    lines.append(f"- OK: `{report.get('ok')}`")
    lines.append(f"- Tables: `{report.get('table_count')}`")
    lines.append(f"- Images: `{report.get('image_count')}`")
    lines.append(f"- SwUFn headings: `{report.get('swufn_heading_count')}`")
    lines.append(f"- FunctionInfo tables: `{report.get('function_info_table_count')}`")
    lines.append(f"- Logic rows: `{report.get('logic_row_count')}`")
    lines.append(f"- Logic rows with image: `{report.get('logic_with_image_count')}`")
    lines.append("")
    issues = report.get("issues") or []
    lines.append("## Issues")
    if issues:
        for issue in issues:
            lines.append(f"- {issue}")
    else:
        lines.append("- none")
    lines.append("")
    lines.append("## Top Headers")
    for row in report.get("top_headers") or []:
        lines.append(f"- {row.get('count')}: {row.get('header')}")
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_called_calling_accuracy_report(
    docx_path: str,
    source_root: str,
    out_path: str,
    relation_mode: str = "code",
) -> str:
    relation_mode = str(relation_mode or "code").strip().lower()
    if relation_mode not in {"code", "document"}:
        relation_mode = "code"
    from report_gen.uds_generator import generate_uds_source_sections  # lazy: heavy module
    source_sections = generate_uds_source_sections(source_root)
    details_by_name = source_sections.get("function_details_by_name", {}) or {}
    fn_to_swcom: Dict[str, str] = {}
    for row in source_sections.get("function_table_rows", []) or []:
        if not isinstance(row, list) or len(row) < 4:
            continue
        swcom = str(row[0] or "").strip()
        name = _normalize_symbol_name(str(row[3] or "")).lower()
        if swcom and name:
            fn_to_swcom[name] = swcom
    def _canon_names(values: List[str]) -> Set[str]:
        out: Set[str] = set()
        for v in values:
            s = _normalize_symbol_name(str(v or ""))
            if not s:
                continue
            if re.search(r"[\s,\[\]\{\}\*]", s):
                continue
            if not re.match(r"^[A-Za-z_]\w*$", s):
                continue
            out.add(s.lower())
        return out

    exp_called: Dict[str, Set[str]] = {}
    exp_calling: Dict[str, Set[str]] = {}
    for name, info in details_by_name.items():
        if not isinstance(info, dict):
            continue
        callee = [str(x).strip() for x in (info.get("calls_list") or []) if str(x).strip()]
        exp_called[str(name).lower()] = _canon_names(callee)
    for caller, callees in exp_called.items():
        for callee in callees:
            exp_calling.setdefault(str(callee).lower(), set()).add(caller)

    try:
        import docx  # type: ignore
    except Exception as exc:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(f"# Called/Calling Accuracy Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)
    doc_map = _extract_function_info_from_docx(docx.Document(str(docx_path)))
    doc_by_name: Dict[str, Dict[str, str]] = {}
    doc_swcom_by_name: Dict[str, str] = {}
    for _, row in doc_map.items():
        fn = _normalize_symbol_name(str(row.get("name") or "")).lower()
        if fn:
            doc_by_name[fn] = row
            fid = str(row.get("id") or "").strip()
            m_sw = re.search(r"SwUFn_(\d{2})\d+", fid, flags=re.I)
            if m_sw:
                doc_swcom_by_name[fn] = f"SwCom_{m_sw.group(1)}"

    def _ratio(a: int, b: int) -> str:
        return "0.0%" if b <= 0 else f"{(a / b) * 100:.1f}%"

    total = 0
    called_match = 0
    calling_match = 0
    total_swcom01 = 0
    called_match_swcom01 = 0
    calling_match_swcom01 = 0
    mismatches: List[str] = []

    known_functions = set(details_by_name.keys()) | set(doc_by_name.keys())

    def _keep_known(names: Set[str]) -> Set[str]:
        if not known_functions:
            return set(names)
        return {n for n in names if n in known_functions}

    for name, row in doc_by_name.items():
        exp_callee = exp_called.get(name, set())
        exp_caller = exp_calling.get(name, set())
        # Keep canonical relation semantics for both modes:
        # called = callee list, calling = caller list.
        exp_called_field = exp_callee
        exp_calling_field = exp_caller
        exp_called_field = _keep_known(exp_called_field)
        exp_calling_field = _keep_known(exp_calling_field)
        act_called = _keep_known(_canon_names(_extract_call_names(str(row.get("called") or ""))))
        act_calling = _keep_known(_canon_names(_extract_call_names(str(row.get("calling") or ""))))
        total += 1
        ok_called = exp_called_field == act_called
        ok_calling = exp_calling_field == act_calling
        if ok_called:
            called_match += 1
        if ok_calling:
            calling_match += 1
        swcom = fn_to_swcom.get(name, "") or doc_swcom_by_name.get(name, "")
        if swcom == "SwCom_01":
            total_swcom01 += 1
            if ok_called:
                called_match_swcom01 += 1
            if ok_calling:
                calling_match_swcom01 += 1
        if (not ok_called or not ok_calling) and len(mismatches) < 50:
            mismatches.append(
                (
                    f"- `{name}` | called exp={sorted(exp_called_field)} act={sorted(act_called)} | "
                    f"calling exp={sorted(exp_calling_field)} act={sorted(act_calling)}"
                )
            )

    lines: List[str] = []
    lines.append("# Called/Calling Accuracy Report")
    lines.append("")
    lines.append(f"- Target DOCX: `{docx_path}`")
    lines.append(f"- Relation mode: `{relation_mode}`")
    lines.append(f"- Total functions compared: `{total}`")
    lines.append(f"- Called exact match: `{called_match}` / `{total}` ({_ratio(called_match, total)})")
    lines.append(f"- Calling exact match: `{calling_match}` / `{total}` ({_ratio(calling_match, total)})")
    if not details_by_name:
        lines.append("- Note: source function details are empty; accuracy result is document-baseline only.")
    lines.append("")
    lines.append("## SwCom_01")
    lines.append(f"- Functions: `{total_swcom01}`")
    lines.append(
        f"- Called exact match: `{called_match_swcom01}` / `{total_swcom01}` ({_ratio(called_match_swcom01, total_swcom01)})"
    )
    lines.append(
        f"- Calling exact match: `{calling_match_swcom01}` / `{total_swcom01}` ({_ratio(calling_match_swcom01, total_swcom01)})"
    )
    lines.append("")
    lines.append("## Sample Mismatches")
    lines.extend(mismatches or ["- none"])
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_swcom_context_report(docx_path: str, out_path: str) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(f"# SwCom Context Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)
    doc = docx.Document(str(docx_path))
    sw_rows: Dict[str, Dict[str, List[List[str]]]] = {}
    current_sw = ""
    pending = ""
    expected_header = "name|type|value range|reset value|description"
    for block in _iter_template_blocks(doc):
        if hasattr(block, "text"):
            text = (getattr(block, "text", "") or "").strip()
            if not text:
                continue
            m = re.search(r"\b(SwCom_\d+)\b", text, flags=re.I)
            if m:
                current_sw = m.group(1).replace("swcom", "SwCom")
                sw_rows.setdefault(current_sw, {"global_variables": [], "static_variables": []})
            if re.search(r"\bGlobal variables\b", text, flags=re.I):
                pending = "global_variables"
            elif re.search(r"\bStatic Variables\b", text, flags=re.I):
                pending = "static_variables"
            elif re.search(r"\bMacro\b|\bCalibration\b|\bInterface Functions\b|\bInternal Functions\b", text, flags=re.I):
                pending = ""
            continue
        if not hasattr(block, "rows"):
            continue
        if not current_sw or not pending:
            continue
        if not getattr(block, "rows", None):
            continue
        header = "|".join([(c.text or "").strip().lower() for c in block.rows[0].cells])
        if expected_header not in header:
            continue
        rows = [[(c.text or "").strip() for c in r.cells] for r in block.rows]
        sw_rows.setdefault(current_sw, {"global_variables": [], "static_variables": []})[pending].extend(rows[1:])
    suspicious: List[str] = []
    gv_total = 0
    sv_total = 0
    for swcom, sections in sw_rows.items():
        gv = sections.get("global_variables", [])
        sv = sections.get("static_variables", [])
        gv_total += len(gv)
        sv_total += len(sv)
        for sec_name, rows in [("global_variables", gv), ("static_variables", sv)]:
            for row in rows:
                name = str(row[0] if len(row) > 0 else "").strip()
                desc = str(row[4] if len(row) > 4 else "").strip()
                if not name or name.upper() == "N/A":
                    continue
                if name.startswith("REG_"):
                    suspicious.append(f"- {swcom} {sec_name}: REG alias `{name}`")
                if re.search(r"\.c$|\.h$", desc, flags=re.I):
                    suspicious.append(f"- {swcom} {sec_name}: description looks like file name `{desc}`")
    lines: List[str] = []
    lines.append("# SwCom Context Report")
    lines.append("")
    lines.append(f"- Target DOCX: `{docx_path}`")
    lines.append(f"- SwCom sections: `{len(sw_rows)}`")
    lines.append(f"- Global variable rows: `{gv_total}`")
    lines.append(f"- Static variable rows: `{sv_total}`")
    lines.append("")
    lines.append("## Suspicious Rows")
    lines.extend(suspicious[:120] if suspicious else ["- none"])
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_swcom_context_diff_report(reference_docx_path: str, target_docx_path: str, out_path: str) -> str:
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        import docx  # type: ignore
    except Exception as exc:
        out.write_text(f"# SwCom Context Diff Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)

    def _collect(path: str) -> Dict[str, Dict[str, int]]:
        doc = docx.Document(str(path))
        rows: Dict[str, Dict[str, int]] = {}
        current_sw = ""
        pending = ""
        expected_header = "name|type|value range|reset value|description"
        for block in _iter_template_blocks(doc):
            if hasattr(block, "text"):
                text = (getattr(block, "text", "") or "").strip()
                if not text:
                    continue
                m = re.search(r"\b(SwCom_\d+)\b", text, flags=re.I)
                if m:
                    current_sw = m.group(1).replace("swcom", "SwCom")
                    rows.setdefault(current_sw, {"global": 0, "static": 0})
                if re.search(r"\bGlobal variables\b", text, flags=re.I):
                    pending = "global"
                elif re.search(r"\bStatic Variables\b", text, flags=re.I):
                    pending = "static"
                elif re.search(r"\bMacro\b|\bCalibration\b|\bInterface Functions\b|\bInternal Functions\b", text, flags=re.I):
                    pending = ""
                continue
            if not hasattr(block, "rows"):
                continue
            if not current_sw or not pending:
                continue
            if not getattr(block, "rows", None):
                continue
            header = "|".join([(c.text or "").strip().lower() for c in block.rows[0].cells])
            if expected_header not in header:
                continue
            cnt = max(0, len(block.rows) - 1)
            rows.setdefault(current_sw, {"global": 0, "static": 0})[pending] += cnt
        return rows

    try:
        ref = _collect(reference_docx_path)
        tgt = _collect(target_docx_path)
    except Exception as exc:
        out.write_text(f"# SwCom Context Diff Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)

    sw_all = sorted(set(ref.keys()) | set(tgt.keys()))
    lines: List[str] = []
    lines.append("# SwCom Context Diff Report")
    lines.append("")
    lines.append(f"- Reference: `{reference_docx_path}`")
    lines.append(f"- Target: `{target_docx_path}`")
    lines.append(f"- SwCom count (reference): `{len(ref)}`")
    lines.append(f"- SwCom count (target): `{len(tgt)}`")
    lines.append("")
    only_ref = sorted(set(ref.keys()) - set(tgt.keys()))
    only_tgt = sorted(set(tgt.keys()) - set(ref.keys()))
    lines.append("## Section Coverage")
    lines.append(f"- Only in reference: `{len(only_ref)}`")
    lines.extend([f"  - {x}" for x in only_ref[:40]] if only_ref else ["  - none"])
    lines.append(f"- Only in target: `{len(only_tgt)}`")
    lines.extend([f"  - {x}" for x in only_tgt[:40]] if only_tgt else ["  - none"])
    lines.append("")
    lines.append("## Differences")
    diff_count = 0
    for sw in sw_all:
        r = ref.get(sw, {"global": 0, "static": 0})
        t = tgt.get(sw, {"global": 0, "static": 0})
        if r != t:
            diff_count += 1
            lines.append(
                f"- {sw}: global `{r.get('global',0)}` -> `{t.get('global',0)}`, "
                f"static `{r.get('static',0)}` -> `{t.get('static',0)}`"
            )
    if diff_count == 0:
        lines.append("- none")
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_uds_field_quality_gate_report(
    docx_path: str,
    out_path: str,
    thresholds: Optional[Dict[str, float]] = None,
) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(f"# UDS Field Quality Gate Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)

    gate = {
        "description_fill_rate": 0.70,
        "input_fill_rate": 0.20,
        "output_fill_rate": 0.10,
        "globals_global_fill_rate": 0.35,
        "globals_static_fill_rate": 0.15,
        "called_fill_rate": 0.50,
        "calling_fill_rate": 0.25,
        "asil_non_tbd_rate": 0.30,
        "related_non_tbd_rate": 0.30,
        "traceability_rate": 0.20,
    }
    _gate_improvement_guide = {
        "description_fill_rate": "코드 주석에 @brief 태그를 추가하거나, SDS 문서의 함수 설명을 보완하세요.",
        "input_fill_rate": "함수 시그니처의 파라미터에 @param 태그를 추가하세요.",
        "output_fill_rate": "함수의 @return 태그를 추가하세요.",
        "globals_global_fill_rate": "전역 변수 사용을 코드 주석에 명시하세요.",
        "globals_static_fill_rate": "정적 변수 사용을 코드 주석에 명시하세요.",
        "called_fill_rate": "소스 코드의 함수 호출 관계를 확인하세요. 간접 호출(함수 포인터)이 감지되지 않을 수 있습니다.",
        "calling_fill_rate": "호출 그래프의 역방향 관계를 확인하세요.",
        "asil_non_tbd_rate": "SRS/SDS 문서에 ASIL 등급을 명시하거나, 코드에 @asil 태그를 추가하세요.",
        "related_non_tbd_rate": "SRS 문서의 요구사항 ID를 함수와 매핑하거나, @requirement 태그를 추가하세요.",
        "traceability_rate": "DID/서비스 매핑을 확인하고, 코드 내 UDS 서비스 핸들러에 요구사항 ID를 연결하세요.",
    }
    if isinstance(thresholds, dict):
        for k, v in thresholds.items():
            if k in gate:
                try:
                    gate[k] = float(v)
                except Exception:
                    pass

    doc = docx.Document(str(docx_path))
    doc_map = _extract_function_info_from_docx(doc)
    payload = _load_uds_payload_for_docx(docx_path)
    payload_by_name = _payload_function_details_by_name(payload)
    payload_details = payload.get("function_details") or {}
    if isinstance(payload_details, dict) and payload_details:
        base_items = [info for info in payload_details.values() if isinstance(info, dict)]
    elif payload_by_name:
        base_items = list(payload_by_name.values())
    else:
        base_items = [row for row in doc_map.values() if isinstance(row, dict)]
    total = len(base_items)

    def _filled(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return False
        if s.upper() in {"N/A", "TBD", "-", "NONE"}:
            return False
        return True

    def _filled_desc(v: Any) -> bool:
        s = str(v or "").strip()
        if not _filled(s):
            return False
        if _is_generic_description(s):
            return False
        return True

    def _filled_list(v: Any) -> bool:
        if isinstance(v, list):
            return any(_filled(x) for x in v)
        return _filled(v)

    desc_ok = 0
    input_ok = 0
    output_ok = 0
    input_applicable = 0
    output_applicable = 0
    gg_ok = 0
    gs_ok = 0
    called_ok = 0
    calling_ok = 0
    tbd_asil = 0
    tbd_related = 0
    desc_high = 0
    desc_med = 0
    desc_low = 0
    direct_called_ok = 0
    indirect_support_ok = 0
    supported_called_ok = 0
    direct_traceable = 0
    supported_traceable = 0
    direct_call_applicable = 0
    leaf_function_count = 0
    for row in base_items:
        name_key = _normalize_symbol_name(str(row.get("name") or "")).lower()
        payload_info = payload_by_name.get(name_key, {}) if name_key else {}
        desc_value = payload_info.get("description") if payload_info else row.get("description")
        if _filled_desc(desc_value):
            desc_ok += 1
        dq = _classify_description_quality(
            str(payload_info.get("description") if payload_info else row.get("description") or ""),
            str(payload_info.get("description_source") if payload_info else row.get("description_source") or ""),
        )
        if dq == "high":
            desc_high += 1
        elif dq == "medium":
            desc_med += 1
        else:
            desc_low += 1
        asil_value = str(payload_info.get("asil") if payload_info else row.get("asil") or "").strip().upper()
        if asil_value == "TBD":
            tbd_asil += 1
        related_value = str(payload_info.get("related") if payload_info else row.get("related") or "").strip()
        if related_value.upper() == "TBD":
            tbd_related += 1
        proto = str(row.get("prototype") or "").strip()
        has_input_slot = _has_doc_input_slot(proto)
        if has_input_slot:
            input_applicable += 1
        if _filled_list(row.get("inputs")):
            input_ok += 1
        if _has_doc_output_slot(proto):
            output_applicable += 1
        if _filled_list(row.get("outputs")):
            output_ok += 1
        gg = row.get("globals_global")
        if gg is None:
            gg = row.get("globals")
        if _filled_list(gg):
            gg_ok += 1
        if _filled_list(row.get("globals_static")):
            gs_ok += 1
        called_names = _valid_call_names(payload_info.get("calls_list") if payload_info else row.get("called"))
        has_direct_calls = bool(called_names)
        if has_direct_calls:
            direct_called_ok += 1
        has_leaf_behavior = False
        if payload_info:
            logic_flow = payload_info.get("logic_flow") or []
            has_leaf_behavior = (
                not has_direct_calls
                and not _valid_call_names(payload_info.get("called_indirect"))
                and not _valid_call_names(payload_info.get("calling_indirect"))
                and all(isinstance(node, dict) and node.get("type") == "return" for node in logic_flow)
                and bool(logic_flow)
            )
        if has_leaf_behavior:
            leaf_function_count += 1
        else:
            direct_call_applicable += 1
        calling_names = _valid_call_names(payload_info.get("calling") if payload_info else row.get("calling"))
        if calling_names:
            calling_ok += 1
        has_indirect_calls = False
        if payload_info:
            if _valid_call_names(payload_info.get("called_indirect")) or _valid_call_names(payload_info.get("calling_indirect")):
                indirect_support_ok += 1
                has_indirect_calls = True
        has_supported_calls = has_direct_calls or has_indirect_calls
        if has_supported_calls:
            called_ok += 1
            supported_called_ok += 1
        related_text = related_value.strip().upper()
        has_related = bool(related_text) and related_text not in {"TBD", "N/A", "-"}
        if has_related and has_direct_calls:
            direct_traceable += 1
        if has_related and has_supported_calls:
            supported_traceable += 1

    def _rate(v: int, base: int) -> float:
        return 0.0 if base <= 0 else float(v) / float(base)

    asil_non_tbd = total - tbd_asil
    related_non_tbd = total - tbd_related
    traceable = supported_traceable

    input_base = max(input_ok, input_applicable)
    output_base = max(output_ok, output_applicable)
    metrics = {
        "description_fill_rate": _rate(desc_ok, total),
        "input_fill_rate": _rate(input_ok, input_base),
        "output_fill_rate": _rate(output_ok, output_base),
        "globals_global_fill_rate": _rate(gg_ok, total),
        "globals_static_fill_rate": _rate(gs_ok, total),
        "called_fill_rate": _rate(called_ok, total),
        "calling_fill_rate": _rate(calling_ok, total),
        "asil_non_tbd_rate": _rate(asil_non_tbd, total),
        "related_non_tbd_rate": _rate(related_non_tbd, total),
        "traceability_rate": _rate(traceable, total),
        "direct_called_fill_rate": _rate(direct_called_ok, total),
        "direct_traceability_rate": _rate(direct_traceable, total),
        "direct_called_fill_applicable_rate": _rate(direct_called_ok, direct_call_applicable),
    }
    failed = [k for k, v in metrics.items() if v < gate.get(k, 0.0)]

    lines: List[str] = []
    lines.append("# UDS Field Quality Gate Report")
    lines.append("")
    lines.append(f"- Target DOCX: `{docx_path}`")
    lines.append(f"- Total functions: `{total}`")
    lines.append(f"- Gate pass: `{'False' if failed else 'True'}`")
    lines.append(f"- Gates: `{len(metrics) - len(failed)}` / `{len(metrics)}` passed")
    lines.append("")
    lines.append("## Metrics")
    lines.append(f"- Description fill: `{desc_ok}` / `{total}` ({metrics['description_fill_rate']*100:.1f}%)")
    lines.append(f"- Input fill: `{input_ok}` / `{input_base}` ({metrics['input_fill_rate']*100:.1f}%)")
    lines.append(f"- Output fill: `{output_ok}` / `{output_base}` ({metrics['output_fill_rate']*100:.1f}%)")
    lines.append(f"- Globals(Global) fill: `{gg_ok}` / `{total}` ({metrics['globals_global_fill_rate']*100:.1f}%)")
    lines.append(f"- Globals(Static) fill: `{gs_ok}` / `{total}` ({metrics['globals_static_fill_rate']*100:.1f}%)")
    lines.append(f"- Called fill (supported): `{called_ok}` / `{total}` ({metrics['called_fill_rate']*100:.1f}%)")
    lines.append(f"- Calling fill: `{calling_ok}` / `{total}` ({metrics['calling_fill_rate']*100:.1f}%)")
    lines.append(f"- Direct called fill: `{direct_called_ok}` / `{total}` ({metrics['direct_called_fill_rate']*100:.1f}%)")
    lines.append(f"- Direct called fill (applicable): `{direct_called_ok}` / `{direct_call_applicable}` ({metrics['direct_called_fill_applicable_rate']*100:.1f}%)")
    lines.append(f"- Leaf / no-call functions: `{leaf_function_count}` / `{total}` ({_rate(leaf_function_count, total)*100:.1f}%)")
    lines.append(f"- Indirect call support: `{indirect_support_ok}` / `{total}` ({_rate(indirect_support_ok, total)*100:.1f}%)")
    lines.append(f"- ASIL non-TBD: `{asil_non_tbd}` / `{total}` ({metrics['asil_non_tbd_rate']*100:.1f}%)")
    lines.append(f"- Related non-TBD: `{related_non_tbd}` / `{total}` ({metrics['related_non_tbd_rate']*100:.1f}%)")
    lines.append(f"- Traceability (Related + Supported Call): `{traceable}` / `{total}` ({metrics['traceability_rate']*100:.1f}%)")
    lines.append(f"- Direct traceability: `{direct_traceable}` / `{total}` ({metrics['direct_traceability_rate']*100:.1f}%)")
    lines.append("")
    lines.append("## TBD Residual")
    lines.append(f"- ASIL TBD: `{tbd_asil}` / `{total}`")
    lines.append(f"- Related TBD: `{tbd_related}` / `{total}`")
    lines.append("")
    lines.append("## Description Quality Grade")
    lines.append(f"- High (comment/SDS/reference): `{desc_high}` ({_rate(desc_high, total)*100:.1f}%)")
    lines.append(f"- Medium (keyword inference): `{desc_med}` ({_rate(desc_med, total)*100:.1f}%)")
    lines.append(f"- Low (generic template): `{desc_low}` ({_rate(desc_low, total)*100:.1f}%)")
    lines.append("")
    lines.append("## Thresholds")
    for k, v in gate.items():
        lines.append(f"- {k}: `{v*100:.1f}%`")
    lines.append("")
    lines.append("## Failed Gates")
    if failed:
        for k in failed:
            guide = _gate_improvement_guide.get(k, "")
            lines.append(f"- **{k}**: {metrics[k]*100:.1f}% < {gate[k]*100:.1f}%")
            if guide:
                lines.append(f"  - 개선 가이드: {guide}")
    else:
        lines.append("- none")
    lines.append("")
    lines.append("## Improvement Recommendations")
    if tbd_asil > total * 0.5:
        lines.append("- SDS 문서를 추가하면 ASIL 정확도를 크게 향상시킬 수 있습니다.")
    if tbd_related > total * 0.5:
        lines.append("- SRS 문서를 추가하면 요구사항 추적성을 크게 향상시킬 수 있습니다.")
    if desc_low > total * 0.3:
        lines.append("- 소스 코드에 Doxygen 주석(@brief)을 추가하면 Description 품질이 향상됩니다.")
    if not failed and not (tbd_asil > total * 0.5 or tbd_related > total * 0.5 or desc_low > total * 0.3):
        lines.append("- 모든 품질 게이트를 통과했습니다. 정기적인 재검증을 권장합니다.")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_uds_delta_report(
    current_docx: str,
    previous_docx: str,
    out_path: str,
) -> str:
    """Compare two UDS DOCX files and generate a delta report."""
    try:
        import docx  # type: ignore
    except Exception as exc:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(f"# UDS Delta Report\n\n- error: {exc}\n", encoding="utf-8")
        return str(out)

    cur_doc = docx.Document(str(current_docx))
    prev_doc = docx.Document(str(previous_docx))
    cur_map = _extract_function_info_from_docx(cur_doc)
    prev_map = _extract_function_info_from_docx(prev_doc)

    cur_names = set(cur_map.keys())
    prev_names = set(prev_map.keys())
    added = sorted(cur_names - prev_names)
    removed = sorted(prev_names - cur_names)
    common = sorted(cur_names & prev_names)

    changed_asil: List[str] = []
    changed_desc: List[str] = []
    changed_related: List[str] = []
    for name in common:
        cur_info = cur_map.get(name, {})
        prev_info = prev_map.get(name, {})
        if not isinstance(cur_info, dict) or not isinstance(prev_info, dict):
            continue
        c_asil = str(cur_info.get("asil") or "").strip().upper()
        p_asil = str(prev_info.get("asil") or "").strip().upper()
        if c_asil != p_asil:
            changed_asil.append(f"  - {name}: `{p_asil}` → `{c_asil}`")
        c_rel = str(cur_info.get("related") or "").strip()
        p_rel = str(prev_info.get("related") or "").strip()
        if c_rel != p_rel:
            changed_related.append(f"  - {name}: `{p_rel[:40]}` → `{c_rel[:40]}`")
        c_desc = str(cur_info.get("description") or "").strip()[:60]
        p_desc = str(prev_info.get("description") or "").strip()[:60]
        if c_desc != p_desc:
            changed_desc.append(f"  - {name}")

    lines: List[str] = [
        "# UDS Delta Report",
        "",
        f"- Current: `{current_docx}`",
        f"- Previous: `{previous_docx}`",
        f"- Current functions: `{len(cur_names)}`",
        f"- Previous functions: `{len(prev_names)}`",
        "",
        "## Summary",
        f"- Added: `{len(added)}`",
        f"- Removed: `{len(removed)}`",
        f"- ASIL Changed: `{len(changed_asil)}`",
        f"- Related Changed: `{len(changed_related)}`",
        f"- Description Changed: `{len(changed_desc)}`",
        "",
    ]
    if added:
        lines.append("## Added Functions")
        lines.extend(f"  - {n}" for n in added[:50])
        lines.append("")
    if removed:
        lines.append("## Removed Functions")
        lines.extend(f"  - {n}" for n in removed[:50])
        lines.append("")
    if changed_asil:
        lines.append("## ASIL Changes")
        lines.extend(changed_asil[:50])
        lines.append("")
    if changed_related:
        lines.append("## Related ID Changes")
        lines.extend(changed_related[:50])
        lines.append("")
    if changed_desc:
        lines.append("## Description Changes")
        lines.extend(changed_desc[:50])
        lines.append("")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def _split_csvish(value: str) -> List[str]:
    items: List[str] = []
    for part in re.split(r"[,/;]+", str(value or "")):
        token = str(part or "").strip()
        if token and token not in items:
            items.append(token)
    return items


def _clean_param_lines(values: Any) -> List[str]:
    rows = values if isinstance(values, list) else [values]
    cleaned: List[str] = []
    for raw in rows:
        line = str(raw or "").strip()
        if not line:
            continue
        low = line.lower()
        if low in {"n/a", "none", "-", "tbd"}:
            continue
        if "no name type" in low and "description" in low:
            continue
        if line not in cleaned:
            cleaned.append(line)
    return cleaned


def _parse_accuracy_summary(text: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "called_exact_match": "",
        "calling_exact_match": "",
        "swcom_01_called_exact_match": "",
        "swcom_01_calling_exact_match": "",
        "total_functions": 0,
    }
    in_swcom01 = False
    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith("## swcom_01"):
            in_swcom01 = True
            continue
        if low.startswith("## ") and "swcom_01" not in low:
            in_swcom01 = False
        m_total = re.search(r"total functions compared:\s*`?(\d+)`?", line, flags=re.I)
        if m_total:
            out["total_functions"] = int(m_total.group(1))
            continue
        m_called = re.search(
            r"called exact match:\s*`?(\d+)`?\s*/\s*`?(\d+)`?\s*\(([^)]+)\)",
            line,
            flags=re.I,
        )
        if m_called:
            ratio = str(m_called.group(3)).strip()
            if in_swcom01:
                out["swcom_01_called_exact_match"] = ratio
            elif not out.get("called_exact_match"):
                out["called_exact_match"] = ratio
            continue
        m_calling = re.search(
            r"calling exact match:\s*`?(\d+)`?\s*/\s*`?(\d+)`?\s*\(([^)]+)\)",
            line,
            flags=re.I,
        )
        if m_calling:
            ratio = str(m_calling.group(3)).strip()
            if in_swcom01:
                out["swcom_01_calling_exact_match"] = ratio
            elif not out.get("calling_exact_match"):
                out["calling_exact_match"] = ratio
            continue
    return out


def _parse_quality_gate_summary(text: str) -> Dict[str, Any]:
    out: Dict[str, Any] = {"gate_pass": "", "metrics": {}}
    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        m_gate = re.search(r"gate pass:\s*`?(true|false)`?", line, flags=re.I)
        if m_gate:
            out["gate_pass"] = str(m_gate.group(1)).lower()
            continue
        m_metric = re.search(r"-\s*([^:]+):\s*`?\d+`?\s*/\s*`?\d+`?\s*\(([^)]+)\)", line)
        if m_metric:
            key = re.sub(r"[^a-z0-9]+", "_", str(m_metric.group(1)).strip().lower()).strip("_")
            out.setdefault("metrics", {})[key] = str(m_metric.group(2)).strip()
    return out


def build_uds_view_payload(
    docx_path: str,
    accuracy_report_path: str = "",
    quality_gate_report_path: str = "",
) -> Dict[str, Any]:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"python-docx import failed: {exc}") from exc
    target = Path(docx_path).expanduser().resolve()
    if not target.exists():
        raise FileNotFoundError(f"UDS DOCX not found: {target}")
    doc = docx.Document(str(target))
    doc_map = _extract_function_info_from_docx(doc)
    functions: List[Dict[str, Any]] = []
    swcom_summary: Dict[str, int] = {}
    traceability: List[Dict[str, Any]] = []
    for fn_id, info in doc_map.items():
        row = dict(info or {})
        fid = str(row.get("id") or fn_id or "").strip()
        name = _normalize_symbol_name(str(row.get("name") or ""))
        swcom = "UNMAPPED"
        m_sw = re.search(r"SwUFn_(\d{2})", fid, flags=re.I)
        if m_sw:
            swcom = f"SwCom_{m_sw.group(1)}"
        swcom_summary[swcom] = swcom_summary.get(swcom, 0) + 1
        related_tokens = _split_csvish(str(row.get("related") or ""))
        for rid in related_tokens:
            if re.match(r"^(Sw|REQ|SRS|SDS)", rid, flags=re.I):
                traceability.append(
                    {
                        "requirement_id": rid,
                        "function_id": fid,
                        "function_name": name,
                        "swcom": swcom,
                    }
                )
        called_text = str(row.get("called") or "")
        calling_text = str(row.get("calling") or "")
        functions.append(
            {
                "id": fid,
                "name": name,
                "swcom": swcom,
                "prototype": str(row.get("prototype") or "").strip(),
                "description": str(row.get("description") or "").strip(),
                "asil": str(row.get("asil") or "").strip(),
                "related": str(row.get("related") or "").strip(),
                "precondition": str(row.get("precondition") or "").strip(),
                "inputs": _clean_param_lines(row.get("inputs") or []),
                "outputs": _clean_param_lines(row.get("outputs") or []),
                # KEY_FN_GLOBALS — fall back to legacy "globals" key in old sidecar JSONs
                "globals_global": _clean_param_lines(
                    row.get("globals_global") or row.get("globals") or []
                ),
                # KEY_FN_STATICS
                "globals_static": _clean_param_lines(row.get("globals_static") or []),
                "called": _extract_call_names(called_text),
                "calling": _extract_call_names(calling_text),
                "called_raw": called_text,
                "calling_raw": calling_text,
            }
        )
    functions.sort(key=lambda x: (str(x.get("id") or ""), str(x.get("name") or "")))
    accuracy_summary: Dict[str, Any] = {}
    quality_gate_summary: Dict[str, Any] = {}
    acc_path = Path(accuracy_report_path).expanduser().resolve() if accuracy_report_path else None
    qg_path = Path(quality_gate_report_path).expanduser().resolve() if quality_gate_report_path else None
    if acc_path and acc_path.exists():
        accuracy_summary = _parse_accuracy_summary(acc_path.read_text(encoding="utf-8", errors="ignore"))
    if qg_path and qg_path.exists():
        quality_gate_summary = _parse_quality_gate_summary(qg_path.read_text(encoding="utf-8", errors="ignore"))
    return {
        "ok": True,
        "filename": target.name,
        "docx_path": str(target),
        "summary": {
            "total_functions": len(functions),
            "swcom_count": len(swcom_summary),
        },
        "functions": functions,
        "traceability": traceability,
        "swcom_summary": swcom_summary,
        "accuracy_summary": accuracy_summary,
        "quality_gate_summary": quality_gate_summary,
    }


def generate_uds_constraints_report(
    uds_payload: Dict[str, Any],
    out_path: str,
    schema_path: str = "",
) -> str:
    payload = _safe_dict(uds_payload)
    details_by_name = payload.get("function_details_by_name") or {}
    if (not isinstance(details_by_name, dict)) or (not details_by_name):
        details = payload.get("function_details") or {}
        rebuilt: Dict[str, Dict[str, Any]] = {}
        if isinstance(details, dict):
            for _, info in details.items():
                if not isinstance(info, dict):
                    continue
                n = str(info.get("name") or "").strip().lower()
                if n:
                    rebuilt[n] = info
        details_by_name = rebuilt

    target = {"functions": []}
    for _, info in (details_by_name or {}).items():
        if not isinstance(info, dict):
            continue
        target["functions"].append(
            {
                "id": str(info.get("id") or ""),
                "name": str(info.get("name") or ""),
                "description": str(info.get("description") or ""),
                "description_source": str(info.get("description_source") or ""),
                "asil": str(info.get("asil") or ""),
                "related": str(info.get("related") or ""),
                "inputs": list(info.get("inputs") or []),
                "outputs": list(info.get("outputs") or []),
                "globals_global": list(info.get("globals_global") or info.get("globals") or []),  # KEY_FN_GLOBALS (+ legacy fallback)
                "globals_static": list(info.get("globals_static") or []),  # KEY_FN_STATICS
                "called": str(info.get("called") or ""),
                "calling": str(info.get("calling") or ""),
            }
        )

    root_dir = Path(__file__).resolve().parent
    if not schema_path:
        schema_path = str(root_dir / "docs" / "uds_constraints.schema.json")
    schema_file = Path(schema_path)
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    payload_json = out.with_suffix(".json")
    payload_json.write_text(json.dumps(target, ensure_ascii=False, indent=2), encoding="utf-8")

    errors: List[str] = []
    valid = False
    try:
        import jsonschema  # type: ignore

        schema = json.loads(schema_file.read_text(encoding="utf-8"))
        validator = jsonschema.Draft202012Validator(schema)
        for err in validator.iter_errors(target):
            p = ".".join([str(x) for x in err.absolute_path])
            errors.append(f"{p}: {err.message}")
            if len(errors) >= 300:
                break
        valid = len(errors) == 0
    except Exception as exc:
        errors.append(str(exc))
        valid = False

    lines: List[str] = []
    lines.append("# UDS Constraint Validation Report")
    lines.append("")
    lines.append(f"- Schema: `{schema_file}`")
    lines.append(f"- Functions: `{len(target.get('functions') or [])}`")
    lines.append(f"- Valid: `{'True' if valid else 'False'}`")
    lines.append("")
    lines.append("## Errors (Top 300)")
    lines.extend([f"- {e}" for e in errors] if errors else ["- none"])
    lines.append("")
    lines.append(f"- Payload JSON: `{payload_json}`")
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)


def generate_asil_related_confidence_report(
    uds_payload: Dict[str, Any],
    out_path: str,
    generated_docx_path: str = "",
) -> str:
    payload = _safe_dict(uds_payload)
    details_by_name = payload.get("function_details_by_name") or {}
    if (not isinstance(details_by_name, dict)) or (not details_by_name):
        details = payload.get("function_details") or {}
        rebuilt: Dict[str, Dict[str, Any]] = {}
        if isinstance(details, dict):
            for _, info in details.items():
                if not isinstance(info, dict):
                    continue
                name = str(info.get("name") or "").strip().lower()
                if name:
                    rebuilt[name] = info
        details_by_name = rebuilt

    # Fallback: when source parser cannot provide function details, rebuild
    # lightweight entries from generated DOCX function tables.
    if (not isinstance(details_by_name, dict)) or (not details_by_name):
        docx_path = str(generated_docx_path or "").strip()
        if docx_path and Path(docx_path).exists():
            try:
                import docx  # type: ignore
                doc = docx.Document(docx_path)
                doc_map = _extract_function_info_from_docx(doc)
            except Exception:
                doc_map = {}
            rebuilt_from_doc: Dict[str, Dict[str, Any]] = {}
            for _, row in (doc_map or {}).items():
                if not isinstance(row, dict):
                    continue
                name = str(row.get("name") or "").strip().lower()
                if not name:
                    continue
                desc = str(row.get("description") or "").strip()
                asil = str(row.get("asil") or "").strip()
                rel = str(row.get("related") or "").strip()
                desc_src = "inference" if (not desc or _is_generic_description(desc)) else "reference"
                asil_src = "inference"
                rel_src = "inference"
                if asil and asil not in {"TBD", "N/A", "-"}:
                    asil_src = "sds"
                if rel and rel not in {"TBD", "N/A", "-"}:
                    if re.search(r"\bSw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+\b", rel):
                        rel_src = "srs"
                    elif re.search(r"\bSwCom_\d+\b", rel, flags=re.I):
                        rel_src = "rule"
                    else:
                        rel_src = "reference"
                rebuilt_from_doc[name] = {
                    "id": str(row.get("id") or ""),
                    "name": str(row.get("name") or ""),
                    "description": desc,
                    "description_source": desc_src,
                    "asil": asil,
                    "asil_source": asil_src,
                    "related": rel,
                    "related_source": rel_src,
                    "comment_description": "",
                    "comment_asil": "",
                    "comment_related": "",
                }
            if rebuilt_from_doc:
                details_by_name = rebuilt_from_doc
    else:
        # Merge generated DOCX function fields when payload sources are weak.
        docx_path = str(generated_docx_path or "").strip()
        if docx_path and Path(docx_path).exists():
            try:
                import docx  # type: ignore
                doc = docx.Document(docx_path)
                doc_map = _extract_function_info_from_docx(doc)
            except Exception:
                doc_map = {}
            if isinstance(doc_map, dict) and doc_map:
                by_name_from_doc: Dict[str, Dict[str, Any]] = {}
                by_id_from_doc: Dict[str, Dict[str, Any]] = {}
                def _norm_name(v: Any) -> str:
                    return re.sub(r"[^a-z0-9_]", "", str(v or "").lower())
                for _, row in doc_map.items():
                    if not isinstance(row, dict):
                        continue
                    n = str(row.get("name") or "").strip().lower()
                    if n:
                        by_name_from_doc[n] = row
                    fid = str(row.get("id") or "").strip()
                    if fid:
                        by_id_from_doc[fid] = row
                for name, info in list(details_by_name.items()):
                    if not isinstance(info, dict):
                        continue
                    row = None
                    fid = str(info.get("id") or "").strip()
                    if fid:
                        row = by_id_from_doc.get(fid)
                    if row is None:
                        row = by_name_from_doc.get(str(name).strip().lower())
                    if row is None:
                        nkey = _norm_name(name)
                        if nkey:
                            # fuzzy name fallback for rows with extra annotations.
                            for dk, dv in by_name_from_doc.items():
                                dkey = _norm_name(dk)
                                if not dkey:
                                    continue
                                if nkey == dkey or nkey in dkey or dkey in nkey:
                                    row = dv
                                    break
                    if not row:
                        continue
                    cur_desc = str(info.get("description") or "").strip()
                    cur_asil = str(info.get("asil") or "").strip()
                    cur_rel = str(info.get("related") or "").strip()
                    cur_desc_src = str(info.get("description_source") or "").strip().lower()
                    cur_rel_src = str(info.get("related_source") or "").strip().lower()
                    doc_desc = str(row.get("description") or "").strip()
                    doc_asil = str(row.get("asil") or "").strip()
                    doc_rel = str(row.get("related") or "").strip()
                    comment_desc = str(info.get("comment_description") or "").strip()
                    weak_desc_src = cur_desc_src in {"", "inference", "rule"}
                    weak_rel_src = cur_rel_src in {"", "inference", "rule"}
                    if comment_desc and not _is_generic_description(comment_desc) and weak_desc_src:
                        info["description"] = comment_desc
                        info["description_source"] = "comment"
                    elif doc_desc and doc_desc.upper() not in {"N/A", "TBD", "-"} and weak_desc_src:
                        # Prioritize explicit document text over inferred prose.
                        if (not cur_desc) or _is_generic_description(cur_desc):
                            info["description"] = doc_desc
                        info["description_source"] = "reference"
                    if (not cur_asil or cur_asil in {"TBD", "N/A", "-"}) and doc_asil:
                        info["asil"] = doc_asil
                        info["asil_source"] = "sds"
                    if (not cur_rel or cur_rel in {"TBD", "N/A", "-"}) and doc_rel and weak_rel_src:
                        info["related"] = doc_rel
                        if re.search(r"\bSw(?:TR|TSR|NTR|NTSR|CNF|EI|ST|STR|Fn|TK)_\d+\b", doc_rel):
                            info["related_source"] = "srs"
                        elif re.search(r"\bSwCom_\d+\b", doc_rel, flags=re.I):
                            info["related_source"] = "rule"
                        else:
                            info["related_source"] = "reference"

    src_labels = {
        "comment": "주석",
        "sds": "SDS",
        "srs": "SRS",
        "reference": "레퍼런스",
        "rule": "룰",
        "ai": "AI",
        "inference": "추론",
    }
    src_score = {
        "comment": 1.00,
        "sds": 0.95,
        "srs": 0.95,
        "reference": 0.90,
        "ai": 0.85,
        "rule": 0.75,
        "inference": 0.60,
    }

    def _norm_src(v: Any) -> str:
        s = str(v or "").strip().lower()
        if s == "sds_match":
            return "sds"
        if s == "hsis":
            return "sds"
        return s if s in src_labels else "inference"

    def _score_for(info: Dict[str, Any]) -> float:
        ds = _norm_src(info.get("description_source"))
        asrc = _norm_src(info.get("asil_source"))
        rsrc = _norm_src(info.get("related_source"))
        return (src_score.get(ds, 0.6) + src_score.get(asrc, 0.6) + src_score.get(rsrc, 0.6)) / 3.0

    def _evidence_for(info: Dict[str, Any], src: str, field_name: str) -> str:
        if src == "comment":
            if field_name == "description":
                return str(info.get("comment_description") or info.get("description") or "").strip()
            if field_name == "asil":
                return str(info.get("comment_asil") or info.get("asil") or "").strip()
            if field_name == "related":
                return str(info.get("comment_related") or info.get("related") or "").strip()
        if src == "sds":
            return "SDS 매핑 규칙에 의해 보강됨"
        if src == "srs":
            return "SRS 요구사항 ID/ASIL 추출 규칙에 의해 보강됨"
        if src == "reference":
            return "레퍼런스 SUDS DOCX에서 상속됨"
        if src == "ai":
            return "AI(Gemini) 모델이 코드 컨텍스트로 생성"
        if src == "rule":
            return "함수명/ID 기반 룰로 할당됨"
        return "코드/문맥 기반 추론"

    def _overall_grade(score: float) -> str:
        if score >= 0.95:
            return "A"
        if score >= 0.90:
            return "B"
        if score >= 0.80:
            return "C"
        return "D"

    total = 0
    desc_count: Dict[str, int] = {}
    asil_count: Dict[str, int] = {}
    rel_count: Dict[str, int] = {}
    desc_evidence_count: Dict[str, int] = {}
    rel_evidence_count: Dict[str, int] = {}
    desc_detail_count: Dict[str, int] = {}
    rel_detail_count: Dict[str, int] = {}
    op_counts: Dict[str, int] = {
        "desc_canonical_doc": 0,
        "desc_backed_by_code": 0,
        "desc_backed_by_hsis": 0,
        "desc_doc_only": 0,
        "rel_canonical_doc": 0,
        "rel_backed_by_code": 0,
        "rel_backed_by_hsis": 0,
        "rel_doc_only": 0,
    }
    low_conf: List[Tuple[float, str, str, str, str, str, str]] = []
    all_rows: List[Tuple[float, str, str, str, str, str, str, str, str, str, str]] = []
    by_swcom: Dict[str, Dict[str, int]] = {}
    for name, info in details_by_name.items():
        if not isinstance(info, dict):
            continue
        total += 1
        fid = str(info.get("id") or "").strip()
        swcom = "UNMAPPED"
        m_sw = re.search(r"SwUFn_(\d{2})", fid, flags=re.I)
        if m_sw:
            swcom = f"SwCom_{m_sw.group(1)}"
        swstats = by_swcom.setdefault(swcom, {"total": 0, "low": 0})
        swstats["total"] += 1
        ds = _norm_src(info.get("description_source"))
        asrc = _norm_src(info.get("asil_source"))
        rsrc = _norm_src(info.get("related_source"))
        desc_count[ds] = desc_count.get(ds, 0) + 1
        asil_count[asrc] = asil_count.get(asrc, 0) + 1
        rel_count[rsrc] = rel_count.get(rsrc, 0) + 1
        desc_detail = str(info.get("description_source_detail") or "").strip().lower()
        rel_detail = str(info.get("related_source_detail") or "").strip().lower()
        if desc_detail:
            desc_detail_count[desc_detail] = desc_detail_count.get(desc_detail, 0) + 1
        if rel_detail:
            rel_detail_count[rel_detail] = rel_detail_count.get(rel_detail, 0) + 1
        for src in info.get("description_evidence_sources") or []:
            s = str(src or "").strip().lower()
            if s:
                desc_evidence_count[s] = desc_evidence_count.get(s, 0) + 1
        for src in info.get("related_evidence_sources") or []:
            s = str(src or "").strip().lower()
            if s:
                rel_evidence_count[s] = rel_evidence_count.get(s, 0) + 1
        desc_evidence_set = {str(src or "").strip().lower() for src in (info.get("description_evidence_sources") or []) if str(src or "").strip()}
        rel_evidence_set = {str(src or "").strip().lower() for src in (info.get("related_evidence_sources") or []) if str(src or "").strip()}
        if ds in {"sds", "comment", "reference"}:
            op_counts["desc_canonical_doc"] += 1
            if "code" in desc_evidence_set:
                op_counts["desc_backed_by_code"] += 1
            if "hsis" in desc_evidence_set:
                op_counts["desc_backed_by_hsis"] += 1
            if "code" not in desc_evidence_set and "hsis" not in desc_evidence_set:
                op_counts["desc_doc_only"] += 1
        if rsrc in {"sds", "srs", "reference"}:
            op_counts["rel_canonical_doc"] += 1
            if "code" in rel_evidence_set:
                op_counts["rel_backed_by_code"] += 1
            if "hsis" in rel_evidence_set:
                op_counts["rel_backed_by_hsis"] += 1
            if "code" not in rel_evidence_set and "hsis" not in rel_evidence_set:
                op_counts["rel_doc_only"] += 1
        score = _score_for(info)
        all_rows.append(
            (
                score,
                str(info.get("id") or ""),
                str(name),
                ds,
                asrc,
                rsrc,
                str(info.get("asil") or ""),
                str(info.get("related") or ""),
                _evidence_for(info, ds, "description"),
                _evidence_for(info, asrc, "asil"),
                _evidence_for(info, rsrc, "related"),
            )
        )
        if score < 0.80:
            swstats["low"] += 1
            low_conf.append(
                (
                    score,
                    str(name),
                    ds,
                    asrc,
                    rsrc,
                    str(info.get("asil") or ""),
                    str(info.get("related") or ""),
                )
            )
    low_conf.sort(key=lambda x: (x[0], x[1], x[5], x[6]))
    all_rows.sort(key=lambda x: (x[0], x[2], x[1]))

    def _ratio(v: int) -> str:
        return "0.0%" if total <= 0 else f"{(v / total) * 100:.1f}%"

    def _dump_counter(title: str, counter: Dict[str, int]) -> List[str]:
        rows: List[str] = [f"## {title}"]
        if not counter:
            rows.append("- none")
            rows.append("")
            return rows
        for key, cnt in sorted(counter.items(), key=lambda kv: (-kv[1], kv[0])):
            rows.append(f"- {src_labels.get(key, key)}: `{cnt}` / `{total}` ({_ratio(cnt)})")
        rows.append("")
        return rows

    lines: List[str] = []
    avg_score = 0.0
    if total > 0:
        avg_score = sum([r[0] for r in all_rows]) / float(total)
    lines.append("# ASIL/Related ID Confidence Report")
    lines.append("")
    lines.append(f"- Total functions: `{total}`")
    lines.append(f"- Overall confidence score: `{avg_score:.3f}` (grade: `{_overall_grade(avg_score)}`)")
    lines.append(f"- Low confidence threshold: `< 0.80`")
    lines.append("- Source categories: `SDS / SRS / 주석 / 추론` (레퍼런스/룰 포함)")
    lines.append("")
    lines.extend(_dump_counter("Description Source", desc_count))
    lines.extend(_dump_counter("ASIL Source", asil_count))
    lines.extend(_dump_counter("Related ID Source", rel_count))
    lines.extend(_dump_counter("Description Evidence Mix", desc_evidence_count))
    lines.extend(_dump_counter("Related ID Evidence Mix", rel_evidence_count))
    lines.extend(_dump_counter("Description Source Detail", desc_detail_count))
    lines.extend(_dump_counter("Related ID Source Detail", rel_detail_count))
    lines.append("## Operating Judgment")
    lines.append("- Canonical policy: `doc-first` with `code/HSIS` as supporting evidence.")
    lines.append(
        f"- Description canonical(doc-backed): `{op_counts['desc_canonical_doc']}` / `{total}` ({_ratio(op_counts['desc_canonical_doc'])})"
    )
    lines.append(
        f"- Description canonical + code evidence: `{op_counts['desc_backed_by_code']}` / `{total}` ({_ratio(op_counts['desc_backed_by_code'])})"
    )
    lines.append(
        f"- Description canonical + HSIS evidence: `{op_counts['desc_backed_by_hsis']}` / `{total}` ({_ratio(op_counts['desc_backed_by_hsis'])})"
    )
    lines.append(
        f"- Description doc-only residual: `{op_counts['desc_doc_only']}` / `{total}` ({_ratio(op_counts['desc_doc_only'])})"
    )
    lines.append(
        f"- Related canonical(doc-backed): `{op_counts['rel_canonical_doc']}` / `{total}` ({_ratio(op_counts['rel_canonical_doc'])})"
    )
    lines.append(
        f"- Related canonical + code evidence: `{op_counts['rel_backed_by_code']}` / `{total}` ({_ratio(op_counts['rel_backed_by_code'])})"
    )
    lines.append(
        f"- Related canonical + HSIS evidence: `{op_counts['rel_backed_by_hsis']}` / `{total}` ({_ratio(op_counts['rel_backed_by_hsis'])})"
    )
    lines.append(
        f"- Related doc-only residual: `{op_counts['rel_doc_only']}` / `{total}` ({_ratio(op_counts['rel_doc_only'])})"
    )
    lines.append("")
    lines.append("## Component (SwCom) Low Confidence Ratio")
    if not by_swcom:
        lines.append("- none")
    else:
        for sw, stats in sorted(by_swcom.items(), key=lambda kv: kv[0]):
            t = int(stats.get("total", 0))
            l = int(stats.get("low", 0))
            ratio = 0.0 if t <= 0 else (l / t) * 100.0
            lines.append(f"- {sw}: low=`{l}` / total=`{t}` ({ratio:.1f}%)")
    lines.append("")
    lines.append("## Low Confidence Samples")
    if not low_conf:
        lines.append("- none")
    else:
        for score, name, ds, asrc, rsrc, asil_v, rel_v in low_conf[:120]:
            lines.append(
                f"- `{name}` score={score:.2f} | desc={src_labels.get(ds, ds)} | "
                f"asil={src_labels.get(asrc, asrc)}({asil_v or 'TBD'}) | "
                f"related={src_labels.get(rsrc, rsrc)}({rel_v or 'TBD'})"
            )
    lines.append("")
    lines.append("## Evidence Samples (Low First)")
    if not all_rows:
        lines.append("- none")
    else:
        for row in all_rows[:120]:
            score, fid, name, ds, asrc, rsrc, asil_v, rel_v, dev, aev, rev = row
            lines.append(
                f"- `{fid or 'N/A'}` `{name}` score={score:.2f} | "
                f"desc={src_labels.get(ds, ds)} | asil={src_labels.get(asrc, asrc)} | "
                f"related={src_labels.get(rsrc, rsrc)}"
            )
            lines.append(f"  - desc evidence: {dev[:180] if dev else 'N/A'}")
            lines.append(f"  - asil evidence: {aev[:180] if aev else (asil_v or 'N/A')}")
            lines.append(f"  - related evidence: {rev[:180] if rev else (rel_v or 'N/A')}")
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(lines), encoding="utf-8")
    return str(out)
