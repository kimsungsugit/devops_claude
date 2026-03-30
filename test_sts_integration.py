"""Integration test: STS generation with real HDPDM01 data."""
import sys
import os
import re
import json
import logging
from pathlib import Path

os.chdir(r"D:\Project\devops\260105")
sys.path.insert(0, ".")

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")

SRS_PATH = Path(r"D:\Project\devops\260105\docs\(HDPDM01_SRS) Software Requirements Specification_v1.05_20230510.docx")
STS_TEMPLATE = Path(r"D:\Project\devops\260105\docs\(HDPDM01_STS) Software Test Specification_v1.02_230116.xlsm")
OUTPUT_PATH = Path(r"D:\Project\devops\260105\output\sts_test_output.xlsm")

REQ_PAT = re.compile(r"\b(Sw(?:TR|TSR|NTR|NTSR|EI|CNF|ST|STR)_\d+)\b")


def extract_srs_text_lines(srs_path: Path) -> list:
    from docx import Document
    doc = Document(str(srs_path))
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            combined = " | ".join(cells)
            if combined.strip():
                lines.append(combined)
    return lines


def build_function_details() -> dict:
    """Try to get real function_details from source, fallback to mock."""
    source_root = Path(r"D:\Project\devops\260105\source")
    if source_root.is_dir():
        try:
            from report_generator import generate_uds_source_sections
            result = generate_uds_source_sections(str(source_root))
            fd = result.get("function_details", {})
            if fd:
                lf_count = sum(1 for v in fd.values() if isinstance(v, dict) and v.get("logic_flow"))
                print(f"[OK] Real source: {len(fd)} functions, {lf_count} with logic_flow")
                return fd
        except Exception as e:
            print(f"[WARN] Source parsing failed: {e}")

    print("[INFO] Using mock function_details (source dir not available)")
    return {
        "F001": {
            "name": "HDPDM01_Init",
            "inputs": ["void"],
            "output": "void",
            "calls_list": ["HAL_Init", "SystemClock_Config"],
            "related": "SwTR_0101 SwTR_0102",
            "logic_flow": [
                {"type": "call", "name": "HAL_Init"},
                {"type": "if", "condition": "status != OK",
                 "true_body": [{"type": "return", "value": "ERROR"}],
                 "false_body": [{"type": "call", "name": "SystemClock_Config"}]},
                {"type": "return", "value": "OK"},
            ],
        },
        "F002": {
            "name": "HDPDM01_MainLoop",
            "inputs": ["void"],
            "output": "void",
            "calls_list": ["ProcessInput", "UpdateOutput"],
            "related": "SwTR_0201 SwTR_0202",
            "logic_flow": [
                {"type": "loop", "kind": "while", "condition": "running",
                 "body": [
                     {"type": "call", "name": "ProcessInput"},
                     {"type": "call", "name": "UpdateOutput"},
                 ]},
            ],
        },
        "F003": {
            "name": "HDPDM01_ErrorHandler",
            "inputs": ["u8 error_code"],
            "output": "void",
            "calls_list": ["LogError", "ResetModule", "SafeState"],
            "related": "SwTR_0301",
            "logic_flow": [
                {"type": "switch", "expr": "error_code",
                 "cases": [
                     {"label": "ERR_COMM", "calls": ["LogError", "ResetModule"]},
                     {"label": "ERR_HW", "calls": ["SafeState"]},
                     {"label": "ERR_SW", "calls": ["LogError"]},
                 ],
                 "default_calls": ["LogError"]},
            ],
        },
    }


def main():
    print("=" * 60)
    print("STS Integration Test - HDPDM01")
    print("=" * 60)

    # Step 1: Test direct SRS DOCX table parsing
    print("\n[Step 1a] SRS DOCX table parsing...")
    from sts_generator import parse_srs_docx_tables
    docx_reqs = parse_srs_docx_tables(str(SRS_PATH))
    print(f"  DOCX table reqs: {len(docx_reqs)}")
    asil_counts = {}
    for r in docx_reqs:
        a = r.get("asil", "") or "N/A"
        asil_counts[a] = asil_counts.get(a, 0) + 1
    print(f"  ASIL distribution: {asil_counts}")
    type_counts = {}
    for r in docx_reqs:
        t = r.get("req_type", "?")
        type_counts[t] = type_counts.get(t, 0) + 1
    print(f"  Req type distribution: {type_counts}")
    safety_reqs = [r for r in docx_reqs if r.get("asil") and r["asil"].upper() not in ("QM", "TBD", "")]
    print(f"  Safety requirements (non-QM): {len(safety_reqs)}")
    if docx_reqs:
        print(f"  Sample: {docx_reqs[0]}")

    # Step 1b: Text fallback
    print("\n[Step 1b] SRS text line extraction (fallback)...")
    srs_lines = extract_srs_text_lines(SRS_PATH)
    from sts_generator import parse_requirements_structured
    text_reqs = parse_requirements_structured(srs_lines)
    print(f"  Text-parsed reqs: {len(text_reqs)} (deduped)")

    # Step 2: Function details
    print("\n[Step 2] Building function details...")
    func_details = build_function_details()

    # Step 3: Full pipeline with srs_docx_path
    print("\n[Step 3] Running full STS pipeline (DOCX table mode)...")
    from sts_generator import generate_sts

    project_config = {
        "project_id": "HDPDM01",
        "doc_id": "HDPDM01_STS",
        "version": "v1.00",
        "asil_level": "ASIL-B",
        "max_tc_per_req": 5,
        "default_test_env": "SwTE_01",
    }

    result = generate_sts(
        requirements_text=srs_lines,
        function_details=func_details,
        output_path=str(OUTPUT_PATH),
        template_path=str(STS_TEMPLATE),
        project_config=project_config,
        srs_docx_path=str(SRS_PATH),
    )

    # Step 4: Results
    print("\n" + "=" * 60)
    print("RESULTS")
    print("=" * 60)
    print(f"  Output file: {result['output_path']}")
    print(f"  Test cases: {result['test_case_count']}")
    print(f"  Elapsed: {result['elapsed_seconds']}s")

    qr = result.get("quality_report", {})
    print(f"\n  Quality Report:")
    print(f"    Total TCs: {qr.get('total_test_cases', 0)}")
    print(f"    Complete TCs (>=2 steps): {qr.get('complete_test_cases', 0)}")
    print(f"    Completeness: {qr.get('completeness_pct', 0)}%")
    print(f"    Safety TCs: {qr.get('safety_test_cases', 0)}")

    cov = result.get("trace_coverage", {})
    print(f"\n  Traceability:")
    print(f"    Total reqs: {cov.get('total_reqs', 0)}")
    print(f"    Covered reqs: {cov.get('covered_reqs', 0)}")
    print(f"    Coverage: {cov.get('pct', 0)}%")

    print(f"\n  Test Method: {qr.get('test_method_distribution', {})}")
    print(f"  Gen Method: {qr.get('gen_method_distribution', {})}")

    # Verify XLSM
    out_path = Path(result["output_path"])
    if out_path.exists():
        size_kb = out_path.stat().st_size // 1024
        print(f"\n  File size: {size_kb} KB")

        import openpyxl
        wb = openpyxl.load_workbook(str(out_path), read_only=True)
        print(f"  Sheets: {wb.sheetnames}")
        for sn in wb.sheetnames:
            ws = wb[sn]
            print(f"    {sn}: {ws.max_row}r x {ws.max_column}c")

        main_ws = wb["3.SW Integration Test Spec"]
        print(f"\n  Main sheet sample (row 7-9):")
        for r in range(7, min(10, main_ws.max_row + 1)):
            cells = []
            for c in range(1, 14):
                v = main_ws.cell(row=r, column=c).value
                cells.append(str(v)[:25] if v else "")
            print(f"    Row {r}: {' | '.join(cells)}")

        trace_ws = wb["5. Traceability(SwRS)"]
        print(f"\n  Traceability sheet: {trace_ws.max_row}r x {trace_ws.max_column}c")

        wb.close()
        print("\n[SUCCESS] STS generation completed successfully!")
    else:
        print("\n[ERROR] Output file not found!")


if __name__ == "__main__":
    main()
