#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import io
import re
from collections import defaultdict

# UTF-8 강제 출력
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

UDS_PATH = r"D:/Project/devops/260105/reports/uds_local/uds_local_20260318_091326.docx"
SUDS_PATH = r"D:/Project/devops/260105/docs/(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"

def load_doc(path):
    doc = Document(path)
    return doc

def get_style(para):
    return para.style.name if para.style else "Unknown"

def get_headings(doc):
    headings = []
    for para in doc.paragraphs:
        s = get_style(para)
        if s.startswith("Heading") and para.text.strip():
            headings.append((s, para.text.strip()))
    return headings

def get_table_headers(table):
    if table.rows:
        return [cell.text.strip() for cell in table.rows[0].cells]
    return []

print("=" * 70)
print("DOCX 파일 비교 분석")
print("=" * 70)

uds = load_doc(UDS_PATH)
suds = load_doc(SUDS_PATH)
print(f"[로드 완료] UDS: {len(uds.paragraphs)}단락, SUDS: {len(suds.paragraphs)}단락")

# ============================================================
# 1. 전체 구조 비교
# ============================================================
print("\n" + "#"*70)
print("# 1. 전체 구조 비교")
print("#"*70)

uds_headings = get_headings(uds)
suds_headings = get_headings(suds)

print(f"\n{'항목':<20} {'우리 UDS':>15} {'레퍼런스 SUDS':>15}")
print("-"*55)
print(f"{'총 단락 수':<20} {len(uds.paragraphs):>15} {len(suds.paragraphs):>15}")
print(f"{'총 테이블 수':<20} {len(uds.tables):>15} {len(suds.tables):>15}")
print(f"{'총 섹션 수':<20} {len(uds.sections):>15} {len(suds.sections):>15}")
print(f"{'헤딩 수':<20} {len(uds_headings):>15} {len(suds_headings):>15}")

# 커버/시작부
print("\n[UDS 시작부 (비어있지 않은 첫 10단락)]")
cnt = 0
for p in uds.paragraphs:
    if p.text.strip():
        print(f"  [{get_style(p)}] {p.text.strip()[:100]}")
        cnt += 1
        if cnt >= 10: break

print("\n[SUDS 시작부 (비어있지 않은 첫 10단락)]")
cnt = 0
for p in suds.paragraphs:
    if p.text.strip():
        print(f"  [{get_style(p)}] {p.text.strip()[:100]}")
        cnt += 1
        if cnt >= 10: break

# UDS 헤딩 구조
print(f"\n[UDS 헤딩 구조 (총 {len(uds_headings)}개, 최대 40개)]")
for s, t in uds_headings[:40]:
    lv = s[-1] if s[-1].isdigit() else "?"
    indent = "  " * (int(lv) - 1)
    print(f"  {indent}[H{lv}] {t[:80]}")
if len(uds_headings) > 40:
    print(f"  ... (총 {len(uds_headings)}개)")

# SUDS 헤딩 구조
print(f"\n[SUDS 헤딩 구조 (총 {len(suds_headings)}개, 전체)]")
for s, t in suds_headings:
    lv = s[-1] if s[-1].isdigit() else "?"
    indent = "  " * (int(lv) - 1)
    print(f"  {indent}[H{lv}] {t[:80]}")

# ============================================================
# 2. 테이블 구조 비교
# ============================================================
print("\n" + "#"*70)
print("# 2. 테이블 구조 비교")
print("#"*70)

print(f"\n[UDS 테이블 목록 (총 {len(uds.tables)}개, 처음 10개)]")
for i, t in enumerate(uds.tables[:10]):
    r, c = len(t.rows), len(t.columns)
    h = get_table_headers(t)
    print(f"  테이블#{i+1}: {r}행x{c}열 | 헤더: {h}")

print(f"\n[SUDS 테이블 목록 (총 {len(suds.tables)}개, 처음 10개)]")
for i, t in enumerate(suds.tables[:10]):
    r, c = len(t.rows), len(t.columns)
    h = get_table_headers(t)
    print(f"  테이블#{i+1}: {r}행x{c}열 | 헤더: {h}")

# 고유 헤더 패턴
print("\n[UDS 고유 테이블 헤더 패턴]")
uds_patterns = defaultdict(int)
for t in uds.tables:
    h = tuple(get_table_headers(t))
    uds_patterns[h] += 1
for h, cnt in sorted(uds_patterns.items(), key=lambda x: -x[1]):
    print(f"  (x{cnt}) {list(h)}")

print("\n[SUDS 고유 테이블 헤더 패턴]")
suds_patterns = defaultdict(int)
for t in suds.tables:
    h = tuple(get_table_headers(t))
    suds_patterns[h] += 1
for h, cnt in sorted(suds_patterns.items(), key=lambda x: -x[1]):
    print(f"  (x{cnt}) {list(h)}")

# ============================================================
# 3. 함수 항목 구조 비교
# ============================================================
print("\n" + "#"*70)
print("# 3. 함수 항목 구조 분석")
print("#"*70)

def extract_func_blocks(doc):
    """헤딩과 그 직후 테이블들을 묶어서 반환"""
    body = doc.element.body
    blocks = []
    cur_heading = None
    cur_level = None
    cur_tables = []

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            style_elem = child.find('.//' + qn('w:pStyle'))
            para_style = style_elem.get(qn('w:val'), '') if style_elem is not None else ''
            texts = [t.text for t in child.findall('.//' + qn('w:t')) if t.text]
            para_text = ''.join(texts).strip()
            if para_style.startswith('Heading') and para_text:
                if cur_heading and cur_tables:
                    blocks.append((cur_level, cur_heading, cur_tables))
                elif cur_heading and not cur_tables:
                    blocks.append((cur_level, cur_heading, []))
                cur_heading = para_text
                cur_level = para_style
                cur_tables = []
        elif tag == 'tbl':
            if cur_heading is not None:
                cur_tables.append(child)

    if cur_heading:
        blocks.append((cur_level, cur_heading, cur_tables))
    return blocks

def get_tbl_info(tbl_elem):
    rows = tbl_elem.findall('.//' + qn('w:tr'))
    if not rows:
        return {"rows": 0, "headers": [], "data": []}
    headers = []
    for cell in rows[0].findall('.//' + qn('w:tc')):
        texts = [t.text for t in cell.findall('.//' + qn('w:t')) if t.text]
        headers.append(''.join(texts).strip())
    data = []
    for row in rows[1:]:
        row_data = []
        for cell in row.findall('.//' + qn('w:tc')):
            texts = [t.text for t in cell.findall('.//' + qn('w:t')) if t.text]
            row_data.append(''.join(texts).strip()[:60])
        data.append(row_data)
    return {"rows": len(rows), "headers": headers, "data": data}

uds_blocks = extract_func_blocks(uds)
suds_blocks = extract_func_blocks(suds)

print(f"\nUDS 헤딩 블록 수: {len(uds_blocks)}")
print(f"SUDS 헤딩 블록 수: {len(suds_blocks)}")

# 함수처럼 보이는 블록
uds_func_re = re.compile(r'SwUFn_|SwFn_|[A-Za-z_]\w*\s*\(')
suds_func_re = re.compile(r'\d+\.\d+\.\d+|[A-Za-z_]\w*\s*\(')

uds_func_blocks = [(lv, h, tbls) for lv, h, tbls in uds_blocks if uds_func_re.search(h)]
suds_func_blocks = [(lv, h, tbls) for lv, h, tbls in suds_blocks if suds_func_re.search(h)]

print(f"\nUDS 함수 블록 수: {len(uds_func_blocks)}")
print(f"SUDS 함수 블록 수: {len(suds_func_blocks)}")

# UDS 샘플 함수 3개
print(f"\n[UDS 샘플 함수 3개 상세]")
for i, (lv, h, tbls) in enumerate(uds_func_blocks[:3]):
    print(f"\n  --- UDS 함수 #{i+1}: [{lv}] {h[:70]} ---")
    print(f"  테이블 수: {len(tbls)}")
    for j, tbl in enumerate(tbls):
        info = get_tbl_info(tbl)
        print(f"  테이블{j+1} ({info['rows']}행): 헤더={info['headers']}")
        for row_idx, row in enumerate(info['data'][:3]):
            print(f"    행{row_idx+1}: {row}")

# SUDS 샘플 함수 3개
print(f"\n[SUDS 샘플 함수 3개 상세]")
for i, (lv, h, tbls) in enumerate(suds_func_blocks[:3]):
    print(f"\n  --- SUDS 함수 #{i+1}: [{lv}] {h[:70]} ---")
    print(f"  테이블 수: {len(tbls)}")
    for j, tbl in enumerate(tbls):
        info = get_tbl_info(tbl)
        print(f"  테이블{j+1} ({info['rows']}행): 헤더={info['headers']}")
        for row_idx, row in enumerate(info['data'][:5]):
            print(f"    행{row_idx+1}: {row}")

# ============================================================
# 4. 레퍼런스 특이 항목
# ============================================================
print("\n" + "#"*70)
print("# 4. 레퍼런스(SUDS) 특이 항목")
print("#"*70)

uds_heading_texts = set(t for _, t in uds_headings)
suds_heading_texts = set(t for _, t in suds_headings)

only_suds = suds_heading_texts - uds_heading_texts
only_uds  = uds_heading_texts - suds_heading_texts
common    = uds_heading_texts & suds_heading_texts

print(f"\n공통 헤딩: {len(common)}개 | SUDS전용: {len(only_suds)}개 | UDS전용: {len(only_uds)}개")

print(f"\n[SUDS에만 있는 헤딩 (총 {len(only_suds)}개)]")
for t in sorted(only_suds):
    print(f"  - {t[:80]}")

print(f"\n[UDS에만 있는 헤딩 (총 {len(only_uds)}개, 최대 30개)]")
for t in sorted(list(only_uds))[:30]:
    print(f"  - {t[:80]}")

print(f"\n[공통 헤딩]")
for t in sorted(common):
    print(f"  - {t[:80]}")

# SUDS 테이블 전체 헤더
print(f"\n[SUDS 전체 테이블 헤더 목록 (총 {len(suds.tables)}개)]")
for i, t in enumerate(suds.tables):
    r, c = len(t.rows), len(t.columns)
    h = get_table_headers(t)
    print(f"  테이블#{i+1} ({r}행x{c}열): {h}")

# SUDS 함수별 테이블 필드 목록
print(f"\n[SUDS 함수별 포함 필드 목록 (처음 5개 함수)]")
for i, (lv, heading, tbls) in enumerate(suds_func_blocks[:5]):
    print(f"\n  함수{i+1}: {heading[:60]}")
    for j, tbl in enumerate(tbls):
        info = get_tbl_info(tbl)
        print(f"    테이블{j+1}: {info['headers']}")
        for row in info['data'][:3]:
            print(f"      -> {row}")

# UDS 함수별 테이블 필드 목록
print(f"\n[UDS 함수별 포함 필드 목록 (처음 5개 함수)]")
for i, (lv, heading, tbls) in enumerate(uds_func_blocks[:5]):
    print(f"\n  함수{i+1}: {heading[:60]}")
    for j, tbl in enumerate(tbls):
        info = get_tbl_info(tbl)
        print(f"    테이블{j+1}: {info['headers']}")
        for row in info['data'][:3]:
            print(f"      -> {row}")

# 스타일 비교
print("\n[사용 단락 스타일 비교]")
uds_styles = set(get_style(p) for p in uds.paragraphs)
suds_styles = set(get_style(p) for p in suds.paragraphs)
print(f"  UDS 스타일: {sorted(uds_styles)}")
print(f"  SUDS 스타일: {sorted(suds_styles)}")
print(f"  SUDS전용 스타일: {sorted(suds_styles - uds_styles)}")
print(f"  UDS전용 스타일: {sorted(uds_styles - suds_styles)}")

print("\n\n[분석 완료]")
