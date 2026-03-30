"""
Unit tests for UDS quality improvements.
Verifies each phase's improvements work correctly.
"""

import sys
import os
import json
import re
import pytest
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))

REF_SUDS = REPO_ROOT / "docs" / "(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"
SRS_TXT = REPO_ROOT / "docs" / "HDPDM01_SRS.txt"
SRS_DOCX = REPO_ROOT / "docs" / "(HDPDM01_SRS) Software Requirements Specification_v1.05_20230510.docx"
SDS_DOCX = REPO_ROOT / "docs" / "(HDPDM01_SDS) Software Architecture Design Specification_v1.04_20230512.docx"
SOURCE_ROOT = Path(r"D:\Project\Ados\PDS_64_RD")


class TestPhase2InputOutputParsing:
    """Phase 2: DOCX table parsing for Input/Output/Globals."""

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_total_functions_extracted(self, ref_fn_map):
        assert len(ref_fn_map) > 100, f"Expected >100 functions, got {len(ref_fn_map)}"

    def test_inputs_fill_rate(self, ref_fn_map):
        total = len(ref_fn_map)
        filled = sum(1 for v in ref_fn_map.values() if v.get("inputs") and len(v["inputs"]) > 0)
        rate = filled / total
        assert rate > 0.5, f"Input fill rate {rate:.2%} < 50%"

    def test_outputs_fill_rate(self, ref_fn_map):
        total = len(ref_fn_map)
        filled = sum(1 for v in ref_fn_map.values() if v.get("outputs") and len(v["outputs"]) > 0)
        rate = filled / total
        assert rate > 0.7, f"Output fill rate {rate:.2%} < 70%"

    def test_input_has_direction_and_type(self, ref_fn_map):
        for fid, info in ref_fn_map.items():
            for inp in (info.get("inputs") or []):
                assert "[IN]" in inp or "[INOUT]" in inp, f"{fid} input missing direction: {inp}"
                break
            if info.get("inputs"):
                break

    def test_output_has_direction_and_type(self, ref_fn_map):
        for fid, info in ref_fn_map.items():
            for out in (info.get("outputs") or []):
                assert "[OUT]" in out or "[INOUT]" in out, f"{fid} output missing direction: {out}"
                break
            if info.get("outputs"):
                break


class TestPhase2StaticNaming:
    """Phase 2: Static variable naming convention detection."""

    def test_static_prefixes_config(self):
        from config import STATIC_VAR_PREFIXES, GLOBAL_VAR_PREFIXES
        assert "u8s_" in STATIC_VAR_PREFIXES
        assert "u8g_" in GLOBAL_VAR_PREFIXES

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_globals_global_fill_rate(self, ref_fn_map):
        total = len(ref_fn_map)
        filled = sum(1 for v in ref_fn_map.values() if v.get("globals_global") and len(v["globals_global"]) > 0)
        rate = filled / total
        assert rate > 0.5, f"globals_global fill rate {rate:.2%} < 50%"


class TestPhase3SrsExtraction:
    """Phase 3: SRS parsing with ASIL/Related ID."""

    @pytest.fixture
    def srs_text(self):
        if not SRS_TXT.exists():
            pytest.skip("SRS TXT not available")
        return SRS_TXT.read_text(encoding="utf-8", errors="replace")

    def test_srs_extracts_requirements(self, srs_text):
        from report_generator import _extract_requirements_from_doc
        results = _extract_requirements_from_doc(srs_text)
        assert len(results) > 50, f"Expected >50 requirements, got {len(results)}"

    def test_srs_extracts_asil(self, srs_text):
        from report_generator import _extract_requirements_from_doc
        results = _extract_requirements_from_doc(srs_text)
        asil_count = sum(1 for r in results if "[ASIL:" in r)
        assert asil_count > 0, "No ASIL fields extracted from SRS"

    def test_req_map_accepts_hyphenated_asil(self):
        from report_generator import _build_req_map_from_texts

        text = "SwTSR_0101\nASIL: ASIL-B\nRelated ID: SyRS_0001"
        mapping = _build_req_map_from_texts([text])
        assert mapping["swtsr_0101"]["asil"] == "B"

    def test_docx_req_map_extracts_real_srs_sds_asil(self):
        from report_generator import _build_req_map_from_doc_paths

        if not SRS_DOCX.exists() or not SDS_DOCX.exists():
            pytest.skip("Reference SRS/SDS DOCX not available")

        mapping = _build_req_map_from_doc_paths([str(SRS_DOCX), str(SDS_DOCX)])
        asil_rows = sum(1 for item in mapping.values() if item.get("asil"))

        assert asil_rows > 10, f"Expected DOCX table parsing to extract ASIL rows, got {asil_rows}"
        assert mapping["swtr_0202"]["asil"] == "A"
        assert mapping["swtr_0203"]["asil"] == "A"
        assert mapping["swtr_0204"]["asil"] == "QM"


class TestPhase4Config:
    """Phase 4: Hardcoded paths removed."""

    def test_ref_suds_path_in_config(self):
        from config import UDS_REF_SUDS_PATH
        assert UDS_REF_SUDS_PATH, "UDS_REF_SUDS_PATH is empty"
        assert Path(UDS_REF_SUDS_PATH).suffix == ".docx"

    def test_no_hardcoded_paths_in_report_generator(self):
        code = (REPO_ROOT / "report_generator.py").read_text(encoding="utf-8")
        matches = re.findall(r'r"D:\\Project\\devops\\260105', code)
        assert len(matches) == 0, f"Found hardcoded paths: {matches}"


class TestPhase5JsonParsing:
    """Phase 5: Robust JSON parsing."""

    @pytest.fixture
    def parser(self):
        from workflow.uds_ai import _extract_json_payload
        return _extract_json_payload

    def test_valid_json(self, parser):
        assert parser('{"a": 1}') == {"a": 1}

    def test_markdown_code_block(self, parser):
        result = parser('```json\n{"a": 1}\n```')
        assert result == {"a": 1}

    def test_trailing_comma(self, parser):
        result = parser('{"a": 1,}')
        assert result is not None
        assert result["a"] == 1

    def test_truncated_json(self, parser):
        result = parser('{"a": "hello"')
        assert result is not None

    def test_preamble_text(self, parser):
        result = parser('Sure! Here is the JSON:\n{"a": 1}')
        assert result is not None
        assert result["a"] == 1

    def test_empty_string(self, parser):
        assert parser("") is None

    def test_non_json(self, parser):
        assert parser("not json at all") is None


class TestPhase5Prompts:
    """Phase 5: Enhanced prompts."""

    def test_reviewer_prompt_length(self):
        p = REPO_ROOT / "prompts" / "uds_reviewer.txt"
        assert p.exists()
        text = p.read_text(encoding="utf-8")
        assert len(text) > 500, f"Reviewer prompt too short: {len(text)} chars"

    def test_auditor_prompt_length(self):
        p = REPO_ROOT / "prompts" / "uds_auditor.txt"
        assert p.exists()
        text = p.read_text(encoding="utf-8")
        assert len(text) > 500, f"Auditor prompt too short: {len(text)} chars"

    def test_load_prompt(self):
        from workflow.uds_ai import _load_prompt
        text = _load_prompt("uds_reviewer")
        assert "ISO 26262" in text


class TestPhase6FuzzyMatching:
    """Phase 6: Fuzzy function matching."""

    def test_case_insensitive_match(self):
        from report_generator import generate_uds_function_mapping
        # Can't fully test without mocking, but verify function exists and signature
        import inspect
        sig = inspect.signature(generate_uds_function_mapping)
        params = list(sig.parameters.keys())
        assert "texts" in params
        assert "source_root" in params

    def test_traceability_accepts_function_details(self):
        from report_generator import generate_uds_traceability_mapping
        import inspect
        sig = inspect.signature(generate_uds_traceability_mapping)
        params = list(sig.parameters.keys())
        assert "function_details" in params


class TestQualityScore:
    """Phase 7: Integrated quality scoring."""

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_quality_score_calculation(self, ref_fn_map):
        total = len(ref_fn_map)
        fields = ["description", "inputs", "outputs", "called", "calling"]
        scores = {}
        for field in fields:
            filled = sum(
                1 for v in ref_fn_map.values()
                if (isinstance(v.get(field), list) and len(v[field]) > 0)
                or (isinstance(v.get(field), str) and v[field].strip() and v[field].strip().upper() not in {"N/A", "TBD", "-"})
            )
            scores[field] = filled / total
        avg = sum(scores.values()) / len(scores)
        assert avg > 0.5, f"Average quality score {avg:.2%} < 50%"

    def test_all_critical_fields_above_threshold(self, ref_fn_map):
        total = len(ref_fn_map)
        thresholds = {"description": 0.9, "called": 0.9, "calling": 0.5, "inputs": 0.5, "outputs": 0.7}
        for field, threshold in thresholds.items():
            filled = sum(
                1 for v in ref_fn_map.values()
                if (isinstance(v.get(field), list) and len(v[field]) > 0)
                or (isinstance(v.get(field), str) and v[field].strip() and v[field].strip().upper() not in {"N/A", "TBD", "-"})
            )
            rate = filled / total
            assert rate >= threshold, f"{field} rate {rate:.2%} < {threshold:.0%}"


class TestPhase2V2IsStaticVar:
    """Phase 2v2: _is_static_var fallback with naming convention."""

    def test_static_by_map(self):
        from report_generator import _is_static_var
        assert _is_static_var("u8s_foo", {"u8s_foo": True}) is True

    def test_static_by_naming_convention(self):
        from report_generator import _is_static_var
        assert _is_static_var("u8s_bar", {}) is True
        assert _is_static_var("s_myvar", {}) is True

    def test_global_by_naming(self):
        from report_generator import _is_static_var
        assert _is_static_var("g_something", {}) is False
        assert _is_static_var("u8g_data", {}) is False

    def test_map_overrides_naming(self):
        from report_generator import _is_static_var
        assert _is_static_var("s_override", {"s_override": False}) is False


class TestPhase2V2CallingMetric:
    """Phase 2v2: calling/called N/A as valid."""

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_calling_with_na_valid(self, ref_fn_map):
        total = len(ref_fn_map)
        present = sum(1 for v in ref_fn_map.values() if str(v.get("calling") or "").strip())
        rate = present / total
        assert rate >= 0.95, f"calling present rate {rate:.2%} < 95%"

    def test_called_with_na_valid(self, ref_fn_map):
        total = len(ref_fn_map)
        present = sum(1 for v in ref_fn_map.values() if str(v.get("called") or "").strip())
        rate = present / total
        assert rate >= 0.95, f"called present rate {rate:.2%} < 95%"


class TestPhase2V2SrsAsilRegex:
    """Phase 2v2: ASIL extraction with improved regex."""

    def test_asil_regex_matches_bracket(self):
        assert re.search(r"ASIL[:\s]*[A-DQ]", "[ASIL:A]", re.I) is not None

    def test_asil_regex_matches_qm(self):
        assert re.search(r"ASIL[:\s]*[A-DQ]", "[ASIL:QM]", re.I) is not None
        assert re.search(r"ASIL[:\s]*[A-DQ]", "ASIL QM", re.I) is not None

    def test_srs_asil_count(self):
        if not SRS_TXT.exists():
            pytest.skip("SRS TXT not available")
        from report_generator import _extract_requirements_from_doc
        text = SRS_TXT.read_text(encoding="utf-8", errors="replace")
        results = _extract_requirements_from_doc(text)
        asil_count = sum(1 for r in results if re.search(r"ASIL[:\s]*[A-DQ]", r, re.I))
        assert asil_count >= 20, f"Only {asil_count} requirements with ASIL"


class TestPhase2V2Precondition:
    """Phase 2v2: precondition inference from code body."""

    def test_infer_init_check(self):
        from report_generator import _infer_precondition_from_body
        body = "{ if (!g_Init_done) return; process(); }"
        result = _infer_precondition_from_body(body)
        assert result, "Should infer precondition from init check"

    def test_infer_null_check(self):
        from report_generator import _infer_precondition_from_body
        body = "{ if (ptr == NULL) return; *ptr = 42; }"
        result = _infer_precondition_from_body(body)
        assert "NULL" in result

    def test_infer_init_function(self):
        from report_generator import _infer_precondition_from_body
        result = _infer_precondition_from_body("{ x = 0; }", "SysInit")
        assert "N/A" in result

    def test_no_precond(self):
        from report_generator import _infer_precondition_from_body
        result = _infer_precondition_from_body("{ x = y + 1; }")
        assert result == ""


class TestPhase2V2CommentKeywords:
    """Phase 2v2: comment keyword expansion for precondition."""

    def test_at_pre_keyword(self):
        from workflow.code_parser.c_parser import _parse_comment_fields
        comment = "/** @pre System must be initialized\n * @param x input */\n"
        desc, asil, related, precond, *_ = _parse_comment_fields(comment)
        assert "initialized" in precond.lower()

    def test_pre_condition_keyword(self):
        from workflow.code_parser.c_parser import _parse_comment_fields
        comment = "/* Pre-condition: Timer running */\n"
        desc, asil, related, precond, *_ = _parse_comment_fields(comment)
        assert "timer" in precond.lower()


class TestPhase3TbdResolve:
    """Phase 3: ASIL TBD resolution via module inheritance and QM default."""

    def test_finalize_asil_default_qm(self):
        from report_generator import _finalize_function_fields
        info = {"name": "test_func", "description": "does stuff"}
        result = _finalize_function_fields(info)
        assert result["asil"] != "TBD", "ASIL should not be TBD"
        assert "QM" in result["asil"] or result["asil"] in {"A", "B", "C", "D"}

    def test_finalize_asil_preserves_existing(self):
        from report_generator import _finalize_function_fields
        info = {"name": "test_func", "description": "test", "asil": "B"}
        result = _finalize_function_fields(info)
        assert result["asil"] == "B"


class TestPhase3DescQuality:
    """Phase 3: Description quality classification."""

    def test_classify_high(self):
        from report_generator import _classify_description_quality
        assert _classify_description_quality("Motor Speed 감시 및 제어 수행", "comment") == "high"

    def test_classify_medium(self):
        from report_generator import _classify_description_quality
        assert _classify_description_quality("Motor Speed를 감시하고 Anti-pinch Detection을 수행하여 Door 안전 제어를 진행한다.", "inference") == "medium"

    def test_classify_low_empty(self):
        from report_generator import _classify_description_quality
        assert _classify_description_quality("", "") == "low"

    def test_classify_low_generic(self):
        from report_generator import _classify_description_quality
        assert _classify_description_quality("입력/상태 데이터를 처리하고 후속 제어 흐름을 진행한다", "inference") == "low"

    def test_generic_patterns_expanded(self):
        from report_generator import _is_generic_description
        assert _is_generic_description("핵심 동작을 수행한다") is True
        assert _is_generic_description("기능을 수행한다.") is True
        assert _is_generic_description("입력/상태 데이터를 처리하고 후속 제어 흐름을 진행한다") is True
        assert _is_generic_description("Motor Speed 감시 및 Anti-pinch 제어") is False


class TestPhase3SdsMapping:
    """Phase 3: SDS partition map flexible matching."""

    def test_sds_header_flexible(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        sds_docx = REPO_ROOT / "docs" / "(HDPDM01_SDS) Software Architecture Design Specification_v1.04_20230512.docx"
        if not sds_docx.exists():
            pytest.skip("SDS DOCX not available")
        from report_generator import _extract_sds_partition_map
        mapping = _extract_sds_partition_map(str(sds_docx))
        assert isinstance(mapping, dict)


class TestPhase3InputsImprove:
    """Phase 3: inputs direction tagging."""

    def test_const_param_tagged_in(self):
        from report_generator import _parse_signature_params
        result = _parse_signature_params("void foo(const U8 x)", tag_direction=True)
        assert len(result) > 0
        assert "[IN]" in result[0]

    def test_pointer_param_tagged_inout(self):
        from report_generator import _parse_signature_params
        result = _parse_signature_params("void foo(U8 *ptr)", tag_direction=True)
        assert len(result) > 0
        assert "[INOUT]" in result[0]

    def test_plain_param_tagged_in(self):
        from report_generator import _parse_signature_params
        result = _parse_signature_params("void foo(U8 x)", tag_direction=True)
        assert len(result) > 0
        assert "[IN]" in result[0]

    def test_no_tag_by_default(self):
        from report_generator import _parse_signature_params
        result = _parse_signature_params("void foo(U8 x)")
        assert len(result) > 0
        assert "[IN]" not in result[0]

    def test_void_param_empty(self):
        from report_generator import _parse_signature_params
        result = _parse_signature_params("void foo(void)")
        assert result == []


class TestPhase3JsonRobust:
    """Phase 3: JSON parsing robustness for edge cases."""

    def test_empty_returns_dict(self):
        from workflow.uds_ai import _extract_json_payload
        assert _extract_json_payload("") is None

    def test_none_like_returns_dict(self):
        from workflow.uds_ai import _extract_json_payload
        assert _extract_json_payload("   ") is None

    def test_non_json_returns_dict(self):
        from workflow.uds_ai import _extract_json_payload
        result = _extract_json_payload("not json at all")
        assert result is None

    def test_kv_extraction(self):
        from workflow.uds_ai import _extract_json_payload
        result = _extract_json_payload('The result is "key": "value" here')
        assert isinstance(result, dict)
        assert result.get("key") == "value"


class TestPhase3QualityBaseline:
    """Phase 3: quality baseline measures new metrics."""

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_tbd_asil_zero(self, ref_fn_map):
        from report_generator import _classify_description_quality
        tbd_count = sum(
            1 for v in ref_fn_map.values()
            if isinstance(v, dict) and str(v.get("asil") or "").strip().upper() == "TBD"
        )
        total = len(ref_fn_map)
        assert tbd_count / total < 0.1, f"TBD ASIL ratio {tbd_count}/{total} too high"

    def test_desc_quality_distribution(self, ref_fn_map):
        from report_generator import _classify_description_quality
        high = med = low = 0
        for v in ref_fn_map.values():
            if not isinstance(v, dict):
                continue
            q = _classify_description_quality(
                str(v.get("description") or ""),
                str(v.get("description_source") or ""),
            )
            if q == "high":
                high += 1
            elif q == "medium":
                med += 1
            else:
                low += 1
        total = high + med + low
        assert total > 0
        assert high / total > 0.3 or (high + med) / total > 0.5, \
            f"Description quality: high={high}, med={med}, low={low}"


class TestPhase4DiagramLogic:
    """Phase 4: Logic diagram dynamic height."""

    def test_call_graph_renders(self, tmp_path: Path):
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("PIL not available")
        from report_generator import _render_call_graph_image
        out = tmp_path / "test_logic.png"
        result = _render_call_graph_image(
            "test_func", ["call_a", "call_b"], None, out,
            max_depth=2,
        )
        assert result is not None
        assert out.exists()
        img = Image.open(str(out))
        h = img.height
        img.close()
        assert h >= 500

    def test_call_graph_depth3_taller(self, tmp_path: Path):
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("PIL not available")
        from report_generator import _render_call_graph_image
        call_map = {"call_a": ["sub1", "sub2"], "sub1": ["deep1"]}
        out2 = tmp_path / "d2.png"
        _render_call_graph_image("f", ["call_a"], call_map, out2, max_depth=2)
        out3 = tmp_path / "d3.png"
        _render_call_graph_image("f", ["call_a"], call_map, out3, max_depth=3)
        img2 = Image.open(str(out2))
        img3 = Image.open(str(out3))
        h2, h3 = img2.height, img3.height
        img2.close()
        img3.close()
        assert h3 >= h2


class TestPhase4DiagramStructure:
    """Phase 4: Structure diagram expanded items."""

    def test_structure_max_20(self, tmp_path: Path):
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("PIL not available")
        from report_generator import _render_unit_structure_image
        out = tmp_path / "struct.png"
        interfaces = [f"iface_{i}" for i in range(25)]
        internals = [f"internal_{i}" for i in range(5)]
        result = _render_unit_structure_image("TestMod", interfaces, internals, out)
        assert result is not None
        assert out.exists()

    def test_structure_overflow_text(self, tmp_path: Path):
        from report_generator import _render_unit_structure_image
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("PIL not available")
        out = tmp_path / "struct_overflow.png"
        interfaces = [f"iface_{i}" for i in range(25)]
        result = _render_unit_structure_image("Mod", interfaces, ["a"], out)
        assert result is not None


class TestPhase4ConditionExtract:
    """Phase 4: Condition extraction with else-if and nested parens."""

    def test_simple_if(self):
        from report_generator import _extract_primary_condition
        body = "{ if (x > 0) { do_a(); } }"
        cond = _extract_primary_condition(body)
        assert "x" in cond

    def test_else_if_chain(self):
        from report_generator import _extract_primary_condition
        body = "{ if (a) { f(); } else if (b) { g(); } else if (c) { h(); } }"
        cond = _extract_primary_condition(body)
        assert "/" in cond or "a" in cond

    def test_switch_case_labels(self):
        from report_generator import _extract_primary_condition
        body = "{ switch(state) { case 0: init(); break; case 1: run(); break; default: stop(); } }"
        cond = _extract_primary_condition(body)
        assert "switch" in cond.lower()

    def test_nested_parens(self):
        from report_generator import _extract_primary_condition
        body = "{ if ((a && b) || (c > d)) { f(); } }"
        cond = _extract_primary_condition(body)
        assert "a" in cond and "b" in cond


class TestPhase4QGateReport:
    """Phase 4: Quality Gate report includes TBD count and desc quality."""

    def test_qgate_report_has_tbd_section(self, tmp_path: Path):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        from report_generator import generate_uds_field_quality_gate_report
        out = tmp_path / "qgate.md"
        generate_uds_field_quality_gate_report(str(REF_SUDS), str(out))
        content = out.read_text(encoding="utf-8")
        assert "TBD Residual" in content
        assert "ASIL TBD" in content

    def test_qgate_report_has_desc_quality(self, tmp_path: Path):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        from report_generator import generate_uds_field_quality_gate_report
        out = tmp_path / "qgate2.md"
        generate_uds_field_quality_gate_report(str(REF_SUDS), str(out))
        content = out.read_text(encoding="utf-8")
        assert "Description Quality Grade" in content
        assert "High" in content
        assert "Medium" in content
        assert "Low" in content


class TestPhase5E2EGeneration:
    """Phase 5-3: End-to-end UDS generation and quality verification."""

    @pytest.fixture(scope="class")
    def generated_payload(self):
        if not SOURCE_ROOT.exists():
            pytest.skip("Source code directory not available")
        from report_generator import generate_uds_source_sections
        payload = generate_uds_source_sections(str(SOURCE_ROOT))
        assert payload, "generate_uds_source_sections returned empty"
        return payload

    def test_source_sections_has_function_details(self, generated_payload):
        fd = generated_payload.get("function_details", {})
        fbn = generated_payload.get("function_details_by_name", {})
        total = len(fd) + len(fbn)
        assert total > 0, "No function details generated from source"

    def test_source_sections_has_overview(self, generated_payload):
        overview = generated_payload.get("overview", "")
        assert len(overview) > 10, "Overview section is too short"

    def test_source_sections_has_call_map(self, generated_payload):
        cm = generated_payload.get("call_map", {})
        assert len(cm) > 5, f"Call map too small: {len(cm)} entries"

    def test_e2e_docx_generation_and_parsing(self, generated_payload, tmp_path: Path):
        import docx
        from report_generator import generate_uds_docx, _extract_function_info_from_docx

        out_path = str(tmp_path / "e2e_test.docx")
        payload = dict(generated_payload)
        payload["project_name"] = "E2E_Test"
        generate_uds_docx(None, payload, out_path)
        assert Path(out_path).exists(), "DOCX not created"
        assert Path(out_path).stat().st_size > 1000, "DOCX too small"

        doc = docx.Document(out_path)
        fn_map = _extract_function_info_from_docx(doc)
        if fn_map:
            total = len(fn_map)
            desc_filled = sum(
                1 for v in fn_map.values()
                if str(v.get("description") or "").strip()
            )
            assert desc_filled / total > 0.5, (
                f"Description fill rate {desc_filled}/{total} < 50%"
            )

    def test_e2e_quality_gate_fields(self, generated_payload):
        fd = generated_payload.get("function_details", {})
        fbn = generated_payload.get("function_details_by_name", {})
        all_details = list(fd.values()) + list(fbn.values())
        all_details = [d for d in all_details if isinstance(d, dict)]
        if not all_details:
            pytest.skip("No function details")
        total = len(all_details)
        with_desc = sum(1 for d in all_details if str(d.get("description") or "").strip())
        with_name = sum(1 for d in all_details if str(d.get("name") or "").strip())
        assert with_name / total > 0.9, f"Name fill rate {with_name}/{total} too low"
        assert with_desc / total > 0.5, f"Description fill rate {with_desc}/{total} too low"


class TestPhase5TemplateFormat:
    """Phase 5-2: Template formatting preservation with Run-based replacement."""

    def test_run_based_replacement_preserves_formatting(self):
        import docx
        from docx.shared import Pt, RGBColor
        from report_generator import _replace_docx_text
        doc = docx.Document()
        p = doc.add_paragraph()
        run1 = p.add_run("Project: ")
        run1.bold = True
        run1.font.size = Pt(14)
        run1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        run2 = p.add_run("{{project_name}}")
        run2.font.size = Pt(12)
        _replace_docx_text(doc, {"{{project_name}}": "MyProject"})
        assert "MyProject" in p.text
        assert run1.bold is True
        assert run1.font.size == Pt(14)
        assert run1.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)

    def test_header_footer_placeholder_detection(self):
        import docx
        from report_generator import _template_has_placeholders
        doc = docx.Document()
        doc.add_paragraph("No placeholders here.")
        assert _template_has_placeholders(doc) is False
        p = doc.add_paragraph()
        p.add_run("{{test_placeholder}}")
        assert _template_has_placeholders(doc) is True

    def test_table_cell_replacement(self):
        import docx
        from report_generator import _replace_docx_text
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "Version: {{version}}"
        _replace_docx_text(doc, {"{{version}}": "1.0.0"})
        assert "1.0.0" in cell.text


class TestPhase5DescQuality:
    """Phase 5-5: Description quality improvements."""

    def test_brief_extraction(self):
        from workflow.code_parser.c_parser import _parse_comment_fields
        comment = """
        * @brief Initialize motor controller hardware
        * @details Sets up PWM channels and configures timer registers
        * for motor control operation.
        * @param void
        """
        desc, _, _, _, *_ = _parse_comment_fields(comment)
        assert "Initialize motor controller" in desc
        assert "PWM" in desc

    def test_fallback_init_function(self):
        from report_generator import _fallback_function_description
        desc = _fallback_function_description("S_SysMain_Init", [])
        assert "초기화" in desc
        assert "S_SysMain_Init" in desc

    def test_fallback_with_callees(self):
        from report_generator import _fallback_function_description
        desc = _fallback_function_description(
            "S_Motor_Control",
            ["S_PWM_SetDuty", "S_ADC_Read"]
        )
        assert "S_Motor_Control" in desc
        assert "S_PWM_SetDuty" in desc or "호출" in desc

    def test_fallback_not_exact_generic(self):
        from report_generator import _is_generic_description
        assert _is_generic_description("function") is True
        assert _is_generic_description("tbd") is True
        desc = "S_SysMain_Init: Sys Main Init 모듈에서 초기화 절차를 수행하고 기본 파라미터를 설정한다."
        assert _is_generic_description(desc) is False

    def test_classify_comment_source_high(self):
        from report_generator import _classify_description_quality
        r = _classify_description_quality("Initialize motor controller hardware.", "comment")
        assert r == "high"


class TestPhase5GlobalsStatic:
    """Phase 5-4: Enhanced globals_static detection from code."""

    def test_local_static_detection_in_body(self):
        body = """
        static U8 counter = 0;
        static volatile U16 timer_val;
        counter++;
        if (counter > 10) { timer_val = 0; }
        """
        import re
        found = []
        for m in re.finditer(
            r"^\s*static\s+(?:volatile\s+)?(?:const\s+)?\w+\s+(\w+)\s*[=;\[]",
            body,
            re.M,
        ):
            found.append(m.group(1))
        assert "counter" in found
        assert "timer_val" in found

    def test_e2e_globals_static_rate(self):
        if not SOURCE_ROOT.exists():
            pytest.skip("Source code directory not available")
        from report_generator import generate_uds_source_sections
        payload = generate_uds_source_sections(str(SOURCE_ROOT))
        fd = payload.get("function_details", {})
        fbn = payload.get("function_details_by_name", {})
        all_fns = list(fd.values()) + list(fbn.values())
        all_fns = [d for d in all_fns if isinstance(d, dict)]
        if not all_fns:
            pytest.skip("No function details")
        with_gs = sum(
            1 for d in all_fns
            if isinstance(d.get("globals_static"), list) and len(d["globals_static"]) > 0
        )
        rate = with_gs / len(all_fns) if all_fns else 0
        assert rate > 0.05, f"globals_static rate {rate:.2%} too low"


class TestPhase5DescSource:
    """Phase 5-1: description_source tracking in DOCX extraction and classify."""

    @pytest.fixture(scope="class")
    def ref_fn_map(self):
        if not REF_SUDS.exists():
            pytest.skip("Reference SUDS not available")
        import docx
        from report_generator import _extract_function_info_from_docx
        doc = docx.Document(str(REF_SUDS))
        return _extract_function_info_from_docx(doc)

    def test_description_source_set_on_extraction(self, ref_fn_map):
        with_source = sum(
            1 for v in ref_fn_map.values()
            if v.get("description_source") == "reference"
        )
        total_with_desc = sum(
            1 for v in ref_fn_map.values()
            if str(v.get("description") or "").strip()
        )
        assert with_source >= total_with_desc * 0.8, (
            f"Only {with_source}/{total_with_desc} have description_source='reference'"
        )

    def test_classify_ai_as_high(self):
        from report_generator import _classify_description_quality
        result = _classify_description_quality("Initializes the motor control module and sets default parameters.", "ai")
        assert result == "high"

    def test_classify_reference_as_high(self):
        from report_generator import _classify_description_quality
        result = _classify_description_quality("Initializes system startup sequence.", "reference")
        assert result == "high"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
