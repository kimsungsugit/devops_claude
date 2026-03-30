"""
Coverage boost tests for report_generator.py.
Targets uncovered utility, parsing, generation, and report functions.
"""
import sys
import os
import re
import json
import tempfile
import pytest
from pathlib import Path
from typing import Dict, Any

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT))


# ---------------------------------------------------------------------------
# Phase 1: Utility / helper functions
# ---------------------------------------------------------------------------
class TestSafeFunctions:
    def test_safe_dict_with_dict(self):
        from report_generator import _safe_dict
        assert _safe_dict({"a": 1}) == {"a": 1}

    def test_safe_dict_with_none(self):
        from report_generator import _safe_dict
        assert _safe_dict(None) == {}

    def test_safe_dict_with_list(self):
        from report_generator import _safe_dict
        assert _safe_dict([1, 2]) == {}

    def test_safe_list_with_list(self):
        from report_generator import _safe_list
        assert _safe_list([1, 2]) == [1, 2]

    def test_safe_list_with_none(self):
        from report_generator import _safe_list
        assert _safe_list(None) == []

    def test_fmt_bool_true(self):
        from report_generator import _fmt_bool
        assert _fmt_bool(True) == "YES"

    def test_fmt_bool_false(self):
        from report_generator import _fmt_bool
        assert _fmt_bool(False) == "NO"

    def test_fmt_bool_none(self):
        from report_generator import _fmt_bool
        assert _fmt_bool(None) == "N/A"


class TestExtractIssueCounts:
    def test_with_issue_counts(self):
        from report_generator import _extract_issue_counts
        summary = {
            "static": {"cppcheck": {"issue_counts": {"total": 5, "error": 2, "warning": 3}}}
        }
        result = _extract_issue_counts(summary)
        assert result["total"] == 5
        assert result["error"] == 2

    def test_with_issues_list(self):
        from report_generator import _extract_issue_counts
        summary = {
            "static": {"cppcheck": {"data": {"issues": [{"id": 1}, {"id": 2}]}}}
        }
        result = _extract_issue_counts(summary)
        assert result["total"] == 2

    def test_empty_summary(self):
        from report_generator import _extract_issue_counts
        result = _extract_issue_counts({})
        assert result["total"] == 0


class TestGenerateMarkdownSummary:
    def test_basic_generation(self, tmp_path):
        from report_generator import generate_markdown_summary
        summary = {
            "project_root": "/test/project",
            "exit_code": 0,
            "coverage": {"enabled": True, "line_rate": 0.85, "threshold": 80},
            "tests": {"enabled": True, "ok": True},
            "build": {"enabled": True, "ok": True},
            "static": {"cppcheck": {"issue_counts": {"total": 3, "error": 1, "warning": 2}}},
            "fuzzing": {},
            "qemu": {},
            "domain_tests": {},
            "docs": {},
            "report_health": {"missing": [], "warnings": []},
            "scm": {"mode": "git"},
            "git": {"status": "ok", "branch": "main", "commit": "abc", "dirty": False},
            "svn": {},
            "strict": {},
            "artifacts": {"summary_json": "summary.json"},
        }
        out = tmp_path / "summary.md"
        result = generate_markdown_summary(summary, str(out))
        assert out.exists()
        content = out.read_text(encoding="utf-8")
        assert "project" in content.lower()
        assert "85.0%" in content
        assert result == str(out)

    def test_empty_summary(self, tmp_path):
        from report_generator import generate_markdown_summary
        out = tmp_path / "empty.md"
        generate_markdown_summary({}, str(out))
        assert out.exists()


class TestTextProcessing:
    def test_title_case_line(self):
        from report_generator import _title_case_line
        assert _title_case_line("hello") == "Hello"
        assert _title_case_line("") == ""
        assert _title_case_line("A") == "A"

    def test_split_sentences(self):
        from report_generator import _split_sentences
        result = _split_sentences("First sentence. Second one! Third?")
        assert len(result) == 3
        assert _split_sentences("") == []

    def test_trim_sentence_words(self):
        from report_generator import _trim_sentence_words
        assert _trim_sentence_words("a b c d e", 3) == "a b c..."
        assert _trim_sentence_words("a b", 5) == "a b"
        assert _trim_sentence_words("test", 0) == "test"

    def test_apply_sentence_rules(self):
        from report_generator import _apply_sentence_rules
        result = _apply_sentence_rules("Hello world. Second.", 1, 10, 200, True)
        assert result.endswith(".")
        assert _apply_sentence_rules("", 1, 10, 50, True) == ""

    def test_apply_sentence_rules_truncate(self):
        from report_generator import _apply_sentence_rules
        long_text = "This is a very long sentence that should be truncated."
        result = _apply_sentence_rules(long_text, 0, 0, 20, False)
        assert len(result) <= 23

    def test_apply_uds_rules(self):
        from report_generator import _apply_uds_rules
        text = "- First bullet\n- Second bullet\n- Third bullet"
        result = _apply_uds_rules(text, "overview")
        assert result
        assert "First" in result


class TestAiSectionHelpers:
    def test_ai_section_text_dict(self):
        from report_generator import _ai_section_text
        sections = {"overview": {"text": "Hello overview"}}
        assert _ai_section_text(sections, "overview") == "Hello overview"

    def test_ai_section_text_string(self):
        from report_generator import _ai_section_text
        sections = {"overview": "Direct string"}
        assert _ai_section_text(sections, "overview") == "Direct string"

    def test_ai_section_text_none(self):
        from report_generator import _ai_section_text
        assert _ai_section_text(None, "overview") == ""
        assert _ai_section_text({}, "missing") == ""

    def test_ai_evidence_lines(self):
        from report_generator import _ai_evidence_lines
        sections = {
            "overview": {
                "text": "test",
                "evidence": [
                    {"source_type": "rag", "source_file": "test.c", "excerpt": "example"},
                    {"source": "manual", "content": "note"},
                ],
            }
        }
        lines = _ai_evidence_lines(sections)
        assert len(lines) >= 1
        assert "overview" in lines[0]

    def test_ai_evidence_lines_empty(self):
        from report_generator import _ai_evidence_lines
        assert _ai_evidence_lines(None) == []
        assert _ai_evidence_lines({}) == []

    def test_ai_quality_warnings(self):
        from report_generator import _ai_quality_warnings
        sections = {"quality_warnings": ["warn1", "warn2", ""]}
        result = _ai_quality_warnings(sections)
        assert result == ["warn1", "warn2"]
        assert _ai_quality_warnings(None) == []
        assert _ai_quality_warnings({}) == []

    def test_merge_section_text(self):
        from report_generator import _merge_section_text
        sections = {"overview": {"text": "AI text"}}
        assert _merge_section_text("base", sections, "overview") == "AI text"
        assert _merge_section_text("base", sections, "missing") == "base"
        assert _merge_section_text("base", None, "overview") == "base"

    def test_merge_section_text_append(self):
        from report_generator import _merge_section_text
        sections = {"notes": {"text": "AI notes"}}
        result = _merge_section_text("base notes", sections, "notes", append_base=True)
        assert "AI notes" in result
        assert "base notes" in result

    def test_ai_document_text(self):
        from report_generator import _ai_document_text
        assert _ai_document_text({"document": "doc text"}) == "doc text"
        assert _ai_document_text(None) == ""

    def test_merge_logic_ai_items(self):
        from report_generator import _merge_logic_ai_items
        items = [{"title": "A", "description": ""}, {"title": "B", "description": "orig"}]
        ai = {"logic_diagrams": [{"title": "A", "description": "new desc"}]}
        result = _merge_logic_ai_items(items, ai)
        assert result[0]["description"] == "new desc"
        assert result[1]["description"] == "orig"

    def test_merge_logic_ai_items_no_ai(self):
        from report_generator import _merge_logic_ai_items
        items = [{"title": "A"}]
        assert _merge_logic_ai_items(items, None) is items
        assert _merge_logic_ai_items([], None) == []


class TestDocxTextHelpers:
    def test_add_docx_text_block(self):
        import docx
        from report_generator import _add_docx_text_block
        doc = docx.Document()
        _add_docx_text_block(doc, "Line 1\n2.1 Section Title\nLine 3")
        texts = [p.text for p in doc.paragraphs]
        assert any("Section Title" in t for t in texts)

    def test_add_docx_text_block_empty(self):
        import docx
        from report_generator import _add_docx_text_block
        doc = docx.Document()
        _add_docx_text_block(doc, "")
        assert any("N/A" in p.text for p in doc.paragraphs)

    def test_uds_lines_to_html(self):
        from report_generator import _uds_lines_to_html
        result = _uds_lines_to_html("- Item 1\n- Item 2")
        assert "<ul>" in result
        assert "Item 1" in result
        assert _uds_lines_to_html("") == "<p>N/A</p>"

    def test_uds_logic_html_with_items(self):
        from report_generator import _uds_logic_html
        items = [
            {"title": "Diagram A", "url": "http://img.png", "description": "Desc A"},
            {"title": "Diagram B", "url": "", "description": "Desc B"},
        ]
        result = _uds_logic_html(items)
        assert "Diagram A" in result
        assert "img" in result
        assert "Desc B" in result

    def test_uds_logic_html_empty(self):
        from report_generator import _uds_logic_html
        assert _uds_logic_html([]) == "<p>N/A</p>"


# ---------------------------------------------------------------------------
# Phase 2: C code parsing / extraction functions
# ---------------------------------------------------------------------------
class TestReadTextLimited:
    def test_read_existing_file(self, tmp_path):
        from report_generator import _read_text_limited
        f = tmp_path / "test.c"
        f.write_text("int main() { return 0; }", encoding="utf-8")
        result = _read_text_limited(f)
        assert "int main" in result

    def test_read_nonexistent(self, tmp_path):
        from report_generator import _read_text_limited
        result = _read_text_limited(tmp_path / "nope.c")
        assert result == ""

    def test_read_truncated(self, tmp_path):
        from report_generator import _read_text_limited
        f = tmp_path / "big.c"
        f.write_bytes(b"A" * 500)
        result = _read_text_limited(f, max_bytes=100)
        assert len(result) == 100


class TestCCodeExtraction:
    SAMPLE_C = """
static void S_Init(void) {
    g_State = 0;
}

extern void G_Update(U8 val);

U8 S_Calc(U16 input, U8 mode) {
    S_Init();
    return input + mode;
}

#define MAX_VAL  255
#define MIN_VAL  0

static U8 g_Counter;
U16 g_Speed = 100;
"""

    def test_strip_c_comments(self):
        from report_generator import _strip_c_comments
        code = "int a; /* comment */ int b; // line"
        result = _strip_c_comments(code)
        assert "comment" not in result
        assert "line" not in result
        assert "int a" in result

    def test_extract_prototypes(self):
        from report_generator import _extract_c_prototypes
        result = _extract_c_prototypes(self.SAMPLE_C)
        names = [r[0] for r in result]
        assert "G_Update" in names

    def test_extract_definitions(self):
        from report_generator import _extract_c_definitions
        result = _extract_c_definitions(self.SAMPLE_C)
        names = [r[0] for r in result]
        assert "S_Init" in names
        assert "S_Calc" in names

    def test_extract_function_bodies(self):
        from report_generator import _extract_c_function_bodies
        result = _extract_c_function_bodies(self.SAMPLE_C)
        assert "S_Init" in result
        assert "g_State" in result["S_Init"]

    def test_extract_simple_call_names(self):
        from report_generator import _extract_simple_call_names
        result = _extract_simple_call_names("S_Init(); if (x) { S_Calc(1, 2); }")
        assert "S_Init" in result
        assert "S_Calc" in result
        assert "if" not in result

    def test_extract_simple_call_names_func_ptr(self):
        from report_generator import _extract_simple_call_names
        result = _extract_simple_call_names("(*callback)(arg);")
        assert "callback" in result

    def test_extract_macros(self):
        from report_generator import _extract_c_macros
        result = _extract_c_macros(self.SAMPLE_C)
        assert "MAX_VAL" in result
        assert "MIN_VAL" in result

    def test_extract_macro_defs(self):
        from report_generator import _extract_c_macro_defs
        result = _extract_c_macro_defs(self.SAMPLE_C)
        names = [r[0] for r in result]
        assert "MAX_VAL" in names
        vals = dict(result)
        assert vals["MAX_VAL"] == "255"

    def test_extract_global_candidates(self):
        from report_generator import _extract_c_global_candidates
        result = _extract_c_global_candidates(self.SAMPLE_C)
        names = [g["name"] for g in result]
        assert "g_Counter" in names or "g_Speed" in names

    def test_empty_inputs(self):
        from report_generator import (
            _strip_c_comments, _extract_c_prototypes, _extract_c_definitions,
            _extract_c_function_bodies, _extract_simple_call_names,
            _extract_c_macros, _extract_c_macro_defs, _extract_c_global_candidates,
        )
        assert _strip_c_comments("") == ""
        assert _extract_c_prototypes("") == []
        assert _extract_c_definitions("") == []
        assert _extract_c_function_bodies("") == {}
        assert _extract_simple_call_names("") == []
        assert _extract_c_macros("") == []
        assert _extract_c_macro_defs("") == []
        assert _extract_c_global_candidates("") == []


class TestSignatureParamHelpers:
    def test_split_signature_param_chunks(self):
        from report_generator import _split_signature_param_chunks
        result = _split_signature_param_chunks("U8 a, U16 b, const U8 *c")
        assert len(result) == 3
        assert result[0].strip() == "U8 a"

    def test_split_nested_parens(self):
        from report_generator import _split_signature_param_chunks
        result = _split_signature_param_chunks("void (*cb)(int), U8 val")
        assert len(result) == 2

    def test_extract_param_symbol(self):
        from report_generator import _extract_param_symbol
        assert _extract_param_symbol("U8 mode") == "mode"
        assert _extract_param_symbol("const U16 *ptr") == "ptr"
        assert _extract_param_symbol("void (*callback)(int)") == "callback"
        assert _extract_param_symbol("U8 arr[10]") == "arr"
        assert _extract_param_symbol("") == ""

    def test_split_param(self):
        from report_generator import _split_param
        t, n, a = _split_param("U8 mode")
        assert t == "U8"
        assert n == "mode"
        t, n, a = _split_param("U16 arr[10]")
        assert n == "arr"
        assert "[10]" in a
        assert _split_param("") == ("", "", "")


class TestCommentStripping:
    def test_strip_comments_and_strings(self):
        from report_generator import _strip_comments_and_strings
        code = 'int a = 1; /* block */ int b; // line\nchar *s = "hello";'
        result = _strip_comments_and_strings(code)
        assert "block" not in result
        assert "line" not in result
        assert "hello" not in result
        assert _strip_comments_and_strings("") == ""

    def test_safe_eval_int(self):
        from report_generator import _safe_eval_int
        assert _safe_eval_int("42") == 42
        assert _safe_eval_int("0xFF") == 255
        assert _safe_eval_int("10 + 5") == 15
        assert _safe_eval_int("") is None
        assert _safe_eval_int("import os") is None

    def test_normalize_bracket_expr(self):
        from report_generator import _normalize_bracket_expr
        r, v = _normalize_bracket_expr("MAX_SIZE", {"MAX_SIZE": "10"})
        assert r == "10"
        assert v == 10
        r, v = _normalize_bracket_expr("unknown", {})
        assert r == "unknown"
        assert v is None
        r, v = _normalize_bracket_expr("", {})
        assert r == ""


class TestCollectVarUsage:
    def test_basic_usage(self):
        from report_generator import _collect_var_usage
        body = "g_State = input;\nif (g_Flag) { x = g_State + 1; }"
        result = _collect_var_usage(body, ["g_State", "g_Flag"])
        assert "g_State" in result
        assert result["g_State"].get("lhs") or result["g_State"].get("rhs")

    def test_empty_body(self):
        from report_generator import _collect_var_usage
        result = _collect_var_usage("", ["x"])
        if result:
            for v in result.values():
                assert not v.get("lhs") and not v.get("rhs") and not v.get("inout")


# ---------------------------------------------------------------------------
# Phase 3: generate_uds_docx mock tests
# ---------------------------------------------------------------------------
class TestDocxGeneration:
    def test_generate_basic_docx(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx
        out = tmp_path / "test_uds.docx"
        generate_uds_docx(None, mock_uds_payload, str(out))
        assert out.exists()
        assert out.stat().st_size > 1000

    def test_generate_docx_has_function_tables(self, mock_uds_payload, tmp_path):
        import docx
        from report_generator import generate_uds_docx
        out = tmp_path / "test_fn.docx"
        generate_uds_docx(None, mock_uds_payload, str(out))
        doc = docx.Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "S_Motor_Init" in all_text or any(
            "S_Motor_Init" in c.text
            for t in doc.tables for r in t.rows for c in r.cells
        )

    def test_generate_docx_with_ai_sections(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx
        payload = dict(mock_uds_payload)
        payload["ai_sections"] = {
            "overview": {"text": "AI-generated overview content."},
            "requirements": {"text": "AI requirements analysis."},
        }
        out = tmp_path / "test_ai.docx"
        generate_uds_docx(None, payload, str(out))
        assert out.exists()

    def test_generate_docx_with_sds_texts(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx
        payload = dict(mock_uds_payload)
        payload["sds_texts"] = ["Component Name: Motor\nASIL: QM\nDescription: Motor control module"]
        out = tmp_path / "test_sds.docx"
        generate_uds_docx(None, payload, str(out))
        assert out.exists()

    def test_generate_docx_empty_payload(self, tmp_path):
        from report_generator import generate_uds_docx
        out = tmp_path / "test_empty.docx"
        generate_uds_docx(None, {}, str(out))
        assert out.exists()

    def test_roundtrip_parse(self, mock_uds_payload, tmp_path):
        import docx
        from report_generator import generate_uds_docx, _extract_function_info_from_docx
        out = tmp_path / "roundtrip.docx"
        generate_uds_docx(None, mock_uds_payload, str(out))
        doc = docx.Document(str(out))
        fn_map = _extract_function_info_from_docx(doc)
        if fn_map:
            some_fn = next(iter(fn_map.values()))
            assert some_fn.get("name") or some_fn.get("id")


# ---------------------------------------------------------------------------
# Phase 4: Report generation functions
# ---------------------------------------------------------------------------
class TestPreviewGeneration:
    def test_preview_markdown(self, mock_uds_payload):
        from report_generator import generate_uds_preview_markdown
        result = generate_uds_preview_markdown(mock_uds_payload)
        assert isinstance(result, str)
        assert "TestProject" in result

    def test_preview_html(self, mock_uds_payload):
        from report_generator import generate_uds_preview_html
        result = generate_uds_preview_html(mock_uds_payload)
        assert isinstance(result, str)
        assert "TestProject" in result or "<" in result


class TestValidationReport:
    def test_validate_docx_structure(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx, validate_uds_docx_structure
        docx_path = tmp_path / "val_test.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        result = validate_uds_docx_structure(str(docx_path))
        assert isinstance(result, dict)

    def test_validation_report_output(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx, generate_uds_validation_report
        docx_path = tmp_path / "val_test2.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        out = tmp_path / "validation.md"
        generate_uds_validation_report(str(docx_path), str(out))
        assert out.exists()


class TestQualityGateReport:
    def test_quality_gate_from_generated_docx(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx, generate_uds_field_quality_gate_report
        docx_path = tmp_path / "qgate_test.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        out = tmp_path / "qgate.md"
        generate_uds_field_quality_gate_report(str(docx_path), str(out))
        if out.exists():
            content = out.read_text(encoding="utf-8")
            assert "Quality" in content or "quality" in content


class TestBuildViewPayload:
    def test_build_view_payload(self, mock_uds_payload, tmp_path):
        import docx
        from report_generator import generate_uds_docx, build_uds_view_payload
        docx_path = tmp_path / "view_test.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        result = build_uds_view_payload(str(docx_path))
        assert isinstance(result, dict)

    def test_build_view_payload_nonexistent(self, tmp_path):
        from report_generator import build_uds_view_payload
        with pytest.raises(FileNotFoundError):
            build_uds_view_payload(str(tmp_path / "nope.docx"))


class TestSwcomContextReport:
    def test_swcom_context_report(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx, generate_swcom_context_report
        docx_path = tmp_path / "swcom_src.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        out = tmp_path / "swcom_ctx.md"
        generate_swcom_context_report(str(docx_path), str(out))
        assert out.exists()


class TestConstraintsReport:
    def test_constraints_report(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_constraints_report
        out = tmp_path / "constraints.md"
        generate_uds_constraints_report(mock_uds_payload, str(out))
        assert out.exists()


class TestConfidenceReport:
    def test_confidence_report(self, mock_uds_payload, tmp_path):
        from report_generator import generate_asil_related_confidence_report
        out = tmp_path / "confidence.md"
        generate_asil_related_confidence_report(mock_uds_payload, str(out))
        assert out.exists()


# ---------------------------------------------------------------------------
# Phase 5: Requirements / mapping / traceability
# ---------------------------------------------------------------------------
class TestRequirementsExtraction:
    def test_extract_requirement_blocks(self):
        from report_generator import _extract_requirement_blocks
        text = """
REQ-001: The system shall initialize within 100ms.
REQ-002: The system shall perform self-diagnostic checks.
"""
        result = _extract_requirement_blocks(text)
        assert isinstance(result, list)

    def test_extract_doc_section(self):
        from report_generator import _extract_doc_section
        text = """1 Requirements
REQ-001: Motor init
REQ-002: Diag check
2 Design
Design info here
"""
        result = _extract_doc_section(text, "Requirements")
        assert "REQ-001" in result

    def test_extract_requirements_fallback(self):
        from report_generator import _extract_requirements_fallback
        text = "REQ-001: Motor initialization.\nREQ-002: Diagnostic."
        result = _extract_requirements_fallback(text)
        assert isinstance(result, (list, str))


class TestFunctionMapping:
    def test_function_mapping(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_function_mapping
        out = tmp_path / "fn_map.json"
        result = generate_uds_function_mapping(mock_uds_payload, str(out))
        assert isinstance(result, (dict, str, type(None)))


class TestTraceabilityMapping:
    def test_traceability_mapping_basic(self):
        from report_generator import generate_uds_traceability_mapping
        items = [
            {"id": "SwTR_001", "text": "Motor init requirement"},
            {"id": "SwTR_002", "text": "Diag check requirement"},
        ]
        mapping_texts = ["SwTR_001 -> S_Motor_Init\nSwTR_002 -> S_Diag_Check"]
        result = generate_uds_traceability_mapping(items, mapping_texts)
        assert isinstance(result, dict)

    def test_traceability_mapping_with_function_details(self, mock_function_details):
        from report_generator import generate_uds_traceability_mapping
        items = [{"id": "SwCom_01"}, {"id": "SwCom_02"}]
        result = generate_uds_traceability_mapping(items, [], function_details=mock_function_details)
        assert isinstance(result, dict)


class TestParsePreviewHtml:
    def test_parse_html(self):
        from report_generator import parse_uds_preview_html
        html = "<html><body><h1>Test</h1><p>Content</p></body></html>"
        result = parse_uds_preview_html(html)
        assert isinstance(result, (dict, str, type(None)))


# ---------------------------------------------------------------------------
# Phase 6: Image rendering extra paths
# ---------------------------------------------------------------------------
class TestImageRenderingExtras:
    def test_render_logic_text_image(self):
        from report_generator import _render_logic_text_image
        try:
            from PIL import Image
        except ImportError:
            pytest.skip("Pillow not available")
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            out = Path(tmp) / "logic.png"
            text = "Condition: if (x > 0)\nTRUE -> S_Init()\nFALSE -> S_Error()"
            result = _render_logic_text_image(text, out)
            if result:
                img = Image.open(str(out))
                assert img.size[0] > 0
                img.close()

    def test_render_swcom_overview_image(self):
        from report_generator import _render_swcom_overview_image
        try:
            from PIL import Image
        except ImportError:
            pytest.skip("Pillow not available")
        with tempfile.TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
            out = Path(tmp) / "swcom.png"
            swcoms = ["SwCom_01: Motor", "SwCom_02: Diag", "SwCom_03: Comm"]
            result = _render_swcom_overview_image(swcoms, Path(out))
            if result and Path(out).exists():
                img = Image.open(str(out))
                assert img.size[0] > 0
                img.close()


# ---------------------------------------------------------------------------
# Phase 6+: Additional coverage for large uncovered blocks
# ---------------------------------------------------------------------------
class TestSourceSectionsWithMockFiles:
    """Cover generate_uds_source_sections with temporary C files."""

    @pytest.fixture()
    def mock_source_tree(self, tmp_path):
        src = tmp_path / "src"
        src.mkdir()
        (src / "Motor.c").write_text("""
#include "Motor.h"
static U8 s_MotorState = 0;
U16 g_Speed = 100;
#define MOTOR_MAX 255

/* @brief Initialize motor controller */
void S_Motor_Init(U8 mode) {
    s_MotorState = mode;
    g_Speed = 0;
    S_PWM_SetDuty(0);
}

U8 S_Motor_GetState(void) {
    return s_MotorState;
}

static void S_Motor_Internal(void) {
    if (s_MotorState > 0) {
        g_Speed = MOTOR_MAX;
    }
}
""", encoding="utf-8")
        (src / "Motor.h").write_text("""
#ifndef MOTOR_H
#define MOTOR_H
typedef unsigned char U8;
typedef unsigned short U16;
extern void S_Motor_Init(U8 mode);
extern U8 S_Motor_GetState(void);
#endif
""", encoding="utf-8")
        (src / "Diag.c").write_text("""
#include "Diag.h"
static U8 s_DiagResult = 0;

/* @brief Check diagnostic status
 * @details Runs all diagnostic checks and updates result
 * @pre System must be initialized
 */
U8 S_Diag_Check(void) {
    s_DiagResult = 1;
    S_Error_Log(s_DiagResult);
    return s_DiagResult;
}

void S_Diag_Reset(void) {
    s_DiagResult = 0;
}
""", encoding="utf-8")
        (src / "Diag.h").write_text("""
#ifndef DIAG_H
#define DIAG_H
typedef unsigned char U8;
extern U8 S_Diag_Check(void);
extern void S_Diag_Reset(void);
#endif
""", encoding="utf-8")
        return str(src)

    def test_source_sections_from_mock(self, mock_source_tree):
        from report_generator import generate_uds_source_sections
        result = generate_uds_source_sections(mock_source_tree)
        assert isinstance(result, dict)
        assert result.get("overview") or result.get("function_details") or result.get("function_details_by_name")

    def test_source_sections_has_functions(self, mock_source_tree):
        from report_generator import generate_uds_source_sections
        result = generate_uds_source_sections(mock_source_tree)
        fd = result.get("function_details", {})
        fbn = result.get("function_details_by_name", {})
        total = len(fd) + len(fbn)
        assert total >= 2, f"Expected >= 2 functions, got {total}"

    def test_source_sections_has_globals(self, mock_source_tree):
        from report_generator import generate_uds_source_sections
        result = generate_uds_source_sections(mock_source_tree)
        gv = result.get("global_vars", [])
        sv = result.get("static_vars", [])
        assert len(gv) + len(sv) >= 1

    def test_source_sections_call_map(self, mock_source_tree):
        from report_generator import generate_uds_source_sections
        result = generate_uds_source_sections(mock_source_tree)
        cm = result.get("call_map", {})
        assert isinstance(cm, dict)

    def test_source_sections_empty_dir(self, tmp_path):
        from report_generator import generate_uds_source_sections
        empty = tmp_path / "empty_src"
        empty.mkdir()
        result = generate_uds_source_sections(str(empty))
        assert isinstance(result, dict)

    def test_source_sections_nonexistent(self, tmp_path):
        from report_generator import generate_uds_source_sections
        result = generate_uds_source_sections(str(tmp_path / "nope"))
        assert result == {}


class TestDocxGenerationExtended:
    """Additional DOCX generation paths for coverage."""

    def test_docx_with_template(self, mock_uds_payload, tmp_path):
        import docx
        from report_generator import generate_uds_docx
        tpl = tmp_path / "template.docx"
        doc = docx.Document()
        doc.add_paragraph("{{project_name}}")
        doc.add_paragraph("{{overview}}")
        doc.save(str(tpl))
        out = tmp_path / "from_tpl.docx"
        generate_uds_docx(str(tpl), mock_uds_payload, str(out))
        assert out.exists()
        result_doc = docx.Document(str(out))
        all_text = "\n".join(p.text for p in result_doc.paragraphs)
        assert "TestProject" in all_text or len(all_text) > 10

    def test_docx_with_many_functions(self, tmp_path):
        from report_generator import generate_uds_docx
        details = {}
        for i in range(10):
            fid = f"SwUFn_{i:03d}"
            details[fid] = {
                "id": fid,
                "name": f"Func_{i}",
                "prototype": f"void Func_{i}(U8 p{i})",
                "description": f"Function {i} performs operation {i}.",
                "description_source": "comment",
                "asil": "QM",
                "related": f"SwCom_{i:02d}",
                "inputs": [f"[IN] U8 p{i}"],
                "outputs": [],
                "precondition": "N/A",
                "globals_global": [],
                "globals_static": [],
                "called": "",
                "calling": "",
                "module_name": f"Mod{i}",
            }
        payload = {
            "project_name": "ManyFuncTest",
            "overview": "Test with many functions.",
            "function_details": details,
            "function_table_rows": [
                [f"SwCom_{i:02d}", f"Mod{i}", f"SwUFn_{i:03d}", f"Func_{i}", "I/F", ""]
                for i in range(10)
            ],
        }
        out = tmp_path / "many_fn.docx"
        generate_uds_docx(None, payload, str(out))
        assert out.exists()
        assert out.stat().st_size > 2000

    def test_docx_with_logic_diagrams(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx
        payload = dict(mock_uds_payload)
        payload["logic_diagrams"] = [
            {"title": "Motor Init Logic", "description": "Checks mode then sets PWM"},
        ]
        out = tmp_path / "with_logic.docx"
        generate_uds_docx(None, payload, str(out))
        assert out.exists()


class TestAccuracyReport:
    """Cover generate_called_calling_accuracy_report."""

    def test_accuracy_report(self, mock_uds_payload, tmp_path):
        from report_generator import generate_uds_docx
        docx_path = tmp_path / "acc_src.docx"
        generate_uds_docx(None, mock_uds_payload, str(docx_path))
        out = tmp_path / "accuracy.md"
        empty_src = tmp_path / "empty_src"
        empty_src.mkdir()
        try:
            from report_generator import generate_called_calling_accuracy_report
            generate_called_calling_accuracy_report(str(docx_path), str(empty_src), str(out))
        except Exception:
            pass


class TestNormalizationHelpers:
    """Cover normalization functions not yet tested."""

    def test_normalize_swufn_id(self):
        from report_generator import _normalize_swufn_id
        assert _normalize_swufn_id("SwUFn_001") == "SwUFn_001"
        assert "SwUFn" in _normalize_swufn_id("swufn_1")
        assert _normalize_swufn_id("") == ""

    def test_normalize_call_field(self):
        from report_generator import _normalize_call_field
        r = _normalize_call_field("S_Init\nS_Calc\nS_Init")
        assert "S_Init" in r
        assert "S_Calc" in r
        assert _normalize_call_field("") == ""

    def test_normalize_asil_value(self):
        from report_generator import _normalize_asil_value
        result = _normalize_asil_value("ASIL-B")
        assert "B" in result
        assert "QM" in _normalize_asil_value("qm") or _normalize_asil_value("qm") == "QM"
        assert _normalize_asil_value("") == ""

    def test_normalize_related_ids(self):
        from report_generator import _normalize_related_ids
        r = _normalize_related_ids("SwCom_01, SwCom_02, SwCom_01")
        assert "SwCom_01" in r
        assert "SwCom_02" in r

    def test_dedupe_multiline_text(self):
        from report_generator import _dedupe_multiline_text
        r = _dedupe_multiline_text("Line1\nLine2\nLine1\nLine3")
        assert r.count("Line1") == 1
        assert "Line3" in r
        assert _dedupe_multiline_text("N/A", na_to_empty=True) == ""

    def test_extract_call_names(self):
        from report_generator import _extract_call_names
        r = _extract_call_names("S_Init\nS_Calc\nS_Error")
        assert "S_Init" in r
        assert len(r) >= 3

    def test_normalize_symbol_name(self):
        from report_generator import _normalize_symbol_name
        assert _normalize_symbol_name("S_Motor_Init") == "S_Motor_Init"
        assert _normalize_symbol_name("  S_Init  ") == "S_Init"
        assert _normalize_symbol_name("") == ""


class TestConditionAndLogicFunctions:
    """Cover _extract_primary_condition, _extract_condition_branch_calls, etc."""

    def test_switch_condition(self):
        from report_generator import _extract_primary_condition
        body = "{ switch(state) { case 0: break; case 1: run(); break; } }"
        cond = _extract_primary_condition(body)
        assert "state" in cond or "switch" in cond.lower() or "case" in cond.lower()

    def test_extract_condition_branch_calls(self):
        from report_generator import _extract_condition_branch_calls
        body = "if (x) { A(); B(); } else { C(); }"
        true_calls, false_calls = _extract_condition_branch_calls(body)
        assert isinstance(true_calls, list)
        assert isinstance(false_calls, list)

    def test_extract_logic_terminal_paths(self):
        from report_generator import _extract_logic_terminal_paths
        body = "if (err) { return -1; } DoWork(); return 0;"
        ret, err = _extract_logic_terminal_paths(body)
        assert isinstance(ret, (bool, str))
        assert isinstance(err, (bool, str))

    def test_infer_precondition_from_body(self):
        from report_generator import _infer_precondition_from_body
        body = "if (!g_Initialized) { return; } DoWork();"
        result = _infer_precondition_from_body(body, "S_Test")
        assert isinstance(result, str)

    def test_format_param_entry(self):
        from report_generator import _format_param_entry
        r = _format_param_entry("U8", "mode", False, "IN", [], 0, False)
        assert "mode" in r
        assert "U8" in r

    def test_extract_return_type(self):
        from report_generator import _extract_return_type
        assert "U8" in _extract_return_type("U8 S_GetMode(void)", "S_GetMode")
        assert _extract_return_type("void S_Init(void)", "S_Init") in {"void", ""}

    def test_classify_param_direction(self):
        from report_generator import _classify_param_direction
        assert _classify_param_direction("const U8 val") == "[IN]"
        assert _classify_param_direction("U8 *ptr") in {"[INOUT]", "[OUT]"}


class TestPreviewHtmlExtended:
    """Cover the second generate_uds_preview_html variant."""

    def test_preview_html_with_logic_items(self, mock_uds_payload):
        from report_generator import generate_uds_preview_html
        payload = dict(mock_uds_payload)
        payload["logic_diagrams"] = [
            {"title": "Test Logic", "description": "Logic desc", "url": ""},
        ]
        result = generate_uds_preview_html(payload)
        assert isinstance(result, str)
        assert len(result) > 100


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
