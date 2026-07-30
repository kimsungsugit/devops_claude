"""UDS DOCX 생성 충실도 — "파일이 있고 0바이트 아님" 이 성공 판정이던 것.

## 왜 이 테스트가 있나 (실측)

`generate_uds_docx` 는 **템플릿 주도**다: 템플릿의 SwUFn heading 을 순회하며 payload 함수를
찾아 채운다. 따라서

- 템플릿에 heading 이 없는 payload 함수는 **문서에 안 들어간다**
- payload 에 없는 heading 은 **빈 껍데기로 남는다**

둘 다 예전엔 어디에도 보고되지 않았고, 프로덕션 성공 판정은
`returncode == 0 and out_path.exists() and out_path.stat().st_size > 0` **뿐**이었다.

실측(HDPDM01 실 템플릿 430 heading + 실 payload 432 함수):

| 항목 | 값 |
|---|---|
| 문서에 반영된 함수 | 336 (77.8%) |
| 템플릿에 heading 이 없어 **미반영** | 96 (22.2%) |
| 내용 없이 남은 heading | 75 |
| 그중 템플릿이 "삭제" 로 표기한 것 | 10 — **갭 아님** |

즉 함수 22%가 빠진 문서가 `status: "success"` 로 기록됐다.

## 설계 제약 (측정으로 확인)

프로덕션은 `backend/helpers/uds.py` 의 exec 문자열을 **서브프로세스**로 돌리고 반환값을
버린다. 그래서 in-process `stats_out` 만으로는 호출자에 닿지 않고 **파일 sidecar 가 필수**다.
"""
from __future__ import annotations

import json
import re

import pytest

pytest.importorskip("docx")


@pytest.fixture(autouse=True)
def _no_reference_suds(monkeypatch, tmp_path):
    """생성기가 저장소 고정 참조 SUDS(40.7MB)를 읽는 것을 막는다.

    실측: `generate_uds_docx` 31.9초 중 **24.3초**가
    `requirements.py::_extract_function_info_from_docx` 였고, 입력은
    `config.UDS_REF_SUDS_PATH` = `docs\\(HDPDM01_SUDS) ..._v1.07_240213.docx` 였다.
    함수 2개짜리 payload 인데도 `table.text` 를 64,312회 부른다.

    테스트가 이걸 읽으면 (a) 3분 이상 걸려 스위트 예산을 먹고 (b) **다른 프로젝트 문서의
    값이 섞여** 계측 단정이 흔들린다. 둘 다 테스트가 의존할 이유가 없다.
    """
    import config
    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_reference.docx"), raising=False)
    # `template_path=None` 은 "템플릿 없음" 이 아니라 `resolve_uds_template_path()` 로
    # **저장소 기본 템플릿(430 heading)** 을 끌어온다. 실측: 그 경로를 탄 테스트 하나가
    # **313.88초** 였다. 기본값을 비워 두고, 폴백을 검증하는 테스트만 명시적으로 되살린다.
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)


# ==============================================================
# 헬퍼 — 실 템플릿 없이 최소 템플릿을 만든다
# ==============================================================

def _make_template(path, headings):
    """SwUFn heading 만 가진 최소 템플릿. 치환 토큰은 넣지 않는다(구조 복제 경로 진입)."""
    import docx
    d = docx.Document()
    d.add_heading("Software Unit Design", level=1)
    for h in headings:
        d.add_heading(h, level=4)
    d.save(str(path))
    return str(path)


def _payload(names, *, substantive=True):
    fd = {}
    for i, n in enumerate(names, start=1):
        fd[f"SwUFn_{i:04d}"] = {
            "id": f"SwUFn_{i:04d}",
            "name": n,
            "prototype": f"void {n}(void);" if substantive else "",
            "description": "설명" if substantive else "",
            "asil": "QM",
            "related": "",
            "inputs": [],
            "outputs": [],
            "precondition": "",
            "globals_global": [],
            "globals_static": [],
            "called": "",
            "logic": "",
        }
    return {
        "project_name": "T",
        "overview": "o",
        "requirements": "r",
        "interfaces": "i",
        "uds_frames": "u",
        "notes": "n",
        "function_details": fd,
    }


def _gen(tmp_path, template, payload):
    from report_gen.docx_builder import gen_stats_path, generate_uds_docx
    out = tmp_path / "uds.docx"
    stats = {}
    generate_uds_docx(template, payload, str(out), stats_out=stats)
    side = json.loads(gen_stats_path(str(out)).read_text(encoding="utf-8"))
    return out, stats, side


# ==============================================================
# 1. sidecar 는 항상 남는다 (서브프로세스 경계)
# ==============================================================

class TestSidecar:
    def test_sidecar_path_is_derived_from_output(self):
        from report_gen.docx_builder import gen_stats_path
        assert gen_stats_path("/a/b/uds.docx").name == "uds.docx.gen_stats.json"

    def test_template_mode_writes_sidecar(self, tmp_path):
        """in-process stats_out 은 서브프로세스 경계를 못 넘는다 — 파일이 정본이다.

        뮤테이션: `_write_gen_stats(output_path, _stats)` 를 없애면 sidecar 가 없어 실패.
        """
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, side = _gen(tmp_path, tpl, _payload(["alpha"]))
        assert side["mode"] == "template"
        assert side == {**side, **{k: stats[k] for k in ("payload_functions", "matched_functions")}}

    def test_template_none_falls_back_to_config_and_says_so(self, tmp_path, monkeypatch):
        """⚠ `template_path=None` 은 "템플릿 없음" 이 아니다 — `config.resolve_uds_template_path()`
        가 **저장소 기본 템플릿**을 해결해 template 경로를 탄다(admin API 와 합의된 의도된 동작).

        측정에서 이걸 몰랐다: `None` 을 넘겼는데 `mode="template"` 이 나왔다. 의도된 동작이라
        고치지 않되, **호출자가 고른 것인지 기본값인지**는 통계에 남긴다 — 반영률이 낮을 때
        원인이 "프로젝트 템플릿 미지정" 인지 "템플릿 자체" 인지 갈리는 지점이다.

        ⚠ 폴백을 **작은** 템플릿으로 되살려 검증한다 — 실 저장소 템플릿(430 heading)을
        쓰면 이 테스트 하나가 313초를 먹는다(실측).

        뮤테이션: `template_source` 를 상수 `"argument"` 로 바꾸면 실패.
        """
        import config
        fallback_tpl = _make_template(tmp_path / "fallback.docx", ["SwUFn_0001: alpha"])
        monkeypatch.setattr(config, "resolve_uds_template_path",
                            lambda: fallback_tpl, raising=False)
        _out, _stats, side = _gen(tmp_path, None, _payload(["alpha", "beta"]))
        assert side["mode"] == "template"
        assert side["template_source"] == "config_fallback"
        assert side["template_path"] == fallback_tpl, "폴백으로 해결된 경로가 기록돼야 한다"
        # 폴백 템플릿엔 alpha 만 있으므로 beta 는 미반영으로 잡혀야 한다
        assert side["unmatched_payload_count"] == 1

    def test_no_template_mode_when_fallback_yields_nothing(self, tmp_path, monkeypatch):
        """폴백도 못 찾으면 진짜 no_template — 그 경로도 통계를 남긴다.

        통계 부재를 "정상" 으로 오독하지 않게 하는 것이 목적이다.
        """
        import config
        monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)
        _out, stats, side = _gen(tmp_path, None, _payload(["alpha", "beta"]))
        assert side["mode"] == "no_template"
        assert side["template_source"] == "none"
        assert side["unmatched_payload_count"] == 0
        assert side["match_pct"] == 100.0

    def test_explicit_template_is_marked_as_argument(self, tmp_path):
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, _stats, side = _gen(tmp_path, tpl, _payload(["alpha"]))
        assert side["template_source"] == "argument"

    def test_sidecar_failure_does_not_break_generation(self, tmp_path, monkeypatch):
        """통계 기록 실패가 산출물을 막아선 안 된다."""
        from report_gen import docx_builder
        monkeypatch.setattr(docx_builder, "gen_stats_path",
                            lambda _p: tmp_path / "nonexistent_dir" / "x.json")
        out = tmp_path / "uds.docx"
        docx_builder.generate_uds_docx(None, _payload(["alpha"]), str(out))
        assert out.exists() and out.stat().st_size > 0


# ==============================================================
# 2. 미반영 함수 / 빈 heading 계측
# ==============================================================

class TestFidelityCounts:
    def test_payload_function_absent_from_template_is_counted(self, tmp_path):
        """템플릿에 heading 이 없는 함수는 문서에 안 들어간다 — 그걸 세야 한다.

        뮤테이션: `_note_fn_match` 호출을 없애면 matched 가 0 이 되어 실패.
        """
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha", "beta", "gamma"]))
        assert stats["payload_functions"] == 3
        assert stats["matched_functions"] == 1
        assert stats["unmatched_payload_count"] == 2
        assert set(stats["unmatched_payload_sample"]) == {"beta", "gamma"}
        assert stats["match_pct"] == pytest.approx(33.33, abs=0.01)

    def test_full_match_reports_no_gap(self, tmp_path):
        """음성 대조군 — 전부 반영되면 갭이 0. 없으면 항상-경고 코드로도 위 테스트가 통과한다."""
        tpl = _make_template(tmp_path / "t.docx",
                             ["SwUFn_0001: alpha", "SwUFn_0002: beta"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha", "beta"]))
        assert stats["unmatched_payload_count"] == 0
        assert stats["empty_heading_count"] == 0
        assert stats["match_pct"] == 100.0

    def test_boilerplate_only_content_is_not_counted_as_reflected(self, tmp_path):
        """이름은 맞아도 내용이 전부 생성기 합성이면 "반영" 이 아니다 — 세면 반영률이 부풀린다.

        실측: `_finalize_function_fields` 는 내용이 완전히 빈 함수에도
        `description="alpha은(는) alpha 관련 연산을 수행하고…"` / `asil="QM"` /
        `related="TBD"` 를 채우고 `description_source="inference"` 를 남긴다. 그래서 텍스트
        모양이 아니라 **provenance** 로 판별한다.

        뮤테이션: `description` 을 hard-content 목록에 되돌리면 matched 가 1 이 되어 실패.
        """
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha"], substantive=False))
        assert stats["matched_functions"] == 0
        assert stats["boilerplate_only_count"] == 1
        assert stats["empty_heading_count"] == 0, "합성만인 것은 '내용 없음' 과 원인이 다르다"

    def test_description_alone_does_not_count_as_reflected(self, tmp_path):
        """⚠ 설명만 있는 함수는 **반영으로 세지 않는다** — 의도된 결정이다.

        측정에서 드러난 이유: 설명이 진짜인지 합성인지 판별할 수단이 둘 다 못 쓴다.
          · `description_source` — `_resolve_related_asil_desc` 가 **출처 미기록을 전부
            `"inference"` 로 확정**한다. 실측: 사람이 쓴 설명을 넣고 생성했더니
            생성 후 payload 엔트리가 `description_source='inference'` 였다(별도 결함).
          · `_is_generic_description` — 합성기 자신의 출력을 generic 으로 보지 않는다.

        고장난 판정을 지표에서 흉내내면 결함이 복제되므로, 생성기가 만들지 않는 필드
        (prototype/inputs/outputs/logic)만 근거로 삼는다. 단위 상세 설계 문서에서
        prototype·I/O 가 전무한 항목을 "반영됨" 으로 세지 않는 건 정당하다.

        뮤테이션: `description` 을 hard-content 목록에 넣으면 matched 가 1 이 되어 실패.
        """
        payload = _payload(["alpha"], substantive=False)
        payload["function_details"]["SwUFn_0001"]["description"] = "사람이 직접 쓴 상세 설명"
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, _side = _gen(tmp_path, tpl, payload)
        assert stats["matched_functions"] == 0
        assert stats["boilerplate_only_count"] == 1

    def test_prototype_alone_counts_as_reflected(self, tmp_path):
        """음성 대조군 — 생성기가 만들지 않는 필드가 하나라도 있으면 반영이다."""
        payload = _payload(["alpha"], substantive=False)
        payload["function_details"]["SwUFn_0001"]["prototype"] = "void alpha(u8 x);"
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, _side = _gen(tmp_path, tpl, payload)
        assert stats["matched_functions"] == 1
        assert stats["boilerplate_only_count"] == 0

    def test_template_only_heading_is_counted_empty(self, tmp_path):
        tpl = _make_template(tmp_path / "t.docx",
                             ["SwUFn_0001: alpha", "SwUFn_0002: ghost"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha"]))
        assert stats["matched_functions"] == 1
        assert stats["empty_heading_count"] == 1
        assert any("ghost" in h for h in stats["empty_heading_sample"])

    def test_zero_payload_functions_reports_none_not_zero_pct(self, tmp_path):
        """분모 0 은 "미측정" 이다 — 0% 로 접으면 "전부 실패" 로 오독된다."""
        tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: alpha"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload([]))
        assert stats["payload_functions"] == 0
        assert stats["match_pct"] is None


# ==============================================================
# 3. 삭제 표기 heading 은 갭이 아니다
# ==============================================================

class TestDeletedHeadings:
    @pytest.mark.parametrize("marker", ["(삭제)", "(New, 삭제)", "(제거)", "(deleted)"])
    def test_deleted_marker_headings_are_separated(self, tmp_path, marker):
        """템플릿이 "삭제" 로 표기한 heading 은 비어 있는 게 정상 — 갭에 섞으면 오탐.

        실측: HDPDM01 템플릿 430 heading 중 `(삭제)` 9건 + `(New, 삭제)` 1건 = 10건.

        뮤테이션: `elif _deleted_marker.search(...)` 분기를 없애면 empty_heading_count 가
        늘고 deleted_heading_count 가 0 이 되어 실패.
        """
        tpl = _make_template(tmp_path / "t.docx",
                             ["SwUFn_0001: alpha", f"SwUFn_0002: gone{marker}"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha"]))
        assert stats["deleted_heading_count"] == 1
        assert stats["empty_heading_count"] == 0
        assert any(marker in h for h in stats["deleted_heading_sample"])

    def test_new_marker_is_not_treated_as_deleted(self, tmp_path):
        """`(New)` 는 삭제가 아니다 — 실측에서 47+10건 존재하므로 오분류하면 갭을 대량 은폐한다."""
        tpl = _make_template(tmp_path / "t.docx",
                             ["SwUFn_0001: alpha", "SwUFn_0002: fresh(New)"])
        _out, stats, _side = _gen(tmp_path, tpl, _payload(["alpha"]))
        assert stats["deleted_heading_count"] == 0
        assert stats["empty_heading_count"] == 1

    def test_deleted_marker_regex_matches_measured_vocabulary(self):
        """실측된 어휘를 회귀로 고정 — 템플릿 괄호 전수조사 결과는 (New)/(NEW)/(삭제) 뿐이었다."""
        pat = re.compile(r"\([^)]*(?:삭제|제거|delete[d]?)[^)]*\)", re.I)
        assert pat.search("SwUFn_0206: s_AdcValidationChk(삭제)")
        assert pat.search("SwUFn_0001: x(New, 삭제)")
        assert pat.search("SwUFn_0002: y (삭제)")
        assert not pat.search("SwUFn_0003: z(New)")
        assert not pat.search("SwUFn_0004: w")


# ==============================================================
# 4. 절단된 sample 로 총량을 되짚지 않는다
# ==============================================================

def test_counts_are_computed_before_the_sample_cap(tmp_path):
    """sample 은 50건으로 잘리지만 총량은 캡 **전**에 센다.

    이 저장소가 반복해 겪은 함정이다("절단을 소비처에서 길이로 되짚지 말 것").

    뮤테이션: `unmatched_payload_count` 를 `len(sample)` 로 바꾸면 50 이 되어 실패.
    """
    from report_gen.docx_builder import _STAT_SAMPLE_CAP
    names = [f"fn{i:03d}" for i in range(_STAT_SAMPLE_CAP + 30)]
    tpl = _make_template(tmp_path / "t.docx", ["SwUFn_0001: fn000"])
    _out, stats, _side = _gen(tmp_path, tpl, _payload(names))
    assert stats["unmatched_payload_count"] == len(names) - 1
    assert len(stats["unmatched_payload_sample"]) == _STAT_SAMPLE_CAP
    assert stats["unmatched_payload_count"] > len(stats["unmatched_payload_sample"])


# ==============================================================
# 5. 프로덕션 호출자가 sidecar 를 읽는다
# ==============================================================

class TestCallerReadsStats:
    def test_missing_sidecar_is_reported_as_unmeasured(self, tmp_path):
        """sidecar 부재를 `{}` 로 받고 호출자가 "미측정" 으로 명시해야 한다 — 침묵 금지."""
        from backend.helpers.uds import _read_gen_stats
        assert _read_gen_stats(tmp_path / "no_such.docx") == {}

    def test_corrupt_sidecar_is_reported_as_unmeasured(self, tmp_path):
        from backend.helpers.uds import _read_gen_stats
        from report_gen.docx_builder import gen_stats_path
        out = tmp_path / "uds.docx"
        gen_stats_path(str(out)).write_text("{ not json", encoding="utf-8")
        assert _read_gen_stats(out) == {}

    def test_non_dict_sidecar_is_rejected(self, tmp_path):
        """JSON 이지만 dict 가 아니면 `.get()` 이 터진다 — shape 을 검사해야 한다."""
        from backend.helpers.uds import _read_gen_stats
        from report_gen.docx_builder import gen_stats_path
        out = tmp_path / "uds.docx"
        gen_stats_path(str(out)).write_text("[1, 2, 3]", encoding="utf-8")
        assert _read_gen_stats(out) == {}

    def test_valid_sidecar_is_read(self, tmp_path):
        from backend.helpers.uds import _read_gen_stats
        from report_gen.docx_builder import gen_stats_path
        out = tmp_path / "uds.docx"
        gen_stats_path(str(out)).write_text(
            json.dumps({"mode": "template", "unmatched_payload_count": 7}), encoding="utf-8")
        assert _read_gen_stats(out)["unmatched_payload_count"] == 7
