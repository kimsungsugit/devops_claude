"""R29 — 사이드카 채점 함수 안의 결함 4건 (계획서 §2.2 Q-1 · Q-4 · Q-5 · Q-3).

한 함수(`generate_uds_field_quality_gate_report`) 안에 네 형태가 같이 있었다:

  Q-1  `Gates: N / M` 의 M 에 **임계 없는 축 3개**(`direct_*`)가 섞였다. `gate.get(k, 0.0)` 비교라
       구조적으로 절대 실패하지 않는 "공짜 통과" 3건이 분모에 들어가 통과율을 올렸다.
  Q-4  "해당 없음(대상 0)" 과 "못 잼(슬롯을 셀 수 없음)" 이 한 버킷(`None`)이라 입출력이 없는
       문서가 영구 `Gate pass: False`. 그리고 `input_base = max(ok, applicable)` 라 applicable 0
       인데 inputs 5개면 **100%** — 분자가 분모를 끌어올렸다.
  Q-5  `asil_non_tbd = total - tbd` 라 **빈 ASIL/Related 가 통과**로 계상됐다. 같은 축을 quick gate
       는 `_has_meaningful_value`(빈 값·N/A 제외)로 세어 두 리포트가 다른 수를 냈다.
  Q-3  임계 10개가 함수 안 리터럴이라 어디에도 공시되지 않았고, 호출부가 `UDS_QUALITY_GATE_THRESHOLDS`
       (키 `*_min`, 0~100)를 넘겨도 키가 달라 **조용히 무시**됐다. `/api/quality/policy` 는 12키 전부
       "판정에 쓰인다" 고 했는데 판정식은 7+3 이고 2키는 사유 전용이었다.

원칙: 미측정 ≠ 0 ≠ 통과 · 해당 없음 ≠ 미측정 · 판정 어휘는 한 곳 · 임계는 값 불변으로 승격만(§8 #2).
"""
from __future__ import annotations

import contextlib
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import pytest

from tests.unit.test_quality_gate_unmeasured import FN_INFO_BANNER, _doc_with

pytest.importorskip("docx")

_ALL_ZERO = {k: 0.0 for k in (
    "description_fill_rate", "input_fill_rate", "output_fill_rate",
    "globals_global_fill_rate", "globals_static_fill_rate", "called_fill_rate",
    "calling_fill_rate", "asil_non_tbd_rate", "related_non_tbd_rate", "traceability_rate")}

# R29 이전 `validation.py:812-823` 리터럴 — "값 불변" 약속의 기준점
_PRE_R29_LITERALS = {
    "description_fill_rate": 0.70,
    "input_fill_rate": 0.20,
    "output_fill_rate": 0.10,
    "globals_global_fill_rate": 0.35,
    "globals_static_fill_rate": 0.15,
    "called_fill_rate": 0.50,
    "calling_fill_rate": 0.25,
    "asil_non_tbd_rate": 0.30,
    "related_non_tbd_rate": 0.30,
    "traceability_rate": 0.20,
}


@contextlib.contextmanager
def _reloaded_config(env: Dict[str, Optional[str]]):
    """env 를 바꿔 `config` 를 재로드하고, 나갈 때 **env 를 먼저 원복한 뒤** 다시 재로드한다.

    리뷰 W3: `finally: importlib.reload(config)` 가 monkeypatch teardown 보다 먼저 돌면 지워진 env 로
    config 가 재빌드돼 세션 나머지 동안 `UDS_SIDECAR_*` 가 무시된다(실증: 후속 테스트가 0.7 를 봤다).
    """
    from tests.unit._config_reload import reexec_config

    saved = {k: os.environ.get(k) for k in env}
    for k, v in env.items():
        if v is None:
            os.environ.pop(k, None)
        else:
            os.environ[k] = v
    try:
        # ⚠ `importlib.reload` 가 아니다 — sys.path 가 외부 트리로 그림자져 있으면(tools/generate_uds_local.py
        #   import 이후) reload 는 **다른 config.py** 를 실행한다(리뷰 I7 2차 원인). 같은 파일을 재실행한다.
        yield reexec_config()
    finally:
        for k, v in saved.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
        reexec_config()


def _doc_per_fn(tmp: Path, per_fn: Sequence[Sequence[Tuple[str, str]]], name: str = "u.docx") -> Path:
    """함수마다 **다른** 행을 넣는 최소 UDS DOCX (`_doc_with` 는 전 함수가 같은 행이다)."""
    import docx

    tmp.mkdir(parents=True, exist_ok=True)
    d = docx.Document()
    for i, rows in enumerate(per_fn, start=1):
        d.add_heading(f"SwUFn_{i:03d}: fn_{i}", level=2)
        t = d.add_table(rows=len(rows) + 2, cols=3)
        t.rows[0].cells[0].text = FN_INFO_BANNER
        t.rows[1].cells[0].text = "ID"
        t.rows[1].cells[2].text = f"SwUFn_{i:03d}"
        for r, (label, value) in enumerate(rows, start=2):
            t.rows[r].cells[0].text = label
            # 함수명은 유일해야 한다 — 같은 이름은 doc_map 에서 한 항목으로 접힌다
            t.rows[r].cells[2].text = f"fn_{i}" if (label == "Name" and value == "fn") else value
    out = tmp / name
    d.save(str(out))
    return out


def _fn(proto: str, inputs: str = "N/A", outputs: str = "N/A", asil: str = "B",
        related: str = "SwFn_001", desc: str = "무언가를 한다") -> List[Tuple[str, str]]:
    return [("Name", "fn"), ("Prototype", proto), ("Description", desc), ("ASIL", asil),
            ("Related ID", related), ("Input Parameters", inputs), ("Output Parameters", outputs)]


def _run(docx_path: Path, out: Path, thresholds: Any = None) -> Tuple[str, Dict[str, Any]]:
    from report_gen.evidence import read_gate_report
    from report_gen.validation import generate_uds_field_quality_gate_report

    generate_uds_field_quality_gate_report(str(docx_path), str(out), thresholds=thresholds)
    return out.read_text(encoding="utf-8"), read_gate_report(out)


def _section(text: str, title: str) -> List[str]:
    from report_gen.evidence import _sections
    return _sections(text).get(title, [])


def _metric_line(text: str, label: str) -> str:
    return next(ln for ln in text.splitlines() if ln.strip().startswith(f"- {label}:"))


# ==============================================================
# Q-1  임계 없는 축은 분모가 아니다
# ==============================================================

class TestThresholdlessAxesAreNotGates:

    @pytest.fixture(scope="class")
    def measured(self, tmp_path_factory):
        tmp = tmp_path_factory.mktemp("q1")
        d = _doc_per_fn(tmp, [_fn("int fn_1(int a, int *out)", "a", "out"), _fn("int fn_2(int b, int *out)", "b", "out")])
        return _run(d, tmp / "u.quality_gate.md")

    def test_denominator_is_at_most_the_number_of_thresholds(self, measured):
        """예전 13 은 `direct_*` 3축을 공짜 통과로 섞은 값이었다."""
        _, got = measured
        assert got["gates_total"] == 10
        assert got["unmeasured_count"] == 0 and got["not_applicable_count"] == 0

    def test_direct_axes_live_in_the_informational_section(self, measured):
        text, _ = measured
        info = "\n".join(_section(text, "Informational (no threshold)"))
        assert "Direct called fill" in info and "Direct traceability" in info
        metrics = "\n".join(_section(text, "Metrics"))
        assert "Direct" not in metrics, metrics

    def test_informational_values_are_still_reported_not_dropped(self, measured):
        """축을 분모에서 뺐다고 값까지 지우면 leaf/indirect 진단이 사라진다."""
        text, got = measured
        assert "- Informational metrics (no threshold): `3`" in text
        assert "direct_called_fill" in got["metrics"]          # 파서는 절과 무관하게 잡는다

    def test_no_threshold_key_is_left_without_a_metric(self, measured):
        """임계 표에만 있고 재지 않는 키가 있으면 `ignored keys` 에 이름이 적힌다 — 기본은 0."""
        text, _ = measured
        assert "- ignored keys:" not in text, text


# ==============================================================
# Q-4  해당 없음 ≠ 미측정, 분자는 분모 안에서만
# ==============================================================

class TestNotApplicableIsNotUnmeasured:

    def test_all_void_prototypes_are_not_applicable_and_do_not_hold_gate_pass(self, tmp_path):
        """모든 Prototype 을 읽었고 슬롯 있는 함수가 0 → 잴 대상이 없다. `Gate pass` 를 붙들지 않는다."""
        # (R30) payload 없는 문서 자기 대조는 설명 출처가 generated_doc 이라 30자 이하 설명은 Low 로 떨어져
        # "Doxygen 주석" 권고가 붙고 "모두 통과" 산문이 안 나온다 — 여기선 산문을 보려고 긴 설명을 준다.
        long_desc = "이 함수는 센서 원시값을 읽어 보정 계수를 곱한 뒤 전역 버퍼에 저장한다"
        d = _doc_per_fn(tmp_path, [_fn("void fn_1(void)", desc=long_desc), _fn("void fn_2(void)", desc=long_desc)])
        text, got = _run(d, tmp_path / "na.quality_gate.md", thresholds=_ALL_ZERO)
        assert got["not_applicable_count"] == 2
        assert got["unmeasured_count"] == 0
        assert "input_fill_rate" in " ".join(got["not_applicable_gates"])
        assert "해당 없음 — 대상 0" in _metric_line(text, "Input fill")
        assert "- Gate pass: `True`" in text, text[:600]
        assert got["gates_total"] == 8                       # 10 - N/A 2
        assert "해당 없음 2개는 판정 밖" in text

    def test_unreadable_prototype_stays_unmeasured_and_holds_gate_pass(self, tmp_path):
        """Prototype 을 못 읽은 함수가 하나라도 있으면 잴 대상이 있을지 모른다 → 미측정(False 유지)."""
        d = _doc_per_fn(tmp_path, [_fn("void fn_1(void)"), _fn("")])
        text, got = _run(d, tmp_path / "unm.quality_gate.md", thresholds=_ALL_ZERO)
        assert got["unmeasured_count"] == 2
        assert got["not_applicable_count"] == 0
        assert "- Gate pass: `False`" in text
        assert "Prototype 을 읽지 못한 함수 1개" in text

    @pytest.mark.parametrize("placeholder", ["N/A", "-", "TBD", "Std_ReturnType Foo", "uint8 g_Flag",
                                             "Std_ReturnType Foo(uint8 a"])
    def test_placeholder_prototypes_are_not_read_so_not_applicable_never_fires(self, tmp_path, placeholder):
        """리뷰 W1: 비어 있지 않음 ≠ 읽었음. `N/A`·괄호 없음·절단이 `void f(void)` 와 같은 리포트를 내면
        Prototype 칸이 전부 placeholder 인 문서가 `Gate pass: True` 로 최종 게이트(merged_pass)까지 통과한다."""
        d = _doc_per_fn(tmp_path, [_fn(placeholder), _fn(placeholder)])
        text, got = _run(d, tmp_path / "ph.quality_gate.md", thresholds=_ALL_ZERO)
        assert got["not_applicable_count"] == 0, placeholder
        assert got["unmeasured_count"] == 2
        assert "- Gate pass: `False`" in text
        assert got["prototype_unreadable"] == {"count": 2, "total": 2}

    def test_partial_measurement_is_in_the_head_for_the_reader(self, tmp_path):
        """리뷰 W2: applicable > 0 이면 못 읽은 함수는 분모에서 조용히 빠진다 — 그 수를 head 로 올려
        리더·보드가 본다(강등 임계는 정책 미결)."""
        d = _doc_per_fn(tmp_path, [_fn("int fn_1(int a)", inputs="a"), _fn("", inputs="y"), _fn("N/A")])
        text, got = _run(d, tmp_path / "pm.quality_gate.md")
        assert "- Prototype unreadable: `2` / `3`" in text
        assert got["prototype_unreadable"] == {"count": 2, "total": 3}
        assert got["metrics"]["input_fill"]["percent"] == 100.0   # 부분집합 채점은 그대로 — 사실만 드러낸다

    def test_numerator_counts_only_functions_that_have_a_slot(self, tmp_path):
        """예전엔 `max(ok, applicable)` — 슬롯 없는 함수의 입력이 분자에 들어가 100% 를 만들었다."""
        d = _doc_per_fn(tmp_path, [_fn("int fn_1(int a)", inputs="a"), _fn("", inputs="y")])
        text, got = _run(d, tmp_path / "num.quality_gate.md")
        line = _metric_line(text, "Input fill")
        assert "`1` / `1`" in line, line                      # 예전엔 `2` / `2`
        assert got["metrics"]["input_fill"]["percent"] == 100.0
        assert "슬롯을 셀 수 없는 함수 1개" in text

    def test_filled_inputs_without_any_slot_are_not_one_hundred_percent(self, tmp_path):
        """applicable 0 + inputs 채움 = 예전 100%. 지금은 N/A 이고 채운 사실은 note 로 남는다."""
        d = _doc_per_fn(tmp_path, [_fn("void fn_1(void)", inputs="x"), _fn("void fn_2(void)", inputs="y")])
        text, got = _run(d, tmp_path / "na2.quality_gate.md")
        line = _metric_line(text, "Input fill")
        assert "100.0%" not in line, line
        assert got["metrics"]["input_fill"]["percent"] is None
        assert "input_fill_rate" in " ".join(got["not_applicable_gates"])
        assert "슬롯을 셀 수 없는 함수 2개" in text

    def test_empty_document_is_unmeasured_not_not_applicable(self, tmp_path):
        """함수 0 개는 "대상이 없다" 가 아니라 "잰 게 없다" 다 — False 유지(R15 계약 그대로)."""
        import docx
        empty = tmp_path / "empty.docx"
        docx.Document().save(str(empty))
        text, got = _run(empty, tmp_path / "e.quality_gate.md")
        assert got["not_applicable_count"] == 0
        assert "- Gate pass: `False`" in text

    def test_legacy_report_without_the_section_reads_none_not_zero(self, tmp_path):
        from report_gen.evidence import read_gate_report

        p = tmp_path / "old.quality_gate.md"
        p.write_text("# UDS Field Quality Gate Report\n\n- Total functions: `4`\n- Gate pass: `False`\n"
                     "- Gates: `5` / `13` passed\n", encoding="utf-8")
        got = read_gate_report(p)
        assert got["not_applicable_count"] is None
        assert got["not_applicable_gates"] == []
        assert got["threshold_source"] is None
        assert got["ungated_count"] is None and got["ungated_gates"] == []
        assert got["prototype_unreadable"] is None


# ==============================================================
# Q-5  빈 ASIL/Related 는 기재가 아니다
# ==============================================================

class TestBlankIsNotFilled:

    def test_blank_asil_and_na_related_are_not_counted(self, tmp_path):
        d = _doc_with(tmp_path, [("Name", "fn_{i}"), ("Prototype", "int fn_{i}(int a)"),
                                 ("Description", "fn_{i} 는 무언가를 한다"), ("ASIL", ""),
                                 ("Related ID", "N/A"), ("Input Parameters", "a"), ("Output Parameters", "int")])
        text, got = _run(d, tmp_path / "blank.quality_gate.md")
        assert "`0` / `2`" in _metric_line(text, "ASIL non-TBD")        # 예전엔 `2` / `2`
        assert "`0` / `2`" in _metric_line(text, "Related non-TBD")
        assert "- ASIL unfilled: `2` / `2`" in text
        assert "- Related unfilled: `2` / `2`" in text
        assert got["tbd_residual"]["asil_unfilled"]["count"] == 2
        assert got["tbd_residual"]["asil_tbd"]["count"] == 0            # TBD 는 여전히 0
        failed = {f["gate"] for f in got["failed_gates"]}
        assert {"asil_non_tbd_rate", "related_non_tbd_rate"} <= failed

    def test_filled_tbd_and_empty_partition_the_total(self, tmp_path):
        """기재 + TBD + 빈 칸 = 전체 — 어느 하나가 두 번 세어지지 않는다.

        ⚠ 실측(R29): DOCX 추출기(`_extract_function_info_from_docx`)는 ASIL `TBD` 를 **빈 값으로
        접는다** — 그래서 DOCX 경로에서 `ASIL TBD` 잔여는 항상 0 이고 `TBD` 는 `empty` 쪽에 선다.
        (payload 경로에서만 TBD 가 보인다 — R30 Q-2 가 되살릴 축.) 여기서는 분할 등식만 잰다.
        """
        d = _doc_per_fn(tmp_path, [_fn("void a(void)", asil="B"), _fn("void b(void)", asil="TBD"),
                                   _fn("void c(void)", asil="")])
        text, got = _run(d, tmp_path / "part.quality_gate.md")
        assert "`1` / `3`" in _metric_line(text, "ASIL non-TBD")
        tbd = got["tbd_residual"]["asil_tbd"]["count"]
        empty = got["tbd_residual"]["asil_unfilled"]["count"]
        assert 1 + tbd + empty == 3, (tbd, empty)
        assert empty >= 1                                     # 빈 칸은 확실히 empty 다

    @pytest.mark.parametrize("asil_json", ['null', '"None"', '""'])
    def test_payload_null_asil_is_unfilled_not_filled(self, tmp_path, asil_json):
        """리뷰 C1: `str(a if p else b or "")` 는 payload 갈래에서 `None` 을 `"None"` 으로 만들어 기재로
        계상했다 — 이번 라운드가 새로 적은 `ASIL unfilled: 0` 이 거짓을 단언할 뻔했다. `"None"` 문자열은
        메모리 규약("none 은 none")대로 **명시된 값**이라 기재다 — null/빈 값과 갈라 잰다."""
        import json

        d = _doc_per_fn(tmp_path, [_fn("void fn_1(void)", asil="B", related="SwFn_001")])
        payload = {"function_details": {"fn_1": json.loads(
            '{"name": "fn_1", "asil": %s, "related": %s, "description": "x"}' % (asil_json, asil_json))}}
        (tmp_path / "u.docx.payload.full.json").write_text(json.dumps(payload), encoding="utf-8")
        text, got = _run(d, tmp_path / "pl.quality_gate.md")
        filled_expected = 1 if asil_json == '"None"' else 0
        assert f"`{filled_expected}` / `1`" in _metric_line(text, "ASIL non-TBD"), asil_json
        assert f"`{filled_expected}` / `1`" in _metric_line(text, "Related non-TBD"), asil_json
        assert got["tbd_residual"]["asil_unfilled"]["count"] == 1 - filled_expected

    def test_payload_missing_asil_key_is_unfilled(self, tmp_path):
        import json

        d = _doc_per_fn(tmp_path, [_fn("void fn_1(void)", asil="B")])
        payload = {"function_details": {"fn_1": {"name": "fn_1", "description": "x"}}}
        (tmp_path / "u.docx.payload.full.json").write_text(json.dumps(payload), encoding="utf-8")
        text, got = _run(d, tmp_path / "pk.quality_gate.md")
        assert "`0` / `1`" in _metric_line(text, "ASIL non-TBD")
        assert got["tbd_residual"]["asil_unfilled"]["count"] == 1

    def test_related_na_is_not_traceable_even_with_calls(self, tmp_path):
        """`has_related` 도 같은 헬퍼다 — N/A 는 추적성 분자에 들어가지 않는다."""
        d = _doc_per_fn(tmp_path, [
            _fn("void a(void)", related="N/A") + [("Called Function", "g_helper")],
            _fn("void b(void)", related="SwFn_002") + [("Called Function", "g_helper")],
        ])
        text, _ = _run(d, tmp_path / "tr.quality_gate.md")
        assert "`1` / `2`" in _metric_line(text, "Traceability (Related + Supported Call)")
        assert "`2` / `2`" in _metric_line(text, "Called fill (supported)")   # 대조군: 호출은 둘 다 잡혔다

    def test_the_two_reports_share_one_helper(self):
        """판정 어휘 단일 출처 — alias 가 아니라 본문을 다시 쓰면 다시 갈린다."""
        from backend.helpers.common import _has_meaningful_value
        from report_gen.gate_report import has_meaningful_value

        assert _has_meaningful_value is has_meaningful_value

    @pytest.mark.parametrize("value,expected", [
        ("", False), ("   ", False), ("N/A", False), ("n/a", False), ("TBD", False), ("tbd", False),
        ("-", False), (None, False), ([], False), ([" "], False),
        ("B", True), ("QM", True), (["SwFn_001"], True),
        # "none" 은 **명시된 값**이다(메모리 규약 "none 은 none, tbd 면 tbd" — 근거 부재를 등급으로 바꾸지
        # 않되 적힌 값은 지우지 않는다). 같은 파일의 `_filled`(설명 칸)는 "NONE" 을 placeholder 로 보는데
        # 그건 설명 축의 어휘다 — 두 축을 합치지 않는 것이 의도(리뷰 I6).
        ("none", True),
    ])
    def test_helper_semantics(self, value, expected):
        from report_gen.gate_report import has_meaningful_value
        assert has_meaningful_value(value) is expected


# ==============================================================
# Q-3  임계는 공시되고, 넘긴 키는 조용히 버려지지 않는다
# ==============================================================

class TestSidecarThresholdsArePublished:

    def test_config_table_and_builtin_fallback_are_lockstep(self):
        import config
        from report_gen.validation import _SIDECAR_GATE_FALLBACK

        assert set(config.UDS_SIDECAR_GATE_THRESHOLDS) == set(_SIDECAR_GATE_FALLBACK)

    def test_values_are_unchanged_from_the_pre_r29_literals(self):
        """승격은 값 불변이다(§8 #2 — 통일은 별도 결정). env 를 전부 걷어낸 실효값으로 잰다."""
        from backend.routers.quality import _SIDECAR_ENV_NAMES

        with _reloaded_config({env: None for env in _SIDECAR_ENV_NAMES.values()}) as reloaded:
            assert reloaded.UDS_SIDECAR_GATE_THRESHOLDS == _PRE_R29_LITERALS

    def test_env_override_changes_the_effective_value(self):
        from backend.routers.quality import _SIDECAR_ENV_NAMES

        with _reloaded_config({_SIDECAR_ENV_NAMES["description_fill_rate"]: "0.99"}) as reloaded:
            assert reloaded.UDS_SIDECAR_GATE_THRESHOLDS["description_fill_rate"] == 0.99

    def test_reload_helper_leaves_the_shell_env_respected(self):
        """리뷰 W3 의 가드 — env 가 설정된 셸에서 이 파일을 돌린 뒤에도 후속 테스트가 env 를 본다."""
        import config as config_mod
        from backend.routers.quality import _SIDECAR_ENV_NAMES

        env = _SIDECAR_ENV_NAMES["description_fill_rate"]
        with _reloaded_config({env: "0.99"}):
            assert config_mod.UDS_SIDECAR_GATE_THRESHOLDS["description_fill_rate"] == 0.99
            with _reloaded_config({env: None}):
                assert config_mod.UDS_SIDECAR_GATE_THRESHOLDS["description_fill_rate"] == 0.70
            # 안쪽 블록이 나간 뒤 — 셸(바깥) env 0.99 가 다시 실효값이어야 한다
            assert config_mod.UDS_SIDECAR_GATE_THRESHOLDS["description_fill_rate"] == 0.99

    def test_env_name_map_matches_config_keys(self):
        import config
        from backend.routers.quality import _SIDECAR_ENV_NAMES

        assert set(_SIDECAR_ENV_NAMES) == set(config.UDS_SIDECAR_GATE_THRESHOLDS)

    def test_report_names_its_threshold_source(self, tmp_path):
        d = _doc_per_fn(tmp_path, [_fn("int a(int x, int *o)", "x", "o")])
        text, got = _run(d, tmp_path / "src.quality_gate.md")
        import config

        # 리뷰 W5: 어느 config **파일**을 읽었는지까지 — `tools/generate_uds_local.py` 는 sys.path 그림자로
        # 외부 트리의 config.py 를 잡고 거기엔 이 표가 없다. "미로딩" 과 "다른 config" 를 가른다.
        assert got["threshold_source"].startswith("config.UDS_SIDECAR_GATE_THRESHOLDS (")
        assert Path(config.__file__).name in got["threshold_source"]
        assert "- Threshold source: `config.UDS_SIDECAR_GATE_THRESHOLDS (" in text
        assert "- source: `config.UDS_SIDECAR_GATE_THRESHOLDS (" in "\n".join(_section(text, "Thresholds"))

    def test_fallback_names_the_config_file_that_lacked_the_table(self, monkeypatch):
        """폴백 문구가 "config 미로딩" 이면 진단자가 "config 가 없는 환경" 으로 오독한다(리뷰 W5)."""
        import config
        from report_gen.validation import _SIDECAR_GATE_FALLBACK, _sidecar_gate_thresholds

        monkeypatch.delattr(config, "UDS_SIDECAR_GATE_THRESHOLDS")
        table, source = _sidecar_gate_thresholds()
        assert table == _SIDECAR_GATE_FALLBACK
        assert source.startswith("builtin fallback (config=") and "표 없음" in source and "env 조정 무효" in source
        assert "미로딩" not in source

    def test_report_thresholds_are_the_config_values(self, tmp_path, monkeypatch):
        """리터럴이 아니라 실효값 — config 를 바꾸면 리포트의 임계와 판정이 따라온다."""
        import config

        d = _doc_per_fn(tmp_path, [_fn("int a(int x)", "x", "int", desc="a 는 무언가를 한다")])
        monkeypatch.setitem(config.UDS_SIDECAR_GATE_THRESHOLDS, "description_fill_rate", 1.5)
        text, got = _run(d, tmp_path / "eff.quality_gate.md")
        assert "- description_fill_rate: `150.0%`" in text
        assert "description_fill_rate" in {f["gate"] for f in got["failed_gates"]}

    def test_measured_axis_without_a_threshold_blocks_gate_pass_and_is_named(self, tmp_path, monkeypatch):
        """임계 표에서 키가 빠지면 그 축은 판정할 수 없다 — 통과로 접지 않고 이름을 적는다."""
        import config

        table = dict(config.UDS_SIDECAR_GATE_THRESHOLDS)
        table.pop("traceability_rate")
        monkeypatch.setattr(config, "UDS_SIDECAR_GATE_THRESHOLDS", table)
        d = _doc_per_fn(tmp_path, [_fn("int a(int x, int *o)", "x", "o")])
        text, got = _run(d, tmp_path / "ung.quality_gate.md",
                         thresholds={k: 0.0 for k in table})            # 잰 축은 전부 통과시킨다
        assert "## Ungated Metrics" in text
        assert "- Gate pass: `False`" in text
        assert got["gates_total"] == 9
        # 리뷰 W4: "재지 못했다" 는 거짓이다 — 값은 Metrics 에 있다. 없는 건 임계. 리더가 세고 산문이 가른다.
        assert got["ungated_count"] == 1
        assert "traceability_rate" in " ".join(got["ungated_gates"])
        assert "임계가 없어 판정할 수 없습니다" in text
        assert "재지 못했습니다" not in text
        assert "- Ungated metrics: `1`" in text

    def test_caller_override_is_named_and_unknown_keys_are_listed(self, tmp_path):
        d = _doc_per_fn(tmp_path, [_fn("int a(int x)", "x", "int")])
        text, got = _run(d, tmp_path / "ov.quality_gate.md",
                         thresholds={"description_fill_rate": 0.0, "called_min": 95.0})
        assert got["threshold_source"].endswith("caller override: description_fill_rate")
        assert "- ignored keys: `called_min`" in text

    def test_quick_gate_table_passed_by_mistake_is_wholly_ignored_and_said_so(self, tmp_path):
        """호출부가 `UDS_QUALITY_GATE_THRESHOLDS` 를 넘기면 12키 전부 무시 — **판정은 기본과 같고**
        그 사실이 리포트에 적힌다(예전엔 조용히 같았다)."""
        import config

        d = _doc_per_fn(tmp_path, [_fn("int a(int x)", "x", "int"), _fn("", asil="", related="")])
        _, base = _run(d, tmp_path / "base.quality_gate.md")
        text, got = _run(d, tmp_path / "wrong.quality_gate.md", thresholds=config.UDS_QUALITY_GATE_THRESHOLDS)
        assert {f["gate"] for f in got["failed_gates"]} == {f["gate"] for f in base["failed_gates"]}
        ignored = next(ln for ln in text.splitlines() if ln.startswith("- ignored keys:"))
        assert all(k in ignored for k in config.UDS_QUALITY_GATE_THRESHOLDS), ignored
        assert "caller override" not in got["threshold_source"]


class TestPolicyPublishesRoles:

    @pytest.fixture
    def client(self):
        pytest.importorskip("fastapi")
        from fastapi import FastAPI
        from fastapi.testclient import TestClient

        from backend.dependencies.auth import require_user
        from backend.routers import quality

        app = FastAPI()
        app.include_router(quality.router)
        app.dependency_overrides[require_user] = lambda: "tester"
        return TestClient(app)

    def test_quick_gate_keys_carry_their_role(self, client):
        """'적용됨' 표 12키 중 판정식에 있는 건 7 + 신뢰도 3, 2키는 사유 전용 — 전부 "판정에 쓰인다" 가 아니다."""
        tables = {t["key"]: t for t in client.get("/api/quality/policy").json()["tables"]}
        entries = {e["key"]: e for e in tables["UDS_QUALITY_GATE_THRESHOLDS"]["entries"]}
        roles = {k: e["role"] for k, e in entries.items()}
        assert sum(1 for r in roles.values() if r == "gate") == 7
        assert sum(1 for r in roles.values() if r == "confidence_gate") == 3
        assert roles["global_min"] == "reason_only" and roles["static_min"] == "reason_only"
        assert "unknown" not in roles.values()
        assert "판정 7 · 신뢰도 판정 3 · 사유 전용 2" in tables["UDS_QUALITY_GATE_THRESHOLDS"]["status_label"]

    def test_sidecar_table_is_published_with_env_names(self, client):
        import config

        tables = {t["key"]: t for t in client.get("/api/quality/policy").json()["tables"]}
        side = tables["UDS_SIDECAR_GATE_THRESHOLDS"]
        assert side["status"] == "applied" and side["adjustable"] == "env"
        entries = {e["key"]: e for e in side["entries"]}
        assert set(entries) == set(config.UDS_SIDECAR_GATE_THRESHOLDS)
        assert entries["description_fill_rate"]["env_name"] == "UDS_SIDECAR_DESCRIPTION_FILL_RATE"
        assert entries["description_fill_rate"]["value"] == config.UDS_SIDECAR_GATE_THRESHOLDS["description_fill_rate"]

    def test_roles_derive_from_the_judgment_tuples(self):
        """공시가 판정식과 같은 튜플을 읽는다 — 목록을 따로 들면 R29 이전 결함이 재발한다."""
        from backend.helpers.uds import CONFIDENCE_GATE_AXES, QUICK_GATE_AXES
        from backend.routers.quality import _gate_key_roles

        roles = _gate_key_roles()
        assert {t for _, t in QUICK_GATE_AXES} == {k for k, r in roles.items() if r == "gate"}
        assert {t for _, t in CONFIDENCE_GATE_AXES} == {k for k, r in roles.items() if r == "confidence_gate"}


class TestQuickGateAxesAreTheSingleSource:

    def test_axes_match_the_reason_code_lockstep_table(self):
        from backend.helpers.uds import QUICK_GATE_AXES
        from tests.unit.test_quality_threshold_zero import GATED_PAIRS

        assert {(r, t) for r, t in QUICK_GATE_AXES} == {(r, t) for r, t, _ in GATED_PAIRS}

    def _payload(self) -> Dict[str, Any]:
        row = {"name": "f", "description": "does x", "asil": "B", "related": "SwFn_001",
               "inputs": ["a"], "outputs": ["r"], "called": ["g"], "calling": ["h"],
               "globals_global": ["G"], "globals_static": ["S"],
               "description_source": "comment", "asil_source": "comment", "related_source": "comment"}
        return {"function_details_by_name": {"f": row}}

    def test_gate_pass_still_reads_every_axis(self, monkeypatch):
        """리팩터 회귀 가드 — 어느 한 축의 임계만 올려도 `gate_pass` 가 떨어져야 한다."""
        import config
        from backend.helpers.uds import QUICK_GATE_AXES, _compute_quick_quality_gate

        base = {k: 0.0 for k in config.UDS_QUALITY_GATE_THRESHOLDS}
        monkeypatch.setattr(config, "UDS_QUALITY_GATE_THRESHOLDS", base)
        assert _compute_quick_quality_gate(self._payload())["gate_pass"] is True
        for _, tkey in QUICK_GATE_AXES:
            monkeypatch.setattr(config, "UDS_QUALITY_GATE_THRESHOLDS", {**base, tkey: 101.0})
            assert _compute_quick_quality_gate(self._payload())["gate_pass"] is False, tkey

    def test_reason_only_keys_do_not_move_gate_pass(self, monkeypatch):
        """`global_min`/`static_min` 은 판정식에 없다 — 101 로 올려도 통과가 유지된다(evaluator 의 의도)."""
        import config
        from backend.helpers.uds import _compute_quick_quality_gate

        base = {k: 0.0 for k in config.UDS_QUALITY_GATE_THRESHOLDS}
        monkeypatch.setattr(config, "UDS_QUALITY_GATE_THRESHOLDS", {**base, "global_min": 101.0, "static_min": 101.0})
        assert _compute_quick_quality_gate(self._payload())["gate_pass"] is True
