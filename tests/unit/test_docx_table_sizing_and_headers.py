"""소유 표의 **크기**와 **헤더 판정** — 둘 다 템플릿이 정하고 있었다.

## 실측이 드러낸 것 (2026-09-01, KJPDS02_PV 정본 53MB · 표준 템플릿 v0.10)

라운드 9가 "생성기가 채우지 않는 표는 원본을 유지" 를 고쳤다. 그 다음 축, 즉 **채우는**
표는 여전히 템플릿이 두 가지를 정하고 있었다.

### 1. 행수 — 데이터가 아니라 템플릿이 정했다

`_add_blank_table(doc, rows, ...)` 의 `rows` 가 템플릿 표의 행수였다. 그래서 두 방향으로
조용히 틀렸다:

| 템플릿 | 증상 | 실측 |
|---|---|---|
| 표준(작은 표) | 데이터가 **잘린다** | 6개 표에서 1,144행 소실 · `Software Unit Tables` 는 함수 57개 중 **15개만** · 전역변수 366개는 **0개** |
| 정본(찬 표) | 빈 행이 **남는다** | 문서 전체 행의 **28.2%** 가 완전 빈 행 · `Software Unit Tables` 1,037행 중 978행 |

잘림에는 경고가 없었다. 같은 파일의 비-템플릿 경로 11곳은 이미 `max(len(rows)+1, 2)` 로
데이터에 맞춰 잡는다 — 이 경로만 예외였다.

### 2. 2행 헤더 판정 — 예시 행이 헤더로 굳었다

`_extract_template_blocks` 는 2행 헤더를 **부분문자열**로 판정했다. 데이터 행은 긴 설명을
달고 다녀서 `type`·`version`·`reset` 이 값 안에 우연히 들어간다::

    row1 = ['CPU_POWER_ON_RESET', 'U16', '0x01', 'The last reset has been caused by ...']
                            ^^^^^                              ^^^^^

그 행이 헤더로 굳으면 `_add_blank_table` 이 **그대로 다시 써서** 산출물에 실린다. 실측:
정본 6곳 + 표준 1곳, 그중 6건이 산출물에 남았다(`VERSION1`, `u8s_InitiComplet_F`,
`CPU_POWER_ON_RESET`, `APP_FIRMWARE_VERSION_ADDR`, `PT6_MOTOR_SPEED_TCapturedValue`,
`u8g_SysOptCtrl_DeviceType`). 분석 결과와 구별되지 않으므로 **출처 세탁**이다.

## 왜 비대칭인가

1행은 **관대하게**(부분문자열), 2행은 **엄격하게**(셀 완전일치) 본다. 둘 다 엄격하게
하면 `[ Function Information ]` 같은 배너 헤더가 판정을 잃어 **정상 표 1,035개가 회귀**
한다(전수 실측). 묻는 질문이 다르기 때문이다 — 1행은 "헤더이긴 한가", 2행은 "헤더의
둘째 줄인가 아니면 데이터인가".
"""
from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

# 실제 템플릿에서 헤더로 오인됐던 예시 데이터 행 (측정본 그대로)
REAL_EXAMPLE_ROWS = [
    ["VERSION1", "", "‘A’", ""],
    ["u8g_SW_DEFVER_YEAR_FIRST", "U8", "0x32", "Software version year first digit (2)"],
    ["u8s_InitiComplet_F", "U8", "0x00 ~ 0x01", "0x00",
     "초기화 완료 상태 플래그로, u8g_ON일 때 _PTT.Bits.PTT3를 big_RESET으로 설정"],
    ["CPU_POWER_ON_RESET", "U16", "0x01",
     "The last reset has been caused by Low level detection at the RESET pin"],
    ["PT6_MOTOR_SPEED_TCapturedValue", "U16", "0x0000 ~ 0xFFFF",
     "Type of the captured value"],
    ["APP_FIRMWARE_VERSION_ADDR", "UINT32", "0x00FF9DF0",
     "애플리케이션 펌웨어 버전이 저장된 주소 0x00FF9DF0UL."],
]

# 실제 템플릿의 **진짜** 2행 헤더 (Software Unit Tables — 가로 병합이라 1행이 반복된다)
REAL_HEADER_ROW0 = ["Component", "Component", "Function", "Function", "Function"]
REAL_HEADER_ROW1 = ["ID", "Name", "ID", "Name", "Type"]
REAL_BANNER_ROW = ["[ Function Information ]"] * 6


@pytest.fixture(autouse=True)
def _no_reference_suds(monkeypatch, tmp_path):
    """저장소 고정 참조 SUDS(40.7MB)를 읽지 않게 한다 — 이 파일과 무관한 축."""
    import config

    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_reference.docx"), raising=False)
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)


class TestHeaderRuleIsAsymmetric:
    """1행은 관대하게, 2행은 엄격하게 — 묻는 질문이 다르다."""

    def test_a_real_label_row_still_counts_as_a_second_header_row(self):
        from report_gen.docx_builder import _is_header_label_row

        assert _is_header_label_row(REAL_HEADER_ROW1), (
            "진짜 2행 헤더까지 데이터로 보면 Software Unit Tables 가 한 행 밀린다")

    @pytest.mark.parametrize("row", REAL_EXAMPLE_ROWS)
    def test_a_template_example_row_is_not_a_header(self, row):
        from report_gen.docx_builder import _is_header_label_row

        assert not _is_header_label_row(row), (
            f"예시 데이터 행이 헤더로 굳었다 — 산출물에 그대로 실린다: {row[0]!r}")

    def test_a_banner_row_still_counts_as_a_first_header_row(self):
        """1행까지 완전일치로 보면 배너 헤더가 판정을 잃어 1,035개 표가 회귀한다."""
        from report_gen.docx_builder import _looks_like_header_row

        assert _looks_like_header_row(REAL_BANNER_ROW)
        assert _looks_like_header_row(REAL_HEADER_ROW0)

    def test_the_two_rules_are_actually_different(self):
        """둘을 같은 규칙으로 만들면 두 결함 중 하나가 반드시 되살아난다."""
        from report_gen.docx_builder import _is_header_label_row, _looks_like_header_row

        loose_only = [r for r in REAL_EXAMPLE_ROWS
                      if _looks_like_header_row(r) and not _is_header_label_row(r)]
        assert loose_only, (
            "두 판정이 같은 답을 낸다 — 비대칭이 사라지면 예시 행이 다시 헤더가 된다")


def _table_doc(path: Path, header, second):
    """헤더 1행 + `second` 행짜리 표 하나."""
    import docx

    d = docx.Document()
    t = d.add_table(rows=2, cols=len(header))
    for c, v in zip(t.rows[0].cells, header, strict=False):
        c.text = v
    for c, v in zip(t.rows[1].cells, second, strict=False):
        c.text = v
    d.save(str(path))
    return path


class TestHeaderRuleThroughTheExtractor:
    """헬퍼가 맞아도 `_extract_template_blocks` 가 안 쓰면 소용없다."""

    def _header_rows(self, tmp_path, header, second):
        import docx

        from report_gen.docx_builder import _extract_template_blocks

        p = _table_doc(tmp_path / "t.docx", header, second)
        blocks = _extract_template_blocks(docx.Document(str(p)))
        tables = [b for b in blocks if b[0] == "table"]
        assert tables, "표 블록이 안 나왔다"
        return tables[0][1][3]

    def test_an_example_row_leaves_one_header_row(self, tmp_path):
        hdr = self._header_rows(
            tmp_path, ["Macro Name", "Type", "Define", "Description"],
            ["VERSION1", "", "A", ""])
        assert len(hdr) == 1, f"예시 행이 헤더로 굳었다: {hdr}"

    def test_a_genuine_two_row_header_is_kept(self, tmp_path):
        hdr = self._header_rows(tmp_path, REAL_HEADER_ROW0, REAL_HEADER_ROW1)
        assert len(hdr) == 2, f"진짜 2행 헤더가 1행으로 깎였다: {hdr}"


# --------------------------------------------------------------------------
# 생성까지 태워서 재는 축 — 표 크기는 산출물에서만 보인다
# --------------------------------------------------------------------------

def _owned_table_template(path: Path, template_rows: int, example_row=None) -> Path:
    """`Common Macro Definition`(생성기 소유) 표 하나짜리 템플릿."""
    import docx

    d = docx.Document()
    d.add_heading("Common Macro Definition", level=1)
    t = d.add_table(rows=template_rows, cols=4)
    for c, v in zip(t.rows[0].cells,
                    ("Macro Name", "Type", "Define", "Description"), strict=False):
        c.text = v
    if example_row and template_rows >= 2:
        for c, v in zip(t.rows[1].cells, example_row, strict=False):
            c.text = v
    d.save(str(path))
    return path


def _payload(macros):
    return {
        "project_name": "KJPDS02",
        "overview": "o", "requirements": "r", "interfaces": "i",
        "uds_frames": "u", "notes": "n",
        "common_macros": list(macros),
        "function_details": {},
    }


def _generate(tmp_path, template_rows, macros, example_row=None):
    import docx

    from report_gen.docx_builder import generate_uds_docx

    tpl = _owned_table_template(
        tmp_path / "(KJPDS02_SwUDS) tpl.docx", template_rows, example_row)
    out = tmp_path / "out.docx"
    stats: dict = {}
    generate_uds_docx(str(tpl), _payload(macros), str(out), stats_out=stats)
    doc = docx.Document(str(out))
    owned = None
    for t in doc.tables:
        if t.rows and "Macro Name" in [c.text.strip() for c in t.rows[0].cells]:
            owned = t
            break
    return doc, owned, stats


class TestOwnedTableIsSizedByTheData:
    """행수를 템플릿이 정하면 데이터가 잘리거나 빈 행이 남는다."""

    def test_more_data_than_template_rows_is_not_dropped(self, tmp_path):
        """실측: 표준 템플릿이 함수 57개 중 42개를 **경고 없이** 버렸다."""
        macros = [f"PAYLOAD_MACRO_{i:02d}" for i in range(12)]
        _doc, owned, _stats = _generate(tmp_path, template_rows=2, macros=macros)
        assert owned is not None, "소유 표를 못 찾았다"
        cells = {c.text.strip() for r in owned.rows for c in r.cells}
        missing = [m for m in macros if m not in cells]
        assert not missing, f"템플릿 행수(2)에 걸려 {len(missing)}개가 잘렸다: {missing[:4]}"

    def test_fewer_data_than_template_rows_leaves_no_blank_flood(self, tmp_path):
        """실측: 정본으로 만든 문서의 28.2%가 완전 빈 행이었다."""
        _doc, owned, _stats = _generate(
            tmp_path, template_rows=30, macros=["A_MACRO", "B_MACRO"])
        assert owned is not None
        assert len(owned.rows) == 3, (
            f"헤더 1 + 데이터 2 여야 하는데 {len(owned.rows)}행 — 빈 행이 남았다")

    def test_the_template_example_row_never_reaches_the_document(self, tmp_path):
        """헤더 오인 + 그대로 재기록 = 남의 값이 분석 결과처럼 실린다."""
        _doc, owned, _stats = _generate(
            tmp_path, template_rows=2, macros=["PAYLOAD_MACRO"],
            example_row=["CPU_POWER_ON_RESET", "U16", "0x01",
                         "The last reset has been caused by Low level detection"])
        assert owned is not None
        cells = {c.text.strip() for r in owned.rows for c in r.cells}
        assert "CPU_POWER_ON_RESET" not in cells, (
            "템플릿 예시 행이 산출물에 실렸다 — 분석 결과와 구별되지 않는다")
        assert "PAYLOAD_MACRO" in cells, "예시를 지우면서 실데이터까지 지웠다"


class TestTheCorrectionIsCounted:
    """고쳤다고 침묵하면 다음에 같은 자리가 새도 모른다."""

    def test_recovered_rows_are_recorded(self, tmp_path):
        macros = [f"M_{i:02d}" for i in range(12)]
        _doc, _owned, stats = _generate(tmp_path, template_rows=2, macros=macros)
        assert stats.get("table_rows_recovered") == 11, (
            f"템플릿 행수였다면 잘렸을 행 수가 안 남았다: "
            f"{stats.get('table_rows_recovered')!r}")

    def test_trimmed_blank_rows_are_recorded(self, tmp_path):
        _doc, _owned, stats = _generate(
            tmp_path, template_rows=30, macros=["A_MACRO", "B_MACRO"])
        assert stats.get("table_rows_blank_trimmed") == 27, (
            f"템플릿 행수였다면 남았을 빈 행 수가 안 남았다: "
            f"{stats.get('table_rows_blank_trimmed')!r}")

    def test_the_counts_reach_the_api_surface(self, tmp_path):
        """사이드카에만 있고 화이트리스트에 없으면 API 응답에서 **조용히 잘린다**."""
        from backend.helpers.uds import _gen_stats_result_fields
        from report_gen.docx_builder import generate_uds_docx

        tpl = _owned_table_template(tmp_path / "(KJPDS02_SwUDS) t.docx", 2)
        out = tmp_path / "o.docx"
        generate_uds_docx(str(tpl), _payload([f"M_{i}" for i in range(12)]), str(out))
        summary = _gen_stats_result_fields(out)["gen_stats_summary"] or {}
        assert summary.get("table_rows_recovered") == 11
        assert "table_rows_blank_trimmed" in summary


class TestUndersizingIsNeverSilent:
    """호출부는 다 고쳤다. 그래도 **다음** 호출부가 다시 자를 수 있다.

    잘린 행은 산출물 어디에도 안 남으므로, 문서만 보면 "원래 그만큼인 프로젝트" 로
    읽힌다 — 실측으로 함수 57개짜리 프로젝트가 15개짜리로 보였다.
    """

    def test_a_caller_that_undersizes_a_table_is_warned(self, caplog):
        import logging

        import docx

        from report_gen.docx_builder import _add_blank_table

        d = docx.Document()
        with caplog.at_level(logging.WARNING, logger="report_generator"):
            _add_blank_table(d, 2, 2, None, [["H1", "H2"]],
                             [["a", "b"], ["c", "d"], ["e", "f"]])
        msgs = " ".join(r.getMessage() for r in caplog.records)
        assert "3행 중 1행" in msgs, f"조용히 2행을 버렸다: {msgs!r}"

    def test_a_correctly_sized_caller_stays_quiet(self, caplog):
        """맞게 부른 호출까지 경고하면 경고가 소음이 된다."""
        import logging

        import docx

        from report_gen.docx_builder import _add_blank_table

        d = docx.Document()
        with caplog.at_level(logging.WARNING, logger="report_generator"):
            _add_blank_table(d, 4, 2, None, [["H1", "H2"]],
                             [["a", "b"], ["c", "d"], ["e", "f"]])
        assert not [r for r in caplog.records if "행만 기록" in r.getMessage()]


# --------------------------------------------------------------------------
# 표 크기를 고치자 드러난 것 — 귀속 실패 시 전역변수 **전부**를 싣던 폴백
# --------------------------------------------------------------------------

_GLOBAL_HDR = ["Name", "Type", "Value Range", "Reset Value", "Description"]


def _swcom_template(path: Path, *, with_swcom_heading: bool) -> Path:
    """`Global Data > Global variables` 표 하나. SwCom 아래 두거나 문서 레벨로 둔다.

    실측: 표준 템플릿의 전역/정적 표 4개는 **전부** 문서 레벨, 정본의 64개는 **전부**
    SwCom 아래다. 그래서 두 경우가 갈려야 한다.
    """
    import docx

    d = docx.Document()
    if with_swcom_heading:
        d.add_heading("SwCom_02 (DRV In)", level=2)
    d.add_heading("Global Data", level=3)
    d.add_heading("Global variables", level=4)
    t = d.add_table(rows=2, cols=5)
    for c, v in zip(t.rows[0].cells, _GLOBAL_HDR, strict=False):
        c.text = v
    d.save(str(path))
    return path


def _globals_payload():
    """payload 는 `SwCom_01` 만 안다 — 템플릿의 `SwCom_02` 는 귀속 불가다."""
    names = [f"u8g_Var_{i:02d}" for i in range(9)]
    return {
        "project_name": "KJPDS02",
        "overview": "o", "requirements": "r", "interfaces": "i",
        "uds_frames": "u", "notes": "n",
        "function_details": {},
        "function_details_by_name": {
            "adc_enable": {"file": "adc_monitor.c", "globals_global": [], "globals_static": []},
        },
        "function_table_rows": [["SwCom_01", "Generated_Code", "SwUFn_0101",
                                 "ADC_Enable", "I/F"]],
        "global_vars": [[n, "U8", "0x00 ~ 0x01", "0x00", "d"] for n in names],
        "globals_info_map": {n: {"type": "U8", "range": "0x00 ~ 0x01",
                                 "init": "0x00", "desc": "d", "file": "other.c"}
                             for n in names},
    }


def _generate_globals(tmp_path, *, with_swcom_heading: bool):
    import docx

    from report_gen.docx_builder import generate_uds_docx

    tpl = _swcom_template(tmp_path / "(KJPDS02_SwUDS) t.docx",
                          with_swcom_heading=with_swcom_heading)
    out = tmp_path / "out.docx"
    stats: dict = {}
    generate_uds_docx(str(tpl), _globals_payload(), str(out), stats_out=stats)
    doc = docx.Document(str(out))
    table = None
    for t in doc.tables:
        if t.rows and "Value Range" in [c.text.strip() for c in t.rows[0].cells]:
            table = t
            break
    return table, stats


class TestGlobalsAreNotAttributedToForeignComponents:
    """"귀속 못 했으니 전부 싣는다" 는 침묵한 오답이다.

    `swcom_function_files` 는 **payload 의** function_table_rows 로만 만들어진다.
    템플릿이 payload 보다 넓으면(정본은 늘 그렇다) 대부분의 컴포넌트가 폴백으로 와서
    전역변수 전체를 받았다. 표 크기가 템플릿에 묶여 있던 동안엔 그게 1~4행으로 잘려
    "그 컴포넌트의 전역변수" 처럼 보였고, 크기를 데이터에 맞추자 **8,562행**으로 드러났다.
    """

    def test_a_component_we_cannot_attribute_gets_no_globals(self, tmp_path):
        table, _stats = _generate_globals(tmp_path, with_swcom_heading=True)
        assert table is not None, "전역변수 표를 못 찾았다"
        body = [c.text.strip() for r in table.rows[1:] for c in r.cells]
        leaked = [v for v in body if v.startswith("u8g_Var_")]
        assert not leaked, (
            f"귀속 못 한 컴포넌트 표에 이 프로젝트 전역변수가 실렸다: {leaked[:4]}")

    def test_a_document_level_table_still_lists_them_all(self, tmp_path):
        """SwCom 문맥이 없는 표는 전체가 맞다 — 여기까지 비우면 문서가 빈다."""
        table, _stats = _generate_globals(tmp_path, with_swcom_heading=False)
        assert table is not None
        body = [c.text.strip() for r in table.rows[1:] for c in r.cells]
        assert len([v for v in body if v.startswith("u8g_Var_")]) == 9, (
            f"문서 레벨 표까지 비웠다: {body[:6]}")

    def test_the_unattributed_components_are_counted(self, tmp_path):
        """빈 표는 '없다' 로 읽힌다 — 실제로는 '모른다' 이므로 수를 남긴다."""
        _table, stats = _generate_globals(tmp_path, with_swcom_heading=True)
        assert stats.get("swcom_globals_unattributed") == 1, (
            f"귀속 실패 컴포넌트 수가 안 남았다: "
            f"{stats.get('swcom_globals_unattributed')!r}")

    def test_a_document_level_table_is_not_counted_as_unattributed(self, tmp_path):
        _table, stats = _generate_globals(tmp_path, with_swcom_heading=False)
        assert stats.get("swcom_globals_unattributed") == 0, (
            "문서 레벨 표를 귀속 실패로 세면 경고가 소음이 된다")

    def test_the_count_reaches_the_api_surface(self, tmp_path):
        from backend.helpers.uds import _gen_stats_result_fields
        from report_gen.docx_builder import generate_uds_docx

        tpl = _swcom_template(tmp_path / "(KJPDS02_SwUDS) t.docx", with_swcom_heading=True)
        out = tmp_path / "o.docx"
        generate_uds_docx(str(tpl), _globals_payload(), str(out))
        summary = _gen_stats_result_fields(out)["gen_stats_summary"] or {}
        assert summary.get("swcom_globals_unattributed") == 1
