# tests/unit/test_asil_no_fabrication.py
"""근거가 없을 때 **ASIL 을 지어내지 않는다** — 네 사이트 한 세트 (2026-07-31).

## 왜

ASIL 이 비어 있거나 `TBD` 일 때 저장소는 네 곳에서 `QM` 을 써 넣었다:

| 사이트 | 조건 | 붙던 출처 |
|---|---|---|
| `report_gen/requirements.py` | SRS 요구에 매칭됐는데 **그 요구에 ASIL 이 없음** | `srs_default_qm` |
| `report_gen/docx_builder.py::_inherit_module_asil` | 모듈 상속도 실패 | `default` |
| `report_gen/function_analyzer.py::_finalize_function_fields` | asil 이 빈 값 | `default` |
| `backend/helpers/uds.py::_enrich_function_quality_fields` | 정규화가 등급을 못 뽑음 | `default` |

ISO 26262 에서 `QM` 은 "안전 요구가 면제된다" 는 **실질 주장**이다. 근거의 부재를
그 주장으로 바꾸면 under-classification 이 된다. `helpers/uds.py` 의 옛 주석은 이걸
*"보수적 기본값"* 이라 불렀지만 방향이 거꾸로다 — QM 은 **최저** 등급이라 보수는
상향이지 하향이 아니다.

실측(2026-07-31, 실 payload 3세트 431함수): `asil_source` 분포는
`sds` 407 / `srs` 18 / **`srs_default_qm` 6**, `asil` 값은 `A` 354 / `QM` 77 로
**빈 값 0건**. 즉 근거 없는 6건이 등급을 가진 것처럼 보였고
`helpers/uds.py:210` 의 `with_asil` 은 431/431(100%)로 만점을 찍고 있었다.

`srs_default_qm` 은 특히 나빴다 — 점수표(`validation.py::_src_aliases`)는 이걸
`default`(0.30, 최약체)로 접는데 `provenance.WEAK_SOURCES` 엔 없어서
`is_weak_source()` 는 **강한 출처**라고 답했다. 그래서 더 나은 근거(주석 `@asil`=1.00,
SDS=0.95)가 와도 이 칸을 덮지 못했다. **최약체가 최강자를 막고 있었다.**

## 이 파일이 잠그는 것

네 사이트는 **한 세트**다. 한 곳만 지우면 하류가 다시 채워 상류 수정이 통째로
no-op 이 된다(실제로 `requirements.py` 만 지우면 `docx_builder` 가 곧바로 되채운다).
그래서 사이트별 테스트를 한 파일에 모으고, 마지막에 AST 계약으로 재도입을 막는다.

⚠ 이 파일의 테스트는 **판정을 복제하지 않는다** — 실제 함수/실제 빌더 루프를 태운다.
   (`test_uds_reference_suds_isolation.py` 가 규칙 복제로 뮤테이션 2건을 놓쳤던 전례)
"""
from __future__ import annotations

import ast
from pathlib import Path

import pytest

_ROOT = Path(__file__).resolve().parents[2]


# ---------------------------------------------------------------------------
# 사이트 3 — report_gen/function_analyzer.py::_finalize_function_fields
# ---------------------------------------------------------------------------

class TestFunctionAnalyzerSite:
    @staticmethod
    def _run(**overrides):
        from report_gen.function_analyzer import _finalize_function_fields

        info = {"name": "Motor_Init", "description": "", "asil": "", "related": "",
                "precondition": "", "inputs": None, "outputs": None,
                "globals_global": None, "globals_static": None, "called": "", "calling": ""}
        info.update(overrides)
        return _finalize_function_fields(info)

    def test_blank_asil_stays_blank(self):
        out = self._run(asil="")
        assert out["asil"] == "", "근거가 없는데 등급을 지어냈다"

    def test_blank_asil_gets_no_invented_source(self):
        out = self._run(asil="")
        assert not str(out.get("asil_source") or "").strip(), (
            "채운 값이 없는데 출처를 적었다 — 없는 근거를 기록한 것이다")

    @pytest.mark.parametrize("token", ["TBD", "N/A", "-"])
    def test_placeholder_token_is_preserved(self, token):
        """"미정"(TBD)과 "아예 없음"(빈칸)은 다른 상태다 — 접으면 구분이 사라진다.

        `_normalize_asil_value` 는 이 토큰들을 빈 문자열로 접는다(계약:
        `tests/unit/test_report_gen.py`). 순위 비교에선 그게 맞지만 **값 칸에
        그대로 대입하면 안 된다**.
        """
        out = self._run(asil=token)
        assert out["asil"] == token

    def test_real_grade_is_still_normalized(self):
        """대조군 — 정규화 자체는 살아 있어야 한다."""
        out = self._run(asil="ASIL-D")
        assert out["asil"] == "D"

    def test_non_safety_placeholders_are_untouched(self):
        """대조군 — `related`/`precondition` 의 자리표시자는 안전 판정이 아니라 서식이다."""
        out = self._run(asil="")
        assert out["related"] == "TBD"
        assert out["precondition"] == "N/A"


# ---------------------------------------------------------------------------
# 사이트 4 — backend/helpers/uds.py::_enrich_function_quality_fields
# ---------------------------------------------------------------------------

class TestQualityFieldsSite:
    @staticmethod
    def _run(info):
        from backend.helpers.uds import _enrich_function_quality_fields

        payload = {"function_details": {"F": dict(info)}, "function_details_by_name": {}}
        _enrich_function_quality_fields(payload)
        return payload["function_details"]["F"]

    def test_blank_asil_is_not_filled(self):
        out = self._run({"name": "f", "asil": ""})
        assert out["asil"] == ""
        assert not str(out.get("asil_source") or "").strip()

    def test_tbd_is_preserved(self):
        out = self._run({"name": "f", "asil": "TBD"})
        assert out["asil"] == "TBD"

    def test_existing_source_is_not_rewritten(self):
        """아무 일도 안 했으면 출처도 그대로다 — 정규화도 승격도 근거가 아니다."""
        out = self._run({"name": "f", "asil": "", "asil_source": "inference"})
        assert out["asil_source"] == "inference"

    def test_real_grade_is_still_normalized(self):
        """대조군 — `"asil d"` → `"D"` 는 계속 동작해야 한다."""
        out = self._run({"name": "f", "asil": "asil d", "asil_source": "inference"})
        assert out["asil"] == "D"
        assert out["asil_source"] == "inference"


# ---------------------------------------------------------------------------
# 사이트 1 — report_gen/requirements.py::enrich_function_details_with_docs
# ---------------------------------------------------------------------------

class TestRequirementsSite:
    """SRS 요구에 매칭됐는데 **그 요구에 ASIL 이 없는** 경우.

    파서(`_build_req_map_from_doc_paths`)만 대체하고 **실제 enrich 루프**를 태운다.
    """

    @staticmethod
    def _run(monkeypatch, *, req_asil, cur_asil="TBD"):
        from report_gen import requirements as rq

        monkeypatch.setattr(
            rq, "_build_req_map_from_doc_paths",
            lambda _paths, texts=None: {"swst_01": {"id": "SwST_01", "asil": req_asil}})
        details = {"F1": {"name": "alpha", "asil": cur_asil, "related": "SwST_01"}}
        rq.enrich_function_details_with_docs(details, [], req_doc_paths=["x.docx"])
        return details["F1"]

    def test_requirement_without_asil_does_not_fabricate_qm(self, monkeypatch):
        out = self._run(monkeypatch, req_asil="")
        assert out["asil"] == "TBD", "SRS 가 침묵했는데 QM 이라고 적었다"
        assert out.get("asil_source") != "srs_default_qm"

    def test_reason_is_recorded_not_lost(self, monkeypatch):
        """왜 미상인지는 잃지 않는다 — SRS 문서에 등급이 없다는 건 상류 결함이다."""
        out = self._run(monkeypatch, req_asil="")
        assert out.get("asil_unresolved") == "srs_req_without_asil"

    def test_diagnostic_is_not_a_provenance_label(self, monkeypatch):
        """진단 키는 `*_source` 가 아니다 — 점수·병합 판정에 참여하면 안 된다."""
        out = self._run(monkeypatch, req_asil="")
        assert "asil_unresolved" in out
        assert not any(k.endswith("_source") and v == "srs_req_without_asil"
                       for k, v in out.items())

    def test_requirement_with_asil_still_applies(self, monkeypatch):
        """대조군 — 실제 등급이 있으면 예전처럼 들어와야 한다."""
        out = self._run(monkeypatch, req_asil="C")
        assert out["asil"] == "C"
        assert out["asil_source"] == "srs"
        assert "asil_unresolved" not in out


# ---------------------------------------------------------------------------
# 사이트 2 — report_gen/docx_builder.py::_inherit_module_asil (실제 빌더 루프)
# ---------------------------------------------------------------------------

@pytest.fixture
def build(tmp_path, monkeypatch):
    """**실제 `generate_uds_docx`** 를 태운다. 규칙 복제가 아니다.

    `_inherit_module_asil` 은 `generate_uds_docx` 안의 중첩 함수라 직접 부를 수 없다.
    """
    import docx

    import config
    from report_gen import docx_builder

    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_ref.docx"), raising=False)
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)

    def _run(functions, *, module_rows=None):
        tpl = tmp_path / "t.docx"
        d = docx.Document()
        d.add_heading("Software Unit Design", level=1)
        for fid, info in functions.items():
            d.add_heading(f"{fid}: {info['name']}", level=4)
        d.save(str(tpl))

        details = {fid: dict(info) for fid, info in functions.items()}
        payload = {"project_name": "PDS64_RD", "overview": "o", "requirements": "r",
                   "interfaces": "i", "uds_frames": "u", "notes": "n",
                   "function_details": details}
        if module_rows:
            payload["function_table_rows"] = module_rows
        docx_builder.generate_uds_docx(str(tpl), payload, str(tmp_path / "out.docx"))
        return details

    return _run


_BASE = {"prototype": "void f(void);", "description": "", "related": "TBD",
         "precondition": "N/A", "inputs": [], "outputs": [],
         "globals_global": [], "globals_static": [], "called": "", "logic": ""}


class TestDocxBuilderSite:
    def test_no_module_evidence_does_not_become_qm(self, build):
        out = build({"SwUFn_0001": dict(_BASE, name="alpha", asil="TBD")})
        assert out["SwUFn_0001"]["asil"] != "QM", (
            "모듈 상속조차 못 찾았는데 QM 이라고 적었다 — 근거의 부재를 등급으로 굳힌다")

    def test_tbd_survives_the_builder(self, build):
        """⚠ 빌더 안에 **다섯 번째 사이트**가 있었다(`docx_builder.py:2109`).

        거기서 `TBD` 를 빈 문자열로 지웠는데, 같은 함수 끝(`:2505`)의
        `UDS TBD residual` 경고가 세는 게 바로 그 값이라 **`asil_tbd` 가 구조적으로
        항상 0** 이었다 — 발화할 수 없는 잔량 카운터.
        """
        out = build({"SwUFn_0001": dict(_BASE, name="alpha", asil="TBD")})
        assert out["SwUFn_0001"]["asil"] == "TBD"

    def test_empty_value_is_labelled_weakest(self, build):
        """대조군 — 값이 없으면 **출처 라벨**은 최약체여야 한다.

        값을 안 지어내는 것과, 없는 값에 정직한 라벨을 붙이는 것은 다른 일이다.
        `default`(0.30)가 사실이고 `inference`(0.60)를 주면 빈 칸이 추론 대접을 받는다.
        """
        from report_gen.provenance import is_weak_source

        out = build({"SwUFn_0001": dict(_BASE, name="alpha", asil="")})
        assert is_weak_source(out["SwUFn_0001"].get("asil_source")), (
            "값이 없는데 강한 출처로 분류됐다 — 더 나은 근거가 와도 못 덮는다")

    def test_module_inheritance_still_works(self, build):
        """대조군 — **실재하는** 근거(같은 모듈의 형제 함수)는 계속 물려받아야 한다.

        이걸 같이 잠그지 않으면 "지어내기 제거" 가 정당한 상속까지 죽인 걸
        아무도 못 본다.
        """
        out = build(
            {"SwUFn_0001": dict(_BASE, name="alpha", asil="C"),
             "SwUFn_0002": dict(_BASE, name="beta", asil="TBD")},
            module_rows=[["SwCom_01", "MotorCtrl", "SwUFn_0001", "alpha"],
                         ["SwCom_01", "MotorCtrl", "SwUFn_0002", "beta"]],
        )
        assert out["SwUFn_0002"]["asil"] == "C"
        assert out["SwUFn_0002"]["asil_source"] == "module_inherit"


# ---------------------------------------------------------------------------
# 재도입 금지 — AST 계약
# ---------------------------------------------------------------------------

_SITES = {
    "report_gen/requirements.py": "srs_default_qm",
    "report_gen/docx_builder.py": "_inherit_module_asil",
    "report_gen/function_analyzer.py": "_finalize_function_fields",
    "backend/helpers/uds.py": "_enrich_function_quality_fields",
}


def _qm_assignments(path: Path) -> list[int]:
    """`... ["asil"] = "QM"` 형태의 **대입**만 센다(문자열·주석 언급은 안 센다)."""
    tree = ast.parse(path.read_text(encoding="utf-8"))
    hits: list[int] = []
    for node in ast.walk(tree):
        if not isinstance(node, ast.Assign):
            continue
        if not (isinstance(node.value, ast.Constant) and node.value.value == "QM"):
            continue
        for tgt in node.targets:
            if (isinstance(tgt, ast.Subscript)
                    and isinstance(tgt.slice, ast.Constant)
                    and tgt.slice.value == "asil"):
                hits.append(node.lineno)
    return hits


@pytest.mark.parametrize("rel", sorted(_SITES))
def test_no_site_reintroduces_the_qm_fill(rel):
    """뮤테이션: 어느 사이트든 `info["asil"] = "QM"` 을 되살리면 실패한다.

    네 사이트는 한 세트다 — 한 곳만 되살아나도 나머지 셋의 수정이 무의미해진다.
    """
    hits = _qm_assignments(_ROOT / rel)
    assert not hits, f"{rel}:{hits} 에서 ASIL 지어내기가 되살아났다 ({_SITES[rel]} 참조)"
