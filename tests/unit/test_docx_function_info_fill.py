# tests/unit/test_docx_function_info_fill.py
"""`_fill_function_info_table` 의 **등가성 + 접근 경로 + 행 종류** 계약.

## 왜 이 테스트가 있나

예전 구현은 바로 윗줄에서 `table.rows[r].cells` 로 행을 이미 해석해 놓고 다음 줄에서
`table.cell(r, 0)` 으로 되돌아갔다. python-docx 의 `table.cell()` 은 호출마다 그리드를
처음부터 훑으므로 행 수에 대해 축이 하나 더 붙는다(프로파일: `table.cell()` 81,510회
→ `get_child_element` 4,390만 회).

바꿀 때 무서운 건 **병합 셀**이다. 이 표는 `_merge_function_info_table` 로 행0 전체와
행1+ 의 `[0-1]`·`[2..cols-1]` 이 병합돼 있어서, "두 접근 경로가 같은 셀을 가리키는가"가
자명하지 않다. 실측으로 확인했다(19행 6열, 실제 병합 모양 재현):

    114칸 전부 `table.cell(r,c)._tc is table.rows[r].cells[c]._tc`
    결과 XML 바이트 동일 · 1,918ms → 791ms (2.42배)

## P2-3 이후 — 행 종류가 셋이 됐다

정본 배치를 따르면서 입력은 `List[List[str]]` 에서 `List[(kind, cells)]` 로 바뀌었다.
`full`(전체폭) / `pair`(라벨+값) / `grid`(6칸 독립) 셋이고, **grid 를 pair 로 잘못 쓰면
파라미터가 통째로 사라진다**. 그래서 이 파일은 두 가지를 함께 고정한다:

1. `full`/`pair` 는 P2-3 **이전 구현과 결과 XML 이 같다**(참조 구현과 직접 비교)
2. `grid` 는 6칸이 각각 독립으로 채워진다

수치와 규칙만 적어 두면 다음 사람이 `table.cell()` 로 되돌리거나 grid 를 접어도
아무도 모른다.
"""
from __future__ import annotations

import pytest

from tests.unit._source_probe import source_of

docx = pytest.importorskip("docx", reason="python-docx 없음")

from report_gen.docx_builder import (  # noqa: E402
    _fill_function_info_table,
    _merge_function_info_table,
)
from report_gen.function_analyzer import (  # noqa: E402
    FN_ROW_FULL,
    FN_ROW_GRID,
    FN_ROW_PAIR,
    PARAM_GRID_HEADER,
)

_COLS, _ROWS = 6, 19          # 정본 Function Information 표 모양


def _build(merged: bool = True, cols: int = _COLS, rows: int = _ROWS):
    d = docx.Document()
    t = d.add_table(rows=rows, cols=cols)
    if merged and cols >= 4:          # `_merge_function_info_table` 는 cols<4 면 아무것도 안 한다
        _merge_function_info_table(t, cols)
    return t


def _data(n: int = _ROWS):
    """P2-3 **이전** 모양의 행 목록 — 참조 구현(`_reference_fill`)의 입력이다."""
    return [[f"L{i}", "", f"V{i}"] for i in range(n)]


def _lay(rows):
    """`_data()` 모양을 새 배치로 옮긴다 — 행0 전체폭 + 나머지 라벨/값(예전과 동일)."""
    out = []
    for idx, r in enumerate(rows):
        if idx == 0:
            out.append((FN_ROW_FULL, [r[0]]))
        else:
            out.append((FN_ROW_PAIR, list(r)))
    return out


def _reference_fill(table, data_rows):
    """`table.cell()` 을 쓰던 **옛 구현**. 등가성 판정의 기준선이다."""
    for r_idx, row in enumerate(data_rows):
        if r_idx >= len(table.rows):
            break
        label = row[0] if len(row) > 0 else ""
        value = row[2] if len(row) > 2 else (row[1] if len(row) > 1 else "")
        for c in table.rows[r_idx].cells:
            c.text = ""
        if r_idx == 0:
            for c_idx in range(len(table.rows[r_idx].cells)):
                table.cell(0, c_idx).text = label
        else:
            table.cell(r_idx, 0).text = label
            table.cell(r_idx, 2).text = value


class TestEquivalence:
    @pytest.mark.parametrize("merged", [True, False], ids=["병합표", "비병합표"])
    def test_xml_matches_the_old_implementation(self, merged):
        """⚠ 이게 핵심 — 병합 여부와 무관하게 **결과 XML 이 같아야** 한다."""
        a, b = _build(merged), _build(merged)
        _reference_fill(a, _data())
        _fill_function_info_table(b, _lay(_data()))
        assert a._tbl.xml == b._tbl.xml

    @pytest.mark.parametrize("cols", [2, 3, 4, 6], ids=lambda c: f"{c}열")
    def test_equivalent_on_narrow_tables_too(self, cols):
        """좁은 표에서 **옛 구현은 예외로 중단**한다(`table.cell(r, 2)` 가 범위를 넘음).
        새 구현은 `len(cells) > 2` 로 건너뛴다 — 경로가 다른데 결과가 같은지 확인한다.

        실측: 2·3·4·6열 전부 텍스트·XML 동일. 2열에서 옛 구현은 IndexError 로 중단되지만,
        중단 직전 행이 이미 클리어돼 있어 최종 상태가 같아진다. **우연히 같은 것이라
        확인 없이 '등가' 라고 적으면 안 된다** — 그래서 여기서 못박는다.
        """
        a, b = _build(cols=cols, rows=3), _build(cols=cols, rows=3)
        rows = [["L0", "", "V0"], ["L1", "", "V1"], ["L2", "", "V2"]]
        try:
            _reference_fill(a, rows)
        except Exception:                      # noqa: BLE001 - 옛 구현은 여기서 죽는 게 정상
            pass
        _fill_function_info_table(b, _lay(rows))
        assert [[c.text for c in r.cells] for r in a.rows] == \
               [[c.text for c in r.cells] for r in b.rows]
        assert a._tbl.xml == b._tbl.xml

    def test_packed_pair_row_keeps_the_old_value_column(self):
        """⚠ 좁은 표는 라벨/값 쌍을 한 행에 여러 개 접어 넣는다(`[k1,v1,k2,v2]`).
        값 칸을 `cells_text[-1]` 로 고르면 첫 쌍의 라벨에 **마지막 쌍의 값**이 붙는다.
        P2-3 이전 계약은 `row[2]`(=k2) 이므로 그대로여야 한다.

        (행 0 은 전체 병합이라 라벨/값 두 칸이 같은 셀이 된다 — 행 1 에서 본다.)"""
        t = _build()
        _fill_function_info_table(
            t, [(FN_ROW_FULL, ["hdr"]), (FN_ROW_PAIR, ["k1", "v1", "k2", "v2"])])
        assert t.cell(1, 0).text == "k1"
        assert t.cell(1, 2).text == "k2", "접힌 행의 값 칸 선택이 바뀌었다"

    def test_merged_cells_resolve_to_the_same_element(self):
        """두 접근 경로가 같은 `<w:tc>` 를 가리키는지 — 등가성의 근거."""
        t = _build()
        bad = [(r, c) for r in range(len(t.rows))
               for c in range(len(t.rows[r].cells))
               if t.cell(r, c)._tc is not t.rows[r].cells[c]._tc]
        assert bad == [], f"불일치 {len(bad)}칸: {bad[:5]}"


class TestFillBehaviour:
    def test_label_and_value_land_in_the_right_columns(self):
        t = _build()
        _fill_function_info_table(t, _lay(_data()))
        assert t.cell(1, 0).text == "L1"
        assert t.cell(1, 2).text == "V1"

    def test_header_row_is_filled_across(self):
        """행 0 은 전체 병합이라 어느 칸을 봐도 라벨이어야 한다."""
        t = _build()
        _fill_function_info_table(t, _lay(_data()))
        assert {c.text for c in t.rows[0].cells} == {"L0"}

    def test_more_data_rows_than_table_rows_is_truncated_not_crashed(self, caplog):
        """표보다 데이터가 많으면 **조용히 잘려야** 한다 — 예외로 끝나면 안 된다.

        ⚠ 결과만 보면 구분이 안 된다: 행 초과 가드를 빼도 앞 행은 이미 채워져 있고
        `trows[r_idx]` 의 IndexError 는 `except` 가 삼킨다(뮤테이션 M7 이 그렇게
        살아남았다). **경고가 안 뜨는 것**이 유일한 관측량이다.
        """
        t = _build()
        with caplog.at_level("WARNING", logger="report_generator"):
            _fill_function_info_table(t, _lay(_data(_ROWS + 40)))
        assert t.cell(_ROWS - 1, 0).text == f"L{_ROWS - 1}"
        assert "채우기 실패" not in caplog.text, (
            "정상 절단인데 예외 경로로 빠졌다 — 행 초과 가드가 사라졌는지 볼 것"
        )

    def test_two_column_row_falls_back_to_index_one(self):
        """`cells` 가 2칸뿐이면 값은 `[1]` 이다(기존 계약). 새 배치의 `pair` 행이 바로 이 모양이다."""
        t = _build()
        _fill_function_info_table(t, [(FN_ROW_PAIR, ["L0", "V0"]), (FN_ROW_PAIR, ["L1", "V1"])])
        assert t.cell(1, 2).text == "V1"

    @pytest.mark.parametrize("bad", [None, []], ids=["표없음", "데이터없음"])
    def test_empty_inputs_are_noops(self, bad):
        """음성 대조군 — 빈 입력에 손대면 안 된다."""
        t = _build()
        before = t._tbl.xml
        _fill_function_info_table(
            t if bad is not None else None,
            bad if bad is not None else _lay(_data()),
        )
        if bad is not None:
            assert t._tbl.xml == before


class TestGridRows:
    """P2-3 — 파라미터 그리드는 **6칸이 각각 독립**이다."""

    def _grid_table(self, cols: int = _COLS):
        d = docx.Document()
        t = d.add_table(rows=3, cols=cols)
        layout = [
            (FN_ROW_FULL, ["[ Input Parameters ]"]),
            (FN_ROW_GRID, list(PARAM_GRID_HEADER)),
            (FN_ROW_GRID, ["1", "u8s_Flag", "U8", "0 ~ 255", "0x00", "설명"]),
        ]
        _merge_function_info_table(t, cols, layout)
        _fill_function_info_table(t, layout)
        return t

    def test_every_grid_cell_is_written(self):
        t = self._grid_table()
        assert [c.text for c in t.rows[1].cells] == list(PARAM_GRID_HEADER)
        assert [c.text for c in t.rows[2].cells] == ["1", "u8s_Flag", "U8", "0 ~ 255", "0x00", "설명"]

    def test_grid_row_is_not_merged_into_label_value(self):
        """⚠ 이게 무너지면 파라미터가 통째로 사라진다 — 정본 6칸이 2칸으로 접힌다."""
        t = self._grid_table()
        distinct = len({id(c._tc) for c in t.rows[2].cells})
        assert distinct == _COLS, f"그리드 행이 {distinct}칸으로 접혔다"

    def test_section_header_row_spans_the_whole_width(self):
        t = self._grid_table()
        assert {c.text for c in t.rows[0].cells} == {"[ Input Parameters ]"}
        assert len({id(c._tc) for c in t.rows[0].cells}) == 1

    def test_seventh_column_is_folded_into_the_last_grid_cell(self):
        """실측 템플릿 415개 중 124개가 7열이다. 정본 그리드는 6칸이므로 꼬리를 합친다."""
        t = self._grid_table(cols=7)
        assert len({id(c._tc) for c in t.rows[2].cells}) == 6
        assert t.rows[2].cells[5].text == "설명"
        assert t.rows[2].cells[6].text == "설명", "꼬리 칸이 마지막 그리드 칸과 병합되지 않았다"


class TestAccessPathContract:
    def test_source_does_not_call_table_cell(self):
        """⚠ `table.cell(` 이 다시 들어오면 O(n²) 축이 되살아난다. 문서/주석이 아니라
        **실제 호출**만 잡도록 소스에서 확인한다."""
        src = source_of(_fill_function_info_table)
        body = src.split('"""', 2)[-1]          # docstring 제외(거기엔 설명으로 등장한다)
        assert "table.cell(" not in body, "table.cell() 로 되돌아갔다"

    @pytest.mark.parametrize("n", [5, 19, 60], ids=lambda n: f"{n}행")
    def test_rows_are_never_indexed_through_Rows_getitem(self, monkeypatch, n):
        """⚠ 진짜 O(n²) 축은 `_Rows.__getitem__` 이다 — 구현이 `list(self)[idx]` 라
        **인덱싱할 때마다 전 행을 새로 materialize** 한다(`Table.rows` 는 lazyproperty
        라 캐시되지만 그건 도움이 안 된다). 행 수가 늘어도 **0회**여야 한다.

        `table.rows` 접근 횟수를 세는 건 소용없다 — lazyproperty 라 인스턴스
        `__dict__` 에 캐시돼 클래스 디스크립터를 아예 안 거친다(실측 0회).
        """
        from docx.table import _Rows

        calls = {"n": 0}
        real = _Rows.__getitem__

        def _counting(self, idx):
            calls["n"] += 1
            return real(self, idx)

        monkeypatch.setattr(_Rows, "__getitem__", _counting, raising=True)
        t = _build()
        calls["n"] = 0                     # 표 구성(_merge…)이 쓴 건 세지 않는다
        _fill_function_info_table(t, _lay(_data(n)))
        assert calls["n"] == 0, (
            f"{n}행에서 _Rows.__getitem__ 을 {calls['n']}회 불렀다 — "
            "`list(table.rows)` 로 한 번만 펼칠 것"
        )


class TestFailureIsNotSilent:
    def test_exception_is_logged(self, caplog):
        """예전엔 `except Exception: pass` 라 표가 통째로 비어도 흔적이 없었다."""
        class _Boom:
            @property
            def rows(self):
                raise RuntimeError("boom")

        with caplog.at_level("WARNING", logger="report_generator"):
            _fill_function_info_table(_Boom(), _lay(_data()))
        assert "함수 정보 표 채우기 실패" in caplog.text, caplog.text
        assert "RuntimeError" in caplog.text, "예외 종류를 안 남기면 원인을 못 짚는다"
