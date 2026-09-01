# tests/unit/test_uds_unmatched_headings.py
"""정본에만 있는 **남의 함수 절**을 남길지 지울지 — 그 선택이 산출물을 실제로 바꾸는가.

## 왜 이 축이 생겼나 (라운드 10 실측)

정본을 템플릿으로 쓰면 heading 이 그 문서 전체만큼 온다(KJPDS02 정본 1,035개). 이번
분석의 payload 는 부분집합(57개)이라 나머지는 빈 `[ Function Information ]` 서식으로
남는다 — **문서 행의 23%**. 세고는 있었지만(`empty_heading_count`) 지우지는 않았고,
"정본을 부분집합으로 쓰는 것이 의도일 수 있어" **판단 축**으로 남겨 뒀다.

## 규약

- 기본은 `keep`(= 종전 동작). **고르기 전까지 산출물은 바뀌지 않는다.**
- 모르는 값은 **안 지우는 쪽**으로 떨어진다. 반대로 두면 오타 하나에 문서가 얇아진다.
- `(삭제)` 표기 heading 은 어느 쪽이든 남는다 — 템플릿이 의도해서 비운 자리다.
- **지운 것과 비워 둔 것은 다른 사실이다.** 두 수의 합은 보존된다(같은 집합의 재분류).
- 지울 땐 heading 만이 아니라 **절 전체**를 지운다 — 표만 남으면 내용이 엉뚱한 절에 붙는다.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Tuple

import pytest

pytest.importorskip("docx", reason="python-docx 없음")

from report_gen.docx_builder import (  # noqa: E402
    UNMATCHED_HEADINGS_DROP,
    UNMATCHED_HEADINGS_KEEP,
    normalize_unmatched_headings,
)

FN_INFO_BANNER = "[ Function Information ]"


SIDE_TABLE_HEADER = "Side Note"


def _template(path: Path) -> Path:
    """SwUFn 절 3개짜리 정본형 템플릿.

    - `alpha`   : payload 에 있다 → 반영
    - `bravo`   : payload 에 없다 → 빈 서식(= 지울 후보)
    - `charlie` : payload 에 없지만 템플릿이 `(삭제)` 로 표기 → 어느 쪽이든 남는다

    ⚠ `bravo` 절에는 **Function Information 이 아닌 표**를 하나 더 둔다. 그게 없으면
      "절 전체를 지우는가" 를 잴 수 없다 — FI 표는 heading 분기 밖에서 어차피
      `continue` 되므로 heading 만 지워도 사라지고, 그래서 `_drop_until_level` 을
      0 으로 죽이는 뮤턴트가 살아남았다(라운드 12 실측 M186).
    ⚠ 문단은 관측량이 못 된다 — 라이터가 heading 아래 para 블록을 원래 버린다.
    """
    import docx

    path.parent.mkdir(parents=True, exist_ok=True)
    d = docx.Document()
    d.add_heading("Software Unit Design", level=1)
    for title in ("SwUFn_001: alpha", "SwUFn_002: bravo", "SwUFn_003: charlie (삭제)"):
        d.add_heading(title, level=2)
        t = d.add_table(rows=4, cols=6)
        t.rows[0].cells[0].text = FN_INFO_BANNER
        t.rows[1].cells[0].text = "Function Name"
        if "bravo" in title:
            side = d.add_table(rows=2, cols=2)
            side.rows[0].cells[0].text = SIDE_TABLE_HEADER
            side.rows[0].cells[1].text = "Value"
            side.rows[1].cells[0].text = "남의 프로젝트 메모"
    d.save(str(path))
    return path


def _payload(mode: Any) -> Dict[str, Any]:
    alpha = {
        "name": "alpha", "prototype": "void alpha(void)",
        "inputs": ["none"], "outputs": ["none"], "logic": "does alpha",
    }
    return {
        "project_name": "KJPDS02",
        "overview": "o", "requirements": "r", "interfaces": "i",
        "uds_frames": "u", "notes": "n",
        "function_details": {"SwUFn_001": alpha},
        "function_details_by_name": {"alpha": alpha},
        "unmatched_headings": mode,
    }


def _generate(tmp_path: Path, mode: Any) -> Tuple[List[str], Dict[str, Any], int, int]:
    """`(heading, gen_stats, FI 표 수, 곁표 수)` — 한 번의 생성으로."""
    import docx

    from report_gen.docx_builder import generate_uds_docx

    tpl = _template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
    out = tmp_path / f"out_{str(mode) or 'default'}.docx"
    stats: Dict[str, Any] = {}
    generate_uds_docx(str(tpl), _payload(mode), str(out), stats_out=stats)
    doc = docx.Document(str(out))
    heads = [p.text.strip() for p in doc.paragraphs
             if (p.style.name or "").lower().startswith("heading")]
    tables = sum(1 for t in doc.tables if t.rows
                 and any(FN_INFO_BANNER in c.text for c in t.rows[0].cells))
    sides = sum(1 for t in doc.tables if t.rows
                and any(SIDE_TABLE_HEADER in c.text for c in t.rows[0].cells))
    return heads, stats, tables, sides


# ⚠ 한 번 생성에 수 초가 걸린다. 기본값·keep·drop·오타 네 갈래만 만들고 **재사용**한다
#   (테스트마다 생성하면 이 파일 하나가 80초를 먹는다 — 실측).
_MODES = ("", UNMATCHED_HEADINGS_KEEP, UNMATCHED_HEADINGS_DROP, "dropp")


@pytest.fixture(scope="module")
def runs(tmp_path_factory: pytest.TempPathFactory) -> Dict[str, Any]:
    base = tmp_path_factory.mktemp("unmatched")
    return {m: _generate(base / (m or "default"), m) for m in _MODES}


# ══════════════════════════════════════════════════════════════════════════
# 정규화 — 모르는 값은 **안 지우는 쪽**
# ══════════════════════════════════════════════════════════════════════════

class TestNormalization:
    @pytest.mark.parametrize("raw,want", [
        ("", UNMATCHED_HEADINGS_KEEP), (None, UNMATCHED_HEADINGS_KEEP),
        ("keep", UNMATCHED_HEADINGS_KEEP), ("KEEP", UNMATCHED_HEADINGS_KEEP),
        ("drop", UNMATCHED_HEADINGS_DROP), (" Drop ", UNMATCHED_HEADINGS_DROP),
    ])
    def test_known_values(self, raw: Any, want: str) -> None:
        assert normalize_unmatched_headings(raw)[0] == want

    @pytest.mark.parametrize("raw", ["dropp", "delete", "remove", "0", "true", "지움"])
    def test_unknown_falls_to_the_safe_side_and_says_so(self, raw: str) -> None:
        """`drop` 쪽으로 떨어지면 오타 하나가 문서를 조용히 얇게 만든다."""
        got, bad = normalize_unmatched_headings(raw)
        assert got == UNMATCHED_HEADINGS_KEEP
        assert bad == raw.strip(), "알 수 없었던 원본을 돌려주지 않으면 화면이 말할 게 없다"


# ══════════════════════════════════════════════════════════════════════════
# 산출물 — 고른 대로 실제로 달라지는가
# ══════════════════════════════════════════════════════════════════════════

class TestTheDocumentActuallyChanges:
    def test_default_keeps_everything(self, runs) -> None:
        """기본값에서 산출물이 바뀌면 이 변경이 회귀다."""
        heads, stats, _fi, _side = runs[""]
        assert any("bravo" in h for h in heads)
        assert stats["dropped_heading_count"] == 0
        assert stats["unmatched_headings_mode"] == UNMATCHED_HEADINGS_KEEP

    def test_drop_removes_only_the_unmatched_section(self, runs) -> None:
        heads, stats, _fi, _side = runs[UNMATCHED_HEADINGS_DROP]
        assert not any("bravo" in h for h in heads), "지운다고 했는데 남아 있다"
        assert any("alpha" in h for h in heads), "반영된 함수까지 지웠다"
        assert stats["dropped_heading_count"] == 1

    def test_deleted_marked_heading_is_never_dropped(self, runs) -> None:
        """템플릿이 의도해서 비운 자리다 — 빈 것과 지울 것을 섞으면 안 된다."""
        heads, _stats, _fi, _side = runs[UNMATCHED_HEADINGS_DROP]
        assert any("charlie" in h for h in heads)

    def test_the_whole_section_goes_not_just_the_heading(self, runs) -> None:
        """표만 남으면 그 내용이 **앞 절에 붙는다** — heading 만 지우면 더 나쁘다.

        ⚠ FI 표 수로는 못 잰다(heading 만 지워도 같이 사라진다). 절 안의 **다른** 표가
          남는지를 봐야 `_drop_until_level` 이 실제로 절을 닫는지 알 수 있다.
        """
        kept_fi, kept_side = runs[UNMATCHED_HEADINGS_KEEP][2:]
        drop_fi, drop_side = runs[UNMATCHED_HEADINGS_DROP][2:]
        assert kept_side == 1, "픽스처의 곁표가 keep 에서 안 나왔다 — 관측량이 죽었다"
        assert drop_side == 0, "heading 만 지우고 절 본문이 남았다(앞 절에 붙는다)"
        assert drop_fi == kept_fi - 1

    def test_counts_are_conserved_not_invented(self, runs) -> None:
        """지운 것은 **빈 것에서 옮겨 온다** — 합이 늘거나 줄면 한쪽이 거짓이다."""
        keep = runs[UNMATCHED_HEADINGS_KEEP][1]
        drop = runs[UNMATCHED_HEADINGS_DROP][1]
        assert (drop["empty_heading_count"] + drop["dropped_heading_count"]
                == keep["empty_heading_count"] + keep["dropped_heading_count"])

    def test_the_main_metric_does_not_move(self, runs) -> None:
        """반영률은 payload 축이라 템플릿을 얼마나 지우든 불변이어야 한다."""
        keep = runs[UNMATCHED_HEADINGS_KEEP][1]
        drop = runs[UNMATCHED_HEADINGS_DROP][1]
        assert drop["matched_functions"] == keep["matched_functions"] == 1
        assert drop["payload_functions"] == keep["payload_functions"]

    def test_unknown_value_generates_like_keep(self, runs) -> None:
        heads, stats, _fi, _side = runs["dropp"]
        assert any("bravo" in h for h in heads)
        assert stats["dropped_heading_count"] == 0


# ══════════════════════════════════════════════════════════════════════════
# 보고가 도달하는가 — sidecar 에만 있으면 아무도 못 본다
# ══════════════════════════════════════════════════════════════════════════

class TestTheNumberReachesTheSurfaces:
    def test_api_whitelist_carries_it(self, tmp_path: Path) -> None:
        from backend.helpers.uds import _gen_stats_result_fields
        from report_gen.docx_builder import generate_uds_docx

        tpl = _template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
        out = tmp_path / "out.docx"
        generate_uds_docx(str(tpl), _payload(UNMATCHED_HEADINGS_DROP), str(out))
        summary = _gen_stats_result_fields(out)["gen_stats_summary"]
        assert summary is not None
        assert summary["dropped_heading_count"] == 1
        assert summary["unmatched_headings_mode"] == UNMATCHED_HEADINGS_DROP


# ══════════════════════════════════════════════════════════════════════════
# 인자가 **생성기까지 흘러가는가** — 받는 것과 넘기는 것은 다르다
# ══════════════════════════════════════════════════════════════════════════

class TestTheValueReachesTheGenerator:
    """핸들러가 폼을 받는 것만으로는 부족하다.

    ⚠ 라운드 7이 같은 자리에서 배운 것이다 — 인자를 "받는지" 만 보고 "넘기는지" 를
      안 봐서 뮤턴트가 살아남았다. 여기서는 `_uds_generate_from_paths` 를 **실제로
      돌려** 생성기에 넘어가는 payload 를 잡는다(서브프로세스만 가짜다).
    """

    def test_helper_puts_it_in_the_payload(self, tmp_path: Path,
                                           monkeypatch: pytest.MonkeyPatch) -> None:
        import docx

        from backend.helpers import uds as U

        build = tmp_path / "build"
        build.mkdir()
        src = tmp_path / "src"
        src.mkdir()
        (src / "a.c").write_text("int f(void){return 0;}", encoding="utf-8")

        seen: Dict[str, Any] = {}

        def _fake_docx(tpl, payload, out_path, retries=3):  # noqa: ANN001
            seen["payload"] = payload
            # 뒤따르는 리포트들이 이 파일을 되읽으므로 **진짜 docx** 를 남긴다
            # (깨진 바이트를 두면 로그가 traceback 3개로 뒤덮인다).
            Path(out_path).parent.mkdir(parents=True, exist_ok=True)
            docx.Document().save(str(out_path))

        monkeypatch.setattr(U, "_resolve_cached_build_root", lambda *a, **k: build)
        monkeypatch.setattr(U, "_generate_docx_with_retry", _fake_docx)
        # 품질 DB 는 저장소 루트의 실파일이다 — 테스트가 행을 쌓으면 안 된다.
        monkeypatch.setattr(U, "_record_uds_run", lambda *a, **k: None)

        U._uds_generate_from_paths(
            job_url="http://ci/job/x/", cache_root=str(tmp_path), build_selector="",
            template_path="", source_root=str(src), source_only=True,
            req_file_paths=[], note_file_paths=[], logic_file_paths=[], req_paths=[],
            unmatched_headings=UNMATCHED_HEADINGS_DROP,
        )
        assert seen.get("payload", {}).get("unmatched_headings") == UNMATCHED_HEADINGS_DROP, (
            "헬퍼가 값을 받고도 생성기에 안 넘겼다 — 폼은 받는데 문서는 그대로다"
        )


# ══════════════════════════════════════════════════════════════════════════
# 게이트 — 고를 자리와 고를 근거
# ══════════════════════════════════════════════════════════════════════════

class TestTheGateOffersTheChoice:
    def _row(self, caps: Dict[str, Any] | None = None) -> Dict[str, Any]:
        from backend.routers.docgen_preflight import PreflightRequest, _compute_preflight

        res = _compute_preflight(PreflightRequest(doc_type="uds", caps=caps or {}))
        rows = [s for s in res["steps"] if s["id"] == "unmatched_headings"]
        assert rows, "결정 행이 없다"
        return rows[0]

    def test_the_row_exists_with_server_supplied_options(self) -> None:
        row = self._row()
        assert row["phase"] == "decision"
        values = [o["value"] for o in (row["measured"]["options"] or [])]
        assert UNMATCHED_HEADINGS_DROP in values, "화면이 옵션을 지어내야 하는 상태다"

    def test_picking_drop_changes_what_the_row_says(self) -> None:
        keep = self._row()
        drop = self._row({"unmatched_headings": UNMATCHED_HEADINGS_DROP})
        assert keep["measured"]["value"] == UNMATCHED_HEADINGS_KEEP
        assert drop["measured"]["value"] == UNMATCHED_HEADINGS_DROP
        assert "지웁니다" in drop["reason"]
        assert keep["reason"] != drop["reason"]

    def test_unknown_stored_value_is_reported_not_silently_reset(self) -> None:
        row = self._row({"unmatched_headings": "dropp"})
        assert row["state"] == "degraded"
        assert "dropp" in row["reason"]

    def test_without_a_record_it_says_it_has_not_measured(self) -> None:
        """근거 없이 물으면 사용자는 답할 수 없다 — 없으면 없다고 말한다."""
        assert "재지 못했습니다" in self._row()["reason"]

    @pytest.mark.parametrize("record,want", [
        # 직전이 `keep` 이었으면 근거는 "남았다".
        ({"stage": "full", "status": "success",
          "gen_stats": {"payload_functions": 1, "matched_functions": 1,
                        "empty_heading_count": 978, "dropped_heading_count": 0}},
         "978개**가 그렇게 남았습니다"),
        # 직전이 이미 `drop` 이었으면 그 수는 `empty` 가 아니라 `dropped` 에 있다.
        # `empty` 만 보면 "직전엔 0개" = "지울 이유가 없다" 로 읽혀 방금 지운 사실이
        # 다음 판단에서 사라진다.
        ({"stage": "full", "status": "success",
          "gen_stats": {"payload_functions": 1, "matched_functions": 1,
                        "empty_heading_count": 0, "dropped_heading_count": 978,
                        "unmatched_headings_mode": "drop"}},
         "978개**를 지웠습니다"),
    ])
    def test_evidence_comes_from_the_right_column(
        self, tmp_path: Path, record: Dict[str, Any], want: str,
    ) -> None:
        from backend.services import docgen_last_run as lr
        from backend.services.jenkins_helpers import _job_slug

        job = "http://192.168.110.40:7000/job/DEMO_PV/"
        out_dir = tmp_path / "exports"
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / f"{lr.ARTIFACT_PREFIX}{_job_slug(job)}_20260901_120000.docx"
                   f"{lr.CHECKPOINT_SUFFIX}").write_text(
            __import__("json").dumps(record, ensure_ascii=False), encoding="utf-8")

        from backend.routers.docgen_preflight import PreflightRequest, _compute_preflight

        res = _compute_preflight(PreflightRequest(
            doc_type="uds", job_url=job, cache_root=str(tmp_path)))
        row = [s for s in res["steps"] if s["id"] == "unmatched_headings"][0]
        assert want in row["reason"], row["reason"]

    @pytest.mark.parametrize("raw,want", [
        ("", UNMATCHED_HEADINGS_KEEP), ("dropp", UNMATCHED_HEADINGS_KEEP),
        ("drop", UNMATCHED_HEADINGS_DROP), ("keep", UNMATCHED_HEADINGS_KEEP),
    ])
    def test_fallback_keeps_the_same_direction(
        self, monkeypatch: pytest.MonkeyPatch, raw: str, want: str,
    ) -> None:
        """생성기를 못 읽어도 **모르는 값은 안 지우는 쪽**이어야 한다.

        폴백이 반대로 떨어지면, 화면이 뜨게 하려고 둔 안전장치가 오히려 문서를 지운다.
        `_suts_normalize_scope` 가 같은 이유로 같은 규약을 갖는다.
        """
        import report_gen.docx_builder as db
        from backend.routers.docgen_preflight import _uds_normalize_unmatched

        # `from X import Y` 는 Y 가 없으면 ImportError — 폴백 가지를 실제로 태운다.
        monkeypatch.delattr(db, "normalize_unmatched_headings")
        assert _uds_normalize_unmatched(raw)[0] == want

    @pytest.mark.parametrize("doc_type", ["sts", "suts", "sits"])
    def test_other_doc_types_have_no_such_row(self, doc_type: str) -> None:
        """이 축은 UDS 템플릿 주도 라이터에만 있다 — 다른 곳에 내면 거짓 통제다."""
        from backend.routers.docgen_preflight import PreflightRequest, _compute_preflight

        res = _compute_preflight(PreflightRequest(doc_type=doc_type))
        assert not [s for s in res["steps"] if s["id"] == "unmatched_headings"]
