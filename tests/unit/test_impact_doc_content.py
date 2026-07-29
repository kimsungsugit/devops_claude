"""문서 내용 매칭(doc_content) 로더 — 예측 대신 실 파싱 내용을 보존/추출하는지 검증.

라운드: 함수별 상세 문서 카드가 '원본 heading 확인' 같은 예측 대신 실제 UDS/SUTS/SDS 내용을
표시하도록, impact 오케스트레이터 로더가 이미 파싱한 내용을 버리지 않고 반환하는지 확인한다.
"""
from __future__ import annotations


class _FakeResolver:
    def __init__(self, data: bytes = b"docx-bytes"):
        self._data = data

    def read_bytes(self, path: str) -> bytes:  # noqa: ARG002
        return self._data


def test_load_suts_fn_tcs_content_sink_keeps_tc_content(monkeypatch):
    """content_sink 제공 시 TC 실내용(action/precondition/inputs/expected)을 채우고,
    회귀 반환({fn:[tc_id]})은 불변임을 확인."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    fake_model = {
        "units": [{
            "unit_name": "s_foo",
            "test_cases": [{
                "base_tc_id": "SwUTC_SwUFn_0001",
                "description": "Call s_foo with boundary x",
                "precondition": "system initialized",
                "inputs": {"x": 1, "empty": ""},   # empty 값은 스킵돼야
                "expected": {"ret": 42},
            }],
        }],
        "export_warnings": [],
    }
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: fake_model)

    sink: dict = {}
    result = _load_suts_fn_tcs("U:/suts.xlsm", ["s_foo"], content_sink=sink)

    # 회귀 반환 불변
    assert result == {"s_foo": ["SwUTC_SwUFn_0001"]}
    # 내용 보존
    assert "s_foo" in sink
    tc = sink["s_foo"][0]
    assert tc["tc_id"] == "SwUTC_SwUFn_0001"
    assert tc["action"] == "Call s_foo with boundary x"
    assert tc["precondition"] == "system initialized"
    assert tc["inputs"] == {"x": "1"}          # 실입력, empty 스킵, 문자열화
    assert tc["expected"] == {"ret": "42"}     # 실기대값


def test_load_suts_fn_tcs_forwards_source_loc(monkeypatch):
    """content_sink 제공 시 exporter가 test_case["source"]로 부여한 시트·행 위치(loc)를
    전달하는지 — 문서 카드가 'TC(행 N)를 이렇게 수정' 앵커를 표시하게 한다.
    회귀 반환({fn:[tc_id]})은 불변(순수 additive)."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{
            "unit_name": "s_foo",
            "test_cases": [{
                "base_tc_id": "SwUTC_SwUFn_0001",
                "expected": {"ret": 42},
                "source": {"sheet": "2.SW Unit Test Spec", "tc_row": 42, "sequence_row": 43},
            }],
        }],
        "export_warnings": [],
    })

    sink: dict = {}
    result = _load_suts_fn_tcs("U:/suts.xlsm", ["s_foo"], content_sink=sink)

    assert result == {"s_foo": ["SwUTC_SwUFn_0001"]}   # 회귀 반환 불변
    assert sink["s_foo"][0]["loc"] == {
        "sheet": "2.SW Unit Test Spec", "tc_row": 42, "sequence_row": 43,
    }


def test_load_suts_fn_tcs_no_source_omits_loc(monkeypatch):
    """source가 없으면 loc 키를 만들지 않는다(행 번호 날조 금지 — 정직 표기)."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{
            "unit_name": "s_foo",
            "test_cases": [{"base_tc_id": "TC1", "expected": {"ret": 1}}],
        }],
        "export_warnings": [],
    })

    sink: dict = {}
    _load_suts_fn_tcs("U:/s.xlsm", ["s_foo"], content_sink=sink)
    assert sink["s_foo"][0]["tc_id"] == "TC1"
    assert "loc" not in sink["s_foo"][0]


def test_load_suts_fn_tcs_without_sink_unchanged(monkeypatch):
    """content_sink 미제공 시 종전과 동일(회귀 반환만) — 순수 additive 확인."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{"unit_name": "s_bar", "test_cases": [{"base_tc_id": "TC1"}]}],
        "export_warnings": [],
    })
    assert _load_suts_fn_tcs("U:/s.xlsm", ["s_bar"]) == {"s_bar": ["TC1"]}


def test_load_suts_fn_tcs_isolates_malformed_unit(monkeypatch):
    """B2: 기형 유닛(예외 유발) 하나가 전체 회귀집합을 무너뜨리지 않고 격리되며, 나머지 정상 유닛은
    집계되고 사유(N건 예외)가 warn_sink에 표면화된다(silent-0 방지)."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [
            "not-a-dict",  # 기형: dict 아님 → 격리
            {"unit_name": "s_ok", "test_cases": [{"base_tc_id": "TC_OK"}]},  # 정상
        ],
        "export_warnings": [],
    })
    warns: list = []
    result = _load_suts_fn_tcs("U:/s.xlsm", ["s_ok"], warn_sink=warns)
    assert result == {"s_ok": ["TC_OK"]}                        # 정상 유닛은 유실 없이 집계
    assert any("예외로 건너뜀" in w for w in warns)              # 격리 사유 표면화(silent 아님)


def test_load_suts_fn_tcs_happy_path_no_exc_warn(monkeypatch):
    """B2: 모든 유닛 정상이면 예외 warn이 없고 결과는 종전과 byte 동일(happy-path 불변)."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{"unit_name": "s_ok", "test_cases": [{"base_tc_id": "TC_OK"}]}],
        "export_warnings": [],
    })
    warns: list = []
    assert _load_suts_fn_tcs("U:/s.xlsm", ["s_ok"], warn_sink=warns) == {"s_ok": ["TC_OK"]}
    assert not any("예외로 건너뜀" in w for w in warns)


def test_load_suts_fn_tcs_counts_malformed_row_in_good_unit(monkeypatch):
    """reviewer #4: 정상 유닛 내부의 기형 TC 행(dict 아님)은 조용히 걸러졌었다 — 이제 개수를 집계해
    '행 N건 형식 이상' warn으로 표면화한다(정직화 완성). 같은 유닛의 정상 행은 손실 없이 유지."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{"unit_name": "s_ok", "test_cases": [
            {"base_tc_id": "TC_OK"},   # 정상 행 — 유지돼야
            "not-a-dict-row",          # 기형 행 — 격리+집계
        ]}],
        "export_warnings": [],
    })
    warns: list = []
    result = _load_suts_fn_tcs("U:/s.xlsm", ["s_ok"], warn_sink=warns)
    assert result == {"s_ok": ["TC_OK"]}                           # 정상 행은 유실 없음
    assert any("행 1건 형식 이상" in w for w in warns)              # 기형 행 표면화(silent 아님)


def test_load_suts_fn_tcs_warns_only_on_unit_drops(monkeypatch):
    """유닛 누락 경고는 unit을 실제로 떨어뜨리는 코드(empty_test_case_block/missing_unit_name)만
    센다. empty_expected(입력전용 시퀀스 등 무해)를 합산하면 오경보 — 전용 파서 컬럼탐지 수정으로
    empty_expected가 다수(994)가 됐을 때 '유닛 누락 가능' 허위 경보가 났던 회귀를 가드."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())

    # (a) empty_expected만 다수 → 유닛 누락 경고 없음(오경보 방지)
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{"unit_name": "s_foo", "test_cases": [{"base_tc_id": "TC1"}]}],
        "export_warnings": [{"code": "empty_expected", "message": "x"}] * 994,
    })
    warns_a: list = []
    _load_suts_fn_tcs("U:/s.xlsm", ["s_foo"], warn_sink=warns_a)
    assert not any("유닛 누락" in w for w in warns_a)

    # (b) 드롭 코드는 경고(개수는 무해 코드 제외 정확)
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{"unit_name": "s_foo", "test_cases": [{"base_tc_id": "TC1"}]}],
        "export_warnings": [
            {"code": "empty_test_case_block", "message": "x"},
            {"code": "missing_unit_name", "message": "y"},
            {"code": "empty_expected", "message": "z"},  # 무해분 — 개수에서 제외돼야
        ],
    })
    warns_b: list = []
    _load_suts_fn_tcs("U:/s.xlsm", ["s_foo"], warn_sink=warns_b)
    drop_warn = [w for w in warns_b if "유닛 누락" in w]
    assert len(drop_warn) == 1
    assert "2건" in drop_warn[0]  # empty_expected 제외, 드롭 2건만


def test_load_suts_fn_tcs_signature_unit_name_keyed_by_bare(monkeypatch):
    """HDPDM01: unit_name이 전체 C 시그니처면 bare 식별자로 키잉해야 프론트/스코프 필터(_direct_lc,
    bare 함수명)가 조인된다. 과거엔 시그니처 그대로 키잉→매칭 실패로 카드 '미파싱'·회귀 TC 0이 됐다.
    KJPDS02(bare)는 idempotent라 무변경."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: {
        "units": [{
            "unit_name": "void g_SysOs_WdiCtrl( void )",   # 시그니처(HDPDM01 실 doc 형태)
            "test_cases": [{"base_tc_id": "SwUTC_SwUFn_0001",
                            "inputs": {"a": 1}, "expected": {"r": 2}}],
        }],
        "export_warnings": [],
    })

    sink: dict = {}
    result = _load_suts_fn_tcs("U:/hdpdm01.xlsm", ["g_SysOs_WdiCtrl"], content_sink=sink)
    # bare 키로 저장 → _direct_lc(bare 함수명) 매칭 가능(과거 시그니처 키라 실패)
    assert result == {"g_SysOs_WdiCtrl": ["SwUTC_SwUFn_0001"]}
    assert "g_SysOs_WdiCtrl" in sink
    assert "void g_SysOs_WdiCtrl( void )" not in sink   # 시그니처 그대로 키잉하지 않음
    assert sink["g_SysOs_WdiCtrl"][0]["inputs"] == {"a": "1"}


def test_load_uds_fn_details_widened(monkeypatch):
    """사이드카 payload의 globals/calls/prototype 필드가 surface에 포함(widen)됨을 확인."""
    from workflow import impact_orchestrator as m
    from workflow.impact_orchestrator import _load_uds_fn_details

    monkeypatch.setattr(m, "_safe_exists", lambda p: True)
    monkeypatch.setattr(m, "_load_json", lambda p: {
        "function_details_by_name": {
            "s_foo": {
                "description": "Foo does X",
                "prototype": "void s_foo(uint8 x)",
                "globals_global": ["g_State"],
                "globals_static": ["s_Cnt"],
                "calls_list": ["s_helper"],
                "called": ["s_main"],
                "asil": "B",
            },
        },
    })
    out = _load_uds_fn_details("gen/uds.docx", ["s_foo"])
    assert out["s_foo"]["description"] == "Foo does X"
    assert out["s_foo"]["prototype"] == "void s_foo(uint8 x)"
    assert set(out["s_foo"]["globals"]) == {"g_State", "s_Cnt"}
    assert set(out["s_foo"]["calls"]) == {"s_helper", "s_main"}


def test_load_uds_fn_content_fallback_extracts_prototype(monkeypatch):
    """링크(cloudium) UDS fallback(parse_swuds_docx)이 표의 Prototype을 채운다.

    사이드카 없는 경로 — 예전엔 {description,heading}만 → prototype 조용히 빔. 이제 표에
    Prototype 행이 있으면 채워, 영향분석 UDS 카드가 실 선언 표시 + 원문→변경안 기준선 확보.
    """
    import io as _io

    import backend.services.file_resolver as fr
    from docx import Document  # type: ignore
    from workflow.impact_orchestrator import _load_uds_fn_content

    # 실제 docx: heading + 표(Name/Prototype/Description) — 사이드카 없음(fallback 경로)
    doc = Document()
    doc.add_paragraph("SwUFn_1150 — s_TunningParamRead_16bitData")
    tbl = doc.add_table(rows=3, cols=2)
    for r, (k, v) in enumerate([
        ("Name", "s_TunningParamRead_16bitData"),
        ("Prototype", "void s_TunningParamRead_16bitData(void)"),
        ("Description", "튜닝 파라미터를 16bit로 읽어 반환"),
    ]):
        tbl.cell(r, 0).text = k
        tbl.cell(r, 1).text = v
    buf = _io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver(docx_bytes))
    # 고유 경로(모듈 캐시 _UDS_CONTENT_CACHE 오염 회피) + 사이드카 부재로 fallback 강제
    out = _load_uds_fn_content("U:/link_proto_test.docx", ["s_tunningparamread_16bitdata"])
    key = "s_tunningparamread_16bitdata"
    assert key in out
    assert out[key]["prototype"] == "void s_TunningParamRead_16bitData(void)"
    assert "튜닝 파라미터를 16bit" in out[key]["description"]


def test_load_sds_fn_desc_extracts_description(monkeypatch):
    """SDS 파티션맵의 함수별 description을 영향 함수와 매칭해 반환."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "s_foo": {"description": "SDS: foo handles motor control", "kind": "function"},
        "swcom_10": {"description": "unrelated component", "kind": "component"},
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["s_foo"])
    assert out == {"s_foo": "SDS: foo handles motor control"}   # 영향 함수만, 무관 컴포넌트 제외


def test_load_sds_fn_desc_empty_on_no_path():
    from workflow.impact_orchestrator import _load_sds_fn_desc
    assert _load_sds_fn_desc("", ["s_foo"]) == {}


def test_load_sds_fn_desc_prefix_tolerant_match(monkeypatch):
    """flagged 's_tunningparamread' vs SDS 인터페이스명 'TunningParamRead'(접두어 없음) 매칭.

    exact miss여도 엔지니어링 접두어(s_/g_ 등) 정규화 대조로 SDS description을 찾는다.
    """
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "tunningparamread_16bitdata": {
            "description": "16비트 튜닝 파라미터를 읽는다", "kind": "function",
        },
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["s_TunningParamRead_16bitData"])
    # 키는 flagged fn(소문자) — 프론트 docContentFor 조회 키와 일치
    assert out == {"s_tunningparamread_16bitdata": "16비트 튜닝 파라미터를 읽는다"}


def test_load_sds_fn_desc_component_description_fallback(monkeypatch):
    """인터페이스 행 description이 비면 소속 컴포넌트 설명으로 폴백(출처 라벨링)."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "s_foo": {
            "description": "", "component_description": "모터 제어 컴포넌트", "kind": "function",
        },
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["s_foo"])
    assert out == {"s_foo": "(컴포넌트 설명) 모터 제어 컴포넌트"}


def test_load_sds_fn_desc_exact_match_wins_over_normalized(monkeypatch):
    """exact match가 접두어 정규화보다 우선(기존 동작 유지 — 정확 함수 설명 채택)."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "s_foo": {"description": "정확 매칭 설명", "kind": "function"},
        "foo": {"description": "정규화 매칭 설명", "kind": "function"},
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["s_foo"])
    assert out == {"s_foo": "정확 매칭 설명"}


def test_load_sds_fn_desc_normalized_collision_excluded(monkeypatch):
    """정규화 후 서로 다른 원본이 겹치면 모호 → 제외(잘못된 설명 붙이지 않음)."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    # 's_bar'와 'g_bar' 모두 정규화하면 'bar' → 충돌 → 접두어 매칭 포기
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "s_bar": {"description": "설명 A", "kind": "function"},
        "g_bar": {"description": "설명 B", "kind": "function"},
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["u8g_bar"])  # exact miss + 정규화 충돌
    assert out == {}


def test_load_sds_fn_desc_normalized_only_matches_functions(monkeypatch):
    """정규화 대조는 함수 엔트리만 — 컴포넌트 키로 오귀속하지 않는다."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "foo": {"description": "컴포넌트 설명", "kind": "component"},
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["s_foo"])  # 정규화하면 'foo'지만 컴포넌트라 제외
    assert out == {}


def test_load_sds_fn_desc_domain_prefix_not_stripped(monkeypatch):
    """도메인 접두어(spi_/adc_)는 벗기지 않는다 — [sgl]_ 조건 불충족(X7 오귀속 차단).

    열린 `^[a-z]{1,4}_`였다면 spi_read·adc_read 둘 다 'read'로 정규화돼 flagged spi_read가
    adc_read 설명을 조용히 집었다. 닫힌 헝가리안 패턴은 둘 다 원형 유지 → exact miss → 미매칭.
    """
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "adc_read": {"description": "ADC 읽기 설명", "kind": "function"},
    })
    # flagged 'spi_read'는 adc_read와 무관 — 닫힌 패턴은 'spi_'를 접두어로 안 봄 → 오귀속 없음
    out = _load_sds_fn_desc("U:/sds.docx", ["spi_read"])
    assert out == {}


def test_load_sds_fn_desc_normalized_3way_collision_excluded(monkeypatch):
    """3개+ 원본이 같은 정규화 키로 충돌해도 정확히 제외(set 기반, 2개째부터 영구 충돌)."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "s_bar": {"description": "설명 A", "kind": "function"},
        "g_bar": {"description": "설명 B", "kind": "function"},
        "l_bar": {"description": "설명 C", "kind": "function"},  # 3번째도 'bar'로 충돌
    })
    out = _load_sds_fn_desc("U:/sds.docx", ["u8g_bar"])  # exact miss + 3-way 정규화 충돌
    assert out == {}


def test_load_sds_fn_desc_hungarian_return_type_prefix(monkeypatch):
    """반환형+저장클래스 접두어(u8g_/s16g_)도 정규화 매칭된다(닫힌 집합 내)."""
    import backend.services.file_resolver as fr
    import report_gen.requirements as rq
    from workflow.impact_orchestrator import _load_sds_fn_desc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(rq, "_extract_sds_partition_map", lambda p: {
        "syseepromctrl_readbyte": {"description": "EEPROM 바이트 읽기", "kind": "function"},
    })
    # u8g_SysEepromCtrl_ReadByte → 'u8g_' 제거 → syseepromctrl_readbyte 매칭
    out = _load_sds_fn_desc("U:/sds.docx", ["u8g_SysEepromCtrl_ReadByte"])
    assert out == {"u8g_syseepromctrl_readbyte": "EEPROM 바이트 읽기"}


def test_load_sits_fn_chains_content_sink_keeps_tc_content(monkeypatch):
    """content_sink 제공 시 중간 JSON의 sub_cases(precondition/inputs/expected)를 TC-ID 키로 채우고,
    회귀 반환({entry_fn:[label]})은 flagged 함수만 유지(불변)."""
    import json
    import backend.services.file_resolver as fr
    from workflow.impact_orchestrator import _load_sits_fn_chains

    intermediate = {
        "integrations": [
            {
                "tc_id": "SwITC_SwUFn_0203",
                "entry_fn": "s_main",
                "call_chain": "s_main -> s_helper",
                "sub_cases": [
                    {"precondition": "system on", "inputs": {"a": 1, "empty": ""}, "expected": {"ret": 0}},
                ],
            },
            {   # entry_fn이 flagged 아님 — 회귀 result엔 없으나 내용(content_sink)엔 담겨야(프론트 TC-ID 조인).
                "tc_id": "SwITC_SwUFn_0999",
                "entry_fn": "s_other",
                "call_chain": "s_other -> x",
                "sub_cases": [],
            },
        ],
    }
    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver(json.dumps(intermediate).encode("utf-8")))

    sink: dict = {}
    result = _load_sits_fn_chains("U:/proj_SITS.xlsm", ["s_main"], content_sink=sink)

    # 회귀 반환: flagged(s_main)만
    assert result == {"s_main": ["SwITC_SwUFn_0203: s_main -> s_helper"]}
    # 내용: 전체 TC(entry_fn 무관), TC-ID 정규화(공백제거+대문자) 키
    assert "SWITC_SWUFN_0203" in sink
    assert "SWITC_SWUFN_0999" in sink
    c = sink["SWITC_SWUFN_0203"]
    assert c["call_chain"] == "s_main -> s_helper"
    assert c["sub_cases"][0]["precondition"] == "system on"
    assert c["sub_cases"][0]["inputs"] == {"a": "1"}    # empty 스킵, 문자열화
    assert c["sub_cases"][0]["expected"] == {"ret": "0"}


def test_load_sits_fn_chains_without_sink_unchanged(monkeypatch):
    """content_sink 미제공 시 회귀 반환만(순수 additive)."""
    import json
    import backend.services.file_resolver as fr
    from workflow.impact_orchestrator import _load_sits_fn_chains

    intermediate = {"integrations": [
        {"tc_id": "T1", "entry_fn": "s_a", "call_chain": "s_a", "sub_cases": []},
    ]}
    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver(json.dumps(intermediate).encode("utf-8")))
    assert _load_sits_fn_chains("U:/s.xlsm", ["s_a"]) == {"s_a": ["T1: s_a"]}


def test_load_testspec_by_tc_parses_and_normalizes(monkeypatch):
    """STS/SITS xlsm을 parse_swuts_xlsm로 파싱해 TC-ID 정규화 키 내용 맵을 반환."""
    import backend.services.file_resolver as fr
    import backend.services.swuts_excel_parser as sp
    from workflow.impact_orchestrator import _load_testspec_by_tc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver(b"xlsm-bytes"))

    class _E:
        description = "STS: verify boundary"
        precondition = "init done"
        test_method = "REQ"
        unit_name = "s_foo"

    class _Res:
        ok = True
        by_tc_id = {"SwTC_ 01": _E()}   # 내부 공백 포함 — 정규화로 제거돼야

    monkeypatch.setattr(sp, "parse_swuts_xlsm", lambda *a, **k: _Res())
    out = _load_testspec_by_tc("U:/sts.xlsm", doc_label="STS")
    assert "SWTC_01" in out
    assert out["SWTC_01"]["description"] == "STS: verify boundary"
    assert out["SWTC_01"]["test_method"] == "REQ"
    assert out["SWTC_01"]["unit_name"] == "s_foo"


def test_load_testspec_by_tc_honest_miss_on_parser_fail(monkeypatch):
    """파서 미매칭(ok=False)이면 빈 dict + 사유 warn(과대추정 금지, 정직 '미파싱')."""
    import backend.services.file_resolver as fr
    import backend.services.swuts_excel_parser as sp
    from workflow.impact_orchestrator import _load_testspec_by_tc

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver(b"xlsm-bytes"))

    class _Res:
        ok = False
        by_tc_id: dict = {}

    monkeypatch.setattr(sp, "parse_swuts_xlsm", lambda *a, **k: _Res())
    warns: list = []
    out = _load_testspec_by_tc("U:/x.xlsm", warn_sink=warns, doc_label="STS")
    assert out == {}
    assert any("미파싱" in w for w in warns)


def test_load_testspec_by_tc_empty_on_no_path():
    from workflow.impact_orchestrator import _load_testspec_by_tc
    assert _load_testspec_by_tc("", doc_label="SITS") == {}


# ── _build_doc_proposal: 문서 생성기(suts/sits/sts) 재사용으로 시험 초안 합성 ──
# 배선 로직(스코프/캡/graceful/예외격리/마커보존)을 생성기 스텁으로 격리 검증한다.
# (실 생성기 통합은 라이브 소스섹션 캐시로 별도 확인 — generate_sequences 실산출 검증)

def _proposal_sections(fd=None, gim=None):
    """generate_uds_source_sections 산출물 형태(function_details/globals_info_map) 최소 모사."""
    return {"function_details": fd or {}, "globals_info_map": gim or {}}


def _stub_generators(monkeypatch, *, suts_seq=None, sits_flows=None, sts_steps=None):
    """suts/sits/sts 생성기를 결정론 스텁으로 대체(배선만 격리)."""
    import generators.sits as gsits
    import generators.sts as gsts
    import generators.suts as gsuts
    monkeypatch.setattr(
        gsuts, "collect_unit_functions",
        lambda fdmap, gim=None: [{"name": info["name"]} for info in fdmap.values() if isinstance(info, dict)],
    )
    monkeypatch.setattr(
        gsuts, "generate_sequences",
        lambda unit, max_seq=6, type_cache=None: (suts_seq if suts_seq is not None
                                                  else [{"strategy": "BV_MIN", "inputs": {"x": 0}, "expected": {"ret": 0}, "description": "d"}]),
    )
    monkeypatch.setattr(
        gsits, "collect_integration_flows",
        lambda fdmap, max_flows=120: (sits_flows if sits_flows is not None
                                      else [{"entry_fn": "s_foo", "call_chain": "s_foo -> s_dep"}]),
    )
    monkeypatch.setattr(
        gsits, "generate_itc_list",
        lambda flows, **k: [{"entry_fn": f["entry_fn"], "call_chain": f["call_chain"],
                             "sub_cases": [{"inputs": {"x": 0}, "expected": {"ret": 0}, "precondition": "p"}]}
                            for f in flows],
    )
    monkeypatch.setattr(
        gsts, "_generate_steps_from_flow",
        lambda lf, info: (sts_steps if sts_steps is not None else [[{"action": "call s_foo", "expected": "ok"}]]),
    )


def test_build_doc_proposal_reuses_generators_scoped_to_changed(monkeypatch):
    """직접 변경 함수에 대해 suts/sits/sts 생성기를 재사용해 콜체인·경계값 입력·기대출력을 합성하고,
    changed_set 밖 함수(s_bar)는 초안에 포함하지 않는다(스코프)."""
    from workflow.impact_orchestrator import _build_doc_proposal

    fd = {
        "f1": {"name": "s_foo", "prototype": "U16 s_foo(U16 x)", "logic_flow": [], "calls_list": []},
        "f2": {"name": "s_bar", "prototype": "void s_bar(void)", "logic_flow": [], "calls_list": []},
    }
    # SUTS 스텁이 s_foo·s_bar 둘 다 반환해도 changed_set(s_foo) 밖은 필터돼야
    _stub_generators(monkeypatch, suts_seq=[{"strategy": "BV_MIN", "inputs": {"x": 0}, "expected": {"ret": 0}, "description": "d"}])
    out = _build_doc_proposal(_proposal_sections(fd), {"s_foo"})

    assert list(out["suts"].keys()) == ["s_foo"]                      # s_bar 제외(스코프)
    # ⚠ 값은 **문자열**이다 — doc_proposal도 doc_content와 같이 `_cap_kv`를 태워 표기를
    #   통일했다(예전엔 생성기만 native int라 화면에서 "0 → 0x0" 같은 허위 변경이 그려졌다).
    assert out["suts"]["s_foo"][0]["inputs"] == {"x": "0"}
    assert out["suts"]["s_foo"][0]["expected"] == {"ret": "0"}         # 기대출력 합성(문자열 통일)
    assert out["sits"]["s_foo"]["call_chain"] == "s_foo -> s_dep"      # 실 통합 콜체인
    assert out["sits"]["s_foo"]["sub_cases"][0]["expected"] == {"ret": "0"}
    assert out["sts"]["s_foo"][0][0]["action"] == "call s_foo"          # 시험 절차 스텝


def test_build_doc_proposal_empty_changed_set_returns_empty(monkeypatch):
    """빈 changed_set이면 생성기를 호출하지 않고 모든 노드가 빈 채로 반환."""
    from workflow.impact_orchestrator import _build_doc_proposal, _empty_doc_proposal
    called = {"n": 0}
    import generators.suts as gsuts
    monkeypatch.setattr(gsuts, "collect_unit_functions", lambda *a, **k: called.__setitem__("n", called["n"] + 1) or [])
    out = _build_doc_proposal(_proposal_sections({"f1": {"name": "s_foo"}}), set())
    assert out == _empty_doc_proposal()
    assert called["n"] == 0


def test_build_doc_proposal_no_function_details_warns():
    """소스도 문서도 없으면 **모든 노드가 빈 채로** 반환 + 사유 warn(silent-0 금지).

    ⚠ 노드가 늘어도(uds/sds/var_types/suts_meta) 이 경로는 전부 비어 있어야 한다 —
    `_empty_doc_proposal()`과 정확히 같은지로 검사해 새 노드가 몰래 채워지는 것을 막는다."""
    from workflow.impact_orchestrator import _build_doc_proposal, _empty_doc_proposal
    warns: list = []
    out = _build_doc_proposal({}, {"s_foo"}, warn_sink=warns)
    assert out == _empty_doc_proposal()
    assert all(not v for k, v in out.items() if k != "source")
    assert out["source"] == ""      # 근거 없음 — generator도 document도 아니다
    assert any("function_details 없음" in w for w in warns)


def test_build_doc_proposal_caps_function_count(monkeypatch):
    """fn_cap을 초과하는 변경 함수는 초안 대상에서 제외(페이로드 억제)."""
    from workflow.impact_orchestrator import _build_doc_proposal
    fd = {f"f{i}": {"name": f"s_fn{i}", "logic_flow": [], "calls_list": []} for i in range(5)}
    _stub_generators(monkeypatch, sits_flows=[])
    changed = {f"s_fn{i}" for i in range(5)}
    out = _build_doc_proposal(_proposal_sections(fd), changed, fn_cap=2)
    # SUTS/STS는 targets(cap=2)만 — 정확히 2개
    assert len(out["suts"]) == 2
    assert len(out["sts"]) == 2


def test_build_doc_proposal_generator_exception_isolated(monkeypatch):
    """한 생성기(SUTS) 예외가 다른 생성기(SITS/STS)를 막지 않고, 사유가 warn으로 표면화된다."""
    from workflow.impact_orchestrator import _build_doc_proposal
    import generators.suts as gsuts
    fd = {"f1": {"name": "s_foo", "logic_flow": [], "calls_list": []}}
    _stub_generators(monkeypatch)

    def _boom(*a, **k):
        raise RuntimeError("suts boom")
    monkeypatch.setattr(gsuts, "collect_unit_functions", _boom)

    warns: list = []
    out = _build_doc_proposal(_proposal_sections(fd), {"s_foo"}, warn_sink=warns)
    assert out["suts"] == {}                                   # SUTS는 예외로 비었으나
    assert out["sits"]["s_foo"]["call_chain"] == "s_foo -> s_dep"   # SITS는 정상
    assert out["sts"]["s_foo"]                                  # STS도 정상
    assert any("SUTS 시퀀스 합성 예외" in w for w in warns)


def test_build_doc_proposal_preserves_verification_marker(monkeypatch):
    """생성기가 판정 불가값을 '[검증 필요]'로 표기하면 초안에 그대로 보존한다(정직성)."""
    from workflow.impact_orchestrator import _build_doc_proposal
    fd = {"f1": {"name": "s_foo", "logic_flow": [], "calls_list": []}}
    _stub_generators(monkeypatch, sits_flows=[], sts_steps=[],
                     suts_seq=[{"strategy": "BV_MAX", "inputs": {"x": 65535},
                                "expected": {"ret": "[검증 필요] saturated"}, "description": "d"}])
    out = _build_doc_proposal(_proposal_sections(fd), {"s_foo"})
    assert out["suts"]["s_foo"][0]["expected"] == {"ret": "[검증 필요] saturated"}


def test_build_doc_proposal_suts_uses_gim_types_not_global():
    """SUTS 경계값이 gim 타입(로컬 주입)으로 산출되고 프로세스 전역 타입캐시를 오염시키지 않는다.

    과거 버그: _build_doc_proposal이 set_globals_type_cache를 호출하지 않아 빈/stale 전역
    타입캐시→uint8 폴백으로 U32 변수의 경계 max를 255로 잘못 산출했다(실측 12개 변수 상이).
    이제 gim→{var:type} 로컬맵을 generate_sequences(type_cache=)로 명시 주입한다:
      (a) 실 문서생성(set_globals_type_cache→전역)과 동일한 타입으로 경계 산출,
      (b) 프로세스 전역은 일절 변이 안 함 → 동시 문서생성과 write-race/오염 차단.
    실 생성기(스텁 아님)로 두 불변식을 함께 잠근다.
    """
    import generators.suts as gsuts
    from workflow.impact_orchestrator import _build_doc_proposal

    # xfer_len 은 gim 에서 U32. 함수 input 으로 노출('[IN] TYPE name' = _parse_signature_params 형식).
    fd = {"f1": {"name": "s_xfer", "prototype": "void s_xfer(U32 xfer_len)",
                 "inputs": ["[IN] U32 xfer_len"], "outputs": [],
                 "logic_flow": [], "calls_list": []}}
    gim = {"xfer_len": {"type": "U32"}}

    # ⚠ 전역 타입캐시를 일부러 틀린 타입(uint8)으로 오염 — 주입이 전역을 '읽지' 않는지 검증
    gsuts._globals_type_cache.clear()
    gsuts._globals_type_cache["xfer_len"] = "uint8_t"
    try:
        out = _build_doc_proposal(_proposal_sections(fd, gim), {"s_xfer"})

        # (a) SUTS 시퀀스의 xfer_len 최대 입력이 U32 max(4294967295) — uint8(255) 폴백 아님
        seqs = out["suts"]["s_xfer"]
        xfer_vals = [s["inputs"]["xfer_len"] for s in seqs if "xfer_len" in (s.get("inputs") or {})]
        assert "4294967295" in xfer_vals, f"U32 경계가 반영돼야(전역 오염 무시): {xfer_vals}"
        assert 255 not in xfer_vals                          # uint8 폴백 아님

        # (b) 전역 타입캐시는 호출 후에도 오염값 그대로 — _build_doc_proposal 이 전역을 변이하지 않음
        assert gsuts._globals_type_cache.get("xfer_len") == "uint8_t"
    finally:
        gsuts._globals_type_cache.clear()   # 다른 테스트 격리(전역 싱글톤 복원)


def test_generate_sequences_type_cache_overrides_global():
    """generate_sequences(type_cache=)가 주어지면 전역 _globals_type_cache 대신 그것을 읽고,
    None(기본)이면 기존대로 전역을 읽는다(실 문서생성 경로 무변경 보증)."""
    import generators.suts as gsuts

    unit = {"name": "u", "input_vars": ["v"], "output_vars": [], "logic_flow": []}
    gsuts._globals_type_cache.clear()
    gsuts._globals_type_cache["v"] = "uint8_t"     # 전역 = uint8
    try:
        # type_cache 명시 주입(U32) → 전역(uint8) 무시하고 U32 경계
        inj = [s for s in gsuts.generate_sequences(unit, type_cache={"v": "U32"})
               if s["strategy"] == "BV_MAX"]
        assert inj and inj[0]["inputs"]["v"] == 4294967295

        # type_cache=None(기본) → 전역(uint8) 사용 = 실 문서생성 경로 무변경
        glob = [s for s in gsuts.generate_sequences(unit)
                if s["strategy"] == "BV_MAX"]
        assert glob and glob[0]["inputs"]["v"] == 255
    finally:
        gsuts._globals_type_cache.clear()


def test_extract_mcdc_conditions_threads_type_cache_into_recursion():
    """_extract_mcdc_conditions의 재귀 자기호출(중첩 if/children)에도 type_cache가 스레드된다.

    중첩 조건은 재귀 경로로만 도달하므로, 재귀에서 type_cache를 흘리면(과거 스냅샷 결함) 안쪽
    var-vs-var 조건의 MC/DC 경계값이 전역(폴백 uint8)을 읽는다. 바깥 if의 children에 안쪽 if를
    두고, 주입(U32) vs 미주입(전역 비움→uint8) 대조로 재귀 스레딩을 보증한다."""
    import generators.suts as gsuts

    gsuts._globals_type_cache.clear()   # 전역 비움 → 미주입 시 uint8 폴백(대조군)
    try:
        # 바깥 if(children) → 안쪽 if 'big > small'(var-vs-var, 재귀로만 도달)
        flow = [{"type": "if", "condition": "outer > 0", "children": [
            {"type": "if", "condition": "big > small", "children": []},
        ]}]
        ivars = ["big", "small", "outer"]

        inj = gsuts._extract_mcdc_conditions(flow, ivars, {"big": "U32"})
        big_inj = [c for c in inj if c[0] == "big"]
        assert big_inj, f"중첩 var-vs-var 조건 미추출: {inj}"
        # tuple=(var, op, rhs, true_val, false_val); '>'→true_val=bmax. U32 max면 재귀 스레드 성공.
        assert big_inj[0][3] == 4294967295, f"재귀에 type_cache 미스레드(uint8 폴백): {big_inj}"

        # 대조군: 미주입(전역 비움) → uint8 폴백(255). 재귀가 None을 흘려도 이 값이라 회귀 검출 가능.
        glob = gsuts._extract_mcdc_conditions(flow, ivars)
        big_glob = [c for c in glob if c[0] == "big"]
        assert big_glob and big_glob[0][3] == 255
    finally:
        gsuts._globals_type_cache.clear()


# ─────────────────────────────────────────────────────────────────────────────
# 문서 작성급 초안 — 원문 grounding 원재료(컬럼/메타/캡 표면화)
#
# 사용자가 본 "TC 3건 × 변수 5개"는 문서의 실제 크기가 아니라 `_seqs[:3]`·`_cap_kv(n=5)`의
# **절단값**이었고 아무 표기가 없었다. 캡을 올리고, 절단이 남았으면 seq_total로 드러낸다.
# 컬럼명은 시트 헤더 원문이라 "재계산 대상 변수집합"의 권위 소스다(시그니처 유추 금지).
# ─────────────────────────────────────────────────────────────────────────────


def _suts_model_with_columns(n_seq: int = 10, n_kv: int = 14) -> dict:
    """입력 컬럼 n_kv개 × 시퀀스 n_seq개를 가진 모델(캡 검증용)."""
    cols = [f"g_arr[{i}]" for i in range(n_kv)]
    return {
        "units": [{
            "unit_name": "void s_foo( void )",     # 시그니처 템플릿 — bare 정규화 확인 겸용
            "component": "SwCom_07\n(diag)",
            "metadata": {"gen_method": "AEC, ABV", "test_method": "FNCT",
                         "related_ids": ["SwRS_0101", "SwDS_02"]},
            "columns": {"inputs": cols, "expected": ["ret"], "sheet": "2.SW Unit Test Spec"},
            "test_cases": [
                {
                    "base_tc_id": "SwUTC_SwUFn_1219",
                    "description": f"seq {i}",
                    "inputs": {c: f"0x{i}" for c in cols},
                    "expected": {"ret": "0x0"},
                    "source": {"sheet": "2.SW Unit Test Spec", "tc_row": 2103, "sequence_row": 2104 + i},
                }
                for i in range(n_seq)
            ],
        }],
        "export_warnings": [],
    }


def test_load_suts_fn_tcs_meta_sink_carries_columns_and_document_meta(monkeypatch):
    """meta_sink에 시트 헤더 원문 컬럼 + component/test_method/gen_method/related_ids가 실린다.

    컬럼명이 곧 '재계산 대상 변수집합'이자 Excel 붙여넣기 열 순서다 — 파서가 이미 뽑아둔
    값을 옮길 뿐(재파싱 없음). 열 **순서 보존**이 계약이라 set 비교가 아니라 list 비교."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: _suts_model_with_columns(2, 3))

    meta: dict = {}
    result = _load_suts_fn_tcs("U:/suts.xlsm", ["s_foo"], meta_sink=meta)

    assert result == {"s_foo": ["SwUTC_SwUFn_1219"]}       # 회귀 반환 불변
    m = meta["s_foo"]                                       # 키는 bare 정규화된 함수명
    assert m["columns"]["inputs"] == ["g_arr[0]", "g_arr[1]", "g_arr[2]"]  # 열 순서 보존
    assert m["columns"]["expected"] == ["ret"]
    assert m["columns"]["sheet"] == "2.SW Unit Test Spec"
    assert m["component"].startswith("SwCom_07")
    assert m["test_method"] == "FNCT"
    assert m["gen_method"] == "AEC, ABV"
    assert m["related_ids"] == ["SwRS_0101", "SwDS_02"]


def test_load_suts_fn_tcs_columns_preserve_array_subscripts(monkeypatch):
    """배열 첨자(`g_sys_error_his[0]`)를 벗기지 않는다.

    생성기(`generators/suts.py:_extract_var_names`)는 첨자를 strip하지만 그건 실 문서생성
    경로다. 초안 grounding은 **문서 컬럼 그대로**여야 원문 셀과 1:1 대응한다."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    cols = [f"g_sys_error_his[{i}]" for i in range(5)]
    model = {
        "units": [{
            "unit_name": "s_updateerrorcode",
            "columns": {"inputs": cols, "expected": cols, "sheet": "2.SW Unit Test Spec"},
            "test_cases": [{"base_tc_id": "T1", "inputs": {c: "0x0" for c in cols},
                            "expected": {c: "0x0" for c in cols}}],
        }],
        "export_warnings": [],
    }
    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: model)

    meta: dict = {}
    content: dict = {}
    _load_suts_fn_tcs("U:/s.xlsm", ["s_updateerrorcode"], content_sink=content, meta_sink=meta)

    assert meta["s_updateerrorcode"]["columns"]["inputs"] == cols
    assert list(content["s_updateerrorcode"][0]["inputs"]) == cols


def test_load_suts_fn_tcs_caps_are_raised_and_truncation_is_surfaced(monkeypatch):
    """기본 캡이 8시퀀스·12변수로 올라가고, 잘린 만큼은 seq_total로 드러난다(silent 절단 금지)."""
    import backend.services.file_resolver as fr
    import tools.export_suts_vectorcast as ev
    from workflow.impact_orchestrator import _load_suts_fn_tcs

    monkeypatch.setattr(fr, "get_resolver", lambda: _FakeResolver())
    monkeypatch.setattr(ev, "build_vectorcast_model", lambda *a, **k: _suts_model_with_columns(10, 14))

    content: dict = {}
    meta: dict = {}
    _load_suts_fn_tcs("U:/s.xlsm", ["s_foo"], content_sink=content, meta_sink=meta)

    rows = content["s_foo"]
    assert len(rows) == 8                       # 과거 3 → 8
    assert len(rows[0]["inputs"]) == 12         # 과거 5 → 12
    assert meta["s_foo"]["seq_total"] == 10     # 문서의 실제 개수
    assert meta["s_foo"]["seq_shown"] == 8      # 표시된 개수 → 프론트가 '10건 중 8건'


def test_cap_kv_default_stays_five_for_shared_callers():
    """`_cap_kv` 기본값 5는 STS/SITS 폴백 캡처와 공용이라 불변 — 상향은 호출부 `n=`로만."""
    from workflow.impact_orchestrator import _cap_kv

    src = {f"k{i}": i for i in range(20)}
    assert len(_cap_kv(src)) == 5
    assert len(_cap_kv(src, n=12)) == 12


def test_load_sits_fn_chains_keeps_case_label_and_sub_total(monkeypatch):
    """SITS 서브케이스의 case_label(AEC 등가분할 라벨)과 절단 전 총 개수를 보존한다."""
    import json as _json

    import backend.services.file_resolver as fr
    from workflow.impact_orchestrator import _load_sits_fn_chains

    payload = {"integrations": [{
        "entry_fn": "s_entry", "call_chain": "s_entry -> Hal_Read", "tc_id": "SwITC_SwUFn_0101",
        "gen_method": "AEC", "related_ids": ["SwDS_11"],
        "sub_cases": [
            {"case_label": f"{i} [EC{i}:무효-하한]", "precondition": f"env{i}",
             "inputs": {"rpm": i}, "expected": {"state": i}}
            for i in range(9)
        ],
    }]}
    monkeypatch.setattr(fr, "get_resolver",
                        lambda: _FakeResolver(_json.dumps(payload).encode("utf-8")))

    sink: dict = {}
    _load_sits_fn_chains("U:/sits.xlsm", ["s_entry"], content_sink=sink)

    tc = sink["SWITC_SWUFN_0101"]               # _normTc(공백제거+대문자)
    assert len(tc["sub_cases"]) == 6            # 과거 3 → 6
    assert tc["sub_total"] == 9                 # 절단 전 실제 개수(표면화)
    assert tc["sub_cases"][0]["case_label"] == "0 [EC0:무효-하한]"
    assert tc["entry_fn"] == "s_entry"
    assert tc["gen_method"] == "AEC"
    assert tc["related_ids"] == ["SwDS_11"]


def test_shrink_doc_content_degrades_and_warns():
    """예산 초과 시 단계적으로 줄이되 **사유를 반드시 남긴다**(절단 침묵 금지)."""
    from workflow.impact_orchestrator import _shrink_doc_content

    big = {
        "suts": {f"fn{i}": [
            {"tc_id": f"T{j}", "inputs": {f"v{k}": "0x1234" * 4 for k in range(12)},
             "expected": {f"e{k}": "0x1234" * 4 for k in range(12)}}
            for j in range(8)
        ] for i in range(40)},
        "sits_by_tc": {f"TC{i}": {"sub_cases": [
            {"inputs": {f"v{k}": "0x1234" * 4 for k in range(12)}, "expected": {}}
            for _ in range(6)
        ]} for i in range(40)},
    }
    warns: list = []
    out = _shrink_doc_content(big, warn_sink=warns, budget=50_000)

    # 경고는 "상한으로 축소"가 아니라 **실제 도달 크기**를 밝혀야 한다 — 지배항(TC 개수)은
    # 축소 대상이 아니라 상한을 못 맞출 수 있고, 그때 "상한으로 축소"는 지킨 것처럼 읽힌다.
    assert any("페이로드 축소" in w and "상한" in w for w in warns), "절단을 침묵시키면 안 된다"
    assert len(out["suts"]["fn0"]) <= 5
    assert len(out["sits_by_tc"]["TC0"]["sub_cases"]) <= 4


def test_shrink_doc_content_noop_under_budget():
    """예산 이내면 손대지 않고 경고도 없다(불필요한 노이즈 금지)."""
    from workflow.impact_orchestrator import _shrink_doc_content

    small = {"suts": {"fn": [{"tc_id": "T1", "inputs": {"x": "1"}, "expected": {"r": "2"}}] * 8},
             "sits_by_tc": {}}
    warns: list = []
    out = _shrink_doc_content(small, warn_sink=warns)

    assert warns == []
    assert len(out["suts"]["fn"]) == 8


# ─────────────────────────────────────────────────────────────────────────────
# doc_proposal 확장 — 버려지던 문서 컬럼 복원 / UDS·SDS 노드 / cloudium 문서 폴백
# ─────────────────────────────────────────────────────────────────────────────


def test_build_doc_proposal_restores_suts_document_columns(monkeypatch):
    """Component/Test Method/Gen.Method/ASIL/Related ID/Precondition을 **함수당 1회** 싣는다.

    생성기가 이미 산출하는데 예전 슬림화가 통째로 버려서 카드가 시험 값만 보여줬다.
    행마다 반복하면 시퀀스 수만큼 페이로드가 부풀므로 suts_meta로 분리한다."""
    import generators.suts as gsuts
    from workflow.impact_orchestrator import _build_doc_proposal

    _stub_generators(monkeypatch, suts_seq=[
        {"strategy": "BV_MIN", "inputs": {"x": 0}, "expected": {"ret": 0}, "description": "d", "seq_num": 1},
        {"strategy": "BV_MAX", "inputs": {"x": 255}, "expected": {"ret": 1}, "description": "d2", "seq_num": 2},
    ])
    monkeypatch.setattr(gsuts, "collect_unit_functions", lambda fdmap, gim=None: [{
        "name": "s_foo", "component": "SwCom_07\n(diag)", "asil": "C",
        "srs_req_ids": ["SwRS_0101"], "precondition": "system initialized",
        "prototype": "void s_foo(U16 x)", "input_vars": ["x"], "output_vars": ["ret"],
    }])
    monkeypatch.setattr(gsuts, "determine_test_method", lambda u: "FNCT")
    monkeypatch.setattr(gsuts, "determine_gen_method", lambda u: "AEC, ABV")

    out = _build_doc_proposal(_proposal_sections({"f1": {"name": "s_foo"}}), {"s_foo"})

    assert out["source"] == "generator"
    m = out["suts_meta"]["s_foo"]
    assert m["component"].startswith("SwCom_07")
    assert (m["test_method"], m["gen_method"]) == ("FNCT", "AEC, ABV")
    assert m["asil"] == "C"
    assert m["srs_req_ids"] == ["SwRS_0101"]
    assert m["precondition"] == "system initialized"
    # 생성기 축은 gen_* — 문서 TC 수(doc_total)와 절대 같은 슬롯에 두지 않는다(C1)
    assert m["gen_total"] == 2 and m["gen_truncated"] is False
    assert "total" not in m, "문서 축과 혼동되는 이름 금지"
    assert out["suts"]["s_foo"][0]["seq_num"] == 1        # 시퀀스 번호도 복원
    # 메타는 행에 반복되지 않는다(페이로드 계약)
    assert "component" not in out["suts"]["s_foo"][0]


def test_build_doc_proposal_restores_sits_case_labels_and_meta(monkeypatch):
    """SITS의 case_label(AEC 등가분할)·tc_id·gen_method·related_ids를 보존."""
    import generators.sits as gsits
    from workflow.impact_orchestrator import _build_doc_proposal

    _stub_generators(monkeypatch)
    monkeypatch.setattr(gsits, "generate_itc_list", lambda flows, **k: [{
        "entry_fn": "s_foo", "call_chain": "s_foo -> s_dep", "tc_id": "SwITC_SwUFn_0101",
        "gen_method": "AEC", "related_ids": ["SwDS_11"], "asil": "B", "module_name": "diag",
        "sub_cases": [{"case_num": i, "case_label": "%d [EC%d:무효-하한]" % (i, i),
                       "inputs": {"x": i}, "expected": {"ret": i}, "precondition": "p"}
                      for i in range(9)],
    }])

    out = _build_doc_proposal(_proposal_sections({"f1": {"name": "s_foo"}}), {"s_foo"})

    s = out["sits"]["s_foo"]
    assert s["tc_id"] == "SwITC_SwUFn_0101"
    assert (s["gen_method"], s["asil"], s["module_name"]) == ("AEC", "B", "diag")
    assert s["related_ids"] == ["SwDS_11"]
    assert len(s["sub_cases"]) == 6          # sub_cap 기본 3 → 6
    assert s["total"] == 9 and s["truncated"] is True   # 절단 표면화
    assert s["sub_cases"][0]["case_label"] == "0 [EC0:무효-하한]"


def test_build_doc_proposal_var_types_omit_unknown(monkeypatch):
    """var_types는 타입이 해상된 변수만 — 미상은 키 부재(uint8_t 기본값 환각 차단)."""
    import generators.suts as gsuts
    from workflow.impact_orchestrator import _build_doc_proposal

    _stub_generators(monkeypatch, suts_seq=[
        {"strategy": "BV_MIN", "inputs": {"g_sys_error_his[0]": 0, "u16t_Data": 0},
         "expected": {"g_sys_error_his[0]": 0}, "description": "d"},
    ])
    monkeypatch.setattr(gsuts, "collect_unit_functions",
                        lambda fdmap, gim=None: [{"name": "s_foo", "input_vars": [], "output_vars": []}])

    sections = _proposal_sections({"f1": {"name": "s_foo"}})
    sections["globals_info_map"] = {"g_sys_error_his": {"type": "U16"}}
    out = _build_doc_proposal(sections, {"s_foo"})

    vt = out["var_types"]["s_foo"]
    assert vt["g_sys_error_his"]["type"] == "uint16_t"   # 첨자 접힘 + 실측 타입
    assert vt["u16t_Data"]["type"] == "uint16_t"         # 이름 규칙
    assert all(v["type"] != "uint8_t" for v in vt.values()), "미상→uint8_t 기본값이 새면 안 된다"


def test_build_doc_proposal_uds_sds_nodes_do_not_invent_prose(monkeypatch):
    """UDS/SDS 노드는 구조 항목만 채우고 산문은 지어내지 않는다(*_source: ai_required)."""
    from workflow.impact_orchestrator import _build_doc_proposal

    _stub_generators(monkeypatch, sits_flows=[])
    fd = {"f1": {"name": "s_foo", "prototype": "void s_foo(U16 x)",
                 "inputs": ["[IN] U16 x"], "outputs": ["[OUT] U8 ret"],
                 "globals_global": ["g_a"], "globals_static": ["s_b"],
                 "calls_list": ["s_dep"], "logic_flow": ["if (x > 0)"],
                 "precondition": "init", "asil": "C", "module_name": "diag"}}
    out = _build_doc_proposal(_proposal_sections(fd), {"s_foo"},
                              change_details={"s_foo": {"after": "void s_foo(U32 x)"}})

    u = out["uds"]["s_foo"]
    assert u["prototype"] == "void s_foo(U16 x)"
    assert u["prototype_after"] == "void s_foo(U32 x)"
    assert u["globals"] == ["g_a", "s_b"]
    assert u["calls"] == ["s_dep"] and u["logic_flow"] == ["if (x > 0)"]
    assert u["description_source"] == "ai_required"      # 산문 창작 금지
    assert "description" not in u

    s = out["sds"]["s_foo"]
    assert s["interface_before"] == "void s_foo(U16 x)"
    assert s["interface_after"] == "void s_foo(U32 x)"
    assert s["behavior_source"] == "ai_required"
    assert s["related_ids"] == [], "신규 요구 ID를 창작하지 않는다"


def test_build_doc_proposal_document_fallback_when_source_unresolved():
    """소스 미해결(cloudium)이어도 문서 원문이 있으면 초안 원재료를 만든다.

    실사용 화면이 골격만 남던 근본 원인 — sections={}면 통째로 빈 dict를 돌려줬다.
    시퀀스(`suts`)는 만들지 않는다: 생성기 없이 시험 값을 지어내면 환각이다."""
    from workflow.impact_orchestrator import _build_doc_proposal

    cols = ["g_sys_error_his[%d]" % i for i in range(5)]
    doc_content = {
        "suts": {"s_updateerrorcode": [
            {"tc_id": "SwUTC_SwUFn_1219", "inputs": {c: "0x0" for c in cols},
             "expected": {c: "0x0" for c in cols},
             "loc": {"sheet": "2.SW Unit Test Spec", "tc_row": 2103}},
        ]},
        "suts_meta": {"s_updateerrorcode": {
            "columns": {"inputs": cols, "expected": cols, "sheet": "2.SW Unit Test Spec"},
            "component": "SwCom_07", "test_method": "FNCT", "gen_method": "AEC, ABV",
            "related_ids": ["SwRS_0101"], "seq_total": 12, "seq_shown": 8,
        }},
        # UDS payload의 어노테이션(`[IN] U16 …`) — 소스가 없는 cloudium에서 유일한 타입 근거
        "uds": {"s_updateerrorcode": {"prototype": "void s_updateerrorcode(U16 u16t_Data)",
                                      "globals": ["[IN] U16 g_sys_error_his"], "calls": []}},
    }
    warns: list = []
    out = _build_doc_proposal(
        {}, {"s_updateerrorcode"}, warn_sink=warns,
        doc_content=doc_content,
        change_details={"s_updateerrorcode": {"after": "void s_updateerrorcode(U8 d)"}})

    assert out["source"] == "document"
    assert any("문서 원문 기준" in w for w in warns), "폴백 사용 사실을 표면화해야 한다"
    assert out["suts"] == {}, "생성기 없이 시퀀스를 지어내면 안 된다"
    # 원문 컬럼(첨자 포함) ↔ UDS 어노테이션 조인 — 배열 첨자는 base로 접힌다
    _vt = out["var_types"]["s_updateerrorcode"]["g_sys_error_his"]
    assert _vt == {"type": "uint16_t", "source": "doc_annotation"}
    m = out["suts_meta"]["s_updateerrorcode"]
    assert (m["component"], m["test_method"], m["gen_method"]) == ("SwCom_07", "FNCT", "AEC, ABV")
    assert m["doc_total"] == 12 and m["doc_shown"] == 8
    assert "total" not in m and "gen_total" not in m
    assert out["uds"]["s_updateerrorcode"]["prototype_after"] == "void s_updateerrorcode(U8 d)"
    assert out["uds"]["s_updateerrorcode"]["logic_flow"] == [], "소스 없이 의사코드를 만들지 않는다"


def test_build_doc_proposal_document_fallback_unknown_type_yields_no_var_type():
    """문서 폴백에서도 미상 타입은 var_types에 넣지 않는다(숫자 제안 억제)."""
    from workflow.impact_orchestrator import _build_doc_proposal

    doc_content = {
        "suts": {"s_foo": [{"tc_id": "T1", "inputs": {"SomeEnum_Mode": "IDLE"}, "expected": {}}]},
        "suts_meta": {"s_foo": {"columns": {"inputs": ["SomeEnum_Mode"], "expected": []}}},
    }
    out = _build_doc_proposal({}, {"s_foo"}, doc_content=doc_content)
    assert "s_foo" not in out["var_types"]


def test_shrink_drops_functions_only_as_last_resort_and_records_them():
    """행·변수를 다 줄여도 초과면 **함수 축**을 자르되, 빠진 함수를 기록한다.

    ⚠ 기록이 없으면 프론트가 '문서에 TC 없음(미파싱)'과 구분하지 못해, 실제로는 문서에 있는
    TC를 "없다"고 표시한다. 무조건 자르지 않는 이유는 기존에 보이던 원문이 사라지는 회귀이기 때문."""
    from workflow.impact_orchestrator import _shrink_doc_content

    big = {
        "suts": {f"fn{i:04d}": [
            {"tc_id": f"T{j}", "inputs": {f"v{k}": "0x1234" * 6 for k in range(12)},
             "expected": {f"e{k}": "0x1234" * 6 for k in range(12)}}
            for j in range(8)
        ] for i in range(400)},
        "suts_meta": {f"fn{i:04d}": {"columns": {"inputs": [f"v{k}" for k in range(12)], "expected": []},
                                     "seq_total": 8, "seq_shown": 8} for i in range(400)},
        "sits_by_tc": {},
    }
    warns: list = []
    out = _shrink_doc_content(big, warn_sink=warns, budget=200_000)

    assert len(out["suts"]) < 400, "지배항인 함수 축이 마지막에 줄어야 한다"
    om = out.get("suts_omitted")
    assert om and om["reason"] == "payload_budget"
    assert om["count"] == 400 - len(out["suts"])
    assert om["functions"], "어떤 함수가 빠졌는지 알 수 있어야 '미파싱'과 구분된다"
    # 남은 함수의 meta도 함께 정리돼야 한다(고아 meta 금지)
    assert set(out["suts_meta"]) == set(out["suts"])
    assert any("함수" in w and "생략" in w for w in warns)


def test_shrink_does_not_claim_reduction_when_nothing_shrank():
    """이미 최소 표본이라 0바이트도 안 줄었는데 "축소"라고 보고하면 안 된다(실측 흔한 경로)."""
    from workflow.impact_orchestrator import _shrink_doc_content

    # 시퀀스 1건·변수 1개 — 어떤 강등 단계로도 더 줄일 게 없다. 크기는 함수 개수에서 온다.
    tight = {
        "suts": {f"fn{i:04d}": [{"tc_id": "T", "inputs": {"v": "0x1"}, "expected": {}}] for i in range(4000)},
        "suts_meta": {}, "sits_by_tc": {},
    }
    warns: list = []
    _shrink_doc_content(tight, warn_sink=warns, budget=100)

    joined = " ".join(warns)
    assert "축소 여지 없음" in joined or "생략" in joined
    assert "시퀀스 3건·변수 5개) 0KB" not in joined


def test_shrink_noop_under_budget_records_nothing():
    from workflow.impact_orchestrator import _shrink_doc_content

    small = {"suts": {"fn": [{"tc_id": "T1", "inputs": {"x": "1"}, "expected": {}}]},
             "suts_meta": {}, "sits_by_tc": {}}
    warns: list = []
    out = _shrink_doc_content(small, warn_sink=warns)
    assert warns == []
    assert "suts_omitted" not in out
