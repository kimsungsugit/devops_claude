"""템플릿이 **이 프로젝트의** 문서인지 표면화한다 — §6 후보 9.

## 계획서 전제가 두 번 뒤집혔다

계획서 표(§6 9행)는 *"소스 900함수 중 271개만 HDPDM01 템플릿에 존재(629건=69.9% 부재).
의도된 부분집합인지 오배치인지는 프로젝트 설정 판단이 필요해 `ok` verdict 는 뒤집지 않고
수치만 노출해 둔 상태"* 라고 적었다.

**① "무손실이라 `ok` 가 정답" 이라는 검증이 무효였다.** 그 검증에 쓴 산출물은 템플릿
모드가 아니라 **무템플릿 폴백**으로 만들어진 것이다(`docx_builder.py:4140` 주석: "반영
누락이 원리적으로 없다", `unmatched_payload_count: 0` 하드코딩). 손실이 구조적으로
불가능한 조건에서 나온 clean 이라 근거가 되지 않는다. 같은 payload 를 템플릿 모드로
돌린 실측은 저장소가 이미 기록해 뒀다 — **payload 432개 중 95개(22.0%) 미반영 + 빈
heading 74개인데 `success`**(`backend/helpers/uds.py:1308-1312`, 커밋 f063d93).

**② 결함의 절반(주입)이 빠져 있었다.** 템플릿 모드에서 미매칭 heading 은 빈칸이 아니라
`_fallback_function_description` 으로 **합성 설명이 붙은 섹션**으로 출력된다
(`docx_builder.py:3088-3092`). 즉 누락뿐 아니라 남의 프로젝트 함수명이 설명까지 달고
납품 문서에 실린다. (ASIL 은 지어내지 않는다 — `_finalize_function_fields` 는 payload
details 에만 적용된다. 과대주장 금지.)

## 그래서 무엇을 했나

`ok`/`success` **판정은 뒤집지 않는다** — 회사 양식이 의도된 부분집합인 경우가 실제로
있고, 그걸 실패로 만들면 정상 산출이 막힌다. 대신 **신원을 표면화**한다:

  - `_stats["template_identity"]` — 참조 SUDS 와 **같은 판정 함수**(`_reference_identity_verdict`)
  - 불일치 시 경고 1건 (누락·빈 heading 수치 동반)
  - `_gen_stats_result_fields` 화이트리스트에 추가 — **여기 안 넣으면 API 응답에서 잘린다**

⚠ 새 판정 함수를 만들지 않았다. 같은 질문("이 문서가 이 프로젝트 것인가")에 답하는
판정이 둘이 되면 이 저장소가 네 번 겪은 "한쪽만 고쳐짐" 이 된다.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

pytest.importorskip("docx")

REPO_ROOT = Path(__file__).resolve().parents[2]


@pytest.fixture(autouse=True)
def _no_reference_suds(monkeypatch, tmp_path):
    """저장소 고정 참조 SUDS(40.7MB)를 읽지 않게 한다 — 이 테스트와 무관한 축."""
    import config

    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_reference.docx"), raising=False)
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)


def _template(path: Path, headings: list[str]) -> Path:
    import docx

    d = docx.Document()
    d.add_heading("Software Unit Design", level=1)
    for h in headings:
        d.add_heading(h, level=4)
    d.save(str(path))
    return path


def _payload(project: str, fns: list[str]) -> dict:
    return {
        "project_name": project,
        "overview": "o", "requirements": "r", "interfaces": "i",
        "uds_frames": "u", "notes": "n",
        "function_details": {
            f"SwUFn_{i:04d}": {
                "id": f"SwUFn_{i:04d}", "name": n, "prototype": f"void {n}(void);",
                "description": "", "asil": "TBD", "related": "TBD",
                "precondition": "N/A", "inputs": [], "outputs": [],
                "globals_global": [], "globals_static": [], "called": "", "logic": "",
            }
            for i, n in enumerate(fns, start=1)
        },
    }


def _gen(tmp_path, template_stem: str, project: str, fns: list[str]):
    from report_gen.docx_builder import generate_uds_docx

    tpl = _template(tmp_path / f"({template_stem}) t.docx",
                    [f"SwUFn_{i:04d}: {n}" for i, n in enumerate(fns, start=1)])
    out = tmp_path / "out.docx"
    stats: dict = {}
    generate_uds_docx(str(tpl), _payload(project, fns), str(out), stats_out=stats)
    return stats


class TestTemplateIdentityIsRecorded:
    def test_foreign_template_is_flagged(self, tmp_path):
        stats = _gen(tmp_path, "HDPDM01_SUDS", "KJPDS02_PV", ["alpha"])

        assert "template_identity" in stats, "템플릿 신원이 기록되지 않았다"
        ident = stats["template_identity"]
        assert ident["same_project"] is not True
        assert ident["reason"] == "token_mismatch"
        assert "HDPDM01" in ident["ref_tokens"]
        assert "KJPDS02" in ident["payload_tokens"]

    def test_matching_template_is_confirmed(self, tmp_path):
        """대조군 — 신원이 맞으면 `True`. 이게 없으면 '항상 불일치' 도 통과한다."""
        stats = _gen(tmp_path, "KJPDS02_SwUDS", "KJPDS02_PV", ["alpha"])

        assert stats["template_identity"]["same_project"] is True
        assert stats["template_identity"]["reason"] == "token_match"

    def test_generation_is_not_blocked(self, tmp_path):
        """판정은 뒤집지 않는다 — 의도된 부분집합인 회사 양식이 실재한다."""
        stats = _gen(tmp_path, "HDPDM01_SUDS", "KJPDS02_PV", ["alpha"])
        assert stats["mode"] == "template"
        assert stats["payload_functions"] == 1

    def test_mismatch_is_warned_with_numbers(self, tmp_path, caplog):
        import logging

        with caplog.at_level(logging.WARNING, logger="report_generator"):
            _gen(tmp_path, "HDPDM01_SUDS", "KJPDS02_PV", ["alpha"])

        hits = [r.getMessage() for r in caplog.records if "템플릿의 프로젝트 신원" in r.getMessage()]
        assert hits, "템플릿 신원 불일치가 경고되지 않았다"
        assert "미반영" in hits[0], f"누락 수치가 경고에 없다: {hits[0]}"


class TestIdentityJudgementIsNotDuplicated:
    def test_reuses_the_reference_verdict_function(self):
        from report_gen import docx_builder

        source = Path(docx_builder.__file__).read_text(encoding="utf-8")
        assert "_template_identity = _reference_identity_verdict(" in source, (
            "템플릿 신원 판정이 참조 SUDS 판정과 다른 함수를 쓴다 — 같은 질문에 답하는 "
            "판정이 둘이 되면 한쪽만 고쳐진다(이 저장소 1위 재발 패턴)"
        )


class TestIdentityReachesTheApiSurface:
    """사이드카에만 있으면 산출물 검토자가 못 본다 — 화이트리스트에 있어야 한다."""

    def test_result_fields_whitelist_includes_it(self, tmp_path):
        from backend.helpers.uds import _gen_stats_result_fields
        from report_gen.docx_builder import gen_stats_path

        out = tmp_path / "out.docx"
        out.write_bytes(b"x")
        side = gen_stats_path(str(out))
        side.parent.mkdir(parents=True, exist_ok=True)
        side.write_text(json.dumps({
            "mode": "template",
            "template_source": "argument",
            "template_identity": {"same_project": False, "reason": "token_mismatch"},
            "payload_functions": 432,
            "matched_functions": 337,
        }, ensure_ascii=False), encoding="utf-8")

        fields = _gen_stats_result_fields(out)
        summary = fields["gen_stats_summary"]
        assert summary is not None
        assert summary.get("template_identity") == {
            "same_project": False, "reason": "token_mismatch",
        }, "새 축이 API 응답 화이트리스트에서 잘렸다"
