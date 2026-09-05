"""`.validation.md` 왕복 — **진짜 라이터 → 진짜 리더** 로 수치가 복원되는가.

## 왜 이 파일이 필요했나

2026-09-01 실측: `evidence.py` 가 영문 키(`Expected functions` 등)를 찾는데 라이터는
한국어(`Payload 함수 수` 등)를 썼다. 저장소 전체에서 영문 키를 쓰는 writer 는 **0곳**
— 세 필드는 한 번도 채워진 적이 없고, 그 값을 조건으로 그리던 화면
(`DocGenStatusBoard` 의 "문서에 빠진 함수 N개")은 **렌더된 적이 없다**.

기존 테스트가 못 잡은 이유가 핵심이다:

- `test_quality_evidence.py::test_fields_absent_in_old_reports_are_none_not_zero` 는
  **손으로 쓴 구판 .md** 를 먹여 `is None` 을 단언한다 → 결함이 있든 없든 통과한다.
- `DocGenStatusBoard.test.jsx` 는 `missing_from_docx: 9` 라는 **파이프라인이 만들 수
  없는 shape** 을 손으로 넣는다 → 실제로는 영원히 `null` 인데 초록이다.

둘 다 관측량을 단언하지 않는 가드다. 빠진 것은 **왕복**이었다: 라이터가 쓴 파일을
리더가 읽어 같은 수를 내는가. 라벨이 갈리는 순간 여기서 죽는다.

⚠ 이 파일은 라벨 문자열을 **테스트에 복제하지 않는다**. 복제하면 상수를 한 곳에 둔
  이유가 없어지고, 양쪽을 같이 고치면 가드가 조용히 따라간다.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, Tuple

import pytest

FN_INFO_BANNER = "Function Information"


# ==============================================================
# 픽스처 — 절 5개 중 payload 엔 alpha 하나. 나머지 4개가 '남의 함수 절'
# ==============================================================

def _template(path: Path) -> Path:
    import docx

    path.parent.mkdir(parents=True, exist_ok=True)
    d = docx.Document()
    d.add_heading("Software Unit Design", level=1)
    for title in ("SwUFn_001: alpha", "SwUFn_002: bravo", "SwUFn_003: charlie",
                  "SwUFn_004: delta", "SwUFn_005: echo (삭제)"):
        d.add_heading(title, level=2)
        t = d.add_table(rows=4, cols=6)
        t.rows[0].cells[0].text = FN_INFO_BANNER
        t.rows[1].cells[0].text = "Function Name"
    d.save(str(path))
    return path


def _payload(mode: str) -> Dict[str, Any]:
    alpha = {"name": "alpha", "prototype": "void alpha(void)",
             "inputs": ["none"], "outputs": ["none"], "logic": "does alpha"}
    return {
        "project_name": "KJPDS02", "overview": "o", "requirements": "r",
        "interfaces": "i", "uds_frames": "u", "notes": "n",
        "function_details": {"SwUFn_001": alpha},
        "function_details_by_name": {"alpha": alpha},
        "unmatched_headings": mode,
    }


def _run(tmp: Path, mode: str) -> Tuple[Dict[str, Any], Dict[str, Any], str]:
    """`(validator dict, reader dict, .md 원문)` — 한 번의 실제 생성으로."""
    from report_gen.docx_builder import generate_uds_docx
    from report_gen.evidence import read_docx_validation
    from report_gen.validation import (
        generate_uds_validation_report,
        validate_uds_docx_structure,
    )

    tpl = _template(tmp / "(KJPDS02_SwUDS) tpl.docx")
    out = tmp / f"out_{mode}.docx"
    pay = _payload(mode)
    generate_uds_docx(str(tpl), pay, str(out))
    # 실 파이프라인과 동일하게 payload 사이드카를 남긴다(`backend/helpers/uds.py`).
    out.with_suffix(".payload.json").write_text(
        json.dumps(pay, ensure_ascii=False), encoding="utf-8")

    rep = validate_uds_docx_structure(str(out))
    md = out.with_suffix(".validation.md")
    generate_uds_validation_report(str(out), str(md))
    return rep, read_docx_validation(md), md.read_text(encoding="utf-8")


# ⚠ 생성 한 번에 수 초가 걸린다. keep/drop 두 갈래만 만들고 재사용한다.
@pytest.fixture(scope="module")
def runs(tmp_path_factory) -> Dict[str, Tuple[Dict[str, Any], Dict[str, Any], str]]:
    pytest.importorskip("docx")
    base = tmp_path_factory.mktemp("valroundtrip")
    return {m: _run(base / m, m) for m in ("keep", "drop")}


# ==============================================================
# 1. 왕복 — 라이터가 쓴 수를 리더가 되찾는가
# ==============================================================

class TestTheReaderRecoversWhatTheWriterWrote:

    @pytest.mark.parametrize("mode", ["keep", "drop"])
    def test_payload_and_matched_counts_survive_the_round_trip(self, runs, mode):
        """예전엔 둘 다 영원히 `None` 이었다 — 라벨이 영문/한국어로 갈려서."""
        rep, got, _ = runs[mode]
        assert got["expected_functions"] == rep["expected_functions"] == 1
        assert got["matched_functions"] == rep["matched_functions"] == 1

    @pytest.mark.parametrize("mode", ["keep", "drop"])
    def test_empty_heading_count_survives(self, runs, mode):
        """리더에 대응 키가 **아예 없어** 화면에 닿은 적이 없던 값."""
        rep, got, _ = runs[mode]
        assert got["headings_without_payload"] == rep["headings_without_payload_count"]
        assert got["headings_without_payload"] is not None

    def test_a_label_change_on_one_side_alone_cannot_pass(self, runs):
        """라벨은 한 곳에서만 온다 — 리더가 자체 문자열을 들면 안 된다.

        ⚠ 값 대조만으로는 부족하다. 양쪽이 같은 상수를 **실제로** 참조하는지 본다.
        """
        import report_gen.validation_labels as VL
        from report_gen import evidence, validation

        assert validation.VL is VL
        assert evidence.VL is VL

    def test_renaming_a_label_is_loud_because_old_artifacts_use_it(self):
        """라벨은 **디스크에 남은 산출물**이 쓰는 문자열이다 — 조용히 못 바꾼다.

        상수를 한 곳에 두면 라이터/리더는 항상 함께 움직인다(그게 목적이다). 그래서
        "한쪽만 바꾸는" 실수는 구조적으로 불가능해졌지만, **양쪽을 같이 바꾸는** 것은
        여전히 가능하고 그 순간 이미 생성돼 있는 `.validation.md` 가 전부 안 읽힌다.
        새 산출물만 보는 왕복 가드로는 그걸 못 잡는다(뮤테이션 M203 이 생존해서 드러났다).

        여기서 문자열을 복제하는 것은 의도다 — 이 테스트의 **목적 자체가** 변경을
        시끄럽게 만드는 것이다. 바꿔야 한다면: 리더에 옛 라벨 별칭을 먼저 넣고,
        그 다음 이 목록을 갱신할 것.
        """
        import report_gen.validation_labels as VL

        assert VL.LABEL_EXPECTED_FUNCTIONS == "Payload 함수 수"
        assert VL.LABEL_MATCHED_FUNCTIONS == "문서에 실린 함수(매칭)"
        assert VL.LABEL_MISSING_FROM_DOCX == "문서에 없는 소스 함수"
        assert VL.LABEL_HEADINGS_WITHOUT_PAYLOAD == "데이터 없는 템플릿 heading"
        assert VL.LABEL_DROPPED_HEADINGS == "문서에서 제거된 heading"
        assert VL.LABEL_UNMATCHED_MODE == "남의 함수 절 처리"

    def test_the_label_text_is_not_duplicated_in_the_reader(self):
        """리더 소스에 라벨 문자열 리터럴이 있으면 상수를 둔 의미가 없다."""
        import report_gen.validation_labels as VL

        src = Path(__import__("report_gen.evidence", fromlist=["x"]).__file__)
        text = src.read_text(encoding="utf-8")
        for name in ("LABEL_EXPECTED_FUNCTIONS", "LABEL_MATCHED_FUNCTIONS",
                     "LABEL_MISSING_FROM_DOCX", "LABEL_HEADINGS_WITHOUT_PAYLOAD",
                     "LABEL_DROPPED_HEADINGS", "LABEL_UNMATCHED_MODE"):
            literal = getattr(VL, name)
            assert f'"{literal}"' not in text, f"{name} 문자열이 리더에 복제돼 있다"


# ==============================================================
# 2. 제거된 절이 증거에 남는가 — 절단의 침묵 차단
# ==============================================================

class TestRemovalIsVisibleInTheEvidence:

    def test_drop_records_how_many_sections_were_removed(self, runs):
        rep, got, md = runs["drop"]
        assert rep["dropped_heading_count"] == 3
        assert got["dropped_headings"] == 3
        assert got["unmatched_headings_mode"] == "drop"

    def test_keep_removes_nothing_and_says_so(self, runs):
        rep, got, _ = runs["keep"]
        assert rep["dropped_heading_count"] == 0
        assert got["dropped_headings"] == 0
        assert got["unmatched_headings_mode"] == "keep"

    def test_the_report_warns_that_the_empty_count_only_counts_survivors(self, runs):
        """`drop` 은 '빈 heading' 을 4 → 1 로 **줄인다**. 그 사실이 없으면 얇아진
        문서가 거의 완결된 것처럼 보인다 — 실측한 바로 그 착시."""
        _, _, md = runs["drop"]
        assert "문서에서 제거" in md
        assert "남은 것만" in md

    def test_drop_really_shrinks_the_visible_empty_count(self, runs):
        """착시가 실재함을 대조군으로 고정한다 — 이 차이가 사라지면 이 가드는 무의미."""
        keep_rep = runs["keep"][0]
        drop_rep = runs["drop"][0]
        assert drop_rep["headings_without_payload_count"] < keep_rep[
            "headings_without_payload_count"]
        # 사라진 만큼이 제거 수로 설명돼야 한다(보존 법칙).
        assert (keep_rep["headings_without_payload_count"]
                - drop_rep["headings_without_payload_count"]
                == drop_rep["dropped_heading_count"])

    def test_keep_report_does_not_claim_a_removal(self, runs):
        _, _, md = runs["keep"]
        assert "남은 것만" not in md


# ==============================================================
# 3. 미측정은 0 이 아니다 (음성 대조군)
# ==============================================================

class TestUnmeasuredIsNotZero:

    def test_absent_gen_stats_yields_none_not_zero(self, tmp_path):
        """사이드카가 없으면 **모드를 추정하지 않는다**.

        `keep` 으로 가정하면 "제거 없음" 이라는 거짓 단언이 되고, 그게 이 라운드가
        고치는 결함 자체다.
        """
        from report_gen.validation import _read_drop_stats

        assert _read_drop_stats(tmp_path / "nope.docx") == (None, None)

    def test_corrupt_gen_stats_is_unmeasured_not_zero(self, tmp_path):
        from report_gen.docx_builder import gen_stats_path
        from report_gen.validation import _read_drop_stats

        docx_path = tmp_path / "x.docx"
        gen_stats_path(str(docx_path)).write_text("{ broken", encoding="utf-8")
        assert _read_drop_stats(docx_path) == (None, None)

    def test_boolean_is_not_a_count(self, tmp_path):
        """`True` 는 `isinstance(x, int)` 를 통과한다 — 1 로 읽히면 거짓 제거 보고."""
        from report_gen.docx_builder import gen_stats_path
        from report_gen.validation import _read_drop_stats

        docx_path = tmp_path / "y.docx"
        gen_stats_path(str(docx_path)).write_text(
            json.dumps({"dropped_heading_count": True, "unmatched_headings_mode": "keep"}),
            encoding="utf-8")
        assert _read_drop_stats(docx_path)[0] is None

    def test_report_omits_the_line_when_unmeasured(self, tmp_path):
        """미측정이면 `0건` 을 적지 않는다 — 0 은 '제거 없음' 이라는 단언이다."""
        import report_gen.validation_labels as VL
        from report_gen.validation import generate_uds_validation_report

        out = tmp_path / "r.md"
        fake = {"docx_path": "x.docx", "ok": True, "table_count": 1,
                "dropped_heading_count": None, "unmatched_headings_mode": None}
        import report_gen.validation as V
        orig = V.validate_uds_docx_structure
        V.validate_uds_docx_structure = lambda _p: fake        # noqa: ARG005
        try:
            generate_uds_validation_report("x.docx", str(out))
        finally:
            V.validate_uds_docx_structure = orig
        text = out.read_text(encoding="utf-8")
        assert VL.LABEL_DROPPED_HEADINGS not in text
        assert VL.LABEL_UNMATCHED_MODE not in text


# ==============================================================
# 4. 값 파싱 — 구판 산출물이 단위를 backtick 안에 넣었다
# ==============================================================

class TestLegacyUnitSuffix:

    def _read(self, tmp_path: Path, body: str) -> Dict[str, Any]:
        from report_gen.evidence import read_docx_validation

        p = tmp_path / "old.validation.md"
        p.write_text(body, encoding="utf-8")
        return read_docx_validation(p)

    def test_unit_inside_backticks_still_parses(self, tmp_path):
        """`` `629건` `` → 629. 키를 고쳐도 값이 안 들어오던 두 번째 겹."""
        import report_gen.validation_labels as VL

        got = self._read(tmp_path, "# R\n- OK: `True`\n"
                         f"- ⚠ {VL.LABEL_MISSING_FROM_DOCX}: `629건` (예: a)\n")
        assert got["missing_from_docx"] == 629

    def test_warning_marker_prefix_does_not_hide_the_key(self, tmp_path):
        """사람이 읽는 `⚠` 가 키에 붙어 있으면 대조가 실패한다."""
        import report_gen.validation_labels as VL

        got = self._read(tmp_path, "# R\n- OK: `True`\n"
                         f"- ⚠ {VL.LABEL_HEADINGS_WITHOUT_PAYLOAD}: `4`건\n")
        assert got["headings_without_payload"] == 4

    def test_non_numeric_stays_none(self, tmp_path):
        """`대조 불가(…)` 를 0 으로 읽으면 '누락 없음' 이 된다."""
        import report_gen.validation_labels as VL

        got = self._read(tmp_path, "# R\n- OK: `True`\n"
                         f"- {VL.LABEL_EXPECTED_FUNCTIONS}: `{VL.VALUE_UNCOMPARABLE}`\n")
        assert got["expected_functions"] is None

    def test_strict_int_is_still_strict_for_plain_fields(self):
        """`_as_int` 까지 느슨해지면 쓰레기를 숫자로 읽는다 — 관용은 건수 전용이다."""
        from report_gen.evidence import _as_count, _as_int

        assert _as_int("12개") is None
        assert _as_count("12개") == 12


# ==============================================================
# 5. 소비처 계약 — 화면이 읽는 키가 실제로 존재하는가
# ==============================================================

class TestTheBoardReadsKeysThatExist:

    def test_every_key_the_board_reads_is_produced_by_the_reader(self, runs):
        """보드가 `val.<키>` 로 읽는 이름이 리더 출력에 있어야 한다.

        `missing_from_docx` 는 **영원히 null** 이었는데도 보드가 그걸 조건으로
        문장을 그리고 있었다(=죽은 코드). 소스에서 키를 뽑아 대조한다.
        """
        import re

        board = Path("frontend-v2/src/components/sections/DocGenStatusBoard.jsx")
        if not board.exists():                       # pragma: no cover - 경로 이동 대비
            pytest.skip(f"보드 파일 없음: {board}")
        text = board.read_text(encoding="utf-8")
        used = set(re.findall(r"\bval\.([a-z_]+)", text))
        produced = set(runs["drop"][1].keys()) | {"reason"}
        missing = sorted(used - produced)
        assert not missing, f"보드가 읽는데 리더가 안 내는 키: {missing}"
