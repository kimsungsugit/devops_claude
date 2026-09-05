# tests/unit/test_provenance_vocabulary.py
"""출처(provenance) 어휘 — 미기록 라벨과 "약한 출처" 판정 단일화 (D1).

## 실측 (2026-07-31, 고치기 전)

`_resolve_related_asil_desc` 가 출처 미기록을 **세 축 모두 무조건 `"inference"`** 로
확정했다. 사람이 쓴 설명·실제 등급 `C`·실제 `SwFn_07` 을 넣고 생성했더니:

| 필드 | 값 | 옛 출처 | 실제 |
|---|---|---|---|
| description | `CAN 수신 버퍼를 검증한다` | `inference` (0.60, 표기 "추론") | 사람이 씀 |
| asil | `C` | `inference` | 주어진 값 |
| related | `SwFn_07` | `inference` | 주어진 값 |

아무것도 추론하지 않았는데 보고서 표에 "추론" 으로 찍힌다. 반대로 **근거 없이 자리만
채운 값**(`asil=""` 로 비우면서 `related_source="inference"`)도 같은 0.60 을 받아
**두 방향으로 동시에** 틀렸다.

`report_gen/validation.py` 의 `src_labels`/`src_score` 에는 `unknown`(미상, 0.30)과
`default`(기본값·근거 없음, 0.30)가 **이미 1급 어휘**로 있었다 — 쓰지 않았을 뿐이다.
생산자(`function_analyzer.py`)는 이미 **자기가 한 행위에 묶어** 라벨한다(합성했을 때만
`inference`, QM 을 채웠을 때만 `default`). 그 규약을 미기록 케이스로 확장한다.

## 왜 `provenance.py` 로 단일화했나

`unknown` 을 도입하려 하니 **"이 출처는 약한가?"** 판정이 저장소 7곳에 서로 다른
리터럴로 복제돼 있었다(`{"", "inference"}` / `{"inference", "rule", ""}` /
`{"inference","module_inherit","default","rule",""}` / `!= "inference"` …).
그대로 뒀으면 5곳이 `unknown` 을 **강한 출처로 오인**해, "출처를 모른다" 가 "확정됐다"
처럼 굳어 더 나은 근거가 덮어쓰지 못했을 것이다 — 라벨 하나 고치려다 다섯 곳에 구멍을
내는, 이 저장소가 반복해 겪은 whack-a-mole 이다.
"""
from __future__ import annotations

import ast
from pathlib import Path

import pytest

from report_gen.provenance import (
    PLACEHOLDER_VALUES,
    WEAK_SOURCES,
    is_weak_source,
    unrecorded_source,
)

# --------------------------------------------------------------
# 1. 미기록 라벨 — 값을 보고 정한다
# --------------------------------------------------------------

class TestUnrecordedSource:
    @pytest.mark.parametrize("value", ["", "TBD", "N/A", "n/a", "-", "None", None, "  "])
    def test_placeholder_value_is_default_not_inference(self, value):
        """근거 없이 칸만 채운 상태 = `default`(0.30). 추론이 아니다.

        뮤테이션: `unrecorded_source` 가 항상 `"inference"` 를 내게 하면 실패.
        """
        assert unrecorded_source(value) == "default"

    @pytest.mark.parametrize("value", ["C", "D", "QM", "SwFn_07", "CAN 수신 버퍼를 검증한다"])
    def test_real_value_without_source_is_unknown(self, value):
        """**핵심** — 실제 값인데 출처 미기록이면 `unknown`(미상). 추론이라 주장하지 않는다."""
        assert unrecorded_source(value) == "unknown"

    def test_generic_text_is_inference(self):
        """생성기가 만든 일반 문구는 실제로 추론이 있었다 — `inference` 가 맞다."""
        assert unrecorded_source("alpha 관련 연산을 수행한다", generic=True) == "inference"

    def test_generic_flag_does_not_override_placeholder(self):
        """값이 없으면 generic 여부와 무관하게 `default` — 없는 것을 추론했다고 하지 않는다."""
        assert unrecorded_source("", generic=True) == "default"

    def test_all_returned_labels_are_first_class_vocabulary(self):
        """반환 라벨은 전부 `validation.py` 의 점수표에 있어야 한다.

        없으면 `_norm_src` 가 조용히 `unknown` 으로 접고 어휘 드리프트 카운터만 오른다.
        """
        src = Path("report_gen/validation.py").read_text(encoding="utf-8")
        for label in ("default", "inference", "unknown"):
            assert f'"{label}":' in src, f"{label} 이 validation.py 어휘에 없다"


# --------------------------------------------------------------
# 2. 약한 출처 판정
# --------------------------------------------------------------

class TestWeakSource:
    @pytest.mark.parametrize("src", ["", "unknown", "default", "inference", "module_inherit", "rule"])
    def test_weak_sources_are_overwritable(self, src):
        assert is_weak_source(src) is True

    @pytest.mark.parametrize("src", ["comment", "sds", "srs", "uds", "swcom", "rag", "call_graph"])
    def test_evidence_backed_sources_are_not_weak(self, src):
        """음성 대조군 — 실제 근거를 본 출처를 약하다고 하면 좋은 값이 덮어써진다."""
        assert is_weak_source(src) is False

    def test_unknown_must_be_weak(self):
        """**이 fix 의 핵심 불변식.**

        `unknown` 이 강한 출처로 취급되면 "출처를 모른다" 가 "확정됐다" 처럼 굳어
        뒤따르는 주석·SDS·SRS 근거가 덮어쓰지 못한다.
        뮤테이션: `WEAK_SOURCES` 에서 `"unknown"` 을 빼면 실패.
        """
        assert is_weak_source("unknown") is True
        assert "unknown" in WEAK_SOURCES

    @pytest.mark.parametrize("src", [None, "  ", "INFERENCE", "Default", "UNKNOWN"])
    def test_case_and_whitespace_insensitive(self, src):
        assert is_weak_source(src) is True

    def test_placeholder_values_are_distinct_from_sources(self):
        """자리표시자(값)와 출처 라벨은 다른 축이다 — 섞으면 `none` 같은 값이 출처가 된다."""
        assert PLACEHOLDER_VALUES & WEAK_SOURCES == {""}


# --------------------------------------------------------------
# 3. 복제 금지 — 판정은 한 곳에만
# --------------------------------------------------------------

_DUP_SITES = [
    "report_gen/docx_builder.py",
    "report_gen/function_analyzer.py",
    "backend/routers/local.py",
]


class TestNoDuplicatedJudgment:
    @pytest.mark.parametrize("path", _DUP_SITES)
    def test_no_hardcoded_weak_source_set(self, path):
        """`{"inference", ...}` 같은 집합 리터럴로 판정을 재현하지 않는다.

        새 라벨이 생길 때 한쪽만 갱신되어 조용히 갈라진다 — 이 모듈이 생긴 이유다.
        뮤테이션: 어느 사이트든 리터럴 집합을 되살리면 실패.
        """
        tree = ast.parse(Path(path).read_text(encoding="utf-8"))
        offenders = []
        for node in ast.walk(tree):
            if not isinstance(node, ast.Set):
                continue
            vals = {e.value for e in node.elts
                    if isinstance(e, ast.Constant) and isinstance(e.value, str)}
            # 출처 라벨 2개 이상을 모아 둔 집합 = 판정 복제 신호
            # ⚠ 예전엔 `& WEAK_SOURCES - {""}` 라 **빈 문자열을 세지 않았다**.
            #   그래서 이 저장소에서 실제로 가장 흔한 복제 형태인 `{"", "inference"}`
            #   (원소 2개, 그중 하나가 `""`)가 1개로 세어져 **영영 안 걸렸다**.
            #   `""` 도 어엿한 출처 라벨(= 아무도 기록 안 함)이므로 함께 센다.
            #   실측(2026-08-03): 조여도 스캔 대상 3파일에서 신규 탐지 0건 = 오탐 없음.
            if len(vals & WEAK_SOURCES) >= 2:
                offenders.append(f"{path}:{node.lineno} {sorted(vals)}")
        assert not offenders, f"약한 출처 판정이 복제됐다: {offenders}"

    @pytest.mark.parametrize("path", _DUP_SITES)
    def test_site_imports_the_shared_predicate(self, path):
        src = Path(path).read_text(encoding="utf-8")
        assert "from report_gen.provenance import" in src, f"{path} 가 단일 출처를 안 쓴다"

    def test_remaining_literal_site_is_intentional(self):
        """⚠ **전환하지 말 것.** `report_gen/requirements.py:1590` 의 `{"", "inference"}` 는
        의도된 잔존이다.

        2026-08-03 실측으로 판정이 **뒤집혔다**. 이 docstring 은 예전에
        *"전환하면 여기서 실패하므로 그때 이 테스트를 지우고 위 목록에 파일을 추가하면
        된다"* 고 **전환을 지시**하고 있었다. 세 가지 이유로 틀렸다:

        1. **방향이 세탁이다.** 이 분기가 붙이는 라벨은 `sds_match` 이고
           `provenance.SOURCE_ALIASES` 가 그걸 `sds`(**0.95**)로 접는다. 리터럴을
           `is_weak_source()` 로 넓히면 `default`/`unknown`/`generated_doc`(0.30)·
           `module_inherit`(0.70)·`rule`(0.75) 이던 행이 0.95 를 받는다.
        2. **`sds_match` 의 뜻이 반대다.** `requirements.py:1586-1591` 을 보면 이건
           설명이 SDS 에서 **오지 않았을 때**의 else 분기다. 그런 행을 더 만드는 게
           개선일 리 없다.
        3. **복제 사이트는 짝이 아니다.** `tools/generate_uds_local.py:206` 의 같은
           리터럴은 `hsis` 를 붙인다(다른 라벨). 게다가 그 파일은 `:16-18` 에서
           `sys.path` 를 `D:\\Project\\devops\\260105` 로 밀어 넣고 거기서 `report_gen` 을
           가져오는데 **그 트리엔 `provenance.py` 가 없다** — 단일 출처를 import 하면
           스크립트가 import 에서 죽는다.

        `sds_match` 별칭이 0.95 를 받는 것 자체가 과대 신용이 아닌지는 별건(계획서 #20).
        그 판단이 먼저다.
        """
        src = Path("report_gen/requirements.py").read_text(encoding="utf-8")
        assert 'in {"", "inference"}' in src, (
            "requirements.py:1590 의 리터럴이 사라졌다. 이건 **의도된 잔존**이다 — "
            "is_weak_source() 로 넓히면 0.30~0.75 이던 행이 sds_match→sds(0.95) 를 받는다. "
            "되돌리거나, 바꿀 근거가 생겼다면 이 테스트의 docstring 부터 갱신할 것"
        )


# --------------------------------------------------------------
# 4. 통합 — 실제 빌더가 이 어휘를 쓴다
# --------------------------------------------------------------

@pytest.fixture
def gen(tmp_path, monkeypatch):
    import docx

    import config
    from report_gen import docx_builder

    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)
    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH", str(tmp_path / "none.docx"), raising=False)

    def _run(**overrides):
        tpl = tmp_path / "t.docx"
        d = docx.Document()
        d.add_heading("Software Unit Design", level=1)
        d.add_heading("SwUFn_0001: alpha", level=4)
        d.save(str(tpl))
        info = {"id": "SwUFn_0001", "name": "alpha", "prototype": "void alpha(void);",
                "description": "", "asil": "", "related": "", "precondition": "",
                "inputs": [], "outputs": [], "globals_global": [], "globals_static": [],
                "called": "", "logic": ""}
        info.update(overrides)
        docx_builder.generate_uds_docx(
            str(tpl),
            {"project_name": "T", "overview": "o", "requirements": "r", "interfaces": "i",
             "uds_frames": "u", "notes": "n", "function_details": {"SwUFn_0001": info}},
            str(tmp_path / "out.docx"))
        return info

    return _run


class TestBuilderIntegration:
    def test_human_written_values_are_not_labelled_inference(self, gen):
        """**이 라운드의 대표 케이스.** 실제로 생성해서 확인한다."""
        info = gen(description="CAN 수신 버퍼를 검증한다", asil="C", related="SwFn_07")
        assert info["asil"] == "C"
        assert info["related"] == "SwFn_07"
        assert info["description_source"] == "unknown"
        assert info["asil_source"] == "unknown"
        assert info["related_source"] == "unknown"

    def test_empty_input_still_reports_default_not_inference(self, gen):
        """값이 없는 축은 `default`(근거 없음) — 0.60 을 받지 않는다."""
        info = gen()
        assert info["asil_source"] == "default"
        assert info["related_source"] == "default"

    def test_synthesized_description_is_still_inference(self, gen):
        """음성 대조군 — 진짜 합성은 `inference` 가 맞다. 전부 unknown 으로 밀지 않는다."""
        info = gen()
        assert info["description_source"] == "inference"

    def test_comment_evidence_still_upgrades(self, gen):
        """**회귀 방지** — `unknown` 이 약한 출처로 유지되어야 주석 근거가 덮어쓴다.

        뮤테이션: `WEAK_SOURCES` 에서 `"unknown"` 을 빼면 여기서 실패.
        """
        info = gen(asil="C", comment_asil="D", comment_related="SwFn_09")
        assert info["asil"] == "D"
        assert info["asil_source"] == "comment"
        assert info["related_source"] == "comment"
