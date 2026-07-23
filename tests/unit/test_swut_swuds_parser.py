"""SwUDS docx parser 단위 테스트 (16차 라운드).

합성 docx로 SwUDS 양식 시뮬레이션 + 함수 ID 추출 검증.
"""
from __future__ import annotations

import io
from pathlib import Path
import sys

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from backend.services.swut_swuds_parser import (  # noqa: E402
    DOCX_MAX_BYTES,
    parse_swuds_docx,
)


def _build_swuds_docx(function_ids: list[str], with_descriptions: bool = False) -> bytes:
    """합성 SwUDS docx 생성 — Hyundai 양식 (heading 'SwUFn_XXXX' + table)."""
    from docx import Document  # type: ignore

    doc = Document()
    for fn_id in function_ids:
        doc.add_paragraph(f"{fn_id} — Sample Description Heading")
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Description"
        if with_descriptions:
            tbl.cell(0, 1).text = f"Function {fn_id} does X, Y, Z."
        else:
            tbl.cell(0, 1).text = ""
        tbl.cell(1, 0).text = "Interface"
        tbl.cell(1, 1).text = "void"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestSwUDSParser:
    def test_empty_bytes_returns_ok_false(self):
        warnings: list[str] = []
        result = parse_swuds_docx(b"", parse_warnings=warnings)
        assert not result.ok
        assert any("비어있음" in w for w in warnings)

    def test_oversize_rejected(self):
        warnings: list[str] = []
        oversize = b"x" * (DOCX_MAX_BYTES + 1)
        result = parse_swuds_docx(oversize, parse_warnings=warnings)
        assert not result.ok
        assert any("한도 초과" in w for w in warnings)

    def test_invalid_docx_returns_ok_false(self):
        warnings: list[str] = []
        result = parse_swuds_docx(b"not a docx file", parse_warnings=warnings)
        assert not result.ok
        assert any("로드 실패" in w for w in warnings)

    def test_no_swufn_heading_returns_ok_false(self):
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_paragraph("Just a regular paragraph")
        doc.add_paragraph("Another one")
        buf = io.BytesIO()
        doc.save(buf)

        warnings: list[str] = []
        result = parse_swuds_docx(buf.getvalue(), parse_warnings=warnings)
        assert not result.ok
        assert any("미발견" in w for w in warnings)

    def test_extracts_3_function_ids(self):
        docx_bytes = _build_swuds_docx(["SwUFn_0001", "SwUFn_0002", "SwUFn_0101"])
        result = parse_swuds_docx(docx_bytes)
        assert result.ok
        assert result.function_ids == {"SwUFn_0001", "SwUFn_0002", "SwUFn_0101"}
        assert len(result.entries) == 3

    def test_description_extracted_from_table(self):
        docx_bytes = _build_swuds_docx(["SwUFn_0101"], with_descriptions=True)
        result = parse_swuds_docx(docx_bytes)
        assert result.ok
        entry = result.entries[0]
        assert entry.function_id == "SwUFn_0101"
        assert "Function SwUFn_0101 does X" in entry.description

    def test_heading_text_preserved(self):
        docx_bytes = _build_swuds_docx(["SwUFn_0501"])
        result = parse_swuds_docx(docx_bytes)
        assert result.ok
        assert "Sample Description Heading" in result.entries[0].heading_text

    def test_to_dict_contains_qualification_meta(self):
        docx_bytes = _build_swuds_docx(["SwUFn_0001"])
        result = parse_swuds_docx(docx_bytes)
        d = result.to_dict()
        assert d["ok"] is True
        assert "tool_qualification" in d
        assert d["tool_qualification"]["evidence_class"] == "auto-generated draft"

    def test_heading_without_following_table_still_captured(self):
        """heading 직후 paragraph만 있어도 (table 없이) entry 보존."""
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_paragraph("SwUFn_9999")
        doc.add_paragraph("Just a paragraph, no table.")
        buf = io.BytesIO()
        doc.save(buf)

        result = parse_swuds_docx(buf.getvalue())
        assert result.ok
        assert "SwUFn_9999" in result.function_ids


class TestSwUDSConsistencyIntegration:
    """16차: SwUDS↔SwUTS 비교가 build_coverage_report로 전달되는지."""

    def test_swuds_function_ids_added_to_consistency_when_all_match(self, tmp_path):
        """SwUTS 모든 함수가 SwUDS에 있으면 PASS."""
        import openpyxl
        from backend.services.swut_coverage_aggregator import build_coverage_report, CoverageBuildMeta
        from backend.services.swut_input_adapter import (
            EnvironmentData, ExecutionRow, FunctionCoverage, SwUTSession,
        )
        # session: SwUFn_0001 만 보유
        env = EnvironmentData(
            env_name="SWTE_01", component_name="X",
            test_cases={"SwUFn_0001.001": [object()]},
            test_results={"SwUFn_0001.001": ExecutionRow(tc_name="SwUFn_0001.001", passed=True)},
            function_coverage=[FunctionCoverage(unit_id="SwUFn_0001", name="X")],
        )
        session = SwUTSession(environments=[env])
        meta = CoverageBuildMeta(release_sw_version="1.0.0", test_date="2024-02-19")

        # 최소 template
        wb_tmpl = openpyxl.Workbook()
        wb_tmpl.remove(wb_tmpl.active)
        wb_tmpl.create_sheet("Cover")
        wb_tmpl.create_sheet("Test Summary")
        wb_tmpl.create_sheet("1.Traceability")
        wb_tmpl.create_sheet("2.Consistency")
        cov3 = wb_tmpl.create_sheet("3. Coverage")
        cov3["A1"] = "Statement Coverage"
        cov3["A6"] = "Unit ID"
        wb_tmpl.create_sheet("History")
        tmpl_buf = io.BytesIO()
        wb_tmpl.save(tmpl_buf)

        # SwUDS function_ids: 동일한 SwUFn_0001
        result = build_coverage_report(
            session, meta, tmpl_buf.getvalue(),
            swuds_function_ids={"SwUFn_0001"},
        )

        wb = openpyxl.load_workbook(io.BytesIO(result.xlsx_bytes))
        cons = wb["2.Consistency"]
        # row 8 = SwUDS↔SwUTS 매핑
        items = [str(cons.cell(r, 1).value or "") for r in range(4, 10)]
        result_col = [str(cons.cell(r, 4).value or "") for r in range(4, 10)]
        # 5번째 row가 SwUDS 매핑
        swuds_rows = [(i, item) for i, item in enumerate(items) if "SwUDS" in item]
        assert swuds_rows, f"SwUDS 매핑 row 없음: {items}"
        idx = swuds_rows[0][0]
        assert result_col[idx] == "PASS"
        # summary 갱신
        assert result.summary.get("consistency_self_check_rows") == 5
        assert result.summary.get("consistency_swuds_compared") is True

    def test_swuds_missing_function_marked_fail(self, tmp_path):
        """SwUDS에는 있는데 SwUTS에 없는 함수가 있으면 FAIL."""
        import openpyxl
        from backend.services.swut_coverage_aggregator import build_coverage_report, CoverageBuildMeta
        from backend.services.swut_input_adapter import (
            EnvironmentData, ExecutionRow, FunctionCoverage, SwUTSession,
        )
        env = EnvironmentData(
            env_name="SWTE_01", component_name="X",
            test_cases={"SwUFn_0001.001": [object()]},
            test_results={"SwUFn_0001.001": ExecutionRow(tc_name="SwUFn_0001.001", passed=True)},
            function_coverage=[FunctionCoverage(unit_id="SwUFn_0001", name="X")],
        )
        session = SwUTSession(environments=[env])
        meta = CoverageBuildMeta(release_sw_version="1.0.0", test_date="2024-02-19")

        wb_tmpl = openpyxl.Workbook()
        wb_tmpl.remove(wb_tmpl.active)
        wb_tmpl.create_sheet("Cover")
        wb_tmpl.create_sheet("Test Summary")
        wb_tmpl.create_sheet("1.Traceability")
        wb_tmpl.create_sheet("2.Consistency")
        cov3 = wb_tmpl.create_sheet("3. Coverage")
        cov3["A1"] = "Statement Coverage"
        cov3["A6"] = "Unit ID"
        wb_tmpl.create_sheet("History")
        tmpl_buf = io.BytesIO()
        wb_tmpl.save(tmpl_buf)

        # SwUDS: SwUFn_0001 + SwUFn_0002 (SwUTS에 0002 없음)
        result = build_coverage_report(
            session, meta, tmpl_buf.getvalue(),
            swuds_function_ids={"SwUFn_0001", "SwUFn_0002"},
        )

        wb = openpyxl.load_workbook(io.BytesIO(result.xlsx_bytes))
        cons = wb["2.Consistency"]
        items = [str(cons.cell(r, 1).value or "") for r in range(4, 10)]
        result_col = [str(cons.cell(r, 4).value or "") for r in range(4, 10)]
        notes = [str(cons.cell(r, 5).value or "") for r in range(4, 10)]
        swuds_rows = [(i, item) for i, item in enumerate(items) if "SwUDS" in item]
        assert swuds_rows
        idx = swuds_rows[0][0]
        assert result_col[idx] == "FAIL"
        assert "SwUFn_0002" in notes[idx], f"누락 함수 표시 누락: {notes[idx]}"

    def test_swuds_not_provided_partial_label_kept(self, tmp_path):
        """swuds_function_ids=None이면 incomplete_sheets에 partial 라벨 유지."""
        import openpyxl
        from backend.services.swut_coverage_aggregator import build_coverage_report, CoverageBuildMeta
        from backend.services.swut_input_adapter import (
            EnvironmentData, ExecutionRow, FunctionCoverage, SwUTSession,
        )
        env = EnvironmentData(
            env_name="SWTE_01", component_name="X",
            test_cases={"SwUFn_0001.001": [object()]},
            test_results={"SwUFn_0001.001": ExecutionRow(tc_name="SwUFn_0001.001", passed=True)},
            function_coverage=[FunctionCoverage(unit_id="SwUFn_0001", name="X")],
        )
        session = SwUTSession(environments=[env])
        meta = CoverageBuildMeta(release_sw_version="1.0.0", test_date="2024-02-19")

        wb_tmpl = openpyxl.Workbook()
        wb_tmpl.remove(wb_tmpl.active)
        wb_tmpl.create_sheet("Cover")
        wb_tmpl.create_sheet("Test Summary")
        wb_tmpl.create_sheet("1.Traceability")
        wb_tmpl.create_sheet("2.Consistency")
        cov3 = wb_tmpl.create_sheet("3. Coverage")
        cov3["A1"] = "Statement Coverage"
        cov3["A6"] = "Unit ID"
        wb_tmpl.create_sheet("History")
        tmpl_buf = io.BytesIO()
        wb_tmpl.save(tmpl_buf)

        result = build_coverage_report(session, meta, tmpl_buf.getvalue())
        assert any("partial" in s for s in result.incomplete_sheets)
        assert result.summary.get("consistency_swuds_compared") is False


class TestSwUDSSchemaIntegration:
    """16차: SwUTBuildRequest에 swuds_docx_path 필드 추가."""

    def test_schema_accepts_empty_swuds_path(self):
        from backend.schemas import SwUTBuildRequest
        req = SwUTBuildRequest(
            project_id="HDPDM01",
            release_sw_version="1.0.0",
            test_date="2024-02-19",
        )
        assert req.swuds_docx_path == ""

    def test_schema_accepts_explicit_swuds_path(self):
        from backend.schemas import SwUTBuildRequest
        req = SwUTBuildRequest(
            project_id="HDPDM01",
            release_sw_version="1.0.0",
            test_date="2024-02-19",
            swuds_docx_path="U:/docs/SwUDS_v3.docx",
        )
        assert req.swuds_docx_path == "U:/docs/SwUDS_v3.docx"

    def test_schema_rejects_swuds_path_with_newline(self):
        """W8 패턴: 줄바꿈 차단 (헤더 인젝션 안전)."""
        from backend.schemas import SwUTBuildRequest
        with pytest.raises(Exception):  # noqa: B017
            SwUTBuildRequest(
                project_id="HDPDM01",
                release_sw_version="1.0.0",
                test_date="2024-02-19",
                swuds_docx_path="U:/docs\r\nX-Injected: evil",
            )

    def test_schema_rejects_swuds_path_too_long(self):
        from backend.schemas import SwUTBuildRequest
        with pytest.raises(Exception):  # noqa: B017
            SwUTBuildRequest(
                project_id="HDPDM01",
                release_sw_version="1.0.0",
                test_date="2024-02-19",
                swuds_docx_path="U:/" + "a" * 600,
            )


def _build_swuds_docx_with_asil(
    function_id: str, asil_raw: str | None, asil_label: str = "ASIL",
) -> bytes:
    """32차 W28: 합성 SwUDS docx — heading 다음 표에 ASIL 라벨/값 row 포함.

    asil_raw=None이면 ASIL row 생략 (라벨 미발견 fail-safe 시뮬레이션).
    """
    from docx import Document  # type: ignore
    doc = Document()
    doc.add_paragraph(f"{function_id} — Sample Heading")
    rows = 3 if asil_raw is not None else 2
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.cell(0, 0).text = "Description"
    tbl.cell(0, 1).text = "Sample fn"
    tbl.cell(1, 0).text = "Interface"
    tbl.cell(1, 1).text = "void"
    if asil_raw is not None:
        tbl.cell(2, 0).text = asil_label
        tbl.cell(2, 1).text = asil_raw
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestSwUDSAsilExtraction32:
    """32차 W28: SwUDS docx 함수별 ASIL 추출 (heading 다음 표 'ASIL' 라벨)."""

    def test_extracts_asil_from_english_label(self):
        """영문 'ASIL' 라벨 옆 'ASIL-B' → 'B'."""
        bytes_ = _build_swuds_docx_with_asil("SwUFn_0101", "ASIL-B")
        result = parse_swuds_docx(bytes_)
        assert result.ok
        assert result.entries[0].asil == "B"
        assert result.function_asil_map == {"SwUFn_0101": "B"}

    def test_extracts_asil_from_korean_label(self):
        """한글 '안전등급' 라벨도 매칭."""
        bytes_ = _build_swuds_docx_with_asil(
            "SwUFn_0102", "ASIL-D", asil_label="안전등급",
        )
        result = parse_swuds_docx(bytes_)
        assert result.ok
        assert result.entries[0].asil == "D"

    def test_extracts_asil_letter_only_value(self):
        """라벨 옆 셀이 'B' 단일 letter — c_parser fallback 패턴 정규화."""
        bytes_ = _build_swuds_docx_with_asil("SwUFn_0103", "B")
        result = parse_swuds_docx(bytes_)
        assert result.ok
        assert result.entries[0].asil == "B"

    def test_no_asil_label_means_blank(self):
        """ASIL row 없는 docx → entry.asil = '' + function_asil_map 빈 dict."""
        bytes_ = _build_swuds_docx_with_asil("SwUFn_0104", None)
        result = parse_swuds_docx(bytes_)
        assert result.ok
        assert result.entries[0].asil == ""
        assert result.function_asil_map == {}

    def test_invalid_asil_value_returns_blank(self):
        """라벨은 있으나 값이 'High'/'Medium' 같은 비표준 — 빈 string."""
        bytes_ = _build_swuds_docx_with_asil("SwUFn_0105", "High")
        result = parse_swuds_docx(bytes_)
        assert result.ok
        assert result.entries[0].asil == ""


class TestRound87HeadingTableFallback:
    """라운드 87 T2102: 라운드 80 regex fallback이 다양한 양식에서 정상 동작 회귀."""

    def test_heading_table_format_extracts_asil_normally(self):
        """기본 heading+table 양식 — regex fallback 발화 안 함 (정상 path)."""
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_paragraph("SwUFn_0201 — heading")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.cell(0, 0).text = "ASIL"
        tbl.cell(0, 1).text = "C"
        buf = io.BytesIO()
        doc.save(buf)
        result = parse_swuds_docx(buf.getvalue())
        assert result.ok
        assert result.entries[0].asil == "C"
        # 정상 path → regex fallback warning 없음
        assert not any("regex fallback" in w for w in result.parse_warnings)

    def test_no_table_asil_triggers_regex_fallback(self):
        """heading+table 추출 0건 → regex fallback 발화 + 매핑 성공."""
        from docx import Document  # type: ignore
        doc = Document()
        # heading + table (ASIL 라벨 없음) → table 추출 0
        doc.add_paragraph("SwUFn_0301 — heading")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.cell(0, 0).text = "Description"
        tbl.cell(0, 1).text = "함수 설명"
        # 본문에 SwUFn_0301 ASIL B 패턴 (regex fallback target)
        doc.add_paragraph("SwUFn_0301 의 ASIL B 등급으로 분류됨.")
        buf = io.BytesIO()
        doc.save(buf)
        result = parse_swuds_docx(buf.getvalue())
        assert result.ok
        # regex fallback 발화 → ASIL B 매핑
        assert result.entries[0].asil == "B"
        assert any("regex fallback" in w for w in result.parse_warnings)

    def test_regex_fallback_silent_when_no_match(self):
        """heading+table 0 + regex 0 — silent skip (fallback warning 안 emit)."""
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_paragraph("SwUFn_0401 — heading")
        tbl = doc.add_table(rows=1, cols=2)
        tbl.cell(0, 0).text = "Other"
        tbl.cell(0, 1).text = "value"
        # 본문에 ASIL 패턴 없음
        buf = io.BytesIO()
        doc.save(buf)
        result = parse_swuds_docx(buf.getvalue())
        assert result.ok
        assert result.entries[0].asil == ""
        # 매핑 0건 → fallback warning 안 emit
        fallback_warnings = [w for w in result.parse_warnings if "regex fallback 적용" in w]
        assert fallback_warnings == []


class TestDescriptionExtractionKJPDS02:
    """Description 추출기 KJPDS02 병합라벨/rows≥5 대응 (ASIL/Name 추출기와 동일 패턴).

    구식 `rows[:5]` + naive `cells[i+1]`는 병합 라벨(반복 셀)·index≥5 Description을
    침묵 실패시켜 UDS Description이 빈 채로 표시됐다(heading만 노출). ASIL/Name 추출기가
    받은 병합라벨 업그레이드를 Description만 못 받은 결함의 회귀 방지.
    """

    @staticmethod
    def _docx(rows_spec, cols, fn_id="SwUFn_1150"):
        from docx import Document  # type: ignore
        doc = Document()
        doc.add_paragraph(f"{fn_id} — s_TunningParamRead_16bitData")
        tbl = doc.add_table(rows=len(rows_spec), cols=cols)
        for r, cells in enumerate(rows_spec):
            for c, txt in enumerate(cells):
                tbl.cell(r, c).text = txt
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_merged_label_repeated_cells(self):
        """['Description','Description','튜닝...'] — 반복 라벨 셀 skip 후 값 채택."""
        b = self._docx(
            [["Description", "Description", "튜닝 파라미터를 16bit로 읽어 반환한다"]], cols=3,
        )
        r = parse_swuds_docx(b)
        assert r.ok
        assert "튜닝 파라미터를 16bit" in r.entries[0].description

    def test_description_row_index_beyond_5(self):
        """Description 행이 index 5 — 구식 rows[:5]는 놓침, rows[:8]로 추출."""
        rows = [
            ["ID", "SwUFn_1150"],
            ["Name", "s_TunningParamRead_16bitData"],
            ["Prototype", "void s_TunningParamRead_16bitData(void)"],
            ["Reuse", "N"],
            ["Cyber", "-"],
            ["Description", "튜닝 파라미터를 16bit로 읽어 반환한다"],
            ["ASIL", "A"],
        ]
        b = self._docx(rows, cols=2)
        r = parse_swuds_docx(b)
        assert r.ok
        assert "튜닝 파라미터를 16bit" in r.entries[0].description

    def test_korean_merged_label_variant(self):
        """한글 '기능설명' 병합 라벨도 매칭."""
        b = self._docx(
            [["기능설명", "기능설명", "16비트 튜닝값을 읽는다"]], cols=3,
        )
        r = parse_swuds_docx(b)
        assert r.ok
        assert "16비트 튜닝값을 읽는다" in r.entries[0].description

    def test_empty_adjacent_cell_skipped(self):
        """['Description','', '실제 설명'] — 빈 인접 셀 skip(구식은 빈 값 반환)."""
        b = self._docx(
            [["Description", "", "실제 설명 문장"]], cols=3,
        )
        r = parse_swuds_docx(b)
        assert r.ok
        assert "실제 설명 문장" in r.entries[0].description

    def test_simple_hdpdm01_layout_no_regression(self):
        """HDPDM01 단순 ['Description','값'] — 무회귀(첫 후속 셀이 곧 값)."""
        b = self._docx(
            [["Description", "간단한 설명"]], cols=2,
        )
        r = parse_swuds_docx(b)
        assert r.ok
        assert "간단한 설명" in r.entries[0].description
