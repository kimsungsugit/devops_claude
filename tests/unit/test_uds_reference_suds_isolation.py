# tests/unit/test_uds_reference_suds_isolation.py
"""참조 SUDS 가 **다른 프로젝트의** 안전 등급을 흘려넣지 못하게 한다 (D2).

## 실측 (2026-07-31, 고치기 전)

`config.UDS_REF_SUDS_PATH` 의 기본값은 저장소 `docs/` 의 **HDPDM01 SUDS**(40.7MB)다.
`generate_uds_docx` 는 프로젝트가 무엇이든 이 문서를 읽어 **함수명만으로** 매칭해
대상 함수에 `asil`·`related`·`description`·`logic`·`inputs`·`outputs`·`called`·`calling`
을 덧씌웠다. 신원 확인은 **없었다**.

| 항목 | 실측 |
|---|---|
| 참조 문서 함수 블록 | **416개** (고유 함수명 408) |
| 그중 `asil` 보유 | **416개 전부** — A 280 / QM 135 / 파싱오류 1 |
| 그중 `related` 보유 | **416개 전부** |
| 파싱 오류 실물 | `asil = 'void s_Init_SystemManagementFunc( void )'` — 프로토타입 문자열이 ASIL 칸에 |

ASIL 하향은 ISO 26262 에서 가장 위험한 방향의 오류다. 참조의 67%가 `A` 이므로,
실제 ASIL C/D 함수가 이름만 겹치면 `A` 를 물려받는다.

같은 패턴을 이 저장소가 이미 두 번 고쳤다 —
`backend/routers/local.py::_pick_doc_path`("지정 문서를 못 읽으면 저장소 docs/ 로
바꿔치기")와 "SUTS ASIL 이 HDPDM01 로 채워짐". **여기가 남은 사이트다.**

## 정책

- **안전·추적성 축**(`asil`, `related`, `ref_related_by_name`) — 프로젝트 신원이
  **확인된 경우에만** 적용. 판정 불가(`None`)는 확인됨이 아니다(fail-closed).
- **서술 축**(description/precondition/logic/inputs/outputs/…) — 틀려도 안전 판정이
  아니라 문서 품질 문제이므로 유지. 대신 출처를 남긴다.
- 무엇을 막았는지는 `stats_out["reference_suds"]` 로 **보고**한다(로그만이면 안 보인다).
"""
from __future__ import annotations

from pathlib import Path

import pytest

from report_gen.docx_builder import (
    _VALID_ASIL,
    _project_tokens,
    _reference_identity_verdict,
)

_REF = Path("docs/(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx")


# --------------------------------------------------------------
# 토큰 추출
# --------------------------------------------------------------

class TestProjectTokens:
    def test_extracts_project_id_from_reference_stem(self):
        assert _project_tokens(_REF.stem) == {"HDPDM01"}

    @pytest.mark.parametrize("name,expected", [
        ("HDPDM01_PDS64_RD", {"HDPDM01", "PDS64"}),
        ("KJPDS02_PV", {"KJPDS02"}),
        ("KJPDS02_NE1AW_PORTING_ToolDev", {"KJPDS02", "NE1AW", "PORTING", "TOOLDEV"}),
    ])
    def test_extracts_from_real_job_names(self, name, expected):
        """이 저장소 캐시에 실재하는 Jenkins job 이름들."""
        assert _project_tokens(name) == expected

    def test_document_kind_words_are_not_identifiers(self):
        """`SUDS`·`SRS` 같은 문서 종류가 식별자로 잡히면 모든 프로젝트가 서로 '일치' 한다."""
        assert _project_tokens("SUDS SRS Software Unit Design Specification") == set()

    @pytest.mark.parametrize("bad", [None, "", 123, [], {}])
    def test_non_string_input_is_safe(self, bad):
        assert _project_tokens(bad) == set()

    def test_short_tokens_are_ignored(self):
        """2~4자 토큰(RD·PV·GW)은 우연 일치가 너무 쉬워 식별자로 쓰지 않는다."""
        assert _project_tokens("RD PV GW USS LC") == set()


# --------------------------------------------------------------
# 신원 판정
# --------------------------------------------------------------

class TestIdentityVerdict:
    def test_same_project_is_verified(self):
        """음성 대조군 — 본인 프로젝트는 막히면 안 된다(과잉 차단 방지)."""
        v = _reference_identity_verdict({"project_name": "HDPDM01_PDS64_RD"}, _REF)
        assert v["same_project"] is True
        assert v["shared_tokens"] == ["HDPDM01"]

    @pytest.mark.parametrize("payload", [
        {"project_name": "KJPDS02_PV"},
        {"module_name": "KJPDS02_DV"},
        {"summary": {"project": "SENSOR_GW_USS_LC"}},
        {"source_docs": ["D:/x/KJPDS02_SRS.txt"]},
    ])
    def test_other_projects_are_rejected(self, payload):
        """뮤테이션: `_reference_identity_verdict` 가 항상 True 를 내게 하면 실패."""
        v = _reference_identity_verdict(payload, _REF)
        assert v["same_project"] is False
        assert v["reason"] == "token_mismatch"

    def test_identity_reads_source_docs(self):
        """job 이름이 없어도 요구문서 파일명에 프로젝트 ID 가 있으면 확인된다."""
        v = _reference_identity_verdict({"source_docs": ["D:/x/HDPDM01_SRS.txt"]}, _REF)
        assert v["same_project"] is True

    @pytest.mark.parametrize("payload", [{}, None, {"project_name": ""}, {"project_name": "RD"}])
    def test_unverifiable_is_not_verified(self, payload):
        """**판정 불가는 '같은 프로젝트' 가 아니다.** 이게 fail-closed 의 핵심이다.

        뮤테이션: `None` 을 통과로 취급(`is not False`)하게 바꾸면 실패.
        """
        v = _reference_identity_verdict(payload, _REF)
        assert v["same_project"] is None
        assert v["same_project"] is not True

    def test_verdict_reports_both_sides(self):
        """왜 막혔는지 조치하려면 양쪽 토큰이 다 보여야 한다."""
        v = _reference_identity_verdict({"project_name": "KJPDS02_PV"}, _REF)
        assert v["ref_tokens"] == ["HDPDM01"]
        assert v["payload_tokens"] == ["KJPDS02"]

    def test_reference_without_token_is_unverifiable(self):
        v = _reference_identity_verdict({"project_name": "HDPDM01"}, Path("reference.docx"))
        assert v["same_project"] is None
        assert v["reason"] == "ref_no_token"


# --------------------------------------------------------------
# ASIL 어휘 검사
# --------------------------------------------------------------

class TestAsilVocabulary:
    def test_vocabulary_is_the_iso_grades(self):
        assert _VALID_ASIL == {"A", "B", "C", "D", "QM"}

    def test_parsed_prototype_string_is_not_a_grade(self):
        """실측된 파싱 오류 — 참조 416건 중 1건이 프로토타입 문자열을 ASIL 로 갖고 있었다."""
        assert "void s_Init_SystemManagementFunc( void )".upper() not in _VALID_ASIL


# --------------------------------------------------------------
# 적용 계층 — 안전축 차단 / 서술축 통과
# --------------------------------------------------------------

@pytest.fixture
def gen(tmp_path, monkeypatch):
    """**실제 `generate_uds_docx` 루프**를 태우는 하네스.

    ⚠ 처음엔 판정 규칙을 테스트에 복제해 검증했는데, 뮤테이션 검증에서 **실제 루프를
    고쳐도 테스트가 통과**했다(7건 중 2건 생존). 규칙 복제는 이 저장소가 반복해 겪은
    실패 모드 그대로다 — 규칙이 아니라 **코드**를 태워야 한다.

    참조 문서 파서(`report_gen/requirements.py::_extract_function_info_from_docx`)는
    40MB 실문서 형식을 요구하므로, 파서만 대체하고 루프는 진짜를 돌린다.
    """
    import docx

    import config
    from report_gen import docx_builder

    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_ref.docx"), raising=False)
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)

    def _run(*, project_name, ref_block, target_overrides=None, ref_stem="HDPDM01_SUDS"):
        # 참조 문서 실물(파서는 대체하지만 `docx.Document()` 는 실제로 열린다)
        ref_path = tmp_path / f"({ref_stem}) x.docx"
        docx.Document().save(str(ref_path))
        monkeypatch.setattr(config, "UDS_REF_SUDS_PATH", str(ref_path), raising=False)
        monkeypatch.setattr(docx_builder, "_extract_function_info_from_docx",
                            lambda _doc: {"REF_0001": dict(ref_block, name="alpha")})

        tpl = tmp_path / "t.docx"
        d = docx.Document()
        d.add_heading("Software Unit Design", level=1)
        d.add_heading("SwUFn_0001: alpha", level=4)
        d.save(str(tpl))

        info = {"id": "SwUFn_0001", "name": "alpha", "prototype": "void alpha(void);",
                "description": "", "asil": "TBD", "related": "TBD", "precondition": "N/A",
                "inputs": [], "outputs": [], "globals_global": [], "globals_static": [],
                "called": "", "logic": ""}
        info.update(target_overrides or {})
        payload = {"project_name": project_name, "overview": "o", "requirements": "r",
                   "interfaces": "i", "uds_frames": "u", "notes": "n",
                   "function_details": {"SwUFn_0001": info}}

        stats = {}
        docx_builder.generate_uds_docx(str(tpl), payload, str(tmp_path / "out.docx"),
                                       stats_out=stats)
        return info, stats.get("reference_suds", {})

    return _run


class TestFieldClassSplit:
    """실제 빌더 루프를 태운다 — 규칙 복제가 아니다."""

    def test_foreign_reference_cannot_set_related_id(self, gen):
        """핵심 — 다른 프로젝트 참조는 Related ID 를 채우지 못한다.

        ⚠ 2026-07-31 갱신 — 예전 이 docstring 은 *"실제로 유출되던 축은 `related`
        뿐이고 `asil` 은 병합 이전에 이미 `QM`(source=default)으로 채워져 참조가
        닿지 못한다(§D3)"* 라고 적었다. 그 선행 `QM` 지어내기를 **제거**했으므로
        (`tests/unit/test_asil_no_fabrication.py`) 이제 `asil` 축도 병합 자격을
        얻는다 → 신원 게이트가 **실제로** 두 축 다 막는다(1건 → 2건).
        차단 수가 는 것은 과대계상이 아니라 **게이트가 비로소 일을 하는 것**이다.

        뮤테이션: 루프의 `if not _ref_safety_ok: ... continue` 를 지우면 실패.
        """
        info, rs = gen(project_name="KJPDS02_PV",
                       ref_block={"asil": "A", "related": "SwFn_99"})
        assert "SwFn_99" not in str(info.get("related") or "")
        assert str(info.get("asil") or "").upper() != "A", "남의 프로젝트 등급이 들어왔다"
        assert rs["safety_fields_blocked"] == 2
        assert rs["safety_fields_applied"] == 0
        assert rs["identity"]["same_project"] is False

    def test_unverifiable_project_is_also_blocked(self, gen):
        """신원 판정 불가도 차단 — '확인 못 함' 은 '같은 프로젝트' 가 아니다."""
        info, rs = gen(project_name="", ref_block={"asil": "A", "related": "SwFn_99"})
        assert "SwFn_99" not in str(info.get("related") or "")
        assert rs["safety_fields_blocked"] == 2
        assert rs["identity"]["same_project"] is None

    def test_same_project_reference_still_fills_related(self, gen):
        """음성 대조군 — 본인 프로젝트 참조의 정당한 보강까지 죽이면 안 된다.

        ⚠ 2026-07-31 — `asil` 도 여기서 들어온다. 예전엔 병합 이전 `QM` 지어내기가
        칸을 선점해 **참조의 ASIL 이 영영 적용되지 않았다**(§D3 로 기록돼 있던 결함).
        지어내기를 제거하니 병합 규칙을 하나도 안 바꾸고 그 결함이 풀렸다 —
        그리고 자격 판정은 여전히 "값이 비었/TBD 인가" 라서 **실등급 하향 경로는
        열리지 않는다**(`test_existing_value_is_never_overwritten` 참조).
        """
        info, rs = gen(project_name="HDPDM01_PDS64_RD",
                       ref_block={"asil": "C", "related": "SwFn_12"})
        assert info["related"] == "SwFn_12"
        assert info["asil"] == "C", "자기 프로젝트 SUDS 의 ASIL 이 아직도 못 들어온다"
        assert rs["safety_fields_applied"] == 2
        assert rs["safety_fields_blocked"] == 0

    def test_invalid_asil_from_reference_is_rejected(self, gen):
        """참조 파싱이 어긋나 **프로토타입 문자열**이 ASIL 칸에 온 경우.

        실측된 실물: `asil = 'void s_Init_SystemManagementFunc( void )'`.
        그게 등급으로 굳으면 ISO 26262 판정 자체가 무의미해진다.

        ⚠ 신원이 **확인된** 프로젝트로 관측한다. 남의 프로젝트로 보면 신원 게이트가
        먼저 막아 어휘 검사가 있는지 없는지 알 수 없다 — 실제로 그래서 뮤테이션
        (`_VALID_ASIL` 검사 제거)이 **생존**했다. 기존 테스트는 상수 집합만 보는
        판정 복제였고, 이 파일 docstring 이 경고하던 바로 그 실패 모드다.
        """
        info, rs = gen(project_name="HDPDM01_PDS64_RD",
                       ref_block={"asil": "void s_Init_SystemManagementFunc( void )",
                                  "related": "SwFn_12"})
        assert info["asil"] != "void s_Init_SystemManagementFunc( void )"
        assert str(info.get("asil") or "").upper() in {"", "TBD"}
        assert rs["invalid_asil_rejected"] == 1
        assert rs["safety_fields_applied"] == 1, "related 만 적용돼야 한다"

    def test_blocked_count_does_not_overstate(self, gen):
        """막았다고 셀 수 있는 건 **실제로 적용됐을** 값뿐이다.

        판정 순서는 `적용 자격 → 값 유효성 → 신원` 이다. 신원을 맨 앞으로 되돌리면
        **어차피 적용되지 않았을 시도까지** 차단으로 세어 막은 양이 부풀려진다.

        ⚠ 2026-07-31 — 예전엔 "선행 `QM` 기본값 때문에 `asil` 이 적용 대상이 아니다"
        로 이걸 관측했는데, 그 지어내기를 제거해 이제 빈 `asil` 은 정당한 적용
        대상이다. 그래서 관측 방법을 바꿨다: 대상이 **이미 실제 등급을 가진** 경우로
        본다. 그 칸은 자격 단계에서 걸러지므로 차단으로 세면 안 된다(related 만 1건).
        """
        _info, rs = gen(project_name="KJPDS02_PV",
                        ref_block={"asil": "A", "related": "SwFn_99"},
                        target_overrides={"asil": "D"})
        assert rs["safety_fields_blocked"] == 1, "적용 불가한 시도까지 차단으로 셌다"

    def test_descriptive_fields_pass_even_for_foreign_reference(self, gen):
        """서술 축은 안전 판정이 아니다 — 차단 대상이 아니다.

        뮤테이션: description 분기에 `_ref_safety_ok` 조건을 붙이면(과잉 차단) 실패.
        (`description` 자체는 이 시점에 이미 합성돼 있어 `precondition` 으로 관측한다.)
        """
        info, rs = gen(project_name="KJPDS02_PV",
                       ref_block={"description": "센서 값을 읽는다",
                                  "precondition": "초기화 완료"})
        assert info["precondition"] == "초기화 완료"
        assert rs["descriptive_fields_applied"] >= 1

    @pytest.mark.parametrize("initial_desc", ["", "Auto-generated from source"])
    def test_reference_description_never_reaches_the_document(self, gen, initial_desc):
        """**측정 결과의 기록** — 참조의 `description` 은 어떤 초기값에서도 안 들어간다.

        `description` 은 참조 병합 **이전에** 합성되어(`description_source='inference'`)
        칸을 선점한다. 그래서 병합 루프의 description 분기는 실질적으로 죽어 있다.
        `asil`(선행 기본값 `QM`)도 같다 — **약한 출처가 강한 근거를 선점**하는 같은 패턴이
        세 축 중 둘에 있다(§D3 으로 기록, 이 라운드에서 고치지 않음).

        이 테스트는 결함을 정당화하지 않는다. 순서를 고쳐 참조가 닿게 되면 여기서
        실패하므로, 그때 §D3 판정을 함께 갱신하라는 신호다.
        """
        info, rs = gen(project_name="HDPDM01_PDS64_RD",
                       ref_block={"description": "센서 값을 읽는다"},
                       target_overrides={"description": initial_desc})
        assert info["description"] != "센서 값을 읽는다"
        assert rs["descriptive_fields_applied"] == 0

    def test_existing_value_is_never_overwritten(self, gen):
        """참조는 **빈 칸만** 채운다 — 이미 판정된 값을 덮으면 안 된다."""
        info, _rs = gen(project_name="HDPDM01_PDS64_RD",
                        ref_block={"asil": "A", "related": "SwFn_99"},
                        target_overrides={"asil": "D", "related": "SwFn_01"})
        assert info["asil"] == "D"
        assert info["related"] == "SwFn_01"

    def test_stats_land_in_the_sidecar(self, gen, tmp_path):
        """차단 사실은 파일로 남아야 한다 — 프로덕션은 서브프로세스라 in-process 는 못 넘는다."""
        import json

        from report_gen.docx_builder import gen_stats_path
        gen(project_name="KJPDS02_PV", ref_block={"asil": "A", "related": "SwFn_99"})
        side = json.loads(gen_stats_path(str(tmp_path / "out.docx")).read_text(encoding="utf-8"))
        assert side["reference_suds"]["safety_fields_blocked"] == 2
        assert side["reference_suds"]["identity"]["reason"] == "token_mismatch"


# --------------------------------------------------------------
# 통합 — 실제 빌더가 이 판정을 쓰는가
# --------------------------------------------------------------

class TestIntegration:
    def test_builder_calls_the_identity_verdict(self):
        """뮤테이션: `generate_uds_docx` 에서 `_reference_identity_verdict` 호출을 지우면 실패.

        판정 함수만 있고 호출되지 않으면 결함은 그대로다 — 이 저장소가 겪은 실패 모드
        (게이트는 있는데 발화하지 않음)를 여기서 막는다.
        """
        import ast

        from report_gen import docx_builder
        tree = ast.parse(Path(docx_builder.__file__).read_text(encoding="utf-8"))
        fn = next(n for n in ast.walk(tree)
                  if isinstance(n, ast.FunctionDef) and n.name == "generate_uds_docx")
        called = {n.func.id for n in ast.walk(fn)
                  if isinstance(n, ast.Call) and isinstance(n.func, ast.Name)}
        assert "_reference_identity_verdict" in called

    def test_stats_key_is_declared(self):
        """차단 사실이 sidecar 로 나가야 한다 — 로그만이면 산출물 검토자가 못 본다."""
        from report_gen import docx_builder
        src = Path(docx_builder.__file__).read_text(encoding="utf-8")
        assert '"reference_suds": _ref_stats' in src

    def test_uds_helper_does_not_substitute_repo_doc(self):
        """`backend/helpers/uds.py` 의 하드코딩 HDPDM01 폴백이 되살아나지 않아야 한다.

        운영자가 지정한 참조 경로가 없을 때 저장소 문서로 조용히 바꾸면, 산출물이
        **다른 프로젝트와의 diff** 가 된다.
        """
        src = Path("backend/helpers/uds.py").read_text(encoding="utf-8")
        marker = "(HDPDM01_SUDS) Software Unit Design Specification_v1.07_240213.docx"
        assign_lines = [
            ln for ln in src.splitlines()
            if marker in ln and "ref_docx" in ln and not ln.lstrip().startswith("#")
        ]
        assert not assign_lines, f"하드코딩 폴백이 되살아났다: {assign_lines}"


class TestStructuralFieldsAreCounted:
    """참조가 덧씌우는 축은 11개인데 계수는 5개만 있었다 (계획서 §6 후보 12).

    `descriptive_fields_applied` 는 description/precondition/logic 만 센다.
    `inputs`·`outputs`·`globals_static`·`globals_global`·`called`·`calling` 6축은
    **무기록으로** 들어갔다 — sidecar 의 `reference_suds` 는 "무엇이 적용됐나" 를 묻는
    기록인데 절반 이상이 안 보였다는 뜻이다.

    이건 정직성 문제이자 **후보 12(성능)의 선행 조건**이다: 신원 불일치 시 40MB 읽기를
    건너뛰어도 되는지는 "그때 적용량이 0인가" 로 갈리는데, 6축을 안 세면 알 수 없다.
    """

    _FULL_BLOCK = {
        "description": "ref desc", "asil": "A", "related": "SwFn_12",
        "precondition": "ref pre", "logic": "ref logic",
        "inputs": [{"name": "a"}], "outputs": [{"name": "b"}],
        "globals_static": ["g_s"], "globals_global": ["g_g"],
        "called": "callee_x", "calling": "caller_y",
    }

    def test_foreign_reference_still_writes_six_structural_axes(self, gen):
        """신원 불일치인데도 구조 축은 그대로 들어온다 — 그래서 읽기를 못 건너뛴다."""
        info, stats = gen(project_name="KJPDS02_PV", ref_block=dict(self._FULL_BLOCK))

        assert stats["identity"]["same_project"] is not True
        struct = stats["structural_fields_applied"]
        assert sum(struct.values()) > 0, (
            "구조 축 적용량이 0이면 신원 불일치 시 읽기를 건너뛸 수 있다는 뜻인데, "
            "실측은 그렇지 않다 — 이 단언이 깨지면 성능 최적화 전제를 다시 재라")
        # 실제로 대상 함수에 남의 프로젝트 값이 들어갔는지 값으로 확인한다
        assert info["called"] == "callee_x"
        assert info["calling"] == "caller_y"

    def test_every_structural_axis_is_counted(self, gen):
        """축 하나라도 계수에서 빠지면 그 축만 다시 보이지 않게 된다."""
        _info, stats = gen(project_name="KJPDS02_PV", ref_block=dict(self._FULL_BLOCK))

        struct = stats["structural_fields_applied"]
        assert set(struct) == {
            "inputs", "outputs", "globals_static", "globals_global", "called", "calling",
        }
        zero = [k for k, v in struct.items() if v == 0]
        assert not zero, f"덧씌웠는데 계수가 0인 축: {zero}"

    def test_counter_stays_zero_when_nothing_is_applied(self, gen):
        """계수가 '읽었다'가 아니라 '적용했다'를 세는지 — 대조군."""
        occupied = {
            "inputs": [{"name": "own"}], "outputs": [{"name": "own"}],
            "globals_static": ["own"], "globals_global": ["own"],
            "called": "own_callee", "calling": "own_caller",
        }
        info, stats = gen(project_name="KJPDS02_PV", ref_block=dict(self._FULL_BLOCK),
                          target_overrides=occupied)

        assert stats["structural_fields_applied"] == {
            "inputs": 0, "outputs": 0, "globals_static": 0,
            "globals_global": 0, "called": 0, "calling": 0,
        }
        assert info["called"] == "own_callee", "이미 값이 있으면 참조가 덮지 않는다"
