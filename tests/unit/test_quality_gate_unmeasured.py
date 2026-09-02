"""품질 게이트 — **분모 0 인 축을 0% 로 채점해 실패시키지 않는다**.

## 왜 이 파일이 필요했나

`_rate(v, base)` 가 `base <= 0` 에서 `0.0` 을 냈고, `failed` 가 그걸 임계 미달로 세었다.
실측(2026-09-02, 저장소에 남은 실 산출물 `.quality_gate.md` **127개** 전수):

    분모 0 지표가 있는 리포트 29개 → **29개 전부 `Gate pass: False`**
    실제 429함수 KJPDS02_PV: `Input fill 0/0` · `Output fill 0/0` 둘 다 0.0% 로 실패 계상
      → `Gates: 5 / 13`, 실패 8건 중 **2건이 근거 없는 것**

원인은 "채울 게 없었다" 가 아니라 **못 쟀다** 였다. 그 문서를 열어 보니 SwUFn 항목 상당수가
전역 변수(`SwUFn_0102 = g_sysos_wdictrl`)라 Prototype 칸이 비어 있고(429행 중 2행만 존재)
입력/출력 슬롯 자체를 셀 수 없다.

상시 실패는 회귀만 못 잡는 게 아니라 **진짜 미달을 가린다** — 이 저장소가 SwITCV 게이트에서
이미 같은 형태를 고쳤다(문서가 재지 않는 축으로 채점 → 영구 FAIL + 진짜 미달 25건 은폐).

⚠ 고치는 방향은 **한쪽으로만** 열려 있다: 미측정을 실패에서 빼되 **통과로도 접지 않는다**.
  `gate_report.py` 의 "판정 불가는 통과가 아님" 과 같은 규약이다.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Tuple

import pytest

FN_INFO_BANNER = "Function Information"


def _doc_with(tmp: Path, rows: List[Tuple[str, str]], n: int = 2) -> Path:
    """Function Information 표 `n`개짜리 최소 UDS DOCX.

    `rows` 는 `(라벨, 값)` — 게이트는 이 표를 읽어 채점한다. `Prototype` 을 비우면
    입력/출력 슬롯을 셀 수 없어 분모가 0 이 된다(실 산출물에서 벌어진 일).
    """
    import docx

    tmp.mkdir(parents=True, exist_ok=True)
    d = docx.Document()
    for i in range(1, n + 1):
        d.add_heading(f"SwUFn_{i:03d}: fn_{i}", level=2)
        t = d.add_table(rows=len(rows) + 2, cols=3)
        t.rows[0].cells[0].text = FN_INFO_BANNER
        t.rows[1].cells[0].text = "ID"
        t.rows[1].cells[2].text = f"SwUFn_{i:03d}"
        for r, (label, value) in enumerate(rows, start=2):
            t.rows[r].cells[0].text = label
            t.rows[r].cells[2].text = value.replace("{i}", str(i))
    out = tmp / "u.docx"
    d.save(str(out))
    return out


_NO_PROTOTYPE = [
    ("Name", "fn_{i}"),
    ("Prototype", ""),                 # ← 비어 있다 = 슬롯을 셀 수 없다
    ("Description", "fn_{i} 는 무언가를 한다"),
    ("ASIL", "B"),
    ("Related ID", "SwFn_001"),
    ("Input Parameters", "N/A"),
    ("Output Parameters", "N/A"),
]

_WITH_PROTOTYPE = [
    ("Name", "fn_{i}"),
    ("Prototype", "int fn_{i}(int a)"),
    ("Description", "fn_{i} 는 무언가를 한다"),
    ("ASIL", "B"),
    ("Related ID", "SwFn_001"),
    ("Input Parameters", "a"),
    ("Output Parameters", "int"),
]


def _gate(tmp: Path, rows: List[Tuple[str, str]]) -> Tuple[str, Dict[str, Any]]:
    """`(리포트 원문, 리더 출력)`."""
    from report_gen.evidence import read_gate_report
    from report_gen.validation import generate_uds_field_quality_gate_report

    docx_path = _doc_with(tmp, rows)
    out = tmp / "u.quality_gate.md"
    generate_uds_field_quality_gate_report(str(docx_path), str(out))
    return out.read_text(encoding="utf-8"), read_gate_report(out)


@pytest.fixture(scope="module")
def no_proto(tmp_path_factory):
    pytest.importorskip("docx")
    return _gate(tmp_path_factory.mktemp("noproto"), _NO_PROTOTYPE)


@pytest.fixture(scope="module")
def with_proto(tmp_path_factory):
    pytest.importorskip("docx")
    return _gate(tmp_path_factory.mktemp("withproto"), _WITH_PROTOTYPE)


# ==============================================================
# 1. 미측정은 0% 가 아니다
# ==============================================================

class TestZeroDenominatorIsNotZeroPercent:

    def test_the_metric_line_says_unmeasured_not_zero(self, no_proto):
        text, _ = no_proto
        line = next(ln for ln in text.splitlines() if ln.startswith("- Input fill:"))
        assert "미측정" in line, line
        assert "0.0%" not in line, line

    def test_the_parser_reports_no_percent_for_it(self, no_proto):
        """하류 계약 — 괄호 안에 `%` 가 없으면 `percent=None` 이고 `rates` 에서 빠진다.

        이 두 갈래가 유지돼야 "0% 로 잰 것" 과 "못 잰 것" 이 하류에서도 구분된다.
        """
        from report_gen.gate_report import parse_gate_report, to_rate_map

        text, _ = no_proto
        parsed = parse_gate_report(text)
        assert parsed["metrics"]["input_fill"]["percent"] is None
        assert parsed["metrics"]["input_fill"]["denominator"] == 0
        assert "input_fill" not in to_rate_map(parsed)

    def test_measured_zero_still_reads_as_zero_percent(self, no_proto):
        """대조군 — 분모가 있는 진짜 0% 는 그대로 0.0% 여야 한다.

        이게 없으면 "전부 미측정으로 돌려 실패를 없앴다" 와 구분되지 않는다.
        """
        from report_gen.gate_report import parse_gate_report

        text, _ = no_proto
        parsed = parse_gate_report(text)
        gg = parsed["metrics"]["globals_global_fill"]
        assert gg["denominator"] > 0
        assert gg["percent"] == 0.0


# ==============================================================
# 2. 미측정은 실패가 아니다 — 그리고 통과도 아니다
# ==============================================================

class TestUnmeasuredIsNeitherPassNorFail:

    def test_unmeasured_gates_are_not_listed_as_failures(self, no_proto):
        _, got = no_proto
        names = " ".join(f["gate"] for f in got["failed_gates"])
        assert "input_fill_rate" not in names
        assert "output_fill_rate" not in names

    def test_they_are_listed_separately_with_a_reason(self, no_proto):
        text, got = no_proto
        assert got["unmeasured_count"] == 2
        joined = " ".join(got["unmeasured_gates"])
        assert "input_fill_rate" in joined and "output_fill_rate" in joined
        assert "잰 적 없음" in joined, joined
        assert "## Unmeasured Gates" in text

    def test_the_denominator_counts_only_measured_gates(self, no_proto):
        """못 잰 축을 분모에 남기면 통과율이 근거 없이 낮아진다."""
        _, got = no_proto
        assert got["gates_total"] == 11        # 13 - 미측정 2
        assert got["gates_passed"] + len(got["failed_gates"]) == got["gates_total"]

    def test_unmeasured_alone_does_not_make_it_pass(self, tmp_path):
        """⚠ fail-open 방지 — 실패가 없어도 미측정이 있으면 통과가 아니다.

        미측정을 실패에서 빼는 수정이 곧바로 "통과" 로 새면 이 라운드는 결함을
        옮긴 것이 된다.
        """
        from report_gen.validation import generate_uds_field_quality_gate_report

        # 모든 임계를 0 으로 내려 **잰 축은 전부 통과**시킨다. 그래도 미측정 2건이 남는다.
        docx_path = _doc_with(tmp_path, _NO_PROTOTYPE)
        out = tmp_path / "z.quality_gate.md"
        generate_uds_field_quality_gate_report(
            str(docx_path), str(out),
            thresholds={k: 0.0 for k in (
                "description_fill_rate", "input_fill_rate", "output_fill_rate",
                "globals_global_fill_rate", "globals_static_fill_rate",
                "called_fill_rate", "calling_fill_rate", "asil_non_tbd_rate",
                "related_non_tbd_rate", "traceability_rate")})
        text = out.read_text(encoding="utf-8")
        assert "- Gate pass: `False`" in text, text[:400]
        assert "- Unmeasured gates: `2`" in text


    def test_it_does_not_say_everything_passed(self, tmp_path):
        """⚠ 두 번째 fail-open 출구 — 개선 권고문. **이 라운드가 만들어 낸** 출구다.

        미측정을 `failed` 에서 빼자 `Total functions: 0` 문서(실 산출물 4건)에서
        `failed` 가 **비게 되고**, "모든 품질 게이트를 통과했습니다" 가 나갈 수 있다.
        (구판에서는 그 문서도 실패 7건을 달고 있어 안 나갔다 — 새로 생길 뻔한 거짓이다.)
        판정 줄만 고치고 산문을 안 고치면 사람이 읽는 결론은 그대로 거짓이다.
        """
        import docx

        from report_gen.validation import generate_uds_field_quality_gate_report

        empty = tmp_path / "empty.docx"
        docx.Document().save(str(empty))          # SwUFn 표 0개 → total 0
        out = tmp_path / "empty.quality_gate.md"
        generate_uds_field_quality_gate_report(str(empty), str(out))
        text = out.read_text(encoding="utf-8")
        assert "모든 품질 게이트를 통과했습니다" not in text, text
        assert "재지 못했습니다" in text, text
        assert "- Gate pass: `False`" in text


# ==============================================================
# 3. 잴 수 있으면 잰다 (대조군)
# ==============================================================

class TestWhenItCanBeMeasuredItIs:

    def test_prototype_present_means_no_unmeasured_input_output(self, with_proto):
        text, got = with_proto
        assert got["unmeasured_count"] == 0
        assert "- Unmeasured Gates" not in text
        line = next(ln for ln in text.splitlines() if ln.startswith("- Input fill:"))
        assert "미측정" not in line, line
        assert "%" in line

    def test_the_two_fixtures_actually_differ(self, no_proto, with_proto):
        """픽스처가 실제로 축을 가르는지 고정 — 안 그러면 위 단언들이 무의미하다."""
        assert no_proto[1]["unmeasured_count"] != with_proto[1]["unmeasured_count"]
        assert no_proto[1]["gates_total"] < with_proto[1]["gates_total"]


# ==============================================================
# 4. 구판 산출물 — 없는 것은 0 이 아니다
# ==============================================================

class TestLegacyReportsStayUnmeasured:

    def test_old_report_without_the_section_is_none_not_zero(self, tmp_path):
        """`## Unmeasured Gates` 가 없는 구판을 0 으로 읽으면 "미측정 없음" 이라는
        거짓 단언이 된다 — 디스크에 그런 리포트가 실제로 있다."""
        from report_gen.evidence import read_gate_report

        p = tmp_path / "old.quality_gate.md"
        p.write_text(
            "# UDS Field Quality Gate Report\n\n"
            "- Total functions: `429`\n- Gate pass: `False`\n"
            "- Gates: `5` / `13` passed\n\n"
            "## Failed Gates\n- **input_fill_rate**: 0.0% < 20.0%\n",
            encoding="utf-8")
        got = read_gate_report(p)
        assert got["unmeasured_count"] is None
        assert got["unmeasured_gates"] == []
        assert got["gates_total"] == 13        # 구판 값을 그대로 보존한다
