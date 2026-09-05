"""Unit tests for report_gen.requirements — parsing, mapping, extraction."""
from __future__ import annotations


class TestExtractTableSection:
    def test_finds_section(self):
        from report_gen.requirements import _extract_table_section

        lines = [
            "Introduction",
            "some intro text",
            "",
            "Function List",
            "FuncA  SwUFn_01",
            "FuncB  SwUFn_02",
            "",
            "Summary",
        ]
        result = _extract_table_section(lines, "Function List", ["Summary"], 10)
        assert len(result) == 2
        assert "FuncA" in result[0]

    def test_header_not_found(self):
        from report_gen.requirements import _extract_table_section

        result = _extract_table_section(["line1", "line2"], "Missing", [], 10)
        assert result == []

    def test_max_rows_limit(self):
        from report_gen.requirements import _extract_table_section

        lines = ["Header"] + [f"row{i}" for i in range(20)]
        result = _extract_table_section(lines, "Header", [], 5)
        assert len(result) == 5


class TestNormalizeTableRow:
    def test_splits_by_whitespace(self):
        from report_gen.requirements import _normalize_table_row

        result = _normalize_table_row("col1  col2\tcol3")
        assert result == ["col1", "col2", "col3"]

    def test_empty(self):
        from report_gen.requirements import _normalize_table_row

        assert _normalize_table_row("") == []


class TestExtractFunctionBlocks:
    def test_parses_blocks(self):
        from report_gen.requirements import _extract_function_blocks

        text = (
            "SwCom_01\n"
            "SwUFn_0101: Motor_Init\n"
            "ID SwUFn_0101\n"
            "Name Motor_Init\n"
            "Description Init motor module\n"
            "ASIL B\n"
            "SwUFn_0102: Motor_Run\n"
            "ID SwUFn_0102\n"
            "Name Motor_Run\n"
        )
        blocks = _extract_function_blocks(text)
        # The parser creates separate blocks for header lines and ID lines
        assert len(blocks) >= 2
        names = [b.get("name") for b in blocks if b.get("name")]
        assert "Motor_Init" in names
        assert "Motor_Run" in names
        # Verify SwCom propagation
        assert all(b.get("swcom") == "SwCom_01" for b in blocks)

    def test_empty(self):
        from report_gen.requirements import _extract_function_blocks

        assert _extract_function_blocks("") == []


class TestSplitDocFunctionBlocks:
    def test_splits_by_swufn(self):
        from report_gen.requirements import _split_doc_function_blocks

        text = (
            "SwUFn_0101: Motor_Init\n"
            "Called Function: Helper_A\n"
            "SwUFn_0102: Motor_Run\n"
            "Called Function: Helper_B\n"
        )
        blocks = _split_doc_function_blocks(text)
        assert len(blocks) == 2
        assert blocks[0]["id"] == "SwUFn_0101"
        assert "Helper_A" in blocks[0]["lines"][0]

    def test_empty(self):
        from report_gen.requirements import _split_doc_function_blocks

        assert _split_doc_function_blocks("") == []


class TestCollectSectionLines:
    def test_collects_until_next_header(self):
        from report_gen.requirements import _collect_section_lines

        lines = [
            "Called Function FuncA",
            "FuncB",
            "Calling Function FuncC",
        ]
        result = _collect_section_lines(lines, "Called Function")
        assert "FuncA" in result
        assert "FuncB" in result
        assert len(result) == 2

    def test_empty_when_no_header(self):
        from report_gen.requirements import _collect_section_lines

        result = _collect_section_lines(["line1", "line2"], "Missing")
        assert result == []


class TestExtractStateTokens:
    def test_finds_st_tokens(self):
        from report_gen.requirements import _extract_state_tokens

        lines = ["transition to ST_IDLE", "if ST_RUNNING then ST_ERROR"]
        result = _extract_state_tokens(lines)
        assert "ST_IDLE" in result
        assert "ST_RUNNING" in result
        assert "ST_ERROR" in result

    def test_deduplicates(self):
        from report_gen.requirements import _extract_state_tokens

        result = _extract_state_tokens(["ST_IDLE", "ST_IDLE again"])
        assert result.count("ST_IDLE") == 1

    def test_empty(self):
        from report_gen.requirements import _extract_state_tokens

        assert _extract_state_tokens([]) == []


class TestExtractRequirementBlocks:
    def test_parses_blocks(self):
        from report_gen.requirements import _extract_requirement_blocks

        text = (
            "ID SwTR_001\n"
            "Name Motor Safety\n"
            "Description The motor shall be safe.\n"
            "Related ID SwTR_002\n"
            "\n"
            "ID SwTR_002\n"
            "Name Sensor Check\n"
            "Description Sensors shall be checked.\n"
        )
        blocks = _extract_requirement_blocks(text)
        assert len(blocks) == 2
        assert blocks[0]["id"] == "SwTR_001"
        assert blocks[0]["name"] == "Motor Safety"
        assert "safe" in blocks[0]["description"]
        # "Related ID SwTR_002" is parsed with split(None, 1)[-1]
        assert "SwTR_002" in blocks[0].get("related_ids", "")

    def test_empty(self):
        from report_gen.requirements import _extract_requirement_blocks

        assert _extract_requirement_blocks("") == []


class TestExtractRequirementsFallback:
    def test_finds_shall_lines(self):
        from report_gen.requirements import _extract_requirements_fallback

        text = "intro\nThe system shall init.\nrandom line\nIt must stop.\n"
        result = _extract_requirements_fallback(text)
        assert len(result) == 2
        assert "shall" in result[0].lower()

    def test_finds_korean_keywords(self):
        from report_gen.requirements import _extract_requirements_fallback

        result = _extract_requirements_fallback("이 기능은 요구사항을 충족해야 한다.")
        assert len(result) >= 1

    def test_empty(self):
        from report_gen.requirements import _extract_requirements_fallback

        assert _extract_requirements_fallback("") == []


class TestExtractDocSection:
    def test_extracts_section(self):
        from report_gen.requirements import _extract_doc_section

        text = (
            "1 Introduction\nSome intro.\n"
            "2 Requirements\nReq content here.\nMore content.\n"
            "3 Summary\nEnd.\n"
        )
        result = _extract_doc_section(text, "Requirements")
        assert "Req content" in result
        assert "End" not in result

    def test_missing_section(self):
        from report_gen.requirements import _extract_doc_section

        assert _extract_doc_section("1 Intro\nText\n", "Missing") == ""

    def test_empty(self):
        from report_gen.requirements import _extract_doc_section

        assert _extract_doc_section("", "Title") == ""


class TestBuildReqMapFromTexts:
    def test_builds_map_with_asil(self):
        from report_gen.requirements import _build_req_map_from_texts

        text = (
            "ID SwTR_001\n"
            "Name Motor Safety\n"
            "Description Some desc.\n"
            "ASIL B\n"
            "Related ID SwTR_002\n"
        )
        result = _build_req_map_from_texts([text])
        assert "swtr_001" in result
        assert result["swtr_001"]["asil"] == "B"


class TestExtractDocFunctionNames:
    def test_finds_prefixed_names(self):
        from report_gen.requirements import _extract_doc_function_names

        result = _extract_doc_function_names(["call g_Motor_Init and s_Sensor_Check"])
        assert "g_Motor_Init" in result
        assert "s_Sensor_Check" in result

    def test_empty(self):
        from report_gen.requirements import _extract_doc_function_names

        assert _extract_doc_function_names([]) == []
        assert _extract_doc_function_names([""]) == []


class TestNormalizeTraceMappingEntry:
    def test_basic_entry(self):
        from report_gen.requirements import _normalize_trace_mapping_entry

        result = _normalize_trace_mapping_entry({
            "requirement_id": "SwTR_001",
            "source_ids": ["file1.c", "file2.c"],
        })
        assert result is not None
        assert result["requirement_id"] == "SwTR_001"
        assert len(result["source_ids"]) == 2

    def test_string_sources(self):
        from report_gen.requirements import _normalize_trace_mapping_entry

        result = _normalize_trace_mapping_entry({
            "requirement_id": "SwTR_001",
            "source_ids": "file1.c, file2.c",
        })
        assert result is not None
        assert len(result["source_ids"]) == 2

    def test_missing_id_returns_none(self):
        from report_gen.requirements import _normalize_trace_mapping_entry

        assert _normalize_trace_mapping_entry({"source_ids": ["x"]}) is None

    def test_no_sources_returns_none(self):
        from report_gen.requirements import _normalize_trace_mapping_entry

        assert _normalize_trace_mapping_entry({"requirement_id": "X"}) is None

    def test_alt_keys(self):
        from report_gen.requirements import _normalize_trace_mapping_entry

        result = _normalize_trace_mapping_entry({
            "req_id": "SwTR_002",
            "source": "file.c",
        })
        assert result is not None
        assert result["requirement_id"] == "SwTR_002"


class TestParseTraceabilityJson:
    def test_list_format(self):
        import json

        from report_gen.requirements import _parse_traceability_json

        data = json.dumps([
            {"requirement_id": "SwTR_001", "source_ids": ["a.c"]},
            {"requirement_id": "SwTR_002", "source_ids": ["b.c"]},
        ])
        result = _parse_traceability_json(data)
        assert len(result) == 2

    def test_mappings_wrapper(self):
        import json

        from report_gen.requirements import _parse_traceability_json

        data = json.dumps({"mappings": [
            {"requirement_id": "SwTR_001", "source_ids": ["a.c"]},
        ]})
        result = _parse_traceability_json(data)
        assert len(result) == 1

    def test_dict_format(self):
        import json

        from report_gen.requirements import _parse_traceability_json

        data = json.dumps({"SwTR_001": ["a.c", "b.c"]})
        result = _parse_traceability_json(data)
        assert len(result) == 1
        assert result[0]["requirement_id"] == "SwTR_001"

    def test_invalid_json(self):
        from report_gen.requirements import _parse_traceability_json

        assert _parse_traceability_json("not json") == []


class TestParseTraceabilityCsv:
    def test_basic_csv(self):
        from report_gen.requirements import _parse_traceability_csv

        csv_text = "requirement_id,source_ids\nSwTR_001,a.c\nSwTR_002,b.c\n"
        result = _parse_traceability_csv(csv_text)
        assert len(result) == 2

    def test_empty(self):
        from report_gen.requirements import _parse_traceability_csv

        assert _parse_traceability_csv("") == []


class TestParseTraceabilityText:
    def test_json_detected(self):
        import json

        from report_gen.requirements import _parse_traceability_text

        data = json.dumps([{"requirement_id": "SwTR_001", "source_ids": ["a.c"]}])
        result = _parse_traceability_text(data)
        assert len(result) == 1

    def test_csv_fallback(self):
        from report_gen.requirements import _parse_traceability_text

        csv_text = "requirement_id,source_ids\nSwTR_001,a.c\n"
        result = _parse_traceability_text(csv_text)
        assert len(result) == 1

    def test_empty(self):
        from report_gen.requirements import _parse_traceability_text

        assert _parse_traceability_text("") == []


class TestExtractRequirementsFromComments:
    def test_finds_req_comment(self):
        from report_gen.requirements import _extract_requirements_from_comments

        text = "// REQ: Motor shall init\n/* Requirement: Sensor check */\n"
        result = _extract_requirements_from_comments(text)
        assert len(result) == 2
        assert "Motor shall init" in result[0]

    def test_empty(self):
        from report_gen.requirements import _extract_requirements_from_comments

        assert _extract_requirements_from_comments("") == []


class TestExtractRequirementsFromDoc:
    def test_structured_doc(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = (
            "ID SwTR_001\n"
            "Name Motor Safety\n"
            "Description The motor shall be safe.\n"
            "ASIL B\n"
            "Related ID SwTR_002\n"
            "\n"
            "ID SwTR_002\n"
            "Name Sensor Monitoring\n"
            "Description Sensors shall be monitored.\n"
        )
        result = _extract_requirements_from_doc(text)
        assert len(result) == 2
        assert "SwTR_001" in result[0]
        assert "ASIL" in result[0]

    def test_inline_id_format(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = "SwTR_003: Speed Limit\nDescription Speed must be limited.\n"
        result = _extract_requirements_from_doc(text)
        assert len(result) >= 1
        assert "SwTR_003" in result[0]

    def test_empty(self):
        from report_gen.requirements import _extract_requirements_from_doc

        assert _extract_requirements_from_doc("") == []

    def test_with_stop_keys(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = (
            "ID SwTR_001\n"
            "Description Motor shall init.\n"
            "Rationale Safety requirement\n"
            "Priority High\n"
        )
        result = _extract_requirements_from_doc(text)
        assert len(result) == 1
        assert "Rationale" not in result[0]


class TestGenerateUdsRequirementsPreview:
    def test_deduplication(self):
        from report_gen.requirements import generate_uds_requirements_preview

        text = (
            "ID SwTR_001\n"
            "Name Motor\n"
            "Description Desc1.\n"
        )
        result = generate_uds_requirements_preview([text, text])
        # Should deduplicate
        items = result.get("items", [])
        assert len(items) == 1

    def test_empty(self):
        from report_gen.requirements import generate_uds_requirements_preview

        result = generate_uds_requirements_preview([])
        assert result.get("items", []) == [] or result.get("count", 0) == 0


class TestReqIdPattern:
    def test_pattern_matches_various_prefixes(self):
        from report_gen.requirements import _REQ_ID_PAT

        assert _REQ_ID_PAT.search("SwTR_001")
        assert _REQ_ID_PAT.search("SwTSR_002")
        assert _REQ_ID_PAT.search("SwNTR_003")
        assert _REQ_ID_PAT.search("SwCNF_004")
        assert _REQ_ID_PAT.search("SwFn_005")
        assert not _REQ_ID_PAT.search("random_text")


class TestExtractFunctionBlocksDetailed:
    """More thorough tests for _extract_function_blocks to cover io/logic states."""

    def test_io_params(self):
        from report_gen.requirements import _extract_function_blocks

        text = (
            "SwUFn_0101: Foo\n"
            "[ Input Parameters ]\n"
            "uint8 x\n"
            "uint16 y\n"
            "[ Output Parameters ]\n"
            "uint8 result\n"
            "[ Logic Diagram ]\n"
            "some logic\n"
        )
        blocks = _extract_function_blocks(text)
        assert len(blocks) >= 1
        # Find the block with IO data
        block = next((b for b in blocks if b.get("inputs")), None)
        assert block is not None
        assert len(block["inputs"]) == 2
        assert block.get("outputs") == ["uint8 result"]
        assert block.get("logic") == "present"

    def test_various_fields(self):
        from report_gen.requirements import _extract_function_blocks

        text = (
            "SwUFn_0201: Bar\n"
            "Prototype void Bar(uint8 x)\n"
            "Description Checks bar.\n"
            "Called Function FuncA\n"
            "Calling Function FuncB\n"
            "선행조건 g_init == TRUE\n"
            "사용 전역변수 g_state\n"
        )
        blocks = _extract_function_blocks(text)
        b = blocks[0]
        assert b.get("prototype") == "void Bar(uint8 x)"
        assert b.get("description") == "Checks bar."
        assert b.get("called") == "Function FuncA"  # "Called" prefix stripped by split
        assert b.get("globals") is not None

    def test_related_id(self):
        from report_gen.requirements import _extract_function_blocks

        text = "SwUFn_0301: Baz\nRelated ID SwTR_001, SwTR_002\n"
        blocks = _extract_function_blocks(text)
        assert any("SwTR_001" in str(b.get("related", "")) for b in blocks)


class TestDocxToText:
    def test_extracts_text(self):
        from report_gen.requirements import _docx_to_text

        class FakePara:
            def __init__(self, t): self.text = t

        class FakeCell:
            def __init__(self, t): self.paragraphs = [FakePara(t)]

        class FakeRow:
            def __init__(self, cells): self.cells = cells

        class FakeTable:
            def __init__(self, rows): self.rows = rows

        class FakeDoc:
            def __init__(self):
                self.paragraphs = [FakePara("Hello"), FakePara("World")]
                self.tables = [FakeTable([FakeRow([FakeCell("Cell1")])])]

        result = _docx_to_text(FakeDoc())
        assert "Hello" in result
        assert "World" in result
        assert "Cell1" in result

    def test_empty_doc(self):
        from report_gen.requirements import _docx_to_text

        class FakeDoc:
            paragraphs = []
            tables = []

        assert _docx_to_text(FakeDoc()) == ""


class TestGenerateUdsRequirementsMapping:
    def test_filters_to_swtr(self):
        from report_gen.requirements import generate_uds_requirements_mapping

        items = [
            {"id": "SwTR_001", "name": "Safety", "related_ids": "SwCom_01, SwFn_01"},
            {"id": "SwCNF_002", "name": "Config"},  # not SwTR/SwTSR
            {"id": "SwTR_003", "name": "Sensor"},  # no related swcom/fn
        ]
        result = generate_uds_requirements_mapping(items)
        assert len(result) == 1
        assert result[0]["requirement_id"] == "SwTR_001"
        assert "SwCom_01" in result[0]["related_swcom"]

    def test_empty(self):
        from report_gen.requirements import generate_uds_requirements_mapping

        assert generate_uds_requirements_mapping([]) == []


class TestGenerateUdsRequirementsPreviewCounts:
    def test_counts_by_prefix(self):
        from report_gen.requirements import generate_uds_requirements_preview

        text = (
            "ID SwTR_001\nName A\nDescription Desc A.\n\n"
            "ID SwTR_002\nName B\nDescription Desc B.\n\n"
            "ID SwTSR_001\nName C\nDescription Desc C.\n"
        )
        result = generate_uds_requirements_preview([text])
        assert len(result["items"]) == 3
        assert result["counts"]["SwTR"] == 2
        assert result["counts"]["SwTSR"] == 1


class TestExtractRequirementsFromDocDetails:
    def test_asil_and_related(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = (
            "ID SwTR_001\n"
            "Name Motor Safety\n"
            "Description The motor shall be safe.\n"
            "ASIL B\n"
            "Related ID SwTR_002, SwTR_003\n"
        )
        result = _extract_requirements_from_doc(text)
        assert len(result) >= 1
        assert "ASIL:B" in result[0]
        assert "Related:" in result[0]

    def test_multiline_desc(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = (
            "ID SwTR_001\n"
            "Description First line of desc.\n"
            "Second line of desc.\n"
            "\n"  # blank line stops description
            "ASIL A\n"
        )
        result = _extract_requirements_from_doc(text)
        assert len(result) >= 1
        assert "First line" in result[0]
        assert "Second line" in result[0]

    def test_safety_level_variant(self):
        from report_gen.requirements import _extract_requirements_from_doc

        text = "ID SwTR_010\nSafety Level B\n"
        result = _extract_requirements_from_doc(text)
        assert len(result) >= 1
        assert "B" in result[0]


class TestNormalizeVcastRows:
    def test_groups_by_req_id(self):
        from report_gen.requirements import _normalize_vcast_rows

        rows = [
            {"requirement_id": "SwTR_001", "testcase": "TC_01", "result": "pass"},
            {"requirement_id": "SwTR_001", "testcase": "TC_02", "result": "fail"},
            {"requirement_id": "SwTR_002", "testcase": "TC_03"},
        ]
        result = _normalize_vcast_rows(rows)
        # _normalize_req_id uppercases keys
        assert len(result["SWTR_001"]) == 2
        assert len(result["SWTR_002"]) == 1

    def test_skips_invalid(self):
        from report_gen.requirements import _normalize_vcast_rows

        result = _normalize_vcast_rows([{}, "string", {"requirement_id": ""}])
        assert result == {}

    def test_empty(self):
        from report_gen.requirements import _normalize_vcast_rows

        assert _normalize_vcast_rows([]) == {}


class TestGenerateUdsTraceabilityMatrix:
    def test_basic_matrix(self):
        from report_gen.requirements import generate_uds_traceability_matrix

        items = [
            {"id": "SwTR_001"},
            {"id": "SwTR_002"},
        ]
        mappings = [
            {"requirement_id": "SwTR_001", "source_ids": ["file.c"]},
        ]
        vcast = [
            {"requirement_id": "SwTR_002", "testcase": "TC_01", "result": "pass"},
        ]
        result = generate_uds_traceability_matrix(items, mappings, vcast)
        assert result["total_requirements"] == 2
        assert result["summary"]["mapped_source_count"] == 1
        assert result["summary"]["mapped_test_count"] == 1
        assert result["has_source_mapping"] is True
        assert result["has_tests"] is True

    def test_empty(self):
        from report_gen.requirements import generate_uds_traceability_matrix

        result = generate_uds_traceability_matrix([])
        assert result["total_requirements"] == 0

    def test_row_asil_from_component_asil(self):
        """ASIL 결합(P5) — 요구사항 ASIL이 연결 컴포넌트의 최고등급으로 도출되는지.

        + _normalize_asil_value 엣지 전파: N/A는 거짓 'A' 격상 차단(''), B(C)는 max C.
        (reviewer WARN-A/테스트 갭 — requirements 레이어 ASIL 도출 단위테스트 0건 해소)
        """
        from report_gen.requirements import generate_uds_traceability_matrix

        items = [{"id": "R1"}, {"id": "R2"}, {"id": "R3"}, {"id": "R4"}]
        sds_pairs = [
            {"requirement_id": "R1", "component_ids": ["compA", "compB"]},  # max(A,B)=B
            {"requirement_id": "R2", "component_ids": ["compNA"]},          # N/A → 미상('')
            {"requirement_id": "R3", "component_ids": ["compBC"]},          # B(C) → max C
            {"requirement_id": "R4", "component_ids": ["compX"]},           # 맵에 없음 → ''
        ]
        component_asil = {"compa": "A", "compb": "B", "compna": "N/A", "compbc": "B(C)"}
        out = generate_uds_traceability_matrix(items, sds_pairs=sds_pairs, component_asil=component_asil)
        by_id = {r["requirement_id"]: r for r in out["rows"]}
        assert by_id["R1"]["asil"] == "B"   # 연결요소 최고
        assert by_id["R2"]["asil"] == ""    # N/A → 거짓 A 격상 차단
        assert by_id["R3"]["asil"] == "C"   # B(C) → 보조등급 max
        assert by_id["R4"]["asil"] == ""    # 매칭 실패 → 미상(graceful)

    def test_component_asil_optional(self):
        """component_asil 미전달(기존 호출) 시 행에 asil='' (회귀 안전)."""
        from report_gen.requirements import generate_uds_traceability_matrix

        out = generate_uds_traceability_matrix([{"id": "R1"}])
        assert out["rows"][0]["asil"] == ""

    def test_vcast_direct_rid_not_double_counted(self):
        """§F: VectorCAST 행이 direct requirement_id(매트릭스 SRS)로 이미 traced면,
        subprogram bridge 실패해도 untraced(unmapped_vcast)로 이중집계되지 않는다."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        vcast = [{
            "source": "VectorCAST",
            "requirement_id": "SwTR_001",   # 매트릭스 SRS = direct traced (2367)
            "subprogram": "unmatched_fn",   # SDS/SwUFn bridge 미매칭
            "testcase": "TC_01",
            "result": "pass",
        }]
        s = generate_uds_traceability_matrix(items, vcast_rows=vcast)["summary"]
        assert s["vcast_traced_rows"] == 1, s      # direct로 traced
        assert s["vcast_untraced_rows"] == 0, s    # 이중집계 해소
        assert s["unmapped_vcast_count"] == 0, s   # unmapped 리스트에서 제외

    def test_vcast_no_rid_unmatched_stays_untraced(self):
        """대조(§F 무영향): orig_rid 없고 subprogram도 미매칭이면 진짜 untraced로 유지된다."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        vcast = [{
            "source": "VectorCAST",
            "subprogram": "unmatched_fn",   # rid 없음 + bridge 미매칭 = 진짜 미추적
            "testcase": "TC_01",
            "result": "pass",
        }]
        s = generate_uds_traceability_matrix(items, vcast_rows=vcast)["summary"]
        assert s["unmapped_vcast_count"] == 1, s   # 진짜 미추적 유지
        assert s["vcast_untraced_rows"] == 1

    def test_vcast_subprogram_revoked_when_traced_in_another_row(self):
        """§G: 같은 subprogram이 한 행에서 bridge 실패(unmapped)해도 다른 행에서 traced되면
        미추적 목록(unmapped_vcast)에서 제거된다(never-revoke 이중집계 해소)."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        sds_pairs = [{"requirement_id": "SwTR_001", "component_ids": ["realfn"]}]
        rows = [
            {"source": "SUTS", "requirement_id": "SwUFn_0101", "unit": "realfn"},  # swufn_to_func 세팅
            # 같은 subprogram, testcase에 SwUFn 없음 → bridge 실패 → unmapped 등록
            {"source": "VectorCAST", "subprogram": "x_sub", "testcase": "TC_plain", "result": "pass"},
            # 같은 subprogram, testcase에 SwUFn_0101 → 2-hop bridge 성공 → traced
            {"source": "VectorCAST", "subprogram": "x_sub", "testcase": "SwUFn_0101", "result": "pass"},
        ]
        out = generate_uds_traceability_matrix(items, vcast_rows=rows, sds_pairs=sds_pairs)
        names = [u["subprogram"] for u in out["unmapped_vcast"]]
        assert "x_sub" not in names, names
        assert out["summary"]["unmapped_vcast_count"] == 0, out["summary"]

    def test_vcast_revoke_order_independent(self):
        """§G 순서 무관: traced 행이 unmapped 행보다 먼저 와도 일괄 필터로 revoke된다."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        sds_pairs = [{"requirement_id": "SwTR_001", "component_ids": ["realfn"]}]
        rows = [
            {"source": "SUTS", "requirement_id": "SwUFn_0101", "unit": "realfn"},
            # traced 행을 먼저, unmapped 행을 나중에 — first-wins 로직이면 놓치는 순서
            {"source": "VectorCAST", "subprogram": "x_sub", "testcase": "SwUFn_0101", "result": "pass"},
            {"source": "VectorCAST", "subprogram": "x_sub", "testcase": "TC_plain", "result": "pass"},
        ]
        out = generate_uds_traceability_matrix(items, vcast_rows=rows, sds_pairs=sds_pairs)
        names = [u["subprogram"] for u in out["unmapped_vcast"]]
        assert "x_sub" not in names, names

    def test_vcast_unresolved_swufn_not_counted_as_design_gap(self):
        """§H: SwUFn ID인데 함수명 해석 불가(resolved 빈)면 UNRESOLVED로 분류돼
        진짜 설계갭(unmapped_app_design_gap)을 부풀리지 않는다."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        vcast = [{"source": "VectorCAST", "subprogram": "SwUFn_9999",   # SUTS 대응 없음 → 해석 불가
                  "testcase": "TC_01", "result": "pass"}]
        out = generate_uds_traceability_matrix(items, vcast_rows=vcast)
        s = out["summary"]
        assert s["unmapped_layer_unresolved"] == 1, s          # UNRESOLVED 버킷
        assert s["unmapped_app_design_gap"] == 0, s      # 진짜 갭에서 제외(오염 방지)
        assert [u["layer"] for u in out["unmapped_vcast"]] == ["UNRESOLVED"]

    def test_vcast_named_leaf_stays_app_design_gap(self):
        """대조(§H 무영향): SwUFn ID 아닌 함수명 subprogram은 APP_LEAF 유지(정당한 설계갭 후보)."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        vcast = [{"source": "VectorCAST", "subprogram": "some_leaf_fn",  # 함수명(SwUFn 아님), 미매칭
                  "testcase": "TC_01", "result": "pass"}]
        out = generate_uds_traceability_matrix(items, vcast_rows=vcast)
        s = out["summary"]
        assert s["unmapped_layer_unresolved"] == 0, s          # UNRESOLVED 아님
        assert s["unmapped_app_design_gap"] == 1, s      # 정당한 APP_LEAF 갭 유지

    def test_test_artifact_gap_field_filters_uds(self):
        """FE-W1(deep-review): unmapped_test_artifact_gap = TEST_ARTIFACT ∩ ¬uds (배너 design축 분해용).
        app_design_gap과 동형 — 전체 layer_test_artifact와 달리 ¬uds 필터라 boundaryGap 이중차감 방지."""
        from report_gen.requirements import generate_uds_traceability_matrix
        items = [{"id": "SwTR_001"}]
        # 'Range'는 non-identifier → TEST_ARTIFACT, UDS 인벤토리 미매칭 → in_uds=False
        vcast = [{"source": "VectorCAST", "subprogram": "Range", "testcase": "TC_01", "result": "pass"}]
        out = generate_uds_traceability_matrix(items, vcast_rows=vcast)
        s = out["summary"]
        assert s["unmapped_layer_test_artifact"] == 1, s
        assert s["unmapped_test_artifact_gap"] == 1, s   # in_uds=False라 design gap에도 포함
        # design 분해 불변식: design_gap = app_gap + unresolved + test_artifact_gap + boundary(여기선 0)
        assert s["unmapped_design_gap"] == 1, s
        assert s["unmapped_app_design_gap"] == 0, s


class TestClassifyUnmappedLayerTokens:
    """ISO 계층 분류 정규식 보정(비앵커 인프라 토큰 추가) — 저수준 드라이버/부트/유틸을
    올바른 계층으로 분류하고, 정상 APP 함수는 오삼키지 않음을 검증(실데이터 검증분 고정)."""

    def test_bsw_tokens(self):
        from report_gen.requirements import _classify_unmapped_layer as clf
        # 타이머/CPU/포트핀/Processor Expert/드라이브IC — BSW로 복구돼야 함
        for nm in [
            "tim0_ch0_isr", "tim0_init", "rti_init", "systicktimer_interrupt",
            "cpu_interrupt", "cpu_machineexception", "u8s_cpupllstatuscheck",
            "pe_initialize_gpio", "pe_low_level_init", "pp1_buzzer_pwm_init",
            "pp2_motor_pwm2_in2_init", "pj0_motor_nhz1_putval", "pl0_wake_up_interrupt",
            "pad4_snsr_pwr_en_putval", "s_driveic_clearfaults", "u16s_driveic_check",
        ]:
            assert clf([nm]) == "BSW_DRIVER", nm

    def test_boot_tokens(self):
        from report_gen.requirements import _classify_unmapped_layer as clf
        # EEPROM(비앵커)/플래시 프로그래밍/리프로그 — BOOT로 복구돼야 함
        for nm in [
            "s_diageepromread", "s_diageepromwrite", "s_write2eeprom_partno",
            "s_uds_wdbi_us_cleareeprom", "s_uds_makeeepromaddr", "g_inlinemodeeepromclear",
            "ftmrz_init", "ftmrz_prog_phrase", "pflash_send_command",
            "eraseflashsector", "erasesectorinternal", "maininitreprogservices",
        ]:
            assert clf([nm]) == "BOOT_REPROG", nm

    def test_lib_tokens(self):
        from report_gen.requirements import _classify_unmapped_layer as clf
        assert clf(["stdutilcopydata"]) == "LIB_UTIL"
        assert clf(["stdutildataset"]) == "LIB_UTIL"

    def test_boot_wins_over_substring_collision(self):
        """eeprom(BOOT)이 'assist' 부분문자열보다 우선 — DTC 이름 속 assist는 앱 로직 아님."""
        from report_gen.requirements import _classify_unmapped_layer as clf
        assert clf(["s_diageepromwrite_assistfail"]) == "BOOT_REPROG"

    def test_app_negative_controls_no_misclassify(self):
        """정상 APP 함수는 신규 토큰에 오삼켜지지 않는다(bare flash/pin/lin 미추가).
        `Ap_*Eeprom*`류의 예외는 앱 앵커 도입으로 해소됐다 — test_app_eeprom_stays_app_not_boot 참조."""
        from report_gen.requirements import _classify_unmapped_layer as clf
        for nm in [
            "s_buzzerstateflashing_on",   # 'flashing' ≠ flash 토큰(미추가)
            "s_antipinchdetect_close",    # 'pin' 토큰 미추가
            "s_manualoperationdetect",
            "s_doorstatectrl", "s_motorspeedctrl", "s_reboundprotection",
            "s_slipcheck_original", "s_overopencheck", "s_playprotectchk",
            # W2(deep-review): _isr 앵커(_isr$|_isr_)로 앱 술어 _IsReady/_IsReset/_IsRunning/_IsRequest는
            # BSW 오매치 배제(bare _isr였다면 전부 BSW로 잘못 분류됐다).
            "s_motorctrl_isready", "s_door_isreset", "s_assist_isrunning",
            "g_window_isrequestpending",
        ]:
            assert clf([nm]) == "APP_LEAF", nm

    def test_app_eeprom_stays_app_not_boot(self):
        """앱 도메인 함수(`Ap_*`)는 이름 **중간**의 'eeprom'으로 BOOT에 삼켜지지 않는다.

        이전 계약은 이들을 BOOT로 재라벨하고 "실데이터 2건은 in_uds=True라 무영향"으로 넘어갔으나,
        그 docstring 자체가 latent 위험을 인정하고 있었다 — in_uds=False인 앱 eeprom 함수가
        생기면 설계 갭이 조용히 은닉된다. ISO 26262에서 허용할 수 없는 침묵이라 `Ap_` 앵커
        우선으로 근본 차단했다.

        ⚠ 인프라 복구(4e449dc)를 되돌리지 않는다 — 실측 955개 함수 중 이동은 아래 1건뿐이고
        BSW 216·LIB 44는 불변이다. 앵커 근거: APP 디렉터리 밖에 `Ap_*` **정의는 0건**.
        """
        from report_gen.requirements import _classify_unmapped_layer as clf
        assert clf(["s_ap_motorctrl_reseteepromparamstate"]) == "APP_LEAF"
        assert clf(["s_ap_previousctrl_reseteepromparams"]) == "APP_LEAF"
        # 앱 앵커가 없는 EEPROM 조작 함수는 그대로 BOOT — 복구분은 살아 있다.
        assert clf(["s_diageepromwrite_assistfail"]) == "BOOT_REPROG"
        assert clf(["s_write2eeprom_partno"]) == "BOOT_REPROG"
        assert clf(["g_syseepromctrl_main"]) == "BOOT_REPROG"
        # 후보에 인프라 이름이 섞이면 인프라 판정 기회를 남긴다(앱 앵커는 전원 일치일 때만).
        assert clf(["s_ap_x_reseteeprom", "s_write2eeprom_partno"]) == "BOOT_REPROG"

    def test_isr_anchored_real_isr_still_bsw(self):
        """W2(deep-review): _isr → _isr$|_isr_ 앵커 후에도 실제 ISR은 BSW 유지."""
        from report_gen.requirements import _classify_unmapped_layer as clf
        assert clf(["sci0_isr"]) == "BSW_DRIVER"            # _isr$ (또는 ^sci)
        assert clf(["tim0_ch0_isr"]) == "BSW_DRIVER"        # _isr$ (또는 ^tim[0-9])
        assert clf(["s_appmodule_isr_entry"]) == "BSW_DRIVER"  # _isr_ (중간 위치)

    def test_preexisting_behavior_unchanged(self):
        from report_gen.requirements import _classify_unmapped_layer as clf
        assert clf([]) == "APP_LEAF"
        assert clf(["range"]) == "TEST_ARTIFACT"
        assert clf(["s_sha256_transform"]) == "LIB_UTIL"
        assert clf(["g_drvin_main"]) == "BSW_DRIVER"


class TestGenerateUdsRequirementsFromDocs:
    def test_deduplicates(self):
        from report_gen.requirements import generate_uds_requirements_from_docs

        text = "ID SwTR_001\nDescription Motor safety.\n"
        result = generate_uds_requirements_from_docs([text, text])
        assert result.count("SwTR_001") == 1

    def test_fallback_to_keywords(self):
        from report_gen.requirements import generate_uds_requirements_from_docs

        text = "The system shall initialize safely.\n"
        result = generate_uds_requirements_from_docs([text])
        assert "shall" in result.lower()

    def test_empty(self):
        from report_gen.requirements import generate_uds_requirements_from_docs

        assert generate_uds_requirements_from_docs([]) == ""


class TestSdsComponentDescriptionAmbiguity:
    """component_description 폴백의 다중 SwCom 모호성 처리 (deep-review Warning #2).

    같은 함수가 2개+ SwCom 인터페이스에 서로 다른 설명으로 등장하면(SwCom echo/교차참조
    과다추출 실측 패턴) 표 순서로 임의 상속하는 오귀속을 막고 component_description을 붙이지
    않는다(honest miss). 단일 SwCom 함수는 정상 부착.
    """

    @staticmethod
    def _sc_table(doc, sc_id, sc_desc, fn_rows):
        """SC-Information 표 1개 추가 — 헤더/SC설명/인터페이스 함수행."""
        # rows: 헤더 + SC ID + SC Description + 인터페이스 헤더 + 함수행들
        tbl = doc.add_table(rows=4 + len(fn_rows), cols=3)
        tbl.cell(0, 0).text = "Software Component Information"
        tbl.cell(1, 0).text = "SC ID"
        tbl.cell(1, 1).text = sc_id
        tbl.cell(2, 0).text = "SC Description"
        tbl.cell(2, 1).text = sc_desc
        # 인터페이스 헤더행: No | Name | Description → in_interface + iface_header
        tbl.cell(3, 0).text = "No"
        tbl.cell(3, 1).text = "Name"
        tbl.cell(3, 2).text = "Description"
        for i, (fname, fdesc) in enumerate(fn_rows):
            r = 4 + i
            tbl.cell(r, 0).text = str(i + 1)  # 첫 셀 digit → 함수행 인식
            tbl.cell(r, 1).text = fname
            tbl.cell(r, 2).text = fdesc

    def _build(self, path):
        from docx import Document  # type: ignore
        doc = Document()
        # SwCom_01: s_shared(desc 비어 폴백 대상) + s_solo
        self._sc_table(doc, "SwCom_01", "첫 컴포넌트 설명",
                       [("s_shared", ""), ("s_solo", "")])
        # SwCom_02: s_shared 재등장(다른 SC 설명) → 모호
        self._sc_table(doc, "SwCom_02", "둘째 컴포넌트 설명",
                       [("s_shared", "")])
        doc.save(str(path))

    def test_multi_swcom_ambiguous_desc_not_attached(self, tmp_path):
        from report_gen.requirements import _extract_sds_partition_map
        p = tmp_path / "sds_multi_swcom.docx"
        self._build(p)
        pm = _extract_sds_partition_map(str(p))
        # s_shared: 두 SwCom에 다른 설명 → 모호 → component_description 미부착(오귀속 방지)
        assert "s_shared" in pm
        assert "component_description" not in pm["s_shared"]
        # s_solo: 단일 SwCom → 정상 부착
        assert pm.get("s_solo", {}).get("component_description") == "첫 컴포넌트 설명"


class TestSdsCompKey:
    """_sds_comp_key: SDS component_id → 함수명 bridge 키 정규화 (§C).

    SDS 추출이 component_id에 C 시그니처 조각을 붙여오는데, 반환형이 **공백으로
    분리돼** 앞에 붙은 형태('void f( void )')를 현재는 통째로 버려 함수→SRS 역추적을
    침묵 누락한다. fix는 선행 토큰이 **전부 알려진 C 반환형/한정자일 때만** 마지막
    토큰을 함수명으로 채택한다 — 설명문의 마지막 단어가 함수명으로 오매칭돼
    over-trace 되는 것(순진한 last-token 채택)을 원천 차단하는 게 핵심 대조.
    """

    def test_signature_return_prefix_recovered(self):
        """(P) 공백분리 반환형/한정자 접두 → 함수명 복구 (현재 코드에선 FAIL 예상)."""
        from report_gen.requirements import _sds_comp_key
        cases = [
            ("void f( void )",           "f"),          # 계획서 명시 케이스(void=자연어안전 타입)
            ("static void bar( void )",  "bar"),        # 한정자 + 타입(void)
            ("u8 s_calc(void)",          "s_calc"),     # 헝가리안 + 저장클래스
            ("const u16 getval( void )", "getval"),     # 한정자 + 헝가리안
            ("u8 * s_getbuf( void )",    "s_getbuf"),   # 포인터 반환형
        ]
        for comp, expected in cases:
            assert _sds_comp_key(comp) == expected, \
                f"{comp!r} → {_sds_comp_key(comp)!r} (기대 {expected!r})"

    def test_description_not_mistaken_for_function(self):
        """(N) over-trace 가드: 설명문(괄호 포함 포함)의 마지막 토큰을 함수명으로 삼지 않는다.

        reset counter / do something( foo ) 는 순진한 last-token fix라면 각각
        'counter'/'something' 이 되어 우연히 같은 이름 함수를 SRS로 잘못 추적한다.
        이 그룹이 ''를 유지하는지가 over-trace 무회귀의 게이트다.
        """
        from report_gen.requirements import _sds_comp_key
        cases = [
            "reset counter",             # last-token이면 'counter' (오매칭)
            "do something( foo )",       # last-token이면 'something' (오매칭)
            "power operation disable",   # 다토큰 설명문
            "power operation disable(swst_09) state 천이",  # 실데이터: 괄호+비화이트 첫토큰
            "mcu 이상 감지(레지스터 미지원)",  # 한글 설명문 + 괄호
            "swcom_35: bootloader\t115", # 표 아티팩트 + 콜론
            "auto close",                # ★실데이터 회귀: 'auto'=저장한정자 겹침 → 괄호 없어 미적용
            "auto closestandby",         # ★실데이터 회귀: 동상
            "auto close( void )",        # ★가정: 괄호 있어도 auto 화이트리스트 제외로 차단(이중방어)
            # ── deep-review W1: 화이트리스트 자연어 충돌 (기본타입/typedef 제거로 첫토큰 미등재) ──
            "long delay( ms )",          # 'long'=자연어 → 미등재 → 미채택 (재현된 over-trace)
            "short circuit( detect )",   # 'short'
            "double click( fast )",      # 'double'
            "int overflow( check )",     # 'int'
            "char array( init )",        # 'char'
            "byte order( swap )",        # 'byte' (벤더 typedef, 충돌 최고)
            "word align( addr )",        # 'word'
            "unsigned int foo(int a)",   # 기본타입 제거 trade-off: 이전 복구 → 이제 미복구
            # ── 타입-최소-1개 게이트: 한정자 단독 선행 차단 ──
            "static delay( ms )",        # 'static'=한정자뿐, 타입 토큰 없음 → 미채택
            "const value( x )",          # 'const'=한정자뿐
            # ── last-token 키워드 차단: 말단이 타입/한정자면 함수명 아님 ──
            "void void( x )",            # 말단 'void'가 타입 키워드 → 배제
            "const void( x )",           # 말단 'void'가 타입 키워드 → 배제
            # ── deep-review 후속: struct/union/enum 제외(집합체 반환형은 태그 동반→차단, 단독은 자연어) ──
            "union select( x )",         # 'union'=영어단어, 화이트 제외 → 미채택
            "struct data( init )",       # 'struct' 제외 → 미채택
            "struct led_state get( void )",  # 진짜 집합체 반환형도 태그(led_state)로 미복구(under-trace 안전측)
        ]
        for comp in cases:
            assert _sds_comp_key(comp) == "", \
                f"{comp!r} → {_sds_comp_key(comp)!r} (기대 '' — over-trace 가드)"

    def test_existing_normalization_unchanged(self):
        """(R) 공백 없는 기존 경로 무회귀 (§A/라운드109·112 케이스 — fix 미발동 경로)."""
        from report_gen.requirements import _sds_comp_key
        cases = [
            ("s_systemhashcalculate( void", "s_systemhashcalculate"),  # 함수명 + 시그니처 꼬리
            ("u8g_x_partnoinfo[10]",        "u8g_x_partnoinfo"),       # 배열 첨자 제거
            ("_entrypoint",                 "entrypoint"),             # 선행 언더스코어
            ("foo",                         "foo"),                    # 순수 식별자
            ("",                            ""),                       # 빈 입력
            (None,                          ""),                       # None
        ]
        for comp, expected in cases:
            assert _sds_comp_key(comp) == expected, \
                f"{comp!r} → {_sds_comp_key(comp)!r} (기대 {expected!r})"
