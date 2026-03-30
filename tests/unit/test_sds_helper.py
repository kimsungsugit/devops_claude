from __future__ import annotations

from pathlib import Path


def _build_sample_sds(path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("SwCom_13 DoorCtrl", level=1)
    doc.add_paragraph("ap_doorctrl_pds handles door state transitions and module coordination.")
    doc.add_heading("Interface Detail", level=2)
    doc.add_paragraph("s_doorctrl_open is used when the door open request is accepted.")
    doc.save(str(path))


def _build_sample_sds_with_contents(path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("1. Introduction 1")
    doc.add_paragraph("2. SwCom_13 DoorCtrl 2")
    doc.add_paragraph("3. ap_doorctrl_pds 3")
    doc.add_paragraph("4. Interface Detail 4")
    doc.add_paragraph("5. s_doorctrl_open 5")
    doc.add_heading("SwCom_13 DoorCtrl", level=1)
    doc.add_paragraph("ap_doorctrl_pds handles door state transitions and module coordination.")
    doc.add_heading("Interface Detail", level=2)
    doc.add_paragraph("s_doorctrl_open is used when the door open request is accepted.")
    doc.save(str(path))


def _build_sample_sds_with_tables(path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Software Component Information", level=2)
    doc.add_heading("SwCom_01: System OS", level=3)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "SC ID"
    table.cell(0, 1).text = "SwCom_01"
    table.cell(1, 0).text = "SC Name"
    table.cell(1, 1).text = "System OS"
    table.cell(2, 0).text = "Description"
    table.cell(2, 1).text = "Core OS component detail"
    doc.add_heading("SwCom_02: Drv In", level=3)
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "SC ID"
    table2.cell(0, 1).text = "SwCom_02"
    table2.cell(1, 0).text = "Description"
    table2.cell(1, 1).text = "Driver input detail"
    doc.save(str(path))


def _build_sample_sds_with_generic_module_sections(path: Path) -> None:
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Software Component Information", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "SC ID"
    table.cell(0, 1).text = "SwCom_01"
    table.cell(1, 0).text = "SC ID"
    table.cell(1, 1).text = "SwCom_02"
    table.cell(2, 0).text = "SC ID"
    table.cell(2, 1).text = "SwCom_03"
    doc.add_heading("SwCom_01: System OS", level=3)
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "SC ID"
    table2.cell(0, 1).text = "SwCom_01"
    table2.cell(1, 0).text = "Description"
    table2.cell(1, 1).text = "Core OS component detail"
    doc.add_heading("Component Folder Struct", level=2)
    table3 = doc.add_table(rows=3, cols=2)
    table3.cell(0, 0).text = "Folder"
    table3.cell(0, 1).text = "SwCom_01"
    table3.cell(1, 0).text = "Folder"
    table3.cell(1, 1).text = "SwCom_02"
    table3.cell(2, 0).text = "Folder"
    table3.cell(2, 1).text = "SwCom_03"
    doc.save(str(path))


def test_build_sds_view_model_extracts_functions_and_modules(tmp_path):
    from backend.helpers.sds import build_sds_view_model

    path = tmp_path / "sample_sds.docx"
    _build_sample_sds(path)

    view = build_sds_view_model(
        str(path),
        changed_functions={"ap_doorctrl_pds": "HEADER"},
        changed_files=["Sources/APP/Ap_DoorCtrl_PDS.c"],
        flagged_modules=["Doorctrl"],
    )

    assert view["path"].endswith("sample_sds.docx")
    assert view["counts"]["functions"] >= 1
    assert view["counts"]["modules"] >= 1
    titles = {item["title"] for item in view["items"]}
    assert "ap_doorctrl_pds" in titles
    item = next(item for item in view["items"] if item["title"] == "ap_doorctrl_pds")
    assert "SwCom_13" in " ".join(item.get("relatedModules") or [])
    assert item["changed"] is True
    assert "HEADER" in item["changeTypes"]
    assert item["sourceFiles"]
    assert item["matchConfidence"] > 0


def test_local_sds_view_route_returns_view_model(tmp_path):
    from backend.routers.local import local_sds_view
    from backend.schemas import SdsViewRequest

    path = tmp_path / "sample_sds.docx"
    _build_sample_sds(path)

    result = local_sds_view(
        SdsViewRequest(
            path=str(path),
            max_items=50,
            changed_functions={"ap_doorctrl_pds": "BODY"},
            changed_files=["Sources/APP/Ap_DoorCtrl_PDS.c"],
            flagged_modules=["Doorctrl"],
        )
    )

    assert result["ok"] is True
    assert result["view"]["counts"]["functions"] >= 1
    changed_items = [item for item in result["view"]["items"] if item.get("changed")]
    assert changed_items


def test_build_sds_view_model_ignores_contents_for_mapping(tmp_path):
    from backend.helpers.sds import build_sds_view_model

    path = tmp_path / "sample_sds_contents.docx"
    _build_sample_sds_with_contents(path)

    view = build_sds_view_model(
        str(path),
        changed_functions={"ap_doorctrl_pds": "HEADER"},
        changed_files=["Sources/APP/Ap_DoorCtrl_PDS.c"],
        flagged_modules=["DoorCtrl"],
    )

    item = next(item for item in view["items"] if item["title"] == "ap_doorctrl_pds")
    assert item["sections"]
    assert item["sections"][0]["heading"] != "Contents"


def test_build_sds_view_model_includes_table_content_under_heading(tmp_path):
    from backend.helpers.sds import build_sds_view_model

    path = tmp_path / "sample_sds_tables.docx"
    _build_sample_sds_with_tables(path)

    view = build_sds_view_model(str(path))

    item = next(item for item in view["items"] if item["title"] == "SwCom_01")
    assert item["sections"]
    text = item["sections"][0]["text"]
    assert "System OS" in text
    assert "Core OS component detail" in text


def test_build_sds_view_model_excludes_generic_multi_module_sections_from_module_cards(tmp_path):
    from backend.helpers.sds import build_sds_view_model

    path = tmp_path / "sample_sds_generic_sections.docx"
    _build_sample_sds_with_generic_module_sections(path)

    view = build_sds_view_model(str(path))

    item = next(item for item in view["items"] if item["title"] == "SwCom_01")
    headings = [section["heading"] for section in item["sections"]]
    assert headings[0] == "SwCom_01: System OS"
    assert "Software Component Information" not in headings
    assert "Component Folder Struct" not in headings
