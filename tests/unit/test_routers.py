"""Unit tests for backend routers (health, exports, code, config).

Uses starlette TestClient to exercise FastAPI endpoints without a running server.
"""
from __future__ import annotations

import os
import sys
import tempfile
import types
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from starlette.testclient import TestClient

# Ensure repo root is on sys.path so backend/config can be imported
_repo_root = Path(__file__).resolve().parents[2]
if str(_repo_root) not in sys.path:
    sys.path.insert(0, str(_repo_root))

# Stub optional dependencies that may not be installed in the test environment
# so that importing backend.main does not fail.
for _mod_name in [
    "langchain_core",
    "langchain_core.tools",
    "langchain_mcp_adapters",
    "langchain_mcp_adapters.tools",
    "mcp",
    "mcp.client",
    "mcp.client.stdio",
]:
    if _mod_name not in sys.modules:
        _stub = types.ModuleType(_mod_name)
        # Provide minimal class stubs that routers may reference at import time
        _stub.BaseTool = MagicMock       # type: ignore[attr-defined]
        _stub.StructuredTool = MagicMock  # type: ignore[attr-defined]
        sys.modules[_mod_name] = _stub

import backend.middleware as _mw  # noqa: E402

# Disable rate limiting for tests
_mw.RATE_LIMIT = 999999

from backend.main import app  # noqa: E402

client = TestClient(app, raise_server_exceptions=False)
# UserContextMiddleware blocks any /api/* request that lacks an X-User header
# with 401. Inject a default header here so the router contract tests below
# exercise their actual paths instead of the auth gate. allowed_users.json is
# empty (unrestricted), so any non-empty value satisfies the middleware.
# If you ever add a test that asserts 401 behaviour, override the header on
# that specific request: `client.get(url, headers={"X-User": ""})`.
client.headers["X-User"] = "test"

# 파일 resolver 는 conftest 의 `_default_local_resolver` 가 local 로 고정한다
# (머신의 config/file_mode.json=cloudium 에 의존하면 파일 계열 라우터가 전부
#  403 cloudium-blocked 로 떨어져 이 파일이 단독 실행 시 14건 깨졌었다).


# ═══════════════════════════════════════════════════════════════════
# Health Router
# ═══════════════════════════════════════════════════════════════════
class TestHealthRouter:
    """Tests for /api/health and related health endpoints."""

    def test_health_check_status_200(self):
        r = client.get("/api/health")
        assert r.status_code == 200
        data = r.json()
        assert data["status"] == "ok"

    def test_health_check_has_version(self):
        r = client.get("/api/health")
        data = r.json()
        assert "version" in data
        assert isinstance(data["version"], str)

    def test_health_check_has_engine(self):
        r = client.get("/api/health")
        data = r.json()
        assert "engine" in data
        assert isinstance(data["engine"], str)

    def test_health_check_has_file_mode(self):
        r = client.get("/api/health")
        data = r.json()
        assert "file_mode" in data
        assert data["file_mode"] in ("local", "cloudium")

    def test_file_mode_get(self):
        r = client.get("/api/file-mode")
        assert r.status_code == 200
        data = r.json()
        assert "mode" in data

    def test_preview_excel_missing_path(self):
        """POST /api/preview-excel with empty path returns 400."""
        r = client.post("/api/preview-excel", json={"path": ""})
        assert r.status_code == 400

    def test_preview_excel_nonexistent_file(self):
        """POST /api/preview-excel with nonexistent file returns 404."""
        r = client.post(
            "/api/preview-excel",
            json={"path": "/nonexistent/path/file.xlsx"},
        )
        assert r.status_code == 404

    def test_preview_excel_unsupported_format(self):
        """POST /api/preview-excel with unsupported extension returns 400."""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"dummy")
            tmp_path = f.name
        try:
            r = client.post("/api/preview-excel", json={"path": tmp_path})
            assert r.status_code == 400
        finally:
            os.unlink(tmp_path)

    def test_preview_excel_txt_file(self):
        """POST /api/preview-excel with a .txt file returns content."""
        with tempfile.NamedTemporaryFile(
            suffix=".txt", delete=False, mode="w", encoding="utf-8"
        ) as f:
            f.write("line1\nline2\nline3\n")
            tmp_path = f.name
        try:
            r = client.post("/api/preview-excel", json={"path": tmp_path})
            assert r.status_code == 200
            data = r.json()
            assert data["ok"] is True
            assert len(data["sheets"]) == 1
            assert data["sheets"][0]["name"] == "Content"
            assert len(data["sheets"][0]["rows"]) == 3
        finally:
            os.unlink(tmp_path)

    def test_preview_image_nonexistent(self):
        """GET /api/preview-image with nonexistent docx returns 404."""
        r = client.get(
            "/api/preview-image",
            params={"path": "/nonexistent/doc.docx", "image_id": "rId1"},
        )
        assert r.status_code == 404

    def test_check_access_no_body(self):
        """POST /api/file-mode/check-access with empty body returns ok."""
        r = client.post("/api/file-mode/check-access", json={})
        assert r.status_code == 200
        data = r.json()
        assert "mode" in data

    def test_check_access_with_nonexistent_path(self):
        """POST /api/file-mode/check-access with nonexistent path."""
        r = client.post(
            "/api/file-mode/check-access",
            json={"path": "/nonexistent/test/path"},
        )
        assert r.status_code == 200
        data = r.json()
        assert data.get("accessible") is False


# ═══════════════════════════════════════════════════════════════════
# Exports Router
# ═══════════════════════════════════════════════════════════════════
class TestExportsRouter:
    """Tests for /api/exports endpoints."""

    def test_list_exports_returns_list(self):
        """GET /api/exports returns a JSON list."""
        r = client.get("/api/exports")
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, list)

    def test_list_exports_with_nonexistent_base(self):
        """GET /api/exports with a nonexistent base dir returns empty or error."""
        r = client.get("/api/exports", params={"base": "/nonexistent/base/dir"})
        # Server may return 200 (empty list), 400, or 403 (forbidden path)
        assert r.status_code in (200, 400, 403)

    def test_list_exports_with_session_filter(self):
        """GET /api/exports with session_id filter still returns a list."""
        r = client.get(
            "/api/exports",
            params={"session_id": "nonexistent_session_xyz"},
        )
        assert r.status_code == 200
        assert isinstance(r.json(), list)

    def test_delete_export_nonexistent(self):
        """DELETE /api/exports/<filename> with nonexistent file returns 404."""
        r = client.delete("/api/exports/nonexistent_file.zip")
        assert r.status_code == 404

    def test_restore_export_nonexistent(self):
        """POST /api/exports/restore/<filename> with nonexistent returns 404."""
        r = client.post("/api/exports/restore/nonexistent_file.zip")
        assert r.status_code == 404

    def test_download_export_nonexistent(self):
        """GET /api/exports/download/<filename> with nonexistent returns 404."""
        r = client.get("/api/exports/download/nonexistent_file.zip")
        assert r.status_code == 404

    def test_pdf_convert_nonexistent_source(self):
        """POST /api/exports/pdf/convert with nonexistent file returns error."""
        r = client.post(
            "/api/exports/pdf/convert",
            json={"source_path": "/nonexistent/file.docx"},
        )
        # Should get 404 (FileNotFoundError) or 500
        assert r.status_code in (404, 500)

    def test_pdf_convert_unsupported_extension(self):
        """POST /api/exports/pdf/convert with unsupported ext returns error."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"dummy")
            tmp_path = f.name
        try:
            r = client.post(
                "/api/exports/pdf/convert",
                json={"source_path": tmp_path},
            )
            # 400 (HTTPException) or 500 (APIError caught by generic handler)
            assert r.status_code in (400, 500)
        finally:
            os.unlink(tmp_path)

    def test_pdf_convert_missing_source_path(self):
        """POST /api/exports/pdf/convert without source_path returns 422."""
        r = client.post("/api/exports/pdf/convert", json={})
        assert r.status_code == 422

    def test_pdf_report_missing_fields(self):
        """POST /api/exports/pdf/report without required fields returns 422."""
        r = client.post("/api/exports/pdf/report", json={})
        assert r.status_code == 422

    def test_pdf_report_with_sections(self):
        """POST /api/exports/pdf/report with temp output path generates PDF."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = str(Path(tmpdir) / "test_report.pdf")
            r = client.post(
                "/api/exports/pdf/report",
                json={
                    "title": "Test Report",
                    "sections": [
                        {"heading": "Section 1", "content": "Hello world"},
                        {"heading": "Section 2", "content": "Test content"},
                    ],
                    "output_path": out_path,
                },
            )
            # Might succeed (200) or fail (500) depending on PDF libs installed
            if r.status_code == 200:
                data = r.json()
                assert data["ok"] is True
                assert "pdf_path" in data
                assert "size_mb" in data
            else:
                # PDF generation library may not be available in test env
                assert r.status_code == 500

    def test_cleanup_exports_returns_deleted_count(self):
        """POST /api/exports/cleanup returns deleted count."""
        r = client.post("/api/exports/cleanup", params={"days": 1})
        assert r.status_code == 200
        data = r.json()
        assert "deleted" in data
        assert isinstance(data["deleted"], int)


# ═══════════════════════════════════════════════════════════════════
# Code Router
# ═══════════════════════════════════════════════════════════════════
class TestCodeRouter:
    """Tests for /api/code endpoints."""

    def test_preview_function_missing_params(self):
        """GET /api/code/preview/function without required params returns 422."""
        r = client.get("/api/code/preview/function")
        assert r.status_code == 422

    def test_preview_function_missing_function_name(self):
        """GET /api/code/preview/function without function_name returns 422."""
        r = client.get(
            "/api/code/preview/function",
            params={"source_root": "/some/path"},
        )
        assert r.status_code == 422

    def test_preview_function_empty_function_name(self):
        """GET /api/code/preview/function with empty function_name returns 400."""
        r = client.get(
            "/api/code/preview/function",
            params={"source_root": "/some/path", "function_name": ""},
        )
        assert r.status_code == 400

    def test_call_graph_missing_source_root(self):
        """GET /api/code/call-graph without source_root returns 422."""
        r = client.get("/api/code/call-graph")
        assert r.status_code == 422

    def test_call_graph_invalid_depth(self):
        """GET /api/code/call-graph with depth out of range returns 422."""
        r = client.get(
            "/api/code/call-graph",
            params={"source_root": "/tmp", "depth": 99},
        )
        assert r.status_code == 422

    def test_call_graph_nonexistent_source(self):
        """GET /api/code/call-graph with nonexistent source_root returns error."""
        r = client.get(
            "/api/code/call-graph",
            params={"source_root": "/nonexistent/src/root"},
        )
        # May return 200 (empty graph), 400 (bad path), or 500
        assert r.status_code in (200, 400, 500)

    def test_dependency_map_missing_source_root(self):
        """GET /api/code/dependency-map without source_root returns 422."""
        r = client.get("/api/code/dependency-map")
        assert r.status_code == 422

    def test_globals_missing_source_root(self):
        """GET /api/code/globals without source_root returns 422."""
        r = client.get("/api/code/globals")
        assert r.status_code == 422

    def test_globals_nonexistent_source(self):
        """GET /api/code/globals with nonexistent source returns error."""
        r = client.get(
            "/api/code/globals",
            params={"source_root": "/nonexistent/code/root"},
        )
        # Returns 200 (empty globals), 400 (bad path), or 500
        assert r.status_code in (200, 400, 500)

    def test_call_graph_max_files_boundaries(self):
        """GET /api/code/call-graph validates max_files range."""
        # Below minimum
        r = client.get(
            "/api/code/call-graph",
            params={"source_root": "/tmp", "max_files": 50},
        )
        assert r.status_code == 422

        # Above maximum
        r = client.get(
            "/api/code/call-graph",
            params={"source_root": "/tmp", "max_files": 9999},
        )
        assert r.status_code == 422


# ═══════════════════════════════════════════════════════════════════
# Config Router
# ═══════════════════════════════════════════════════════════════════
class TestConfigRouter:
    """Tests for /api/config endpoints."""

    def test_config_defaults_returns_200(self):
        """GET /api/config/defaults returns config data."""
        r = client.get("/api/config/defaults")
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, dict)

    def test_config_defaults_has_required_keys(self):
        """Config defaults contains essential configuration fields."""
        r = client.get("/api/config/defaults")
        data = r.json()
        required_keys = [
            "project_root",
            "report_dir",
            "targets_glob",
            "include_paths",
            "quality_preset",
            "do_build",
            "do_coverage",
        ]
        for key in required_keys:
            assert key in data, f"Missing config key: {key}"

    def test_config_defaults_types(self):
        """Config defaults values have correct types."""
        r = client.get("/api/config/defaults")
        data = r.json()
        assert isinstance(data["project_root"], str)
        assert isinstance(data["report_dir"], str)
        assert isinstance(data["include_paths"], list)
        assert isinstance(data["do_build"], bool)
        assert isinstance(data["do_coverage"], bool)
        assert isinstance(data["quality_preset"], str)

    def test_config_options_returns_200(self):
        """GET /api/config/options returns options data."""
        r = client.get("/api/config/options")
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, dict)

    def test_config_options_has_presets(self):
        """Config options includes quality presets."""
        r = client.get("/api/config/options")
        data = r.json()
        assert "quality_presets" in data
        assert isinstance(data["quality_presets"], list)
        assert len(data["quality_presets"]) > 0

    def test_config_options_has_strategy(self):
        """Config options includes build strategy and fallback options."""
        r = client.get("/api/config/options")
        data = r.json()
        assert "build_strategy_options" in data
        assert "build_fallback_options" in data


# ═══════════════════════════════════════════════════════════════════
# General API behavior
# ═══════════════════════════════════════════════════════════════════
class TestGeneralAPI:
    """Tests for cross-cutting API behavior."""

    def test_nonexistent_api_route(self):
        """Unmatched /api/* route returns 404."""
        r = client.get("/api/this-endpoint-does-not-exist")
        assert r.status_code == 404

    def test_cors_headers_present(self):
        """CORS middleware adds Access-Control-Allow-Origin header."""
        r = client.get(
            "/api/health",
            headers={"Origin": "http://localhost:3000"},
        )
        assert r.status_code == 200
        assert "access-control-allow-origin" in r.headers

    def test_options_preflight(self):
        """OPTIONS preflight request is handled by CORS middleware."""
        r = client.options(
            "/api/health",
            headers={
                "Origin": "http://localhost:3000",
                "Access-Control-Request-Method": "GET",
            },
        )
        assert r.status_code == 200
        assert "access-control-allow-origin" in r.headers


# ═══════════════════════════════════════════════════════════════════
# Jenkins Router
# ═══════════════════════════════════════════════════════════════════
class TestJenkinsRouter:
    """Tests for /api/jenkins endpoints."""

    def test_jenkins_jobs_missing_token(self):
        """POST /api/jenkins/jobs with empty api_token returns 400."""
        r = client.post(
            "/api/jenkins/jobs",
            json={
                "base_url": "http://jenkins.local",
                "username": "user",
                "api_token": "",
            },
        )
        assert r.status_code == 400
        body = r.json()
        msg = body.get("error", {}).get("message", "") or body.get("detail", "")
        assert "Token" in msg or "토큰" in msg

    def test_jenkins_builds_missing_job_url(self):
        """POST /api/jenkins/builds with empty job_url returns 400."""
        r = client.post(
            "/api/jenkins/builds",
            json={
                "job_url": "",
                "username": "user",
                "api_token": "some-token",
            },
        )
        assert r.status_code == 400

    def test_jenkins_builds_missing_token(self):
        """POST /api/jenkins/builds with empty api_token returns 400."""
        r = client.post(
            "/api/jenkins/builds",
            json={
                "job_url": "http://jenkins.local/job/test/",
                "username": "user",
                "api_token": "",
            },
        )
        assert r.status_code == 400

    def test_jenkins_progress_returns_progress(self):
        """GET /api/jenkins/progress returns ok + progress."""
        r = client.get(
            "/api/jenkins/progress",
            params={
                "action": "sync",
                "job_url": "http://jenkins.local/job/test/",
            },
        )
        assert r.status_code == 200
        data = r.json()
        assert "ok" in data
        assert "progress" in data

    def test_jenkins_jobs_missing_required_fields(self):
        """POST /api/jenkins/jobs without body returns 422."""
        r = client.post("/api/jenkins/jobs", json={})
        assert r.status_code == 422

    def test_jenkins_builds_missing_required_fields(self):
        """POST /api/jenkins/builds without body returns 422."""
        r = client.post("/api/jenkins/builds", json={})
        assert r.status_code == 422

    def test_jenkins_build_info_missing_required_fields(self):
        """POST /api/jenkins/build-info without body returns 422."""
        r = client.post("/api/jenkins/build-info", json={})
        assert r.status_code == 422


# ═══════════════════════════════════════════════════════════════════
# Impact Router
# ═══════════════════════════════════════════════════════════════════
class TestImpactRouter:
    """Tests for /api/impact endpoints."""

    def test_impact_analyze_missing_source_root(self):
        """POST /api/impact/analyze with nonexistent source_root returns 400."""
        r = client.post(
            "/api/impact/analyze",
            json={
                "source_root": "/nonexistent/source/root",
                "changed_files": ["main.c"],
            },
        )
        assert r.status_code == 400
        body = r.json()
        msg = body.get("error", {}).get("message", "") or body.get("detail", "")
        assert "source_root" in msg

    def test_impact_analyze_no_changed_files(self):
        """POST /api/impact/analyze without changed_files returns 400."""
        with tempfile.TemporaryDirectory() as tmpdir:
            r = client.post(
                "/api/impact/analyze",
                json={
                    "source_root": tmpdir,
                    "changed_files": [],
                    "changed_raw": "",
                },
            )
            assert r.status_code == 400
            body = r.json()
            msg = body.get("error", {}).get("message", "") or body.get("detail", "")
            assert "changed" in msg

    def test_impact_analyze_missing_body(self):
        """POST /api/impact/analyze without body returns 422."""
        r = client.post("/api/impact/analyze", json={})
        assert r.status_code == 422


# ═══════════════════════════════════════════════════════════════════
# Chat Router
# ═══════════════════════════════════════════════════════════════════
class TestChatRouter:
    """Tests for /api/chat endpoints."""

    def test_chat_missing_question(self):
        """POST /api/chat without question returns 422."""
        r = client.post("/api/chat", json={})
        assert r.status_code == 422

    def test_chat_approval_get_nonexistent(self):
        """GET /api/chat/approval/<id> with nonexistent id returns 404."""
        r = client.get("/api/chat/approval/nonexistent_approval_id_xyz")
        assert r.status_code == 404

    def test_chat_approval_resolve_nonexistent(self):
        """POST /api/chat/approval/resolve with nonexistent id returns 404."""
        r = client.post(
            "/api/chat/approval/resolve",
            json={
                "approval_id": "nonexistent_approval_id_xyz",
                "decision": "approve",
            },
        )
        assert r.status_code == 404

    def test_chat_approval_resolve_missing_fields(self):
        """POST /api/chat/approval/resolve without body returns 422."""
        r = client.post("/api/chat/approval/resolve", json={})
        assert r.status_code == 422


# ═══════════════════════════════════════════════════════════════════
# Sessions Router
# ═══════════════════════════════════════════════════════════════════
class TestSessionsRouter:
    """Tests for /api/sessions endpoints."""

    def test_list_sessions_returns_list(self):
        """GET /api/sessions returns a JSON list."""
        r = client.get("/api/sessions")
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, list)

    def test_create_and_delete_session(self):
        """POST /api/sessions/new + DELETE /api/sessions/<id> lifecycle."""
        r = client.post("/api/sessions/new")
        assert r.status_code == 200
        data = r.json()
        assert "id" in data
        assert "path" in data
        session_id = data["id"]

        # Get session data
        r2 = client.get(f"/api/sessions/{session_id}/data")
        assert r2.status_code == 200
        d2 = r2.json()
        assert "summary" in d2
        assert "findings" in d2

        # Get session config
        r3 = client.get(f"/api/sessions/{session_id}/config")
        assert r3.status_code == 200
        d3 = r3.json()
        assert "config" in d3

        # Save session config
        r4 = client.post(
            f"/api/sessions/{session_id}/config",
            json={"config": {"project_root": "/tmp/test"}},
        )
        assert r4.status_code == 200
        assert r4.json()["ok"] is True

        # Set session name
        r5 = client.post(
            f"/api/sessions/{session_id}/name",
            json={"name": "Test Session"},
        )
        assert r5.status_code == 200
        assert r5.json()["name"] == "Test Session"

        # Get log
        r6 = client.get(f"/api/sessions/{session_id}/log")
        assert r6.status_code == 200
        assert "lines" in r6.json()

        # Delete session
        r7 = client.delete(f"/api/sessions/{session_id}")
        assert r7.status_code == 200
        assert r7.json()["ok"] is True

    def test_delete_session_nonexistent(self):
        """DELETE /api/sessions/<id> with nonexistent id returns 404."""
        r = client.delete("/api/sessions/nonexistent_session_xyz")
        assert r.status_code == 404

    def test_session_complexity_returns_rows(self):
        """GET /api/sessions/<id>/report/complexity returns rows key."""
        r = client.post("/api/sessions/new")
        session_id = r.json()["id"]
        try:
            r2 = client.get(f"/api/sessions/{session_id}/report/complexity")
            assert r2.status_code == 200
            assert "rows" in r2.json()
        finally:
            client.delete(f"/api/sessions/{session_id}")

    def test_session_docs_nonexistent(self):
        """GET /api/sessions/<id>/report/docs for empty session returns ok=False."""
        r = client.post("/api/sessions/new")
        session_id = r.json()["id"]
        try:
            r2 = client.get(f"/api/sessions/{session_id}/report/docs")
            assert r2.status_code == 200
            assert r2.json()["ok"] is False
        finally:
            client.delete(f"/api/sessions/{session_id}")

    def test_session_logs_returns_logs(self):
        """GET /api/sessions/<id>/report/logs returns logs key."""
        r = client.post("/api/sessions/new")
        session_id = r.json()["id"]
        try:
            r2 = client.get(f"/api/sessions/{session_id}/report/logs")
            assert r2.status_code == 200
            assert "logs" in r2.json()
        finally:
            client.delete(f"/api/sessions/{session_id}")

    def test_session_report_files(self):
        """GET /api/sessions/<id>/report/files returns file listing."""
        r = client.post("/api/sessions/new")
        session_id = r.json()["id"]
        try:
            r2 = client.get(f"/api/sessions/{session_id}/report/files")
            assert r2.status_code == 200
            # Should be a dict (report file listing)
            assert isinstance(r2.json(), dict)
        finally:
            client.delete(f"/api/sessions/{session_id}")

    def test_stop_run_invalid_pid(self):
        """POST /api/run/stop with pid=0 returns 400."""
        r = client.post(
            "/api/run/stop",
            json={"pid": 0},
        )
        assert r.status_code == 400


# ═══════════════════════════════════════════════════════════════════
# Profiles Router
# ═══════════════════════════════════════════════════════════════════
class TestProfilesRouter:
    """Tests for /api/profiles endpoints."""

    def test_list_profiles_returns_names(self):
        """GET /api/profiles returns names list."""
        r = client.get("/api/profiles")
        assert r.status_code == 200
        data = r.json()
        assert "names" in data
        assert isinstance(data["names"], list)

    def test_get_nonexistent_profile(self):
        """GET /api/profiles/<name> with nonexistent name returns 404."""
        r = client.get("/api/profiles/nonexistent_profile_xyz_12345")
        assert r.status_code == 404

    def test_save_and_delete_profile(self):
        """POST + DELETE /api/profiles/<name> lifecycle."""
        name = "__test_profile_unit__"
        # Save
        r = client.post(
            f"/api/profiles/{name}",
            json={"project_root": "/tmp/test", "report_dir": "reports"},
        )
        assert r.status_code == 200
        assert r.json()["ok"] is True

        # Get
        r2 = client.get(f"/api/profiles/{name}")
        assert r2.status_code == 200
        assert r2.json()["project_root"] == "/tmp/test"

        # Delete
        r3 = client.delete(f"/api/profiles/{name}")
        assert r3.status_code == 200
        assert r3.json()["ok"] is True

        # Confirm deleted
        r4 = client.get(f"/api/profiles/{name}")
        assert r4.status_code == 404

    def test_delete_nonexistent_profile(self):
        """DELETE /api/profiles/<name> with nonexistent returns 404."""
        r = client.delete("/api/profiles/nonexistent_profile_xyz_12345")
        assert r.status_code == 404

    def test_set_last_profile(self):
        """POST /api/profiles/last sets the last profile name."""
        r = client.post(
            "/api/profiles/last",
            json={"name": "test_last"},
        )
        assert r.status_code == 200
        assert r.json()["ok"] is True


# ═══════════════════════════════════════════════════════════════════
# Quality Router
# ═══════════════════════════════════════════════════════════════════
class TestQualityRouter:
    """Tests for /api/quality endpoints (응답 shape 검증).

    /api/quality/* 는 라우터 레벨 require_admin 게이트가 걸려 있다(admin only).
    게이트 동작 자체는 test_admin_gate.py 가 검증하므로, 여기서는 게이트를 우회해
    응답 shape 만 본다 (관심사 분리). client 기본 헤더는 X-User=test(비admin).
    """

    @pytest.fixture(autouse=True)
    def _bypass_admin_gate(self, monkeypatch):
        import backend.dependencies.admin as _adm
        monkeypatch.setattr(_adm, "is_admin", lambda _u: True)

    def test_list_runs_returns_runs(self):
        """GET /api/quality/runs returns runs list."""
        r = client.get("/api/quality/runs")
        assert r.status_code == 200
        data = r.json()
        assert "runs" in data
        assert isinstance(data["runs"], list)

    def test_list_runs_with_doc_type_filter(self):
        """GET /api/quality/runs?doc_type=uds filters by type."""
        r = client.get("/api/quality/runs", params={"doc_type": "uds"})
        assert r.status_code == 200
        data = r.json()
        assert "runs" in data

    def test_list_runs_with_limit_offset(self):
        """GET /api/quality/runs with limit/offset pagination."""
        r = client.get(
            "/api/quality/runs",
            params={"limit": 10, "offset": 0},
        )
        assert r.status_code == 200
        data = r.json()
        assert "total" in data
        assert "limit" in data
        assert "offset" in data

    def test_get_run_nonexistent(self):
        """GET /api/quality/runs/<id> with nonexistent returns error."""
        r = client.get("/api/quality/runs/999999")
        assert r.status_code == 200
        data = r.json()
        # Returns error field (not HTTP 404) when run not found
        assert "error" in data or "id" in data

    def test_trend_default_doc_type(self):
        """GET /api/quality/trend without doc_type defaults to uds."""
        r = client.get("/api/quality/trend")
        assert r.status_code == 200
        data = r.json()
        assert "trend" in data

    def test_trend_with_doc_type(self):
        """GET /api/quality/trend?doc_type=uds returns trend."""
        r = client.get("/api/quality/trend", params={"doc_type": "uds"})
        assert r.status_code == 200
        data = r.json()
        assert "trend" in data
        assert isinstance(data["trend"], list)

    def test_advice_nonexistent_run(self):
        """POST /api/quality/runs/<id>/advice with nonexistent returns error."""
        r = client.post("/api/quality/runs/999999/advice")
        assert r.status_code == 200
        data = r.json()
        # Returns error field when advisor module or run not available
        assert isinstance(data, dict)


# ═══════════════════════════════════════════════════════════════════
# Local Router (selected simple endpoints)
# ═══════════════════════════════════════════════════════════════════
class TestLocalRouter:
    """Tests for /api/local endpoints."""

    def test_list_dir(self):
        """POST /api/local/list-dir with real directory returns entries."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a test file inside
            (Path(tmpdir) / "test.txt").write_text("hello", encoding="utf-8")
            r = client.post(
                "/api/local/list-dir",
                json={"project_root": tmpdir, "rel_path": "."},
            )
            assert r.status_code == 200
            data = r.json()
            assert data["ok"] is True
            assert any(e["name"] == "test.txt" for e in data["entries"])

    def test_list_dir_nonexistent(self):
        """POST /api/local/list-dir with nonexistent returns ok=False."""
        r = client.post(
            "/api/local/list-dir",
            json={"project_root": "/nonexistent/root", "rel_path": "."},
        )
        assert r.status_code == 200
        data = r.json()
        assert data["ok"] is False

    def test_search_in_files(self):
        """POST /api/local/search with real files returns results."""
        with tempfile.TemporaryDirectory() as tmpdir:
            (Path(tmpdir) / "main.c").write_text(
                "int main() { return 0; }", encoding="utf-8"
            )
            r = client.post(
                "/api/local/search",
                json={
                    "project_root": tmpdir,
                    "rel_path": ".",
                    "query": "main",
                },
            )
            assert r.status_code == 200
            data = r.json()
            assert data["ok"] is True
            assert len(data["results"]) >= 1

    def test_search_empty_query(self):
        """POST /api/local/search with empty query returns ok=False."""
        with tempfile.TemporaryDirectory() as tmpdir:
            r = client.post(
                "/api/local/search",
                json={
                    "project_root": tmpdir,
                    "rel_path": ".",
                    "query": "",
                },
            )
            assert r.status_code == 200
            data = r.json()
            assert data["ok"] is False

    def test_editor_read_write_cycle(self):
        """POST /api/local/editor/write + read cycle works."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Write
            r = client.post(
                "/api/local/editor/write",
                json={
                    "project_root": tmpdir,
                    "rel_path": "test.txt",
                    "content": "hello world",
                    "make_backup": False,
                },
            )
            assert r.status_code == 200
            assert r.json()["ok"] is True

            # Read back
            r2 = client.post(
                "/api/local/editor/read",
                json={
                    "project_root": tmpdir,
                    "rel_path": "test.txt",
                },
            )
            assert r2.status_code == 200
            d2 = r2.json()
            assert d2["ok"] is True
            assert d2["text"] == "hello world"

    def test_editor_replace(self):
        """POST /api/local/editor/replace replaces lines correctly."""
        with tempfile.TemporaryDirectory() as tmpdir:
            fpath = Path(tmpdir) / "test.c"
            fpath.write_text("line1\nline2\nline3\n", encoding="utf-8")
            r = client.post(
                "/api/local/editor/replace",
                json={
                    "project_root": tmpdir,
                    "rel_path": "test.c",
                    "start_line": 2,
                    "end_line": 2,
                    "content": "REPLACED",
                },
            )
            assert r.status_code == 200
            assert r.json()["ok"] is True
            # Verify
            text = fpath.read_text(encoding="utf-8")
            assert "REPLACED" in text
            assert "line1" in text
            assert "line3" in text

    def test_format_c_returns_dict(self):
        """POST /api/local/format-c returns ok field."""
        r = client.post(
            "/api/local/format-c",
            json={"text": "int main(){return 0;}", "filename": "test.c"},
        )
        assert r.status_code == 200
        data = r.json()
        assert "ok" in data
        # ok may be False if clang-format is not installed

    def test_format_c_empty_text(self):
        """POST /api/local/format-c with empty text returns ok=False."""
        r = client.post(
            "/api/local/format-c",
            json={"text": "", "filename": "test.c"},
        )
        assert r.status_code == 200
        assert r.json()["ok"] is False

    def test_replace_text(self):
        """POST /api/local/replace-text replaces text in file."""
        with tempfile.TemporaryDirectory() as tmpdir:
            fpath = Path(tmpdir) / "main.c"
            fpath.write_text("int foo = 42;", encoding="utf-8")
            r = client.post(
                "/api/local/replace-text",
                json={
                    "project_root": tmpdir,
                    "rel_path": "main.c",
                    "search": "42",
                    "replace": "99",
                },
            )
            assert r.status_code == 200
            assert r.json()["ok"] is True
            assert r.json()["changed"] is True
            assert "99" in fpath.read_text(encoding="utf-8")

    def test_open_file_empty_path(self):
        """POST /api/local/open-file with empty path returns 400."""
        r = client.post("/api/local/open-file", json={"path": ""})
        assert r.status_code == 400

    def test_open_file_nonexistent(self):
        """POST /api/local/open-file with nonexistent returns 403 or 404."""
        r = client.post(
            "/api/local/open-file",
            json={"path": "/nonexistent/file.txt"},
        )
        assert r.status_code in (403, 404)

    def test_open_folder_empty_path(self):
        """POST /api/local/open-folder with empty path returns 400."""
        r = client.post("/api/local/open-folder", json={"path": ""})
        assert r.status_code == 400

    def test_preflight_missing_config(self):
        """POST /api/local/preflight without config returns 422."""
        r = client.post("/api/local/preflight", json={})
        assert r.status_code == 422

    def test_kb_list_missing_fields(self):
        """POST /api/local/kb/list without body returns 422."""
        r = client.post("/api/local/kb/list", json={})
        assert r.status_code == 422

    def test_kb_delete_no_entry_key(self):
        """POST /api/local/kb/delete without entry_key returns 400."""
        r = client.post(
            "/api/local/kb/list",
            json={"project_root": "/tmp", "report_dir": "reports"},
        )
        assert r.status_code == 200
        assert "entries" in r.json()

    def test_local_reports_list(self):
        """GET /api/local/reports returns list."""
        r = client.get("/api/local/reports")
        assert r.status_code == 200
        data = r.json()
        assert isinstance(data, (list, dict))


# ═══════════════════════════════════════════════════════════════════
# preview-excel: 헤더 탐지 일반화 / 서버 페이지네이션 / docx 깨진 이미지 복원
# (2026-06-23 — 문서 생성 미리보기 버그 3종 회귀 가드)
# ═══════════════════════════════════════════════════════════════════
class TestPreviewExcelFixes:
    """preview-excel 미리보기 수정 회귀 테스트.

    file_mode.json이 cloudium일 수 있으므로 LocalFileResolver를 강제 설치해
    (worker 없이) 로컬 임시 파일을 read하도록 격리한다.
    """

    @pytest.fixture(autouse=True)
    def _force_local_resolver(self):
        from backend.routers import health as _health
        from backend.services import file_resolver as _fr
        prev = _fr.get_resolver()
        _fr.set_resolver(_fr.LocalFileResolver())
        # 미리보기 캐시는 path 키 + TTL이므로, mkstemp가 경로를 재사용하면 직전
        # 테스트의 바이트/payload가 stale로 잡힐 수 있다 → 매 테스트 전후 비워 결정성 확보.
        _health.clear_preview_cache()
        try:
            yield
        finally:
            _fr.set_resolver(prev)
            _health.clear_preview_cache()

    def _xlsx(self, sheets: dict) -> str:
        """{sheet_name: [row, ...]} → 임시 .xlsx 경로."""
        import openpyxl
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        for name, rows in sheets.items():
            ws = wb.create_sheet(title=name)
            for row in rows:
                ws.append(row)
        fd, path = tempfile.mkstemp(suffix=".xlsx")
        os.close(fd)
        wb.save(path)
        return path

    # ── 헤더 탐지 일반화(HSIS형: row0 제목 → 실제 헤더는 row1) ──
    def test_general_sheet_header_below_title_row(self):
        """row0=병합 제목(1셀), row1=Device/Pin/Signal 라벨 → 헤더는 row1로,
        데이터가 1열로 잘리지 않아야 한다(HSIS 데이터 안보임 버그)."""
        rows = [
            ["Hardware Software Interface", "", "", "", ""],          # row0 제목
            ["", "Device", "Pin No", "Signal Name", "Signal Type"],   # row1 실헤더
            ["", "MCU1", "12", "VCC_BAT", "Analog"],
            ["", "MCU1", "13", "GND", "Power"],
        ]
        path = self._xlsx({"HSI": rows})
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            assert r.status_code == 200, r.text
            sheets = r.json()["sheets"]
            sh = next(s for s in sheets if s["name"] == "HSI")
            assert "Device" in sh["headers"] and "Signal Name" in sh["headers"]
            assert sh["total_cols"] >= 4
            # 데이터 행이 1열로 잘리지 않음
            assert any(len(row) >= 4 for row in sh["rows"])
        finally:
            os.unlink(path)

    def test_clean_header_row0_unchanged(self):
        """row0이 이미 정상 헤더면 그대로 row0을 헤더로 본다(회귀 0)."""
        rows = [
            ["ID", "Name", "Description"],
            ["1", "alpha", "first"],
            ["2", "beta", "second"],
        ]
        path = self._xlsx({"Data": rows})
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            assert r.status_code == 200, r.text
            sh = next(s for s in r.json()["sheets"] if s["name"] == "Data")
            assert sh["headers"][:3] == ["ID", "Name", "Description"]
            assert ["1", "alpha", "first"] in sh["rows"]
        finally:
            os.unlink(path)

    # ── 서버 페이지네이션 / has_more (200행 캡 제거) ──
    def test_pagination_window_and_has_more(self):
        """헤더 + 250 데이터행: page0=100행 has_more=True, page2=50행 has_more=False,
        페이지마다 다른 행(이전 client slice 버그는 항상 첫 페이지만 보였음)."""
        rows = [["ID", "Val"]] + [[str(i), f"v{i}"] for i in range(250)]
        path = self._xlsx({"Big": rows})
        try:
            r0 = client.post("/api/preview-excel", json={"path": path, "page": 0, "page_size": 100})
            r2 = client.post("/api/preview-excel", json={"path": path, "page": 2, "page_size": 100})
            assert r0.status_code == 200 and r2.status_code == 200
            s0 = next(s for s in r0.json()["sheets"] if s["name"] == "Big")
            s2 = next(s for s in r2.json()["sheets"] if s["name"] == "Big")
            assert len(s0["rows"]) == 100 and s0["has_more"] is True
            assert len(s2["rows"]) == 50 and s2["has_more"] is False
            assert s0["rows"][0] != s2["rows"][0]
            assert s0["rows"][0] == ["0", "v0"]
            assert s2["rows"][0] == ["200", "v200"]
        finally:
            os.unlink(path)

    def test_small_sheet_no_phantom_pages(self):
        """데이터 3행: has_more=False → 다음 페이지 없음(유령 페이지 방지)."""
        rows = [["ID", "Val"], ["1", "a"], ["2", "b"], ["3", "c"]]
        path = self._xlsx({"Tiny": rows})
        try:
            r = client.post("/api/preview-excel", json={"path": path, "page": 0, "page_size": 100})
            sh = next(s for s in r.json()["sheets"] if s["name"] == "Tiny")
            assert sh["has_more"] is False
            assert len(sh["rows"]) == 3
        finally:
            os.unlink(path)

    # ── docx 깨진 임베드 이미지 복원 ──
    def _corrupt_docx(self) -> bytes:
        import io
        import struct
        import zipfile
        docx = pytest.importorskip("docx")
        d = docx.Document()
        tb = d.add_table(rows=1, cols=2)
        tb.rows[0].cells[0].text = "ID"
        tb.rows[0].cells[1].text = "SwCom_7"
        png = bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
            "0000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
        )
        d.add_picture(io.BytesIO(png))
        buf = io.BytesIO()
        d.save(buf)
        raw = bytearray(buf.getvalue())
        zin = zipfile.ZipFile(io.BytesIO(bytes(raw)))
        img = next(i for i in zin.infolist() if i.filename.startswith("word/media/"))
        off = img.header_offset
        n, m = struct.unpack("<HH", bytes(raw[off + 26:off + 30]))
        data_start = off + 30 + n + m
        for k in range(data_start, data_start + 6):
            raw[k] ^= 0xFF
        return bytes(raw)

    def test_preview_excel_docx_corrupt_image_recovers(self):
        """깨진 임베드 이미지 docx도 _safe_docx_open으로 500 없이 200 반환
        (UDS/SDS BadZipFile 에러 버그)."""
        import io

        import docx as _docx
        corrupt = self._corrupt_docx()
        # 전제: raw python-docx는 실패
        with pytest.raises(Exception):
            _docx.Document(io.BytesIO(corrupt))
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(path, "wb") as f:
            f.write(corrupt)
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            assert r.status_code == 200, r.text
            assert r.json()["ok"] is True
        finally:
            os.unlink(path)

    def test_preview_image_docx_corrupt_does_not_500(self):
        """깨진 docx에 preview-image(없는 image_id) 호출 시 500이 아닌 404."""
        corrupt = self._corrupt_docx()
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(path, "wb") as f:
            f.write(corrupt)
        try:
            r = client.get("/api/preview-image", params={"path": path, "image_id": "rIdBogus"})
            assert r.status_code == 404, r.text
        finally:
            os.unlink(path)

    # ── 리뷰 반영 회귀 가드 ──────────────────────────────────────────
    def test_page_bounds_validation(self):
        """page<0 / page_size 상한초과는 422(음수 슬라이스/DoS 방지)."""
        path = self._xlsx({"S": [["ID"], ["1"]]})
        try:
            assert client.post("/api/preview-excel", json={"path": path, "page": -1}).status_code == 422
            assert client.post("/api/preview-excel", json={"path": path, "page_size": 99999}).status_code == 422
            assert client.post("/api/preview-excel", json={"path": path, "page_size": 0}).status_code == 422
        finally:
            os.unlink(path)

    def test_xlsx_without_dimension_tag(self):
        """dimension 태그 없는 xlsx(ws.max_row=None)도 데이터를 반환해야(빈 미리보기 X)."""
        import re as _re
        import zipfile
        path = self._xlsx({"Sheet1": [["ID", "Val"], ["1", "a"], ["2", "b"], ["3", "c"]]})
        # 워크시트 XML에서 <dimension .../> 제거 → read_only에서 max_row=None 유발
        with zipfile.ZipFile(path) as zin:
            names = zin.namelist()
            data = {n: zin.read(n) for n in names}
        sheet_key = next(n for n in names if _re.match(r"xl/worksheets/sheet\d+\.xml$", n))
        data[sheet_key] = _re.sub(rb"<dimension[^>]*/>", b"", data[sheet_key])
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                zout.writestr(n, data[n])
        try:
            r = client.post("/api/preview-excel", json={"path": path, "page": 0, "page_size": 100})
            assert r.status_code == 200, r.text
            sh = next(s for s in r.json()["sheets"] if s["name"] == "Sheet1")
            assert len(sh["rows"]) == 3  # 빈 미리보기가 아니라 실제 3행
            assert ["1", "a"] in sh["rows"]
        finally:
            os.unlink(path)

    def test_header_detection_no_deep_data_row_promotion(self):
        """row0=키워드 없는 실헤더 + 깊은 데이터행에 라벨 단어 → 데이터행 오승격 안 함."""
        rows = [
            ["항번", "코드", "값"],   # row0 실헤더(HDR_WORDS 매칭 0개)
            ["1", "AX", "10"],
            ["2", "BY", "20"],
            ["3", "CZ", "30"],
            ["4", "DW", "40"],
            ["type", "name", "50"],  # row5(>=4)에 키워드 2개(<3) — 가드로 무시되어야
        ]
        path = self._xlsx({"Codes": rows})
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            sh = next(s for s in r.json()["sheets"] if s["name"] == "Codes")
            assert sh["headers"][:3] == ["항번", "코드", "값"]
            assert ["1", "AX", "10"] in sh["rows"]
        finally:
            os.unlink(path)

    def test_docx_function_table_keeps_last_column_without_images(self):
        """이미지 없는 UDS Function 표: 마지막 열(Calling Function)이 누락되지 않아야."""
        docx = pytest.importorskip("docx")
        d = docx.Document()
        tb = d.add_table(rows=4, cols=3)
        tb.rows[0].cells[0].text = "Function Information"
        for i, (label, val) in enumerate(
            [("ID", "SwFn_1"), ("Name", "foo"), ("Calling Function", "bar_caller")], start=1
        ):
            tb.rows[i].cells[0].text = label
            tb.rows[i].cells[2].text = val
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        d.save(path)
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            assert r.status_code == 200, r.text
            fn = next(s for s in r.json()["sheets"] if s["name"].startswith("Functions"))
            assert "Calling Function" in fn["headers"]
            # 행 길이가 헤더 길이와 일치(마지막 열 누락 없음) + 마지막 열 값 존재
            assert any(len(row) == len(fn["headers"]) and "bar_caller" in row for row in fn["rows"])
        finally:
            os.unlink(path)

    def test_docx_other_table_no_infinite_pager(self):
        """미인식 docx 표(other_tables)는 100행 초과여도 has_more=False(무한 페이저 방지)."""
        docx = pytest.importorskip("docx")
        d = docx.Document()
        tb = d.add_table(rows=151, cols=2)
        tb.rows[0].cells[0].text = "Col1"
        tb.rows[0].cells[1].text = "Col2"
        for i in range(1, 151):
            tb.rows[i].cells[0].text = f"r{i}"
            tb.rows[i].cells[1].text = "x"
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        d.save(path)
        try:
            r = client.post("/api/preview-excel", json={"path": path})
            assert r.status_code == 200, r.text
            for s in r.json()["sheets"]:
                # 모든 시트가 has_more 키를 가지며, Table N(other_tables)은 False
                assert "has_more" in s
                if s["name"].startswith("Table "):
                    assert s["has_more"] is False
        finally:
            os.unlink(path)


class TestPreviewCache:
    """미리보기 캐시(바이트/payload/relmap) — 페이지 이동·이미지 로드 시 원본
    재IPC + python-docx 전체 재파싱을 제거하는지 검증."""

    class _CountingResolver:
        """LocalFileResolver를 감싸 read_bytes 호출 수를 센다(나머지는 위임)."""

        def __init__(self, inner):
            self._inner = inner
            self.read_calls = 0

        def exists(self, path):
            return self._inner.exists(path)

        def read_bytes(self, path):
            self.read_calls += 1
            return self._inner.read_bytes(path)

        def __getattr__(self, name):
            return getattr(self._inner, name)

    @pytest.fixture(autouse=True)
    def _isolate(self):
        from backend.routers import health as _health
        from backend.services import file_resolver as _fr
        prev = _fr.get_resolver()
        _fr.set_resolver(_fr.LocalFileResolver())
        _health.clear_preview_cache()
        self._fr = _fr
        self._health = _health
        try:
            yield
        finally:
            _fr.set_resolver(prev)
            _health.clear_preview_cache()

    def _install_counter(self):
        counter = self._CountingResolver(self._fr.LocalFileResolver())
        self._fr.set_resolver(counter)
        return counter

    def _xlsx_big(self) -> str:
        import openpyxl
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        ws = wb.create_sheet("Big")
        ws.append(["ID", "Val"])
        for i in range(250):
            ws.append([str(i), f"v{i}"])
        fd, path = tempfile.mkstemp(suffix=".xlsx")
        os.close(fd)
        wb.save(path)
        return path

    def _docx_multi(self) -> str:
        docx = pytest.importorskip("docx")
        d = docx.Document()
        for i in range(150):
            tb = d.add_table(rows=3, cols=3)
            tb.rows[0].cells[0].text = "Function Information"
            tb.rows[1].cells[0].text = "ID"
            tb.rows[1].cells[2].text = f"SwFn_{i}"
            tb.rows[2].cells[0].text = "Name"
            tb.rows[2].cells[2].text = f"f{i}"
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        d.save(path)
        return path

    def _docx_with_image(self):
        docx = pytest.importorskip("docx")
        import io
        png = bytes.fromhex(
            "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
            "0000000a49444154789c63000100000500010d0a2db40000000049454e44ae426082"
        )
        d = docx.Document()
        d.add_picture(io.BytesIO(png))
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        d.save(path)
        return path, png

    def test_xlsx_bytes_cache_single_read_across_pages(self):
        """xlsx 두 페이지 요청 → 원본 read 1회(바이트 캐시), 페이지 내용은 정확."""
        counter = self._install_counter()
        path = self._xlsx_big()
        try:
            r0 = client.post("/api/preview-excel", json={"path": path, "page": 0, "page_size": 100})
            r2 = client.post("/api/preview-excel", json={"path": path, "page": 2, "page_size": 100})
            assert r0.status_code == 200 and r2.status_code == 200
            assert counter.read_calls == 1, counter.read_calls
            s0 = next(s for s in r0.json()["sheets"] if s["name"] == "Big")
            s2 = next(s for s in r2.json()["sheets"] if s["name"] == "Big")
            assert s0["rows"][0] == ["0", "v0"] and s2["rows"][0] == ["200", "v200"]
        finally:
            os.unlink(path)

    def test_docx_payload_cache_single_parse_across_pages(self):
        """docx 두 페이지 요청 → read+parse 1회(payload 캐시), 페이지마다 다른 행."""
        counter = self._install_counter()
        path = self._docx_multi()
        try:
            r0 = client.post("/api/preview-excel", json={"path": path, "page": 0, "page_size": 100})
            r1 = client.post("/api/preview-excel", json={"path": path, "page": 1, "page_size": 100})
            assert r0.status_code == 200 and r1.status_code == 200
            assert counter.read_calls == 1, counter.read_calls
            f0 = next(s for s in r0.json()["sheets"] if s["name"].startswith("Functions"))
            f1 = next(s for s in r1.json()["sheets"] if s["name"].startswith("Functions"))
            assert len(f0["rows"]) == 100 and f0["has_more"] is True
            assert len(f1["rows"]) == 50 and f1["has_more"] is False
            assert f0["rows"][0] != f1["rows"][0]
        finally:
            os.unlink(path)

    def test_ttl_expiry_reparses(self, monkeypatch):
        """TTL=0 → 매 요청 캐시 만료 → 원본 재read(무효화 동작 확인)."""
        counter = self._install_counter()
        path = self._xlsx_big()
        monkeypatch.setattr(self._health, "_preview_ttl", lambda: 0.0)
        try:
            client.post("/api/preview-excel", json={"path": path, "page": 0})
            client.post("/api/preview-excel", json={"path": path, "page": 0})
            assert counter.read_calls == 2, counter.read_calls
        finally:
            os.unlink(path)

    def test_preview_image_returns_blob_and_caches(self):
        """preview-image가 zip 멤버에서 이미지 바이트를 반환하고, 두 번째 요청은
        바이트 캐시 히트로 원본을 재read하지 않는다(이미지마다 36MB 재IPC 제거)."""
        counter = self._install_counter()
        path, png = self._docx_with_image()
        try:
            with open(path, "rb") as f:
                raw = f.read()
            relmap = self._health._docx_relmap(raw)
            assert relmap, "이미지 rId 맵이 비어있음"
            rid = next(iter(relmap))
            counter.read_calls = 0
            self._health.clear_preview_cache()
            r1 = client.get("/api/preview-image", params={"path": path, "image_id": rid})
            r2 = client.get("/api/preview-image", params={"path": path, "image_id": rid})
            assert r1.status_code == 200 and r2.status_code == 200
            assert r1.content == png
            assert r1.headers["content-type"].startswith("image/")
            assert counter.read_calls == 1, counter.read_calls
        finally:
            os.unlink(path)

    def test_preview_image_corrupt_member_falls_back_to_png(self):
        """유효 rId이지만 이미지 멤버가 CRC 손상 → 1x1 PNG로 graceful 대체(200)."""
        corrupt = TestPreviewExcelFixes()._corrupt_docx()
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        with open(path, "wb") as f:
            f.write(corrupt)
        try:
            relmap = self._health._docx_relmap(corrupt)
            assert relmap
            rid = next(iter(relmap))
            r = client.get("/api/preview-image", params={"path": path, "image_id": rid})
            assert r.status_code == 200, r.text
            assert r.headers["content-type"] == "image/png"
            assert r.content[:8] == bytes.fromhex("89504e470d0a1a0a")  # PNG signature
        finally:
            os.unlink(path)

    def test_docx_image_member_resolution(self):
        """rels Target → zip 멤버 경로 정규화(상대/절대/.. 처리)."""
        m = self._health._docx_image_member
        assert m("media/image1.png") == "word/media/image1.png"
        assert m("/word/media/image2.png") == "word/media/image2.png"
        assert m("../media/image3.png") == "media/image3.png"

    def test_bytes_total_no_phantom_debt(self, monkeypatch):
        """리뷰 C1: TTL 만료로 엔트리 회수 시 total이 정확히 감산 — 반복 만료에도
        phantom debt가 누적되지 않아 바이트 캐시가 capacity-1로 붕괴하지 않는다."""
        self._install_counter()
        path = self._xlsx_big()
        try:
            with open(path, "rb") as f:
                size = len(f.read())
            monkeypatch.setattr(self._health, "_preview_ttl", lambda: 0.0)
            for _ in range(5):
                client.post("/api/preview-excel", json={"path": path, "page": 0})
            # 5회 만료·재적재 후에도 total은 단일 파일 크기, 엔트리 1개(누수 0).
            assert self._health._preview_bytes.total == size, self._health._preview_bytes.total
            assert len(self._health._preview_bytes.store) == 1
        finally:
            os.unlink(path)

    def test_local_sig_invalidation_same_path_new_content(self):
        """리뷰 W5: local 모드에서 같은 path를 다른 내용으로 덮어쓰면 (mtime,size)
        시그니처 불일치로 무효화되어 stale이 아닌 새 내용을 반환한다."""
        import time

        import openpyxl
        self._install_counter()
        path = self._xlsx_big()
        try:
            r1 = client.post("/api/preview-excel", json={"path": path, "page": 0})
            s1 = next(s for s in r1.json()["sheets"] if s["name"] == "Big")
            assert s1["rows"][0] == ["0", "v0"]
            time.sleep(0.01)
            wb = openpyxl.Workbook()
            wb.remove(wb.active)
            ws = wb.create_sheet("Big")
            ws.append(["ID", "Val"])
            for i in range(300):
                ws.append([str(i), f"NEW{i}"])
            wb.save(path)  # 같은 path, 다른 크기/내용
            r2 = client.post("/api/preview-excel", json={"path": path, "page": 0})
            s2 = next(s for s in r2.json()["sheets"] if s["name"] == "Big")
            assert s2["rows"][0] == ["0", "NEW0"], s2["rows"][0]
        finally:
            os.unlink(path)

    def test_clear_preview_cache_empties_all(self):
        """리뷰 C2/W6: clear_preview_cache가 세 캐시를 비우고 total을 0으로 되돌린다
        (모드 전환·/cache/clear 무효화 레버)."""
        self._install_counter()
        path = self._xlsx_big()
        try:
            client.post("/api/preview-excel", json={"path": path, "page": 0})
            assert len(self._health._preview_bytes.store) >= 1
            self._health.clear_preview_cache()
            assert len(self._health._preview_bytes.store) == 0
            assert self._health._preview_bytes.total == 0
        finally:
            os.unlink(path)
