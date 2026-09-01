"""재작성 경로가 템플릿의 **표지·이력·컨트롤 값**을 잃던 것을 막는다.

## 실측이 드러낸 것 (2026-08-31)

`docgen_template_source` 는 정본을 템플릿으로 쓰는 이유로 "표지·이력·Introduction 이
납품본과 같아집니다" 를 공시한다. 그런데 UDS(docx)는 시트 복사가 아니라 본문
**재작성**이라, 실제로 만들어 보니 셋 다 사실이 아니었다
(KJPDS02_PV 정본 53MB · 표준 템플릿 v0.10, 소스 25파일/57함수):

| | 템플릿 | 산출물(고치기 전) |
|---|---|---|
| 표지(`txbxContent`) | 32 | **0** |
| 콘텐츠 컨트롤(`w:sdt`) | 13 | **0** |
| `[Project Name]` | 8 | **0** (채워지지도 않고 사라졌다) |
| 이력 표 데이터 행 | 29행 중 실데이터 | **전부 빈칸** |

원인은 셋이다.

1. Word 는 표지를 **body 직속 `w:sdt`**(`docPartGallery="Cover Pages"`)로 넣는데
   `_iter_template_blocks` 가 `w:p`/`w:tbl` 만 냈다 → 표지가 통째로 누락.
2. 표를 `(행수, 열수, style, 헤더)` 로만 담고 `_add_blank_table` 로 새로 만들었다.
   생성기가 채우지 않는 표(이력·참조)는 **빈칸**으로 나갔고, 빈칸은 "이력 없음" 으로
   읽힌다.
3. `Paragraph.text` 가 `w:p` **직속** `w:r` 만 읽어 `w:sdt` 안의 값을 놓쳤다 →
   `이 문서는 [Project Name] 프로젝트를 위한…` 이 `이 문서는  프로젝트를 위한…` 으로.
   구멍은 오타로 읽혀 검토에서 넘어간다.

## 왜 이렇게 시험하나

문자열이 아니라 **문서를 만들어** 잰다. 이 라이터는 payload 를 템플릿 heading 에
맞춰 채우는 구조라, 헬퍼 단위 시험만으로는 "실제 산출물에 남는가" 를 못 본다.
합성 템플릿에 표지 sdt / 목차 sdt / 이력 표 / 생성기 소유 표 / sdt 문단 /
텍스트박스 / 하이퍼링크를 한꺼번에 넣고 한 번 생성해 전부 확인한다.
"""
from __future__ import annotations

from pathlib import Path

import pytest

pytest.importorskip("docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@pytest.fixture(autouse=True)
def _no_reference_suds(monkeypatch, tmp_path):
    """저장소 고정 참조 SUDS(40.7MB)를 읽지 않게 한다 — 이 파일과 무관한 축."""
    import config

    monkeypatch.setattr(config, "UDS_REF_SUDS_PATH",
                        str(tmp_path / "no_such_reference.docx"), raising=False)
    monkeypatch.setattr(config, "resolve_uds_template_path", lambda: "", raising=False)


def _sdt(gallery: str, inner_text: str) -> object:
    """body 직속 구조 블록 — Word 가 표지/목차를 이 모양으로 넣는다."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    return parse_xml(
        f'<w:sdt {nsdecls("w")}>'
        f"<w:sdtPr><w:docPartObj>"
        f'<w:docPartGallery w:val="{gallery}"/><w:docPartUnique/>'
        f"</w:docPartObj></w:sdtPr>"
        f"<w:sdtContent><w:p><w:r><w:t>{inner_text}</w:t></w:r></w:p></w:sdtContent>"
        f"</w:sdt>"
    )


def _build_template(path: Path) -> Path:
    """표지·목차·이력·소유표·sdt문단·텍스트박스·하이퍼링크를 한 문서에 담는다."""
    import docx
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    d = docx.Document()
    body = d.element.body

    # ① 표지(보존 대상) · ② 목차 sdt(제외 대상)
    body.insert(0, _sdt("Table of Contents", "TOCMARK"))
    body.insert(0, _sdt("Cover Pages", "COVERMARK [Project Name]"))

    # ③ 생성기가 소유하지 않는 표 — 데이터가 그대로 남아야 한다
    d.add_heading("Revision History", level=1)
    t = d.add_table(rows=3, cols=2)
    for c, v in zip(t.rows[0].cells, ("Version", "Date"), strict=False):
        c.text = v
    for c, v in zip(t.rows[1].cells, ("v1.0", "2026-01-01"), strict=False):
        c.text = v
    for c, v in zip(t.rows[2].cells, ("v1.1", "2026-02-02"), strict=False):
        c.text = v

    # ④ Introduction/Purpose — 값이 `w:sdt` 안에 있는 문단
    d.add_heading("Introduction", level=1)
    d.add_heading("Purpose", level=2)
    p = d.add_paragraph()
    p.add_run("이 문서는 ")
    p._p.append(parse_xml(
        f'<w:sdt {nsdecls("w")}><w:sdtPr/><w:sdtContent>'
        f"<w:r><w:t>[Project Name]</w:t></w:r></w:sdtContent></w:sdt>"))
    p.add_run(" 프로젝트를 위한 설계서다.")

    # ⑤ 생성기가 소유하는 표 — payload 로 다시 채워져야 한다(템플릿 예시는 사라진다)
    d.add_heading("Common Macro Definition", level=1)
    t2 = d.add_table(rows=2, cols=4)
    for c, v in zip(t2.rows[0].cells, ("Macro name", "Type", "Define", "Description"),
                   strict=False):
        c.text = v
    # ⚠ 예시 행에 헤더 낱말이 들어가면 2행 헤더 판정이 그 행을 헤더로 오인해
    #   산출물에 남던 결함이 있었다(2026-09-01 수정 — 판정이 비대칭이 됐다).
    #   그 축은 `test_docx_table_sizing_and_headers.py` 가 실제 템플릿 값으로 잰다.
    #   여기서는 축을 하나만 두려고 계속 걸리지 않는 값을 쓴다.
    for c, v in zip(t2.rows[1].cells, ("EXAMPLE_ROW", "U8", "1", "보기"), strict=False):
        c.text = v

    # ⑥ 텍스트박스(작성 지침) · ⑦ 하이퍼링크(옛 목차) — 둘 다 본문으로 오면 안 된다
    d.add_heading("Software Unit Design", level=1)
    pg = d.add_paragraph()
    pg._p.append(parse_xml(
        f'<w:r {nsdecls("w")}><w:pict>'
        f'<v:shape xmlns:v="urn:schemas-microsoft-com:vml" style="width:1pt">'
        f"<v:textbox><w:txbxContent><w:p><w:r><w:t>GUIDEBOX</w:t></w:r></w:p>"
        f"</w:txbxContent></v:textbox></v:shape></w:pict></w:r>"))
    ph = d.add_paragraph()
    ph._p.append(parse_xml(
        f'<w:hyperlink {nsdecls("w")}><w:r><w:t>HYPERTEXT</w:t></w:r></w:hyperlink>'))

    d.save(str(path))
    return path


def _payload(project: str = "KJPDS02") -> dict:
    return {
        "project_name": project,
        "overview": "o", "requirements": "r", "interfaces": "i",
        "uds_frames": "u", "notes": "n",
        "common_macros": ["PAYLOAD_MACRO"],
        "function_details": {},
    }


@pytest.fixture
def generated(tmp_path):
    import docx

    from report_gen.docx_builder import generate_uds_docx

    tpl = _build_template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
    out = tmp_path / "out.docx"
    generate_uds_docx(str(tpl), _payload(), str(out))
    doc = docx.Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    xml = doc.element.body.xml
    return doc, text, xml


class TestCoverPageSurvives:
    """표지는 body 직속 `w:sdt` 다 — `w:p`/`w:tbl` 만 보면 통째로 사라진다."""

    def test_cover_block_is_kept(self, generated):
        _doc, _text, xml = generated
        assert "COVERMARK" in xml, "표지 블록이 산출물에서 사라졌다"
        assert "Cover Pages" in xml, "표지 sdt 가 문단으로 풀려 서식·바인딩을 잃었다"

    def test_cover_comes_first(self, generated):
        doc, _text, _xml = generated
        tags = [c.tag.split("}")[-1] for c in doc.element.body]
        assert tags and tags[0] == "sdt", f"표지가 맨 앞이 아니다: {tags[:4]}"

    def test_section_properties_stay_last(self, generated):
        """`body.append()` 는 `w:sectPr` **뒤**로 간다 — 되살린 블록이 문서 끝으로 밀린다."""
        doc, _text, _xml = generated
        tags = [c.tag.split("}")[-1] for c in doc.element.body]
        assert tags[-1] == "sectPr", f"구역 속성 뒤에 내용이 붙었다: {tags[-3:]}"

    def test_table_of_contents_block_is_not_duplicated(self, generated):
        """목차 sdt 까지 보존하면 생성기가 넣는 목차와 **두 벌**이 된다."""
        _doc, _text, xml = generated
        assert "TOCMARK" not in xml, "옛 목차 블록이 그대로 남아 목차가 두 벌이 됐다"


class TestTablesTheGeneratorDoesNotOwn:
    """모양만 복제하면 데이터 행이 빈칸으로 나가고, 빈칸은 '이력 없음' 으로 읽힌다."""

    def test_revision_history_rows_survive(self, generated):
        doc, _text, _xml = generated
        cells = [c.text.strip() for t in doc.tables for r in t.rows for c in r.cells]
        assert "v1.0" in cells and "2026-01-01" in cells, (
            f"이력 표 데이터 행이 사라졌다: {cells[:12]}")
        assert "v1.1" in cells, "이력 표의 둘째 행만 사라졌다"

    def test_a_table_the_generator_owns_is_rebuilt_from_the_payload(self, generated):
        """소유 표까지 원본을 남기면 **남의 값이 산출물에 실린다** — 반대 방향 결함."""
        doc, _text, _xml = generated
        cells = [c.text.strip() for t in doc.tables for r in t.rows for c in r.cells]
        assert "EXAMPLE_ROW" not in cells, (
            "생성기가 채우는 표인데 템플릿 예시 값이 그대로 남았다")


class TestContentControlValues:
    """`p.text` 는 `w:sdt` 안을 못 본다 — 문장 한복판이 빈 채로 나갔다."""

    def test_the_sentence_has_no_hole(self, generated):
        _doc, text, _xml = generated
        assert "이 문서는 KJPDS02 프로젝트를 위한" in text, (
            f"프로젝트명 자리가 비었거나 표식이 그대로다: {text[:200]!r}")

    def test_the_bracket_marker_is_gone(self, generated):
        _doc, _text, xml = generated
        assert "[Project Name]" not in xml, "표식이 채워지지 않았다"

    def test_the_project_name_reaches_the_cover(self, generated):
        _doc, _text, xml = generated
        assert "COVERMARK KJPDS02" in xml, "표지의 프로젝트명 자리가 안 채워졌다"


class TestWhatMustStayOut:
    """되살릴 것과 **되살리면 안 되는 것**은 다르다."""

    def test_guidance_text_boxes_do_not_leak_into_the_body(self, generated):
        """표준 템플릿의 '■ 작성 내용' 상자가 본문에 실리면 납품 문서에 지침이 남는다."""
        _doc, text, _xml = generated
        assert "GUIDEBOX" not in text, "작성 지침 텍스트박스가 본문 문단으로 들어왔다"

    def test_hyperlink_text_is_not_inlined(self, generated):
        """템플릿의 옛 목차 항목은 하이퍼링크 안에 산다 — 평문으로 복제되면 목차가 두 벌."""
        _doc, text, _xml = generated
        assert "HYPERTEXT" not in text, "하이퍼링크 텍스트가 평문 문단으로 복제됐다"


class TestHelpers:
    """산출물 시험이 못 짚는 경계는 헬퍼 단위로."""

    def test_para_text_reads_sdt_but_not_textbox_or_hyperlink(self):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        from report_gen.docx_builder import _para_text

        p = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f"<w:r><w:t>A</w:t></w:r>"
            f"<w:sdt><w:sdtContent><w:r><w:t>B</w:t></w:r></w:sdtContent></w:sdt>"
            f"<w:hyperlink><w:r><w:t>C</w:t></w:r></w:hyperlink>"
            f'<w:r><w:pict><v:shape xmlns:v="urn:schemas-microsoft-com:vml">'
            f"<v:textbox><w:txbxContent>"
            f"<w:p><w:r><w:t>D</w:t></w:r></w:p>"
            f"</w:txbxContent></v:textbox></v:shape></w:pict></w:r>"
            f"</w:p>")
        assert _para_text(p) == "AB", "sdt 를 읽고 하이퍼링크·텍스트박스는 빼야 한다"

    def test_toc_sdt_is_recognised_by_the_word_marker(self):
        from report_gen.docx_builder import _is_toc_sdt

        assert _is_toc_sdt(_sdt("Table of Contents", "x")) is True
        assert _is_toc_sdt(_sdt("Cover Pages", "x")) is False

    def test_project_name_is_not_invented(self, tmp_path):
        """값이 없으면 표식을 **그대로 둔다** — 없는 이름을 지어내지 않는다."""
        import docx

        from report_gen.docx_builder import _fill_bracket_project_name

        d = docx.Document()
        d.add_paragraph("이 문서는 [Project Name] 프로젝트")
        assert _fill_bracket_project_name(d, "") == 0
        assert _fill_bracket_project_name(d, "UDS Spec") == 0, (
            "폴백 기본값으로 채우면 문서가 프로젝트명을 지어낸 셈이 된다")
        assert _fill_bracket_project_name(d, "KJPDS02") == 1

    def test_body_block_goes_before_the_section_properties(self):
        import docx
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        from report_gen.docx_builder import _append_body_block

        d = docx.Document()
        d.add_paragraph("x")
        _append_body_block(d, parse_xml(f'<w:p {nsdecls("w")}><w:r><w:t>Z</w:t></w:r></w:p>'))
        tags = [c.tag.split("}")[-1] for c in d.element.body]
        assert tags[-1] == "sectPr", f"구역 속성이 마지막이 아니다: {tags}"


class TestWhatCameFromTheTemplateIsCounted:
    """보존은 옳지만 **침묵**하면 안 된다.

    템플릿이 남의 프로젝트 문서면 그 이력·참조가 산출물에 그대로 실린다. 신원 판정
    (`template_identity`)은 이미 경고하지만, "몇 개를 그대로 가져왔는가" 가 없으면
    검토자가 영향 범위를 못 잰다.
    """

    def _stats(self, tmp_path):
        from report_gen.docx_builder import generate_uds_docx

        tpl = _build_template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
        stats: dict = {}
        generate_uds_docx(str(tpl), _payload(), str(tmp_path / "o.docx"), stats_out=stats)
        return stats

    def test_counts_are_recorded(self, tmp_path):
        stats = self._stats(tmp_path)
        assert stats.get("restored_template_blocks") == 1, (
            f"표지 블록 복원 수가 안 남았다: {stats.get('restored_template_blocks')!r}")
        assert stats.get("preserved_template_tables") == 1, (
            f"원본 유지한 표 수가 안 남았다: {stats.get('preserved_template_tables')!r}")

    def test_the_counts_reach_the_api_surface(self, tmp_path):
        """사이드카에만 있고 화이트리스트에 없으면 API 응답에서 **조용히 잘린다**."""
        from backend.helpers.uds import _gen_stats_result_fields
        from report_gen.docx_builder import generate_uds_docx

        tpl = _build_template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
        out = tmp_path / "o.docx"
        generate_uds_docx(str(tpl), _payload(), str(out))
        summary = _gen_stats_result_fields(out)["gen_stats_summary"] or {}
        assert summary.get("restored_template_blocks") == 1
        assert summary.get("preserved_template_tables") == 1


def test_an_owned_table_with_nothing_to_fill_comes_out_empty(tmp_path):
    """`None`(생성기가 안 건드림)과 `[]`(건드리는데 채울 게 없음)은 **다르다**.

    둘을 접으면, 채울 게 없는 소유 표에 템플릿 **예시 행**이 남아 산출물이 남의
    값을 사실처럼 싣는다. 반대로 접으면 이력 표가 빈칸으로 나간다.
    """
    from report_gen.docx_builder import generate_uds_docx

    tpl = _build_template(tmp_path / "(KJPDS02_SwUDS) tpl.docx")
    payload = _payload()
    payload["common_macros"] = []          # 소유 표인데 채울 것이 없다
    out = tmp_path / "o.docx"
    generate_uds_docx(str(tpl), payload, str(out))

    import docx
    cells = [c.text.strip() for t in docx.Document(str(out)).tables
             for r in t.rows for c in r.cells]
    assert "EXAMPLE_ROW" not in cells, (
        "채울 게 없다고 템플릿 예시 행을 남기면 산출물이 남의 값을 싣는다")
    assert "v1.0" in cells, "같은 판정으로 이력 표까지 비우면 안 된다"


class TestSharedHelperContract:
    """공유 헬퍼의 계약을 바꾸면 **다른 소비처가 조용히 빈 결과를 낸다**.

    `_iter_template_blocks` 를 `(kind, obj)` 튜플로 바꿨더니
    `report_gen/validation.py` 의 두 루프가 duck typing(`hasattr(block, "text")`)으로
    전부 건너뛰어, SwCom 정본 diff 가 **분모 0 짜리 빈 리포트**를 냈다. 예외도 로그도
    없었다 — 가드 9건이 잡아서 알았다.
    """

    def test_the_old_name_still_yields_objects(self, tmp_path):
        import docx

        from report_gen.docx_builder import _iter_template_blocks

        d = docx.Document(str(_build_template(tmp_path / "t.docx")))
        items = _iter_template_blocks(d)
        assert items, "블록이 하나도 안 나온다"
        assert all(hasattr(x, "text") or hasattr(x, "rows") for x in items), (
            "튜플을 내면 duck typing 소비처(validation.py)가 전부 건너뛴다")

    def test_the_new_name_carries_the_kind_and_the_raw_blocks(self, tmp_path):
        import docx

        from report_gen.docx_builder import _iter_body_blocks

        d = docx.Document(str(_build_template(tmp_path / "t.docx")))
        kinds = [k for k, _ in _iter_body_blocks(d)]
        assert "raw" in kinds, "표지 블록이 새 반복자에서도 안 나온다"
        assert {"para", "table"} <= set(kinds)
