"""라운드 80 T1405-2 — ISO 26262 doc ASIL extractor 회귀."""
from __future__ import annotations

import io

import pytest
from docx import Document  # type: ignore

from backend.services.iso26262_doc_asil_extractor import (
    extract_component_asil_from_sds,
    extract_function_asil_from_suds,
    extract_function_name_to_swufn_from_suds,
    extract_supplementary_asil_from_srs,
)


def _make_docx_with_text(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(headers: list[str], rows: list[list[str]]) -> bytes:
    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.cell(ri, ci).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestSudsExtractor:
    """SUDS extractor — 역방향 매칭 (각 ASIL의 직전 SwUFn)."""

    def test_simple_inline_pair(self):
        text = "SwUFn_0101 some content ASIL A"
        result = extract_function_asil_from_suds(_make_docx_with_text(text))
        assert result == {"SwUFn_0101": "A"}

    def test_section_distance_within_1500(self):
        text = "SwUFn_0101\n" + ("filler " * 100) + "\nASIL B"
        result = extract_function_asil_from_suds(_make_docx_with_text(text))
        assert result.get("SwUFn_0101") == "B"

    def test_section_distance_exceeds_1500_skip(self):
        text = "SwUFn_0101\n" + ("filler " * 250) + "\nASIL D"
        result = extract_function_asil_from_suds(_make_docx_with_text(text))
        assert "SwUFn_0101" not in result

    def test_reverse_match_picks_nearest_swufn(self):
        text = "SwUFn_0101 spacer SwUFn_0102 spacer ASIL C"
        result = extract_function_asil_from_suds(_make_docx_with_text(text))
        # ASIL C 직전 SwUFn은 0102 — 0101은 이전 단계로 밀림
        assert result.get("SwUFn_0102") == "C"
        # 0101은 직전 SwUFn 부재 (ASIL 라벨 직전이 0102라서)
        assert "SwUFn_0101" not in result

    def test_empty_bytes(self):
        warns: list[str] = []
        result = extract_function_asil_from_suds(b"", warns)
        assert result == {}
        assert any("비어있음" in w for w in warns)


class TestSdsExtractor:
    """SDS extractor — 표에서 컴포넌트 ASIL 추출."""

    def test_component_asil_table(self):
        bytes_ = _make_docx_with_table(
            ["ID", "Name", "ASIL"],
            [
                ["SwCom_01", "System OS", "A"],
                ["SwCom_02", "DRV IN", "QM"],
                ["SwCom_03", "API", "B"],
            ],
        )
        result = extract_component_asil_from_sds(bytes_)
        assert result.get("SwCom_01") == "A"
        assert result.get("SwCom_02") == "QM"
        assert result.get("SwCom_03") == "B"
        # 별칭도 등록
        assert result.get("System OS") == "A"

    def test_no_asil_header_returns_empty(self):
        bytes_ = _make_docx_with_table(
            ["Name", "Description"],
            [["SwCom_01", "System OS"]],
        )
        warns: list[str] = []
        result = extract_component_asil_from_sds(bytes_, warns)
        assert result == {}
        assert any("매핑 0건" in w for w in warns)


class TestSrsExtractor:
    """SRS extractor — 함수명 보조 ASIL."""

    def test_function_asil_pair(self):
        text = "The function g_DoorState handles ASIL A door operations."
        result = extract_supplementary_asil_from_srs(_make_docx_with_text(text))
        assert result.get("g_DoorState") == "A"

    def test_korean_postposition_stripped(self):
        text = "g_ApiIn_MotorPosition과 같은 함수는 ASIL A 등급이다."
        result = extract_supplementary_asil_from_srs(_make_docx_with_text(text))
        assert result.get("g_ApiIn_MotorPosition") == "A"

    def test_no_match_returns_empty(self):
        text = "Plain text without ASIL annotations."
        warns: list[str] = []
        result = extract_supplementary_asil_from_srs(_make_docx_with_text(text), warns)
        assert result == {}
        assert any("0건" in w for w in warns)


class TestRound85SudsReverseMap:
    """라운드 85 T1901: SUDS docx 함수명 ↔ SwUFn reverse map 추출."""

    def test_basic_function_name_to_swufn_extraction(self):
        """`SwUFn_NNNN: 함수명` 패턴 인식 — 함수명 keyed reverse map."""
        text = "SwUFn_0101: main\nSwUFn_0201: g_DrvIn_Main\nSwUFn_0301: s_SystemOp"
        result = extract_function_name_to_swufn_from_suds(_make_docx_with_text(text))
        assert result.get("main") == "SwUFn_0101"
        assert result.get("g_DrvIn_Main") == "SwUFn_0201"
        assert result.get("s_SystemOp") == "SwUFn_0301"

    def test_duplicate_function_name_first_match_wins(self):
        """동일 함수명 다중 매핑 — 첫 매칭 우선 + parse_warning 누적."""
        text = "SwUFn_0101: main\nSwUFn_0902: main"
        warns: list[str] = []
        result = extract_function_name_to_swufn_from_suds(_make_docx_with_text(text), warns)
        # 첫 매칭 우선
        assert result.get("main") == "SwUFn_0101"
        # warning 누적
        assert any("중복 매핑" in w for w in warns)


class TestDocxBoundaries:
    """공통 boundary — DOCX_MAX_BYTES, fail-safe."""

    def test_oversize_skip(self):
        from backend.services.iso26262_doc_asil_extractor import DOCX_MAX_BYTES
        warns: list[str] = []
        result = extract_function_asil_from_suds(b"x" * (DOCX_MAX_BYTES + 1), warns)
        assert result == {}
        assert any("한도 초과" in w for w in warns)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
