"""Phase 3 — v3.02 SwUDS 세로 key-value 표 ASIL 추출기 검증.

실 KJPDS02 SwUDS v3.02는 함수마다 '[ Function Information ]' 세로 표를 두고 행 라벨(col0)
'Name'/'ASIL'에 C 함수명·등급을 담는다(ASIL이 컬럼 헤더가 아니라 행 라벨 → 기존 파서 0건).
extract_function_asil_from_kv_tables는 lxml로 이를 추출한다. fail-closed(둘 다 있고 C 식별자·
유효등급일 때만)·max-merge(안전측)·python-docx 미사용을 검증한다.
"""
from __future__ import annotations

import io

from backend.services.iso26262_doc_asil_extractor import extract_function_asil_from_kv_tables


def _make_uds_docx(func_tables: list[dict[str, str]]) -> bytes:
    """func_tables: 함수별 {행라벨: 값} — 각각 세로 2열 표 하나로 stamp(실 v3.02 레이아웃 모사)."""
    from docx import Document  # backend/.venv에 존재(extractor도 사용)

    doc = Document()
    for ft in func_tables:
        t = doc.add_table(rows=0, cols=2)
        for label, value in ft.items():
            cells = t.add_row().cells
            cells[0].text = label
            cells[1].text = value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_kv_extract_basic():
    """Name/ASIL 행 라벨 세로 표 → {함수명(소문자): 등급}. 사용자 예시 g_drvin_drv8706sq_init 포함."""
    data = _make_uds_docx([
        {"ID": "SwUFn_0101", "Name": "g_drvin_drv8706sq_init", "Prototype": "void x(void)", "ASIL": "A"},
        {"ID": "SwUFn_0102", "Name": "main", "ASIL": "QM"},
    ])
    m = extract_function_asil_from_kv_tables(data)
    assert m == {"g_drvin_drv8706sq_init": "A", "main": "QM"}


def test_kv_extract_normalizes_asil_prefix():
    """값 셀이 'ASIL C'처럼 접두사 포함이어도 정규 등급 'C'로 정규화."""
    data = _make_uds_docx([{"Name": "s_foo", "ASIL": "ASIL C"}])
    assert extract_function_asil_from_kv_tables(data) == {"s_foo": "C"}


def test_kv_extract_requires_both_name_and_asil():
    """ASIL 행 없는 표 → 미추출(fail-closed — 프로즈 표/입력파라미터 표 오귀속 방지)."""
    data = _make_uds_docx([{"ID": "SwUFn_1", "Name": "s_foo", "Prototype": "void s_foo(void)"}])
    assert extract_function_asil_from_kv_tables(data) == {}


def test_kv_extract_rejects_prose_name():
    """Name이 다단어 프로즈면 C 식별자 아님 → 거부(오귀속·오등급 차단)."""
    data = _make_uds_docx([{"Name": "power operation disable", "ASIL": "B"}])
    assert extract_function_asil_from_kv_tables(data) == {}


def test_kv_extract_rejects_invalid_grade():
    """오타/범위밖 등급('B(잠정)')은 무효 → 거부(under-report 위험 차단, _norm_asil_grade)."""
    data = _make_uds_docx([{"Name": "s_foo", "ASIL": "B(잠정)"}])
    assert extract_function_asil_from_kv_tables(data) == {}


def test_kv_extract_max_merge_on_duplicate():
    """같은 함수명이 여러 표에 나오면 max 등급(안전측 — 낮은 등급 채택은 under-report)."""
    data = _make_uds_docx([
        {"Name": "s_dup", "ASIL": "A"},
        {"Name": "s_dup", "ASIL": "C"},
    ])
    assert extract_function_asil_from_kv_tables(data) == {"s_dup": "C"}


def test_kv_extract_fail_safe_on_bad_bytes():
    """빈/비-zip 바이트는 예외 없이 빈 맵(impact 본류 비차단)."""
    assert extract_function_asil_from_kv_tables(b"") == {}
    assert extract_function_asil_from_kv_tables(b"not a zip at all") == {}


def test_kv_extract_rejects_oversize_input():
    """96MB(DOCX_MAX_BYTES) 초과 입력은 파싱 없이 빈 맵 — OOM/과대작업 가드(reviewer W1)."""
    from backend.services.iso26262_doc_asil_extractor import DOCX_MAX_BYTES

    assert extract_function_asil_from_kv_tables(b"x" * (DOCX_MAX_BYTES + 1)) == {}


def test_kv_extract_rejects_non_ascii_name():
    """비-ASCII 이름(g_foo한)은 순수 C 식별자 아님 → 거부(reviewer I2, _C_IDENT_FULL_RE re.ASCII)."""
    data = _make_uds_docx([{"Name": "g_foo한", "ASIL": "B"}])
    assert extract_function_asil_from_kv_tables(data) == {}


def test_asil_rank_dicts_consistent_across_modules():
    """iso26262._ASIL_GRADE_RANK ≡ impact._ASIL_RANK — 두 모듈의 ASIL 순위 정의가 drift하면
    kv 파서 max-merge와 enrich max-merge가 불일치(잠재 오등급). 값·집합 동일성 강제(F4 drift 가드)."""
    from backend.services.iso26262_doc_asil_extractor import _ASIL_GRADE_RANK
    from workflow.impact_orchestrator import _ASIL_RANK

    assert _ASIL_GRADE_RANK == _ASIL_RANK


# ── SwCom(Related ID) 추출기 — 컴포넌트 ASIL 상속 폴백의 함수→SwCom 소스 ──

def _make_uds_with_related(func_rows):
    """func_rows: [{Name, ASIL?, Related ID}] — 각각 세로표(실 v3.02 레이아웃)."""
    from docx import Document

    doc = Document()
    for ft in func_rows:
        t = doc.add_table(rows=0, cols=2)
        for label, value in ft.items():
            c = t.add_row().cells
            c[0].text, c[1].text = label, value
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_swcom_extract_basic():
    """Name + Related ID → {함수명(소문자): [SwCom_NN]}. SwCom_NN만 뽑고 다른 ID(SwST/SwSTR) 무시."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_swcom_from_kv_tables

    data = _make_uds_with_related([
        {"ID": "SwUFn_1", "Name": "s_antipinch_x", "ASIL": "N/A",
         "Related ID": "SwST_01, SwCom_13, SwSTR_02"},
    ])
    assert extract_function_swcom_from_kv_tables(data) == {"s_antipinch_x": ["SwCom_13"]}


def test_swcom_extract_rejects_prose_name():
    """Name이 다단어 프로즈면 C 식별자 아님 → 제외(오귀속 차단)."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_swcom_from_kv_tables

    data = _make_uds_with_related([{"Name": "power operation disable", "Related ID": "SwCom_13"}])
    assert extract_function_swcom_from_kv_tables(data) == {}


def test_swcom_extract_unions_duplicate_names():
    """동명 함수가 여러 표에 나오면 SwCom 리스트 union(reviewer #1 — last-wins면 낮은등급만 남아
    enrich max가 최고등급 못 골라 under-report). 회귀 고정."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_swcom_from_kv_tables

    data = _make_uds_with_related([
        {"Name": "s_dup", "Related ID": "SwCom_50"},   # 표1
        {"Name": "s_dup", "Related ID": "SwCom_13"},   # 표2 — last-wins면 SwCom_13만 남음
    ])
    out = extract_function_swcom_from_kv_tables(data)
    assert out == {"s_dup": ["SwCom_13", "SwCom_50"]}   # union(정렬)


def test_kv_extract_details_description_and_prototype():
    """세로 kv 표 Name/Description/Prototype → {함수명(소문자): {description, prototype}}.

    링크(cloudium) v3.02 UDS를 python-docx(50MB에 41s) 대신 lxml(초 단위)로 뽑아 영향분석
    UDS 카드 실 내용 + 원문→변경안 기준선을 채운다. Name과 (desc 또는 proto) 둘 다 있어야 채택.
    """
    from backend.services.iso26262_doc_asil_extractor import extract_function_details_from_kv_tables
    data = _make_uds_docx([
        {"ID": "SwUFn_1596", "Name": "prv_FindBracketIdx",
         "Prototype": "static U16 prv_FindBracketIdx(const S16 val)",
         "Description": "배열에서 특정 값이 속한 인덱스를 찾아 반환하는 함수."},
        {"Name": "s_foo", "Description": "foo 설명"},  # proto 없이 desc만
    ])
    out = extract_function_details_from_kv_tables(data)
    assert out["prv_findbracketidx"]["description"] == "배열에서 특정 값이 속한 인덱스를 찾아 반환하는 함수."
    assert out["prv_findbracketidx"]["prototype"] == "static U16 prv_FindBracketIdx(const S16 val)"
    assert out["s_foo"]["description"] == "foo 설명"
    assert out["s_foo"]["prototype"] == ""


def test_kv_extract_details_korean_labels():
    """한글 라벨('함수원형'/'설명')도 매칭."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_details_from_kv_tables
    data = _make_uds_docx([{"Name": "s_init", "함수원형": "void s_init(void)", "설명": "초기화"}])
    out = extract_function_details_from_kv_tables(data)
    assert out["s_init"]["prototype"] == "void s_init(void)"
    assert out["s_init"]["description"] == "초기화"


def test_kv_extract_details_requires_name_and_content():
    """Name만 있고 desc/proto 없으면 미채택(입력파라미터 표 등 오귀속 방지)."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_details_from_kv_tables
    data = _make_uds_docx([{"ID": "SwUFn_1", "Name": "s_foo", "ASIL": "A"}])
    assert extract_function_details_from_kv_tables(data) == {}


def test_kv_extract_details_rejects_prose_name():
    """Name이 다단어 프로즈면 C 식별자 아님 → 거부."""
    from backend.services.iso26262_doc_asil_extractor import extract_function_details_from_kv_tables
    data = _make_uds_docx([{"Name": "power operation disable", "Description": "설명"}])
    assert extract_function_details_from_kv_tables(data) == {}
