"""Verify that the tokenized UDS template + docx_builder substitution produces
the expected 1장(Introduction) body for an arbitrary project payload.

Runs _replace_docx_text() and _build_uds_reference_text() directly on a
copy of docs/(HDPDM01_SUDS)_template_tokenized.docx and prints the resulting
1.1~1.4 paragraphs so the user can eyeball the substitution output.
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import docx

from report_gen.docx_builder import (
    _build_uds_reference_lines,
    _build_uds_reference_text,
    _replace_docx_text,
    _replace_reference_table_paragraph,
)


TEMPLATE = ROOT / "docs" / "(HDPDM01_SUDS)_template_tokenized.docx"
OUT = ROOT / "docs" / "_uds_token_verify_sample.docx"


def run(project_name: str, module_name: str | None, payload_extra: dict) -> None:
    payload = {"project_name": project_name, **payload_extra}
    if module_name is not None:
        payload["module_name"] = module_name
    # Replicate docx_builder's resolution rule so the test reflects runtime.
    effective_module = (
        payload.get("module_name")
        or payload.get("module")
        or project_name
    )
    module_name = str(effective_module)
    ref_lines = _build_uds_reference_lines(payload)
    ref_text = _build_uds_reference_text(payload)
    replacements = {
        "{{project_name}}": project_name,
        "{{PROJECT_NAME}}": project_name,
        "{{MODULE_NAME}}": module_name,
        "{{REFERENCE_TABLE}}": ref_text,
    }

    doc = docx.Document(str(TEMPLATE))
    # Same order as generate_uds_docx: paragraph-aware reference list first,
    # then generic flat-text token substitution.
    replaced = _replace_reference_table_paragraph(doc, ref_lines)
    _replace_docx_text(doc, replacements)
    doc.save(str(OUT))

    out_doc = docx.Document(str(OUT))
    print(f"\n=== PROJECT={project_name}  MODULE={module_name} ===")
    print(f"reference lines ({len(ref_lines)}): {ref_lines}")
    print(f"reference_table paragraph replaced: {replaced}")

    # Count w:br tags on paragraphs containing reference bracket markers.
    br_count = 0
    for p in out_doc.paragraphs:
        if "[1]" in p.text and "]" in p.text:
            xml = p._element.xml
            br_count = xml.count("<w:br")
            break
    print(f"soft line breaks (w:br) in reference paragraph: {br_count}")
    capturing = False
    for i, p in enumerate(out_doc.paragraphs[:230]):
        style = (p.style.name or "")
        text = p.text.strip()
        if style.startswith("Heading"):
            if text in (
                "Introduction",
                "Purpose",
                "Scope",
                "Terms, Abbreviations and Definitions",
                "Reference",
            ):
                capturing = True
            elif text == "Software Unit Design":
                break
        if capturing and text:
            print(f"  [{i}] <{style}> {text[:140]}")


if __name__ == "__main__":
    run(
        "KJPDS02_DV",
        "Door Control",
        {
            "srs_path": r"D:/prj/KJPDS02/(KJPDS02_SRS) Software Requirements_v1.00.docx",
            "sds_path": r"D:/prj/KJPDS02/(KJPDS02_SDS) Software Architecture_v1.00.docx",
            "hsis_path": r"D:/prj/KJPDS02/(KJPDS02_HSIS) Hardware Software Interface_v1.00.xlsx",
        },
    )
    run(
        "HDPDM01",
        "BSD",
        {
            "reference_docs": [
                {"title": "Stakeholder Requirements", "version": "2025"},
                {"title": "Software Requirements Specification", "version": "1.05"},
            ],
        },
    )
    # Pipeline realistic case: only `source_docs` populated (what
    # _uds_generate_from_paths now injects via req_file_paths).
    run(
        "KJPDS02_PV",
        "Body Control",
        {
            "source_docs": [
                r"D:/prj/KJPDS02/(KJPDS02_SRS) Software Requirements_v1.02.docx",
                r"D:/prj/KJPDS02/(KJPDS02_SRS) Safety Requirements_v1.00.docx",
            ],
        },
    )
    # module_name omitted entirely → docx_builder must fall back to
    # project_name (no more hardcoded "BSD").
    run("HDPDM_MODULE_FALLBACK", None, {})
    print(f"\nsaved → {OUT}")
