"""P3 verification: check whether Introduction subsection spacing (blank
paragraphs between sibling Heading 2 sections) is preserved in the
rebuild path output."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document


def inspect(path: str) -> None:
    d = Document(path)
    start = None
    for i, p in enumerate(d.paragraphs):
        style = str(getattr(p.style, "name", "") or "")
        text = (p.text or "").strip()
        if style.startswith("Heading") and text.lower() == "introduction":
            start = i
            break
    if start is None:
        print("Introduction not found")
        return

    print(f"=== {Path(path).name} ===")
    print(f"Introduction idx: {start}")
    blank_pairs = []
    prev_heading = None
    for j in range(start, min(start + 30, len(d.paragraphs))):
        p = d.paragraphs[j]
        text = (p.text or "").strip()
        style = str(getattr(p.style, "name", "") or "")
        marker = "(empty)" if not text else repr(text)[:100]
        print(f"[{j}] <{style}> {marker}")
        if not text and prev_heading:
            blank_pairs.append((prev_heading, j))
        if text and style.startswith("Heading"):
            prev_heading = (j, text, style)

    print()
    print(f"Blank paragraphs in Introduction zone: "
          f"{sum(1 for j in range(start, min(start + 30, len(d.paragraphs))) if not (d.paragraphs[j].text or '').strip())}")


if __name__ == "__main__":
    for p in sys.argv[1:]:
        inspect(p)
        print()
