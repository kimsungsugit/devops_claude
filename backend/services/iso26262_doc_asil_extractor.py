"""ISO 26262 추적성 체인 문서 ASIL extractor (라운드 80 T1405-2).

CLAUDE.md ASIL 탐지 기준 #2 "SRS/SDS 문서의 안전 요구사항 매핑" 구현.
SUDS는 ``swut_swuds_parser`` 가 함수 단위 매핑 담당. 본 모듈은:

- **SDS** (Software Design Specification) docx에서 **컴포넌트 단위** ASIL 추출
  (`SwCom_01: A`, `SwCom_02: A` 형식 표). 라운드 77 ``FunctionCoverage.component_name``
  필드와 매칭.
- **SRS** (Software Requirements Specification) docx에서 **함수명 보조** 추출
  (`g_*` / `s_*` / `u_*` 함수명 ↔ ASIL regex).

## ISO 26262 Tool Qualification

- 본 extractor 결과는 SUDS 미발견 함수에 한해 fallback. SUDS / c_source 우선.
- ASIL A: reviewer 검토 후 evidence 사용 가능.
- ASIL B/C/D: manual 검증 후 evidence 확정 의무.

## 양식 가정

- SDS: 컴포넌트 정의 표에 'ASIL' header col 존재. 같은 row의 다른 col에 컴포넌트명.
- SRS: 본문 또는 표 cell에 함수명 + ASIL 등급 인접 (200자 내).
"""
from __future__ import annotations

import io
import re
import zipfile

try:
    from docx import Document  # type: ignore
    _HAS_DOCX = True
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[assignment]
    _HAS_DOCX = False

# 라운드 80 T1405-2: regex pair 근접 거리 (100자 ~ 1500자).
# - SRS/표 단위: 100자 (보수적, false positive 방지)
# - SUDS section 단위: 1500자 (실 양식 SwUFn~ASIL 거리 500자~1000자 측정 결과)
_ASIL_PROXIMITY_TIGHT = 100
_ASIL_PROXIMITY_SECTION = 1500
_REGEX_FN_ASIL = re.compile(
    r"\b(g_\w+|s_\w+|u_\w+|SwUFn_\d+)[\s\S]{{0,{n}}}?ASIL[\s_-]*([ABCD]|QM)\b".format(
        n=_ASIL_PROXIMITY_TIGHT,
    ),
    re.IGNORECASE,
)
_REGEX_COMPONENT_PREFIX = re.compile(r"SwCom_\d+|Sw\s*Com\s*\d+", re.IGNORECASE)
_REGEX_SWUFN_ONLY = re.compile(r"SwUFn_\d+")
_REGEX_ASIL_ONLY = re.compile(r"ASIL[\s_-]*([ABCD]|QM)\b", re.IGNORECASE)

DOCX_MAX_BYTES = 96 * 1024 * 1024  # 라운드 87 T2103: 96MB (SUDS v1.10+ 대비 마진, swut_swuds_parser와 동일)


def _load_doc(docx_bytes: bytes, warnings: list[str]):
    """docx bytes 로드 — fail-safe."""
    if not _HAS_DOCX:
        warnings.append("python-docx 미설치 — ISO 26262 doc ASIL extractor skip")
        return None
    if not docx_bytes:
        warnings.append("docx bytes 비어있음 — extractor skip")
        return None
    if len(docx_bytes) > DOCX_MAX_BYTES:
        warnings.append(
            f"docx {len(docx_bytes):,} bytes > {DOCX_MAX_BYTES:,} 한도 초과 — skip"
        )
        return None
    try:
        return Document(io.BytesIO(docx_bytes))
    except Exception as e:
        warnings.append(f"docx 로드 실패: {type(e).__name__}: {e}")
        return None


# ── v3.02 SwUDS 세로 key-value 함수 표 ASIL 추출 ──
# 실 KJPDS02 SwUDS v3.02는 함수마다 '[ Function Information ]' 세로 표를 두고 행 라벨(col0)
# 'Name'/'ASIL'에 실제 C 함수명·등급을 담는다(ASIL이 컬럼 헤더가 아니라 행 라벨). 기존
# swut_swuds_parser(SwUFn heading 직후 표의 'ASIL' 라벨셀 옆값)·iso26262 reverse-corpus는
# 이 레이아웃에서 0건이었다. 또한 document.xml이 87MB라 python-docx 전체 로드가 ~400s.
# 여기선 lxml iterparse로 표만 스트리밍(초 단위)해 (Name, ASIL) 쌍만 뽑는다.
_WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_ASIL_GRADE_RANK = {"QM": 0, "A": 1, "B": 2, "C": 3, "D": 4}
# re.ASCII: \w를 ASCII로 고정해 'g_foo한' 같은 비-ASCII 이름을 거부(reviewer I2 — "순수 C 식별자" 보장).
_C_IDENT_FULL_RE = re.compile(r"[A-Za-z_]\w{2,}\Z", re.ASCII)
# 압축폭탄 방지 — document.xml 비압축 크기 상한(실 v3.02=88MB, ~4.3× 마진, reviewer W1).
_KV_XML_MAX_BYTES = 384 * 1024 * 1024


def _norm_asil_grade(v) -> str:
    """'A'/'ASIL B'/'QM' → 정규 등급('A'/'B'/'QM'), 무효(오타·범위밖)면 ''."""
    a = re.sub(r"^ASIL[\s_-]*", "", str(v or "").strip().upper()).strip()
    return a if a in _ASIL_GRADE_RANK else ""


def _read_suds_document_xml(docx_bytes: bytes) -> bytes | None:
    """SwUDS docx bytes에서 word/document.xml을 크기 가드와 함께 꺼낸다(없으면 None).

    W1: 입력 96MB(DOCX_MAX_BYTES)·비압축 384MB(_KV_XML_MAX_BYTES) 상한으로 OOM/압축폭탄 방어.
    kv-table ASIL 추출기와 SwCom(Related ID) 추출기가 공유(둘 다 세로표 세션 대상).
    """
    if not docx_bytes or len(docx_bytes) > DOCX_MAX_BYTES:
        return None
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            try:
                _xi = z.getinfo("word/document.xml")
            except KeyError:
                return None   # docx 아님/구조 이상 → 폴백에 위임
            if _xi.file_size > _KV_XML_MAX_BYTES:
                return None   # 압축폭탄(비압축 폭증) 방어
            return z.read("word/document.xml")
    except Exception:
        return None


def extract_function_asil_from_kv_tables(docx_bytes: bytes) -> dict[str, str]:
    """세로 key-value SwUDS 함수 표에서 {함수명(소문자): ASIL} 추출.

    표당 행 라벨(col0)이 'Name'인 행의 값(col1)=C 함수명, 'ASIL'인 행의 값=등급.
    둘 다 있고 함수명이 순수 C 식별자·등급이 유효할 때만 채택(fail-closed — 프로즈/다단어/
    무효등급 거부해 under-report 위험 차단). 같은 함수명 충돌은 max 등급(안전측). lxml 직접
    파싱(python-docx 미사용 → 87MB docx도 초 단위). 접근/파싱 실패는 빈 맵(비차단).
    """
    xml = _read_suds_document_xml(docx_bytes)
    if xml is None:
        return {}
    try:
        from lxml import etree  # type: ignore
    except Exception:
        return {}
    w = "{%s}" % _WML_NS

    def _cell_text(tc) -> str:
        return "".join(t.text or "" for t in tc.iter(w + "t")).strip()

    out: dict[str, str] = {}
    try:
        for _ev, tbl in etree.iterparse(io.BytesIO(xml), events=("end",), tag=w + "tbl"):
            name = ""
            grade = ""
            for tr in tbl.findall(w + "tr"):
                cells = tr.findall(w + "tc")
                if len(cells) < 2:
                    continue
                label = _cell_text(cells[0]).lower().rstrip(":").strip()
                if label == "name" and not name:
                    name = _cell_text(cells[1])
                elif label == "asil" and not grade:
                    grade = _norm_asil_grade(_cell_text(cells[1]))
            if name and grade:
                nl = name.strip().lower()
                if _C_IDENT_FULL_RE.match(nl):
                    ex = out.get(nl)
                    if ex is None or _ASIL_GRADE_RANK[grade] > _ASIL_GRADE_RANK.get(ex, -1):
                        out[nl] = grade
            tbl.clear()  # 처리 후 즉시 비워 메모리 바운드(대용량 document.xml)
    except Exception:
        return out  # 부분 파싱분 보존(비차단)
    return out


# SwUDS `Related ID` 칸의 요소 ID 토큰. 접두 길이를 1~5 로 묶어 `Software_12` 같은
# 산문 토큰이 딸려오는 것을 막는다(실 어휘: SwCom·SwFn·SwSTR·SwST·SwTK·SwUFn — 최대 3).
_REGEX_RELATED_ID = re.compile(r"\bSw[A-Za-z]{1,5}_\d+\b", re.IGNORECASE)
# 접두 정규 케이스 — 문서 표기 흔들림(`SWCOM_13`)을 한 형태로 모아 하류 정확매칭을 지킨다.
# ⚠ 키 **전체**를 대문자로 비교한다. 부분문자열로 판정하면 `SwSTR_01` 이 `SwST` 로 접혀
#   서로 다른 요소가 한 ID 로 뭉개진다(이 저장소가 SUTS 에서만 4번 겪은 실패 형태다).
_RELATED_PREFIX_CANON = {
    "SWCOM": "SwCom", "SWFN": "SwFn", "SWSTR": "SwSTR",
    "SWST": "SwST", "SWTK": "SwTK", "SWUFN": "SwUFn",
}


def _canon_related_token(tok: str) -> str:
    """`SWCOM_13` → `SwCom_13`. 모르는 접두는 **원문 그대로** 둔다(버리지 않는다)."""
    pre, _sep, num = str(tok or "").partition("_")
    if not num:
        return ""
    return f"{_RELATED_PREFIX_CANON.get(pre.upper(), pre)}_{num}"

# 세로 kv 표 Description/Prototype 행 라벨 후보(swut_swuds_parser의 가로표 후보와 동일 원칙).
_KV_DESC_LABELS = (
    "description", "기능 설명", "설명", "기능설명", "function description", "요약", "개요",
)
_KV_PROTO_LABELS = (
    "prototype", "프로토타입", "함수원형", "함수 원형", "function prototype", "함수 프로토타입",
)


def extract_function_details_from_kv_tables(docx_bytes: bytes) -> dict[str, dict[str, str]]:
    """세로 key-value SwUDS 함수 표에서 {함수명(소문자): {description, prototype}} 추출.

    표당 행 라벨(col0) 'Name'=C 함수명, 'Description'/'Prototype'(한글 변종 포함)=값.
    lxml iterparse로 표만 스트리밍(python-docx 미사용 → 50~87MB docx도 초 단위) — v3.02류
    세로표 레이아웃. `_load_uds_fn_content` fallback이 python-docx `parse_swuds_docx`(50MB에
    22~41s)를 쓰던 것을 대체해 속도·신뢰성(중단 시 빈 캐시 굳는 위험) 개선.
    ASIL 추출기(`extract_function_asil_from_kv_tables`)와 동일 스캔·메모리 가드. 실패는 빈 맵.
    """
    xml = _read_suds_document_xml(docx_bytes)
    if xml is None:
        return {}
    try:
        from lxml import etree  # type: ignore
    except Exception:  # silent-ok — lxml 미설치 시 빈 맵(비차단, ASIL 추출기와 동일 옵셔널 import 가드)
        return {}
    w = "{%s}" % _WML_NS

    def _cell_text(tc) -> str:
        return "".join(t.text or "" for t in tc.iter(w + "t")).strip()

    out: dict[str, dict[str, str]] = {}
    try:
        for _ev, tbl in etree.iterparse(io.BytesIO(xml), events=("end",), tag=w + "tbl"):
            name = ""
            desc = ""
            proto = ""
            for tr in tbl.findall(w + "tr"):
                cells = tr.findall(w + "tc")
                if len(cells) < 2:
                    continue
                label = _cell_text(cells[0]).lower().rstrip(":").strip()
                if label == "name" and not name:
                    name = _cell_text(cells[1])
                elif label in _KV_DESC_LABELS and not desc:
                    desc = _cell_text(cells[1])
                elif label in _KV_PROTO_LABELS and not proto:
                    proto = _cell_text(cells[1])
            if name and (desc or proto):
                nl = name.strip().lower()
                if _C_IDENT_FULL_RE.match(nl) and nl not in out:
                    out[nl] = {"description": desc[:500], "prototype": proto[:200]}
            tbl.clear()  # 처리 후 즉시 비워 메모리 바운드(대용량 document.xml)
    except Exception:  # silent-ok — 부분 파싱분 보존(비차단, ASIL 추출기와 동일 fail-safe)
        return out
    return out


def extract_function_related_ids_from_kv_tables(docx_bytes: bytes) -> dict[str, list[str]]:
    """세로 kv SwUDS 함수 표에서 {함수명(소문자): [Related ID 토큰 전체]} 추출.

    `Related ID` 칸은 **SwCom 만 담고 있지 않다**. 정본 대조(KJPDS02_PV SwUDS v3.02) 실측:

        SwCom 1,052 · SwFn 135 · SwSTR 64 · SwST 44 · SwTK 3   (총 1,298 토큰 / 1,026 함수)

    즉 `SwCom_` 으로 좁히면 **19%(246 토큰)를 통째로 버린다**. SITS `Related ID` 칸의
    정본 어휘가 바로 이 다섯 종이고(SwCom 170 · SwFn 69 · SwSTR 62 · SwST 38 · SwTK 8),
    좁힌 탓에 그 칸의 원소 재현율이 18.8% 에 머물렀다(넓히면 41.8% — 과잉은 4건뿐).

    함수명은 순수 C 식별자만(fail-closed). Related ID 가 없으면 그 함수는 결과에서 제외.
    동명 함수가 여러 표에 나오면 union(SwCom 병합 철학과 동일 — 안전측).
    접근/파싱 실패는 빈 맵(비차단).
    """
    xml = _read_suds_document_xml(docx_bytes)
    if xml is None:
        return {}
    try:
        from lxml import etree  # type: ignore
    except Exception:  # silent-ok — lxml 미설치 시 빈 맵(비차단, 형제 추출기와 동일 가드)
        return {}
    w = "{%s}" % _WML_NS

    def _cell_text(tc) -> str:
        return "".join(t.text or "" for t in tc.iter(w + "t")).strip()

    out: dict[str, list[str]] = {}
    try:
        for _ev, tbl in etree.iterparse(io.BytesIO(xml), events=("end",), tag=w + "tbl"):
            name = ""
            related = ""
            for tr in tbl.findall(w + "tr"):
                cells = tr.findall(w + "tc")
                if len(cells) < 2:
                    continue
                label = _cell_text(cells[0]).lower().rstrip(":").strip()
                if label == "name" and not name:
                    name = _cell_text(cells[1])
                elif label in ("related id", "relatedid") and not related:
                    related = _cell_text(cells[1])
            if name and related:
                nl = name.strip().lower()
                if _C_IDENT_FULL_RE.match(nl):
                    toks = {t for t in (_canon_related_token(m)
                                        for m in _REGEX_RELATED_ID.findall(related)) if t}
                    if toks:
                        out[nl] = sorted(set(out.get(nl, [])) | toks)
            tbl.clear()  # 처리 후 즉시 비워 메모리 바운드(대용량 document.xml)
    except Exception:  # silent-ok — 부분 파싱분 보존(비차단, 형제 추출기와 동일 fail-safe)
        return out
    return out


def extract_function_swcom_from_kv_tables(docx_bytes: bytes) -> dict[str, list[str]]:
    """세로 kv SwUDS 함수 표에서 {함수명(소문자): [SwCom_NN, ...]} 추출.

    함수 ASIL이 N/A(미등급)인데 소속 SwCom은 SDS에 등급이 있을 때, **컴포넌트 ASIL 상속**
    폴백에 쓴다(ISO 26262 표준 원칙 — 함수는 소속 SW 컴포넌트의 ASIL을 상속). SDS 컴포넌트
    맵의 키가 `SwCom_NN` 이라 이 축은 SwCom 만 필요하다.

    ⚠ 스캔은 `extract_function_related_ids_from_kv_tables` **하나**가 하고 여기서는 거르기만
    한다. 예전엔 같은 iterparse 루프를 따로 들고 있었는데, 그러면 라벨 변종(`RelatedID` 등)
    대응이 한쪽에만 반영돼 두 축의 답이 갈라진다(이 저장소가 ruff/eslint ratchet 에서 이미
    겪은 복제 결함과 같은 형태다).
    """
    out: dict[str, list[str]] = {}
    for name, toks in extract_function_related_ids_from_kv_tables(docx_bytes).items():
        coms = sorted(t for t in toks if t.startswith("SwCom_"))
        if coms:
            out[name] = coms
    return out


def extract_component_asil_from_sds(
    docx_bytes: bytes, warnings: list[str] | None = None,
) -> dict[str, str]:
    """SDS docx에서 컴포넌트별 ASIL 추출.

    회사 양식 (HDPDM01 SDS v1.04 검증): 표에 'ASIL' header col + 컴포넌트명 col.
    같은 row에서 두 값 매칭 → dict[component_name, asil_letter] 반환.

    Args:
        docx_bytes: SDS docx raw bytes.
        warnings: 외부 누적 list (router warnings 와 통합).

    Returns:
        ``{"SwCom_01": "A", "System OS": "A", ...}`` — 컴포넌트명 다양한 키로 저장
        (SwCom_NN prefix + 별칭 텍스트 모두). 매칭 0건 시 빈 dict.
    """
    warns = warnings if warnings is not None else []
    doc = _load_doc(docx_bytes, warns)
    if doc is None:
        return {}

    try:
        from backend.services.swut_asil_resolver import _normalize_asil
    except ImportError:  # pragma: no cover
        warns.append("swut_asil_resolver import 실패 — SDS extractor skip")
        return {}

    result: dict[str, str] = {}
    table_count = 0
    asil_table_count = 0

    for tbl in doc.tables:
        table_count += 1
        if not tbl.rows:
            continue
        header = [cell.text.strip().upper() for cell in tbl.rows[0].cells]
        asil_col = next((i for i, h in enumerate(header) if "ASIL" in h), None)
        if asil_col is None:
            continue
        asil_table_count += 1
        for row in tbl.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= asil_col:
                continue
            asil_letter = _normalize_asil(cells[asil_col])
            if not asil_letter:
                continue
            # 컴포넌트 이름 후보 — SwCom_NN prefix + 다른 col text
            for idx, val in enumerate(cells):
                if idx == asil_col or not val:
                    continue
                # SwCom_NN 매칭
                m_comp = _REGEX_COMPONENT_PREFIX.search(val)
                if m_comp:
                    key = m_comp.group(0).replace(" ", "")
                    if key not in result:
                        result[key] = asil_letter
                # 별칭 (System OS / DRV IN 등) — 짧은 컴포넌트명만 등록
                v = val.strip()
                if v and len(v) <= 50 and "\n" not in v and v not in result:
                    result[v] = asil_letter

    if not result:
        warns.append(
            f"SDS docx ASIL 컴포넌트 매핑 0건 "
            f"(표 {table_count}개 중 ASIL header {asil_table_count}건)"
        )
    return result


def extract_supplementary_asil_from_srs(
    docx_bytes: bytes, warnings: list[str] | None = None,
) -> dict[str, str]:
    """SRS docx에서 함수명 보조 ASIL 추출.

    SRS는 함수 단위 ASIL 매핑이 표준이 아님 — 본문/표 텍스트에 함수명 ↔ ASIL
    인접 패턴이 산발적. 보수적으로 100자 근접 regex 추출.

    Args:
        docx_bytes: SRS docx raw bytes.
        warnings: 외부 누적 list.

    Returns:
        ``{"g_DrvIn_Main": "A", "s_SystemOperation": "A", ...}`` — 함수명 키.
        매칭 0건 시 빈 dict.
    """
    warns = warnings if warnings is not None else []
    doc = _load_doc(docx_bytes, warns)
    if doc is None:
        return {}

    try:
        from backend.services.swut_asil_resolver import _normalize_asil
    except ImportError:  # pragma: no cover
        warns.append("swut_asil_resolver import 실패 — SRS extractor skip")
        return {}

    # 전체 corpus — paragraph + table cell
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    corpus = "\n".join(parts)

    result: dict[str, str] = {}
    for m in _REGEX_FN_ASIL.finditer(corpus):
        fn_name = m.group(1)
        # 한글 조사 자동 strip (예: g_ApiIn_MotorPosition과 → g_ApiIn_MotorPosition)
        fn_name = re.sub(r"[가-힣]+$", "", fn_name).strip("._")
        if not fn_name:
            continue
        asil_letter = _normalize_asil(m.group(2))
        if not asil_letter or fn_name in result:
            continue
        result[fn_name] = asil_letter

    if not result:
        warns.append("SRS docx 함수 보조 ASIL 매핑 0건")
    return result


def extract_function_asil_from_suds(
    docx_bytes: bytes, warnings: list[str] | None = None,
) -> dict[str, str]:
    """SUDS docx에서 함수 단위 ASIL 직접 추출 (heading 의존 X).

    ``swut_swuds_parser.parse_swuds_docx`` 는 'SwUFn_NNNN heading + 다음 table'
    Hyundai 양식 가정. v1.07 양식은 heading 부재 → entries 빈 → ASIL 매핑 0.
    본 함수는 본문 전체 corpus에서 **역방향 매칭** — 각 ASIL 라벨의 직전 SwUFn
    채택 (1500자 이내). SUDS section 양식 SwUFn~ASIL 거리 500~1000자 측정 결과
    반영. 라이브 검증 SUDS v1.07 = 415건 ASIL 라벨 → 약 380건 매핑.

    Args:
        docx_bytes: SUDS docx raw bytes.
        warnings: 외부 누적 list.

    Returns:
        ``{"SwUFn_0101": "A", "SwUFn_0102": "QM", ...}`` 매핑. 매칭 0건 시 빈 dict.
    """
    warns = warnings if warnings is not None else []
    doc = _load_doc(docx_bytes, warns)
    if doc is None:
        return {}

    try:
        from backend.services.swut_asil_resolver import _normalize_asil
    except ImportError:  # pragma: no cover
        warns.append("swut_asil_resolver import 실패 — SUDS extractor skip")
        return {}

    # 전체 corpus — paragraph + table cell
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    corpus = "\n".join(parts)

    # 역방향 매칭: ASIL 라벨 각각의 직전 SwUFn 찾기 (1500자 내).
    # 1) SwUFn 위치 모두 수집 (정렬 — finditer는 본래 순서 보장)
    swufn_positions: list[tuple[int, str]] = [
        (m.start(), m.group(0)) for m in _REGEX_SWUFN_ONLY.finditer(corpus)
    ]
    # 2) ASIL 라벨 각각의 직전 SwUFn 찾기 — pointer 1-pass (O(N+M))
    result: dict[str, str] = {}
    swufn_idx = 0
    last_swufn: tuple[int, str] | None = None
    for am in _REGEX_ASIL_ONLY.finditer(corpus):
        asil_pos = am.start()
        asil_letter = _normalize_asil(am.group(1))
        if not asil_letter:
            continue
        # 직전 SwUFn 갱신 — asil_pos 이하의 마지막 swufn으로 진행
        while swufn_idx < len(swufn_positions) and swufn_positions[swufn_idx][0] <= asil_pos:
            last_swufn = swufn_positions[swufn_idx]
            swufn_idx += 1
        if last_swufn is None:
            continue
        # 거리 체크
        dist = asil_pos - last_swufn[0]
        if dist > _ASIL_PROXIMITY_SECTION:
            continue
        fn_id = last_swufn[1]
        if fn_id in result:
            continue
        result[fn_id] = asil_letter

    if not result:
        warns.append("SUDS docx 함수 ASIL 매핑 0건 (역방향)")
    return result


_REGEX_SWUFN_TO_NAME = re.compile(r"(SwUFn_\d+)[:\s]+([a-zA-Z_]\w+)")


def extract_function_name_to_swufn_from_suds(
    docx_bytes: bytes, warnings: list[str] | None = None,
) -> dict[str, str]:
    """라운드 85 T1901: SUDS docx에서 함수명 ↔ SwUFn_NNNN reverse map 추출.

    SUDS 본문 'SwUFn_0101: main', 'SwUFn_0201: g_DrvIn_Main' 형식 (Hyundai/Mobis
    양식). 라이브 진단 v1.07 = 1706건 pair (unique 440).

    Args:
        docx_bytes: SUDS docx raw bytes.
        warnings: 외부 누적 list.

    Returns:
        ``{"main": "SwUFn_0101", "g_DrvIn_Main": "SwUFn_0201", ...}``
        함수명 keyed reverse map. 첫 매칭 우선 (중복 시 후속 entry는 parse_warnings).
        매칭 0건 시 빈 dict.

    카드: vcast 추출 fc.unit_id/fc.name (함수명) → 본 map → SwUFn_NNNN →
    function_asil_from_suds (ASIL 등급) chain 완성.
    """
    warns = warnings if warnings is not None else []
    doc = _load_doc(docx_bytes, warns)
    if doc is None:
        return {}

    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text or "")
    corpus = "\n".join(parts)

    result: dict[str, str] = {}
    duplicates: list[str] = []
    for m in _REGEX_SWUFN_TO_NAME.finditer(corpus):
        sw_fn_id = m.group(1)
        fn_name = m.group(2)
        if fn_name in result:
            if result[fn_name] != sw_fn_id and len(duplicates) < 10:
                duplicates.append(f"{fn_name}: {result[fn_name]} vs {sw_fn_id}")
            continue
        result[fn_name] = sw_fn_id

    if duplicates:
        warns.append(
            f"SUDS 함수명↔SwUFn 중복 매핑 {len(duplicates)}건 (첫 매칭 우선): "
            f"{', '.join(duplicates[:5])}"
        )
    if not result:
        warns.append("SUDS docx 함수명↔SwUFn reverse map 0건")
    return result


__all__ = [
    "extract_function_asil_from_suds",
    "extract_function_name_to_swufn_from_suds",
    "extract_component_asil_from_sds",
    "extract_function_asil_from_kv_tables",
    "extract_function_swcom_from_kv_tables",
    "extract_supplementary_asil_from_srs",
    "DOCX_MAX_BYTES",
]
