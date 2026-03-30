from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple

TOP_N_DEFAULT = 20

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None

try:
    from openpyxl import Workbook  # type: ignore
except Exception:  # pragma: no cover
    Workbook = None


def _fmt(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, bool):
        return "Y" if value else "N"
    return str(value)


def _issue_count(section: Dict[str, Any]) -> int:
    try:
        data = section.get("data", {}) if isinstance(section, dict) else {}
        issues = data.get("issues", [])
        return len(issues) if isinstance(issues, list) else 0
    except Exception:
        return 0


def _build_summary_rows(summary: Dict[str, Any]) -> List[Tuple[str, str, str]]:
    tests = summary.get("tests", {}) if isinstance(summary.get("tests"), dict) else {}
    coverage = summary.get("coverage", {}) if isinstance(summary.get("coverage"), dict) else {}
    static = summary.get("static", {}) if isinstance(summary.get("static"), dict) else {}

    rows: List[Tuple[str, str, str]] = [
        ("meta", "generated_at", _fmt(summary.get("generated_at"))),
        ("meta", "project_root", _fmt(summary.get("project_root"))),
        ("meta", "exit_code", _fmt(summary.get("exit_code"))),
        ("meta", "failure_stage", _fmt(summary.get("failure_stage"))),
        ("meta", "change_mode", _fmt(summary.get("change_mode"))),
    ]

    rows += [
        ("tests", "enabled", _fmt(tests.get("enabled"))),
        ("tests", "mode", _fmt(tests.get("mode") or tests.get("generation_mode"))),
        ("tests", "total", _fmt(tests.get("total") or tests.get("count") or tests.get("tests_total"))),
        ("tests", "passed", _fmt(tests.get("passed") or tests.get("passed_count") or tests.get("tests_exec_passed"))),
        ("tests", "failed", _fmt(tests.get("failed") or tests.get("failed_count") or tests.get("tests_exec_failed"))),
        ("tests", "compile_failed", _fmt(tests.get("tests_compile_failed"))),
        ("tests", "generated_count", _fmt(tests.get("generated_count"))),
        ("tests", "ok_count", _fmt(tests.get("ok_count"))),
        ("tests", "failed_count", _fmt(tests.get("failed_count"))),
        ("tests", "compile_failed_count", _fmt(tests.get("compile_failed_count"))),
        ("tests", "syntax_failed_count", _fmt(tests.get("syntax_failed_count"))),
        ("tests", "missing_main_count", _fmt(tests.get("missing_main_count"))),
        ("tests", "invalid_output_count", _fmt(tests.get("invalid_output_count"))),
        ("tests", "plan_ok_count", _fmt(tests.get("plan_ok_count"))),
    ]
    execution = tests.get("execution", {}) if isinstance(tests.get("execution"), dict) else {}
    rows += [
        ("tests_exec", "ok", _fmt(execution.get("ok"))),
        ("tests_exec", "count", _fmt(execution.get("count"))),
        ("tests_exec", "passed", _fmt(execution.get("passed"))),
        ("tests_exec", "failed", _fmt(execution.get("failed"))),
        ("tests_exec", "note", _fmt(execution.get("note"))),
    ]

    rows += [
        ("coverage", "enabled", _fmt(coverage.get("enabled"))),
        ("coverage", "ok", _fmt(coverage.get("ok"))),
        ("coverage", "line_rate_pct", _fmt(coverage.get("line_rate_pct") or coverage.get("line_rate"))),
        ("coverage", "branch_rate_pct", _fmt(coverage.get("branch_rate_pct") or coverage.get("branch_rate"))),
        ("coverage", "threshold", _fmt(coverage.get("threshold"))),
        ("coverage", "reason", _fmt(coverage.get("reason") or coverage.get("parse_error"))),
    ]

    rows += [
        ("static", "cppcheck_issues", _fmt(_issue_count(static.get("cppcheck", {}) if isinstance(static, dict) else {}))),
        ("static", "clang_tidy_issues", _fmt(_issue_count(static.get("clang_tidy", {}) if isinstance(static, dict) else {}))),
        ("static", "semgrep_issues", _fmt(_issue_count(static.get("semgrep", {}) if isinstance(static, dict) else {}))),
    ]

    return rows


def _severity_rank(value: str) -> int:
    raw = str(value or "").lower()
    if raw in ("error", "fatal", "critical", "high"):
        return 0
    if raw in ("warning", "warn", "medium"):
        return 1
    return 2


def _collect_static_issues(summary: Dict[str, Any], top_n: int) -> List[Dict[str, Any]]:
    static = summary.get("static", {}) if isinstance(summary.get("static"), dict) else {}
    rows: List[Dict[str, Any]] = []
    for tool in ("cppcheck", "clang_tidy", "semgrep"):
        block = static.get(tool, {})
        data = block.get("data", {}) if isinstance(block, dict) else {}
        issues = data.get("issues", [])
        if not isinstance(issues, list):
            continue
        for it in issues:
            if not isinstance(it, dict):
                continue
            rows.append(
                {
                    "tool": tool,
                    "severity": it.get("severity") or it.get("level") or "-",
                    "message": it.get("message") or it.get("msg") or it.get("id") or "-",
                    "file": it.get("file") or it.get("path") or "-",
                    "line": it.get("line") or it.get("line_number") or "-",
                    "rule": it.get("rule") or it.get("rule_id") or it.get("id") or "-",
                }
            )
    rows.sort(key=lambda r: (_severity_rank(r.get("severity")), str(r.get("file")), str(r.get("line"))))
    return rows[: max(1, int(top_n))]


def _collect_test_rows(summary: Dict[str, Any], top_n: int) -> List[Dict[str, Any]]:
    tests = summary.get("tests", {}) if isinstance(summary.get("tests"), dict) else {}
    results = tests.get("results", [])
    if not isinstance(results, list):
        return []
    rows: List[Dict[str, Any]] = []
    for it in results:
        if not isinstance(it, dict):
            continue
        rows.append(
            {
                "file": it.get("file") or "-",
                "ok": it.get("ok"),
                "reason": it.get("reason") or "-",
                "test_file": it.get("test_file") or "-",
                "plan_ok": it.get("plan_ok"),
            }
        )
    rows.sort(key=lambda r: (0 if r.get("ok") is False else 1, str(r.get("file"))))
    return rows[: max(1, int(top_n))]


def _collect_coverage_rows(summary: Dict[str, Any]) -> List[Tuple[str, str]]:
    coverage = summary.get("coverage", {}) if isinstance(summary.get("coverage"), dict) else {}
    return [
        ("enabled", _fmt(coverage.get("enabled"))),
        ("ok", _fmt(coverage.get("ok"))),
        ("line_rate_pct", _fmt(coverage.get("line_rate_pct") or coverage.get("line_rate"))),
        ("branch_rate_pct", _fmt(coverage.get("branch_rate_pct") or coverage.get("branch_rate"))),
        ("threshold", _fmt(coverage.get("threshold"))),
        ("xml", _fmt(coverage.get("xml"))),
        ("html", _fmt(coverage.get("html"))),
        ("reason", _fmt(coverage.get("reason") or coverage.get("parse_error"))),
    ]


def generate_local_docx(summary: Dict[str, Any], output_path: Path) -> None:
    if docx is None:
        raise ImportError("python-docx 미설치로 DOCX 생성 불가")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    doc.add_heading("Local Workflow Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")

    rows = _build_summary_rows(summary)
    sections = {}
    for section, key, value in rows:
        sections.setdefault(section, []).append((key, value))

    for section, items in sections.items():
        doc.add_heading(section.upper(), level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Key"
        hdr[1].text = "Value"
        for key, value in items:
            row = table.add_row().cells
            row[0].text = key
            row[1].text = value

    doc.add_heading("TESTS DETAIL (TOP N)", level=2)
    test_rows = _collect_test_rows(summary, TOP_N_DEFAULT)
    if test_rows:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Source"
        hdr[1].text = "OK"
        hdr[2].text = "Reason"
        hdr[3].text = "Test File"
        hdr[4].text = "Plan OK"
        for item in test_rows:
            row = table.add_row().cells
            row[0].text = _fmt(item.get("file"))
            row[1].text = _fmt(item.get("ok"))
            row[2].text = _fmt(item.get("reason"))
            row[3].text = _fmt(item.get("test_file"))
            row[4].text = _fmt(item.get("plan_ok"))
    else:
        doc.add_paragraph("No test results.")

    doc.add_heading("COVERAGE DETAIL", level=2)
    cov_rows = _collect_coverage_rows(summary)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Key"
    hdr[1].text = "Value"
    for key, value in cov_rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    doc.add_heading("STATIC TOP N", level=2)
    issue_rows = _collect_static_issues(summary, TOP_N_DEFAULT)
    if issue_rows:
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Tool"
        hdr[1].text = "Severity"
        hdr[2].text = "Message"
        hdr[3].text = "File"
        hdr[4].text = "Line"
        hdr[5].text = "Rule"
        for item in issue_rows:
            row = table.add_row().cells
            row[0].text = _fmt(item.get("tool"))
            row[1].text = _fmt(item.get("severity"))
            row[2].text = _fmt(item.get("message"))
            row[3].text = _fmt(item.get("file"))
            row[4].text = _fmt(item.get("line"))
            row[5].text = _fmt(item.get("rule"))
    else:
        doc.add_paragraph("No static issues.")

    doc.save(str(output_path))


def generate_local_xlsx(summary: Dict[str, Any], output_path: Path) -> None:
    if Workbook is None:
        raise ImportError("openpyxl 미설치로 XLSX 생성 불가")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["Section", "Key", "Value"])

    for section, key, value in _build_summary_rows(summary):
        ws.append([section, key, value])

    ws_tests = wb.create_sheet("Tests")
    ws_tests.append(["Source", "OK", "Reason", "Test File", "Plan OK"])
    for item in _collect_test_rows(summary, TOP_N_DEFAULT):
        ws_tests.append(
            [
                _fmt(item.get("file")),
                _fmt(item.get("ok")),
                _fmt(item.get("reason")),
                _fmt(item.get("test_file")),
                _fmt(item.get("plan_ok")),
            ]
        )

    ws_cov = wb.create_sheet("Coverage")
    ws_cov.append(["Key", "Value"])
    for key, value in _collect_coverage_rows(summary):
        ws_cov.append([key, value])

    ws_static = wb.create_sheet("StaticTop")
    ws_static.append(["Tool", "Severity", "Message", "File", "Line", "Rule"])
    for item in _collect_static_issues(summary, TOP_N_DEFAULT):
        ws_static.append(
            [
                _fmt(item.get("tool")),
                _fmt(item.get("severity")),
                _fmt(item.get("message")),
                _fmt(item.get("file")),
                _fmt(item.get("line")),
                _fmt(item.get("rule")),
            ]
        )

    wb.save(str(output_path))
