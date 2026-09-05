"""SwUDS 문서에서 **설계 ID**(`SwUFn_xxxx`)를 읽는다.

## 무엇을 푸는가

시험 규격서의 `Related ID / SUDS` 칸과 `TC_ID` 는 **SwUDS 가 부여한 설계 ID** 를
가리켜야 한다. 정본 실측(KJPDS02_SwUTS v1.02): `TC_ID = "SwUTC_" + SUDS` 가
**1,013 / 1,014** 에서 성립한다.

생성기는 그동안 `SwUFn_{소스파싱순번:04d}` 를 만들어 두 칸에 넣었다. 모양은 같지만
**다른 설계 요소를 가리킨다** — 정본과 교집합이 251개 중 178개뿐이었다. 틀린 ID 가
추적성으로 보이는 것은 빈칸보다 나쁘다(`[[project_provenance_laundering]]` 계열).

## 근거는 문서 안에 있다

SwUDS 본문에 `SwUFn_0101: main` 형태의 문단이 함수마다 있다(실측 1,035개).
그대로 읽어 `함수명 → 설계 ID` 로 만든다.

## 동명이인은 채우지 않는다

같은 함수명이 서로 다른 설계 ID 를 갖는 경우가 있다(다른 모듈의 static 함수 —
실측 9건: `SCI0_Init` 이 `SwUFn_2901` 과 `SwUFn_3515` 양쪽). 이름만으로는 어느
쪽인지 정할 수 없으므로 **후보에서 제외한다**. 임의로 하나를 고르면 0.9% 를 채우는
대신 그 0.9% 가 조용히 틀린다.

실측 정합: 유일 이름 기준 정본 1,014건 중 **1,001건 일치(98.7%)** · 이름 미발견 4 ·
동명이인 9(제외 대상).
"""
from __future__ import annotations

import logging
import re
import time
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

_logger = logging.getLogger(__name__)

# 본문 문단 `SwUFn_0101: main` — 콜론 뒤 첫 토큰이 함수명이다.
_DESIGN_ID_PAT = re.compile(r"^(SwUFn_\d+)\s*[:：]\s*(\S+)")

# (경로, mtime_ns, size) → 결과. 53MB docx 파싱이 수 초라 매 호출 재파싱은 못 쓴다.
_CACHE: Dict[str, Tuple[Tuple[int, int], float, Dict[str, Any]]] = {}
_CACHE_TTL_SEC = 600.0


def _signature(p: Path) -> Optional[Tuple[int, int]]:
    try:
        st = p.stat()
        return (st.st_mtime_ns, st.st_size)
    except OSError as exc:
        _logger.debug("uds design-id: stat failed for %s: %s", p, exc)
        return None


def load_uds_design_ids(uds_path: str) -> Dict[str, Any]:
    """`{"by_name": {함수명: 설계ID}, "ambiguous": [...], "total": n}` 를 돌려준다.

    ⚠ 호출부는 **로컬로 materialize 된 경로**를 준다. cloudium 경로를 직접 주면
    `Document()` 가 열지 못한다(worker 만 권한을 가진다).

    파싱 실패는 빈 맵이다 — 그 경우 호출부는 ID 칸을 비워야 하고, **순번으로
    대체하면 안 된다**(그게 원래 결함이다).
    """
    empty: Dict[str, Any] = {"by_name": {}, "ambiguous": [], "total": 0, "source": ""}
    raw = str(uds_path or "").strip()
    if not raw:
        return empty
    p = Path(raw)
    sig = _signature(p)
    if sig is None:
        return empty

    key = str(p.resolve()) if p.exists() else raw
    hit = _CACHE.get(key)
    now = time.monotonic()
    if hit and hit[0] == sig and (now - hit[1]) < _CACHE_TTL_SEC:
        return hit[2]

    if p.suffix.lower() != ".docx":
        _logger.info("uds design-id: unsupported suffix %s (%s)", p.suffix, p.name)
        return empty
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
    except Exception as exc:  # noqa: BLE001 — python-docx 는 예외 종류가 넓다
        _logger.warning("uds design-id: cannot open %s: %s", p.name, exc)
        return empty

    seen: Dict[str, str] = {}
    ambiguous: set[str] = set()
    total = 0
    for para in doc.paragraphs:
        g = _DESIGN_ID_PAT.match((para.text or "").strip())
        if not g:
            continue
        total += 1
        design_id, fn_name = g.group(1), g.group(2)
        prev = seen.get(fn_name)
        if prev is None:
            seen[fn_name] = design_id
        elif prev != design_id:
            # 동명이인 — 이름만으로 못 고른다. 아래에서 통째로 뺀다.
            ambiguous.add(fn_name)

    for name in ambiguous:
        seen.pop(name, None)

    result: Dict[str, Any] = {
        "by_name": seen,
        "ambiguous": sorted(ambiguous),
        "total": total,
        "source": p.name,
    }
    _CACHE[key] = (sig, now, result)
    _logger.info(
        "uds design-id: %s → %d ids (문단 %d · 동명이인 %d 제외)",
        p.name, len(seen), total, len(ambiguous),
    )
    return result


def resolve_design_id(design_ids: Optional[Dict[str, Any]], fn_name: Any) -> str:
    """함수명 → 설계 ID. 못 찾으면 **빈 문자열**(순번으로 대체하지 않는다)."""
    if not design_ids:
        return ""
    by_name = design_ids.get("by_name") or {}
    name = str(fn_name or "").strip()
    if not name:
        return ""
    got = by_name.get(name)
    if got:
        return str(got)
    # 대소문자만 다른 표기는 같은 함수로 본다(문서 표기 흔들림). 그래도 못 찾으면 빈칸.
    lowered = name.lower()
    for k, v in by_name.items():
        if k.lower() == lowered:
            return str(v)
    return ""
