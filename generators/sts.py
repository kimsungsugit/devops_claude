"""STS (Software Test Specification) auto-generation engine.

Generates XLSM output from SRS requirements, UDS function details,
SDS component mapping, and source code analysis.
"""
from __future__ import annotations

import logging
import os
import re
import threading
import time
from collections import OrderedDict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from generators._artifact_check import apply_write_back_check
from generators.safety_marks import is_safety_asil
from generators.safety_marks import resolve_safety_related as _safety_mark_impl
from report_gen.doc_kind import is_sds_filename
from report_gen.requirements import (
    _extract_sds_partition_map,
    is_sds_placeholder_key,
    normalize_sds_key,
)

_logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_REQ_ID_PAT = re.compile(
    r"\b(Sw(?:TR|TSR|NTR|NTSR|EI|CNF|ST|STR)_\d+)\b"
)

# 정본(KJPDS02_SwTS v1.02) 시트명. SwITS 의 `3.SW Integration Test Spec` 과 다르다.
_SPEC_SHEET_NAME = "3.SW Test Spec"

# ── 값 어휘 — **STS 정본의 Introduction 1.5/1.6** 이 출처다 ──────────────────
#
# ⚠ 같은 개념이라도 **문서마다 약어가 다르다**. 실측(2026-08-11):
#     SwTS  1.5 Test Method  : Requirement-based test=RBT · Fault injection test=FIT
#     SwTS  1.6 Generation   : Analysis of requirements=AOR · Equivalent class=ECA
#                              · Boundary value analysis=BAA
#     SwUTS 1.5 Test Method  : REQ · IFT · FI          ← 다른 약어
#     SwUTS 1.6 Generation   : AOR · AEC · ABV · ERG   ← 다른 약어
#   통일하지 말 것. 예전 판은 `FNCT`/`RVW`/`ELCT`/`ERG` 를 썼는데 **SwTS Introduction
#   표에 없는 값**이라 문서를 읽는 사람이 대조할 표가 없었다.
_TEST_METHODS = {"RBT", "FIT"}
_GEN_METHODS = {"AOR", "ECA", "BAA"}
_DEFAULT_TEST_METHOD = "RBT"   # 정본 실측: 102건 전부 RBT
_DEFAULT_GEN_METHOD_STS = "AOR"  # 정본 실측: 102건 전부 AOR

# 실행 산출물이 없는 검증방법. RVW 는 "소스 코드에서 구현부 확인" 같은 **사람이 읽는**
# 활동이라(`_generate_review_steps`) 실행 시험과 증거 성격이 다르다.
# 커버리지를 방법 구분 없이 한 숫자로 내면 "100%"가 실행시험 100%인지 리뷰 포함인지
# 구분되지 않는다 — 실측(HDPDM01 SRS 63건): 보고 100.0% vs 실행시험 87.3%.
_REVIEW_ONLY_METHODS = {"RVW"}

_DEFAULT_TEST_ENV = "SwTE_01"
_MAX_TC_PER_REQ = 5
_MAX_STEPS_PER_TC = 15

_HEADER_ROW = 6

# ── STS 시트 열 스키마 (SSOT) ────────────────────────────────────────────────
# (열 번호 1-indexed, 필드 키, row 6 헤더 라벨).
# writer(generate_sts_xlsm)와 validator(validate_sts_xlsm)가 **같은 출처**를 봐야 한다.
# 과거엔 validator가 SUTS 레이아웃 상수를 그대로 재사용해 5/6/4열(TestEnv·TestMethod·
# SafetyRelated)을 Action·Expected·요구ID로 읽었다 → 실제 Action/Expected가 전부 비어도
# "정상"으로 통과하고, 요구 링크율은 Safety Related 채움률을 보고했다.
_STS_SCHEMA: List[Tuple[int, str, str]] = [
    (1,  "seq",              ""),
    (2,  "tc_id",            "Test Case ID"),
    (3,  "title",            "Title"),
    (4,  "safety_related",   "Safety\nRelated"),
    (5,  "test_environment", "Test\nEnvironment"),
    (6,  "test_method",      "Test\nMethod"),
    (7,  "gen_method",       "Test Case\nGen. Method"),
    (8,  "fs_req",           "FS_REQ"),
    (9,  "description",      "Description"),
    (10, "precondition",     "Pre-condition"),
    (11, "action",           "Test Action\n(Sequence)"),
    (12, "expected",         "Expected Result"),
    (13, "srs",              "SRS"),
]
STS_COL: Dict[str, int] = {key: col for col, key, _ in _STS_SCHEMA}
_COL_HEADERS = [label for _, _, label in _STS_SCHEMA]
_LAST_COL = _STS_SCHEMA[-1][0]
# 산출물이 이 필드들을 비운 채 나오면 시험 명세로서 의미가 없다 — validator 필수 축.
_STS_REQUIRED_FIELDS = ("tc_id", "action", "expected", "srs")

_COL_WIDTHS = [4.0, 20.5, 52.0, 10.0, 12.0, 10.6, 10.4, 13.0, 61.5, 36.0, 61.0, 77.0, 14.0]
_MERGE_COLS = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 12]  # 0-indexed → cols A,B,C,D,E,F,G,H,I,J,M (K,L per-step)
# Columns that get center alignment (not wrap)
_CENTER_COLS = {1, 4, 5, 6, 7, 13}  # #, Safety, TestEnv, TestMethod, GenMethod, SRS

_SDS_MAP_CACHE: Optional[Dict[str, Dict[str, str]]] = None

# ── HSIS 파서 캐시 ───────────────────────────────────────────────────────────
# 키는 **파일 정체성**(정규화 경로, mtime_ns, size)이다. 과거엔 경로를 무시하는 단일
# 전역이라, 장기 기동 서버(`backend/routers/local.py`)가 프로젝트 A의 HSIS를 한 번 읽으면
# 이후 프로젝트 B의 STS/SUTS/SITS 생성이 A의 HW 신호를 그대로 받았다 — 경고 한 줄 없이.
# 실측: 전혀 다른 빈 문서 경로로 재호출해도 첫 파일의 signal 20건이 **동일 객체로** 반환됐다.
# mtime_ns/size 를 키에 넣으므로 파일이 바뀌면 자동 무효화된다.
_HSIS_CACHE_MAX = 8
_HSIS_SIGNALS_CACHE: "OrderedDict[Tuple[str, int, int], Dict[str, Any]]" = OrderedDict()
_HSIS_CACHE_LOCK = threading.Lock()

# HSIS 데이터 행 판정에 쓰는 ID 패턴. 두 파서가 공유한다(아래 `_is_hsis_data_row`).
_HSI_ID_PAT = re.compile(r"HSI_?\d+", re.I)
_HSIS_REQ_ID_PAT = re.compile(r"S[wy][A-Za-z]{1,}_?\d+")


def _is_hsis_data_row(sig_id: Any, related_id: Any) -> bool:
    """HSIS 데이터 행인가 — `_load_hsis_signals`/`parse_hsis_signals` 공용 판정.

    HSI ID가 있거나 Related 열에 Sw/Sy 요구 ID가 있으면 데이터 행이다.
    과거 `_load_hsis_signals`는 HSI ID만 봐서 ID 열이 빈 행을 통째로 버렸다
    (실측 HDPDM01 v5.00: 21건 중 1건 — `Battery Power / u16g_ApiIn_Vsup / SyEI_01`).
    판정을 파서마다 따로 두면 한쪽만 고쳐지고 다른 쪽이 잠복하므로 여기로 묶는다.
    """
    return bool(
        _HSI_ID_PAT.match(str(sig_id or "").strip())
        or _HSIS_REQ_ID_PAT.search(str(related_id or ""))
    )


def _hsis_cache_key(p: Path) -> Optional[Tuple[str, int, int]]:
    """(정규화 경로, mtime_ns, size). stat 실패 시 None → 캐시를 아예 쓰지 않는다."""
    try:
        st = p.stat()
        resolved = str(p.resolve())
    except OSError as exc:
        _logger.debug("HSIS 캐시 키 산출 실패(%s) — 캐시 없이 진행: %s", exc, p)
        return None
    return (os.path.normcase(resolved), st.st_mtime_ns, st.st_size)


def _hsis_cache_get(key: Optional[Tuple[str, int, int]]) -> Optional[Dict[str, Any]]:
    if key is None:
        return None
    with _HSIS_CACHE_LOCK:
        hit = _HSIS_SIGNALS_CACHE.get(key)
        if hit is not None:
            _HSIS_SIGNALS_CACHE.move_to_end(key)
        return hit


def _hsis_cache_put(key: Optional[Tuple[str, int, int]], value: Dict[str, Any]) -> None:
    if key is None:
        return
    with _HSIS_CACHE_LOCK:
        _HSIS_SIGNALS_CACHE[key] = value
        _HSIS_SIGNALS_CACHE.move_to_end(key)
        while len(_HSIS_SIGNALS_CACHE) > _HSIS_CACHE_MAX:
            _HSIS_SIGNALS_CACHE.popitem(last=False)


def _load_default_sds_map() -> Dict[str, Dict[str, str]]:
    global _SDS_MAP_CACHE
    if _SDS_MAP_CACHE is not None:
        return _SDS_MAP_CACHE
    docs_dir = Path(__file__).resolve().parents[1] / "docs"
    merged: Dict[str, Dict[str, str]] = {}
    picked: List[str] = []
    if docs_dir.exists():
        for path in docs_dir.glob("*.docx"):
            # `"sds" in name` 은 `SwDS` 표기를 놓친다("swds" 에 "sds" 없음) — 단일 출처 사용.
            if not is_sds_filename(path.name):
                continue
            picked.append(path.name)
            data = _extract_sds_partition_map(str(path))
            for key, value in data.items():
                if key not in merged:
                    merged[key] = dict(value)
                    continue
                for field in ("asil", "related", "description"):
                    if value.get(field) and not merged[key].get(field):
                        merged[key][field] = value[field]
    if merged:
        # ⚠ 침묵 금지 — 이 맵은 **프로젝트 무관**인데 실측상 요구-함수 링크 전량을
        #   좌우한다(HDPDM01 기준 5,992건 100%). 어느 문서가 쓰였는지 남긴다.
        _logger.warning(
            "SDS 미지정 — 저장소 docs/ 글롭 폴백 사용(**프로젝트 무관**): %s (%d 엔트리). "
            "대상 프로젝트의 SDS 를 넘기면 이 폴백은 쓰이지 않는다",
            ", ".join(picked) or "(없음)", len(merged))
    _SDS_MAP_CACHE = merged
    return merged


def _function_sds_candidates(info: Dict[str, Any]) -> List[str]:
    """SDS 파티션을 찾을 후보 이름들 — **함수 이름이 첫 후보**다.

    ⚠ 예전 판은 `module_name` 파생만 냈다. 그런데 이 프로젝트 SDS 의 871 파티션 중
    **588개가 `kind='function'`**(전부 `related` 보유)이고 키가 곧 함수 이름이다.
    후보에 함수명이 없으면 그 588개는 **모듈명이 우연히 닮았을 때만** 걸린다.
    실측(KJPDS02_PV): 함수명을 넣자 356개가 **정확 키**로 붙고, 요구 매핑이
    43/68 → 48/68 로, 어느 요구에도 못 붙는 함수가 202 → 151 로 줄었다.
    """
    module_name = str(info.get("module_name") or "").strip()
    candidates: List[str] = []
    fn_name = str(info.get("name") or "").strip()
    if fn_name:
        candidates.append(fn_name)
    if module_name:
        candidates.append(module_name)
        base = re.sub(r"_pds$", "", module_name, flags=re.I)
        candidates.append(base)
        tokenized = re.sub(r"([a-z])([A-Z])", r"\1 \2", base.replace("_", " "))
        tokenized = re.sub(r"\bctrl\b", "control", tokenized, flags=re.I)
        tokenized = re.sub(r"\bdiag\b", "diagnostic", tokenized, flags=re.I)
        words = [w for w in tokenized.split() if w.lower() not in {"ap", "drv", "sys", "pds", "main", "func"}]
        if words:
            candidates.append(" ".join(words))
    return [c for c in dict.fromkeys([c.strip() for c in candidates if c and c.strip()])]


def _lookup_sds_related_ids(info: Dict[str, Any], sds_map: Dict[str, Dict[str, str]]) -> List[str]:
    candidates = _function_sds_candidates(info)
    for candidate in candidates:
        direct = sds_map.get(candidate.lower())
        if direct and direct.get("related"):
            return [m.group(1) for m in _REQ_ID_PAT.finditer(str(direct.get("related") or ""))]
    for candidate in candidates:
        nc = normalize_sds_key(candidate)
        if not nc:
            continue
        for key, value in sds_map.items():
            nk = normalize_sds_key(key)
            if not nk or is_sds_placeholder_key(nk):
                continue
            if nc == nk or nc in nk or nk in nc:
                ids = [m.group(1) for m in _REQ_ID_PAT.finditer(str(value.get("related") or ""))]
                if ids:
                    return ids
                # `related` 가 빈 칸이면 **여기서 끝내지 않는다**. 실측 41건이 전부
                # `(swdsg) software architecture design guideline….docx`(SDS 안의
                # 문서 목록 행)에 걸려 탐색이 멈췄다 — `Lin` 이 `guide**lin**e` 에
                # 걸린 것이다. 빈 칸은 "요구가 없다"가 아니라 **그 행이 파티션이
                # 아니라는** 뜻이므로 다음 후보를 계속 본다.
                continue
    return []


# ---------------------------------------------------------------------------
# Additional document loaders (SDS summary, UDS descriptions, STP context)
# ---------------------------------------------------------------------------

def _load_sds_summary(sds_path: str) -> str:
    """Extract a concise design summary from an SDS DOCX for AI prompt context.

    Returns a short text block describing components/modules found in the SDS.
    """
    if not sds_path:
        return ""
    p = Path(sds_path)
    if not p.exists():
        return ""
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
    except Exception:
        return ""

    lines: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or para.runs and para.runs[0].bold:
            lines.append(f"[Section] {text}")
        elif len(text) > 20:
            lines.append(text)
        if len(lines) >= 60:
            break

    # Also scan tables for component/description pairs
    for table in doc.tables[:8]:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        hj = " ".join(header)
        if not any(k in hj for k in ("component", "module", "sw", "function", "design", "description")):
            continue
        for row in table.rows[1:6]:
            cells = [c.text.strip() for c in row.cells]
            non_empty = [c for c in cells if c]
            if non_empty:
                lines.append(" | ".join(non_empty[:4]))

    return "\n".join(lines[:80])


def _load_uds_descriptions(uds_path: str) -> Dict[str, str]:
    """Parse a UDS DOCX/XLSM and extract function_name → description mapping.

    Used to enrich function_details with AI-written descriptions from UDS.
    """
    if not uds_path:
        return {}
    p = Path(uds_path)
    if not p.exists():
        return {}

    result: Dict[str, str] = {}
    suffix = p.suffix.lower()

    # --- DOCX path ---
    if suffix == ".docx":
        try:
            import docx as _docx  # type: ignore
            doc = _docx.Document(str(p))
        except Exception:
            return {}
        current_func: Optional[str] = None
        desc_lines: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            # Headings are often function names
            if "heading" in style or (para.runs and para.runs[0].bold and len(text) < 80):
                if current_func and desc_lines:
                    result[current_func.lower()] = " ".join(desc_lines[:3])
                current_func = text
                desc_lines = []
            elif current_func:
                desc_lines.append(text)
        if current_func and desc_lines:
            result[current_func.lower()] = " ".join(desc_lines[:3])
        # Also scan tables: look for (Name | Description) structure
        for table in doc.tables:
            if not table.rows:
                continue
            header = [c.text.strip().lower() for c in table.rows[0].cells]
            name_col = next((i for i, h in enumerate(header) if "name" in h or "function" in h), -1)
            desc_col = next((i for i, h in enumerate(header) if "desc" in h or "summary" in h or "설명" in h), -1)
            if name_col == -1 or desc_col == -1:
                continue
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) > max(name_col, desc_col):
                    fname = cells[name_col].strip()
                    fdesc = cells[desc_col].strip()
                    if fname and fdesc:
                        result[fname.lower()] = fdesc
        return result

    # --- XLSM / XLSX path ---
    if suffix in (".xlsm", ".xlsx"):
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        except Exception:
            return {}
        for ws in wb.worksheets:
            headers: List[str] = []
            name_col = desc_col = -1
            for ri, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c or "").strip() for c in row]
                if ri == 0:
                    headers = [c.lower() for c in cells]
                    name_col = next((i for i, h in enumerate(headers) if "name" in h or "function" in h), -1)
                    desc_col = next((i for i, h in enumerate(headers) if "desc" in h or "summary" in h or "설명" in h), -1)
                    if name_col == -1 or desc_col == -1:
                        break
                    continue
                if len(cells) > max(name_col, desc_col):
                    fname = cells[name_col]
                    fdesc = cells[desc_col]
                    if fname and fdesc:
                        result[fname.lower()] = fdesc
                if ri > 500:
                    break
        try:
            wb.close()
        except Exception:
            pass
        return result

    return {}


def _load_stp_context(stp_path: str) -> str:
    """Extract test strategy/scope text from an STP document (.docx/.pdf/.txt)."""
    if not stp_path:
        return ""
    p = Path(stp_path)
    if not p.exists():
        return ""

    suffix = p.suffix.lower()

    # ── Plain text ────────────────────────────────────────────────────────
    if suffix == ".txt":
        try:
            return p.read_text(encoding="utf-8", errors="replace")[:8000]
        except Exception:
            return ""

    # ── PDF ───────────────────────────────────────────────────────────────
    if suffix == ".pdf":
        text = ""
        try:
            from pdfminer.high_level import extract_text as _pdf_extract  # type: ignore
            text = _pdf_extract(str(p))
        except ImportError:
            try:
                import pypdf  # type: ignore
                reader = pypdf.PdfReader(str(p))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception:
                pass
        except Exception:
            pass
        return text[:8000]

    # ── DOCX ──────────────────────────────────────────────────────────────
    try:
        import docx as _docx  # type: ignore
        doc = _docx.Document(str(p))
    except Exception:
        return ""

    # Keywords that indicate relevant STP sections
    _STRATEGY_KEYWORDS = {
        "test strategy", "test scope", "test approach", "test method",
        "test environment", "entry criteria", "exit criteria",
        "테스트 전략", "테스트 범위", "테스트 방법", "시험 환경", "시험 범위",
    }

    lines: List[str] = []
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or (para.runs and para.runs[0].bold):
            capture = any(kw in text.lower() for kw in _STRATEGY_KEYWORDS)
            if capture:
                lines.append(f"[{text}]")
        elif capture:
            lines.append(text)
        if len(lines) >= 40:
            break

    return "\n".join(lines)


def _load_hsis_signals(hsis_path: str) -> Dict[str, Any]:
    """Parse HSIS xlsx and return structured signal data for test generation.

    Expected sheet '2.HSIS' columns (0-indexed, row 3 = header):
      F(5)=ID, G(6)=Signal Name, H(7)=Signal Type, I(8)=Pin Name,
      L(11)=Direction, M(12)=Characteristics, T(19)=SW Variable Name, U(20)=Related ID

    Returns:
        {
          'sw_var_names': List[str],          # for _HW_SIGNAL_PAT extension
          'signals': List[Dict],              # full signal info
          'pat': re.Pattern,                  # extended hw signal pattern
        }

    결과는 파일 정체성(경로+mtime_ns+size)으로 캐시된다 — `_hsis_cache_key` 주석 참조.
    """
    empty: Dict[str, Any] = {"sw_var_names": [], "signals": [], "pat": _HW_SIGNAL_PAT}

    if not hsis_path:
        return empty
    p = Path(hsis_path)
    if not p.exists():
        _logger.warning("HSIS file not found: %s", hsis_path)
        return empty

    cache_key = _hsis_cache_key(p)
    cached = _hsis_cache_get(cache_key)
    if cached is not None:
        return cached

    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    except Exception as e:
        _logger.warning("Cannot open HSIS xlsx: %s", e)
        return empty

    # Find HSIS sheet
    sheet_name = None
    for name in wb.sheetnames:
        if "hsis" in name.lower() or "2." in name:
            sheet_name = name
            break
    if sheet_name is None and wb.sheetnames:
        sheet_name = wb.sheetnames[0]

    ws = wb[sheet_name]
    signals: List[Dict[str, Any]] = []

    # Column indices (0-based from the HSIS structure):
    # F=5(ID), G=6(Signal Name), H=7(Signal Type), L=11(Direction),
    # M=12(Characteristics), T=19(SW Variable Name), U=20(Related ID)
    _COL_ID = 5
    _COL_SIG_NAME = 6
    _COL_SIG_TYPE = 7
    _COL_DIRECTION = 11
    _COL_CHARACTERISTICS = 12
    _COL_SW_VAR = 19
    _COL_RELATED = 20

    header_found = False
    for ri, row in enumerate(ws.iter_rows(values_only=True)):
        if ri > 100:  # HSIS sheets are not that long
            break
        cells = [str(c or "").strip() for c in row]
        if len(cells) < 7:
            continue

        # Detect header row (contains "Signal Name" or "SW Variable")
        row_text = " ".join(cells).lower()
        if not header_found:
            if "signal" in row_text and ("sw variable" in row_text or "variable name" in row_text):
                header_found = True
            continue

        # Data rows: HSI ID가 있거나 Related에 Sw/Sy 요구 ID가 있으면 데이터 행
        sig_id = cells[_COL_ID] if len(cells) > _COL_ID else ""
        related = cells[_COL_RELATED] if len(cells) > _COL_RELATED else ""
        if not _is_hsis_data_row(sig_id, related):
            continue

        sig_name = cells[_COL_SIG_NAME] if len(cells) > _COL_SIG_NAME else ""
        sig_type = cells[_COL_SIG_TYPE] if len(cells) > _COL_SIG_TYPE else ""
        direction = cells[_COL_DIRECTION] if len(cells) > _COL_DIRECTION else ""
        characteristics = cells[_COL_CHARACTERISTICS] if len(cells) > _COL_CHARACTERISTICS else ""
        sw_var = cells[_COL_SW_VAR] if len(cells) > _COL_SW_VAR else ""

        if not sig_name and not sw_var:
            continue

        signals.append({
            "id": sig_id,
            "signal_name": sig_name,
            "signal_type": sig_type,
            "direction": direction,
            "characteristics": characteristics,
            "sw_var_name": sw_var,
            "related_id": related,
        })

    try:
        wb.close()
    except Exception:
        pass

    if not signals:
        _logger.info("HSIS: no signals parsed from %s (sheet=%s)", hsis_path, sheet_name)
        _hsis_cache_put(cache_key, empty)
        return empty

    # Collect SW variable names for pattern building
    # Some cells have multiple names separated by newlines or commas
    raw_sw_vars: List[str] = []
    for s in signals:
        if s["sw_var_name"]:
            for tok in re.split(r"[\n,\s]+", s["sw_var_name"]):
                tok = tok.strip().strip(",")
                if tok and re.match(r"^[A-Za-z_]\w+$", tok):
                    raw_sw_vars.append(tok)
    sw_var_names = list(dict.fromkeys(raw_sw_vars))  # deduplicate, preserve order

    # Also collect HW pin names that look like C identifiers (e.g. PTP_PTP4, PIEL_PIEL0)
    sig_names_id = []
    for s in signals:
        n = re.sub(r"[^A-Za-z0-9_]", "_", s["signal_name"])
        if len(n) > 2 and re.match(r"^[A-Za-z_]", n):
            sig_names_id.append(n)

    # Build extended HW signal pattern
    extra_terms = [re.escape(v) for v in sw_var_names if v]
    extra_terms += [re.escape(n) for n in sig_names_id if len(n) > 2]
    if extra_terms:
        extended_pat = re.compile(
            _HW_SIGNAL_PAT.pattern + "|" + "|".join(f"\\b{t}\\b" for t in extra_terms),
            re.I,
        )
    else:
        extended_pat = _HW_SIGNAL_PAT

    result: Dict[str, Any] = {
        "sw_var_names": sw_var_names,
        "signals": signals,
        "pat": extended_pat,
    }
    _hsis_cache_put(cache_key, result)
    _logger.info("HSIS loaded: %d signals, %d SW var names from %s",
                 len(signals), len(sw_var_names), hsis_path)
    return result


def parse_hsis_signals(hsis_path: str) -> Dict[str, Any]:
    """Cache-free HSIS xlsx parser with header auto-detection (layout-variant safe).

    `_load_hsis_signals`(위)는 모듈 캐시 + 고정 0-based 컬럼이라 (1) 멀티파일 요청 시
    첫 결과 오염, (2) HSIS 버전별 컬럼 오프셋(실측: 260105 v5.00 SwVar=20/Related=21 vs
    hiMA계약 23/26)에 깨진다. 추적성 매트릭스 추출 엔드포인트는 이 함수를 쓴다 — 캐시 없이
    헤더 라벨로 컬럼을 동적 탐지한다.

    Returns: {signals:[{id, signal_name, sw_var_name, related_id, direction}], sw_var_names:[...]}.
    헤더 탐지 실패 시 signals=[] + available_columns 힌트(STS available_sheets 패턴).
    """
    empty: Dict[str, Any] = {"sw_var_names": [], "signals": []}
    if not hsis_path:
        return empty
    p = Path(hsis_path)
    if not p.exists():
        _logger.warning("HSIS file not found: %s", hsis_path)
        return empty
    try:
        import openpyxl  # type: ignore
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    except Exception as e:
        _logger.warning("Cannot open HSIS xlsx: %s", e)
        return empty

    sheet_name = None
    for name in wb.sheetnames:
        if "hsis" in name.lower() or "2." in name:
            sheet_name = name
            break
    if sheet_name is None and wb.sheetnames:
        sheet_name = wb.sheetnames[0]
    if not sheet_name:
        try:
            wb.close()
        except Exception:
            pass
        return empty
    ws = wb[sheet_name]

    rows: List[List[str]] = []
    for ri, row in enumerate(ws.iter_rows(values_only=True)):
        if ri > 5000:  # zip-bomb / runaway guard
            break
        rows.append([str(c or "").strip() for c in row])
    try:
        wb.close()
    except Exception:
        pass

    def _norm(s: Any) -> str:
        return re.sub(r"[\s_]+", "", str(s or "").strip().lower())

    # 헤더 행 탐지: 'related id' + ('sw variable'|'variable name'|'signal name') 동시 포함.
    hdr_idx = -1
    cols: Dict[str, int] = {}
    for ri in range(min(20, len(rows))):
        normed = [_norm(c) for c in rows[ri]]
        joined = " ".join(normed)
        if "relatedid" in joined and ("swvariable" in joined or "variablename" in joined or "signalname" in joined):
            for ci, n in enumerate(normed):
                if n == "id" and "id" not in cols:
                    cols["id"] = ci
                if "relatedid" in n and "related" not in cols:
                    cols["related"] = ci
                if ("swvariable" in n or "variablename" in n) and "swvar" not in cols:
                    cols["swvar"] = ci
                if "signalname" in n and "signame" not in cols:
                    cols["signame"] = ci
                if n == "direction" and "dir" not in cols:
                    cols["dir"] = ci
            hdr_idx = ri
            break
    if hdr_idx < 0 or "related" not in cols:
        _logger.info("HSIS: header not detected in %s (sheet=%s)", hsis_path, sheet_name)
        return {"sw_var_names": [], "signals": [],
                "available_columns": rows[2][:30] if len(rows) > 2 else []}

    data_rows = rows[hdr_idx + 1:]
    # ID 컬럼 disambiguation — 헤더 'id'가 여러 개(Arch Element ID·Connector ID)라
    # 데이터에서 HSI_\d+ 빈도 최대 컬럼을 ID로 확정.
    id_col = cols.get("id", -1)
    _hsi = _HSI_ID_PAT
    if id_col < 0 or not any(_hsi.match(dr[id_col]) for dr in data_rows[:30] if id_col < len(dr)):
        best, best_cnt = -1, 0
        max_c = max((len(dr) for dr in data_rows[:50]), default=0)
        for ci in range(max_c):
            cnt = sum(1 for dr in data_rows[:50] if ci < len(dr) and _hsi.match(dr[ci]))
            if cnt > best_cnt:
                best, best_cnt = ci, cnt
        if best_cnt:
            id_col = best

    related_col = cols["related"]
    swvar_col = cols.get("swvar", -1)
    signame_col = cols.get("signame", -1)
    dir_col = cols.get("dir", -1)

    def _get(dr: List[str], ci: int) -> str:
        return dr[ci] if 0 <= ci < len(dr) else ""

    signals: List[Dict[str, Any]] = []
    for dr in data_rows:
        sid = _get(dr, id_col)
        related = _get(dr, related_col)
        swv = _get(dr, swvar_col)
        # 데이터 행: HSI ID가 있거나 Related에 Sw/Sy ID가 있어야(설명/공백 행 스킵)
        if not _is_hsis_data_row(sid, related):
            continue
        if not related and not swv:
            continue
        signals.append({
            "id": sid,
            "signal_name": _get(dr, signame_col),
            "sw_var_name": swv,
            "related_id": related,
            "direction": _get(dr, dir_col),
        })

    sw_var_names: List[str] = []
    for s in signals:
        for tok in re.split(r"[\n,\s]+", s["sw_var_name"] or ""):
            tok = tok.strip().strip(",")
            if tok and re.match(r"^[A-Za-z_]\w+$", tok):
                sw_var_names.append(tok)
    sw_var_names = list(dict.fromkeys(sw_var_names))

    _logger.info("HSIS parsed(headerless): %d signals from %s", len(signals), hsis_path)
    return {"sw_var_names": sw_var_names, "signals": signals}


def _merge_uds_into_function_details(
    function_details: Dict[str, Dict[str, Any]],
    uds_descriptions: Dict[str, str],
) -> None:
    """Overwrite function_details[fid]['description'] with UDS-sourced text if available."""
    if not uds_descriptions:
        return
    for fid, info in function_details.items():
        if not isinstance(info, dict):
            continue
        fname = str(info.get("name") or "").lower()
        uds_desc = uds_descriptions.get(fname) or uds_descriptions.get(fid.lower())
        if uds_desc and len(uds_desc) > len(str(info.get("description") or "")):
            info["description"] = uds_desc


# ---------------------------------------------------------------------------
# Phase 1: Data extraction helpers
# ---------------------------------------------------------------------------

def parse_srs_docx_tables(srs_path: str) -> List[Dict[str, Any]]:
    """Parse SRS DOCX directly from tables to extract structured requirements.

    Each requirement detail table in the SRS has rows like:
      ID | SwTR_0101
      Name | Auto Close
      Description | ...
      ASIL | A
      Related ID | SyTR_0701, ...
      Verification criteria | ...
    """
    try:
        from docx import Document
    except (ImportError, Exception) as _docx_err:
        _logger.warning("python-docx not available; cannot parse SRS DOCX tables: %s", _docx_err)
        return []

    try:
        doc = Document(srs_path)
    except Exception as _open_err:
        _logger.warning("parse_srs_docx_tables: cannot open %s: %s", srs_path, _open_err)
        return []
    results: List[Dict[str, Any]] = []
    seen_ids: set = set()

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 3:
            continue
        cells_map: Dict[str, str] = {}
        for row in rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                key = cells[0].strip()
                val = cells[1].strip()
                if key and val:
                    cells_map[key] = val

        rid = cells_map.get("ID", "")
        if not _REQ_ID_PAT.match(rid):
            continue
        if rid in seen_ids:
            continue
        seen_ids.add(rid)

        results.append({
            "id": rid,
            "name": cells_map.get("Name", ""),
            "description": cells_map.get("Description", ""),
            "asil": cells_map.get("ASIL", ""),
            "related_id": cells_map.get("Related ID", ""),
            "software_state": cells_map.get("Software State", ""),
            "verification": cells_map.get("Verification criteria", ""),
            "priority": cells_map.get("Priority", ""),
            "req_type": _classify_req_type(rid),
        })

    _logger.info("Parsed %d requirements from SRS DOCX tables", len(results))
    return results


def parse_requirements_structured(req_texts: List[str]) -> List[Dict[str, Any]]:
    """Parse raw requirement text lines into structured dicts.

    Deduplicates by requirement ID, keeping the entry with the richest data.
    Each returned dict has: id, name, description, asil, related_id, req_type
    """
    seen: Dict[str, Dict[str, Any]] = {}
    for text in req_texts:
        for m_id in _REQ_ID_PAT.finditer(text):
            rid = m_id.group(1)
            if rid in seen:
                existing = seen[rid]
                if len(text) > len(existing.get("_raw", "")):
                    pass
                else:
                    continue

            req_type = _classify_req_type(rid)
            name = ""
            description = ""
            asil = ""
            related = ""

            m_asil = re.search(r"ASIL\s*[:|]\s*((?:ASIL-)?(?:A|B|C|D)|QM|TBD)", text, re.I)
            if m_asil:
                asil = m_asil.group(1).strip().replace("ASIL-", "")
            m_rel = re.search(r"Related\s*(?:ID)?\s*[:|]\s*([^\n]+)", text, re.I)
            if m_rel:
                related = m_rel.group(1).strip()

            remainder = _REQ_ID_PAT.sub("", text).strip()
            remainder = re.sub(r"ASIL\s*[:|]\s*\S+", "", remainder, flags=re.I)
            remainder = re.sub(r"Related\s*(?:ID)?\s*[:|][^\n]+", "", remainder, flags=re.I)
            remainder = remainder.strip(" -:|")

            if " - " in remainder:
                name, description = remainder.split(" - ", 1)
            elif " | " in remainder:
                parts = remainder.split(" | ", 1)
                name = parts[0].strip()
                description = parts[1].strip() if len(parts) > 1 else ""
            elif remainder:
                first_sent = re.split(r"[.。\n]", remainder, 1)
                name = first_sent[0].strip()[:120]
                description = remainder

            seen[rid] = {
                "id": rid,
                "name": name.strip(),
                "description": description.strip()[:500],
                "asil": asil,
                "related_id": related,
                "req_type": req_type,
                "_raw": text,
            }

    results = []
    for entry in seen.values():
        entry.pop("_raw", None)
        results.append(entry)
    return results


def _classify_req_type(req_id: str) -> str:
    if req_id.startswith("SwEI_"):
        return "EI"
    if req_id.startswith("SwTSR_"):
        return "TSR"
    if req_id.startswith("SwNTSR_"):
        return "NTSR"
    if req_id.startswith("SwNTR_"):
        return "NTR"
    if req_id.startswith("SwTR_"):
        return "TR"
    return "OTHER"


# ---------------------------------------------------------------------------
# Phase 1: Requirement -> Function mapping
# ---------------------------------------------------------------------------

def load_uds_design_ids(uds_path: str) -> Dict[str, List[str]]:
    """SwUDS 함수표 → ``함수 이름(lower) → [설계 ID]``. 설계-ID 브리지의 좌측 끝.

    ## 왜 필요한가 (실측 2026-08-18, KJPDS02_PV)

    함수 이름·모듈 이름으로 SwDS 파티션을 찾는 기존 사슬은 `kind='function'` 파티션
    588개만 닿는다. 그런데 **어떤 요구는 그 kind 에 아예 없다** — 68 요구 중 20 이
    미매핑이었고, 그중 16 이 걸린 SwDS 파티션의 kind 는::

        design_id 19 · table_row 12 · design_element 4   ← `function` 0

    즉 `swfn_35`(설계 ID) 나 `차속에 따른 도어 open 방지`(한글 기능명)가 키다.
    함수 이름이 그런 키를 닮을 리 없으므로 **구조적으로 못 닿는다**. 이름을 더 세게
    비벼도 안 되고, 비비면 오히려 유령 매칭이 는다.

    SwUDS 문서는 함수마다 "이 함수가 구현하는 설계 요소"를 `Related ID` 로 적어 둔다.
    그 설계 ID 로 SwDS 설계 파티션을 찾으면 요구에 닿는다 — 추적성 매트릭스가 이미
    쓰는 브리지와 **같은 구조**다(`report_gen/requirements.py::design_to_reqs`).

    ## ⚠ SwCom 을 뺀다

    `_DESIGN_ID_BRIDGE_RE`(SwFn/SwSTR/SwST/SwTK)만 통과시킨다. SwCom 은 컴포넌트
    레벨이라 fan-out 만 폭증시킨다 — 실측: 요구당 링크 중앙 138 → **4**, 최대
    1068 → 110, 합 16,461 → 766. 그러면서 위 16 건은 **16/16 그대로** 닿는다.

    ## ⚠ 이름으로만 잇는다

    반환 키가 `SwUFn` ID 가 아니라 **함수 이름**인 이유는 `report_gen/uds_related.py`
    모듈 docstring 에 있다(문서와 소스의 SwUFn 번호 체계가 다르다 — 43쌍 중 35쌍 불일치).
    """
    raw = str(uds_path or "").strip()
    if not raw:
        return {}
    p = Path(raw)
    try:
        data = p.read_bytes()
    except OSError as exc:
        _logger.warning("SwUDS 를 읽지 못해 설계-ID 브리지가 꺼진다 — %s (%s)", raw, exc)
        return {}
    from report_gen.requirements import _DESIGN_ID_BRIDGE_RE
    from report_gen.uds_related import docx_tables_text, extract_function_related_rows

    tables = docx_tables_text(data)
    if tables is None:
        # ⚠ 못 읽은 것을 "설계 ID 가 없다" 로 접지 않는다.
        _logger.warning("SwUDS 표를 파싱하지 못해 설계-ID 브리지가 꺼진다(문서 손상 가능): %s", raw)
        return {}
    out: Dict[str, List[str]] = {}
    total = kept = 0
    for row in extract_function_related_rows(tables):
        name = str(row.get("name") or "").strip()
        if not name:
            continue
        ids = [d for d in (row.get("design_ids") or [])]
        total += len(ids)
        tight = [d for d in ids if _DESIGN_ID_BRIDGE_RE.match(str(d).upper())]
        kept += len(tight)
        if tight:
            bucket = out.setdefault(name.lower(), [])
            for d in tight:
                if d not in bucket:
                    bucket.append(d)
    _logger.info(
        "SwUDS 설계-ID 브리지: 함수 %d개 · 설계 ID %d개 채택(SwCom 등 %d개 제외) — 출처=%s",
        len(out), kept, total - kept, raw,
    )
    return out


def map_requirements_to_functions(
    requirements: List[Dict[str, Any]],
    function_details: Dict[str, Dict[str, Any]],
    sds_map: Optional[Dict[str, Dict[str, str]]] = None,
    uds_design_ids: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, List[str]]:
    """Map requirement IDs to lists of function IDs (fid).

    Uses the `related` field in function_details to find reverse mapping.

    Args:
        sds_map: `related` 필드로 못 잇는 함수를 요구에 잇는 **폴백 매핑 출처**.
            None이면 저장소 `docs/` 글롭(`_load_default_sds_map`)을 쓰는데 이는
            **프로젝트 무관**이다 — 실측(HDPDM01): 요구-함수 링크 5,992건이 100%
            이 폴백에서 나왔다(폴백을 끄면 0/63). 요구 ID(`SwTR_0101` 등)는
            프로젝트 간 네임스페이스가 겹쳐 오매핑이 걸러지지도 않으므로,
            호출자가 대상 프로젝트의 SDS를 알고 있으면 반드시 넘길 것.
        uds_design_ids: `함수 이름(lower) → [설계 ID]` (`load_uds_design_ids`).
            이름으로 SwDS 를 못 찾는 요구를 **설계 ID 경유**로 잇는 3티어.
            None/빈 dict 면 그 티어는 **꺼진다** — 없는 것을 있는 척하지 않는다.
    """
    req_to_fids: Dict[str, List[str]] = {r["id"]: [] for r in requirements}
    if sds_map is None:
        sds_map = _load_default_sds_map()

    by_comment = by_sds = linkless = 0
    for fid, info in function_details.items():
        if not isinstance(info, dict):
            continue
        related = str(info.get("related") or info.get("comment_related") or "")
        matched = False
        for m in _REQ_ID_PAT.finditer(related):
            rid = m.group(1)
            if rid in req_to_fids and fid not in req_to_fids[rid]:
                req_to_fids[rid].append(fid)
                matched = True
        if matched:
            by_comment += 1
            continue
        hit = False
        for rid in _lookup_sds_related_ids(info, sds_map):
            if rid in req_to_fids and fid not in req_to_fids[rid]:
                req_to_fids[rid].append(fid)
            if rid in req_to_fids:
                hit = True
        by_sds += 1 if hit else 0
        linkless += 0 if hit else 1

    # ── 3티어: 설계-ID 브리지 (SwUDS Related ID → 설계 ID → SwDS → 요구) ──────
    # ⚠ 위 두 티어를 **건드리지 않고 별도 패스**로 돈다. 위 루프는 주석 매칭 시
    #   `continue` 로 SDS 티어를 건너뛰므로, 그 안에 끼워 넣으면 기존 링크 구성이
    #   조용히 달라진다. 여기서는 기존 결과에 **더하기만** 한다.
    #   실측(KJPDS02_PV): 요구 48/68 → **64/68** · 링크 8,397 → 8,667(+3.2%) ·
    #   요구당 링크 중앙 76 → 58(내려간다 — 작은 링크 집합을 가진 요구가 늘어서).
    by_design = 0
    if uds_design_ids:
        for fid, info in function_details.items():
            if not isinstance(info, dict):
                continue
            # ⚠ **이름**으로만 조인한다. `fid`(SwUFn 번호)로 조인하면 문서와 소스의
            #   번호 체계가 달라 오귀속이 된다(실측 43쌍 중 35쌍 불일치, 오귀속 링크
            #   276건) — `report_gen/uds_related.py` 모듈 docstring 참조.
            name = str(info.get("name") or "").strip().lower()
            if not name:
                continue
            gained = False
            for did in uds_design_ids.get(name) or ():
                entry = sds_map.get(str(did).lower())
                if not entry:
                    continue
                for m in _REQ_ID_PAT.finditer(str(entry.get("related") or "")):
                    rid = m.group(1)
                    if rid in req_to_fids and fid not in req_to_fids[rid]:
                        req_to_fids[rid].append(fid)
                        gained = True
            if gained:
                by_design += 1

    # ⚠ 침묵 금지 — 어느 요구가 **함수 근거 없이** TC 를 받는지 남긴다.
    #   `generate_test_cases` 는 매핑이 빈 요구에도 TC 를 낸다(`_generate_review_steps`).
    #   그래서 요구 커버리지는 100% 로 보이는데 그중 일부는 소스 근거가 0 이다.
    #   실측(KJPDS02_PV): 브리지 전엔 68 요구 중 20 이 여기 해당했고, 그 20 중 16 은
    #   SDS 의 `related` **에는 있었다**(우리가 그 파티션에 못 닿은 것). 설계-ID 브리지
    #   도입 후 **4** 로 줄었고, 남은 4 는 SwDS 어디에도 없다 = 문서 간 추적 부재라
    #   코드로 고칠 것이 아니다.
    unmapped = [r["id"] for r in requirements if not req_to_fids.get(r["id"])]
    if unmapped:
        _logger.warning(
            "STS 요구-함수 매핑: %d/%d 요구가 함수에 안 붙었다 — 이 요구들의 TC 는 "
            "소스 근거 없이 리뷰 절차로만 만들어진다: %s%s",
            len(unmapped), len(requirements), ", ".join(unmapped[:12]),
            " …" if len(unmapped) > 12 else "",
        )
    _logger.info(
        "STS 요구-함수 매핑 경로: 주석 related %d · SDS 파티션 %d · 설계-ID 브리지 %d "
        "(브리지 %s) · 이름/주석으로는 어느 요구에도 못 붙은 함수 %d",
        by_comment, by_sds, by_design,
        "켜짐" if uds_design_ids else "꺼짐(SwUDS 미지정)", linkless,
    )

    return req_to_fids


# ---------------------------------------------------------------------------
# Phase 1: Test method / gen-method determination
# ---------------------------------------------------------------------------

_HW_SIGNAL_PAT = re.compile(
    r"\bREG_|\blin_|\bPS\.|\bDiagData\.|\bADC|\bPWM|\bGPIO|\bCAN|\bLIN|\bSPI|\bI2C",
    re.I,
)

_ERROR_GUARD_PAT = re.compile(
    r"\b(error|fault|fail|invalid|null|timeout|overflow|underflow|out.of.range)\b",
    re.I,
)


# 내부 휴리스틱이 내는 라벨 → **정본 Introduction 표의 약어**.
# ⚠ 휴리스틱 자체는 유용하다(어떤 성격의 시험인지 구분한다). 문제는 그 결과 라벨이
#   SwTS Introduction 1.5/1.6 표에 없는 값이라는 것이었다 — 문서를 읽는 사람이
#   대조할 표가 없으면 그 칸은 근거가 아니라 장식이다.
_METHOD_TO_STS_VOCAB = {
    "RBT": "RBT", "FIT": "FIT",
    "FNCT": "RBT",   # 기능 시험 = 요구 기반 시험
    "RVW": "RBT",    # 리뷰도 요구 기반 확인으로 기록한다(정본에 RVW 칸이 없다)
    "ELCT": "RBT",   # 전기/HW 신호 확인도 요구 기반
}
_GEN_TO_STS_VOCAB = {
    "AOR": "AOR",
    "AEC": "ECA", "ECA": "ECA",   # 등가 분할 — SwTS 는 ECA, SwUTS 는 AEC 로 쓴다
    "ABV": "BAA", "BAA": "BAA",   # 경계값 분석 — SwTS 는 BAA, SwUTS 는 ABV
    "ERG": "AOR", "AFD": "AOR", "ADF": "AOR", "STA": "AOR",
    "AOI": "AOR", "AUC": "AOR", "ASV": "AOR",
}


# `Safety Related` 칸 — 구현은 `generators/safety_marks.py` 가 단일 출처다.
# ⚠ 이 파일에 다시 쓰지 말 것. 안전 판정을 고친 커밋 3건(fe9481e·e69b9dd·fb385d8)이
#   **여기에는 한 번도 안 닿았다** — 복제가 있으면 그 다음 수정도 같은 길을 간다.
_safety_mark = _safety_mark_impl


def _to_sts_vocab(method: str, gen: str) -> Tuple[str, str]:
    """휴리스틱 라벨을 정본 어휘로 좁힌다. 모르는 값은 정본 최빈값으로 떨어뜨린다."""
    return (
        _METHOD_TO_STS_VOCAB.get(str(method or "").strip().upper(), _DEFAULT_TEST_METHOD),
        _GEN_TO_STS_VOCAB.get(str(gen or "").strip().upper(), _DEFAULT_GEN_METHOD_STS),
    )


def _determine_test_method(
    req: Dict[str, Any],
    func_info: Optional[Dict[str, Any]] = None,
    logic_flow: Optional[List[Dict[str, Any]]] = None,
    hsis_signals: Optional[Dict[str, Any]] = None,
) -> Tuple[str, str]:
    """(test_method, gen_method) — **정본 어휘로 정규화해서** 돌려준다."""
    return _to_sts_vocab(*_determine_test_method_raw(req, func_info, logic_flow, hsis_signals))


def _determine_test_method_raw(
    req: Dict[str, Any],
    func_info: Optional[Dict[str, Any]] = None,
    logic_flow: Optional[List[Dict[str, Any]]] = None,
    hsis_signals: Optional[Dict[str, Any]] = None,
) -> Tuple[str, str]:
    """Return (test_method, gen_method) based on requirement type and function analysis."""
    rtype = req.get("req_type", "")
    asil = str(req.get("asil") or "").upper()

    # Use HSIS-extended pattern if available, else base pattern
    _hw_pat = (hsis_signals or {}).get("pat") or _HW_SIGNAL_PAT

    # ── Hardware/Electrical requirement → ELCT + AFD ─────────────────────
    if rtype == "EI":
        # Check if it involves hardware signals → ELCT, else fault injection
        req_desc = str(req.get("description") or req.get("name") or "")
        func_text = ""
        if func_info:
            func_text = " ".join([
                str(func_info.get("name") or ""),
                str(func_info.get("description") or ""),
                str(func_info.get("module_name") or ""),
            ])
        if _hw_pat.search(req_desc + " " + func_text):
            return ("ELCT", "AFD")
        return ("FIT", "ERG")  # EI without HW signals → error generation

    if rtype in ("TSR", "NTSR"):
        return ("FIT", "ERG")  # Safety requirements → error generation method

    if not func_info and not logic_flow:
        # NTR/NTSR with no function: use RBT (requirements-based test)
        if rtype in ("NTR", "NTSR"):
            return ("RBT", "ADF")
        return ("RVW", "ADF")

    has_switch = False
    has_if = False
    has_loop = False
    has_boundary = False
    has_error_guard = False

    if logic_flow:
        for node in logic_flow:
            ntype = node.get("type", "")
            cond = str(node.get("condition") or "")
            if ntype == "switch":
                has_switch = True
            elif ntype == "if":
                has_if = True
                if _ERROR_GUARD_PAT.search(cond):
                    has_error_guard = True
            elif ntype == "loop":
                has_loop = True

    if func_info:
        inputs = func_info.get("inputs") or []
        for inp in inputs:
            inp_str = str(inp).lower()
            if any(k in inp_str for k in ["range", "min", "max", "limit", "bound"]):
                has_boundary = True
            if re.search(r"\bu8\b|\bu16\b|\bu32\b|\bs8\b|\bs16\b|\bs32\b", inp_str):
                has_boundary = True
        # Hardware register access in function → ELCT
        fn_text = str(func_info.get("name") or "") + str(func_info.get("description") or "")
        if _hw_pat.search(fn_text):
            return ("ELCT", "AFD")

    if has_switch:
        return ("FNCT", "STA")
    if has_error_guard:
        return ("FIT", "ERG")   # Guard/error conditions → error generation
    if has_boundary:
        return ("FIT", "ABV")
    if has_if:
        return ("FNCT", "AEC")
    if has_loop:
        return ("FNCT", "AOR")

    return ("FIT", "AOR")


def _format_gen_method(gen: str) -> str:
    """Format gen method as numbered list when multiple (matches reference: '1. AOR\\n2. AOI')."""
    parts = [p.strip() for p in re.split(r"[,;/]", gen) if p.strip()]
    if len(parts) > 1:
        return "\n".join(f"{i + 1}. {p}" for i, p in enumerate(parts))
    return gen


# ---------------------------------------------------------------------------
# Phase 2: Test case / step generation from logic_flow
# ---------------------------------------------------------------------------

def _generate_steps_from_flow(
    logic_flow: List[Dict[str, Any]],
    func_info: Dict[str, Any],
) -> List[List[Dict[str, str]]]:
    """Generate multiple test-case step-lists from a function's logic flow.

    Handles nested if/else-if chains, switch-case, loops, and error-path branches.
    Returns a list of test cases, each being a list of {"action", "expected"} dicts.
    """
    if not logic_flow:
        return _generate_simple_steps(func_info)

    test_cases: List[List[Dict[str, str]]] = []
    normal_steps: List[Dict[str, str]] = []
    branch_tcs: List[List[Dict[str, str]]] = []

    _walk_flow_nodes(logic_flow, normal_steps, branch_tcs, depth=0)

    # Generate an error-path TC if any guard-like condition exists
    error_tc = _generate_error_path_tc(logic_flow, normal_steps)
    if error_tc:
        branch_tcs.append(error_tc)

    if branch_tcs:
        test_cases.extend(branch_tcs)
    elif normal_steps:
        test_cases.append(normal_steps)
    else:
        test_cases = _generate_simple_steps(func_info)

    for tc in test_cases:
        tc[:] = tc[:_MAX_STEPS_PER_TC]

    return test_cases[:_MAX_TC_PER_REQ]


def _walk_flow_nodes(
    nodes: List[Dict[str, Any]],
    prefix_steps: List[Dict[str, str]],
    branch_tcs: List[List[Dict[str, str]]],
    depth: int,
) -> None:
    """Recursively walk logic flow nodes, expanding nested branches into TCs."""
    max_depth = 4
    for node in nodes:
        ntype = node.get("type", "")

        if ntype == "call":
            prefix_steps.append({
                "action": f"{node['name']}() 호출",
                "expected": f"{node['name']} 정상 실행 확인",
            })

        elif ntype == "return":
            val = node.get("value", "")
            prefix_steps.append({
                "action": "함수 반환값 확인",
                "expected": f"반환값: {val}" if val else "정상 반환",
            })

        elif ntype == "if":
            cond = node.get("condition", "조건")
            true_body = node.get("true_body", [])
            false_body = node.get("false_body", [])
            elif_chains = node.get("elif_chains", [])

            true_steps = list(prefix_steps)
            true_steps.append({
                "action": f"조건 충족 설정: {cond}",
                "expected": "조건 분기 → True 경로 진입",
            })
            _expand_branch_body(true_body, true_steps, branch_tcs, depth, max_depth)
            branch_tcs.append(true_steps)

            for ei, elif_node in enumerate(elif_chains[:_MAX_TC_PER_REQ - 2]):
                econd = elif_node.get("condition", f"else-if #{ei+1}")
                ebody = elif_node.get("body", elif_node.get("true_body", []))
                elif_steps = list(prefix_steps)
                elif_steps.append({
                    "action": f"else-if 조건 설정: {econd}",
                    "expected": f"else-if 분기 #{ei+1} 진입",
                })
                if isinstance(ebody, list):
                    _expand_branch_body(ebody, elif_steps, branch_tcs, depth, max_depth)
                branch_tcs.append(elif_steps)

            false_steps = list(prefix_steps)
            false_steps.append({
                "action": f"조건 미충족 설정: NOT ({cond})",
                "expected": "조건 분기 → False/else 경로 진입",
            })
            if false_body:
                _expand_branch_body(false_body, false_steps, branch_tcs, depth, max_depth)
                branch_tcs.append(false_steps)

        elif ntype == "switch":
            expr = node.get("expr", "변수")
            cases = node.get("cases", [])
            default_calls = node.get("default_calls", [])

            for case in cases[:_MAX_TC_PER_REQ]:
                case_steps = list(prefix_steps)
                label = case.get("label", "?")
                case_steps.append({
                    "action": f"{expr} = {label} 설정",
                    "expected": f"switch 분기 → case {label} 진입",
                })
                case_body = case.get("body", case.get("calls", []))
                if isinstance(case_body, list) and case_body:
                    if isinstance(case_body[0], str):
                        for cn in case_body[:4]:
                            case_steps.append({
                                "action": f"{cn}() 호출 확인",
                                "expected": f"{cn} 정상 실행",
                            })
                    elif isinstance(case_body[0], dict) and depth < max_depth:
                        _walk_flow_nodes(case_body, case_steps, branch_tcs, depth + 1)
                branch_tcs.append(case_steps)

            if default_calls:
                def_steps = list(prefix_steps)
                def_steps.append({
                    "action": f"{expr} = 정의되지 않은 값 설정",
                    "expected": "switch 분기 → default 진입",
                })
                for call_name in (default_calls if isinstance(default_calls, list) else [])[:4]:
                    if isinstance(call_name, str):
                        def_steps.append({
                            "action": f"{call_name}() 호출 확인",
                            "expected": f"{call_name} 정상 실행",
                        })
                branch_tcs.append(def_steps)

        elif ntype == "loop":
            kind = node.get("kind", "loop")
            cond = node.get("condition", "")
            body = node.get("body", [])

            # Normal iteration TC
            loop_steps = list(prefix_steps)
            loop_steps.append({
                "action": f"{kind} 루프 조건 설정: {cond}",
                "expected": "루프 정상 진입",
            })
            for sub in body:
                if isinstance(sub, dict) and sub.get("type") == "call":
                    loop_steps.append({
                        "action": f"루프 내 {sub['name']}() 호출 확인",
                        "expected": f"{sub['name']} 반복 실행 확인",
                    })
            branch_tcs.append(loop_steps)

            # Zero-iteration TC
            zero_steps = list(prefix_steps)
            zero_steps.append({
                "action": f"{kind} 루프 즉시 종료 조건 설정: NOT ({cond})",
                "expected": "루프 미진입 (0회 반복)",
            })
            branch_tcs.append(zero_steps)


def _expand_branch_body(
    body: List[Dict[str, Any]],
    steps: List[Dict[str, str]],
    branch_tcs: List[List[Dict[str, str]]],
    depth: int,
    max_depth: int,
) -> None:
    """Expand sub-nodes inside a branch body, recursing into nested branches."""
    for sub in body:
        st = sub.get("type", "")
        if st == "call":
            steps.append({
                "action": f"{sub['name']}() 호출 확인",
                "expected": f"{sub['name']} 정상 실행",
            })
        elif st == "return":
            v = sub.get("value", "")
            steps.append({
                "action": "반환값 확인",
                "expected": f"반환값: {v}" if v else "정상 반환",
            })
        elif st == "if" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)
        elif st == "switch" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)
        elif st == "loop" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)
        elif st == "assign":
            var = sub.get("var", "")
            val = sub.get("value", "")
            if var:
                steps.append({
                    "action": f"{var} = {val} 설정 확인",
                    "expected": f"{var} 값 변경 정상",
                })
        elif st == "if" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)
        elif st == "switch" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)
        elif st == "loop" and depth < max_depth:
            _walk_flow_nodes([sub], steps, branch_tcs, depth + 1)


def _generate_error_path_tc(
    logic_flow: List[Dict[str, Any]],
    prefix_steps: List[Dict[str, str]],
) -> Optional[List[Dict[str, str]]]:
    """Generate a dedicated error-path TC from guard-like conditions."""
    guard_conds = _collect_guard_conds(logic_flow)
    if not guard_conds:
        return None
    steps = list(prefix_steps)
    for gc in guard_conds[:4]:
        steps.append({
            "action": f"에러 조건 설정: {gc}",
            "expected": "에러 처리 경로 진입",
        })
    steps.append({
        "action": "에러 처리 결과 확인",
        "expected": "에러 상태 반환 또는 안전 동작 수행",
    })
    return steps


def _collect_guard_conds(
    nodes: List[Dict[str, Any]], depth: int = 0,
) -> List[str]:
    """Recursively collect guard-like conditions from nested logic flow."""
    if depth > 4:
        return []
    _GUARD_KW = (
        "null", "err", "fail", "invalid", "< 0", "!= 0",
        "== null", "error", "fault", "status", "overflow",
        "underflow", "timeout", "limit", "range", "bound",
        "max", "min", "> 0xff", "> 255", "< 0x", "== 0",
        "!= ok", "nok", "e_not_ok", "e_ok", "std_return",
        "det_report", "dem_report", "safety", "diag",
    )
    result: List[str] = []
    for node in nodes:
        ntype = node.get("type", "")
        if ntype == "if":
            cond = str(node.get("condition", ""))
            if any(k in cond.lower() for k in _GUARD_KW):
                result.append(cond)
            result.extend(_collect_guard_conds(node.get("true_body", []), depth + 1))
            result.extend(_collect_guard_conds(node.get("false_body", []), depth + 1))
        elif ntype == "switch":
            for case in node.get("cases", []):
                body = case.get("body", case.get("calls", []))
                if isinstance(body, list) and body and isinstance(body[0], dict):
                    result.extend(_collect_guard_conds(body, depth + 1))
        elif ntype == "loop":
            result.extend(_collect_guard_conds(node.get("body", []), depth + 1))
    return result


def _ensure_min_steps(
    steps: List[Dict[str, str]],
    func_info: Dict[str, Any],
    min_count: int = 3,
) -> List[Dict[str, str]]:
    """Guarantee every TC has at least `min_count` steps.

    If the step list is short (e.g. only one branch-condition step), we append:
      1. A function-call step (if not already present)
      2. An output-verification step
      3. A state-check step (if still below min_count)
    """
    result = list(steps)
    name = func_info.get("name", "function") if func_info else "function"
    outputs = func_info.get("output") if func_info else None

    has_call = any("() 호출" in s.get("action", "") for s in result)
    if not has_call:
        result.append({
            "action": f"{name}() 호출",
            "expected": f"{name} 정상 실행 확인",
        })

    has_output_check = any(
        any(kw in s.get("action", "") for kw in ("출력", "반환값", "확인"))
        for s in result
    )
    if not has_output_check or len(result) < min_count:
        out_hint = f" ({outputs})" if outputs and str(outputs).strip() not in ("void", "None", "") else ""
        result.append({
            "action": "출력/반환값 확인",
            "expected": f"기대 결과와 일치{out_hint}",
        })

    if len(result) < min_count:
        result.append({
            "action": "시스템 상태 확인",
            "expected": "함수 실행 후 시스템 상태 정상 유지",
        })

    return result


def _generate_simple_steps(
    func_info: Dict[str, Any],
    _import_cache: Dict[str, Any] = {},  # noqa: B006 — intentional one-time init
) -> List[List[Dict[str, str]]]:
    """Fallback: generate 1~3 TCs from function info (no logic_flow).

    TC1 (NORMAL):   normal input values → call → verify output
    TC2 (BOUNDARY): boundary input values (min/max) → call → boundary behavior  [if inputs]
    TC3 (ERROR):    invalid input values (min_inv/max_inv) → call → error handling [if inputs]
    """
    # Lazy import once (shared via mutable default)
    if "ready" not in _import_cache:
        try:
            from generators.suts import get_boundary_values, infer_variable_type
            _import_cache["get_bv"] = get_boundary_values
            _import_cache["infer_type"] = infer_variable_type
        except Exception:
            _import_cache["get_bv"] = None
            _import_cache["infer_type"] = None
        _import_cache["ready"] = True

    get_boundary_values = _import_cache["get_bv"]
    infer_variable_type = _import_cache["infer_type"]

    name = func_info.get("name", "function")
    inputs = func_info.get("inputs") or []
    calls = func_info.get("calls_list") or []
    outputs_hint = func_info.get("output") or ""

    # Pre-compute boundary values once per variable (reused by TC1/TC2/TC3)
    var_cache: Dict[str, Dict[str, Any]] = {}  # vname → boundary dict
    if inputs and get_boundary_values and infer_variable_type:
        for inp in inputs[:5]:
            vname = str(inp).split(":")[0].strip()
            if vname not in var_cache:
                try:
                    vtype = infer_variable_type(vname)
                    var_cache[vname] = get_boundary_values(vtype)
                except Exception:
                    var_cache[vname] = {}

    # ── TC1: Normal path ──────────────────────────────────────────────────
    tc1: List[Dict[str, str]] = []
    if inputs and var_cache:
        mid_parts = []
        for inp in inputs[:5]:
            vname = str(inp).split(":")[0].strip()
            bnd = var_cache.get(vname)
            if bnd and "mid" in bnd:
                mid_parts.append(f"{vname}={bnd['mid']}")
            else:
                mid_parts.append(str(inp))
        in_str = ", ".join(mid_parts)
        tc1.append({"action": f"입력 설정 (정상값): {in_str}", "expected": "입력 파라미터가 유효 범위 내 정상 설정됨"})
    elif inputs:
        in_str = ", ".join(str(i) for i in inputs[:5])
        tc1.append({"action": f"입력 설정: {in_str}", "expected": "입력 파라미터 정상 설정"})
    tc1.append({"action": f"{name}() 호출", "expected": f"{name} 정상 실행 확인"})
    if calls:
        call_str = ", ".join(calls[:4])
        tc1.append({"action": f"내부 호출 확인: {call_str}", "expected": "하위 함수 정상 호출"})
    if outputs_hint and str(outputs_hint).strip() not in ("void", "None", ""):
        tc1.append({"action": "반환값 / 출력 확인", "expected": f"출력: {str(outputs_hint)[:80]} (정상 범위 내 값)"})
    elif calls:
        tc1.append({"action": "반환값 / 출력 확인",
                     "expected": f"{name} 정상 완료, {', '.join(calls[:3])} 호출 후 글로벌 상태 정상 유지"})
    else:
        tc1.append({"action": "반환값 / 출력 확인", "expected": f"{name} 정상 완료, 부작용 없음"})

    # For void functions with no explicit inputs, add globals-based verification
    globals_list = func_info.get("globals_global") or func_info.get("globals") or []
    if not inputs and globals_list:
        # Extract global variable names for concrete expected
        gvars = []
        for g in globals_list[:4]:
            gn = str(g).split("[")[0].strip().split(":")[-1].strip()
            if gn and len(gn) > 2:
                gvars.append(gn)
        if gvars:
            gvar_str = ", ".join(gvars[:3])
            tc1.append({"action": f"글로벌 상태 확인: {gvar_str}",
                         "expected": f"글로벌 변수 {gvar_str} 값이 함수 실행 전후 예상 범위 내 유지"})

    if not inputs or not var_cache:
        return [tc1]

    # ── TC2: Boundary (min/max) ───────────────────────────────────────────
    tc2: List[Dict[str, str]] = []
    bnd_parts_min: List[str] = []
    bnd_parts_max: List[str] = []
    for inp in inputs[:5]:
        vname = str(inp).split(":")[0].strip()
        bnd = var_cache.get(vname)
        if bnd and "min" in bnd:
            bnd_parts_min.append(f"{vname}={bnd['min']}")
            bnd_parts_max.append(f"{vname}={bnd['max']}")
        else:
            bnd_parts_min.append(str(inp))
            bnd_parts_max.append(str(inp))
    tc2.append({"action": f"입력 설정 (경계 최솟값): {', '.join(bnd_parts_min)}", "expected": "입력 경계 최솟값 설정"})
    tc2.append({"action": f"{name}() 호출", "expected": f"{name} 경계 최솟값 조건 실행 확인"})
    tc2.append({"action": f"입력 설정 (경계 최댓값): {', '.join(bnd_parts_max)}", "expected": "입력 경계 최댓값 설정"})
    tc2.append({"action": f"{name}() 호출", "expected": f"{name} 경계 최댓값 조건 실행 확인"})
    tc2.append({"action": "경계값 출력 확인",
                "expected": f"최솟값({', '.join(bnd_parts_min)}), 최댓값({', '.join(bnd_parts_max)}) 입력 시 유효 범위 내 정상 처리"})

    # ── TC3: Invalid (min_inv/max_inv) ────────────────────────────────────
    tc3: List[Dict[str, str]] = []
    inv_parts: List[str] = []
    for inp in inputs[:5]:
        vname = str(inp).split(":")[0].strip()
        bnd = var_cache.get(vname)
        if bnd and "max_inv" in bnd:
            inv_parts.append(f"{vname}={bnd['max_inv']}")
        else:
            inv_parts.append(str(inp))
    tc3.append({"action": f"입력 설정 (유효 범위 초과): {', '.join(inv_parts)}", "expected": "유효 범위 초과 입력 설정"})
    tc3.append({"action": f"{name}() 호출", "expected": f"{name} 범위 초과 입력 방어 처리 확인"})
    # Build concrete saturation expectation from cached boundaries
    sat_parts = []
    for inp in inputs[:3]:
        vname = str(inp).split(":")[0].strip()
        bnd = var_cache.get(vname)
        if bnd and "max" in bnd:
            sat_parts.append(f"{vname} 초과 시 출력 포화={bnd['max']} 또는 하한 클램프={bnd['min']}")
        else:
            sat_parts.append(f"{vname} 초과 시 안전값 유지")
    sat_str = "; ".join(sat_parts) if sat_parts else "출력 포화 또는 에러 플래그 설정"
    tc3.append({"action": "에러 처리 / 포화 출력 확인", "expected": f"방어 처리 결과: {sat_str}"})

    return [tc1, tc2, tc3]


def _generate_review_steps(req: Dict[str, Any]) -> List[List[Dict[str, str]]]:
    """Generate review-based TC steps when no function is mapped."""
    desc = req.get("description") or req.get("name") or req.get("id", "")
    verification = req.get("verification", "")

    steps: List[Dict[str, str]] = [
        {"action": "소스 코드에서 해당 요구사항 구현부 확인", "expected": "요구사항 구현 코드 존재 확인"},
        {"action": f"요구사항 내용 리뷰: {desc[:120]}", "expected": "구현이 요구사항과 일치함을 확인"},
    ]

    if verification:
        v_lines = [ln.strip() for ln in verification.split("\n") if ln.strip()]
        for vl in v_lines[:5]:
            vl_clean = re.sub(r"^\d+[\.\)]\s*", "", vl).strip()
            if vl_clean:
                steps.append({
                    "action": f"검증: {vl_clean[:120]}",
                    "expected": "검증 기준 충족 확인",
                })

    precond = req.get("software_state", "")
    if precond:
        steps.insert(0, {
            "action": f"사전조건 설정: S/W State = {precond}",
            "expected": f"시스템 {precond} 상태 진입 확인",
        })

    return [steps[:_MAX_STEPS_PER_TC]]


# ---------------------------------------------------------------------------
# Phase 2: Main TC generation
# ---------------------------------------------------------------------------

def generate_test_cases(
    requirements: List[Dict[str, Any]],
    function_details: Dict[str, Dict[str, Any]],
    req_to_fids: Dict[str, List[str]],
    project_config: Optional[Dict[str, Any]] = None,
    hsis_signals: Optional[Dict[str, Any]] = None,
    stats_out: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """Generate all test cases from requirements and function details.

    Args:
        stats_out: 주면 생성 통계를 채워 넣는다(절단 표면화용). 요구당 TC 상한
            `max_tc_per_req`(기본 5)은 **함수 루프 자체를 끊으므로**, 요구에 매핑된
            함수가 많으면 대부분이 시험 없이 남는다. 그 사실이 어디에도 안 남으면
            요구 단위 커버리지 100%가 "다 시험됨"으로 읽힌다.
            실측(HDPDM01): 매핑 함수 747개 중 TC 를 얻은 것 48개(93.6% 무시험),
            요구 35/37 이 상한에 도달, 그런데 보고 커버리지는 100.0%였다.
    """
    config = project_config or {}
    max_tc = config.get("max_tc_per_req", _MAX_TC_PER_REQ)
    test_env = config.get("default_test_env", _DEFAULT_TEST_ENV)

    # 캡 **전** 총량을 먼저 센다 — 소비처에서 결과 길이로 되짚으면 절단을 못 본다.
    mapped_fids: set = set()
    used_fids: set = set()
    truncated_reqs: List[str] = []

    project_asil = str(config.get("asil_level") or config.get("asil") or "").strip()
    _proj_is_safety = is_safety_asil(project_asil)

    all_tcs: List[Dict[str, Any]] = []

    for req in requirements:
        rid = req["id"]
        fids = req_to_fids.get(rid, [])
        rtype = req.get("req_type", "")

        req_asil = str(req.get("asil") or "").strip()
        if not req_asil and _proj_is_safety:
            req_asil = project_asil
            req["asil"] = project_asil
        is_safety = is_safety_asil(req_asil)

        if not fids:
            method, gen = _determine_test_method(req, hsis_signals=hsis_signals)
            step_sets = _generate_review_steps(req)
            for idx, steps in enumerate(step_sets[:max_tc]):
                tc_id = _make_tc_id(rid, idx + 1)
                all_tcs.append(_build_tc_dict(
                    tc_id=tc_id, req=req, steps=steps,
                    test_method=method, gen_method=gen,
                    test_env=test_env, is_safety=is_safety,
                ))
            continue

        mapped_fids.update(fids)
        tc_counter = 0
        for fid in fids:
            if tc_counter >= max_tc:
                # 남은 함수는 시험 없이 버려진다 — 이 사실을 반드시 남긴다
                truncated_reqs.append(rid)
                break
            info = function_details.get(fid, {})
            if not isinstance(info, dict):
                continue
            logic_flow = info.get("logic_flow") or []
            method, gen = _determine_test_method(req, info, logic_flow, hsis_signals=hsis_signals)
            step_sets = _generate_steps_from_flow(logic_flow, info)

            for steps in step_sets:
                if tc_counter >= max_tc:
                    break
                tc_counter += 1
                used_fids.add(fid)
                tc_id = _make_tc_id(rid, tc_counter)
                all_tcs.append(_build_tc_dict(
                    tc_id=tc_id, req=req, steps=_ensure_min_steps(steps, info),
                    test_method=method, gen_method=gen,
                    test_env=test_env, is_safety=is_safety,
                    func_name=info.get("name"),
                ))

    if stats_out is not None:
        stats_out.update({
            "max_tc_per_req": max_tc,
            "mapped_functions": len(mapped_fids),
            "functions_with_tc": len(used_fids),
            "functions_without_tc": len(mapped_fids - used_fids),
            "function_tc_coverage_pct": round(
                len(used_fids) / max(len(mapped_fids), 1) * 100, 1),
            "requirements_truncated": sorted(set(truncated_reqs)),
            "requirements_truncated_count": len(set(truncated_reqs)),
        })

    return all_tcs


def _make_tc_id(req_id: str, seq: int) -> str:
    return f"SwTC_{req_id}_{seq:02d}"


def _build_tc_dict(
    tc_id: str,
    req: Dict[str, Any],
    steps: List[Dict[str, str]],
    test_method: str,
    gen_method: str,
    test_env: str,
    is_safety: bool,
    func_name: Optional[str] = None,
    _bv_cache: Dict[str, Any] = {},  # noqa: B006 — intentional mutable default for lazy init
) -> Dict[str, Any]:
    # Lazy-init boundary helpers once (shared across all calls via mutable default)
    if "ready" not in _bv_cache:
        try:
            from generators.suts import get_boundary_values, infer_variable_type
            _bv_cache["get_bv"] = get_boundary_values
            _bv_cache["infer_type"] = infer_variable_type
        except Exception:
            _bv_cache["get_bv"] = None
            _bv_cache["infer_type"] = None
        _bv_cache["ready"] = True

    title = req.get("name") or req.get("description", "")[:60] or req["id"]
    if func_name:
        title = f"{title} - {func_name}"
    desc = req.get("description") or ""

    # Build precondition from req + function context
    precond_parts: List[str] = []
    sw_state = req.get("software_state", "")
    if sw_state:
        precond_parts.append(f"S/W State: {sw_state}")
    if func_name:
        precond_parts.append("시스템 초기화 완료")
        precond_parts.append(f"{func_name}() 호출 가능 상태")
    asil_val = str(req.get("asil") or "").strip()
    if is_safety_asil(asil_val):
        precond_parts.append(f"ASIL {asil_val} 안전 조건 충족")

    # Extract variable names from step actions
    input_vars: List[str] = []
    for step in steps:
        action = step.get("action", "")
        m_inp = re.search(r"입력 설정[^:]*:\s*(.+)", action)
        if m_inp:
            vars_str = m_inp.group(1).strip()
            for v in re.split(r",\s*", vars_str):
                vname = re.split(r"[=\s]", v.strip())[0].strip()
                if vname and len(vname) < 40 and not vname.startswith("("):
                    input_vars.append(vname)
    if input_vars:
        _get_bv = _bv_cache.get("get_bv")
        _infer_t = _bv_cache.get("infer_type")

        def _resolve_init(vn: str) -> str:
            if _get_bv and _infer_t:
                try:
                    return str(_get_bv(_infer_t(vn))["mid"])
                except Exception:
                    pass
            return "초기값"

        precond_parts.append("입력: " + ", ".join(
            f"{v}={_resolve_init(v)}" for v in input_vars[:4]
        ))

    precond = ", ".join(precond_parts)

    fs_req = ""
    asil = str(req.get("asil") or "").strip()
    if is_safety_asil(asil):
        related_id = req.get("related_id", "")
        sys_ids = re.findall(r"Sy\w+_\d+", related_id)
        if sys_ids:
            fs_req = ", ".join(sys_ids[:3])
        else:
            m = re.search(r"(\d+)(?:_(\d+))?", req["id"])
            if m:
                fs_req = f"{m.group(1)}_{m.group(2)}" if m.group(2) else m.group(1)

    return {
        "id": tc_id,
        "title": title[:120],
        # ⚠ 정본은 `O`=안전 관련 / `X`=비안전 이다(실측 X 86 · O 15). 예전엔
        #   `"X" if is_safety else ""` 라 **의미가 정반대**였다. SUTS 도 같은 결함이었다.
        "safety_related": _safety_mark(req.get("asil")),
        "test_environment": test_env,
        "test_method": test_method,
        "gen_method": _format_gen_method(gen_method),  # numbered list if multiple
        "fs_req": fs_req,
        "description": desc[:300],
        "precondition": precond,
        "srs_id": req["id"],
        "steps": steps,
    }


# ---------------------------------------------------------------------------
# Traceability matrix
# ---------------------------------------------------------------------------

def generate_traceability_matrix(
    test_cases: List[Dict[str, Any]],
    requirements: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """Generate traceability matrix data.

    Returns:
        {"req_ids": [...], "tc_ids": [...], "matrix": {tc_id: {req_id: 1/0}},
         "coverage": {"total_reqs": N, "covered_reqs": N, "pct": float,
                      "executable_covered_reqs": N, "executable_pct": float,
                      "review_only_reqs": [...], "review_only_count": N}}

    `covered_reqs`/`pct`는 **검증방법을 가리지 않은** 값이다(기존 계약 유지).
    거기에 검증방법 축을 더한다 — 아래 `_REVIEW_ONLY_METHODS` 주석 참조.
    """
    req_ids = sorted(set(r["id"] for r in requirements))
    tc_ids = [tc["id"] for tc in test_cases]
    matrix: Dict[str, Dict[str, int]] = {}
    covered_reqs: set = set()
    exec_covered: set = set()

    for tc in test_cases:
        tid = tc["id"]
        srs = tc.get("srs_id", "")
        is_executable = str(tc.get("test_method") or "").upper() not in _REVIEW_ONLY_METHODS
        row: Dict[str, int] = {}
        for rid in req_ids:
            if rid == srs:
                row[rid] = 1
                covered_reqs.add(rid)
                if is_executable:
                    exec_covered.add(rid)
            else:
                row[rid] = 0
        matrix[tid] = row

    total = len(req_ids)
    covered = len(covered_reqs)
    review_only = sorted(covered_reqs - exec_covered)
    return {
        "req_ids": req_ids,
        "tc_ids": tc_ids,
        "matrix": matrix,
        "coverage": {
            "total_reqs": total,
            "covered_reqs": covered,
            "pct": round(covered / max(total, 1) * 100, 1),
            # ── 검증방법 축 (additive) ──
            "executable_covered_reqs": len(exec_covered),
            "executable_pct": round(len(exec_covered) / max(total, 1) * 100, 1),
            "review_only_reqs": review_only,
            "review_only_count": len(review_only),
        },
    }


# ---------------------------------------------------------------------------
# Quality report
# ---------------------------------------------------------------------------

def generate_quality_report(
    test_cases: List[Dict[str, Any]],
    trace: Dict[str, Any],
    generation_stats: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    total_tc = len(test_cases)
    complete = sum(1 for tc in test_cases if tc.get("steps") and len(tc["steps"]) >= 2)
    safety_tc = sum(1 for tc in test_cases if tc.get("safety_related") == "X")
    methods: Dict[str, int] = {}
    gen_methods: Dict[str, int] = {}
    for tc in test_cases:
        m = tc.get("test_method", "?")
        methods[m] = methods.get(m, 0) + 1
        g = tc.get("gen_method", "?")
        gen_methods[g] = gen_methods.get(g, 0) + 1

    cov = trace.get("coverage", {})

    # 리뷰로만 덮인 요구는 숫자 옆에 반드시 말로 남긴다 — `requirement_coverage.pct`만
    # 읽는 소비자에게는 실행시험 100%와 구분되지 않기 때문이다.
    coverage_warnings: List[str] = []
    review_only = list(cov.get("review_only_reqs") or [])
    if review_only:
        shown = ", ".join(review_only[:10])
        suffix = f" 외 {len(review_only) - 10}건" if len(review_only) > 10 else ""
        coverage_warnings.append(
            f"[coverage] 요구 {len(review_only)}건은 실행 시험 없이 코드 리뷰(RVW)로만 덮였다 — "
            f"보고 커버리지 {cov.get('pct')}%는 리뷰를 포함한 값이고 "
            f"실행 시험 기준은 {cov.get('executable_pct')}%다 ({shown}{suffix})"
        )

    # TC 상한에 걸려 시험 없이 남은 함수 — 요구 단위 커버리지는 이 절단을 반영하지 않는다.
    gen_stats = generation_stats or {}
    without_tc = int(gen_stats.get("functions_without_tc") or 0)
    if without_tc:
        trunc = list(gen_stats.get("requirements_truncated") or [])
        shown = ", ".join(trunc[:8])
        suffix = f" 외 {len(trunc) - 8}건" if len(trunc) > 8 else ""
        coverage_warnings.append(
            f"[coverage] 요구당 TC 상한(max_tc_per_req="
            f"{gen_stats.get('max_tc_per_req')})에 걸려 매핑된 함수 "
            f"{gen_stats.get('mapped_functions')}개 중 {gen_stats.get('functions_with_tc')}개만 "
            f"TC 를 가진다(함수 기준 {gen_stats.get('function_tc_coverage_pct')}%, "
            f"무시험 {without_tc}개). 요구 커버리지 {cov.get('pct')}%는 요구 단위 값이라 "
            f"이 절단을 반영하지 않는다"
            + (f" — 상한 도달 요구: {shown}{suffix}" if trunc else "")
        )

    return {
        "total_test_cases": total_tc,
        "complete_test_cases": complete,
        "completeness_pct": round(complete / max(total_tc, 1) * 100, 1),
        "safety_test_cases": safety_tc,
        "requirement_coverage": cov,
        "test_method_distribution": methods,
        "gen_method_distribution": gen_methods,
        "coverage_warnings": coverage_warnings,
        "generation_stats": gen_stats,
    }


# ---------------------------------------------------------------------------
# XLSM output generation
# ---------------------------------------------------------------------------

def generate_sts_xlsm(
    template_path: Optional[str],
    test_cases: List[Dict[str, Any]],
    trace: Dict[str, Any],
    output_path: str,
    project_config: Optional[Dict[str, Any]] = None,
) -> str:
    """Generate STS XLSM file."""
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        _logger.error("openpyxl not installed")
        raise

    config = project_config or {}
    project_id = config.get("project_id", "PROJECT")
    doc_id = config.get("doc_id", f"{project_id}-STS")
    version = config.get("version", "v1.00")
    asil_level = config.get("asil_level", "")

    if template_path and Path(template_path).is_file():
        wb = openpyxl.load_workbook(template_path, keep_vba=True)
        _logger.info("Loaded STS template: %s", template_path)
    else:
        wb = openpyxl.Workbook()
        _create_cover_sheet(wb, project_id, doc_id, version, asil_level)
        _create_history_sheet(wb)
        _create_intro_sheet(wb)
        _create_sts_test_env_sheet(wb)
        _logger.info("Created new STS workbook (no template)")

    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    header_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    safety_fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
    header_font = Font(name="맑은 고딕", size=10, bold=True)
    data_font = Font(name="맑은 고딕", size=9)
    wrap_align = Alignment(wrap_text=True, vertical="top")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # --- Main test spec sheet ---
    # ⚠ 정본(KJPDS02_SwTS v1.02)의 시트명은 `3.SW Test Spec` 이다. 예전엔
    #   `3.SW Integration Test Spec`(= SwITS 의 시트명)에 SW 시험 명세를 쓰고 있었다.
    #   시트명이 다르면 정본을 읽는 쪽이 이 시트를 못 찾거나 통합시험으로 오독한다.
    sheet_name = _SPEC_SHEET_NAME
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(sheet_name)

    for ci, w in enumerate(_COL_WIDTHS, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Rows 1-5: document header block
    title_font = Font(name="맑은 고딕", size=13, bold=True)
    label_font = Font(name="맑은 고딕", size=9, bold=True)
    meta_fill = PatternFill(start_color="EBF3FB", end_color="EBF3FB", fill_type="solid")
    # Row 1: title merged A1:M1 (all 13 columns — matches reference)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=13)
    ws.cell(row=1, column=1, value="Software Test Specification").font = title_font
    ws.cell(row=1, column=1).alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[1].height = 26
    # Meta items: label at col 11, value merged cols 12-13 — one per row
    meta_items = [
        (2, "Doc. ID",    doc_id),
        (3, "Version",    version),
        (4, "Project",    project_id),
    ]
    for r, lbl, val in meta_items:
        c_lbl = ws.cell(row=r, column=11, value=lbl)
        c_lbl.font = label_font
        c_lbl.fill = meta_fill
        c_lbl.border = thin_border
        c_lbl.alignment = center_align
        ws.merge_cells(start_row=r, start_column=12, end_row=r, end_column=13)
        c_val = ws.cell(row=r, column=12, value=val)
        c_val.font = data_font
        c_val.border = thin_border
        c_val.alignment = wrap_align

    # Row 5: group sub-headers matching reference layout
    #   B5:K5 merged → "Test Case"
    #   L5           → "Expected Result"
    #   M5           → "Related ID"
    group_fill = PatternFill(start_color="B8CCE4", end_color="B8CCE4", fill_type="solid")
    group_font = Font(name="맑은 고딕", size=9, bold=True)
    group_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.merge_cells(start_row=5, start_column=2, end_row=5, end_column=11)  # B5:K5
    for ci in range(2, 14):
        ws.cell(row=5, column=ci).fill = group_fill
        ws.cell(row=5, column=ci).border = thin_border
    ws.cell(row=5, column=2, value="Test Case").font = group_font
    ws.cell(row=5, column=2).alignment = group_align
    ws.cell(row=5, column=12, value="Expected Result").font = group_font
    ws.cell(row=5, column=12).alignment = group_align
    ws.cell(row=5, column=13, value="Related ID").font = group_font
    ws.cell(row=5, column=13).alignment = group_align
    ws.row_dimensions[5].height = 18

    # Header row
    for ci, hdr in enumerate(_COL_HEADERS, 1):
        cell = ws.cell(row=_HEADER_ROW, column=ci, value=hdr)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_align
    ws.row_dimensions[_HEADER_ROW].height = 50.25  # matches reference row 6 height

    row_num = _HEADER_ROW + 1
    tc_counter = 0

    for tc in test_cases:
        tc_counter += 1
        steps = tc.get("steps") or [{"action": "", "expected": ""}]
        n_steps = len(steps)
        start_row = row_num
        end_row = row_num + n_steps - 1
        is_safety = tc.get("safety_related") == "X"

        for si, step in enumerate(steps):
            r = row_num + si
            ws.cell(row=r, column=1, value=tc_counter).font = data_font
            ws.cell(row=r, column=1).alignment = center_align

            # 열 번호는 전부 STS_COL(SSOT) 경유 — validator가 같은 상수를 본다.
            ws.cell(row=r, column=STS_COL["tc_id"], value=tc["id"] if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["title"], value=tc["title"] if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["safety_related"], value=tc.get("safety_related", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["test_environment"], value=tc.get("test_environment", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["test_method"], value=tc.get("test_method", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["gen_method"], value=tc.get("gen_method", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["fs_req"], value=tc.get("fs_req", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["description"], value=tc.get("description", "") if si == 0 else None).font = data_font
            ws.cell(row=r, column=STS_COL["precondition"], value=tc.get("precondition", "") if si == 0 else None).font = data_font

            # step 단위 열(병합하지 않는다 — TC당 여러 행)
            ws.cell(row=r, column=STS_COL["action"], value=step.get("action", "")).font = data_font
            ws.cell(row=r, column=STS_COL["expected"], value=step.get("expected", "")).font = data_font

            ws.cell(row=r, column=STS_COL["srs"], value=tc.get("srs_id", "") if si == 0 else None).font = data_font

            for ci in range(1, _LAST_COL + 1):
                ws.cell(row=r, column=ci).border = thin_border
                ws.cell(row=r, column=ci).alignment = (
                    center_align if ci in _CENTER_COLS else wrap_align
                )
                if is_safety:
                    ws.cell(row=r, column=ci).fill = safety_fill

        if n_steps > 1:
            for mc in _MERGE_COLS:
                col = mc + 1
                try:
                    ws.merge_cells(
                        start_row=start_row, start_column=col,
                        end_row=end_row, end_column=col,
                    )
                except Exception:
                    pass

        row_num = end_row + 1

    # --- Traceability sheet ---
    _write_traceability_sheet(wb, trace, thin_border, header_fill, header_font, data_font)

    # --- Save ---
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out))
    _logger.info("STS saved: %s (%d test cases, %d rows)", out, len(test_cases), row_num - 1)
    return str(out)


def _write_traceability_sheet(wb, trace, border, header_fill, header_font, data_font):
    from openpyxl.styles import Alignment
    from openpyxl.utils import get_column_letter

    sheet_name = "5. Traceability(SwRS)"
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(sheet_name)

    req_ids = trace.get("req_ids", [])
    tc_ids = trace.get("tc_ids", [])
    matrix = trace.get("matrix", {})

    center = Alignment(horizontal="center", vertical="center")

    ws.cell(row=1, column=1, value="Traceability Between [STS] and [SRS]").font = header_font

    ws.cell(row=3, column=3, value="Requirement IDs →").font = header_font
    for ci, rid in enumerate(req_ids):
        col = ci + 5
        ws.cell(row=4, column=col, value=rid).font = data_font
        ws.cell(row=4, column=col).border = border
        ws.cell(row=4, column=col).alignment = center
        ws.column_dimensions[get_column_letter(col)].width = max(6, len(rid) + 1)

    ws.cell(row=4, column=3, value="Test Case ID").font = header_font
    ws.cell(row=4, column=3).border = border
    ws.cell(row=4, column=4, value="Count").font = header_font
    ws.cell(row=4, column=4).border = border
    ws.column_dimensions["C"].width = 28
    ws.column_dimensions["D"].width = 8

    # Count row
    for ci, rid in enumerate(req_ids):
        col = ci + 5
        count = sum(1 for row in matrix.values() if row.get(rid, 0) == 1)
        ws.cell(row=5, column=col, value=count).font = data_font
        ws.cell(row=5, column=col).border = border
        ws.cell(row=5, column=col).alignment = center

    for ri, tid in enumerate(tc_ids):
        row = ri + 6
        ws.cell(row=row, column=3, value=tid).font = data_font
        ws.cell(row=row, column=3).border = border
        row_data = matrix.get(tid, {})
        mapped = sum(row_data.values())
        ws.cell(row=row, column=4, value=mapped).font = data_font
        ws.cell(row=row, column=4).border = border
        ws.cell(row=row, column=4).alignment = center
        for ci, rid in enumerate(req_ids):
            col = ci + 5
            val = row_data.get(rid, 0)
            if val:
                ws.cell(row=row, column=col, value=1).font = data_font
            ws.cell(row=row, column=col).border = border
            ws.cell(row=row, column=col).alignment = center


def _create_cover_sheet(wb, project_id, doc_id, version, asil_level):
    ws = wb.active
    ws.title = "Cover"
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    title_font = Font(name="맑은 고딕", size=24, bold=True)
    label_font = Font(name="맑은 고딕", size=9, bold=True)
    data_font = Font(name="맑은 고딕", size=9)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    hdr_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Column widths matching reference: A:2.875, B:6.875, C-G:13.0, H:4.625, I:6.875, J:13.0, K:10.625
    col_widths = {"A": 2.875, "B": 6.875, "C": 13.0, "D": 13.0, "E": 13.0,
                  "F": 13.0, "G": 13.0, "H": 4.625, "I": 6.875, "J": 13.0, "K": 10.625}
    for col, w in col_widths.items():
        ws.column_dimensions[col].width = w

    # B5:K5 merged — main title block (height=123 matching reference)
    ws.merge_cells("B5:K5")
    ws["B5"] = "Software Test Specification\n(소프트웨어 테스트 명세서)"
    ws["B5"].font = title_font
    ws["B5"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[5].height = 123.0

    # I2 = "Doc. ID" label, J2:K2 merged = value
    ws["I2"] = "Doc. ID"
    ws["I2"].font = label_font
    ws["I2"].alignment = center
    ws.merge_cells("J2:K2")
    ws["J2"] = doc_id
    ws["J2"].font = data_font
    ws["J2"].alignment = center

    # I3 = "Version" label, J3:K3 merged = value
    ws["I3"] = "Version"
    ws["I3"].font = label_font
    ws["I3"].alignment = center
    ws.merge_cells("J3:K3")
    ws["J3"] = version
    ws["J3"].font = data_font
    ws["J3"].alignment = center

    info_rows = [
        ("Project", project_id),
        ("ASIL Level", asil_level),
        ("Status", "Draft"),
        ("Date", datetime.now().strftime("%Y-%m-%d")),
    ]
    for i, (lbl, val) in enumerate(info_rows):
        r = 21 + i
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=5)
        ws.merge_cells(start_row=r, start_column=6, end_row=r, end_column=11)
        ws.cell(row=r, column=2, value=lbl).font = label_font
        ws.cell(row=r, column=2).fill = hdr_fill
        ws.cell(row=r, column=2).border = thin
        ws.cell(row=r, column=2).alignment = center
        ws.cell(row=r, column=6, value=val).font = data_font
        ws.cell(row=r, column=6).border = thin
        ws.cell(row=r, column=6).alignment = left


def _create_history_sheet(wb):
    ws = wb.create_sheet("History")
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    hdr_font = Font(name="맑은 고딕", size=10, bold=True)
    data_font = Font(name="맑은 고딕", size=10)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    hdr_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Column widths matching reference: A:1.25, B:8.375, C:9.125, D:35.5, E:8.625, F:13.0, G:13.0, H:1.25
    ws.column_dimensions["A"].width = 1.25
    ws.column_dimensions["H"].width = 1.25
    ws.row_dimensions[2].height = 18.0
    ws.row_dimensions[3].height = 14.25

    ws.merge_cells("B2:G2")
    ws["B2"] = "▶ Revision History"
    ws["B2"].font = Font(name="맑은 고딕", size=12, bold=True)
    ws["B2"].alignment = Alignment(horizontal="left", vertical="center")

    headers = ["Version", "Date", "Description", "Author", "Reviewer", "Approver"]
    widths = [8.375, 9.125, 35.5, 8.625, 13.0, 13.0]
    for ci, (h, w) in enumerate(zip(headers, widths), 2):
        col = get_column_letter(ci)
        cell = ws.cell(row=4, column=ci, value=h)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.border = thin
        cell.alignment = center
        ws.column_dimensions[col].width = w

    init_row = ["v1.00", datetime.now().strftime("%Y-%m-%d"), "Initial draft", "", "", ""]
    for ci, val in enumerate(init_row, 2):
        cell = ws.cell(row=5, column=ci, value=val)
        cell.font = data_font
        cell.border = thin
        cell.alignment = center if ci in (2, 3) else left


def _create_intro_sheet(wb):
    ws = wb.create_sheet("1.Introduction")
    ws["A1"] = "1. Introduction"
    ws["A3"] = "1.1 Purpose"
    ws["A4"] = "본 문서는 소프트웨어 테스트 사양을 기술한다."
    ws["A6"] = "1.5 Test Method"
    methods = [
        ("FNCT", "Functional test - 기능 테스트"),
        ("FIT", "Fault Injection test - 결함 주입 테스트"),
        ("ELCT", "Electrical test - 전기적 테스트"),
        ("RVW", "Review - 코드 리뷰"),
        ("RBT", "Requirements Based test - 요구사항 기반 테스트"),
    ]
    for i, (code, desc) in enumerate(methods):
        ws.cell(row=8 + i, column=1, value=code)
        ws.cell(row=8 + i, column=2, value=desc)


def _create_sts_test_env_sheet(wb):
    """Create 2.Test Environment sheet matching reference structure."""
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    ws = wb.create_sheet("2.Test Environment")
    hdr_font = Font(name="맑은 고딕", size=10, bold=True)
    data_font = Font(name="맑은 고딕", size=9)
    hdr_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws["A1"] = "2. Test Environment"
    ws["A1"].font = Font(name="맑은 고딕", size=13, bold=True)
    ws.row_dimensions[1].height = 26

    ws["A3"] = "2.1 SW Test Environment"
    ws["A3"].font = hdr_font

    headers = ["ID", "Name", "Description", "HW", "OS", "Tool"]
    widths = [12, 20, 50, 20, 20, 20]
    for ci, (h, w) in enumerate(zip(headers, widths), 1):
        from openpyxl.utils import get_column_letter
        col = get_column_letter(ci)
        cell = ws.cell(row=5, column=ci, value=h)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.border = thin
        cell.alignment = center
        ws.column_dimensions[col].width = w

    env_rows = [
        ("SwTE_01", "Host PC Simulation", "PC 기반 소프트웨어 시뮬레이션 환경", "x86 PC", "Windows", "MATLAB/Simulink"),
        ("SwTE_02", "HIL (HW-in-the-Loop)", "실제 ECU 하드웨어 기반 테스트 환경", "Target ECU", "AUTOSAR OS", "CANoe/dSPACE"),
        ("SwTE_03", "SIL (SW-in-the-Loop)", "소프트웨어 루프 테스트 환경", "x86 PC", "Linux", "GCC/GCOV"),
        ("SwTE_04", "MIL (Model-in-the-Loop)", "모델 기반 테스트 환경", "x86 PC", "Windows", "MATLAB"),
        ("SyTE_02", "System Test Env", "시스템 레벨 테스트 환경", "Vehicle/Rig", "—", "CANoe"),
    ]
    for ri, row_data in enumerate(env_rows, 6):
        for ci, val in enumerate(row_data, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font = data_font
            cell.border = thin
            cell.alignment = center if ci in (1,) else left


# ---------------------------------------------------------------------------
# AI Enhancement Module
# ---------------------------------------------------------------------------

_STS_AI_SYSTEM_PROMPT = (
    "You are a software test engineer writing test specifications for automotive ECU software (ISO 26262).\n"
    "Given a test case skeleton, enhance the Description, Pre-condition, Test Actions, and Expected Results.\n"
    "Rules:\n"
    "- Be specific and technical. Use actual signal names, function names, and parameter values.\n"
    "- For boundary tests, include specific boundary values.\n"
    "- For state transition tests, specify the exact states and transitions.\n"
    "- Keep Korean language for descriptions.\n"
    "- Return JSON: {\"description\":\"...\", \"precondition\":\"...\", \"steps\":[{\"action\":\"...\",\"expected\":\"...\"}]}\n"
    "- Do NOT invent requirements or functions not in the input."
)


_AI_TIMEOUT_SEC = 30
_AI_MAX_RETRIES = 2


def _sts_ai_call_with_retry(agent_call_fn, ai_config, messages, *,
                              max_retries: int = _AI_MAX_RETRIES,
                              timeout: int = _AI_TIMEOUT_SEC) -> str:
    """AI call wrapper with timeout and retry for STS enhancement."""
    import threading

    last_err: Optional[Exception] = None
    for attempt in range(1, max_retries + 1):
        result_holder: Dict[str, Any] = {}
        exc_holder: list = []

        def _invoke():
            try:
                r = agent_call_fn(
                    ai_config, messages,
                    role="writer", stage="sts_enhance",
                    settings={"temperature": 0.3},
                )
                result_holder["val"] = r
            except Exception as ex:
                exc_holder.append(ex)

        t = threading.Thread(target=_invoke, daemon=True)
        t.start()
        t.join(timeout=timeout)

        if t.is_alive():
            _logger.warning("STS AI call timed out (attempt %d/%d)", attempt, max_retries)
            last_err = TimeoutError(f"AI timed out after {timeout}s")
            continue
        if exc_holder:
            last_err = exc_holder[0]
            _logger.warning("STS AI call error (attempt %d/%d): %s", attempt, max_retries, last_err)
            continue

        raw = result_holder.get("val")
        reply = raw.get("output", "") if isinstance(raw, dict) else ""
        if reply:
            return reply
        last_err = ValueError("Empty AI response")

    if last_err:
        _logger.warning("STS AI call exhausted retries: %s", last_err)
    return ""


def _parse_sts_ai_response(reply: str) -> Optional[Dict[str, Any]]:
    """Parse and validate STS AI JSON response."""
    import json as _json
    if not reply:
        return None
    try:
        payload = _json.loads(reply) if isinstance(reply, str) else reply
    except Exception:
        m_json = re.search(r"\{[\s\S]*\}", reply)
        if m_json:
            try:
                payload = _json.loads(m_json.group())
            except Exception:
                return None
        else:
            return None

    if not isinstance(payload, dict):
        return None

    valid: Dict[str, Any] = {}
    if isinstance(payload.get("description"), str) and len(payload["description"]) <= 1000:
        valid["description"] = payload["description"][:500]
    if isinstance(payload.get("precondition"), str) and len(payload["precondition"]) <= 600:
        valid["precondition"] = payload["precondition"][:300]
    ai_steps = payload.get("steps")
    if isinstance(ai_steps, list) and ai_steps:
        cleaned = []
        for s in ai_steps[:_MAX_STEPS_PER_TC]:
            if isinstance(s, dict) and isinstance(s.get("action"), str) and isinstance(s.get("expected"), str):
                cleaned.append({"action": s["action"], "expected": s["expected"]})
        if cleaned:
            valid["steps"] = cleaned
    return valid if valid else None


def enhance_test_cases_with_ai(
    test_cases: List[Dict[str, Any]],
    function_details: Dict[str, Dict[str, Any]],
    ai_config: Optional[Dict[str, Any]] = None,
    max_batch: int = 50,
    sds_summary: str = "",
    stp_context: str = "",
    hsis_signals: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """Enhance test case descriptions using Gemini AI with timeout/retry."""
    if not ai_config:
        _logger.info("AI enhancement skipped (no config)")
        return test_cases

    try:
        from workflow.ai import agent_call
    except ImportError:
        _logger.warning("workflow.ai not available; skipping AI enhancement")
        return test_cases

    enhanced = 0
    batch_size = min(max_batch, len(test_cases))
    candidates = [tc for tc in test_cases if tc.get("steps") and len(tc["steps"]) <= 3]
    candidates = candidates[:batch_size]

    for tc in candidates:
        func_name = ""
        for s in tc.get("steps", []):
            m = re.search(r"(\w+)\(\)", s.get("action", ""))
            if m:
                func_name = m.group(1)
                break

        func_ctx = ""
        if func_name:
            for fid, info in function_details.items():
                if isinstance(info, dict) and info.get("name") == func_name:
                    func_ctx = (
                        f"Function: {func_name}\n"
                        f"Description: {info.get('description', '')}\n"
                        f"Inputs: {info.get('inputs', [])}\n"
                        f"Output: {info.get('output', 'void')}\n"
                        f"Calls: {info.get('calls_list', [])}\n"
                        f"Component: {info.get('module_name', '')}\n"
                    )
                    break

        extra_ctx_parts: List[str] = []
        if sds_summary:
            extra_ctx_parts.append(f"[SDS Design Context]\n{sds_summary[:800]}")
        if stp_context:
            extra_ctx_parts.append(f"[STP Test Strategy]\n{stp_context[:600]}")
        if hsis_signals and hsis_signals.get("signals"):
            hsis_lines = []
            for sig in hsis_signals["signals"][:15]:
                # ID 열이 빈 HSIS 행도 이제 채택되므로(`_is_hsis_data_row`) 라벨을 비워두지
                # 않는다 — 빈 라벨은 LLM이 ID를 지어내게 만든다.
                _label = sig["id"] or sig["related_id"] or "(HSI ID 미기재)"
                hsis_lines.append(
                    f"  {_label}: {sig['signal_name']} "
                    f"(SW: {sig['sw_var_name']}, Dir: {sig['direction']}, "
                    f"Char: {sig['characteristics'][:40]})"
                )
            extra_ctx_parts.append("[HSIS Hardware Signals]\n" + "\n".join(hsis_lines))
        extra_ctx = "\n\n".join(extra_ctx_parts)

        user_msg = (
            f"Requirement: {tc.get('srs_id', '')} - {tc.get('title', '')}\n"
            f"Description: {tc.get('description', '')}\n"
            f"Test Method: {tc.get('test_method', '')}, Gen Method: {tc.get('gen_method', '')}\n"
            f"{func_ctx}\n"
            f"{extra_ctx}\n"
            f"Current steps:\n"
        )
        for i, s in enumerate(tc.get("steps", []), 1):
            user_msg += f"  {i}. Action: {s.get('action', '')} | Expected: {s.get('expected', '')}\n"

        reply = _sts_ai_call_with_retry(
            agent_call, ai_config,
            [
                {"role": "system", "content": _STS_AI_SYSTEM_PROMPT},
                {"role": "user", "content": user_msg},
            ],
        )

        validated = _parse_sts_ai_response(reply)
        if validated:
            if "description" in validated:
                tc["description"] = validated["description"]
            if "precondition" in validated:
                tc["precondition"] = validated["precondition"]
            if "steps" in validated:
                tc["steps"] = validated["steps"]
            enhanced += 1

    _logger.info("AI enhanced %d/%d test cases", enhanced, len(candidates))
    return test_cases


# ---------------------------------------------------------------------------
# Top-level pipeline function
# ---------------------------------------------------------------------------

def generate_sts(
    requirements_text: List[str],
    function_details: Dict[str, Dict[str, Any]],
    output_path: str,
    template_path: Optional[str] = None,
    project_config: Optional[Dict[str, Any]] = None,
    srs_docx_path: Optional[str] = None,
    sds_docx_path: Optional[str] = None,
    uds_path: Optional[str] = None,
    stp_path: Optional[str] = None,
    hsis_path: Optional[str] = None,
    ai_config: Optional[Dict[str, Any]] = None,
    on_progress: Optional[Any] = None,
    source_root: Optional[str] = None,
) -> Dict[str, Any]:
    """Top-level STS generation pipeline.

    Args:
        requirements_text: Raw requirement text lines (from SRS parsing)
        function_details: UDS function_details dict (fid -> info)
        output_path: Path for output XLSM
        template_path: Optional STS template XLSM
        project_config: Optional config dict
        srs_docx_path: Optional path to SRS DOCX for direct table parsing
        sds_docx_path: Optional path to SDS DOCX for design context
        uds_path: Optional path to generated UDS DOCX/XLSM for function descriptions
        stp_path: Optional path to STP DOCX for test strategy context
        hsis_path: Optional path to HSIS xlsx for hardware signal enrichment
        ai_config: Optional AI config dict for Gemini enhancement
        on_progress: Optional callback(pct: int, message: str) for progress updates
        source_root: Optional C source root(s) — 콤마 구분 복수 경로 허용.
            품질 DB(record_run)의 project_root 로만 쓰인다. generate_suts/generate_sits
            와 동일한 역할이며, 과거 이 파라미터가 없는 채로 record_run 이
            source_root 를 참조해 NameError → except 로 삼켜져 **STS 품질 기록이
            통째로 유실**되고 있었다.

    Returns:
        Dict with keys: output_path, quality_report, trace_coverage
    """
    def _progress(pct: int, msg: str):
        _logger.info("[%d%%] %s", pct, msg)
        if on_progress:
            try:
                on_progress(pct, msg)
            except Exception:
                pass

    _progress(5, "STS 생성 시작")
    t0 = time.time()

    # ── Load supplementary document context ──────────────────────────────
    sds_summary = ""
    stp_ctx = ""

    # 이 맵은 function_details 보강뿐 아니라 **요구-함수 매핑의 폴백 출처**로도 쓰인다
    # (`map_requirements_to_functions`). None으로 두면 저장소 `docs/` 글롭(프로젝트 무관)이
    # 대신하는데, 실측상 요구-함수 링크 전량이 그 폴백에서 나온다.
    sds_partition_map: Optional[Dict[str, Dict[str, str]]] = None

    if sds_docx_path:
        _progress(7, "SDS 설계 컨텍스트 로드 중")
        sds_summary = _load_sds_summary(sds_docx_path)
        if sds_summary:
            _logger.info("SDS summary loaded (%d chars)", len(sds_summary))
        # ⚠ 파티션 맵 추출을 summary 유무에 종속시키지 않는다 — 과거엔 `if sds_summary:`
        # 안에 있어서, 요약 절이 안 잡히는 SDS면 파티션 표가 멀쩡해도 맵을 아예 안 만들고
        # 조용히 저장소 폴백으로 넘어갔다.
        try:
            from report_gen.requirements import _extract_sds_partition_map
            sds_partition_map = _extract_sds_partition_map(sds_docx_path) or None
        except Exception as _e:
            _logger.warning("SDS 파티션 맵 추출 실패 — 요구-함수 매핑이 저장소 docs/ "
                            "폴백(프로젝트 무관)으로 넘어간다: %s (%s)", sds_docx_path, _e)
        if sds_partition_map:
            _logger.info("SDS 파티션 %d건 로드 — 출처=%s", len(sds_partition_map), sds_docx_path)
            for fid, info in function_details.items():
                if not isinstance(info, dict):
                    continue
                for cand in _function_sds_candidates(info):
                    entry = sds_partition_map.get(cand.lower())
                    if entry:
                        if entry.get("asil") and not info.get("asil"):
                            info["asil"] = entry["asil"]
                        if entry.get("description") and not info.get("sds_description"):
                            info["sds_description"] = entry["description"]
                        break
        else:
            _logger.warning("SDS를 지정했으나 파티션 0건 — 요구-함수 매핑이 저장소 docs/ "
                            "폴백(프로젝트 무관)으로 넘어간다: %s", sds_docx_path)

    # 설계-ID 브리지의 좌측 끝. SwUDS 를 안 주면 **꺼진다** — 실측상 그 상태에서
    # 요구 매핑은 48/68 이고, 브리지가 켜지면 64/68 이다(`load_uds_design_ids` 참조).
    uds_design_ids: Dict[str, List[str]] = {}

    if uds_path:
        _progress(8, "UDS 함수 설명 로드 중")
        uds_descs = _load_uds_descriptions(uds_path)
        if uds_descs:
            _logger.info("UDS descriptions loaded (%d entries)", len(uds_descs))
            _merge_uds_into_function_details(function_details, uds_descs)
        uds_design_ids = load_uds_design_ids(uds_path)
    else:
        _logger.warning(
            "SwUDS 미지정 — 설계-ID 브리지가 꺼진다. SwDS 의 설계 파티션"
            "(`design_id`/`design_element`)에만 걸린 요구는 함수 근거 없이 리뷰 TC 로만 "
            "만들어진다(실측 KJPDS02_PV: 16 요구)")

    if stp_path:
        _progress(9, "STP 시험 전략 로드 중")
        stp_ctx = _load_stp_context(stp_path)
        if stp_ctx:
            _logger.info("STP context loaded (%d chars)", len(stp_ctx))

    # ── HSIS hardware signal enrichment ──────────────────────────────────
    hsis_signals: Dict[str, Any] = {}
    if hsis_path:
        _progress(10, "HSIS 하드웨어 신호 로드 중")
        hsis_signals = _load_hsis_signals(hsis_path)
        if hsis_signals.get("signals"):
            _logger.info(
                "HSIS signals loaded: %d signals, SW vars: %s",
                len(hsis_signals["signals"]),
                ", ".join(hsis_signals["sw_var_names"][:5]),
            )

    # ── Requirements parsing ──────────────────────────────────────────────
    reqs: List[Dict[str, Any]] = []
    _progress(10, "요구사항 파싱 중")
    if srs_docx_path and Path(srs_docx_path).is_file():
        reqs = parse_srs_docx_tables(srs_docx_path)

    if not reqs and requirements_text:
        reqs = parse_requirements_structured(requirements_text)

    _progress(25, f"요구사항 {len(reqs)}개 파싱 완료")

    _progress(30, "요구사항-함수 매핑 중")
    req_to_fids = map_requirements_to_functions(reqs, function_details,
                                                sds_map=sds_partition_map,
                                                uds_design_ids=uds_design_ids)
    mapped = sum(1 for v in req_to_fids.values() if v)
    _progress(40, f"{mapped}/{len(reqs)}개 요구사항 매핑 완료")

    _progress(45, "테스트 케이스 생성 중")
    gen_stats: Dict[str, Any] = {}
    test_cases = generate_test_cases(
        reqs, function_details, req_to_fids, project_config,
        hsis_signals=hsis_signals or None,
        stats_out=gen_stats,
    )
    _progress(60, f"테스트 케이스 {len(test_cases)}개 생성 완료")
    if gen_stats.get("functions_without_tc"):
        _logger.warning(
            "STS: TC 상한(%s)에 걸려 매핑 함수 %s개 중 %s개만 시험한다 (무시험 %s개, "
            "상한 도달 요구 %s건) — 요구 커버리지는 이 절단을 반영하지 않는다",
            gen_stats.get("max_tc_per_req"), gen_stats.get("mapped_functions"),
            gen_stats.get("functions_with_tc"), gen_stats.get("functions_without_tc"),
            gen_stats.get("requirements_truncated_count"),
        )

    if ai_config:
        _progress(65, "AI 향상 적용 중")
        enhance_test_cases_with_ai(
            test_cases, function_details, ai_config,
            sds_summary=sds_summary,
            stp_context=stp_ctx,
            hsis_signals=hsis_signals or None,
        )
        _progress(75, "AI 향상 완료")

    _progress(78, "추적성 매트릭스 생성 중")
    trace = generate_traceability_matrix(test_cases, reqs)

    _progress(82, "품질 리포트 생성 중")
    quality = generate_quality_report(test_cases, trace, generation_stats=gen_stats)

    _progress(85, "XLSM 파일 생성 중")
    out = generate_sts_xlsm(template_path, test_cases, trace, output_path, project_config)

    _progress(92, "생성 문서 자동 검증 중")
    try:
        validation = validate_sts_xlsm(out)
        # 생성 수 ↔ 파일 기록 수 대조(세 생성기 공용 단일 출처). 이게 없으면 아래
        # 반환값의 test_case_count 는 **파일이 아니라 생성기가 세어준 값**이라,
        # 라이터가 흘려도 호출자는 끝까지 모른다.
        validation = apply_write_back_check(validation, {"tc_count": len(test_cases)})
        if validation.get("issues"):
            _logger.warning("STS validation issues: %s", validation["issues"])
    except Exception as _ve:
        _logger.warning("STS validation skipped: %s", _ve)
        # B7 — 검증이 크래시했으면 valid:True(통과)로 위장하지 않는다. 검증을 **못 한** 것을
        # 통과로 쓰는 fail-open 이다(미검증 ≠ 유효). valid:False + 사유를 warnings 로 표면화.
        validation = {
            "valid": False, "issues": [], "stats": {},
            "warnings": [f"검증 실행 실패(미검증): {_ve}"],
        }

    validation_report_path = ""
    try:
        validation_report_path = generate_sts_validation_report(out, quality)
        _logger.info("STS validation report: %s", validation_report_path)
    except Exception as _vr:
        _logger.warning("STS validation report generation skipped: %s", _vr)

    elapsed = time.time() - t0
    _progress(100, f"STS 생성 완료 ({elapsed:.1f}초)")

    # Quality DB recording (non-fatal)
    try:
        from workflow.quality.recorder import record_run
        record_run(
            "sts", quality,
            project_root=str(source_root or ""),
            elapsed_sec=elapsed,
            output_path=out,
            ai_model=str((ai_config or {}).get("model", "")),
        )
    except Exception:
        # non-fatal 은 유지하되 **침묵은 금지**. 이 `except: pass` 가 source_root
        # NameError 를 몇 년간 삼켜 STS 품질 기록이 통째로 유실된 걸 아무도 몰랐다.
        # recorder 내부엔 _logger.exception 이 있지만, 인자 평가에서 터지면
        # record_run 진입 자체가 없어 거기까지 못 간다 → 호출부에서 남긴다.
        _logger.exception("STS quality record skipped (non-fatal)")

    return {
        "output_path": out,
        "quality_report": quality,
        "trace_coverage": trace.get("coverage"),
        "test_case_count": len(test_cases),
        "elapsed_seconds": round(elapsed, 1),
        "validation": validation,
        "validation_report_path": validation_report_path,
    }


# ---------------------------------------------------------------------------
# Document validation
# ---------------------------------------------------------------------------

# 읽는 쪽도 **writer 와 같은 상수**를 본다. 예전엔 시트명이 문자열로 박혀 있어
# writer 만 고치면 validator 가 "Missing sheet" 를 내며 조용히 갈라졌다.
_STS_SHEET_CANDIDATES = (_SPEC_SHEET_NAME, "3.SW Integration Test Spec", "2.SW Test Spec")


def _normalize_header(value: Any) -> str:
    """헤더 라벨 비교용 정규화 — 개행·공백·구두점 제거 후 소문자."""
    return re.sub(r"[^a-z0-9]", "", str(value or "").lower())


def _detect_sts_columns(ws: Any, header_row: int = _HEADER_ROW) -> Tuple[Dict[str, int], List[str]]:
    """header_row의 라벨로 필드→열 번호를 찾는다. 못 찾은 필드는 스키마 상수로 폴백.

    열 번호 하드코딩만 쓰면 템플릿이 한 칸만 밀려도 **엉뚱한 열을 조용히 읽는다**(이 파일
    상단 _STS_SCHEMA 주석의 실제 사고). 반대로 헤더 탐지만 쓰면 헤더가 없는 구 산출물에서
    아무것도 못 읽는다. 둘을 합치고, 폴백을 썼다는 사실은 호출자에게 알린다.

    Returns: (필드→열, 폴백을 쓴 필드 목록)
    """
    label_to_key = {_normalize_header(label): key for _, key, label in _STS_SCHEMA if label}
    found: Dict[str, int] = {}
    try:
        # read_only 워크북은 .cell() 호출마다 시트를 재파싱한다 — 헤더 한 줄도 iter_rows로.
        max_col = min(int(ws.max_column or 0), 64)   # 폭주 방지 — STS는 13열
        header_vals = next(
            ws.iter_rows(min_row=header_row, max_row=header_row,
                         max_col=max_col, values_only=True),
            (),
        )
        for col, raw in enumerate(header_vals, 1):
            key = label_to_key.get(_normalize_header(raw))
            if key and key not in found:             # 중복 헤더는 첫 번째만 채택
                found[key] = col
    except Exception as exc:
        _logger.warning("STS header 탐지 실패(상수 폴백): %s", exc)

    fallback_fields = [k for k in STS_COL if k not in found]
    cols = {**STS_COL, **found}
    return cols, fallback_fields


def validate_sts_xlsm(xlsm_path: str) -> Dict[str, Any]:
    """생성된 STS XLSM의 구조·데이터 품질 검증.

    ⚠ 이 함수는 **STS 전용**이다. 과거엔 generators.suts에 있으면서 SUTS 레이아웃 상수를
    그대로 썼다 — writer가 11/12/13열(Action/Expected/SRS)에 쓰는데 validator는 5/6/4열
    (TestEnvironment/TestMethod/SafetyRelated)을 읽어, Action·Expected가 전부 비어도
    "이상 없음"이 되고 요구 링크율은 Safety Related 채움률을 보고했다.
    """
    issues: List[str] = []
    warnings: List[str] = []
    stats: Dict[str, Any] = {}

    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"valid": False, "issues": ["openpyxl not installed"], "warnings": [], "stats": {}}

    p = Path(xlsm_path)
    if not p.exists():
        return {"valid": False, "issues": [f"File not found: {xlsm_path}"], "warnings": [], "stats": {}}

    try:
        wb = load_workbook(str(p), read_only=True, data_only=True)
    except Exception as e:
        return {"valid": False, "issues": [f"Cannot open: {e}"], "warnings": [], "stats": {}}

    try:
        stats["sheets"] = wb.sheetnames
        stats["sheet_count"] = len(wb.sheetnames)

        for s in ("Cover", "History", "1.Introduction"):
            if s not in wb.sheetnames:
                warnings.append(f"Optional sheet missing: {s}")

        sts_sheet = next((c for c in _STS_SHEET_CANDIDATES if c in wb.sheetnames), None)
        if not sts_sheet:
            issues.append("No STS main sheet found")
            return {"valid": False, "issues": issues, "warnings": warnings, "stats": stats}

        ws = wb[sts_sheet]
        max_row = ws.max_row or 0
        stats["max_row"] = max_row
        stats["max_col"] = ws.max_column or 0
        stats["sheet"] = sts_sheet

        cols, fallback_fields = _detect_sts_columns(ws)
        stats["columns"] = dict(cols)
        if fallback_fields:
            # 필수 축이 헤더에서 안 잡히면 잘못된 열을 읽고 있을 수 있다 — 침묵 금지.
            missing_required = [f for f in _STS_REQUIRED_FIELDS if f in fallback_fields]
            if missing_required:
                warnings.append(
                    "헤더에서 못 찾아 상수 위치로 판독한 필수 열: "
                    + ", ".join(missing_required)
                )

        tc_count = 0
        empty_title_tcs = 0
        no_step_tcs = 0
        no_expected_tcs = 0
        reqs_linked = 0

        # ⚠ read_only 워크북에서 ws.cell(row=N, ...) 랜덤 접근은 매 호출마다 시트를 다시
        # 훑어 O(행²)가 된다(이 저장소 실측 전례: 75분 → iter_rows 0.9초). 한 번만 순회한다.
        needed = ("tc_id", "title", "action", "expected", "srs")
        max_needed_col = max(cols[f] for f in needed)

        def _val(row_vals: Tuple[Any, ...], field: str) -> str:
            idx = cols[field] - 1
            return str(row_vals[idx] or "").strip() if idx < len(row_vals) else ""

        # TC ID는 첫 스텝 행에만 있고 Action/Expected는 스텝마다 있다 — TC 블록 단위로 본다.
        cur_has_action = cur_has_expected = False
        has_open_tc = False

        def _close_tc() -> None:
            nonlocal no_step_tcs, no_expected_tcs
            if not cur_has_action:
                no_step_tcs += 1
            if not cur_has_expected:
                no_expected_tcs += 1

        for row_vals in ws.iter_rows(
            min_row=_HEADER_ROW + 1, max_row=max_row,
            max_col=max_needed_col, values_only=True,
        ):
            tc_id = _val(row_vals, "tc_id")
            if tc_id:
                if has_open_tc:
                    _close_tc()
                tc_count += 1
                has_open_tc = True
                cur_has_action = cur_has_expected = False
                if not _val(row_vals, "title"):
                    empty_title_tcs += 1
                if _val(row_vals, "srs"):
                    reqs_linked += 1
            if not has_open_tc:
                continue   # TC 시작 전 잔여 행 — 어느 TC에도 귀속되지 않는다
            if _val(row_vals, "action"):
                cur_has_action = True
            if _val(row_vals, "expected"):
                cur_has_expected = True
        if has_open_tc:
            _close_tc()

        stats["tc_count"] = tc_count
        stats["empty_title_tcs"] = empty_title_tcs
        stats["no_step_tcs"] = no_step_tcs
        stats["no_expected_tcs"] = no_expected_tcs
        stats["reqs_linked"] = reqs_linked
        stats["req_linkage_pct"] = round(reqs_linked / tc_count * 100, 1) if tc_count else 0

        if tc_count == 0:
            issues.append("No test cases found")
        if empty_title_tcs > tc_count * 0.3:
            issues.append(f"Over 30% TCs lack titles ({empty_title_tcs}/{tc_count})")
        # Action/Expected 부재는 "시험을 수행할 수 없다"는 뜻이라 경고가 아니라 결함이다.
        if tc_count and no_step_tcs == tc_count:
            issues.append(f"All TCs lack action steps ({no_step_tcs}/{tc_count})")
        elif no_step_tcs > tc_count * 0.5:
            warnings.append(f"Over 50% TCs lack action steps ({no_step_tcs}/{tc_count})")
        if tc_count and no_expected_tcs == tc_count:
            issues.append(f"All TCs lack expected results ({no_expected_tcs}/{tc_count})")
        elif no_expected_tcs > tc_count * 0.5:
            warnings.append(f"Over 50% TCs lack expected results ({no_expected_tcs}/{tc_count})")
        if tc_count > 0 and reqs_linked == 0:
            warnings.append("No TCs linked to requirements")
    finally:
        wb.close()

    return {"valid": len(issues) == 0, "issues": issues, "warnings": warnings, "stats": stats}


def validate_sts_output(xlsm_path: str) -> Dict[str, Any]:
    """Validate a generated STS XLSM for structural completeness.

    Returns dict with 'valid' bool, 'issues' list, 'warnings' list, and 'stats' dict.
    """
    try:
        return validate_sts_xlsm(xlsm_path)
    except ImportError:
        pass

    from openpyxl import load_workbook
    wb = load_workbook(xlsm_path, read_only=True, data_only=True)
    issues: List[str] = []
    stats: Dict[str, Any] = {"sheets": wb.sheetnames, "sheet_count": len(wb.sheetnames)}

    _sheet = next((n for n in _STS_SHEET_CANDIDATES if n in wb.sheetnames), None)
    if _sheet:
        ws = wb[_sheet]
        tc_count = 0
        for r in range(_HEADER_ROW + 1, (ws.max_row or _HEADER_ROW) + 1):
            tc_id = ws.cell(row=r, column=2).value
            if tc_id and str(tc_id).strip():
                tc_count += 1
        stats["tc_count"] = tc_count
        if tc_count == 0:
            issues.append("No test cases found in main sheet")
    else:
        issues.append(f"Missing sheet: {_SPEC_SHEET_NAME}")

    wb.close()
    stats["issues"] = issues
    stats["valid"] = len(issues) == 0
    return stats


def generate_sts_validation_report(
    xlsm_path: str,
    quality_report: Optional[Dict[str, Any]] = None,
) -> str:
    """Generate a validation report markdown for STS XLSM.

    Writes a .validation.md file next to the XLSM and returns its path.
    """
    validation = validate_sts_xlsm(xlsm_path)
    stats = validation.get("stats", {})
    issues = validation.get("issues", [])
    warnings = validation.get("warnings", [])
    qr = quality_report or {}

    lines = [
        "# STS 생성 문서 자동 검증 리포트",
        "",
        f"**파일**: `{Path(xlsm_path).name}`  ",
        f"**검증 시각**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  ",
        f"**결과**: {'PASS' if validation.get('valid') else 'FAIL'}",
        "",
        "---",
        "",
        "## 1. 구조 검증",
        "",
        "| 항목 | 값 |",
        "|------|-----|",
        f"| 시트 수 | {stats.get('sheet_count', 0)} |",
        f"| 시트 목록 | {', '.join(stats.get('sheets', []))} |",
        f"| TC 수 | {stats.get('tc_count', 0)} |",
        f"| 빈 제목 TC | {stats.get('empty_title_tcs', 0)} |",
        f"| 스텝 없는 TC | {stats.get('no_step_tcs', 0)} |",
        f"| 기대값 없는 TC | {stats.get('no_expected_tcs', 0)} |",
        f"| 요구사항 연결 TC | {stats.get('reqs_linked', 0)} |",
        f"| 요구사항 연결률 | {stats.get('req_linkage_pct', 0)}% |",
        "",
    ]

    if qr:
        lines.extend([
            "## 2. 품질 지표",
            "",
            "| 항목 | 값 |",
            "|------|-----|",
            f"| 총 TC 수 | {qr.get('total_test_cases', 0)} |",
            f"| 완전한 TC 수 | {qr.get('complete_test_cases', 0)} ({qr.get('completeness_pct', 0)}%) |",
            f"| 안전 관련 TC | {qr.get('safety_test_cases', 0)} |",
            "",
        ])
        if qr.get("test_method_distribution"):
            lines.extend([
                "### 테스트 메서드 분포",
                "",
                "| 메서드 | 수 |",
                "|--------|-----|",
            ])
            for k, v in qr["test_method_distribution"].items():
                lines.append(f"| {k} | {v} |")
            lines.append("")

    gate_items = [
        ("TC 존재", stats.get("tc_count", 0) > 0),
        ("빈 제목 < 30%", stats.get("empty_title_tcs", 0) <= stats.get("tc_count", 1) * 0.3 if stats.get("tc_count") else True),
        ("스텝 존재 > 50%", stats.get("no_step_tcs", 0) < stats.get("tc_count", 1) * 0.5 if stats.get("tc_count") else True),
        ("기대값 존재 > 50%", stats.get("no_expected_tcs", 0) < stats.get("tc_count", 1) * 0.5 if stats.get("tc_count") else True),
        ("요구사항 연결 존재", stats.get("reqs_linked", 0) > 0 if stats.get("tc_count") else True),
    ]
    passed = sum(1 for _, ok in gate_items if ok)

    lines.extend([
        f"## 3. Quality Gate ({passed}/{len(gate_items)})",
        "",
        "| 항목 | 결과 |",
        "|------|------|",
    ])
    for name, ok in gate_items:
        lines.append(f"| {name} | {'PASS' if ok else 'FAIL'} |")
    lines.append("")

    if issues:
        lines.extend(["## 4. Issues", ""])
        for i in issues:
            lines.append(f"- {i}")
        lines.append("")

    if warnings:
        lines.extend(["## 5. Warnings", ""])
        for w in warnings:
            lines.append(f"- {w}")
        lines.append("")

    out_path = Path(xlsm_path).with_suffix(".validation.md")
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return str(out_path)
